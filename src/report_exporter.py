"""
Word document report exporter for Spec Critic.

Generates a formatted .docx report from a PipelineResult:
    - Title block with generation metadata
    - Files reviewed list
    - Summary table (severity counts) with colored cell shading
    - LEED and placeholder alerts
    - Per-spec findings grouped by severity, sorted by confidence
    - Verification verdicts and corrections inline with findings
    - Cross-spec coordination findings (if cross-check was enabled)

Collapsible findings (v2.5.0):
    Each finding header uses Word Heading 3 style. In Word 2016+ and
    365, hovering over any heading shows a collapse triangle. Clicking
    it hides everything between that heading and the next heading of
    the same or higher level. This means:
    - Collapse a severity heading (Heading 1) to hide all its findings
    - Collapse a single finding heading (Heading 3) to hide its details
    No macros or special XML required — this is native Word behavior.

Usage:
    from report_exporter import export_report

    export_report(
        pipeline_result=result,
        output_path=Path("review-report.docx"),
    )
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from .report_status import (
    EDIT_ACTION_DISPLAY_ORDER,
    EditActionLabel,
    ReportStatus,
    STATUS_DISPLAY_ORDER,
    classify_edit_action,
    classify_status,
    edit_action_label,
    status_glyph,
    status_label,
    summarize_edit_actions,
    summarize_statuses,
)



SEVERITY_COLORS = {
    "CRITICAL": RGBColor(192, 0, 0),
    "HIGH": RGBColor(255, 102, 0),
    "MEDIUM": RGBColor(192, 152, 0),
    "GRIPES": RGBColor(128, 0, 128),
}

SEVERITY_SHADING = {
    "CRITICAL": "C00000",
    "HIGH": "FF6600",
    "MEDIUM": "C09800",
    "GRIPES": "800080",
}

VERDICT_COLORS = {
    "CONFIRMED": RGBColor(0, 128, 0),
    "CORRECTED": RGBColor(204, 132, 0),
    "UNVERIFIED": RGBColor(128, 128, 128),
    "DISPUTED": RGBColor(192, 0, 0),
}

VERDICT_ICONS = {
    "CONFIRMED": "✓",
    "CORRECTED": "✎",
    "DISPUTED": "✗",
    "UNVERIFIED": "—",
}

CONFIDENCE_COLORS = {
    "high": RGBColor(0, 128, 0),
    "moderate": RGBColor(204, 132, 0),
    "low": RGBColor(192, 0, 0),
}

COORDINATION_COLOR = RGBColor(6, 182, 212)

SEVERITY_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "GRIPES"]

STATUS_COLORS: dict[ReportStatus, RGBColor] = {
    ReportStatus.VERIFIED_SUPPORTED: RGBColor(0, 128, 0),
    ReportStatus.VERIFIED_CONTRADICTED: RGBColor(204, 132, 0),
    ReportStatus.DISPUTED: RGBColor(192, 0, 0),
    ReportStatus.INSUFFICIENT_EVIDENCE: RGBColor(128, 128, 128),
    ReportStatus.LOCALLY_CLASSIFIED: RGBColor(59, 130, 246),
    ReportStatus.NOT_CHECKED: RGBColor(100, 100, 100),
    ReportStatus.MANUAL_REVIEW_REQUIRED: RGBColor(255, 102, 0),
}

STATUS_SHADING: dict[ReportStatus, str] = {
    ReportStatus.VERIFIED_SUPPORTED: "008000",
    ReportStatus.VERIFIED_CONTRADICTED: "CC8400",
    ReportStatus.DISPUTED: "C00000",
    ReportStatus.INSUFFICIENT_EVIDENCE: "808080",
    ReportStatus.LOCALLY_CLASSIFIED: "3B82F6",
    ReportStatus.NOT_CHECKED: "646464",
    ReportStatus.MANUAL_REVIEW_REQUIRED: "FF6600",
}

EDIT_ACTION_COLORS: dict[EditActionLabel, RGBColor] = {
    EditActionLabel.AUTO_EDIT_CANDIDATE: RGBColor(0, 128, 0),
    EditActionLabel.MANUAL_EDIT_CANDIDATE: RGBColor(204, 132, 0),
    EditActionLabel.REPORT_ONLY: RGBColor(100, 100, 100),
    EditActionLabel.SUPPRESSED: RGBColor(192, 0, 0),
}



def _confidence_tier(confidence: float) -> str:
    """Return 'high', 'moderate', or 'low' for a confidence score."""
    if confidence >= 0.85:
        return "high"
    elif confidence >= 0.60:
        return "moderate"
    return "low"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"


def _set_paragraph_collapsed(paragraph) -> None:
    """Mark a heading paragraph as collapsed-by-default when the doc is opened.

    Emits ``<w15:collapsed w15:val="1"/>`` per [MS-DOCX] §2.5.1.3. The element
    MUST be in the Word 2012 extension namespace
    (http://schemas.microsoft.com/office/word/2012/wordml). Using the base
    w: namespace causes Word to silently ignore the element, which is why
    earlier versions of this code rendered Sources headings expanded on open.

    Semantics (per the spec): when this element is set on a heading of level N,
    immediately subsequent paragraphs with a higher heading level number appear
    collapsed when the document is opened. In our use, it's set on the
    "Sources" Heading 4; the URL paragraph below carries outlineLvl=8, which
    satisfies "higher heading level number" and gets collapsed. The next
    finding's Heading 3 (outlineLvl=2) is lower and terminates the zone.

    python-docx's default nsmap does NOT register w15, so we construct the
    element via lxml directly instead of using OxmlElement / qn helpers.
    """
    pPr = paragraph._p.get_or_add_pPr()
    collapsed = etree.SubElement(
        pPr,
        f"{{{_W15_NS}}}collapsed",
        nsmap={"w15": _W15_NS},
    )
    collapsed.set(f"{{{_W15_NS}}}val", "1")


def _set_paragraph_outline_level(paragraph, level: int) -> None:
    """Set <w:outlineLvl> on a paragraph so Word includes it in a preceding
    heading's open-time collapse zone without changing its visual style."""
    pPr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), str(level))
    pPr.append(outline)


def _add_styled_paragraph(doc: Document, text: str, style: str | None = None,
                          bold: bool = False, color: RGBColor | None = None,
                          size: int | None = None, space_after: int | None = None,
                          italic: bool = False):
    """Add a paragraph with optional styling. Returns the paragraph."""
    para = doc.add_paragraph()
    if style:
        para.style = style

    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)

    return para



def _write_title_block(doc: Document, review, files_reviewed: list[str],
                       cycle_label: str = "2025") -> None:
    """Write the report title and metadata.

    Uses separate paragraphs instead of \\n within runs to ensure
    reliable rendering across all Word versions and viewers.
    """
    title = doc.add_heading("Spec Critic — M&P Specification Review Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_lines = [
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Model: {review.model}",
        f"Files Reviewed: {len(files_reviewed)}",
        f"Code Cycle: California {cycle_label}",
    ]

    for line in meta_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(2)



def _write_files_reviewed(doc: Document, files_reviewed: list[str]) -> None:
    """Write the files reviewed section with a bullet list."""
    doc.add_heading("Files Reviewed", level=1)
    for filename in files_reviewed:
        doc.add_paragraph(filename, style='List Bullet')




def _summarize_verification_outcomes(findings: list) -> dict[str, object]:
    """Roll up the trust-model statuses + raw verdict counts for the methodology note.

    Chunk N: the status histogram (computed via
    :func:`report_status.summarize_statuses`) drives the methodology
    narrative and the new trust-model summary table. The verdict-count
    breakdown is preserved alongside it so existing summary lines
    (``CONFIRMED / CORRECTED / DISPUTED / UNVERIFIED``) keep working.
    """
    stats = {
        "total_findings": len(findings),
        "with_verification": 0,
        "verdict_counts": {"CONFIRMED": 0, "CORRECTED": 0, "DISPUTED": 0, "UNVERIFIED": 0},
        "status_counts": summarize_statuses(findings),
        "edit_action_counts": summarize_edit_actions(findings),
    }
    for finding in findings:
        verification = getattr(finding, "verification", None)
        if not verification:
            continue
        stats["with_verification"] += 1
        verdict = str(getattr(verification, "verdict", "UNVERIFIED") or "UNVERIFIED").upper()
        if verdict not in stats["verdict_counts"]:
            verdict = "UNVERIFIED"
        stats["verdict_counts"][verdict] += 1
    verified_non_unverified = stats["verdict_counts"]["CONFIRMED"] + stats["verdict_counts"]["CORRECTED"] + stats["verdict_counts"]["DISPUTED"]
    stats["all_unverified"] = stats["with_verification"] > 0 and verified_non_unverified == 0 and stats["verdict_counts"]["UNVERIFIED"] == stats["with_verification"]
    stats["partial_unverified"] = stats["verdict_counts"]["UNVERIFIED"] > 0 and verified_non_unverified > 0
    return stats


def _write_methodology_note(doc, cross_check_enabled: bool = False, cycle_label: str = "2025", cross_check_status: str | None = None, cross_check_reason: str = "", verification_stats: dict[str, object] | None = None) -> None:
    """Write a brief methodology note explaining how the review was produced."""
    doc.add_heading("About This Review", level=1)

    doc.add_paragraph(
        "This report was generated by Spec Critic, an AI-assisted specification "
        "review tool. Each specification was analyzed by Claude for "
        "code compliance issues, coordination problems, and technical errors "
        "relevant to California K-12 DSA projects. Findings are classified by "
        "severity (Critical, High, Medium, Gripe) and assigned a confidence score "
        "reflecting the model\u2019s certainty."
    )

    verification_stats = verification_stats or {}
    all_unverified = bool(verification_stats.get("all_unverified", False))
    partial_unverified = bool(verification_stats.get("partial_unverified", False))
    with_verification = int(verification_stats.get("with_verification", 0) or 0)

    if all_unverified:
        para2_text = (
            "Verification was attempted but did not return usable results. "
            "Findings have not been independently verified."
        )
    elif partial_unverified:
        para2_text = (
            "Findings were checked in a secondary AI verification pass with web search access. "
            "Some findings could not be verified — see individual verdicts."
        )
    elif with_verification > 0:
        para2_text = (
            "All findings were checked in a secondary AI verification pass with web search access. "
            "Verification verdicts (Confirmed, Corrected, Disputed, or Unverified) reflect the verifier model's assessment and should be treated as advisory."
        )
    else:
        para2_text = (
            "No verification outcomes were recorded for this run. "
            "Findings should be treated as unverified unless noted otherwise."
        )

    para2_text += f" This review used California {cycle_label} code cycle references."

    if cross_check_enabled and cross_check_status == "completed":
        para2_text += " Cross-check completed."
    elif cross_check_enabled and cross_check_status == "skipped":
        para2_text += f" Cross-check was skipped: {cross_check_reason}"
    elif cross_check_enabled and cross_check_status == "failed":
        para2_text += f" Cross-check failed: {cross_check_reason}"

    para2_text += (
        " This report is advisory \u2014 findings should be reviewed by the "
        "engineer of record before acting on them."
    )

    doc.add_paragraph(para2_text)

    doc.add_paragraph(
        "Tip: In Word, hover over any heading to reveal a collapse triangle. "
        "Click it to hide the content beneath that heading. Use this to "
        "collapse individual findings or entire severity groups."
    )



def _write_summary_table(doc: Document, review, cross_check_result, *, total_elapsed_seconds: float | None = None) -> None:
    """Write the summary section with a styled severity counts table."""
    doc.add_heading("Summary", level=1)

    cc_count = (len(cross_check_result.findings)
                if cross_check_result and cross_check_result.findings else 0)

    columns = [
        ("CRITICAL", review.critical_count, "C00000"),
        ("HIGH", review.high_count, "FF6600"),
        ("MEDIUM", review.medium_count, "C09800"),
        ("GRIPES", review.gripe_count, "800080"),
        ("TOTAL", review.total_count, "333333"),
    ]
    if cc_count > 0:
        columns.append(("CROSS-CHECK", cc_count, "06B6D4"))

    table = doc.add_table(rows=2, cols=len(columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_idx, (label, _count, hex_color) in enumerate(columns):
        cell = table.rows[0].cells[col_idx]
        _set_cell_shading(cell, hex_color)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        if label == "MEDIUM":
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for col_idx, (_label, count, _hex) in enumerate(columns):
        cell = table.rows[1].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(count))
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run("Review Stage Tokens: ").bold = True
    para.add_run(
        f"{review.input_tokens:,} input → "
        f"{review.output_tokens:,} output"
    )

    para = doc.add_paragraph()
    para.add_run("Note: ").bold = True
    para.add_run(
        "Cross-check and verification token usage are not included in the totals above."
    )

    para = doc.add_paragraph()
    para.add_run("Processing Time: ").bold = True
    processing_seconds = total_elapsed_seconds if total_elapsed_seconds is not None else review.elapsed_seconds
    para.add_run(f"{processing_seconds:.1f} seconds")

    all_findings = list(review.findings)
    if cross_check_result and cross_check_result.findings:
        all_findings.extend(cross_check_result.findings)

    verdicts: dict[str, int] = {}
    for f in all_findings:
        if f.verification:
            v = f.verification.verdict
            verdicts[v] = verdicts.get(v, 0) + 1

    if verdicts:
        para = doc.add_paragraph()
        para.add_run("Verification: ").bold = True
        verdict_parts = [
            f"{verdicts[v]} {v.lower()}"
            for v in ["CONFIRMED", "CORRECTED", "DISPUTED", "UNVERIFIED"]
            if v in verdicts
        ]
        para.add_run(", ".join(verdict_parts))



def _write_estimated_cost(doc: Document, estimated_cost: dict | None) -> None:
    """Render the conservative API-cost estimate for the run.

    Chunk 10 — users see "Estimated API Cost" alongside the severity
    table so they can size future runs without having to interpret
    raw token counts. The wording is intentionally conservative: this
    is the planning estimate, not the invoiced amount, and the
    disclaimer is rendered as italic gray context so the reader
    understands it is advisory.
    """
    if not estimated_cost:
        return

    from .cost_estimator import format_usd

    doc.add_heading("Estimated API Cost", level=1)
    if not estimated_cost.get("available"):
        para = doc.add_paragraph()
        run = para.add_run(
            "Estimated API cost is unavailable for this run — pricing data "
            "could not be matched to the models that were called."
        )
        run.font.size = Pt(10)
        run.font.italic = True
        missing = estimated_cost.get("missing_pricing_models") or []
        if missing:
            line = doc.add_paragraph()
            r = line.add_run("Models without pricing: " + ", ".join(missing))
            r.font.size = Pt(10)
            r.font.italic = True
            r.font.color.rgb = RGBColor(100, 100, 100)
        return

    total_para = doc.add_paragraph()
    label_run = total_para.add_run("Total Estimate: ")
    label_run.bold = True
    total_run = total_para.add_run(format_usd(estimated_cost.get("total_usd", 0.0)))
    total_run.bold = True
    total_run.font.size = Pt(14)

    disclaimer = doc.add_paragraph()
    d_run = disclaimer.add_run(
        f"Estimated API cost only — Anthropic's invoiced amount may differ. "
        f"Pricing snapshot: {estimated_cost.get('pricing_as_of', '')}."
    )
    d_run.font.size = Pt(9)
    d_run.font.italic = True
    d_run.font.color.rgb = RGBColor(100, 100, 100)

    by_phase = estimated_cost.get("by_phase") or {}
    if by_phase:
        sub = doc.add_paragraph()
        sub.add_run("By phase:").bold = True
        for phase_name, bucket in by_phase.items():
            line = doc.add_paragraph(style="List Bullet")
            line.add_run(f"{phase_name}: ").bold = True
            parts = [f"total {format_usd(bucket.get('total_usd', 0.0))}"]
            parts.append(f"input {format_usd(bucket.get('input_usd', 0.0))}")
            parts.append(f"output {format_usd(bucket.get('output_usd', 0.0))}")
            if bucket.get("cache_write_usd") or bucket.get("cache_read_usd"):
                parts.append(
                    "cache write "
                    f"{format_usd(bucket.get('cache_write_usd', 0.0))} / "
                    f"read {format_usd(bucket.get('cache_read_usd', 0.0))}"
                )
            if bucket.get("web_search_usd"):
                parts.append(
                    f"web search {format_usd(bucket.get('web_search_usd', 0.0))}"
                )
            line.add_run(" • ".join(parts))

    if estimated_cost.get("missing_pricing_calls"):
        warn = doc.add_paragraph()
        wrun = warn.add_run(
            f"Note: {estimated_cost['missing_pricing_calls']} call(s) on "
            f"unknown model(s) "
            f"({', '.join(estimated_cost.get('missing_pricing_models') or [])}) "
            f"are not included in the total above."
        )
        wrun.font.size = Pt(9)
        wrun.font.italic = True
        wrun.font.color.rgb = RGBColor(150, 100, 0)



def _write_trust_model_summary(
    doc: Document,
    status_counts: dict,
    edit_action_counts: dict,
) -> None:
    """Render the per-status / per-edit-action histograms.

    Chunk N Directive 1+4: every finding receives one status and one
    edit-action label. The table here gives a top-of-report at-a-glance
    picture of how much of the run is supported vs. uncertain vs.
    suppressed, before the reader gets to individual findings. The
    severity table above answers "how many issues are critical?"; this
    one answers "how many of them are actually trustworthy?"
    """
    total_status = sum(status_counts.values())
    if total_status == 0:
        return

    doc.add_heading("Trust Model Summary", level=1)
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Every finding is tagged with a trust status and an edit-action "
        "label. The two histograms below give the at-a-glance picture; "
        "individual findings carry the same labels inline."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(100, 100, 100)
    intro.paragraph_format.space_after = Pt(6)

    visible_statuses = [
        s for s in STATUS_DISPLAY_ORDER if status_counts.get(s, 0) > 0
    ]
    if visible_statuses:
        table = doc.add_table(rows=2, cols=len(visible_statuses))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col_idx, status in enumerate(visible_statuses):
            hex_color = STATUS_SHADING[status]
            header_cell = table.rows[0].cells[col_idx]
            _set_cell_shading(header_cell, hex_color)
            header_cell.text = ""
            p = header_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(status_label(status))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)

            count_cell = table.rows[1].cells[col_idx]
            count_cell.text = ""
            p = count_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(status_counts[status]))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = STATUS_COLORS[status]

    visible_actions = [
        a for a in EDIT_ACTION_DISPLAY_ORDER if edit_action_counts.get(a, 0) > 0
    ]
    if visible_actions:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        label_run = para.add_run("Edit eligibility: ")
        label_run.bold = True
        for i, action in enumerate(visible_actions):
            if i > 0:
                sep = para.add_run("  •  ")
                sep.font.color.rgb = RGBColor(160, 160, 160)
            count_run = para.add_run(f"{edit_action_counts[action]} ")
            count_run.bold = True
            count_run.font.color.rgb = EDIT_ACTION_COLORS[action]
            name_run = para.add_run(edit_action_label(action).lower())
            name_run.font.color.rgb = EDIT_ACTION_COLORS[action]
        para.paragraph_format.space_after = Pt(6)



def _write_alert_section(
    doc: Document,
    title: str,
    description: str,
    alerts: list[dict],
) -> None:
    """Render one named alert section with per-file bullet groups.

    Chunk O — every deterministic alert category goes through this helper
    so the section layout stays consistent. Each section also gets a
    "(deterministic check)" suffix so the reader can tell at a glance which
    alerts came from local rules vs. LLM findings — Chunk O Directive 2
    asks for deterministic findings to be clearly labeled.
    """
    if not alerts:
        return
    doc.add_heading(f"{title} (deterministic check)", level=2)
    _add_styled_paragraph(doc, description, size=10, space_after=6)
    by_file: dict[str, list[dict]] = {}
    for alert in alerts:
        by_file.setdefault(alert.get("filename", ""), []).append(alert)
    for filename, items in by_file.items():
        para = doc.add_paragraph()
        para.add_run(f"{filename}").bold = True
        for alert in items[:5]:
            context = alert.get("context", alert.get("match", ""))
            doc.add_paragraph(context, style='List Bullet')
        if len(items) > 5:
            doc.add_paragraph(
                f"... and {len(items) - 5} more",
                style='List Bullet',
            )


def _write_alerts(
    doc: Document,
    leed_alerts: list[dict],
    placeholder_alerts: list[dict],
    *,
    code_cycle_alerts: list[dict] | None = None,
    structural_alerts: list[dict] | None = None,
    naming_alerts: list[dict] | None = None,
    template_marker_alerts: list[dict] | None = None,
    invalid_code_cycle_alerts: list[dict] | None = None,
    duplicate_paragraph_alerts: list[dict] | None = None,
) -> None:
    """Write the Alerts section with every deterministic-check category.

    Chunk O — previously this rendered only ``leed_alerts`` and
    ``placeholder_alerts``; ``code_cycle_alerts`` / ``structural_alerts`` /
    ``naming_alerts`` were collected during preflight but silently dropped
    before the report saw them, so users had to read the log to discover
    them. The new optional kwargs default to ``None`` (treated as empty) so
    legacy callers — including older snapshot tests — keep working.
    """
    code_cycle_alerts = code_cycle_alerts or []
    structural_alerts = structural_alerts or []
    naming_alerts = naming_alerts or []
    template_marker_alerts = template_marker_alerts or []
    invalid_code_cycle_alerts = invalid_code_cycle_alerts or []
    duplicate_paragraph_alerts = duplicate_paragraph_alerts or []
    if not any((
        leed_alerts,
        placeholder_alerts,
        code_cycle_alerts,
        structural_alerts,
        naming_alerts,
        template_marker_alerts,
        invalid_code_cycle_alerts,
        duplicate_paragraph_alerts,
    )):
        return

    doc.add_heading("Alerts", level=1)
    _add_styled_paragraph(
        doc,
        "These items were detected locally by deterministic rules; no LLM "
        "tokens were spent on them. They sit alongside the LLM findings "
        "below.",
        size=10,
        space_after=6,
    )
    _write_alert_section(
        doc,
        "LEED References Detected",
        "The following LEED references were found. Since this is not a "
        "LEED project, these should be removed:",
        leed_alerts,
    )
    _write_alert_section(
        doc,
        "Unresolved Placeholders",
        "The following editorial placeholders need to be resolved:",
        placeholder_alerts,
    )
    _write_alert_section(
        doc,
        "Unresolved Template Markers",
        "The following TODO / FIXME / XXX / ??? markers are still in the "
        "spec and should be resolved before issuing:",
        template_marker_alerts,
    )
    _write_alert_section(
        doc,
        "Stale California Code Cycle References",
        "The following references cite a historical California code cycle "
        "rather than the one selected for this review:",
        code_cycle_alerts,
    )
    _write_alert_section(
        doc,
        "Invalid California Code Cycle Years",
        "The following references cite a year that is not a real California "
        "code cycle (California publishes cycles every 3 years: 2010, 2013, "
        "2016, 2019, 2022, 2025). These are likely typos:",
        invalid_code_cycle_alerts,
    )
    _write_alert_section(
        doc,
        "Structural Issues",
        "Empty sections and duplicate section headings detected in the "
        "spec body:",
        structural_alerts,
    )
    _write_alert_section(
        doc,
        "Duplicate Paragraphs",
        "These paragraphs appear verbatim more than once in the same spec — "
        "likely copy-paste mistakes:",
        duplicate_paragraph_alerts,
    )
    _write_alert_section(
        doc,
        "Inconsistent Filenames",
        "These files use a CSI naming style that differs from the project's "
        "dominant style:",
        naming_alerts,
    )



def _write_finding_entry(doc: Document, finding, index: int) -> None:
    """Write a single finding as a collapsible block.

    The finding header is rendered as a Heading 3 paragraph, which enables
    Word's native heading-collapse feature. Users can click the collapse
    triangle that appears on hover to hide the finding's body content.

    Chunk N changes:
        - Adds a "Status" line right under the header so the trust-model
          status is the first thing readers see (Directive 5: avoid
          presenting all findings as equally certain).
        - Adds an "Edit eligibility" line so readers can tell at a glance
          whether the finding is an auto-edit candidate, a manual edit
          candidate, report-only, or suppressed.
        - Renames the spec quote / web sources / rationale / rejected
          sources sub-labels so the four evidence concepts (Directive 3)
          are explicit rather than implied.

    Layout:
        Heading 3:  [SEVERITY] 92% — filename.docx — Section ref
        Normal:     Status: <status>  •  Edit: <edit_action>
        Normal:     Issue: ...
        Normal:     Action: ...
        Normal:     Spec evidence: ... (existingText, red)
        Normal:     Proposed replacement: ... (green)
        Normal:     Reference: ... (blue)
        Normal:     Verification verdict: ... (if applicable)
        Normal:     Verification rationale: ... (if applicable)
        Heading 4:  Sources (collapsed by default)
        Normal:       Web/code evidence: <accepted_sources>
        Normal:       Unsupported / rejected sources: <rejected_sources>
    """
    severity_color = SEVERITY_COLORS.get(finding.severity, RGBColor(0, 0, 0))
    conf_tier = _confidence_tier(finding.confidence)
    conf_color = CONFIDENCE_COLORS[conf_tier]

    para = doc.add_paragraph()
    para.style = doc.styles['Heading 3']
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


    run = para.add_run(f"{index}. [{finding.severity}] ")
    run.bold = True
    run.font.color.rgb = severity_color
    run.font.size = Pt(11)
    run = para.add_run(f"{finding.confidence:.0%} ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = conf_color
    run = para.add_run("— ")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(11)
    run = para.add_run(finding.fileName or "Unknown")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if finding.section:
        run = para.add_run(f" — {finding.section}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)


    status = classify_status(finding)
    edit_action = classify_edit_action(finding)
    status_color = STATUS_COLORS[status]
    action_color = EDIT_ACTION_COLORS[edit_action]

    status_para = doc.add_paragraph()
    status_label_run = status_para.add_run("Status: ")
    status_label_run.bold = True
    status_label_run.font.size = Pt(10)
    glyph_run = status_para.add_run(f"{status_glyph(status)} ")
    glyph_run.bold = True
    glyph_run.font.color.rgb = status_color
    glyph_run.font.size = Pt(10)
    value_run = status_para.add_run(status_label(status))
    value_run.bold = True
    value_run.font.color.rgb = status_color
    value_run.font.size = Pt(10)
    sep_run = status_para.add_run("  •  ")
    sep_run.font.color.rgb = RGBColor(160, 160, 160)
    sep_run.font.size = Pt(10)
    edit_label_run = status_para.add_run("Edit: ")
    edit_label_run.bold = True
    edit_label_run.font.size = Pt(10)
    edit_value_run = status_para.add_run(edit_action_label(edit_action))
    edit_value_run.bold = True
    edit_value_run.font.color.rgb = action_color
    edit_value_run.font.size = Pt(10)
    status_para.paragraph_format.space_after = Pt(3)

    para = doc.add_paragraph()
    para.add_run("Issue: ").bold = True
    para.add_run(finding.issue or "")
    para.paragraph_format.space_after = Pt(3)

    proposal = finding.as_edit_proposal()
    if proposal is None:
        para = doc.add_paragraph()
        run = para.add_run("Action: REPORT_ONLY")
        run.bold = True
        para.paragraph_format.space_after = Pt(3)

        demotion = (getattr(finding, "demotion_reason", None) or "").strip()
        if demotion:
            note_text = (
                "Edit proposal demoted to REPORT_ONLY at parse time: "
                f"{demotion}. The underlying finding is preserved; manual "
                "review required to determine a clean textual fix."
            )
        else:
            note_text = (
                "No edit proposal — surfaced for review only (coordination, "
                "interpretation, or multi-paragraph rewrite required)."
            )
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(note_text)
        note_run.font.italic = True
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(100, 100, 100)
        note_para.paragraph_format.space_after = Pt(3)
    else:
        para = doc.add_paragraph()
        para.add_run("Action: ").bold = True
        para.add_run(proposal.action_type or "")
        para.paragraph_format.space_after = Pt(3)

        if proposal.existing_text:
            para = doc.add_paragraph()
            para.add_run("Spec evidence: ").bold = True
            run = para.add_run(proposal.existing_text)
            run.font.color.rgb = RGBColor(192, 0, 0)
            para.paragraph_format.space_after = Pt(3)

        if proposal.replacement_text:
            para = doc.add_paragraph()
            para.add_run("Proposed replacement: ").bold = True
            run = para.add_run(proposal.replacement_text)
            run.font.color.rgb = RGBColor(0, 128, 0)
            para.paragraph_format.space_after = Pt(3)

    if finding.codeReference:
        para = doc.add_paragraph()
        para.add_run("Reference: ").bold = True
        run = para.add_run(finding.codeReference)
        run.font.color.rgb = RGBColor(59, 130, 246)
        para.paragraph_format.space_after = Pt(3)

    if finding.verification:
        vr = finding.verification
        verdict_color = VERDICT_COLORS.get(vr.verdict, VERDICT_COLORS["UNVERIFIED"])
        verdict_icon = VERDICT_ICONS.get(vr.verdict, "—")

        para = doc.add_paragraph()
        run = para.add_run(f"Verification verdict: {verdict_icon} {vr.verdict}")
        run.bold = True
        run.font.color.rgb = verdict_color
        para.paragraph_format.space_after = Pt(3)

        if vr.explanation:
            para = doc.add_paragraph()
            label_run = para.add_run("Verification rationale: ")
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(100, 100, 100)
            body_run = para.add_run(vr.explanation)
            body_run.font.size = Pt(10)
            body_run.font.color.rgb = RGBColor(100, 100, 100)
            para.paragraph_format.space_after = Pt(3)

        if vr.verdict == "CORRECTED" and vr.correction:
            para = doc.add_paragraph()
            para.add_run("Correction: ").bold = True
            run = para.add_run(vr.correction)
            run.font.color.rgb = RGBColor(204, 132, 0)
            para.paragraph_format.space_after = Pt(3)

        accepted = list(vr.sources or [])
        rejected = list(getattr(vr, "rejected_sources", []) or [])
        if accepted or rejected:
            sources_heading = doc.add_heading("Sources", level=4)
            _set_paragraph_collapsed(sources_heading)

            if accepted:
                label_para = doc.add_paragraph()
                _set_paragraph_outline_level(label_para, 8)
                label_run = label_para.add_run("Web/code evidence (cited and found in search results):")
                label_run.font.size = Pt(9)
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(0, 100, 0)

                para = doc.add_paragraph()
                _set_paragraph_outline_level(para, 8)
                para.paragraph_format.space_after = Pt(3)
                for i, url in enumerate(accepted):
                    if i > 0:
                        para.add_run("  •  ")
                    run = para.add_run(url)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(59, 130, 246)

            if rejected:
                label_para = doc.add_paragraph()
                _set_paragraph_outline_level(label_para, 8)
                label_run = label_para.add_run(
                    "Unsupported / rejected sources (cited by the model but not present in web_search results):"
                )
                label_run.font.size = Pt(9)
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(192, 0, 0)

                rej_para = doc.add_paragraph()
                _set_paragraph_outline_level(rej_para, 8)
                rej_para.paragraph_format.space_after = Pt(3)
                for i, entry in enumerate(rejected):
                    if i > 0:
                        rej_para.add_run("  •  ")
                    url = entry.get("url") if isinstance(entry, dict) else str(entry)
                    reason = entry.get("reason") if isinstance(entry, dict) else ""
                    url_run = rej_para.add_run(url or "(empty)")
                    url_run.font.size = Pt(9)
                    url_run.font.color.rgb = RGBColor(128, 128, 128)
                    url_run.font.italic = True
                    if reason:
                        reason_run = rej_para.add_run(f" [{reason}]")
                        reason_run.font.size = Pt(9)
                        reason_run.font.color.rgb = RGBColor(192, 0, 0)



def _write_findings_section(doc: Document, review) -> None:
    """Write per-spec findings grouped by severity, sorted by confidence.

    Uses heading hierarchy for Word-native collapse support:
    - Title (level 0): "Findings"
    - Heading 1: Severity group (e.g., "CRITICAL (1)")
    - Heading 3: Individual finding header (collapsible)
    - Normal: Finding body content

    """
    doc.add_heading("Findings", level=0)

    if review.total_count == 0:
        _add_styled_paragraph(
            doc,
            "No issues found.",
            size=12,
            color=RGBColor(0, 128, 0),
        )
        return

    finding_number = 0

    for severity in SEVERITY_ORDER:
        severity_findings = sorted(
            [f for f in review.findings if f.severity == severity],
            key=lambda f: f.confidence,
            reverse=True,
        )
        if not severity_findings:
            continue

        heading = doc.add_heading(
            f"{severity} ({len(severity_findings)})", level=1,
        )
        for run in heading.runs:
            run.font.color.rgb = SEVERITY_COLORS.get(severity, RGBColor(0, 0, 0))

        for finding in severity_findings:
            finding_number += 1
            _write_finding_entry(doc, finding, finding_number)



def _write_dependency_note(
    doc: Document,
    finding,
    upstream_lookup: dict,
) -> None:
    """Render the Chunk M dependency annotation for a kept cross-check finding.

    When a cross-check finding cites ``upstream_finding_ids`` or
    ``independent_evidence_ids``, surface that information in the report so
    a reviewer can trace the coordination claim back to its sources without
    having to inspect the raw structured payload. The cited review findings
    are looked up in ``upstream_lookup`` (finding_id → review Finding); when
    a cited id is unknown — for example, a stale resume payload referencing
    an id from a prior run — it is rendered verbatim so the gap is visible.
    """
    upstream_ids = [uid for uid in (getattr(finding, "upstream_finding_ids", []) or []) if uid]
    independent_ids = [
        eid for eid in (getattr(finding, "independent_evidence_ids", []) or []) if eid
    ]
    if not upstream_ids and not independent_ids:
        return

    if upstream_ids:
        para = doc.add_paragraph()
        run = para.add_run("Depends on review finding(s): ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        for i, uid in enumerate(upstream_ids):
            if i > 0:
                sep = para.add_run("; ")
                sep.font.size = Pt(10)
                sep.font.color.rgb = RGBColor(100, 100, 100)
            upstream = upstream_lookup.get(uid)
            if upstream is None:
                label = f"[{uid}] (not found in current review)"
            else:
                verdict = ""
                if upstream.verification and upstream.verification.verdict:
                    verdict = f" — {upstream.verification.verdict}"
                file_section = " — ".join(
                    p for p in [upstream.fileName, upstream.section] if p
                )
                label = f"[{upstream.severity}] {file_section}{verdict}"
            entry = para.add_run(label)
            entry.font.size = Pt(10)
            entry.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(3)

    if independent_ids:
        para = doc.add_paragraph()
        run = para.add_run("Independent spec evidence: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        entry = para.add_run(", ".join(independent_ids))
        entry.font.size = Pt(10)
        entry.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(3)


def _write_suppression_reason(doc: Document, finding) -> None:
    """Render the Chunk M suppression reason underneath a finding entry."""
    reason = getattr(finding, "suppression_reason", None)
    if not reason:
        return
    para = doc.add_paragraph()
    label = para.add_run("Suppressed: ")
    label.bold = True
    label.font.size = Pt(10)
    label.font.color.rgb = RGBColor(192, 0, 0)
    body = para.add_run(reason)
    body.font.size = Pt(10)
    body.font.italic = True
    body.font.color.rgb = RGBColor(100, 100, 100)
    para.paragraph_format.space_after = Pt(3)


def _write_cross_check_section(doc: Document, cross_check_result, review_result=None) -> None:
    """Write cross-spec coordination section and explicit status.

    Cross-check findings are rendered with the same collapsible structure
    as per-spec findings.

    Chunk M: kept findings that cite ``upstream_finding_ids`` get a
    "Depends on review finding(s)" annotation so readers can trace the
    coordination claim. Findings dropped by the suppression filter are
    rendered under a dedicated "Suppressed Coordination Findings" sub-
    heading along with the recorded reason, so the decision is visible
    instead of silently making the finding disappear.
    """
    if not cross_check_result:
        return

    doc.add_page_break()

    heading = doc.add_heading("Cross-Spec Coordination", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    status = getattr(cross_check_result, "cross_check_status", None)
    count = len(cross_check_result.findings)
    suppressed = list(getattr(cross_check_result, "suppressed_findings", []) or [])
    subtitle = doc.add_paragraph()
    if status == "skipped":
        run = subtitle.add_run(f"Cross-check was skipped: {cross_check_result.thinking}")
    elif status == "failed":
        run = subtitle.add_run(f"Cross-check failed: {cross_check_result.error}")
    elif status == "completed" and count == 0 and not suppressed:
        run = subtitle.add_run("Cross-check completed — no coordination issues found.")
    else:
        suppressed_note = ""
        if suppressed:
            suppressed_note = (
                f" ({len(suppressed)} suppressed by upstream-disputed filter)"
            )
        run = subtitle.add_run(
            f"Sonnet 4.6 coordination analysis — "
            f"{count} issue{'s' if count != 1 else ''} found{suppressed_note}."
        )
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle.paragraph_format.space_after = Pt(12)

    if status in ("skipped", "failed"):
        return

    upstream_lookup: dict = {}
    if review_result is not None:
        for f in getattr(review_result, "findings", []) or []:
            fid = getattr(f, "finding_id", "")
            if fid:
                upstream_lookup[fid] = f

    severity_rank = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "GRIPES": 3}
    sorted_findings = sorted(
        cross_check_result.findings,
        key=lambda f: (severity_rank.get(f.severity, 99), -f.confidence),
    )

    for idx, finding in enumerate(sorted_findings, 1):
        _write_finding_entry(doc, finding, idx)
        _write_dependency_note(doc, finding, upstream_lookup)

    if cross_check_result.thinking:
        doc.add_heading("Coordination Summary", level=2)
        _write_narrative_text(doc, cross_check_result.thinking)

    if suppressed:
        doc.add_heading("Suppressed Coordination Findings", level=2)
        intro = doc.add_paragraph()
        intro_run = intro.add_run(
            "The findings below were dropped by the upstream-disputed filter "
            "after the per-spec review verification verdicts came in. They "
            "are shown for traceability — they were not re-verified."
        )
        intro_run.font.size = Pt(10)
        intro_run.font.italic = True
        intro_run.font.color.rgb = RGBColor(100, 100, 100)
        intro.paragraph_format.space_after = Pt(8)

        sorted_suppressed = sorted(
            suppressed,
            key=lambda f: (severity_rank.get(f.severity, 99), -f.confidence),
        )
        for idx, finding in enumerate(sorted_suppressed, 1):
            _write_finding_entry(doc, finding, idx)
            _write_suppression_reason(doc, finding)


def _sanitize_markdown_line(line: str) -> str:
    """Strip markdown header prefixes from a line."""
    stripped = line
    while stripped.startswith('#'):
        stripped = stripped[1:]
    return stripped.strip() if line.startswith('#') else line


def _write_narrative_text(doc: Document, text: str) -> None:
    """Write multi-paragraph narrative text, splitting on double newlines.

    Strips markdown header formatting (## ...) since the cross-check
    prompt sometimes produces markdown despite instructions not to.
    """
    if '\n\n' in text:
        paragraphs = text.split('\n\n')
    else:
        paragraphs = text.split('\n')

    for para_text in paragraphs:
        para_text = _sanitize_markdown_line(para_text.strip())
        if para_text:
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.size = Pt(11)
            para.paragraph_format.space_after = Pt(8)



def export_report(
    pipeline_result,
    output_path: Path,
    *,
    estimated_cost: dict | None = None,
) -> Path:
    """Export a complete review report to a Word document.

    Generates a formatted .docx file containing files reviewed, summary
    grid, alerts, per-spec findings, and cross-check findings.

    Each per-spec finding uses Heading 3 for its header line, enabling
    Word's native heading-collapse feature.

    Args:
        pipeline_result: PipelineResult from the review pipeline
        output_path: Path where the .docx file should be saved
        estimated_cost: Chunk 10 — when supplied, render the
            "Estimated API Cost" section after the summary table.
            Pass the ``estimated_cost`` value from
            :meth:`DiagnosticsReport.summary`. Omit (or pass ``None``)
            to skip the section gracefully — useful for legacy callers
            and tests that build a stub result without a diagnostics
            report.

    Returns:
        The output_path (for convenience / confirmation)

    Raises:
        ValueError: If pipeline_result has no review_result
        OSError: If the file cannot be written
    """
    if pipeline_result.review_result is None:
        raise ValueError("Cannot export report: no review results available")

    review = pipeline_result.review_result
    cross_check = pipeline_result.cross_check_result
    all_findings = list(review.findings)
    if cross_check and cross_check.findings:
        all_findings.extend(cross_check.findings)
    if cross_check and getattr(cross_check, "suppressed_findings", None):
        all_findings.extend(cross_check.suppressed_findings)
    verification_stats = _summarize_verification_outcomes(all_findings)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(11)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(4)
    h3_style.font.color.rgb = RGBColor(0, 0, 0)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    cycle_label = getattr(pipeline_result, "cycle_label", "2025") or "2025"
    _write_title_block(
        doc,
        review,
        pipeline_result.files_reviewed,
        cycle_label=cycle_label,
    )
    
    _write_files_reviewed(doc, pipeline_result.files_reviewed)
    cross_check_status = getattr(cross_check, "cross_check_status", None) if cross_check else None
    cross_check_reason = ""
    if cross_check and cross_check_status == "skipped":
        cross_check_reason = cross_check.thinking or ""
    elif cross_check and cross_check_status == "failed":
        cross_check_reason = cross_check.error or ""
    _write_methodology_note(
        doc,
        cross_check_enabled=(cross_check is not None),
        cycle_label=cycle_label,
        cross_check_status=cross_check_status,
        cross_check_reason=cross_check_reason,
        verification_stats=verification_stats,
    )
    _write_summary_table(
        doc,
        review,
        cross_check,
        total_elapsed_seconds=getattr(pipeline_result, "total_elapsed_seconds", None),
    )

    _write_estimated_cost(doc, estimated_cost)

    _write_trust_model_summary(
        doc,
        verification_stats.get("status_counts", {}),
        verification_stats.get("edit_action_counts", {}),
    )

    _write_alerts(
        doc,
        pipeline_result.leed_alerts,
        pipeline_result.placeholder_alerts,
        code_cycle_alerts=getattr(pipeline_result, "code_cycle_alerts", None),
        structural_alerts=getattr(pipeline_result, "structural_alerts", None),
        naming_alerts=getattr(pipeline_result, "naming_alerts", None),
        template_marker_alerts=getattr(pipeline_result, "template_marker_alerts", None),
        invalid_code_cycle_alerts=getattr(pipeline_result, "invalid_code_cycle_alerts", None),
        duplicate_paragraph_alerts=getattr(pipeline_result, "duplicate_paragraph_alerts", None),
    )
    _write_findings_section(doc, review)
    _write_cross_check_section(doc, cross_check, review_result=review)


    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return output_path