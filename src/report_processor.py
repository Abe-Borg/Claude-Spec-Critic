"""Process rejection checkboxes in exported Spec Critic Word reports.

Workflow:
    1. User exports a report via 'Export Report' — each per-spec finding
       gets a clickable w14 checkbox in its heading row.
    2. User opens the .docx in Word (2013+), reviews findings, and checks
       the checkbox next to any finding they want to reject.
    3. User returns to Spec Critic and clicks 'Process Rejections'.
    4. This module reads the checked state of each checkbox, moves rejected
       findings to an appendix, strips all checkboxes from the output, and
       saves a clean triaged report.

The checkbox is a w14 structured document tag (SDT) content control.
When clicked in Word, it toggles between ☐ (unchecked) and ☒ (checked).
This module reads the w14:checked XML attribute to determine state, with
a fallback to checking the display character for compatibility.

Finding identification:
    Each checkbox SDT is tagged with 'reject_f001', 'reject_f002', etc.
    The tag maps 1:1 to the finding's sequential number in the report.
    A finding's paragraph range starts at its Heading 3 paragraph (which
    contains the checkbox) and extends through all subsequent Normal
    paragraphs until the next Heading 3 or higher-level heading.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lxml import etree


# ---------------------------------------------------------------------------
# XML namespace constants (must match report_exporter.py)
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

REJECT_TAG_PREFIX = "reject_f"
INSTRUCTIONS_TAG = "sc_rejection_instructions"

# Heading style IDs that mark the boundary of a finding's paragraph range.
# If we hit any of these while scanning forward from a Heading3, the finding ends.
_BOUNDARY_STYLES = {"Heading1", "Heading2", "Heading0", "Title"}


# ---------------------------------------------------------------------------
# Result dataclass
# ---------------------------------------------------------------------------

@dataclass
class ProcessResult:
    """Result of processing rejections in a report."""
    total_findings: int
    rejected_count: int
    kept_count: int
    output_path: Path


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _get_style_id(p_elem: etree._Element) -> str:
    """Get the style ID (e.g. 'Heading3') from a <w:p> element."""
    pPr = p_elem.find(f"{{{W_NS}}}pPr")
    if pPr is None:
        return ""
    pStyle = pPr.find(f"{{{W_NS}}}pStyle")
    if pStyle is None:
        return ""
    return pStyle.get(f"{{{W_NS}}}val", "")


def _find_reject_tag(p_elem: etree._Element) -> str | None:
    """Find a reject_f* tag in any SDT nested inside this paragraph."""
    for sdt in p_elem.iter(f"{{{W_NS}}}sdt"):
        sdtPr = sdt.find(f"{{{W_NS}}}sdtPr")
        if sdtPr is None:
            continue
        tag_elem = sdtPr.find(f"{{{W_NS}}}tag")
        if tag_elem is None:
            continue
        tag_val = tag_elem.get(f"{{{W_NS}}}val", "")
        if tag_val.startswith(REJECT_TAG_PREFIX):
            return tag_val
    return None


def _is_checkbox_checked(p_elem: etree._Element, tag: str) -> bool:
    """Check if a rejection checkbox with the given tag is checked.

    Primary check: reads the w14:checked/@w14:val attribute.
    Fallback: checks if the SDT content text contains ☒ (U+2612),
    which handles cases where the XML attribute is missing but the
    display character was toggled.
    """
    for sdt in p_elem.iter(f"{{{W_NS}}}sdt"):
        sdtPr = sdt.find(f"{{{W_NS}}}sdtPr")
        if sdtPr is None:
            continue
        tag_elem = sdtPr.find(f"{{{W_NS}}}tag")
        if tag_elem is None:
            continue
        if tag_elem.get(f"{{{W_NS}}}val", "") != tag:
            continue
        # Primary: read w14:checkbox > w14:checked
        checkbox = sdtPr.find(f"{{{W14_NS}}}checkbox")
        if checkbox is not None:
            checked_elem = checkbox.find(f"{{{W14_NS}}}checked")
            if checked_elem is not None:
                return checked_elem.get(f"{{{W14_NS}}}val", "0") == "1"
        # Fallback: check display character
        content = sdt.find(f"{{{W_NS}}}sdtContent")
        if content is not None:
            for t_elem in content.iter(f"{{{W_NS}}}t"):
                if t_elem.text and "\u2612" in t_elem.text:
                    return True
        return False
    return False


def _strip_reject_sdts(elem: etree._Element) -> None:
    """Remove all rejection checkbox SDTs from an element tree.

    Walks the element and its descendants, removing any <w:sdt> whose
    tag starts with the rejection prefix. This cleans up the checkboxes
    from both kept and rejected findings in the triaged output.
    """
    for sdt in list(elem.iter(f"{{{W_NS}}}sdt")):
        sdtPr = sdt.find(f"{{{W_NS}}}sdtPr")
        if sdtPr is None:
            continue
        tag_elem = sdtPr.find(f"{{{W_NS}}}tag")
        if tag_elem is None:
            continue
        if tag_elem.get(f"{{{W_NS}}}val", "").startswith(REJECT_TAG_PREFIX):
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)


def _find_instructions_paragraph(body: etree._Element) -> etree._Element | None:
    """Find the rejection instructions paragraph by its SDT tag."""
    for p_elem in body.iter(f"{{{W_NS}}}p"):
        for sdt in p_elem.iter(f"{{{W_NS}}}sdt"):
            sdtPr = sdt.find(f"{{{W_NS}}}sdtPr")
            if sdtPr is None:
                continue
            tag_elem = sdtPr.find(f"{{{W_NS}}}tag")
            if tag_elem is None:
                continue
            if tag_elem.get(f"{{{W_NS}}}val", "") == INSTRUCTIONS_TAG:
                return p_elem
    return None


# ---------------------------------------------------------------------------
# Core processing
# ---------------------------------------------------------------------------

@dataclass
class _FindingRange:
    """A contiguous range of XML elements belonging to one finding."""
    tag: str
    is_checked: bool
    elements: list[etree._Element]


def _identify_finding_ranges(body: etree._Element) -> list[_FindingRange]:
    """Walk the document body and identify paragraph ranges for each finding.

    A finding's range starts at a Heading3 paragraph that contains a
    rejection checkbox SDT and extends through all subsequent paragraphs
    (and any interleaved non-paragraph elements like tables) until the
    next Heading3 or a higher-level heading boundary.
    """
    children = list(body)
    ranges: list[_FindingRange] = []
    current: _FindingRange | None = None

    for child in children:
        is_paragraph = child.tag == f"{{{W_NS}}}p"

        if is_paragraph:
            style = _get_style_id(child)

            # Check if this paragraph is a finding header (Heading3 with checkbox)
            if style == "Heading3":
                tag = _find_reject_tag(child)
                if tag is not None:
                    # Close previous range
                    if current is not None:
                        ranges.append(current)
                    # Start new range
                    checked = _is_checkbox_checked(child, tag)
                    current = _FindingRange(tag=tag, is_checked=checked, elements=[child])
                    continue
                else:
                    # Heading3 without a checkbox — could be a cross-check finding
                    # or other non-rejectable heading. Closes the current range.
                    if current is not None:
                        ranges.append(current)
                        current = None
                    continue

            # A higher-level heading closes the current range
            if style in _BOUNDARY_STYLES:
                if current is not None:
                    ranges.append(current)
                    current = None
                continue

        # Body content within a finding range
        if current is not None:
            current.elements.append(child)

    # Close the final range
    if current is not None:
        ranges.append(current)

    return ranges


def process_rejections(
    input_path: Path,
    output_path: Path | None = None,
) -> ProcessResult:
    """Read a Spec Critic report .docx, move rejected findings to an appendix.

    Args:
        input_path: Path to the exported report with checkboxes.
        output_path: Where to save the triaged report. Defaults to
            ``{input_stem}-triaged.docx`` alongside the input file.

    Returns:
        ProcessResult with counts and the output path.

    Raises:
        FileNotFoundError: If input_path does not exist.
        ValueError: If no rejection checkboxes are found (not a Spec Critic report).
    """
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"Report not found: {input_path}")
    if output_path is None:
        output_path = input_path.with_stem(input_path.stem + "-triaged")
    output_path = Path(output_path)

    doc = Document(str(input_path))
    body = doc.element.body

    # Identify finding ranges
    finding_ranges = _identify_finding_ranges(body)
    total_findings = len(finding_ranges)

    if total_findings == 0:
        raise ValueError(
            "No rejection checkboxes found in this document. "
            "Make sure you are processing a Spec Critic report exported with "
            "the 'Export Report' output mode."
        )

    rejected_ranges = [r for r in finding_ranges if r.is_checked]
    rejected_count = len(rejected_ranges)

    # Strip all checkboxes from the entire body (kept + rejected findings)
    _strip_reject_sdts(body)

    # Remove the rejection instructions paragraph
    instructions_p = _find_instructions_paragraph(body)
    if instructions_p is not None:
        body.remove(instructions_p)

    if rejected_count == 0:
        # No rejections — just save without checkboxes
        doc.save(str(output_path))
        return ProcessResult(
            total_findings=total_findings,
            rejected_count=0,
            kept_count=total_findings,
            output_path=output_path,
        )

    # Remove rejected finding elements from their current position
    # (iterate in reverse to preserve element indices during removal)
    collected_rejected: list[list[etree._Element]] = []
    for r in reversed(rejected_ranges):
        for elem in r.elements:
            body.remove(elem)
        collected_rejected.insert(0, r.elements)

    # Build the appendix using python-docx's high-level API (appends to end)
    doc.add_page_break()
    appendix_heading = doc.add_heading(
        f"Appendix: Rejected Findings ({rejected_count})", level=0
    )
    for run in appendix_heading.runs:
        run.font.color.rgb = RGBColor(128, 128, 128)

    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        f"{rejected_count} finding{'s' if rejected_count != 1 else ''} "
        f"rejected by the reviewer during triage. These findings were assessed "
        f"and determined to be not applicable, already addressed, or otherwise "
        f"not actionable for this project."
    )
    note_run.font.size = Pt(11)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(128, 128, 128)
    note_para.paragraph_format.space_after = Pt(12)

    # Re-insert rejected finding elements into the appendix
    for elements in collected_rejected:
        for elem in elements:
            body.append(elem)

    # Insert a triage note after the Summary heading so it's visible at the top
    _insert_triage_note(body, rejected_count, total_findings)

    doc.save(str(output_path))

    return ProcessResult(
        total_findings=total_findings,
        rejected_count=rejected_count,
        kept_count=total_findings - rejected_count,
        output_path=output_path,
    )


def _insert_triage_note(
    body: etree._Element,
    rejected_count: int,
    total_findings: int,
) -> None:
    """Insert a triage note paragraph after the 'Summary' heading.

    Scans for a Heading1 element whose text contains 'Summary' and inserts
    a styled note paragraph immediately after it. If no Summary heading is
    found, the note is skipped (non-critical).
    """
    kept = total_findings - rejected_count
    note_text = (
        f"Triage: {rejected_count} of {total_findings} findings were rejected "
        f"by the reviewer ({kept} findings retained). "
        f"Rejected findings are in the appendix at the end of this document."
    )

    # Find the Summary heading
    target: etree._Element | None = None
    for p_elem in body.iter(f"{{{W_NS}}}p"):
        if _get_style_id(p_elem) == "Heading1":
            # Check if its text content contains "Summary"
            text = "".join(
                t.text or "" for t in p_elem.iter(f"{{{W_NS}}}t")
            )
            if "Summary" in text:
                target = p_elem
                break

    if target is None:
        return

    # Build a note paragraph element
    note_p = etree.SubElement(body, f"{{{W_NS}}}p")
    pPr = etree.SubElement(note_p, f"{{{W_NS}}}pPr")
    spacing = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
    spacing.set(f"{{{W_NS}}}after", "120")  # 6pt after
    run = etree.SubElement(note_p, f"{{{W_NS}}}r")
    rPr = etree.SubElement(run, f"{{{W_NS}}}rPr")
    i_elem = etree.SubElement(rPr, f"{{{W_NS}}}i")
    color_elem = etree.SubElement(rPr, f"{{{W_NS}}}color")
    color_elem.set(f"{{{W_NS}}}val", "F59E0B")  # Warning amber
    sz = etree.SubElement(rPr, f"{{{W_NS}}}sz")
    sz.set(f"{{{W_NS}}}val", "22")  # 11pt
    t_elem = etree.SubElement(run, f"{{{W_NS}}}t")
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t_elem.text = note_text

    # Move note_p to right after the Summary heading
    # (SubElement appended it at the end; we need to reposition it)
    body.remove(note_p)
    target_index = list(body).index(target)
    body.insert(target_index + 1, note_p)