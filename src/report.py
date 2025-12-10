"""Word report generation module."""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .reviewer import ReviewResult, Finding


# Severity colors (RGB)
SEVERITY_COLORS = {
    "CRITICAL": RGBColor(192, 0, 0),      # Dark red
    "HIGH": RGBColor(255, 102, 0),         # Orange
    "MEDIUM": RGBColor(192, 152, 0),       # Dark yellow/gold
    "LOW": RGBColor(0, 112, 192),          # Blue
    "GRIPES": RGBColor(128, 0, 128),       # Purple
}

SEVERITY_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "GRIPES"]


def set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc: Document, text: str, style: str = None, bold: bool = False, 
                         color: RGBColor = None, size: int = None, space_after: int = None):
    """Add a paragraph with optional styling."""
    para = doc.add_paragraph()
    if style:
        para.style = style
    
    run = para.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    
    return para


def create_summary_table(doc: Document, review_result: ReviewResult):
    """Create the summary counts table."""
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "GRIPES", "TOTAL"]
    counts = [
        review_result.critical_count,
        review_result.high_count,
        review_result.medium_count,
        review_result.low_count,
        review_result.gripes_count,
        review_result.total_count
    ]
    
    for i, (header, count) in enumerate(zip(headers, counts)):
        cell = header_cells[i]
        cell.text = f"{header}\n{count}"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Color the header based on severity
        if header in SEVERITY_COLORS:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
        elif header == "TOTAL":
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)


def create_alerts_section(doc: Document, leed_alerts: list, placeholder_alerts: list):
    """Create the alerts section for LEED and placeholders."""
    if not leed_alerts and not placeholder_alerts:
        return
    
    doc.add_heading("Alerts", level=1)
    
    if leed_alerts:
        doc.add_heading("LEED References Detected", level=2)
        add_styled_paragraph(
            doc, 
            "The following LEED references were found. Since this is not a LEED project, these should be removed:",
            size=10,
            space_after=6
        )
        
        # Group by filename
        by_file = {}
        for alert in leed_alerts:
            fname = alert['filename']
            if fname not in by_file:
                by_file[fname] = []
            by_file[fname].append(alert)
        
        for filename, alerts in by_file.items():
            para = doc.add_paragraph()
            para.add_run(f"{filename}").bold = True
            for alert in alerts[:5]:  # Limit to 5 per file
                doc.add_paragraph(
                    f"Line {alert['line']}: {alert['text']}", 
                    style='List Bullet'
                )
            if len(alerts) > 5:
                doc.add_paragraph(
                    f"... and {len(alerts) - 5} more",
                    style='List Bullet'
                )
    
    if placeholder_alerts:
        doc.add_heading("Unresolved Placeholders", level=2)
        add_styled_paragraph(
            doc,
            "The following placeholders need to be resolved:",
            size=10,
            space_after=6
        )
        
        # Group by filename
        by_file = {}
        for alert in placeholder_alerts:
            fname = alert['filename']
            if fname not in by_file:
                by_file[fname] = []
            by_file[fname].append(alert)
        
        for filename, alerts in by_file.items():
            para = doc.add_paragraph()
            para.add_run(f"{filename}").bold = True
            for alert in alerts[:5]:
                doc.add_paragraph(
                    f"Line {alert['line']}: {alert['text']}",
                    style='List Bullet'
                )
            if len(alerts) > 5:
                doc.add_paragraph(
                    f"... and {len(alerts) - 5} more",
                    style='List Bullet'
                )


def create_finding_entry(doc: Document, finding: Finding, index: int):
    """Create a single finding entry."""
    # Finding header with severity color
    severity_color = SEVERITY_COLORS.get(finding.severity, RGBColor(0, 0, 0))
    
    para = doc.add_paragraph()
    run = para.add_run(f"{index}. [{finding.severity}] ")
    run.bold = True
    run.font.color.rgb = severity_color
    
    run = para.add_run(f"{finding.fileName}")
    run.bold = True
    
    # Section
    if finding.section:
        para = doc.add_paragraph()
        para.add_run("Section: ").bold = True
        para.add_run(finding.section)
        para.paragraph_format.space_after = Pt(3)
    
    # Issue description
    para = doc.add_paragraph()
    para.add_run("Issue: ").bold = True
    para.add_run(finding.issue)
    para.paragraph_format.space_after = Pt(3)
    
    # Action type
    para = doc.add_paragraph()
    para.add_run("Action: ").bold = True
    para.add_run(finding.actionType)
    para.paragraph_format.space_after = Pt(3)
    
    # Existing text (if applicable)
    if finding.existingText:
        para = doc.add_paragraph()
        para.add_run("Existing Text: ").bold = True
        run = para.add_run(finding.existingText)
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        para.paragraph_format.space_after = Pt(3)
    
    # Replacement text (if applicable)
    if finding.replacementText:
        para = doc.add_paragraph()
        para.add_run("Replace With: ").bold = True
        run = para.add_run(finding.replacementText)
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        para.paragraph_format.space_after = Pt(3)
    
    # Code reference (if applicable)
    if finding.codeReference:
        para = doc.add_paragraph()
        para.add_run("Reference: ").bold = True
        para.add_run(finding.codeReference)
        para.paragraph_format.space_after = Pt(3)
    
    # Add spacing after finding
    doc.add_paragraph()


def create_findings_section(doc: Document, findings: list[Finding], severity: str):
    """Create a section for findings of a specific severity."""
    severity_findings = [f for f in findings if f.severity == severity]
    
    if not severity_findings:
        return
    
    # Section heading with color
    heading = doc.add_heading(f"{severity} ({len(severity_findings)})", level=1)
    for run in heading.runs:
        run.font.color.rgb = SEVERITY_COLORS.get(severity, RGBColor(0, 0, 0))
    
    # Add each finding
    for i, finding in enumerate(severity_findings, 1):
        create_finding_entry(doc, finding, i)


def generate_report(
    review_result: ReviewResult,
    files_reviewed: list[str],
    leed_alerts: list,
    placeholder_alerts: list,
    output_path: Path
) -> Path:
    """
    Generate a Word document report from review results.
    
    Args:
        review_result: The ReviewResult from the API
        files_reviewed: List of filenames that were reviewed
        leed_alerts: List of LEED alert dicts
        placeholder_alerts: List of placeholder alert dicts
        output_path: Path to save the report
        
    Returns:
        Path to the generated report
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading("M&P Specification Review Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    para.add_run(f"Model: {review_result.model}\n")
    para.add_run(f"Files Reviewed: {len(files_reviewed)}")
    
    # Files list
    doc.add_heading("Files Reviewed", level=1)
    for filename in files_reviewed:
        doc.add_paragraph(filename, style='List Bullet')
    
    # Summary
    doc.add_heading("Summary", level=1)
    create_summary_table(doc, review_result)
    doc.add_paragraph()  # Spacing
    
    # Token usage
    para = doc.add_paragraph()
    para.add_run("Token Usage: ").bold = True
    if review_result.thinking_tokens > 0:
        para.add_run(
            f"{review_result.input_tokens:,} input → "
            f"{review_result.thinking_tokens:,} thinking + "
            f"{review_result.output_tokens:,} output"
        )
    else:
        para.add_run(
            f"{review_result.input_tokens:,} input → "
            f"{review_result.output_tokens:,} output"
        )
    
    para = doc.add_paragraph()
    para.add_run("Processing Time: ").bold = True
    para.add_run(f"{review_result.elapsed_seconds:.1f} seconds")
    
    # Alerts section
    create_alerts_section(doc, leed_alerts, placeholder_alerts)
    
    # Findings by severity
    doc.add_heading("Findings", level=0)
    
    if review_result.total_count == 0:
        add_styled_paragraph(
            doc,
            "No issues found.",
            size=12,
            color=RGBColor(0, 128, 0)
        )
    else:
        for severity in SEVERITY_ORDER:
            create_findings_section(doc, review_result.findings, severity)
    
    # Save document
    report_path = output_path / "report.docx"
    doc.save(report_path)
    
    return report_path
