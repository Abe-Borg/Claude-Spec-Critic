"""
Word document report exporter for Spec Critic.

Generates a formatted .docx report from a PipelineResult:
    - Title block with generation metadata
    - Files reviewed list
    - Summary table (severity counts) with colored cell shading
    - LEED and placeholder alerts
    - Per-spec findings grouped by severity, then by spec file, then confidence
    - Verification verdicts and corrections inline with findings
    - Cross-spec coordination findings (if cross-check was enabled)

Collapsible findings (v2.5.0):
    Each finding header uses Word Heading 3 style. In Word 2016+ and
    365, hovering over any heading shows a collapse triangle. Clicking
    it hides everything between that heading and the next heading of
    the same or higher level. This means:
    - Collapse a severity heading (Heading 1) to hide all its findings
    - Collapse a single finding heading (Heading 3) to hide its details
    No macros or special XML required — this is native Word behavior.

Usage:
    from report_exporter import export_report

    export_report(
        pipeline_result=result,
        output_path=Path("review-report.docx"),
    )
"""

from __future__ import annotations

import time
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from ..core.api_config import (
    CROSS_CHECK_MODEL_DEFAULT,
    web_search_max_uses_for_severity,
)
from ..core.pricing import price_for
from ..core.code_cycles import CodeCycle
from ..core.project_profile import ProjectProfile
from ..modules import ReviewModule, get_module
from ..research import (
    PROFILE_CATEGORY_SECTIONS,
    PROFILE_SECTION_ORDER,
    RequirementsProfile,
)
from ..verification.verification_cache import default_cache_path
from .report_status import (
    EDIT_ACTION_DISPLAY_ORDER,
    EditActionLabel,
    ReportStatus,
    STATUS_DISPLAY_ORDER,
    classify_edit_action,
    classify_status,
    edit_action_label,
    is_budget_exhausted,
    status_glyph,
    status_label,
    summarize_budget_exhausted,
    summarize_edit_actions,
    summarize_statuses,
    verdict_supersedes_confidence,
)


# ---------------------------------------------------------------------------
# Evidence panel labels
# ---------------------------------------------------------------------------

# Human-readable labels for the verification-mode strings stamped on
# VerificationResult.verification_mode. Keys match VerificationMode.value
# (lowercase string form); unknown modes fall back to a Title-cased
# version of the raw string so a future mode does not render as blank.
_VERIFICATION_MODE_LABELS: dict[str, str] = {
    "local_skip": "Local skip",
    "strict_structured": "Strict structured",
    "standard_reasoning": "Standard reasoning",
    "deep_reasoning": "Deep reasoning",
}


def _verification_mode_label(mode: str) -> str:
    """Map a raw VerificationMode string to a display label."""
    mode = (mode or "").strip().lower()
    if not mode:
        return ""
    return _VERIFICATION_MODE_LABELS.get(mode, mode.replace("_", " ").title())


# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------

SEVERITY_COLORS = {
    "CRITICAL": RGBColor(192, 0, 0),      # Dark red
    "HIGH": RGBColor(255, 102, 0),         # Orange
    "MEDIUM": RGBColor(192, 152, 0),       # Dark yellow/gold
    "GRIPES": RGBColor(128, 0, 128),       # Purple
}

VERDICT_COLORS = {
    "CONFIRMED": RGBColor(0, 128, 0),     # Green
    "CORRECTED": RGBColor(204, 132, 0),    # Amber
    "UNVERIFIED": RGBColor(128, 128, 128), # Gray
    "DISPUTED": RGBColor(192, 0, 0),       # Red
}

VERDICT_ICONS = {
    "CONFIRMED": "✓",
    "CORRECTED": "✎",
    "DISPUTED": "✗",
    "UNVERIFIED": "—",
}

CONFIDENCE_COLORS = {
    "high": RGBColor(0, 128, 0),           # Green
    "moderate": RGBColor(204, 132, 0),     # Amber
    "low": RGBColor(192, 0, 0),            # Red
}

SEVERITY_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "GRIPES"]

# Closed-set status display colors. To avoid presenting all findings as
# equally certain, each status
# gets a distinct color so a quick scroll of the report makes the
# evidence picture visible. Status color hex strings double as the
# summary-table cell shading values.
STATUS_COLORS: dict[ReportStatus, RGBColor] = {
    ReportStatus.VERIFIED_SUPPORTED: RGBColor(0, 128, 0),          # Green
    ReportStatus.VERIFIED_CONTRADICTED: RGBColor(204, 132, 0),     # Amber
    ReportStatus.DISPUTED: RGBColor(192, 0, 0),                    # Red
    ReportStatus.INSUFFICIENT_EVIDENCE: RGBColor(128, 128, 128),   # Gray
    ReportStatus.LOCALLY_CLASSIFIED: RGBColor(59, 130, 246),       # Blue
    ReportStatus.NOT_CHECKED: RGBColor(100, 100, 100),             # Dark gray
    ReportStatus.MANUAL_REVIEW_REQUIRED: RGBColor(255, 102, 0),    # Orange
    # Dark red-orange, distinct from the red
    # used for DISPUTED (C00000) and the orange used for
    # MANUAL_REVIEW_REQUIRED (FF6600). Operational failures need a
    # visually distinct treatment so a quick scroll of the report shows
    # which findings need re-verification vs. those the verifier ran
    # cleanly on.
    ReportStatus.VERIFICATION_FAILED: RGBColor(178, 34, 34),       # Firebrick / dark red-orange
    # Purple, distinct from every other
    # status color above. Indicates "two verifiers, different verdicts"
    # — a quality signal that survives the verdict-based rendering so
    # the report can flag the disagreement without burying it inside a
    # nominally-supported (green) finding.
    ReportStatus.VERIFIED_CONTESTED: RGBColor(128, 0, 128),        # Purple
}

STATUS_SHADING: dict[ReportStatus, str] = {
    ReportStatus.VERIFIED_SUPPORTED: "008000",
    ReportStatus.VERIFIED_CONTRADICTED: "CC8400",
    ReportStatus.DISPUTED: "C00000",
    ReportStatus.INSUFFICIENT_EVIDENCE: "808080",
    ReportStatus.LOCALLY_CLASSIFIED: "3B82F6",
    ReportStatus.NOT_CHECKED: "646464",
    ReportStatus.MANUAL_REVIEW_REQUIRED: "FF6600",
    ReportStatus.VERIFICATION_FAILED: "B22222",
    ReportStatus.VERIFIED_CONTESTED: "800080",
}

EDIT_ACTION_COLORS: dict[EditActionLabel, RGBColor] = {
    EditActionLabel.EDIT_SUGGESTED: RGBColor(0, 128, 0),            # Green
    EditActionLabel.REPORT_ONLY: RGBColor(100, 100, 100),           # Gray
}


# Cache-age badge color tiers. Cache replays carry
# evidence that may have drifted since the original verdict was produced;
# the badge color signals "how stale is this verdict?" at a glance.
#   < 30 days  → amber  (recent — likely still accurate)
#   30-90 days → orange (worth a second look on high-stakes findings)
#   > 90 days  → red    (likely stale; consider re-verifying)
CACHE_AGE_COLORS: dict[str, RGBColor] = {
    "fresh": RGBColor(204, 132, 0),       # Amber
    "stale": RGBColor(255, 102, 0),       # Orange
    "very_stale": RGBColor(192, 0, 0),    # Red
}


def _cache_age_tier(age_days: int) -> str:
    """Bucket a cache entry's age in days into the badge-color tier."""
    if age_days < 30:
        return "fresh"
    if age_days <= 90:
        return "stale"
    return "very_stale"


def _cache_entry_age_days(verification) -> int | None:
    """Return the age in days of the cache entry behind a cache-hit result.

    Returns ``None`` for non-hit results, for results without a recorded
    ``cache_entry_created_ts`` (legacy resume payloads that predate the badge),
    or for timestamps in the future (clock-skew anomaly). The caller
    suppresses the badge in those cases.
    """
    if verification is None:
        return None
    if (getattr(verification, "cache_status", "") or "") != "hit":
        return None
    created_ts = float(getattr(verification, "cache_entry_created_ts", 0.0) or 0.0)
    if created_ts <= 0.0:
        return None
    age_seconds = time.time() - created_ts
    if age_seconds < 0:
        return None
    return int(age_seconds // 86400)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _confidence_tier(confidence: float) -> str:
    """Return 'high', 'moderate', or 'low' for a confidence score."""
    if confidence >= 0.85:
        return "high"
    elif confidence >= 0.60:
        return "moderate"
    return "low"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


# Word 2012 extension namespace — NOT the base w: namespace.
# Per [MS-DOCX] §2.5.1.3, the `collapsed` element lives here.
_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"


def _set_paragraph_collapsed(paragraph) -> None:
    """Mark a heading paragraph as collapsed-by-default when the doc is opened.

    Emits ``<w15:collapsed w15:val="1"/>`` per [MS-DOCX] §2.5.1.3. The element
    MUST be in the Word 2012 extension namespace
    (http://schemas.microsoft.com/office/word/2012/wordml). Using the base
    w: namespace causes Word to silently ignore the element, which is why
    earlier versions of this code rendered Sources headings expanded on open.

    Semantics (per the spec): when this element is set on a heading of level N,
    immediately subsequent paragraphs with a higher heading level number appear
    collapsed when the document is opened. In our use, it's set on the
    "Sources" Heading 4; the URL paragraph below carries outlineLvl=8, which
    satisfies "higher heading level number" and gets collapsed. The next
    finding's Heading 3 (outlineLvl=2) is lower and terminates the zone.

    python-docx's default nsmap does NOT register w15, so we construct the
    element via lxml directly instead of using OxmlElement / qn helpers.
    """
    pPr = paragraph._p.get_or_add_pPr()
    collapsed = etree.SubElement(
        pPr,
        f"{{{_W15_NS}}}collapsed",
        nsmap={"w15": _W15_NS},
    )
    collapsed.set(f"{{{_W15_NS}}}val", "1")


def _set_paragraph_outline_level(paragraph, level: int) -> None:
    """Set <w:outlineLvl> on a paragraph without changing its visual style.

    Word reads the outline level in both directions: a deep level (e.g. 8)
    pulls a body paragraph INTO a preceding heading's open-time collapse
    zone, while a shallow level (e.g. 0) makes a styled-but-unleveled
    paragraph (such as a Title-styled section header, which carries no
    native outline level) TERMINATE any open collapse zone and appear in
    Word's Navigation Pane."""
    pPr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), str(level))
    pPr.append(outline)


def _add_styled_paragraph(doc: Document, text: str, style: str | None = None,
                          bold: bool = False, color: RGBColor | None = None,
                          size: int | None = None, space_after: int | None = None,
                          italic: bool = False):
    """Add a paragraph with optional styling. Returns the paragraph."""
    para = doc.add_paragraph()
    if style:
        para.style = style

    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)

    return para


# ---------------------------------------------------------------------------
# Title block
# ---------------------------------------------------------------------------

def _write_title_block(doc: Document, review, files_reviewed: list[str],
                       cycle_label: str = "2025",
                       failed_review_count: int = 0,
                       module: ReviewModule | None = None,
                       profile: ProjectProfile | None = None) -> None:
    """Write the report title and metadata.

    Uses separate paragraphs instead of \\n within runs to ensure
    reliable rendering across all Word versions and viewers.

    When ``failed_review_count`` > 0 the "Files Reviewed" line reports
    "{reviewed} of {submitted} ({failed} failed review)" so the metadata
    cannot read as a clean complete run when some specs never produced a
    review. A clean run keeps the original "Files Reviewed: {N}" form
    byte-for-byte.

    The title text and the code-cycle line's jurisdiction wording come
    from the run ``module`` (``None`` resolves to the default module, so
    legacy callers keep the original California rendering byte-for-byte).

    ``profile`` (a per-run :class:`ProjectProfile`) appends two centered
    metadata lines — ``Project: {city}, {state}, {country}`` and
    ``Client: {client}`` (D-13). It is only ever present when the run's
    module opted into a profile, so a profile-less run's title block is
    byte-identical to today.
    """
    module = module if module is not None else get_module(None)
    title = doc.add_heading(module.report_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    submitted = len(files_reviewed)
    if failed_review_count > 0:
        reviewed = max(0, submitted - failed_review_count)
        files_reviewed_line = (
            f"Files Reviewed: {reviewed} of {submitted} "
            f"({failed_review_count} failed review)"
        )
    else:
        files_reviewed_line = f"Files Reviewed: {submitted}"

    jurisdiction = module.detector_vocabulary.jurisdiction_label.strip()
    code_cycle_line = (
        f"Code Cycle: {jurisdiction} {cycle_label}" if jurisdiction
        else f"Code Cycle: {cycle_label}"
    )

    # Metadata as separate centered paragraphs (not \n in a single para)
    meta_lines = [
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Model: {review.model}",
        files_reviewed_line,
        code_cycle_line,
    ]
    # Project / client identity lines only when a profile is present (D-13).
    if profile is not None:
        meta_lines.extend(profile.project_meta_lines())

    for line in meta_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Files reviewed
# ---------------------------------------------------------------------------

def _write_files_reviewed(doc: Document, files_reviewed: list[str],
                          failed_review_specs: set[str] | None = None) -> None:
    """Write the files reviewed section with a bullet list.

    Any filename in ``failed_review_specs`` is annotated in red with a
    "— review failed (not reviewed)" suffix so the per-file list stays
    honest: a spec that failed review is visually distinct from one that
    was reviewed clean.
    """
    failed = set(failed_review_specs or ())
    doc.add_heading("Files Reviewed", level=1)
    for filename in files_reviewed:
        if filename in failed:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(f"{filename} — review failed (not reviewed)")
            run.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)
        else:
            doc.add_paragraph(filename, style='List Bullet')



# ---------------------------------------------------------------------------
# Run Diagnostics banner
# ---------------------------------------------------------------------------

def _summarize_run_diagnostics(
    *,
    findings: list,
    status_counts: dict,
    edit_action_counts: dict,
    cross_check_result,
    pipeline_result=None,
    compliance_result=None,
) -> dict:
    """Roll up operational counts for the Run Diagnostics banner.

    Surfaces at-a-glance operational health:
    edit-action histogram, cache-replay count + oldest age, verification
    failures, parse-time REPORT_ONLY demotions, extraction warnings,
    and cross-check status. Every value is
    derived from data already present on the findings / status counts /
    pipeline result; no new persistence is needed.

    Args:
        findings: All findings included in the report (review + cross-
            check). Used to count cache replays, find the oldest cache
            entry age, and count parse-time edit-shape demotions.
        status_counts: Pre-computed ``ReportStatus`` histogram from
            :func:`summarize_statuses` over the same finding list.
        edit_action_counts: Pre-computed ``EditActionLabel`` histogram.
        cross_check_result: The cross-check ``ReviewResult`` or ``None``
            when cross-check was disabled / never ran.
        pipeline_result: The pipeline result, queried opportunistically
            for ``extracted_specs`` so the extraction-warning slot can
            be populated. ``getattr`` with default
            so legacy callers and test doubles work.

    Returns a dict with the rolled-up values the renderer consumes.
    """
    edit_suggested = int(edit_action_counts.get(EditActionLabel.EDIT_SUGGESTED, 0) or 0)
    report_only = int(edit_action_counts.get(EditActionLabel.REPORT_ONLY, 0) or 0)
    verification_failed = int(status_counts.get(ReportStatus.VERIFICATION_FAILED, 0) or 0)

    # Cache replays: count findings whose verification carries a cache
    # hit. Track the oldest age so a reviewer can see "the staleness
    # picture" without expanding individual findings. Legacy resume
    # payloads without a recorded cache-entry timestamp produce ``None`` from
    # :func:`_cache_entry_age_days` — they count toward the total but
    # cannot contribute to the oldest-age display.
    cache_replay_count = 0
    oldest_age_days: int | None = None
    for finding in findings:
        vr = getattr(finding, "verification", None)
        if vr is None:
            continue
        if (getattr(vr, "cache_status", "") or "") != "hit":
            continue
        cache_replay_count += 1
        age = _cache_entry_age_days(vr)
        if age is not None and (oldest_age_days is None or age > oldest_age_days):
            oldest_age_days = age

    # Parse-time REPORT_ONLY demotions: findings stamped with a
    # ``demotion_reason`` are EDIT/ADD/DELETE proposals that were
    # rejected by :func:`validate_edit_shape` at parse time and routed
    # to REPORT_ONLY. Surfaced separately from the general REPORT_ONLY
    # count because they signal model-output shape issues (a
    # potentially-actionable finding that the model emitted with
    # missing fields), not a deliberate coordination/interpretation
    # finding.
    demotion_count = sum(
        1
        for finding in findings
        if (getattr(finding, "demotion_reason", None) or "").strip()
    )

    # Specs that failed review (not reviewed). Sourced from
    # ``PipelineResult.failed_review_specs`` (carried from
    # ``CollectedBatchState.truncated_specs``). This is the headline
    # honesty signal: a spec whose review truncated / parse-errored /
    # errored produces zero findings, exactly like a genuinely-clean
    # spec — so the banner must call it out explicitly or a partially-
    # failed run reads as fully clean. Defensive ``getattr`` keeps legacy
    # callers and test doubles (which may not set the field) at 0.
    failed_review_specs = [
        str(name)
        for name in (getattr(pipeline_result, "failed_review_specs", None) or [])
    ]
    failed_review_count = len(failed_review_specs)

    # Extraction warnings. Looks for an
    # ``extracted_specs`` attribute on the pipeline result with per-spec
    # ``extraction_warnings`` lists. Resolves to 0 when no specs carry
    # warnings — the row still renders so the banner shape stays stable.
    extraction_warning_count = 0
    extracted_specs = getattr(pipeline_result, "extracted_specs", None) or []
    for spec in extracted_specs:
        warnings = getattr(spec, "extraction_warnings", None) or []
        if warnings:
            extraction_warning_count += 1

    # Specs that carried pending Word "Track Changes" markup. Extraction
    # resolves these to the Accept-All view (insertions kept, deletions
    # removed); the advisory tells reviewers the spec was read as accept-all
    # so they can confirm that is the version they meant to review. This is
    # informational (the spec WAS reviewed), not a failure — surfaced in its
    # own calm-amber row/hint, distinct from the red content-loss warning row.
    # Defensive getattr keeps legacy callers / test doubles at 0.
    tracked_changes_spec_count = 0
    for spec in extracted_specs:
        if getattr(spec, "tracked_changes_detected", False):
            tracked_changes_spec_count += 1

    # Cross-check state: None means cross-check was not requested (or
    # disabled); otherwise the status string is "completed" / "skipped" /
    # "failed". The renderer treats "skipped" / "failed" as the actionable
    # signals the plan calls out.
    cross_check_state: dict | None = None
    if cross_check_result is not None:
        cc_status = (
            getattr(cross_check_result, "cross_check_status", None) or "completed"
        )
        cross_check_state = {
            "status": cc_status,
            "finding_count": len(getattr(cross_check_result, "findings", []) or []),
            # Chunked-pass telemetry: chunks that failed / skipped while the
            # overall status is still "completed" (TRUST_AUDIT P1-3 follow-up).
            # Defensive getattr keeps non-chunked results / test doubles at 0.
            "chunk_failures": int(getattr(cross_check_result, "chunk_failures", 0) or 0),
            "chunk_skips": int(getattr(cross_check_result, "chunk_skips", 0) or 0),
            "reason": (
                getattr(cross_check_result, "thinking", "")
                if cc_status == "skipped"
                else (
                    getattr(cross_check_result, "error", "")
                    if cc_status == "failed"
                    else ""
                )
            ),
        }

    # Findings whose verifier consumed the full
    # mode-scaled search budget without producing a grounded verdict.
    # Surfaced in the banner with a hint pointing operators at the
    # severity-tiered budget knob so they can choose to re-run with more
    # headroom. The flag round-trips through resume state so a resumed
    # run shows the same count.
    budget_exhausted_count = summarize_budget_exhausted(findings)

    # WS-4 conditional states: ``None`` when the phase never ran (flag-off
    # module — every profile-less run), so those banners are byte-identical.
    research_state: dict | None = None
    profile = RequirementsProfile.from_dict(
        getattr(pipeline_result, "requirements_profile", None)
    )
    if profile is not None:
        research_state = {
            "dimensions_total": len(profile.dimension_statuses),
            "dimensions_completed": profile.completed_dimensions,
            "dimensions_failed": profile.failed_dimensions,
            "item_count": len(profile.items),
            "ungrounded_count": sum(1 for i in profile.items if not i.grounded),
        }

    compliance_state: dict | None = None
    if compliance_result is not None:
        comp_status = (
            getattr(compliance_result, "cross_check_status", None) or "completed"
        )
        coverage = list(getattr(compliance_result, "coverage", None) or [])
        compliance_state = {
            "status": comp_status,
            "finding_count": len(getattr(compliance_result, "findings", []) or []),
            "missing": sum(1 for c in coverage if c.get("status") == "missing"),
            "contradicted": sum(
                1 for c in coverage if c.get("status") == "contradicted"
            ),
            "chunk_failures": int(getattr(compliance_result, "chunk_failures", 0) or 0),
            "chunk_skips": int(getattr(compliance_result, "chunk_skips", 0) or 0),
            "reason": (
                getattr(compliance_result, "thinking", "")
                if comp_status == "skipped"
                else (
                    getattr(compliance_result, "error", "")
                    if comp_status == "failed"
                    else ""
                )
            ),
        }

    # WS-5 drawing-impact synthesis state. ``None`` when no construction
    # drawings were attached (no drawing-impact result), so the banner row is
    # absent and a drawing-less report stays byte-identical.
    drawing_impact_state: dict | None = None
    di = getattr(pipeline_result, "drawing_impact_result", None)
    if di is not None:
        drawing_impact_state = {
            "status": str(getattr(di, "status", "") or ""),
            "impact_level": str(getattr(di, "impact_level", "") or ""),
            "linked_finding_count": int(getattr(di, "linked_finding_count", 0) or 0),
            "error": str(getattr(di, "error", "") or ""),
        }

    return {
        "edit_suggested": edit_suggested,
        "report_only": report_only,
        "failed_review_count": failed_review_count,
        "failed_review_specs": failed_review_specs,
        "verification_failed": verification_failed,
        "cache_replay_count": cache_replay_count,
        "oldest_cache_age_days": oldest_age_days,
        "demotion_count": demotion_count,
        "extraction_warning_count": extraction_warning_count,
        "tracked_changes_spec_count": tracked_changes_spec_count,
        "cross_check": cross_check_state,
        "budget_exhausted_count": budget_exhausted_count,
        "research": research_state,
        "compliance": compliance_state,
        "drawing_impact": drawing_impact_state,
    }


def _write_run_diagnostics_banner(doc: Document, summary: dict) -> None:
    """Render the Run Diagnostics banner.

    A styled table that surfaces operational health right after the
    title block. Reviewers can scan the banner to answer "did anything
    operationally bad happen on this run?" without scrolling through
    every finding. The table layout (label | value) makes label/value
    pairs greppable in the resulting .docx and keeps the visual style
    consistent with the existing severity / trust-model tables above.

    The banner highlights verification-failure and extraction-warning
    rows in red whenever the count is non-zero (the plan's only hard
    highlight requirement); skipped/failed cross-check status also
    renders in red so a reviewer can spot "the coordination pass did
    not run" at a glance. All other rows render neutral.

    A failure recovery hint paragraph appears below the table when
    verification-failure count > 0, pointing the reviewer at the
    workflow for re-running the failed findings. The actual re-run-
    failed-only mechanism is deferred; this only surfaces the visibility.
    """
    doc.add_heading("Run Diagnostics", level=1)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Operational summary of this run. Use this section to spot at a "
        "glance whether any findings failed verification, were replayed "
        "from a stale cache, or had model-output shape issues that "
        "needed parse-time demotion."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(100, 100, 100)
    intro.paragraph_format.space_after = Pt(6)

    verification_failed = int(summary.get("verification_failed", 0) or 0)
    extraction_warnings = int(summary.get("extraction_warning_count", 0) or 0)
    tracked_changes_spec_count = int(summary.get("tracked_changes_spec_count", 0) or 0)
    cache_count = int(summary.get("cache_replay_count", 0) or 0)
    oldest_age = summary.get("oldest_cache_age_days")
    cross_check = summary.get("cross_check")
    budget_exhausted_count = int(summary.get("budget_exhausted_count", 0) or 0)
    failed_review_specs = list(summary.get("failed_review_specs", []) or [])
    failed_review_count = int(
        summary.get("failed_review_count", len(failed_review_specs)) or 0
    )

    # Build row tuples: (label, value, highlight). ``highlight=True``
    # paints the value cell with light-red shading + dark-red text so
    # the row pops from the surrounding neutral grid.
    rows: list[tuple[str, str, bool]] = [
        ("Edit suggested", str(summary.get("edit_suggested", 0)), False),
        ("Report-only", str(summary.get("report_only", 0)), False),
    ]

    # Specs that failed review (not reviewed). THE headline honesty row:
    # a spec whose review truncated / parse-errored / errored produces
    # zero findings, indistinguishable from a genuinely-clean spec, so it
    # is called out explicitly and highlighted red when > 0. Placed at the
    # top of the operational-health rows because "did we actually review
    # everything we said we did?" is the first question a reviewer of a
    # compliance tool must be able to answer.
    rows.append(
        (
            "Specs that failed review (not reviewed)",
            str(failed_review_count),
            failed_review_count > 0,
        )
    )

    # Cache replays: when there are any, show oldest age too so a
    # reviewer doesn't have to expand each Sources panel to find the
    # staleness picture.
    if cache_count > 0 and oldest_age is not None:
        cache_text = f"{cache_count} (oldest {oldest_age}d old)"
    else:
        cache_text = str(cache_count)
    rows.append(("Cache replays", cache_text, False))

    # Verification failures: only row the plan mandates highlight on
    # when > 0. Distinct from INSUFFICIENT_EVIDENCE — these are
    # operational errors that re-running can fix.
    rows.append(
        (
            "Verification failures (operational)",
            str(verification_failed),
            verification_failed > 0,
        )
    )

    rows.append(
        (
            "REPORT_ONLY demotions at parse time",
            str(summary.get("demotion_count", 0)),
            False,
        )
    )

    # Spec content extraction warnings.
    # Highlight in red when > 0 so a drawing-heavy spec that may have
    # lost text content stands out. Reads from
    # ExtractedSpec.extraction_warnings.
    rows.append(
        (
            "Spec content extraction warnings",
            str(extraction_warnings),
            extraction_warnings > 0,
        )
    )

    # Budget-exhausted findings. Highlight in
    # red when > 0 because it's an actionable signal — the operator can
    # raise the severity of the affected findings to grant more search
    # headroom. The hint paragraph below the table explains the action;
    # the row count gives the at-a-glance number.
    rows.append(
        (
            "Budget-exhausted findings",
            str(budget_exhausted_count),
            budget_exhausted_count > 0,
        )
    )

    # Specs with pending tracked changes (read as accept-all). Informational,
    # not a failure — the spec WAS reviewed, just resolved to the Accept-All
    # view — so the row renders neutral (not the red used for problems) and is
    # only added when > 0, keeping a clean run's banner byte-identical. The
    # amber hint below the table explains the resolution.
    if tracked_changes_spec_count > 0:
        rows.append(
            (
                "Specs with tracked changes (read as accept-all)",
                str(tracked_changes_spec_count),
                False,
            )
        )

    # Cross-spec coordination: surface when the pass ran (or didn't).
    # "skipped" / "failed" are the actionable signals the plan calls out;
    # "completed" is also rendered for transparency so a reader sees
    # "yes, the pass ran" without parsing the trust-model table.
    if cross_check is not None:
        cc_status = str(cross_check.get("status", "completed") or "completed")
        cc_count = int(cross_check.get("finding_count", 0) or 0)
        if cc_status == "skipped":
            cc_value = "skipped"
            cc_highlight = True
        elif cc_status == "failed":
            cc_value = "failed"
            cc_highlight = True
        else:
            cc_value = f"{cc_count} finding{'s' if cc_count != 1 else ''}"
            # A "completed" chunked pass can still have left a division
            # un-analyzed (a failed/skipped chunk). Flag it red so the green
            # finding count doesn't read as a clean, complete pass.
            incomplete = int(cross_check.get("chunk_failures", 0) or 0) + int(
                cross_check.get("chunk_skips", 0) or 0
            )
            if incomplete:
                cc_value += (
                    f" — {incomplete} chunk{'s' if incomplete != 1 else ''} not analyzed"
                )
                cc_highlight = True
            else:
                cc_highlight = False
        rows.append(("Cross-spec coordination", cc_value, cc_highlight))

    # WS-4 conditional rows — rendered only when the phases ran (a
    # profile-less run's banner is byte-identical to before).
    research = summary.get("research")
    if research is not None:
        completed = int(research.get("dimensions_completed", 0) or 0)
        total = int(research.get("dimensions_total", 0) or 0)
        failed = int(research.get("dimensions_failed", 0) or 0)
        items = int(research.get("item_count", 0) or 0)
        ungrounded = int(research.get("ungrounded_count", 0) or 0)
        research_value = (
            f"{completed} of {total} dimensions completed; "
            f"{items} item{'s' if items != 1 else ''} ({ungrounded} ungrounded)"
        )
        rows.append(("Location/client research", research_value, failed > 0))

    compliance = summary.get("compliance")
    if compliance is not None:
        comp_status = str(compliance.get("status", "completed") or "completed")
        comp_count = int(compliance.get("finding_count", 0) or 0)
        missing = int(compliance.get("missing", 0) or 0)
        contradicted = int(compliance.get("contradicted", 0) or 0)
        if comp_status in ("skipped", "failed"):
            comp_value = comp_status
            comp_highlight = True
        else:
            comp_value = (
                f"{comp_count} finding{'s' if comp_count != 1 else ''} — "
                f"{missing} missing / {contradicted} contradicted"
            )
            incomplete_chunks = int(compliance.get("chunk_failures", 0) or 0) + int(
                compliance.get("chunk_skips", 0) or 0
            )
            if incomplete_chunks:
                comp_value += (
                    f" — {incomplete_chunks} chunk"
                    f"{'s' if incomplete_chunks != 1 else ''} not analyzed"
                )
            comp_highlight = bool(missing or contradicted or incomplete_chunks)
        rows.append(("Local-code compliance", comp_value, comp_highlight))

    # WS-5 drawing-impact row — rendered only when construction drawings were
    # attached (so a drawing-less run's banner is byte-identical). "failed" is
    # highlighted red; a completed pass shows the impact level + link count.
    drawing_impact = summary.get("drawing_impact")
    if drawing_impact is not None:
        di_status = str(drawing_impact.get("status", "") or "")
        if di_status == "completed":
            di_level = str(drawing_impact.get("impact_level", "") or "minimal")
            di_linked = int(drawing_impact.get("linked_finding_count", 0) or 0)
            di_value = (
                f"{di_level} — {di_linked} finding{'s' if di_linked != 1 else ''} linked"
            )
            di_highlight = False
        else:
            di_value = "analysis failed"
            di_highlight = True
        rows.append(("Drawing analysis impact", di_value, di_highlight))

    # Render as a 2-column table. Label cells get a light-gray
    # shading so the visual rhythm matches the existing summary table
    # above; value cells get red shading when ``highlight=True``.
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row_idx, (label, value, highlight) in enumerate(rows):
        label_cell = table.rows[row_idx].cells[0]
        _set_cell_shading(label_cell, "F0F0F0")
        label_cell.text = ""
        p = label_cell.paragraphs[0]
        label_run = p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)

        value_cell = table.rows[row_idx].cells[1]
        if highlight:
            _set_cell_shading(value_cell, "FFE5E5")
        value_cell.text = ""
        p = value_cell.paragraphs[0]
        value_run = p.add_run(value)
        value_run.bold = True
        value_run.font.size = Pt(10)
        if highlight:
            value_run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Failed-review recovery hint ---
    # The headline trust signal: name the specs that were NOT reviewed so
    # a reviewer cannot mistake a partially-failed run for a clean one.
    # Rendered first (above the verification/budget hints) and in
    # failure-red, matching the ⚠ glyph used on the per-spec GUI log line.
    # The cause/remedy differs from a verification failure: these specs
    # never produced findings at all, so the absence of findings carries
    # no information about whether the spec is compliant.
    if failed_review_count > 0:
        names = ", ".join(failed_review_specs) if failed_review_specs else "(names unavailable)"
        plural = failed_review_count != 1
        hint_para = doc.add_paragraph()
        hint_para.paragraph_format.space_before = Pt(6)
        hint_para.paragraph_format.space_after = Pt(8)
        hint_run = hint_para.add_run(
            f"⚠ {failed_review_count} spec{'s' if plural else ''} failed "
            f"review and {'were' if plural else 'was'} NOT reviewed: {names}. "
            f"{'Their' if plural else 'Its'} review truncated, failed to "
            f"parse, or errored, so {'they' if plural else 'it'} produced no "
            "findings — the absence of findings does NOT mean "
            f"{'they are' if plural else 'it is'} compliant. Re-run "
            f"{'these specs' if plural else 'this spec'} individually to "
            "obtain a review."
        )
        hint_run.font.size = Pt(10)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Failure recovery hint ---
    # The re-run-failed-only mechanism is deferred; for now we just
    # describe what re-running will do. The ⚠ glyph match is the same
    # one stamped on individual VERIFICATION_FAILED findings (see
    # STATUS_GLYPHS in report_status.py) so a reviewer can grep for it.
    if verification_failed > 0:
        hint_para = doc.add_paragraph()
        hint_para.paragraph_format.space_before = Pt(6)
        hint_para.paragraph_format.space_after = Pt(8)
        hint_run = hint_para.add_run(
            f"{verification_failed} finding"
            f"{'s' if verification_failed != 1 else ''} failed verification "
            "due to operational errors (network, rate limit, parse failures). "
            "These are visually marked with the ⚠ glyph in the findings "
            "below. Re-running the review will re-attempt verification for "
            "these findings; the cache does not persist operational-failure "
            "results so a re-run sees them fresh."
        )
        hint_run.font.size = Pt(10)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(178, 34, 34)

    # --- Budget-exhaustion recovery hint ---
    # Distinct from the failure hint above because the cause and the
    # remedy differ: failures are transient (re-run sees them fresh),
    # but budget exhaustion is a policy outcome (re-run with the same
    # severity will exhaust the same budget). The actionable hint is to
    # raise the finding's severity so the routing decision allocates
    # more searches. The per-severity budgets are rendered from
    # api_config._SEVERITY_MAX_USES (via web_search_max_uses_for_severity)
    # so this hint cannot drift from the actual policy. Rendered in a
    # calmer amber rather than the failure-red so a reader can tell the
    # two situations apart at a glance.
    if budget_exhausted_count > 0:
        budget_str = ", ".join(
            f"{sev} {web_search_max_uses_for_severity(sev)}"
            for sev in SEVERITY_ORDER
        )
        hint_para = doc.add_paragraph()
        hint_para.paragraph_format.space_before = Pt(6)
        hint_para.paragraph_format.space_after = Pt(8)
        hint_run = hint_para.add_run(
            f"{budget_exhausted_count} finding"
            f"{'s' if budget_exhausted_count != 1 else ''} exhausted the "
            "verifier's web_search budget without grounding a verdict "
            "(rendered as 'Insufficient evidence (search budget exhausted)' "
            "inline). Per-severity web_search budgets are "
            f"{budget_str}; re-running a lower-severity finding at a higher "
            "severity grants more headroom (CRITICAL findings already "
            "receive the maximum budget)."
        )
        hint_run.font.size = Pt(10)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(204, 132, 0)

    # --- Tracked-changes advisory ---
    # Not a failure: the spec WAS reviewed, but as the Accept-All-Changes view
    # (insertions kept, deletions removed). Rendered in the same calm amber as
    # the budget hint so it reads as informational, distinct from the failure-
    # red hints above. Prompts the reviewer to confirm the accept-all view is
    # the one they intend to review (vs. the pre-redline text).
    if tracked_changes_spec_count > 0:
        plural = tracked_changes_spec_count != 1
        hint_para = doc.add_paragraph()
        hint_para.paragraph_format.space_before = Pt(6)
        hint_para.paragraph_format.space_after = Pt(8)
        hint_run = hint_para.add_run(
            f"{tracked_changes_spec_count} spec{'s' if plural else ''} contained "
            f"pending tracked changes (Word 'Track Changes'). "
            f"{'They were' if plural else 'It was'} reviewed as if all changes "
            "were accepted — insertions kept, deletions removed — i.e. the text "
            "that will remain once the redline is accepted. Confirm that is the "
            "version you intend to review."
        )
        hint_run.font.size = Pt(10)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(204, 132, 0)

    # --- Research partial-failure hint (WS-4, D-13) ---
    # Amber (informational, not transient): the run continued on a partial
    # profile per the D-3 failure policy, so the reader must know which
    # coverage the requirements analysis is missing.
    if research is not None and int(research.get("dimensions_failed", 0) or 0) > 0:
        failed = int(research.get("dimensions_failed", 0) or 0)
        hint_para = doc.add_paragraph()
        hint_para.paragraph_format.space_before = Pt(6)
        hint_para.paragraph_format.space_after = Pt(8)
        hint_run = hint_para.add_run(
            f"{failed} location/client research dimension"
            f"{'s' if failed != 1 else ''} failed, so the Project Requirements "
            "Profile — and every check made against it — is missing that "
            "coverage. The Jurisdiction & Client Requirements section names "
            "the failed dimension(s); re-running the review retries them."
        )
        hint_run.font.size = Pt(10)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(204, 132, 0)

    # --- Compliance hint (WS-4, D-13) ---
    # Red when the pass didn't run (skipped/failed — invariant 8: never
    # silent) or when it found missing/contradicted requirements.
    if compliance is not None:
        comp_status = str(compliance.get("status", "completed") or "completed")
        missing = int(compliance.get("missing", 0) or 0)
        contradicted = int(compliance.get("contradicted", 0) or 0)
        if comp_status in ("skipped", "failed"):
            hint_para = doc.add_paragraph()
            hint_para.paragraph_format.space_before = Pt(6)
            hint_para.paragraph_format.space_after = Pt(8)
            hint_run = hint_para.add_run(
                f"⚠ The local-code compliance evaluation {comp_status} "
                f"({compliance.get('reason') or 'no reason recorded'}). The "
                "specs were NOT evaluated against the researched location/"
                "client requirements — absence of compliance findings carries "
                "no information."
            )
            hint_run.font.size = Pt(10)
            hint_run.font.italic = True
            hint_run.font.color.rgb = RGBColor(192, 0, 0)
        elif missing or contradicted:
            hint_para = doc.add_paragraph()
            hint_para.paragraph_format.space_before = Pt(6)
            hint_para.paragraph_format.space_after = Pt(8)
            hint_run = hint_para.add_run(
                f"The compliance evaluation classified {missing} researched "
                f"requirement{'s' if missing != 1 else ''} as missing from the "
                f"package and {contradicted} as contradicted. See the "
                "Requirements Coverage table and the Local-Code Compliance "
                "findings for the specifics."
            )
            hint_run.font.size = Pt(10)
            hint_run.font.italic = True
            hint_run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Drawing-impact failure hint (WS-5) ---
    # Amber, not red: a failed synthesis pass does NOT mean the drawings were
    # ignored — they still rode on every review call as context; the run just
    # couldn't produce the explanatory summary. Distinct cause/remedy from the
    # verification and compliance hints above.
    if drawing_impact is not None and str(drawing_impact.get("status", "")) != "completed":
        err = str(drawing_impact.get("error", "") or "")
        hint_para = doc.add_paragraph()
        hint_para.paragraph_format.space_before = Pt(6)
        hint_para.paragraph_format.space_after = Pt(8)
        hint_run = hint_para.add_run(
            "The drawing-impact analysis did not complete"
            + (f" ({err})" if err else "")
            + ". The attached drawings still informed the review as context; "
            "only the explanatory summary is missing. Re-running the review "
            "re-attempts it."
        )
        hint_run.font.size = Pt(10)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(204, 132, 0)

    doc.add_paragraph()  # Spacer between banner and the next section.


# ---------------------------------------------------------------------------
# Jurisdiction & Client Requirements section (WS-4, D-13)
# ---------------------------------------------------------------------------

# Coverage-status display: label, value-cell shading, text color.
_COVERAGE_STATUS_STYLES: dict[str, tuple[str, str, RGBColor]] = {
    "represented": ("Represented", "C6EFCE", RGBColor(0x00, 0x61, 0x00)),
    "missing": ("MISSING", "FFE5E5", RGBColor(0xC0, 0x00, 0x00)),
    "contradicted": ("CONTRADICTED", "FFE5E5", RGBColor(0xC0, 0x00, 0x00)),
    "unclear": ("Unclear", "FFF2CC", RGBColor(0x99, 0x66, 0x00)),
}

# Categories whose items pin an edition (the at-a-glance "which edition do I
# cite?" table, D-13 [FT]).
_EDITION_CATEGORIES = ("governing_code", "referenced_standard", "local_amendment")


def _requirements_detail_line(item) -> str:
    parts = []
    if item.authority:
        parts.append(f"Authority: {item.authority}")
    if item.code_reference:
        parts.append(f"Ref: {item.code_reference}")
    if item.accepted_sources:
        parts.append("Sources: " + ", ".join(item.accepted_sources))
    parts.append(f"confidence {round(item.confidence * 100)}%")
    return " • ".join(parts)


def _write_requirements_section(
    doc: Document,
    requirements_profile: RequirementsProfile,
    compliance_result,
    module: ReviewModule,
) -> None:
    """Render the "Jurisdiction & Client Requirements" section (D-13).

    Placed between "Files Reviewed" and the methodology note. Contents, in
    order: project identity + research provenance, the per-category
    requirement items (grounded/[UNVERIFIED] marked), the edition reference
    table, the compliance coverage matrix (``represented`` rows rendered as
    visibly as ``missing`` ones — a critic that can say "this part is
    right" earns trust), and the Process & Schedule Advisories subsection.
    Rendered only when the run produced a requirements profile, so
    profile-less reports are byte-identical.
    """
    doc.add_heading("Jurisdiction & Client Requirements", level=1)

    project = ProjectProfile.from_dict(requirements_profile.project)
    intro = doc.add_paragraph()
    total = len(requirements_profile.dimension_statuses)
    searches = sum(
        s.web_search_requests for s in requirements_profile.dimension_statuses
    )
    identity = (
        f"Project: {project.city}, {project.state_display}, "
        f"{project.country_display} — Client: {project.client_name}. "
        if project is not None
        else ""
    )
    intro_run = intro.add_run(
        f"{identity}Requirements researched via location/client web research "
        f"({requirements_profile.completed_dimensions} of {total} dimensions "
        f"completed, {searches} web searches), researched "
        f"{requirements_profile.research_date}. Edition and process facts are "
        "as-of that date. Items marked [UNVERIFIED] could not be grounded in "
        "retrieved sources and are never treated as controlling."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(100, 100, 100)

    failed_dimensions = [
        s for s in requirements_profile.dimension_statuses if s.status != "completed"
    ]
    if failed_dimensions:
        warn = doc.add_paragraph()
        warn_run = warn.add_run(
            f"⚠ {len(failed_dimensions)} research dimension(s) failed "
            f"({', '.join(s.dimension_id for s in failed_dimensions)}); the "
            "profile below is missing their requirements."
        )
        warn_run.font.size = Pt(10)
        warn_run.font.italic = True
        warn_run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Requirement items, grouped by the same sections as the rendered
    # context block so the report and the model saw the same organization.
    spec_items = [i for i in requirements_profile.items if not i.is_process_advisory]
    sections: dict[str, list] = {name: [] for name in PROFILE_SECTION_ORDER}
    for item in spec_items:
        sections[PROFILE_CATEGORY_SECTIONS.get(item.category, "OTHER")].append(item)
    for section_name in PROFILE_SECTION_ORDER:
        section_items = sections[section_name]
        if not section_items:
            continue
        doc.add_heading(section_name.title(), level=2)
        for item in section_items:
            para = doc.add_paragraph(style="List Bullet")
            id_run = para.add_run(f"[{item.item_id}] ")
            id_run.font.size = Pt(9)
            id_run.font.color.rgb = RGBColor(128, 128, 128)
            req_run = para.add_run(item.requirement)
            req_run.font.size = Pt(10)
            if not item.grounded:
                marker = para.add_run("  [UNVERIFIED]")
                marker.font.size = Pt(9)
                marker.bold = True
                marker.font.color.rgb = RGBColor(192, 0, 0)
            detail = para.add_run(f"\n{_requirements_detail_line(item)}")
            detail.font.size = Pt(8)
            detail.font.italic = True
            detail.font.color.rgb = RGBColor(120, 120, 120)

    # --- Edition reference table (D-13 [FT]): the single most-reused
    # artifact — "which edition do I cite?" at a glance.
    edition_items = [i for i in spec_items if i.category in _EDITION_CATEGORIES]
    if edition_items:
        doc.add_heading("Adopted & Referenced Editions", level=2)
        note = doc.add_paragraph()
        note_run = note.add_run(
            f"Edition facts researched {requirements_profile.research_date}; "
            "verify against the adopting instrument before relying on them."
        )
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(120, 120, 120)
        table = doc.add_table(rows=len(edition_items) + 1, cols=3)
        table.style = "Table Grid"
        for col, header in enumerate(("Topic", "Requirement / edition", "Authority & reference")):
            cell = table.rows[0].cells[col]
            _set_cell_shading(cell, "F0F0F0")
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(9)
        for row_idx, item in enumerate(edition_items, start=1):
            cells = table.rows[row_idx].cells
            values = (
                item.topic or item.category,
                item.requirement + ("" if item.grounded else "  [UNVERIFIED]"),
                "; ".join(p for p in (item.authority, item.code_reference) if p),
            )
            for col, value in enumerate(values):
                cells[col].text = ""
                run = cells[col].paragraphs[0].add_run(value)
                run.font.size = Pt(9)

    # --- Compliance coverage matrix.
    coverage = list(getattr(compliance_result, "coverage", None) or [])
    if coverage:
        doc.add_heading("Requirements Coverage", level=2)
        requirement_by_id = {i.item_id: i for i in requirements_profile.items}
        table = doc.add_table(rows=len(coverage) + 1, cols=3)
        table.style = "Table Grid"
        for col, header in enumerate(("Requirement", "Coverage", "Evidence")):
            cell = table.rows[0].cells[col]
            _set_cell_shading(cell, "F0F0F0")
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(9)
        for row_idx, entry in enumerate(coverage, start=1):
            cells = table.rows[row_idx].cells
            rid = entry.get("requirement_id") or ""
            item = requirement_by_id.get(rid)
            requirement_text = item.requirement if item is not None else "(unknown requirement)"
            cells[0].text = ""
            req_run = cells[0].paragraphs[0].add_run(f"[{rid}] {requirement_text}")
            req_run.font.size = Pt(9)

            status = str(entry.get("status") or "unclear")
            label, shading, color = _COVERAGE_STATUS_STYLES.get(
                status, _COVERAGE_STATUS_STYLES["unclear"]
            )
            cells[1].text = ""
            _set_cell_shading(cells[1], shading)
            status_run = cells[1].paragraphs[0].add_run(label)
            status_run.bold = True
            status_run.font.size = Pt(9)
            status_run.font.color.rgb = color

            evidence_parts = [p for p in (entry.get("evidence"), entry.get("fileName")) if p]
            cells[2].text = ""
            ev_run = cells[2].paragraphs[0].add_run(" — ".join(evidence_parts))
            ev_run.font.size = Pt(8)
    elif compliance_result is not None:
        status = getattr(compliance_result, "cross_check_status", None) or "completed"
        if status in ("skipped", "failed"):
            para = doc.add_paragraph()
            reason = (
                getattr(compliance_result, "thinking", "")
                if status == "skipped"
                else getattr(compliance_result, "error", "")
            )
            run = para.add_run(
                f"⚠ Compliance coverage unavailable — the compliance pass "
                f"{status}: {reason}"
            )
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Process & Schedule Advisories (D-7 [FT]): real project-team
    # deliverables that are NOT spec content and never coverage rows.
    advisories = [i for i in requirements_profile.items if i.is_process_advisory]
    if advisories:
        doc.add_heading("Process & Schedule Advisories", level=2)
        note = doc.add_paragraph()
        note_run = note.add_run(
            "Permit, schedule, and process facts the project team must act on. "
            "These are not specification content and are never counted as "
            "missing spec coverage."
        )
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(120, 120, 120)
        for item in advisories:
            para = doc.add_paragraph(style="List Bullet")
            req_run = para.add_run(item.requirement)
            req_run.font.size = Pt(10)
            if not item.grounded:
                marker = para.add_run("  [UNVERIFIED]")
                marker.font.size = Pt(9)
                marker.bold = True
                marker.font.color.rgb = RGBColor(192, 0, 0)
            detail = para.add_run(f"\n{_requirements_detail_line(item)}")
            detail.font.size = Pt(8)
            detail.font.italic = True
            detail.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()  # Spacer before the next section.


# ---------------------------------------------------------------------------
# Drawing-impact section (WS-5)
# ---------------------------------------------------------------------------

# Overall-impact badge styles: display label + text color, keyed by the
# ``DrawingImpactResult.impact_level`` value.
_DRAWING_IMPACT_LEVEL_STYLES: dict[str, tuple[str, RGBColor]] = {
    "substantial": ("Substantial", RGBColor(0, 128, 0)),      # Green
    "moderate": ("Moderate", RGBColor(204, 132, 0)),          # Amber
    "minimal": ("Minimal", RGBColor(128, 128, 128)),          # Gray
    "none": ("None", RGBColor(100, 100, 100)),                # Dark gray
}

# Per-link relationship styles: display label + text color.
_DRAWING_RELATIONSHIP_STYLES: dict[str, tuple[str, RGBColor]] = {
    "corroborated": ("Corroborated by drawings", RGBColor(0, 128, 0)),    # Green
    "contradicted": ("Contradicted by drawings", RGBColor(192, 0, 0)),    # Red
    "contextualized": ("Contextualized by drawings", RGBColor(59, 130, 246)),  # Blue
}


def _write_drawing_impact_section(doc: Document, drawing_impact, findings_by_id: dict) -> None:
    """Render "How the Drawings Informed This Review" (WS-5).

    Rendered only when the run carried a drawing-impact result — i.e. a
    construction-drawing digest was in Project Context. A run without attached
    drawings passes ``drawing_impact=None`` and this returns immediately, so
    the report is byte-identical. All fields are read defensively via
    ``getattr`` so a legacy / test-double result object cannot crash export.

    ``findings_by_id`` maps ``finding_id`` → the report finding, so each link
    renders enough context (severity, file, issue snippet) to stand on its own
    above the findings list.
    """
    if drawing_impact is None:
        return

    doc.add_heading("How the Drawings Informed This Review", level=1)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Construction drawings were attached to this run and analyzed into a "
        "text digest that was supplied as reference context to every stage of "
        "the review. This section reports how that drawing content bore on the "
        "findings — where the drawings corroborate, contradict, or add context "
        "to a finding, cited to the drawing sheet pages."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(100, 100, 100)

    status = str(getattr(drawing_impact, "status", "") or "")
    if status != "completed":
        err = str(getattr(drawing_impact, "error", "") or "")
        para = doc.add_paragraph()
        run = para.add_run(
            "⚠ The drawing-impact analysis did not complete"
            + (f" ({err})" if err else "")
            + ". The drawings were still provided to the review as context; "
            "this run simply could not produce the explanatory summary. "
            "Re-running the review will re-attempt it."
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        doc.add_paragraph()
        return

    # Overall impact badge.
    level = str(getattr(drawing_impact, "impact_level", "") or "").strip().lower()
    label, color = _DRAWING_IMPACT_LEVEL_STYLES.get(
        level, _DRAWING_IMPACT_LEVEL_STYLES["minimal"]
    )
    badge = doc.add_paragraph()
    prefix_run = badge.add_run("Overall drawing impact: ")
    prefix_run.bold = True
    prefix_run.font.size = Pt(11)
    level_run = badge.add_run(label)
    level_run.bold = True
    level_run.font.size = Pt(11)
    level_run.font.color.rgb = color

    # Narrative — split on blank lines into paragraphs so multi-theme prose
    # keeps its structure.
    narrative = str(getattr(drawing_impact, "narrative", "") or "").strip()
    if narrative:
        for chunk in (c.strip() for c in narrative.split("\n\n")):
            if not chunk:
                continue
            para = doc.add_paragraph()
            run = para.add_run(chunk)
            run.font.size = Pt(11)

    # Per-finding links.
    links = list(getattr(drawing_impact, "finding_links", None) or [])
    if links:
        doc.add_heading("Findings the drawings bear on", level=2)
        for link in links:
            fid = str(getattr(link, "finding_id", "") or "")
            relationship = str(getattr(link, "relationship", "") or "").strip().lower()
            rel_label, rel_color = _DRAWING_RELATIONSHIP_STYLES.get(
                relationship, _DRAWING_RELATIONSHIP_STYLES["contextualized"]
            )
            finding = findings_by_id.get(fid)
            severity = str(getattr(finding, "severity", "") or "") if finding is not None else ""
            issue = str(getattr(finding, "issue", "") or "") if finding is not None else ""
            file_name = str(getattr(finding, "fileName", "") or "") if finding is not None else ""

            head = doc.add_paragraph(style="List Bullet")
            rel_run = head.add_run(rel_label)
            rel_run.bold = True
            rel_run.font.size = Pt(10)
            rel_run.font.color.rgb = rel_color
            meta_bits = [b for b in (severity, file_name) if b]
            meta = f"  [{fid}]" + (f" — {' · '.join(meta_bits)}" if meta_bits else "")
            meta_run = head.add_run(meta)
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(120, 120, 120)
            if issue:
                snippet = issue.strip()
                if len(snippet) > 200:
                    snippet = snippet[:199] + "…"
                issue_run = head.add_run(f"\n{snippet}")
                issue_run.font.size = Pt(10)

            explanation = str(getattr(link, "explanation", "") or "").strip()
            refs = [str(r) for r in (getattr(link, "sheet_references", None) or []) if str(r).strip()]
            detail_parts = []
            if explanation:
                detail_parts.append(explanation)
            if refs:
                detail_parts.append("Sheets: " + ", ".join(refs))
            if detail_parts:
                detail = head.add_run("\n" + " • ".join(detail_parts))
                detail.font.size = Pt(9)
                detail.font.italic = True
                detail.font.color.rgb = RGBColor(90, 90, 90)
    else:
        note = doc.add_paragraph()
        note_run = note.add_run(
            "No individual finding turned on drawing content — on this run the "
            "drawings served as background context rather than the basis for a "
            "specific finding."
        )
        note_run.font.size = Pt(10)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Spacer before the next section.


# ---------------------------------------------------------------------------
# Methodology note
# ---------------------------------------------------------------------------

def _summarize_verification_outcomes(findings: list) -> dict[str, object]:
    """Roll up the trust-model statuses + raw verdict counts for the methodology note.

    The status histogram (computed via
    :func:`report_status.summarize_statuses`) drives the methodology
    narrative and the new trust-model summary table. The verdict-count
    breakdown is preserved alongside it so existing summary lines
    (``CONFIRMED / CORRECTED / DISPUTED / UNVERIFIED``) keep working.
    """
    stats = {
        "total_findings": len(findings),
        "with_verification": 0,
        "verdict_counts": {"CONFIRMED": 0, "CORRECTED": 0, "DISPUTED": 0, "UNVERIFIED": 0},
        "status_counts": summarize_statuses(findings),
        "edit_action_counts": summarize_edit_actions(findings),
    }
    for finding in findings:
        verification = getattr(finding, "verification", None)
        if not verification:
            continue
        stats["with_verification"] += 1
        verdict = str(getattr(verification, "verdict", "UNVERIFIED") or "UNVERIFIED").upper()
        if verdict not in stats["verdict_counts"]:
            verdict = "UNVERIFIED"
        stats["verdict_counts"][verdict] += 1
    verified_non_unverified = stats["verdict_counts"]["CONFIRMED"] + stats["verdict_counts"]["CORRECTED"] + stats["verdict_counts"]["DISPUTED"]
    stats["all_unverified"] = stats["with_verification"] > 0 and verified_non_unverified == 0 and stats["verdict_counts"]["UNVERIFIED"] == stats["with_verification"]
    stats["partial_unverified"] = stats["verdict_counts"]["UNVERIFIED"] > 0 and verified_non_unverified > 0
    return stats


def _write_methodology_note(doc, cross_check_enabled: bool = False, cycle_label: str = "2025", cross_check_status: str | None = None, cross_check_reason: str = "", verification_stats: dict[str, object] | None = None, module: ReviewModule | None = None) -> None:
    """Write a brief methodology note explaining how the review was produced.

    The run ``module`` supplies the domain phrase (``report_context_phrase``),
    the jurisdiction wording of the code-cycle sentence, and the cycle whose
    pinned editions are enumerated. ``None`` resolves to the default module,
    so legacy callers keep the original California sentences byte-for-byte.
    """
    module = module if module is not None else get_module(None)
    domain_phrase = module.report_context_phrase
    doc.add_heading("About This Review", level=1)

    doc.add_paragraph(
        "This report was generated by Spec Critic, an AI-assisted specification "
        "review tool. Each specification was analyzed by Claude for "
        "code compliance issues, coordination problems, and technical errors "
        f"relevant to {domain_phrase}. Findings are classified by "
        "severity (Critical, High, Medium, Gripe) and assigned a confidence score "
        "reflecting the review model\u2019s certainty at review time, before "
        "verification. The confidence score and the verification verdict are "
        "distinct signals: confidence is the review model\u2019s own pre-verification "
        "estimate, while the verdict is a separate verification pass\u2019s grounded "
        "assessment. For any finding the verifier went on to confirm, correct, "
        "contest, or dispute, the verification verdict \u2014 not the review "
        "confidence \u2014 is the authoritative trust signal, so the report "
        "de-emphasizes the confidence on those findings (shown small beside the "
        "status, marked \u201cpre-verification\u201d) and lets the verdict stand as "
        "the headline. The confidence score stays prominent only where verification "
        "did not reach a verdict (for example, not checked or insufficient evidence)."
    )

    verification_stats = verification_stats or {}
    all_unverified = bool(verification_stats.get("all_unverified", False))
    partial_unverified = bool(verification_stats.get("partial_unverified", False))
    with_verification = int(verification_stats.get("with_verification", 0) or 0)

    if all_unverified:
        para2_text = (
            "Verification was attempted but did not return usable results. "
            "Findings have not been independently verified."
        )
    elif partial_unverified:
        para2_text = (
            "Findings were checked in a secondary AI verification pass with web search access. "
            "Some findings could not be verified — see individual verdicts."
        )
    elif with_verification > 0:
        para2_text = (
            "All findings were checked in a secondary AI verification pass with web search access. "
            "Verification verdicts (Confirmed, Corrected, Disputed, or Unverified) reflect the verifier model's assessment and should be treated as advisory."
        )
    else:
        para2_text = (
            "No verification outcomes were recorded for this run. "
            "Findings should be treated as unverified unless noted otherwise."
        )

    jurisdiction = module.detector_vocabulary.jurisdiction_label.strip()
    cycle_phrase = f"{jurisdiction} {cycle_label}" if jurisdiction else cycle_label
    para2_text += f" This review used {cycle_phrase} code cycle references."

    if cross_check_enabled and cross_check_status == "completed":
        para2_text += " Cross-check completed."
    elif cross_check_enabled and cross_check_status == "skipped":
        para2_text += f" Cross-check was skipped: {cross_check_reason}"
    elif cross_check_enabled and cross_check_status == "failed":
        para2_text += f" Cross-check failed: {cross_check_reason}"

    para2_text += (
        " This report is advisory \u2014 findings should be reviewed by the "
        "engineer of record before acting on them."
    )

    doc.add_paragraph(para2_text)

    # Surface the pinned standards editions
    # that drove the verifier prompt for this cycle. Reviewers reading
    # the report can see which editions were treated as authoritative
    # without opening the source; if the spec cites a different edition
    # for one of these standards, the finding's relevance to the
    # current cycle should be re-checked.
    pinning_text = _render_pinned_editions_note(module.cycle, jurisdiction)
    if pinning_text:
        doc.add_paragraph(pinning_text)

    # Collapsibility tip
    doc.add_paragraph(
        "Tip: In Word, hover over any heading to reveal a collapse triangle. "
        "Click it to hide the content beneath that heading. Use this to "
        "collapse individual findings or entire severity groups."
    )


def _render_pinned_editions_note(cycle: CodeCycle, jurisdiction: str) -> str:
    """Render the methodology paragraph that enumerates pinned editions.

    Renders from the run module's own :class:`CodeCycle`, so the label in
    the sentence is always the label of the cycle whose standards are
    enumerated. (The legacy label\u2192``AVAILABLE_CYCLES`` lookup silently
    fell back to the California default for any unknown label, which would
    have rendered California standards into another module's report.)
    Pinning details only render when the cycle has populated edition
    fields \u2014 a cycle with no pinning falls back to an empty string so
    the methodology note degrades gracefully.
    """
    entries = [std for std in cycle.standards if std.edition]
    if not entries:
        return ""
    # Join with semicolons because an individual description can itself contain a
    # comma (e.g. "NFPA 13 2025, as amended by California").
    rendered = "; ".join(std.description for std in entries)
    cycle_phrase = (
        f"{cycle.label} {jurisdiction} cycle" if jurisdiction
        else f"{cycle.label} cycle"
    )
    return (
        f"This review pinned the following standards editions per the "
        f"{cycle_phrase}: {rendered}. Findings referencing other editions should "
        "be reviewed for relevance to the current cycle."
    )


# ---------------------------------------------------------------------------
# Summary table
# ---------------------------------------------------------------------------

def _write_summary_table(doc: Document, review, cross_check_result, *, total_elapsed_seconds: float | None = None) -> None:
    """Write the summary section with a styled severity counts table."""
    doc.add_heading("Summary", level=1)

    cc_count = (len(cross_check_result.findings)
                if cross_check_result and cross_check_result.findings else 0)

    # Build column definitions
    columns = [
        ("CRITICAL", review.critical_count, "C00000"),
        ("HIGH", review.high_count, "FF6600"),
        ("MEDIUM", review.medium_count, "C09800"),
        ("GRIPES", review.gripe_count, "800080"),
        ("TOTAL", review.total_count, "333333"),
    ]
    if cc_count > 0:
        columns.append(("CROSS-CHECK", cc_count, "06B6D4"))

    # Create table: header row + count row
    table = doc.add_table(rows=2, cols=len(columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row (colored backgrounds with white text)
    for col_idx, (label, _count, hex_color) in enumerate(columns):
        cell = table.rows[0].cells[col_idx]
        _set_cell_shading(cell, hex_color)
        # Clear default paragraph and write header text
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        # White text on all dark backgrounds, black on medium yellow
        if label == "MEDIUM":
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Count row
    for col_idx, (_label, count, _hex) in enumerate(columns):
        cell = table.rows[1].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(count))
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph()  # Spacer

    # Review-stage token usage only (excludes cross-check + verification)
    para = doc.add_paragraph()
    para.add_run("Review Stage Tokens: ").bold = True
    para.add_run(
        f"{review.input_tokens:,} input → "
        f"{review.output_tokens:,} output"
    )

    para = doc.add_paragraph()
    para.add_run("Note: ").bold = True
    para.add_run(
        "Cross-check and verification token usage are not included in the totals above."
    )

    # Processing time
    para = doc.add_paragraph()
    para.add_run("Processing Time: ").bold = True
    processing_seconds = total_elapsed_seconds if total_elapsed_seconds is not None else review.elapsed_seconds
    para.add_run(f"{processing_seconds:.1f} seconds")

    # Verification summary
    all_findings = list(review.findings)
    if cross_check_result and cross_check_result.findings:
        all_findings.extend(cross_check_result.findings)

    verdicts: dict[str, int] = {}
    for f in all_findings:
        if f.verification:
            v = f.verification.verdict
            verdicts[v] = verdicts.get(v, 0) + 1

    if verdicts:
        para = doc.add_paragraph()
        para.add_run("Verification: ").bold = True
        verdict_parts = [
            f"{verdicts[v]} {v.lower()}"
            for v in ["CONFIRMED", "CORRECTED", "DISPUTED", "UNVERIFIED"]
            if v in verdicts
        ]
        para.add_run(", ".join(verdict_parts))
        # The UNVERIFIED *verdict* lumps locally-classified findings
        # (resolved without a web search — placeholders, stale cycles, etc.)
        # together with genuinely-unverifiable ones, so the raw count
        # overstates "unverified". Break the bucket down by trust status,
        # using the same labels as the Trust Model Summary below, so a reader
        # doesn't see "N unverified" here and "M insufficient evidence" there
        # and wonder which is real.
        if verdicts.get("UNVERIFIED", 0) > 0:
            unv_status_counts: dict = {}
            for f in all_findings:
                ver = getattr(f, "verification", None)
                if ver and str(getattr(ver, "verdict", "") or "").upper() == "UNVERIFIED":
                    st = classify_status(f)
                    unv_status_counts[st] = unv_status_counts.get(st, 0) + 1
            breakdown = [
                f"{unv_status_counts[s]} {status_label(s).lower()}"
                for s in STATUS_DISPLAY_ORDER
                if unv_status_counts.get(s, 0) > 0
            ]
            if breakdown:
                note = doc.add_paragraph()
                note_run = note.add_run(
                    "Of the unverified-verdict findings: " + ", ".join(breakdown) + "."
                )
                note_run.italic = True
                note_run.font.size = Pt(9)
                note_run.font.color.rgb = RGBColor(100, 100, 100)


# ---------------------------------------------------------------------------
# Trust-model summary
# ---------------------------------------------------------------------------

def _write_trust_model_summary(
    doc: Document,
    status_counts: dict,
    edit_action_counts: dict,
) -> None:
    """Render the per-status / per-edit-action histograms.

    Every finding receives one status and one
    edit-action label. The table here gives a top-of-report at-a-glance
    picture of how much of the run is supported vs. uncertain, before
    the reader gets to individual findings. The
    severity table above answers "how many issues are critical?"; this
    one answers "how many of them are actually trustworthy?"
    """
    total_status = sum(status_counts.values())
    if total_status == 0:
        return

    doc.add_heading("Trust Model Summary", level=1)
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Every finding is tagged with a trust status and an edit-action "
        "label. The two histograms below give the at-a-glance picture; "
        "individual findings carry the same labels inline."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(100, 100, 100)
    intro.paragraph_format.space_after = Pt(6)

    # --- Status histogram table ---
    visible_statuses = [
        s for s in STATUS_DISPLAY_ORDER if status_counts.get(s, 0) > 0
    ]
    if visible_statuses:
        table = doc.add_table(rows=2, cols=len(visible_statuses))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col_idx, status in enumerate(visible_statuses):
            hex_color = STATUS_SHADING[status]
            header_cell = table.rows[0].cells[col_idx]
            _set_cell_shading(header_cell, hex_color)
            header_cell.text = ""
            p = header_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(status_label(status))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)

            count_cell = table.rows[1].cells[col_idx]
            count_cell.text = ""
            p = count_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(status_counts[status]))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = STATUS_COLORS[status]

    # --- Edit-action histogram (compact inline form) ---
    visible_actions = [
        a for a in EDIT_ACTION_DISPLAY_ORDER if edit_action_counts.get(a, 0) > 0
    ]
    if visible_actions:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        label_run = para.add_run("Edit eligibility: ")
        label_run.bold = True
        for i, action in enumerate(visible_actions):
            if i > 0:
                sep = para.add_run("  •  ")
                sep.font.color.rgb = RGBColor(160, 160, 160)
            count_run = para.add_run(f"{edit_action_counts[action]} ")
            count_run.bold = True
            count_run.font.color.rgb = EDIT_ACTION_COLORS[action]
            name_run = para.add_run(edit_action_label(action).lower())
            name_run.font.color.rgb = EDIT_ACTION_COLORS[action]
        para.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Alerts
# ---------------------------------------------------------------------------

def _write_alert_section(
    doc: Document,
    title: str,
    description: str,
    alerts: list[dict],
) -> None:
    """Render one named alert section with per-file bullet groups.

    Every deterministic alert category goes through this helper
    so the section layout stays consistent. Each section also gets a
    "(deterministic check)" suffix so the reader can tell at a glance which
    alerts came from local rules vs. LLM findings — deterministic findings
    are clearly labeled.
    """
    if not alerts:
        return
    doc.add_heading(f"{title} (deterministic check)", level=2)
    _add_styled_paragraph(doc, description, size=10, space_after=6)
    by_file: dict[str, list[dict]] = {}
    for alert in alerts:
        by_file.setdefault(alert.get("filename", ""), []).append(alert)
    for filename, items in by_file.items():
        para = doc.add_paragraph()
        para.add_run(f"{filename}").bold = True
        for alert in items[:5]:
            context = alert.get("context", alert.get("match", ""))
            doc.add_paragraph(context, style='List Bullet')
        if len(items) > 5:
            doc.add_paragraph(
                f"... and {len(items) - 5} more",
                style='List Bullet',
            )


def _uniform_cycle_cadence(years: tuple[str, ...]) -> int | None:
    """Return the uniform year gap between consecutive cycle years, or None.

    Drives the "publishes cycles every N years" parenthetical in the
    invalid-cycle alert intro. ``None`` (non-numeric years, fewer than two
    years, or an irregular sequence) drops the cadence claim and the intro
    falls back to a plain "known cycle years" list.
    """
    try:
        values = sorted(int(y) for y in years)
    except (TypeError, ValueError):
        return None
    if len(values) < 2:
        return None
    gaps = {b - a for a, b in zip(values, values[1:])}
    if len(gaps) != 1:
        return None
    gap = gaps.pop()
    return gap if gap > 1 else None


def _write_alerts(
    doc: Document,
    leed_alerts: list[dict],
    placeholder_alerts: list[dict],
    *,
    code_cycle_alerts: list[dict] | None = None,
    structural_alerts: list[dict] | None = None,
    naming_alerts: list[dict] | None = None,
    template_marker_alerts: list[dict] | None = None,
    invalid_code_cycle_alerts: list[dict] | None = None,
    duplicate_paragraph_alerts: list[dict] | None = None,
    polity_alerts: list[dict] | None = None,
    module: ReviewModule | None = None,
) -> None:
    """Write the Alerts section with every deterministic-check category.

    Previously this rendered only ``leed_alerts`` and
    ``placeholder_alerts``; ``code_cycle_alerts`` / ``structural_alerts`` /
    ``naming_alerts`` were collected during preflight but silently dropped
    before the report saw them, so users had to read the log to discover
    them. The new optional kwargs default to ``None`` (treated as empty) so
    legacy callers — including older snapshot tests — keep working.

    The stale/invalid code-cycle section headings and intros render from
    the run ``module``'s detector vocabulary (jurisdiction wording + the
    published-cycle-year list). ``None`` resolves to the default module,
    so legacy callers keep the original California wording byte-for-byte.
    """
    code_cycle_alerts = code_cycle_alerts or []
    structural_alerts = structural_alerts or []
    naming_alerts = naming_alerts or []
    template_marker_alerts = template_marker_alerts or []
    invalid_code_cycle_alerts = invalid_code_cycle_alerts or []
    duplicate_paragraph_alerts = duplicate_paragraph_alerts or []
    polity_alerts = polity_alerts or []
    if not any((
        leed_alerts,
        placeholder_alerts,
        code_cycle_alerts,
        structural_alerts,
        naming_alerts,
        template_marker_alerts,
        invalid_code_cycle_alerts,
        duplicate_paragraph_alerts,
        polity_alerts,
    )):
        return

    doc.add_heading("Alerts", level=1)
    _add_styled_paragraph(
        doc,
        "These items were detected locally by deterministic rules; no LLM "
        "tokens were spent on them. They sit alongside the LLM findings "
        "below.",
        size=10,
        space_after=6,
    )
    _write_alert_section(
        doc,
        "LEED References Detected",
        "The following LEED references were found. Since this is not a "
        "LEED project, these should be removed:",
        leed_alerts,
    )
    _write_alert_section(
        doc,
        "Unresolved Placeholders",
        "The following editorial placeholders need to be resolved:",
        placeholder_alerts,
    )
    _write_alert_section(
        doc,
        "Unresolved Template Markers",
        "The following TODO / FIXME / XXX / ??? markers are still in the "
        "spec and should be resolved before issuing:",
        template_marker_alerts,
    )
    vocab = (module if module is not None else get_module(None)).detector_vocabulary
    jurisdiction = vocab.jurisdiction_label.strip()
    j_prefix = f"{jurisdiction} " if jurisdiction else ""
    _write_alert_section(
        doc,
        f"Stale {j_prefix}Code Cycle References",
        f"The following references cite a historical {j_prefix}code cycle "
        "rather than the one selected for this review:",
        code_cycle_alerts,
    )
    # The parenthetical lists the *published* cycles — the vocabulary's
    # plausible_cycle_years. (valid_cycle_years additionally admits
    # anticipated future cycles like 2028, which aren't published and
    # would falsify the "publishes cycles" claim.)
    published_years = ", ".join(vocab.plausible_cycle_years)
    cadence = _uniform_cycle_cadence(vocab.plausible_cycle_years)
    if published_years and jurisdiction and cadence:
        known_years = (
            f" ({jurisdiction} publishes cycles every {cadence} years: "
            f"{published_years})"
        )
    elif published_years:
        known_years = f" (known cycle years: {published_years})"
    else:
        known_years = ""
    _write_alert_section(
        doc,
        f"Invalid {j_prefix}Code Cycle Years",
        f"The following references cite a year that is not a real "
        f"{j_prefix}code cycle{known_years}. These are likely typos:",
        invalid_code_cycle_alerts,
    )
    _write_alert_section(
        doc,
        "Structural Issues",
        "Empty sections and duplicate section headings detected in the "
        "spec body:",
        structural_alerts,
    )
    _write_alert_section(
        doc,
        "Duplicate Paragraphs",
        "These paragraphs appear verbatim more than once in the same spec — "
        "likely copy-paste mistakes:",
        duplicate_paragraph_alerts,
    )
    _write_alert_section(
        doc,
        "Inconsistent Filenames",
        "These files use a CSI naming style that differs from the project's "
        "dominant style:",
        naming_alerts,
    )
    # WS-4 (D-15): tokens whose suspiciousness is a pure function of the
    # project's country. Rendered only when a profile-bearing run produced
    # them, so profile-less reports are unchanged.
    _write_alert_section(
        doc,
        "Wrong-Polity Tokens",
        "These tokens belong to a different country's code/regulatory regime "
        "than the project location. Each entry notes why the token is "
        "suspicious for this project's country:",
        polity_alerts,
    )


# ---------------------------------------------------------------------------
# Per-finding evidence panel
# ---------------------------------------------------------------------------

def _write_evidence_panel(doc: Document, finding, vr) -> None:
    """Render the verifier-evidence panel under the collapsed Sources heading.

    Replaces the old URL-only Sources block.
    For every finding with a verification result, surface enough audit
    trail that a reviewer can answer "why did the verifier reach this
    verdict?" without leaving the report:

    - Verifier model (Sonnet 4.6 / Opus 4.8 / local).
    - Verification mode (LOCAL_SKIP / STRICT_STRUCTURED / STANDARD_REASONING
      / DEEP_REASONING) in human-readable form.
    - Search budget used ("N of M searches used") computed from
      ``web_search_requests`` and the severity-based ceiling.
    - Source quote: the verbatim snippet the verifier said
      it relied on. Rendered as an indented italic blockquote so it is
      visually distinct from prose paragraphs.
    - Verifier rationale: the model's explanation. Moved here from
      its old location above the Sources heading so it sits adjacent
      to the supporting quote.
    - Escalation history (when applicable): "Initial verdict … → Final
      verdict …" so a reviewer can see when the two models disagreed.
    - Accepted source URLs ("Web/code evidence").
    - Rejected source URLs.

    The Sources heading is collapsed-by-default — same behavior as
    before. Every paragraph inside the panel carries
    ``outlineLvl=8`` so Word's open-time collapse treats them all as
    part of the heading's zone.
    """
    # Locate the dataclass-level fields once with safe defaults so a
    # legacy resume payload (or a test-stub VerificationResult) can be
    # rendered without breaking.
    model_used = (getattr(vr, "model_used", "") or "").strip()
    mode_label = _verification_mode_label(getattr(vr, "verification_mode", "") or "")
    web_search_requests = int(getattr(vr, "web_search_requests", 0) or 0)
    severity_budget = web_search_max_uses_for_severity(
        getattr(finding, "severity", None)
    )
    source_quote = (getattr(vr, "source_quote", "") or "").strip()
    explanation = (getattr(vr, "explanation", "") or "").strip()
    escalation_attempted = bool(getattr(vr, "escalation_attempted", False))
    accepted = list(vr.sources or [])
    rejected = list(getattr(vr, "rejected_sources", []) or [])
    # web_fetch evidence. STRICT_STRUCTURED /
    # LOCAL_SKIP modes never attach the fetch tool, so this is 0/[] for
    # them; STANDARD/DEEP modes attach the tool but the model may not
    # have used it, so 0/[] is also the common case there.
    web_fetch_requests = int(getattr(vr, "web_fetch_requests", 0) or 0)
    fetched_sources = list(getattr(vr, "fetched_sources", []) or [])

    # The panel is rendered whenever ``finding.verification`` exists.
    # Even local-skip findings get the panel — they carry
    # ``verification_mode="local_skip"`` and an explanation, both of
    # which are auditable signal.
    sources_heading = doc.add_heading("Sources", level=4)
    _set_paragraph_collapsed(sources_heading)

    # --- Verifier model ---
    if model_used:
        para = doc.add_paragraph()
        _set_paragraph_outline_level(para, 8)
        label = para.add_run("Verifier model: ")
        label.bold = True
        label.font.size = Pt(9)
        label.font.color.rgb = RGBColor(100, 100, 100)
        body = para.add_run(model_used)
        body.font.size = Pt(9)
        body.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(2)

    # --- Verification mode ---
    if mode_label:
        para = doc.add_paragraph()
        _set_paragraph_outline_level(para, 8)
        label = para.add_run("Verification mode: ")
        label.bold = True
        label.font.size = Pt(9)
        label.font.color.rgb = RGBColor(100, 100, 100)
        body = para.add_run(mode_label)
        body.font.size = Pt(9)
        body.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(2)

    # --- Search budget used ---
    # Local-skip findings never invoke web_search (budget 0/M), but the
    # line is still useful — it makes "no external check ran" explicit.
    # We render the line whenever web_search was at least attempted
    # (web_search_requests > 0) OR when the verifier ran in a mode that
    # could have used the search tool. For local_skip we suppress
    # because "0 of N searches used" is misleading — the mode doesn't
    # use the search tool by design.
    mode_raw = (getattr(vr, "verification_mode", "") or "").strip().lower()
    if mode_raw != "local_skip":
        para = doc.add_paragraph()
        _set_paragraph_outline_level(para, 8)
        label = para.add_run("Search budget used: ")
        label.bold = True
        label.font.size = Pt(9)
        label.font.color.rgb = RGBColor(100, 100, 100)
        # When the verifier used web_fetch,
        # append the fetch count in the same line so a reviewer sees
        # "Searches: N, Full-page fetches: M" at a glance. The fetch
        # count is only rendered when > 0 because most verifications
        # never need a fetch — keeping the line short for the common
        # path matters more than uniformity.
        if web_fetch_requests > 0:
            usage_text = (
                f"Searches: {web_search_requests} of {severity_budget}, "
                f"Full-page fetches: {web_fetch_requests}"
            )
        else:
            usage_text = (
                f"{web_search_requests} of {severity_budget} searches used"
            )
        body = para.add_run(usage_text)
        body.font.size = Pt(9)
        body.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(2)

    # --- Source quote ---
    if source_quote:
        label_para = doc.add_paragraph()
        _set_paragraph_outline_level(label_para, 8)
        label_run = label_para.add_run("Source quote (verbatim from search result):")
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(100, 100, 100)
        label_para.paragraph_format.space_after = Pt(2)

        quote_para = doc.add_paragraph()
        _set_paragraph_outline_level(quote_para, 8)
        # Indent the blockquote so it stands apart from the
        # surrounding labels; italic + slightly smaller font signals
        # "this is verbatim source content, not commentary".
        quote_para.paragraph_format.left_indent = Inches(0.35)
        quote_run = quote_para.add_run(source_quote)
        quote_run.italic = True
        quote_run.font.size = Pt(9)
        quote_run.font.color.rgb = RGBColor(80, 80, 80)
        quote_para.paragraph_format.space_after = Pt(4)

    # --- Verifier rationale (moved from above) ---
    # Label kept as "Verification rationale:" so existing report
    # consumers and the label-rename invariants continue to
    # find the field.
    if explanation:
        para = doc.add_paragraph()
        _set_paragraph_outline_level(para, 8)
        label = para.add_run("Verification rationale: ")
        label.bold = True
        label.font.size = Pt(9)
        label.font.color.rgb = RGBColor(100, 100, 100)
        body = para.add_run(explanation)
        body.font.size = Pt(9)
        body.font.color.rgb = RGBColor(100, 100, 100)
        para.paragraph_format.space_after = Pt(3)

    # --- Escalation history ---
    if escalation_attempted:
        initial_verdict = (getattr(vr, "initial_verdict", "") or "").strip()
        initial_model = (getattr(vr, "initial_model", "") or "").strip()
        final_verdict = (getattr(vr, "verdict", "") or "").strip()
        final_model = model_used
        escalation_reason = (getattr(vr, "escalation_reason", "") or "").strip()
        changed = bool(getattr(vr, "escalation_changed_verdict", False))
        # The stricter "both grounded AND
        # verdicts differ" flag (vs. ``escalation_changed_verdict``
        # which also fires on initial-UNVERIFIED-then-CONFIRMED). When
        # set, the finding renders as VERIFIED_CONTESTED at the
        # top-level status badge and the panel below adds the
        # initial-pass citations side-by-side with the final-pass
        # citations so a reviewer can see "here are the sources Sonnet
        # cited, here are the sources Opus cited" without leaving the
        # finding entry.
        models_disagreed = bool(getattr(vr, "models_disagreed", False))
        initial_sources = list(getattr(vr, "initial_sources", []) or [])

        # Bold inline label.
        para = doc.add_paragraph()
        _set_paragraph_outline_level(para, 8)
        label = para.add_run("Escalation history: ")
        label.bold = True
        label.font.size = Pt(9)
        # Highlight in red-orange when the escalation actually changed
        # the verdict — that's the "two models disagreed" signal a
        # reviewer most wants to see. The contested case
        # (both grounded, verdicts differ) gets the purple
        # VERIFIED_CONTESTED color so the panel matches the top-level
        # status badge.
        if models_disagreed:
            label_color = RGBColor(128, 0, 128)  # Purple — VERIFIED_CONTESTED
        elif changed:
            label_color = RGBColor(178, 34, 34)  # Firebrick — verdict changed
        else:
            label_color = RGBColor(100, 100, 100)  # Gray — neutral
        label.font.color.rgb = label_color
        parts = []
        if initial_verdict:
            initial_summary = (
                f"{initial_verdict} from {initial_model}"
                if initial_model
                else initial_verdict
            )
            parts.append(f"Initial verdict: {initial_summary}")
        if final_verdict:
            final_summary = (
                f"{final_verdict} from {final_model}"
                if final_model
                else final_verdict
            )
            parts.append(f"Final verdict: {final_summary}")
        sentence = " → ".join(parts) if parts else "escalated"
        if escalation_reason:
            sentence += f". Reason: {escalation_reason}"
        # Surface the contested state explicitly so the
        # inline sentence stays self-explanatory even when the panel
        # is read in isolation (resume-state JSON dumps, exported
        # report scanned without the top-level status badge nearby).
        if models_disagreed:
            sentence += (
                " The two models disagreed on this finding while both "
                "produced grounded verdicts; manual review recommended."
            )
        elif changed:
            sentence += " (models disagreed)"
        sentence += "."
        body = para.add_run(sentence)
        body.font.size = Pt(9)
        body.font.color.rgb = label_color
        para.paragraph_format.space_after = Pt(3)

        # When the disagreement is real (both
        # grounded), render the initial verifier's accepted citations
        # as a follow-up line. The final verifier's citations already
        # render below under "Web/code evidence" so a reviewer reading
        # the panel top-to-bottom sees Sonnet's sources here, Opus's
        # sources below, and can compare side-by-side. Omitted when
        # there are no initial citations to show (legacy results, or
        # the stricter ``models_disagreed`` flag was never set).
        if models_disagreed and initial_sources:
            initial_label_para = doc.add_paragraph()
            _set_paragraph_outline_level(initial_label_para, 8)
            initial_label_run = initial_label_para.add_run(
                f"Initial verifier sources ({initial_model or 'initial pass'}):"
            )
            initial_label_run.font.size = Pt(9)
            initial_label_run.bold = True
            initial_label_run.font.color.rgb = RGBColor(128, 0, 128)
            initial_label_para.paragraph_format.space_after = Pt(2)

            initial_sources_para = doc.add_paragraph()
            _set_paragraph_outline_level(initial_sources_para, 8)
            initial_sources_para.paragraph_format.space_after = Pt(3)
            for i, url in enumerate(initial_sources):
                if i > 0:
                    initial_sources_para.add_run("  •  ")
                run = initial_sources_para.add_run(url)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(59, 130, 246)

    # --- Accepted source URLs ("Web/code evidence") ---
    if accepted:
        label_para = doc.add_paragraph()
        _set_paragraph_outline_level(label_para, 8)
        label_run = label_para.add_run(
            "Web/code evidence (cited and found in search results):"
        )
        label_run.font.size = Pt(9)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor(0, 100, 0)

        para = doc.add_paragraph()
        _set_paragraph_outline_level(para, 8)
        para.paragraph_format.space_after = Pt(3)
        for i, url in enumerate(accepted):
            if i > 0:
                para.add_run("  •  ")
            run = para.add_run(url)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(59, 130, 246)

    # --- Rejected source URLs ---
    if rejected:
        label_para = doc.add_paragraph()
        _set_paragraph_outline_level(label_para, 8)
        label_run = label_para.add_run(
            "Unsupported / rejected sources (cited by the model but not present in web_search results):"
        )
        label_run.font.size = Pt(9)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor(192, 0, 0)

        rej_para = doc.add_paragraph()
        _set_paragraph_outline_level(rej_para, 8)
        rej_para.paragraph_format.space_after = Pt(3)
        for i, entry in enumerate(rejected):
            if i > 0:
                rej_para.add_run("  •  ")
            url = entry.get("url") if isinstance(entry, dict) else str(entry)
            reason = entry.get("reason") if isinstance(entry, dict) else ""
            url_run = rej_para.add_run(url or "(empty)")
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(128, 128, 128)
            url_run.font.italic = True
            if reason:
                reason_run = rej_para.add_run(f" [{reason}]")
                reason_run.font.size = Pt(9)
                reason_run.font.color.rgb = RGBColor(192, 0, 0)

    # --- Full-text sources consulted ---
    # When the verifier used web_fetch, list the URLs it pulled in full
    # in a dedicated sub-section so a reviewer can tell at a glance which
    # sources were skimmed (web_search snippets) vs. read in depth
    # (web_fetch). The accepted-citations block above ("Web/code
    # evidence") already mixes both kinds; this block answers the
    # narrower question "which pages did the verifier read in full?"
    # Suppressed when no fetches happened so the panel stays compact
    # on the common path.
    if fetched_sources:
        label_para = doc.add_paragraph()
        _set_paragraph_outline_level(label_para, 8)
        label_run = label_para.add_run(
            "Full-text sources consulted (retrieved via web_fetch):"
        )
        label_run.font.size = Pt(9)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor(0, 100, 0)

        fetch_para = doc.add_paragraph()
        _set_paragraph_outline_level(fetch_para, 8)
        fetch_para.paragraph_format.space_after = Pt(3)
        for i, url in enumerate(fetched_sources):
            if i > 0:
                fetch_para.add_run("  •  ")
            url_run = fetch_para.add_run(url)
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(59, 130, 246)

    # --- Force-refresh hint for cache replays ---
    # A workflow hint, not a programmatic feature: tells the reviewer
    # exactly where to delete the entry if they want fresh verification.
    # Rendered only for cache-hit results so non-replayed findings stay
    # uncluttered.
    if (getattr(vr, "cache_status", "") or "") == "hit":
        hint_para = doc.add_paragraph()
        _set_paragraph_outline_level(hint_para, 8)
        hint_run = hint_para.add_run(
            "To force re-verification of this finding, delete its entry from "
            f"{default_cache_path()}."
        )
        hint_run.font.size = Pt(9)
        hint_run.font.italic = True
        hint_run.font.color.rgb = RGBColor(128, 128, 128)
        hint_para.paragraph_format.space_after = Pt(3)


# ---------------------------------------------------------------------------
# Single finding entry (collapsible via Heading 3)
# ---------------------------------------------------------------------------

def _write_finding_entry(doc: Document, finding, index: int) -> None:
    """Write a single finding as a collapsible block.

    The finding header is rendered as a Heading 3 paragraph, which enables
    Word's native heading-collapse feature. Users can click the collapse
    triangle that appears on hover to hide the finding's body content.

    Trust-model rendering:
        - Adds a "Status" line right under the header so the trust-model
          status is the first thing readers see (avoid presenting all
          findings as equally certain).
        - Adds an "Edit eligibility" line so readers can tell at a glance
          whether the finding carries a suggested edit or is report-only.
        - Renames the spec quote / web sources / rationale / rejected
          sources sub-labels so the four evidence concepts are explicit
          rather than implied.

    Layout:
        Heading 3:  [SEVERITY] 92% — filename.docx — Section ref
        Normal:     Status: <status>  •  Edit: <edit_action>
        Normal:     Issue: ...
        Normal:     Action: ...
        Normal:     Spec evidence: ... (existingText, red)
        Normal:     Proposed replacement: ... (green)
        Normal:     Reference: ... (blue)
        Normal:     Verification verdict: ... (if applicable)
        Normal:     Verification rationale: ... (if applicable)
        Heading 4:  Sources (collapsed by default)
        Normal:       Web/code evidence: <accepted_sources>
        Normal:       Unsupported / rejected sources: <rejected_sources>
    """
    severity_color = SEVERITY_COLORS.get(finding.severity, RGBColor(0, 0, 0))
    conf_tier = _confidence_tier(finding.confidence)
    conf_color = CONFIDENCE_COLORS[conf_tier]

    # --- Finding header as Heading 3 (enables Word collapse) ---
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 3']
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


    # Index + severity badge
    run = para.add_run(f"{index}. [{finding.severity}] ")
    run.bold = True
    run.font.color.rgb = severity_color
    run.font.size = Pt(11)
    # Confidence (the review model's pre-verification certainty).
    # Once verification reaches a verdict that supersedes it
    # (Verified — supported / contradicted / contested / Disputed), the
    # confidence % is dropped from the header: that finding's trust signal
    # is the verdict — shown on the Status line and the Verification
    # verdict line below — not this pre-verification number, which can read
    # misleadingly low on a confirmed finding (or high on a disputed one).
    # The number is preserved as a de-emphasized footnote on the Status
    # line. For not-yet-verified findings the % stays the prominent,
    # primary signal in the header.
    confidence_superseded = verdict_supersedes_confidence(finding)
    if not confidence_superseded:
        run = para.add_run(f"{finding.confidence:.0%} ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = conf_color
    # Separator
    run = para.add_run("— ")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(11)
    # Filename
    run = para.add_run(finding.fileName or "Unknown")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Section (inline in header for compact view)
    if finding.section:
        run = para.add_run(f" — {finding.section}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Body content (Normal paragraphs, hidden when heading is collapsed) ---

    # --- Status + edit-action line ---
    # The trust-model status renders right under the header so readers
    # see "Verified — supported" / "Disputed" / "Insufficient evidence"
    # before they read the issue. The edit-action label sits on the
    # same line so a reader scanning by finding can also see whether
    # there's an actionable edit attached.
    status = classify_status(finding)
    edit_action = classify_edit_action(finding)
    status_color = STATUS_COLORS[status]
    action_color = EDIT_ACTION_COLORS[edit_action]

    status_para = doc.add_paragraph()
    status_label_run = status_para.add_run("Status: ")
    status_label_run.bold = True
    status_label_run.font.size = Pt(10)
    glyph_run = status_para.add_run(f"{status_glyph(status)} ")
    glyph_run.bold = True
    glyph_run.font.color.rgb = status_color
    glyph_run.font.size = Pt(10)
    value_run = status_para.add_run(status_label(status))
    value_run.bold = True
    value_run.font.color.rgb = status_color
    value_run.font.size = Pt(10)
    # When the verifier consumed its full
    # mode-scaled search budget without producing a grounded verdict,
    # append a "(search budget exhausted)" sub-label so the reviewer
    # sees the actionable signal inline. The status itself stays
    # INSUFFICIENT_EVIDENCE — the trust level is the same as any other
    # unground UNVERIFIED — but the sub-label distinguishes "verifier
    # had no headroom" from "verifier ran out of evidence at search 2
    # of 7". Same color as the status so the badge reads as part of
    # the status, not a separate field.
    if is_budget_exhausted(finding):
        budget_run = status_para.add_run(" (search budget exhausted)")
        budget_run.font.size = Pt(10)
        budget_run.font.italic = True
        budget_run.font.color.rgb = status_color
    # Separator + edit-action.
    sep_run = status_para.add_run("  •  ")
    sep_run.font.color.rgb = RGBColor(160, 160, 160)
    sep_run.font.size = Pt(10)
    edit_label_run = status_para.add_run("Edit: ")
    edit_label_run.bold = True
    edit_label_run.font.size = Pt(10)
    edit_value_run = status_para.add_run(edit_action_label(edit_action))
    edit_value_run.bold = True
    edit_value_run.font.color.rgb = action_color
    edit_value_run.font.size = Pt(10)

    # --- Cache-replay badge ---
    # When the verifier result came from a cache hit, render an inline
    # badge showing the entry's age so a reviewer can spot stale verdicts
    # without expanding the Sources panel. Color tier:
    #   amber  for <30 days, orange for 30-90 days, red for >90 days.
    # Suppressed for non-hit results, for legacy resume payloads with
    # no ``cache_entry_created_ts`` recorded, and for clock-skew
    # cases where the recorded timestamp is in the future.
    age_days = _cache_entry_age_days(getattr(finding, "verification", None))
    if age_days is not None:
        badge_color = CACHE_AGE_COLORS[_cache_age_tier(age_days)]
        cache_sep_run = status_para.add_run("  •  ")
        cache_sep_run.font.color.rgb = RGBColor(160, 160, 160)
        cache_sep_run.font.size = Pt(10)
        cache_badge_run = status_para.add_run(
            f"Cache replay — {age_days}d old"
        )
        cache_badge_run.bold = True
        cache_badge_run.font.color.rgb = badge_color
        cache_badge_run.font.size = Pt(10)

    # Review-confidence footnote. When the header suppressed the
    # confidence % (a verdict supersedes it), surface the review model's
    # pre-verification confidence here as a small, gray, explicitly
    # labeled footnote — the number is preserved for anyone who wants it,
    # but rendered so it can't be mistaken for the post-verification trust
    # signal carried by the status/verdict.
    if confidence_superseded:
        conf_sep_run = status_para.add_run("  •  ")
        conf_sep_run.font.color.rgb = RGBColor(170, 170, 170)
        conf_sep_run.font.size = Pt(9)
        conf_note_run = status_para.add_run(
            f"review confidence {finding.confidence:.0%} (pre-verification)"
        )
        conf_note_run.italic = True
        conf_note_run.font.size = Pt(9)
        conf_note_run.font.color.rgb = RGBColor(150, 150, 150)

    status_para.paragraph_format.space_after = Pt(3)

    # --- Issue ---
    para = doc.add_paragraph()
    para.add_run("Issue: ").bold = True
    para.add_run(finding.issue or "")
    para.paragraph_format.space_after = Pt(3)

    # --- Action / edit-proposal block ---
    # The report distinguishes findings that carry an edit proposal
    # from ones that don't. REPORT_ONLY findings render an explicit
    # "No edit proposal — surfaced for review only" line so readers see
    # the finding without expecting an edit; findings with a proposal
    # keep the original Action / Existing / Replace With layout.
    #
    # The "Existing Text" label becomes "Spec evidence" so the
    # quoted-from-the-spec source is explicitly the *spec evidence*
    # concept, distinct from web/code evidence (sources)
    # and verification rationale (explanation).
    proposal = finding.as_edit_proposal()
    if proposal is None:
        para = doc.add_paragraph()
        run = para.add_run("Action: REPORT_ONLY")
        run.bold = True
        para.paragraph_format.space_after = Pt(3)

        # When the parser demoted an EDIT/DELETE/ADD because a required
        # field was missing, surface the specific reason inline so a
        # reader sees "the model claimed EDIT but no existingText was
        # provided" instead of the generic coordination/interpretation
        # explanation. Native REPORT_ONLY emissions keep the original
        # note text.
        demotion = (getattr(finding, "demotion_reason", None) or "").strip()
        if demotion:
            note_text = (
                "Edit proposal demoted to REPORT_ONLY at parse time: "
                f"{demotion}. The underlying finding is preserved; manual "
                "review required to determine a clean textual fix."
            )
        else:
            note_text = (
                "No edit proposal — surfaced for review only (coordination, "
                "interpretation, or multi-paragraph rewrite required)."
            )
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(note_text)
        note_run.font.italic = True
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(100, 100, 100)
        note_para.paragraph_format.space_after = Pt(3)
    else:
        para = doc.add_paragraph()
        para.add_run("Action: ").bold = True
        para.add_run(proposal.action_type or "")
        para.paragraph_format.space_after = Pt(3)

        # --- Spec evidence (red) — the text quoted from the source spec ---
        if proposal.existing_text:
            para = doc.add_paragraph()
            para.add_run("Spec evidence: ").bold = True
            run = para.add_run(proposal.existing_text)
            run.font.color.rgb = RGBColor(192, 0, 0)
            para.paragraph_format.space_after = Pt(3)

        # --- Proposed replacement (green) ---
        if proposal.replacement_text:
            para = doc.add_paragraph()
            para.add_run("Proposed replacement: ").bold = True
            run = para.add_run(proposal.replacement_text)
            run.font.color.rgb = RGBColor(0, 128, 0)
            para.paragraph_format.space_after = Pt(3)

    # --- Code reference (blue) ---
    if finding.codeReference:
        para = doc.add_paragraph()
        para.add_run("Reference: ").bold = True
        run = para.add_run(finding.codeReference)
        run.font.color.rgb = RGBColor(59, 130, 246)
        para.paragraph_format.space_after = Pt(3)

    # --- Verification verdict + correction ---
    if finding.verification:
        vr = finding.verification
        verdict_color = VERDICT_COLORS.get(vr.verdict, VERDICT_COLORS["UNVERIFIED"])
        verdict_icon = VERDICT_ICONS.get(vr.verdict, "—")

        para = doc.add_paragraph()
        run = para.add_run(f"Verification verdict: {verdict_icon} {vr.verdict}")
        run.bold = True
        run.font.color.rgb = verdict_color
        para.paragraph_format.space_after = Pt(3)

        if vr.verdict == "CORRECTED" and vr.correction:
            para = doc.add_paragraph()
            para.add_run("Correction: ").bold = True
            run = para.add_run(vr.correction)
            run.font.color.rgb = RGBColor(204, 132, 0)  # Amber
            para.paragraph_format.space_after = Pt(3)

        # --- Evidence panel ---
        # Rendered under the existing collapsed-by-default "Sources" Heading
        # 4. Order: verifier model → verification mode → search budget →
        # source quote (blockquote) → verifier rationale → escalation
        # history → accepted source URLs → rejected source URLs.
        # All paragraphs in the panel carry outlineLvl=8 so Word's
        # open-time collapse zone hides them along with the URLs that
        # used to live alone under this heading.
        _write_evidence_panel(doc, finding, vr)


# ---------------------------------------------------------------------------
# Findings section
# ---------------------------------------------------------------------------

def _write_findings_section(doc: Document, review) -> None:
    """Write per-spec findings grouped by severity, then spec file, then confidence.

    Uses heading hierarchy for Word-native collapse support:
    - Title (level 0): "Findings" (stamped outlineLvl=0 — Title has no
      native outline level, so without it the header is body text to
      Word's collapse logic and invisible to the Navigation Pane)
    - Heading 1: Severity group (e.g., "CRITICAL (1)")
    - Heading 3: Individual finding header (collapsible)
    - Normal: Finding body content

    Within a severity group, findings are grouped by spec file so one spec's
    issues stay contiguous (the prior confidence-only sort scattered a single
    spec's findings across the whole severity block); confidence orders the
    findings within each file.
    """
    findings_heading = doc.add_heading("Findings", level=0)
    _set_paragraph_outline_level(findings_heading, 0)

    if review.total_count == 0:
        _add_styled_paragraph(
            doc,
            "No issues found.",
            size=12,
            color=RGBColor(0, 128, 0),
        )
        return

    finding_number = 0  # Running counter across all severities

    for severity in SEVERITY_ORDER:
        # Group by spec file (the fileName prefix is the CSI section number,
        # so a lexical sort yields CSI order), then confidence descending.
        # Keeps one spec's findings contiguous instead of interleaving every
        # spec by confidence across the severity block.
        severity_findings = sorted(
            [f for f in review.findings if f.severity == severity],
            key=lambda f: (f.fileName or "", -f.confidence),
        )
        if not severity_findings:
            continue

        # Severity sub-heading with colored text
        heading = doc.add_heading(
            f"{severity} ({len(severity_findings)})", level=1,
        )
        for run in heading.runs:
            run.font.color.rgb = SEVERITY_COLORS.get(severity, RGBColor(0, 0, 0))

        for finding in severity_findings:
            finding_number += 1
            _write_finding_entry(doc, finding, finding_number)


# ---------------------------------------------------------------------------
# Cross-spec coordination section
# ---------------------------------------------------------------------------

def _write_cross_check_section(doc: Document, cross_check_result) -> None:
    """Write cross-spec coordination section and explicit status.

    Cross-check findings are rendered with the same collapsible structure
    as per-spec findings.

    The section header must carry an explicit outline level and own its
    page break: the finding immediately before this section ends with a
    collapsed-by-default "Sources" Heading 4, whose collapse zone runs
    until the next paragraph with an outline level at or above it. A bare
    Title-styled header (no native outline level) and a standalone
    page-break paragraph are both body text to that logic, so Word folded
    the entire section banner into the last finding's collapsed Sources
    panel — hidden on open, absent from the Navigation Pane, with the
    coordination findings dangling under the last severity group.
    """
    if not cross_check_result:
        return

    heading = doc.add_heading("Cross-Spec Coordination", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.page_break_before = True
    _set_paragraph_outline_level(heading, 0)

    status = getattr(cross_check_result, "cross_check_status", None)
    count = len(cross_check_result.findings)
    subtitle = doc.add_paragraph()
    if status == "skipped":
        run = subtitle.add_run(f"Cross-check was skipped: {cross_check_result.thinking}")
    elif status == "failed":
        run = subtitle.add_run(f"Cross-check failed: {cross_check_result.error}")
    elif status == "completed" and count == 0:
        run = subtitle.add_run("Cross-check completed — no coordination issues found.")
    else:
        # Render the cross-check model's human label from the pricing table
        # so the report can't drift from the configured model (the string
        # was previously hardcoded and went stale on model bumps).
        price = price_for(CROSS_CHECK_MODEL_DEFAULT)
        model_label = price.label if price else CROSS_CHECK_MODEL_DEFAULT
        run = subtitle.add_run(
            f"{model_label} coordination analysis — "
            f"{count} issue{'s' if count != 1 else ''} found."
        )
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle.paragraph_format.space_after = Pt(12)

    if status in ("skipped", "failed"):
        return

    # Sort by severity, then group by spec file (fileName prefix is the CSI
    # section number, so a lexical sort yields CSI order), then confidence
    # descending — mirrors _write_findings_section so a spec's coordination
    # findings stay contiguous instead of scattering by confidence.
    severity_rank = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "GRIPES": 3}
    sorted_findings = sorted(
        cross_check_result.findings,
        key=lambda f: (severity_rank.get(f.severity, 99), f.fileName or "", -f.confidence),
    )

    for idx, finding in enumerate(sorted_findings, 1):
        _write_finding_entry(doc, finding, idx)

    # Coordination summary narrative
    if cross_check_result.thinking:
        doc.add_heading("Coordination Summary", level=2)
        _write_narrative_text(doc, cross_check_result.thinking)


def _write_compliance_section(doc: Document, compliance_result) -> None:
    """Write the Local-Code Compliance findings section (WS-4).

    Mirrors :func:`_write_cross_check_section`: explicit outline level +
    page break so the previous finding's collapsed Sources panel cannot
    swallow the banner, findings rendered through the shared
    ``_write_finding_entry``, and an explicit red status line when the pass
    skipped or failed (invariant 8 — never silent). ``None`` (the phase
    never applied — every profile-less run) renders nothing.
    """
    if not compliance_result:
        return

    heading = doc.add_heading("Local-Code Compliance", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.page_break_before = True
    _set_paragraph_outline_level(heading, 0)

    status = getattr(compliance_result, "cross_check_status", None)
    count = len(compliance_result.findings)
    subtitle = doc.add_paragraph()
    if status == "skipped":
        run = subtitle.add_run(
            f"Compliance evaluation was skipped: {compliance_result.thinking}"
        )
    elif status == "failed":
        run = subtitle.add_run(
            f"Compliance evaluation failed: {compliance_result.error}"
        )
    elif status == "completed" and count == 0:
        run = subtitle.add_run(
            "Compliance evaluation completed — no missing or contradicted "
            "requirements found."
        )
    else:
        run = subtitle.add_run(
            "Evaluation against the researched location/client requirements — "
            f"{count} issue{'s' if count != 1 else ''} found."
        )
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle.paragraph_format.space_after = Pt(12)

    if status in ("skipped", "failed"):
        return

    severity_rank = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "GRIPES": 3}
    sorted_findings = sorted(
        compliance_result.findings,
        key=lambda f: (severity_rank.get(f.severity, 99), f.fileName or "", -f.confidence),
    )
    for idx, finding in enumerate(sorted_findings, 1):
        _write_finding_entry(doc, finding, idx)

    if compliance_result.thinking:
        doc.add_heading("Compliance Summary", level=2)
        _write_narrative_text(doc, compliance_result.thinking)


def _sanitize_markdown_line(line: str) -> str:
    """Strip markdown header prefixes from a line."""
    stripped = line
    while stripped.startswith('#'):
        stripped = stripped[1:]
    return stripped.strip() if line.startswith('#') else line


def _write_narrative_text(doc: Document, text: str) -> None:
    """Write multi-paragraph narrative text, splitting on double newlines.

    Strips markdown header formatting (## ...) since the cross-check
    prompt sometimes produces markdown despite instructions not to.
    """
    if '\n\n' in text:
        paragraphs = text.split('\n\n')
    else:
        paragraphs = text.split('\n')

    for para_text in paragraphs:
        para_text = _sanitize_markdown_line(para_text.strip())
        if para_text:
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.size = Pt(11)
            para.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def export_report(
    pipeline_result,
    output_path: Path,
) -> Path:
    """Export a complete review report to a Word document.

    Generates a formatted .docx file containing files reviewed, summary
    grid, alerts, per-spec findings, and cross-check findings.

    Each per-spec finding uses Heading 3 for its header line, enabling
    Word's native heading-collapse feature.

    Args:
        pipeline_result: PipelineResult from the review pipeline
        output_path: Path where the .docx file should be saved

    Returns:
        The output_path (for convenience / confirmation)

    Raises:
        ValueError: If pipeline_result has no review_result
        OSError: If the file cannot be written
    """
    if pipeline_result.review_result is None:
        raise ValueError("Cannot export report: no review results available")

    review = pipeline_result.review_result
    cross_check = pipeline_result.cross_check_result
    # WS-4 compliance output (``None`` on every profile-less run). Its
    # findings join the trust-model / banner statistics alongside review +
    # cross-check findings; the structured requirements profile drives the
    # Jurisdiction & Client Requirements section.
    compliance = getattr(pipeline_result, "compliance_result", None)
    requirements_profile = RequirementsProfile.from_dict(
        getattr(pipeline_result, "requirements_profile", None)
    )
    all_findings = list(review.findings)
    if cross_check and cross_check.findings:
        all_findings.extend(cross_check.findings)
    if compliance is not None and compliance.findings:
        all_findings.extend(compliance.findings)
    verification_stats = _summarize_verification_outcomes(all_findings)

    doc = Document()

    # Set default font (Arial 11pt — clean and professional)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Configure Heading 3 style for finding entries
    # Keep it compact so findings don't dominate vertical space
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(11)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(4)
    # Remove the default Heading 3 color so our per-run colors show through
    h3_style.font.color.rgb = RGBColor(0, 0, 0)

    # Set margins (1 inch sides, 0.75 top/bottom)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Build the report
    cycle_label = getattr(pipeline_result, "cycle_label", "2025") or "2025"
    # The run's module owns every domain-worded report surface (title,
    # jurisdiction phrasing, pinned-standards cycle, alert headings).
    # Legacy results / test doubles without a module_id resolve to the
    # default module.
    module = get_module(getattr(pipeline_result, "module_id", None))
    # Per-run project identity (city/state/country/client), if the run's
    # module collected one. Defensive: legacy results / test doubles have no
    # attribute, and a malformed/empty dict degrades to None (profile-less).
    project_profile = ProjectProfile.from_dict(
        getattr(pipeline_result, "project_profile", None)
    )
    # Specs whose review failed/truncated (never produced findings).
    # Defensive getattr keeps legacy callers / test doubles at empty.
    failed_review_specs = [
        str(name)
        for name in (getattr(pipeline_result, "failed_review_specs", None) or [])
    ]
    failed_review_count = len(failed_review_specs)
    _write_title_block(
        doc,
        review,
        pipeline_result.files_reviewed,
        cycle_label=cycle_label,
        failed_review_count=failed_review_count,
        module=module,
        profile=project_profile,
    )

    # Run Diagnostics banner. Renders right
    # after the title block so the operational picture (edit-suggested
    # counts, cache replays, verification failures, parse-time
    # demotions, cross-check status) is the first thing a reviewer
    # sees. Derived from data already on the findings + verification
    # stats; no resume state or persistence changes needed.
    run_diagnostics = _summarize_run_diagnostics(
        findings=all_findings,
        status_counts=verification_stats.get("status_counts", {}),
        edit_action_counts=verification_stats.get("edit_action_counts", {}),
        cross_check_result=cross_check,
        pipeline_result=pipeline_result,
        compliance_result=compliance,
    )
    _write_run_diagnostics_banner(doc, run_diagnostics)

    _write_files_reviewed(
        doc,
        pipeline_result.files_reviewed,
        failed_review_specs=set(failed_review_specs),
    )
    # WS-4 "Jurisdiction & Client Requirements" — between Files Reviewed and
    # the methodology note (D-13); renders only when the run researched a
    # requirements profile, so profile-less reports are byte-identical.
    if requirements_profile is not None:
        _write_requirements_section(doc, requirements_profile, compliance, module)
    cross_check_status = getattr(cross_check, "cross_check_status", None) if cross_check else None
    cross_check_reason = ""
    if cross_check and cross_check_status == "skipped":
        cross_check_reason = cross_check.thinking or ""
    elif cross_check and cross_check_status == "failed":
        cross_check_reason = cross_check.error or ""
    _write_methodology_note(
        doc,
        cross_check_enabled=(cross_check is not None),
        cycle_label=cycle_label,
        cross_check_status=cross_check_status,
        cross_check_reason=cross_check_reason,
        verification_stats=verification_stats,
        module=module,
    )
    # WS-5 "How the Drawings Informed This Review" — rendered only when the
    # run carried attached construction drawings (a drawing-impact result),
    # so a run without drawings is byte-identical. Placed right after the
    # methodology note (both answer "how was this review produced?") and
    # above the findings its links reference.
    drawing_impact = getattr(pipeline_result, "drawing_impact_result", None)
    if drawing_impact is not None:
        findings_by_id = {
            f.finding_id: f
            for f in all_findings
            if getattr(f, "finding_id", "")
        }
        _write_drawing_impact_section(doc, drawing_impact, findings_by_id)
    _write_summary_table(
        doc,
        review,
        cross_check,
        total_elapsed_seconds=getattr(pipeline_result, "total_elapsed_seconds", None),
    )

    # Trust-model histogram. Renders right after the severity
    # summary so the reader sees "how many issues are critical?" and
    # "how many of them are actually trustworthy?" together.
    _write_trust_model_summary(
        doc,
        verification_stats.get("status_counts", {}),
        verification_stats.get("edit_action_counts", {}),
    )

    # Fall back to ``getattr`` for the new alert lists so
    # ``_StubPipelineResult`` style ad-hoc test doubles (and any legacy
    # callers that build the result by hand without the new fields)
    # keep working. The real ``PipelineResult`` dataclass always has them.
    _write_alerts(
        doc,
        pipeline_result.leed_alerts,
        pipeline_result.placeholder_alerts,
        code_cycle_alerts=getattr(pipeline_result, "code_cycle_alerts", None),
        structural_alerts=getattr(pipeline_result, "structural_alerts", None),
        naming_alerts=getattr(pipeline_result, "naming_alerts", None),
        template_marker_alerts=getattr(pipeline_result, "template_marker_alerts", None),
        invalid_code_cycle_alerts=getattr(pipeline_result, "invalid_code_cycle_alerts", None),
        duplicate_paragraph_alerts=getattr(pipeline_result, "duplicate_paragraph_alerts", None),
        polity_alerts=getattr(pipeline_result, "polity_alerts", None),
        module=module,
    )
    _write_findings_section(doc, review)
    _write_cross_check_section(doc, cross_check)
    _write_compliance_section(doc, compliance)


    # Save
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return output_path