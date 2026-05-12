"""Surgical DOCX edit application utilities."""

from __future__ import annotations

from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from .edit_candidates import (
    SAFETY_AUTO_SAFE,
    SAFETY_AUTO_WITH_CAUTION,
    SAFETY_MANUAL_REVIEW,
    SAFETY_REPORT_ONLY,
)
from .edit_locator import EditLocation, LocatorResult


_WHITESPACE_RE = re.compile(r"[\s\u00A0]+")
_HEADING_HINT_RE = re.compile(r"^\s*(PART\s+\d+|SECTION\s+\d+(\.\d+)*)\b", flags=re.IGNORECASE)


def _normalize_text_for_add(text: str) -> str:
    normalized = unicodedata.normalize("NFC", text)
    normalized = _WHITESPACE_RE.sub(" ", normalized)
    return normalized.strip().casefold()


def _split_insert_paragraphs(text: str) -> list[str]:
    return [chunk.strip() for chunk in re.split(r"\n\s*\n+", text) if chunk.strip()]


def _paragraph_style_id(paragraph_element) -> str | None:
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        return None
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        return None
    return pstyle.get(qn("w:val"))


def _is_heading_like_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    has_alpha = any(ch.isalpha() for ch in stripped)
    is_upper = has_alpha and stripped == stripped.upper()
    short_text = len(stripped) <= 36 and not any(p in stripped for p in ".;:")
    return is_upper or short_text or bool(_HEADING_HINT_RE.match(stripped))


def _reference_style_for_text(anchor_index: int, body_children: list, text: str) -> str | None:
    anchor_style = _paragraph_style_id(body_children[anchor_index])
    if _is_heading_like_text(text):
        return anchor_style

    for idx in range(anchor_index + 1, len(body_children)):
        elem = body_children[idx]
        if not elem.tag.endswith("}p"):
            continue
        para_text = "".join(elem.itertext())
        if _is_heading_like_text(para_text):
            continue
        return _paragraph_style_id(elem) or anchor_style

    return anchor_style


def _build_paragraph_element(text: str, style_id: str | None, anchor_element) -> object:
    paragraph_element = OxmlElement("w:p")
    if style_id:
        ppr = OxmlElement("w:pPr")
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), style_id)
        ppr.append(pstyle)
        paragraph_element.append(ppr)
    else:
        anchor_ppr = anchor_element.find(qn("w:pPr"))
        if anchor_ppr is not None:
            paragraph_element.append(deepcopy(anchor_ppr))

    run_element = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    text_element.set(qn("xml:space"), "preserve")
    run_element.append(text_element)
    paragraph_element.append(run_element)
    return paragraph_element


def _insert_paragraphs_before(anchor_element, texts: list[str], style_id: str | None) -> int:
    inserted = 0
    for text in texts:
        paragraph_element = _build_paragraph_element(text, style_id, anchor_element)
        anchor_element.addprevious(paragraph_element)
        inserted += 1
    return inserted


def _insert_paragraphs_after(anchor_element, texts: list[str], style_id: str | None) -> int:
    inserted = 0
    for text in reversed(texts):
        paragraph_element = _build_paragraph_element(text, style_id, anchor_element)
        anchor_element.addnext(paragraph_element)
        inserted += 1
    return inserted


@dataclass
class EditAction:
    locator_result: LocatorResult
    location: EditLocation
    replacement_text: str | None
    action_type: str
    finding_index: int


@dataclass
class EditOutcome:
    action: EditAction
    status: str
    detail: str
    original_text: str
    new_text: str | None
    # Chunk 9 — tag outcomes that were refused because the target paragraph or
    # table cell carried unsafe Word markup (hyperlinks, field codes,
    # drawings, comments, tracked changes, bookmarks, etc.). The status is
    # still ``"skipped"`` so the existing applied/skipped/failed accounting
    # is unaffected, but the flag lets the report layer surface "auto-edit
    # refused due to unsafe Word markup" rather than burying it in a generic
    # skip reason. Defaults to False so legacy outcomes are unchanged.
    refused_unsafe_markup: bool = False


@dataclass
class EditReport:
    source_path: Path
    output_path: Path
    total_edits_attempted: int
    edits_applied: int
    edits_skipped: int
    edits_failed: int
    outcomes: list[EditOutcome]
    warnings: list[str]
    # Chunk 9 — Boolean signal that the edit pass was aborted before any
    # output was written because at least one auto-edit failed under the
    # configured all-or-none transactional policy. Default False keeps
    # legacy reports unchanged.
    aborted_transactional: bool = False


def _build_run_offset_map(paragraph: Paragraph) -> list[tuple[int, int, int]]:
    """Return (run_index, char_start, char_end) offsets for each run in a paragraph."""
    offsets: list[tuple[int, int, int]] = []
    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        start = cursor
        end = start + len(run.text)
        offsets.append((idx, start, end))
        cursor = end
    return offsets


def _replace_in_paragraph(paragraph: Paragraph, match_start: int, match_end: int, replacement: str) -> tuple[bool, str]:
    """Replace text slice [match_start:match_end] in paragraph without removing runs."""
    full_text = paragraph.text
    if match_start < 0 or match_end < match_start or match_end > len(full_text):
        return False, "Invalid match offsets for paragraph text length."

    expected = full_text[:match_start] + replacement + full_text[match_end:]
    if match_start == match_end and not replacement:
        return True, "No-op replacement."

    run_map = _build_run_offset_map(paragraph)
    affected: list[tuple[int, int, int]] = [
        entry for entry in run_map if entry[1] < match_end and entry[2] > match_start
    ]

    if not affected:
        if match_start == match_end:
            if not paragraph.runs:
                paragraph.add_run(replacement)
                return (paragraph.text == expected), "Inserted replacement into empty paragraph."
            first_run = paragraph.runs[0]
            local = max(0, min(len(first_run.text), match_start))
            first_run.text = first_run.text[:local] + replacement + first_run.text[local:]
            return (paragraph.text == expected), "Inserted replacement at run boundary."
        return False, "Could not map target range to paragraph runs."

    first_idx, first_start, _ = affected[0]
    last_idx, last_start, _ = affected[-1]

    first_run = paragraph.runs[first_idx]
    first_prefix_len = max(0, min(len(first_run.text), match_start - first_start))
    first_prefix = first_run.text[:first_prefix_len]

    if first_idx == last_idx:
        suffix_start = max(0, min(len(first_run.text), match_end - first_start))
        suffix = first_run.text[suffix_start:]
        first_run.text = first_prefix + replacement + suffix
    else:
        last_run = paragraph.runs[last_idx]
        suffix_start = max(0, min(len(last_run.text), match_end - last_start))
        suffix = last_run.text[suffix_start:]

        first_run.text = first_prefix + replacement
        for run_idx, _, _ in affected[1:-1]:
            paragraph.runs[run_idx].text = ""
        last_run.text = suffix

    if paragraph.text != expected:
        return False, "Run-level replacement verification failed."
    return True, "Replacement applied successfully."


def _delete_paragraph(paragraph: Paragraph) -> bool:
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        return False
    parent.remove(element)
    return True


def _is_whole_paragraph_delete(action: EditAction) -> bool:
    location = action.location
    return (
        action.action_type == "DELETE"
        and location.mapping.element_type == "paragraph"
        and location.match_start == 0
        and location.match_end == len(location.mapping.text)
    )


# ---------------------------------------------------------------------------
# Chunk 9 — unsafe WordprocessingML markup detection.
#
# Run-level surgery on paragraphs that carry hyperlinks, field codes,
# drawings/images, comments, tracked changes, bookmarks, or content controls
# can silently corrupt the underlying XML: removing or rewriting a run can
# leave dangling field characters, orphaned bookmark/comment ranges, broken
# hyperlink relationships, or inline drawings with no anchor. Per the plan,
# the safe behavior is to refuse the auto-edit and route the finding to
# manual review rather than risk corruption.
#
# The detector returns a structured :class:`UnsafeMarkupResult` so the
# caller can attach a refusal reason to the EditOutcome. Detection works on
# raw lxml elements and is intentionally cheap — it scans the paragraph's
# subtree for known unsafe tags via :func:`qn` lookups; no heuristics, no
# regex over rendered text.
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class UnsafeMarkupResult:
    """Outcome of scanning a paragraph/cell for edit-unsafe Word markup."""

    unsafe: bool
    reasons: tuple[str, ...]

    @property
    def detail(self) -> str:
        if not self.unsafe:
            return ""
        joined = ", ".join(self.reasons)
        return (
            "Auto-edit refused: paragraph contains unsafe Word markup "
            f"({joined}); routed to manual review to avoid document corruption."
        )


# Mapping from a WordprocessingML local tag name to a human-readable reason.
# Order matters only for the reasons list — we report every distinct hit so
# the EditOutcome detail is useful for an operator triaging the report.
_UNSAFE_TAGS: tuple[tuple[str, str], ...] = (
    ("w:hyperlink", "hyperlink"),
    ("w:fldChar", "field character"),
    ("w:instrText", "field instruction text"),
    ("w:fldSimple", "simple field"),
    ("w:drawing", "drawing/image"),
    ("w:pict", "embedded picture/object"),
    ("w:object", "OLE object"),
    ("w:commentRangeStart", "comment range"),
    ("w:commentRangeEnd", "comment range"),
    ("w:commentReference", "comment reference"),
    ("w:ins", "tracked insertion"),
    ("w:del", "tracked deletion"),
    ("w:moveFrom", "tracked move-from"),
    ("w:moveTo", "tracked move-to"),
    ("w:moveFromRangeStart", "tracked move range"),
    ("w:moveToRangeStart", "tracked move range"),
    ("w:bookmarkStart", "bookmark range"),
    ("w:bookmarkEnd", "bookmark range"),
    ("w:sdt", "content control / smart tag"),
    ("w:footnoteReference", "footnote reference"),
    ("w:endnoteReference", "endnote reference"),
    ("w:smartTag", "smart tag"),
    ("w:customXml", "custom XML element"),
)


def detect_unsafe_markup(element) -> UnsafeMarkupResult:
    """Return :class:`UnsafeMarkupResult` for the given paragraph/cell element.

    The detector walks the element's subtree once and records every distinct
    unsafe-tag class it sees. Returning the full set (instead of stopping at
    the first hit) lets the resulting EditOutcome explain exactly why the
    edit was refused — useful when the same paragraph carries, say, both a
    hyperlink and a tracked change.

    ``element`` may be the lxml element directly or any object that exposes
    one via ``_element`` (Paragraph / _Cell from python-docx both do).
    """
    if element is None:
        return UnsafeMarkupResult(False, ())
    raw = element if hasattr(element, "iter") else getattr(element, "_element", None)
    if raw is None:
        return UnsafeMarkupResult(False, ())

    found: list[str] = []
    seen: set[str] = set()
    for tag, reason in _UNSAFE_TAGS:
        qname = qn(tag)
        # ``iter(tag)`` short-circuits as soon as a match is found if we
        # break, but we want to know whether *any* descendant matches; the
        # presence check is enough.
        if next(raw.iter(qname), None) is not None:
            if reason not in seen:
                found.append(reason)
                seen.add(reason)
    if not found:
        return UnsafeMarkupResult(False, ())
    return UnsafeMarkupResult(True, tuple(found))


def _is_table_cell_mapping(mapping) -> bool:
    return getattr(mapping, "element_type", None) == "table_cell"


def _table_cell_auto_edit_enabled() -> bool:
    """Whether table-cell auto-edits are allowed. Always True."""
    return True


def _edit_transactional_enabled() -> bool:
    """Whether edit application enforces all-or-none output writes.

    Always True: if any auto-edit produced a ``failed`` outcome, the
    serialized output is suppressed so the user does not silently receive
    a partially mutated file.

    Skipped outcomes — including unsafe-markup refusals — are deliberate
    refusals, not failures, and do not abort the transactional write. The
    visible signal stays in ``EditOutcome.refused_unsafe_markup`` and the
    corresponding ``EditReport.warnings`` entry.
    """
    return True


def _refuse_unsafe_outcome(
    action: "EditAction",
    *,
    element,
    original_text: str,
) -> "EditOutcome | None":
    """Return an unsafe-markup refusal outcome, or ``None`` if the element is safe.

    Encapsulates the per-mutation safety gate so the four mutation sites in
    :func:`apply_edits_to_spec` (paragraph in-place, table cell, ADD anchor,
    whole-paragraph DELETE) share one decision point.
    """
    result = detect_unsafe_markup(element)
    if not result.unsafe:
        return None
    return EditOutcome(
        action=action,
        status="skipped",
        detail=result.detail,
        original_text=original_text,
        new_text=None,
        refused_unsafe_markup=True,
    )


@dataclass(frozen=True)
class PreconditionResult:
    """Outcome of revalidating a recorded edit precondition (Chunk F).

    Carries the offsets that should be used for the actual mutation. If the
    recorded offsets still match the live text, ``match_start`` /
    ``match_end`` are returned unchanged. If the live text shifted but the
    expected text is uniquely present at a different offset, the corrected
    offsets are returned so the caller mutates the right span instead of the
    stale one. ``ok`` is False for missing or duplicated text — the caller
    must skip the edit rather than guess which occurrence to mutate.
    """

    ok: bool
    match_start: int
    match_end: int
    detail: str


def _precondition_holds_for_paragraph(
    paragraph: Paragraph,
    match_start: int,
    match_end: int,
    expected_text: str,
) -> PreconditionResult:
    """Verify the live paragraph still contains the expected slice (audit 8.3, Chunk F).

    Strategy:

    1. If the recorded offsets still slice out the expected text, accept
       and return the recorded offsets unchanged.
    2. Otherwise, if the expected text is uniquely present somewhere else
       in the live paragraph, return ``ok=True`` with the corrected
       offsets so the caller can mutate the actual occurrence rather than
       a stale slice that may now contain unrelated characters.
    3. If the expected text is missing or appears more than once, return
       ``ok=False``. The caller must skip the edit — guessing between
       multiple occurrences risks replacing the wrong span, and silently
       trusting the stale offsets risks corrupting the paragraph.
    """
    current = paragraph.text
    if 0 <= match_start <= match_end <= len(current):
        if current[match_start:match_end] == expected_text:
            return PreconditionResult(
                True,
                match_start,
                match_end,
                "Precondition matched at recorded offsets.",
            )

    if not expected_text:
        return PreconditionResult(
            False,
            match_start,
            match_end,
            "Precondition revalidation failed: paragraph text no longer matches "
            "the recorded edit target and no expected text is recorded.",
        )

    occurrences = current.count(expected_text)
    if occurrences == 1:
        new_start = current.find(expected_text)
        new_end = new_start + len(expected_text)
        return PreconditionResult(
            True,
            new_start,
            new_end,
            "Precondition matched via unique substring presence; offsets corrected "
            f"from [{match_start}, {match_end}) to [{new_start}, {new_end}).",
        )

    if occurrences == 0:
        return PreconditionResult(
            False,
            match_start,
            match_end,
            "Precondition revalidation failed: expected text is no longer present in the paragraph.",
        )

    return PreconditionResult(
        False,
        match_start,
        match_end,
        f"Precondition revalidation failed: expected text appears {occurrences} times "
        "in the paragraph; manual review required to avoid wrong-span replacement.",
    )


def _precondition_holds_for_anchor(anchor_paragraph: Paragraph, expected_text: str) -> tuple[bool, str]:
    if not expected_text:
        return True, "No anchor precondition recorded."
    if expected_text in anchor_paragraph.text or anchor_paragraph.text in expected_text:
        return True, "Anchor precondition holds."
    return (
        False,
        "Precondition revalidation failed: anchor paragraph text no longer matches the recorded anchor.",
    )


def _action_confidence(action: EditAction) -> float:
    """Locator match confidence in [0.0, 1.0].

    Populated by :mod:`edit_locator` when an edit's ``existingText`` was
    resolved against the source document. 1.0 = exact byte-for-byte match
    of the model's ``existingText`` against a single paragraph;
    intermediate values reflect normalized / fuzzy / section-anchored
    matches (and those non-exact methods are already gated to manual
    review elsewhere). Used here only to break ties between two edits
    whose spans are identical — *not* to legitimize picking one edit
    over another when their intents partially overlap.
    """
    return action.location.match_confidence


def _severity_rank(action: EditAction) -> int:
    """Lower is more important; missing/unknown severities sort last (99)."""
    severity = (action.locator_result.finding.severity or "").upper()
    return {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "GRIPES": 3}.get(severity, 99)


def _resolve_overlap_winner(action_a: EditAction, action_b: EditAction) -> EditAction | None:
    """Pick a clear winner for two overlapping edits, or return None if ambiguous.

    Resolution rules:

    * Strict containment (one edit's span is fully inside the other and they
      are not identical): the broader edit wins. The narrower edit's intent
      is included in the broader edit's replacement span, so applying the
      broader edit is not a misapplication.
    * Identical span: fall back to severity → confidence → first-arg tie-
      breakers (in that order) so duplicate findings collapse to a single
      applied edit rather than both being thrown away.
      - Severity: CRITICAL < HIGH < MEDIUM < GRIPES (lower rank wins). The
        more important finding's edit is preferred when two identical-span
        edits disagree on what to write.
      - Confidence: only used when severities tie. Higher locator-match
        confidence wins because higher confidence is more likely to mean
        the model's ``existingText`` lines up with the source exactly.
      - Final fallback: return ``action_a``. Deterministic ordering so the
        same input always produces the same applied edit; the actual
        choice is harmless when both edits have identical span, severity,
        and confidence.
    * Partial overlap (neither strictly contains the other): return ``None``.
      The caller skips both edits with a manual-review detail; auto-applying
      either one would discard the other's distinct intent. Severity and
      confidence are intentionally *not* consulted here — they were only
      ever a tie-break for identical-span duplicates, and treating them as
      a precedence rule across partial overlaps silently discarded the
      losing edit's content.
    """
    a_range = (action_a.location.match_start, action_a.location.match_end)
    b_range = (action_b.location.match_start, action_b.location.match_end)

    a_contains_b = a_range[0] <= b_range[0] and a_range[1] >= b_range[1]
    b_contains_a = b_range[0] <= a_range[0] and b_range[1] >= a_range[1]

    if a_contains_b and b_contains_a:
        # Identical spans — duplicate or near-duplicate findings. Keep the
        # heuristic tie-breaker so dedup-survivors still collapse to one
        # applied edit rather than both being lost to "ambiguous".
        rank_a = _severity_rank(action_a)
        rank_b = _severity_rank(action_b)
        if rank_a != rank_b:
            return action_a if rank_a < rank_b else action_b

        conf_a = _action_confidence(action_a)
        conf_b = _action_confidence(action_b)
        if conf_a != conf_b:
            return action_a if conf_a > conf_b else action_b

        return action_a

    if a_contains_b and not b_contains_a:
        return action_a
    if b_contains_a and not a_contains_b:
        return action_b

    # Partial overlap with no containment is ambiguous. Caller must skip both.
    return None


def _action_group_key(action: EditAction) -> tuple[int, str, int | None]:
    mapping = action.location.mapping
    return mapping.body_index, mapping.element_type, mapping.row_index


def _detect_and_resolve_conflicts(actions: list[EditAction]) -> tuple[list[EditAction], list[EditOutcome]]:
    grouped: dict[tuple[int, str, int | None], list[EditAction]] = defaultdict(list)
    for action in actions:
        grouped[_action_group_key(action)].append(action)

    to_apply: list[EditAction] = []
    skipped: list[EditOutcome] = []

    for group in grouped.values():
        if len(group) == 1:
            to_apply.extend(group)
            continue

        whole_deletes = [action for action in group if _is_whole_paragraph_delete(action)]
        if whole_deletes:
            winner = max(whole_deletes, key=_action_confidence)
            to_apply.append(winner)
            for action in group:
                if action is winner:
                    continue
                skipped.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail="Skipped due to whole-paragraph DELETE conflict in same target.",
                        original_text=action.location.matched_text,
                        new_text=None,
                    )
                )
            continue

        # Process the group in descending-start order so the higher-offset
        # edit is checked first. Track tainted ranges from ambiguous-overlap
        # resolutions so a third edit overlapping with either side of an
        # already-discarded ambiguous pair is also skipped (rather than
        # slipping through because the original conflicting actions were
        # removed from ``accepted``).
        sorted_group = sorted(group, key=lambda item: item.location.match_start, reverse=True)
        accepted: list[EditAction] = []
        ambiguous_ranges: list[tuple[int, int]] = []

        def _spans_overlap(start_a: int, end_a: int, start_b: int, end_b: int) -> bool:
            return start_a < end_b and end_a > start_b

        for action in sorted_group:
            a_start = action.location.match_start
            a_end = action.location.match_end

            # If this action overlaps any previously-tainted region, it is
            # also ambiguous: the original conflicting pair is gone from
            # ``accepted`` but their span is still untrustworthy.
            tainted = next(
                (
                    (start, end)
                    for start, end in ambiguous_ranges
                    if _spans_overlap(a_start, a_end, start, end)
                ),
                None,
            )
            if tainted is not None:
                skipped.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail=(
                            "Skipped: overlaps a same-paragraph region already "
                            "flagged for manual review due to ambiguous "
                            f"conflicting edits [{tainted[0]}, {tainted[1]})."
                        ),
                        original_text=action.location.matched_text,
                        new_text=None,
                    )
                )
                continue

            overlap = None
            for existing in accepted:
                if _spans_overlap(
                    a_start,
                    a_end,
                    existing.location.match_start,
                    existing.location.match_end,
                ):
                    overlap = existing
                    break

            if overlap is None:
                accepted.append(action)
                continue

            winner = _resolve_overlap_winner(action, overlap)
            if winner is None:
                # Ambiguous partial overlap — skip both edits and
                # taint the union range so any further action overlapping
                # this region is also routed to manual review.
                accepted.remove(overlap)
                detail = (
                    "Skipped due to ambiguous overlapping edits in the same "
                    "paragraph; manual review required to avoid silently "
                    "picking one intent over the other."
                )
                skipped.append(
                    EditOutcome(
                        action=overlap,
                        status="skipped",
                        detail=detail,
                        original_text=overlap.location.matched_text,
                        new_text=None,
                    )
                )
                skipped.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail=detail,
                        original_text=action.location.matched_text,
                        new_text=None,
                    )
                )
                ambiguous_ranges.append(
                    (
                        min(a_start, overlap.location.match_start),
                        max(a_end, overlap.location.match_end),
                    )
                )
                continue

            if winner is action:
                accepted.remove(overlap)
                skipped.append(
                    EditOutcome(
                        action=overlap,
                        status="skipped",
                        detail="Skipped due to overlapping conflict with broader/higher-priority edit.",
                        original_text=overlap.location.matched_text,
                        new_text=None,
                    )
                )
                accepted.append(action)
            else:
                skipped.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail="Skipped due to overlapping conflict with broader/higher-priority edit.",
                        original_text=action.location.matched_text,
                        new_text=None,
                    )
                )

        to_apply.extend(sorted(accepted, key=lambda item: item.location.match_start, reverse=True))

    ordered = sorted(
        to_apply,
        key=lambda item: (
            item.location.mapping.body_index,
            -item.location.match_start,
        ),
    )
    return ordered, skipped


def _resolve_cell_and_offsets(
    action: EditAction, row
) -> tuple[Paragraph | None, int | None, int | None, str, str]:
    """Locate the target paragraph and span for a table-cell edit.

    Returns ``(paragraph, start, end, detail, status)`` where ``status`` is
    one of ``"resolved"``, ``"failed"`` (data shape problem the caller
    should treat as failed), or ``"skipped"`` (a deliberate Chunk F safety
    refusal — duplicated or missing target text). The caller uses
    ``status`` to record the right outcome.
    """
    mapping = action.location.mapping
    row_text_parts: list[tuple[object, str]] = []
    for cell in row.cells:
        text = cell.text.strip()
        if text:
            row_text_parts.append((cell, text))

    if not row_text_parts:
        return None, None, None, "Row had no non-empty cells for mapped text.", "failed"

    start = action.location.match_start
    end = action.location.match_end
    cursor = 0
    matched_cell = None

    for idx, (cell, text) in enumerate(row_text_parts):
        seg_start = cursor
        seg_end = seg_start + len(text)
        if start >= seg_start and end <= seg_end:
            matched_cell = cell
            break
        cursor = seg_end + 3
        if idx == len(row_text_parts) - 1:
            cursor = seg_end

    if matched_cell is None:
        return (
            None,
            None,
            None,
            "Matched range crossed table-cell boundaries; skipping for safety.",
            "failed",
        )

    # Chunk F: enumerate every occurrence of the expected text across the
    # cell's paragraphs and require uniqueness. The previous implementation
    # picked the first ``paragraph.text.find()`` hit, which silently guessed
    # when the text appeared multiple times in the cell.
    expected = action.location.matched_text
    if not expected:
        return None, None, None, "No expected text recorded for table-cell edit.", "failed"

    candidates: list[tuple[Paragraph, int]] = []
    for paragraph in matched_cell.paragraphs:
        text = paragraph.text
        scan = 0
        while True:
            hit = text.find(expected, scan)
            if hit == -1:
                break
            candidates.append((paragraph, hit))
            scan = hit + 1

    if not candidates:
        return (
            None,
            None,
            None,
            "Expected text is no longer present in any paragraph of the target cell.",
            "skipped",
        )
    if len(candidates) > 1:
        return (
            None,
            None,
            None,
            (
                f"Expected text appears {len(candidates)} times across this cell's "
                "paragraphs; manual review required to avoid wrong-span replacement."
            ),
            "skipped",
        )

    paragraph, hit = candidates[0]
    return paragraph, hit, hit + len(expected), "Resolved table-cell paragraph target.", "resolved"


def _apply_add_action(
    action: EditAction,
    doc: Document,
    *,
    original_body_children: list | None = None,
) -> EditOutcome:
    mapping = action.location.mapping
    if mapping.element_type != "paragraph":
        return EditOutcome(
            action=action,
            status="failed",
            detail="ADD actions are only supported for paragraph mappings.",
            original_text=action.location.matched_text,
            new_text=None,
        )

    # Use the pre-mutation snapshot so DELETE actions earlier in the same
    # apply pass do not shift the body_index used by ADD (audit Issue 6).
    body_children = (
        list(original_body_children)
        if original_body_children is not None
        else list(doc.element.body)
    )
    if mapping.body_index < 0 or mapping.body_index >= len(body_children):
        return EditOutcome(
            action=action,
            status="failed",
            detail="Body index is out of range in current document.",
            original_text=action.location.matched_text,
            new_text=None,
        )

    anchor_element = body_children[mapping.body_index]
    if not anchor_element.tag.endswith("}p"):
        return EditOutcome(
            action=action,
            status="failed",
            detail="ADD mapping expected paragraph but body element was not paragraph.",
            original_text=action.location.matched_text,
            new_text=None,
        )

    # If the anchor was already removed by a DELETE earlier in this run,
    # there is no longer a parent to insert beside. Fail safely instead of
    # silently writing into an orphaned XML node.
    if anchor_element.getparent() is None:
        return EditOutcome(
            action=action,
            status="failed",
            detail="ADD anchor paragraph was deleted earlier in this edit pass; skip.",
            original_text=action.location.matched_text,
            new_text=None,
        )

    anchor_paragraph = Paragraph(anchor_element, doc)

    # Chunk 9: refuse anchors that carry unsafe Word markup. ADD inserts a
    # sibling paragraph adjacent to the anchor, which is structurally
    # simpler than mutating the anchor itself, but inserting beside a
    # paragraph whose XML carries field characters or bookmark ranges can
    # still break those structures if the surrounding context relies on
    # contiguous run order. Conservative refusal keeps the document safe.
    unsafe = detect_unsafe_markup(anchor_element)
    if unsafe.unsafe:
        return EditOutcome(
            action=action,
            status="skipped",
            detail=unsafe.detail,
            original_text=anchor_paragraph.text,
            new_text=None,
            refused_unsafe_markup=True,
        )

    # Phase 4 (audit Section 8.3): revalidate anchor before mutating. If a
    # prior edit changed the anchor paragraph text in a way that no longer
    # matches the recorded anchor, do not insert beside it.
    ok_anchor, anchor_detail = _precondition_holds_for_anchor(
        anchor_paragraph,
        action.location.matched_text,
    )
    if not ok_anchor:
        return EditOutcome(
            action=action,
            status="skipped",
            detail=anchor_detail,
            original_text=anchor_paragraph.text,
            new_text=None,
        )

    replacement = (action.replacement_text or "").strip()
    if not replacement:
        return EditOutcome(
            action=action,
            status="skipped",
            detail="ADD action had empty replacement text; nothing to insert.",
            original_text=anchor_paragraph.text,
            new_text=None,
        )

    anchor_text = action.location.matched_text or anchor_paragraph.text
    replacement_norm = _normalize_text_for_add(replacement)
    anchor_norm = _normalize_text_for_add(anchor_text)

    # If the model provided an explicit insertPosition, trust it instead of
    # guessing from text overlap (audit Issue 5). This makes ADD edits
    # deterministic when the prompt produced the structured anchor model.
    explicit_position = (
        getattr(action.locator_result.finding, "insertPosition", None) or ""
    ).strip().lower()

    if explicit_position in {"before", "after"}:
        position = explicit_position
        new_content = replacement
    else:
        position = "replace"
        new_content = replacement
        if anchor_norm and replacement_norm.startswith(anchor_norm):
            position = "after"
            new_content = replacement[len(anchor_text):].strip()
        elif anchor_norm and replacement_norm.endswith(anchor_norm):
            position = "before"
            new_content = replacement[: max(0, len(replacement) - len(anchor_text))].strip()
        elif anchor_norm and anchor_norm not in replacement_norm:
            position = "after"
            new_content = replacement

    if position == "replace":
        ok, detail = _replace_in_paragraph(
            anchor_paragraph,
            action.location.match_start,
            action.location.match_end,
            replacement,
        )
        return EditOutcome(
            action=action,
            status="applied" if ok else "failed",
            detail=f"ADD fallback to replace: {detail}",
            original_text=anchor_text,
            new_text=anchor_paragraph.text if ok else None,
        )

    paragraphs = _split_insert_paragraphs(new_content)
    if not paragraphs:
        return EditOutcome(
            action=action,
            status="skipped",
            detail="ADD replacement contained no additional content beyond anchor.",
            original_text=anchor_paragraph.text,
            new_text=None,
        )

    style_id = _reference_style_for_text(mapping.body_index, body_children, paragraphs[0])
    inserted_count = (
        _insert_paragraphs_after(anchor_element, paragraphs, style_id)
        if position == "after"
        else _insert_paragraphs_before(anchor_element, paragraphs, style_id)
    )
    return EditOutcome(
        action=action,
        status="applied",
        detail=f"Inserted {inserted_count} paragraph(s) {position} anchor paragraph.",
        original_text=anchor_paragraph.text,
        new_text="\n\n".join(paragraphs),
    )


def apply_edits_to_spec(source_path: Path, output_path: Path, edit_actions: list[EditAction]) -> EditReport:
    source_path = Path(source_path)
    output_path = Path(output_path)

    if source_path.resolve() == output_path.resolve():
        raise ValueError("output_path must differ from source_path; refusing to overwrite source document.")

    doc = Document(source_path)
    actions_to_apply, pre_skipped = _detect_and_resolve_conflicts(edit_actions)
    outcomes: list[EditOutcome] = list(pre_skipped)
    warnings: list[str] = []

    for skipped in pre_skipped:
        warnings.append(skipped.detail)

    # Phase 4 (audit Section 8.4): apply edits in a deterministic safety
    # order — in-place replacements first, then ADDs, then whole-paragraph
    # DELETEs in descending body_index. This keeps anchor elements live for
    # ADDs and avoids any ordering surprise where a DELETE shifts the
    # document structure before later edits run.
    replacement_actions: list[EditAction] = []
    whole_delete_actions: list[EditAction] = []
    for action in actions_to_apply:
        if action.action_type == "ADD":
            continue
        if _is_whole_paragraph_delete(action):
            whole_delete_actions.append(action)
        else:
            replacement_actions.append(action)

    add_actions = sorted(
        (action for action in actions_to_apply if action.action_type == "ADD"),
        key=lambda item: item.location.mapping.body_index,
        reverse=True,
    )
    whole_delete_actions.sort(
        key=lambda item: item.location.mapping.body_index,
        reverse=True,
    )

    body_children = list(doc.element.body)
    for action in replacement_actions:
        mapping = action.location.mapping
        original_text = action.location.matched_text

        if mapping.body_index < 0 or mapping.body_index >= len(body_children):
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="failed",
                    detail="Body index is out of range in current document.",
                    original_text=original_text,
                    new_text=None,
                )
            )
            continue

        element = body_children[mapping.body_index]

        if mapping.element_type == "paragraph":
            if not element.tag.endswith("}p"):
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="failed",
                        detail="Mapping expected paragraph but body element was not paragraph.",
                        original_text=original_text,
                        new_text=None,
                    )
                )
                continue
            paragraph = Paragraph(element, doc)
            paragraph_before = paragraph.text

            # Chunk 9: refuse paragraphs that carry unsafe Word markup
            # (hyperlinks, field codes, drawings, comments, tracked changes,
            # bookmarks, content controls, footnote/endnote refs). Run-level
            # surgery on those structures can silently break the underlying
            # XML — better to skip the auto-edit and route the finding to
            # manual review than to risk a corrupted spec document.
            refusal = _refuse_unsafe_outcome(
                action,
                element=element,
                original_text=paragraph_before,
            )
            if refusal is not None:
                outcomes.append(refusal)
                warnings.append(refusal.detail)
                continue

            # Phase 4 (audit Section 8.3) + Chunk F: revalidate immediately
            # before mutating. If a previous edit in this pass changed the
            # paragraph such that the recorded slice no longer matches, the
            # precondition returns corrected offsets when the expected text
            # is uniquely present elsewhere; if it is missing or duplicated
            # we skip the edit instead of replacing a stale span.
            precondition = _precondition_holds_for_paragraph(
                paragraph,
                action.location.match_start,
                action.location.match_end,
                action.location.matched_text,
            )
            if not precondition.ok:
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail=precondition.detail,
                        original_text=paragraph_before,
                        new_text=None,
                    )
                )
                continue

            replacement = action.replacement_text or ""
            ok, detail = _replace_in_paragraph(
                paragraph,
                precondition.match_start,
                precondition.match_end,
                replacement,
            )
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="applied" if ok else "failed",
                    detail=detail,
                    original_text=paragraph_before,
                    new_text=paragraph.text if ok else None,
                )
            )
            continue

        if mapping.element_type == "table_cell":
            if not element.tag.endswith("}tbl"):
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="failed",
                        detail="Mapping expected table but body element was not table.",
                        original_text=original_text,
                        new_text=None,
                    )
                )
                continue

            # Chunk 9 operator switch: refuse every table-cell auto-edit when
            # ``SPEC_CRITIC_TABLE_CELL_AUTO_EDIT=0``. The finding still flows
            # through the report path; only the silent in-place mutation is
            # suppressed.
            if not _table_cell_auto_edit_enabled():
                detail = (
                    "Auto-edit refused: table-cell auto-edit is disabled "
                    "(SPEC_CRITIC_TABLE_CELL_AUTO_EDIT=0); routed to manual review."
                )
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail=detail,
                        original_text=original_text,
                        new_text=None,
                        refused_unsafe_markup=True,
                    )
                )
                warnings.append(detail)
                continue

            table = DocxTable(element, doc)
            row_index = mapping.row_index
            if row_index is None or row_index < 0 or row_index >= len(table.rows):
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="failed",
                        detail="Row index for table-cell mapping was invalid.",
                        original_text=original_text,
                        new_text=None,
                    )
                )
                continue
            target_paragraph, start, end, detail, resolve_status = _resolve_cell_and_offsets(
                action, table.rows[row_index]
            )
            if target_paragraph is None or start is None or end is None:
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="skipped" if resolve_status == "skipped" else "failed",
                        detail=detail,
                        original_text=original_text,
                        new_text=None,
                    )
                )
                continue

            # Chunk 9: refuse table cells whose target paragraph (or any
            # ancestor cell content) carries unsafe markup. Cell-scoped
            # check is conservative on purpose: hyperlinks/fields/etc.
            # anywhere in the same cell tend to share runs/relationships
            # with the target paragraph.
            cell_element = target_paragraph._element.getparent()
            unsafe_target = _refuse_unsafe_outcome(
                action,
                element=cell_element if cell_element is not None else target_paragraph._element,
                original_text=target_paragraph.text,
            )
            if unsafe_target is not None:
                outcomes.append(unsafe_target)
                warnings.append(unsafe_target.detail)
                continue

            paragraph_before = target_paragraph.text

            # Chunk F: same offset-safety contract as the paragraph path.
            # The table-cell resolver finds the expected text by substring
            # search, but a prior edit in this pass could have shifted or
            # duplicated it; revalidate and use the precondition's
            # (possibly corrected) offsets for the actual replacement.
            precondition = _precondition_holds_for_paragraph(
                target_paragraph,
                start,
                end,
                action.location.matched_text,
            )
            if not precondition.ok:
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="skipped",
                        detail=precondition.detail,
                        original_text=paragraph_before,
                        new_text=None,
                    )
                )
                continue

            cell_start = precondition.match_start
            cell_end = precondition.match_end

            if action.action_type == "DELETE" and cell_start == 0 and cell_end == len(paragraph_before):
                ok, replace_detail = _replace_in_paragraph(target_paragraph, cell_start, cell_end, "")
                outcomes.append(
                    EditOutcome(
                        action=action,
                        status="applied" if ok else "failed",
                        detail=replace_detail,
                        original_text=paragraph_before,
                        new_text=target_paragraph.text if ok else None,
                    )
                )
                continue

            replacement = action.replacement_text or ""
            ok, replace_detail = _replace_in_paragraph(target_paragraph, cell_start, cell_end, replacement)
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="applied" if ok else "failed",
                    detail=replace_detail,
                    original_text=paragraph_before,
                    new_text=target_paragraph.text if ok else None,
                )
            )
            continue

        outcomes.append(
            EditOutcome(
                action=action,
                status="skipped",
                detail=f"Unsupported mapping element type: {mapping.element_type}",
                original_text=original_text,
                new_text=None,
            )
        )

    for action in add_actions:
        outcomes.append(
            _apply_add_action(action, doc, original_body_children=body_children)
        )

    # Whole-paragraph DELETEs run last (audit Section 8.4). Descending body
    # order keeps the snapshot indices stable and avoids any chance that a
    # remove() upstream of an ADD anchor could orphan that anchor before the
    # ADD applied.
    for action in whole_delete_actions:
        mapping = action.location.mapping
        if mapping.body_index < 0 or mapping.body_index >= len(body_children):
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="failed",
                    detail="Body index is out of range in current document.",
                    original_text=action.location.matched_text,
                    new_text=None,
                )
            )
            continue
        element = body_children[mapping.body_index]
        if not element.tag.endswith("}p"):
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="failed",
                    detail="Mapping expected paragraph but body element was not paragraph.",
                    original_text=action.location.matched_text,
                    new_text=None,
                )
            )
            continue
        if element.getparent() is None:
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="skipped",
                    detail="Paragraph already removed by an earlier edit; skipping DELETE.",
                    original_text=action.location.matched_text,
                    new_text=None,
                )
            )
            continue
        paragraph = Paragraph(element, doc)
        paragraph_before = paragraph.text

        # Chunk 9: refuse to delete a paragraph that owns unsafe markup. A
        # whole-paragraph delete that strips a hyperlink or field can leave
        # orphan relationships/bookmarks in document.xml.rels and other
        # ancillary parts; safer to route to manual review.
        refusal = _refuse_unsafe_outcome(
            action,
            element=element,
            original_text=paragraph_before,
        )
        if refusal is not None:
            outcomes.append(refusal)
            warnings.append(refusal.detail)
            continue

        precondition = _precondition_holds_for_paragraph(
            paragraph,
            action.location.match_start,
            action.location.match_end,
            action.location.matched_text,
        )
        if not precondition.ok:
            outcomes.append(
                EditOutcome(
                    action=action,
                    status="skipped",
                    detail=precondition.detail,
                    original_text=paragraph_before,
                    new_text=None,
                )
            )
            continue

        ok = _delete_paragraph(paragraph)
        outcomes.append(
            EditOutcome(
                action=action,
                status="applied" if ok else "failed",
                detail="Deleted full paragraph." if ok else "Failed to delete paragraph element.",
                original_text=paragraph_before,
                new_text="" if ok else None,
            )
        )

    # Chunk 9 — transactional all-or-none output. Serialize to a buffer
    # first; if *any* individual edit ended in ``failed``, suppress the
    # output write entirely so the user does not silently receive a
    # partially mutated file. Skipped outcomes (precondition revalidation,
    # unsafe-markup refusal, ambiguous overlap) are deliberate refusals,
    # not failures, and do NOT abort the write. Operators can opt out via
    # ``SPEC_CRITIC_EDIT_TRANSACTIONAL=0`` if they need the legacy
    # best-effort behavior.
    buf = BytesIO()
    try:
        doc.save(buf)
    except Exception as exc:
        failed_outcomes = [
            EditOutcome(
                action=outcome.action,
                status="failed",
                detail=f"Document serialization failed after edits: {exc}",
                original_text=outcome.original_text,
                new_text=None,
                refused_unsafe_markup=outcome.refused_unsafe_markup,
            )
            for outcome in outcomes
        ]
        return EditReport(
            source_path=source_path,
            output_path=output_path,
            total_edits_attempted=len(edit_actions),
            edits_applied=0,
            edits_skipped=0,
            edits_failed=len(failed_outcomes),
            outcomes=failed_outcomes,
            warnings=warnings + [f"Serialization check failed: {exc}"],
            aborted_transactional=True,
        )

    # Validate the buffer reopens cleanly so we never write a file Word
    # cannot parse. The serialize-then-reopen step is feasible even for
    # large specs because python-docx parses lxml lazily on read.
    buf.seek(0)
    try:
        Document(buf)
    except Exception as exc:
        warnings.append(
            f"Aborted output write: serialized document failed reopen validation: {exc}"
        )
        return EditReport(
            source_path=source_path,
            output_path=output_path,
            total_edits_attempted=len(edit_actions),
            edits_applied=0,
            edits_skipped=sum(1 for o in outcomes if o.status == "skipped"),
            edits_failed=sum(1 for o in outcomes if o.status == "failed") + sum(
                1 for o in outcomes if o.status == "applied"
            ),
            outcomes=[
                EditOutcome(
                    action=o.action,
                    status="failed",
                    detail=(
                        f"Output suppressed: serialized document failed reopen "
                        f"validation ({exc}). Original outcome was: {o.detail}"
                    )
                    if o.status == "applied"
                    else o.detail,
                    original_text=o.original_text,
                    new_text=None,
                    refused_unsafe_markup=o.refused_unsafe_markup,
                )
                for o in outcomes
            ],
            warnings=warnings,
            aborted_transactional=True,
        )

    failed_count = sum(1 for outcome in outcomes if outcome.status == "failed")

    if failed_count > 0 and _edit_transactional_enabled():
        # Demote every ``applied`` outcome to skipped-due-to-abort so the
        # report makes clear nothing was written. The originally-failed
        # outcomes keep their ``failed`` status.
        rewritten: list[EditOutcome] = []
        for outcome in outcomes:
            if outcome.status == "applied":
                rewritten.append(
                    EditOutcome(
                        action=outcome.action,
                        status="skipped",
                        detail=(
                            "Output suppressed under all-or-none policy: "
                            f"{failed_count} edit(s) in this file failed. "
                            f"Original outcome was: {outcome.detail}"
                        ),
                        original_text=outcome.original_text,
                        new_text=None,
                        refused_unsafe_markup=outcome.refused_unsafe_markup,
                    )
                )
            else:
                rewritten.append(outcome)
        warnings.append(
            "Edit output suppressed: all-or-none policy aborted the write because "
            f"{failed_count} edit(s) failed. Set SPEC_CRITIC_EDIT_TRANSACTIONAL=0 "
            "to fall back to best-effort writes."
        )
        applied = 0
        skipped_count = sum(1 for o in rewritten if o.status == "skipped")
        failed = sum(1 for o in rewritten if o.status == "failed")
        return EditReport(
            source_path=source_path,
            output_path=output_path,
            total_edits_attempted=len(edit_actions),
            edits_applied=applied,
            edits_skipped=skipped_count,
            edits_failed=failed,
            outcomes=rewritten,
            warnings=warnings,
            aborted_transactional=True,
        )

    # Either the all-or-none policy passed (no failures) or the operator
    # opted into best-effort writes. Stream the validated buffer to disk.
    buf.seek(0)
    output_path.write_bytes(buf.read())

    applied = sum(1 for outcome in outcomes if outcome.status == "applied")
    skipped_count = sum(1 for outcome in outcomes if outcome.status == "skipped")
    failed = sum(1 for outcome in outcomes if outcome.status == "failed")

    return EditReport(
        source_path=source_path,
        output_path=output_path,
        total_edits_attempted=len(edit_actions),
        edits_applied=applied,
        edits_skipped=skipped_count,
        edits_failed=failed,
        outcomes=outcomes,
        warnings=warnings,
        aborted_transactional=False,
    )


def build_edit_actions(
    locator_results: list[LocatorResult],
    *,
    allow_caution: bool = True,
) -> list[EditAction]:
    """Convert locator results into mutating edit actions.

    Phase 4: gate auto-application on locator-level safety categories. Only
    AUTO_SAFE results are accepted by default; AUTO_WITH_CAUTION is included
    when allow_caution is True (preserves existing behavior). MANUAL_REVIEW
    and REPORT_ONLY locator results never produce actions.
    """
    actions: list[EditAction] = []
    for finding_index, result in enumerate(locator_results):
        action_type = result.action_type.upper()
        if action_type not in {"EDIT", "DELETE", "ADD"}:
            continue
        if result.status == "not_found" or not result.locations:
            continue

        # Ambiguous locator results have multiple plausible targets; previous
        # behavior silently picked the highest-confidence candidate, which
        # could mutate the wrong paragraph. Per audit Issue 4, ambiguous
        # matches must be manual-review-only.
        if result.status == "ambiguous":
            result.warning = (
                "Ambiguous locator result; multiple targets matched. "
                "Review and apply manually instead of auto-editing."
            )
            continue

        category = (result.safety_category or "").upper()
        if category == SAFETY_MANUAL_REVIEW:
            if result.warning is None:
                result.warning = (
                    "Locator classified as manual review; not eligible for auto-apply."
                )
            continue
        if category == SAFETY_REPORT_ONLY:
            continue
        if category == SAFETY_AUTO_WITH_CAUTION and not allow_caution:
            if result.warning is None:
                result.warning = (
                    "Locator classified AUTO_WITH_CAUTION; auto-apply suppressed by caller."
                )
            continue

        best_location = max(result.locations, key=lambda location: location.match_confidence)
        if best_location.mapping.element_type in {"header", "footer"}:
            if result.warning is None:
                result.warning = "Header/footer findings are review-only and cannot be auto-applied yet."
            continue

        actions.append(
            EditAction(
                locator_result=result,
                location=best_location,
                replacement_text=result.replacement_text,
                action_type=action_type,
                finding_index=finding_index,
            )
        )
    return actions


def apply_edits_to_specs(edit_plan: list[tuple[Path, Path, list[EditAction]]]) -> list[EditReport]:
    return [apply_edits_to_spec(source, output, actions) for source, output, actions in edit_plan]
