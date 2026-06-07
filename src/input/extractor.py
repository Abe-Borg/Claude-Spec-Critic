"""Text extraction for Spec Critic specs (DOCX) and Project Context
attachments (DOCX / PDF / Markdown / plain text)."""

from pathlib import Path
from dataclasses import dataclass, field
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table as DocxTable

SUPPORTED_EXTENSIONS = {".docx"}

# Project-context attachments are reviewed as background reference material
# (not edited by the spec pipeline), so several read-only formats are accepted
# in addition to DOCX: PDFs, and plain Markdown / text. The Markdown / text
# path is what lets a drawing-context digest saved by the standalone analyzer
# (``python -m src.drawings``) be attached as Project Context.
CONTEXT_ATTACHMENT_EXTENSIONS = {".docx", ".pdf", ".md", ".txt"}


@dataclass
class ParagraphMapping:
    body_index: int
    element_type: str
    text: str
    table_index: int | None
    row_index: int | None
    cell_index: int | None
    section_index: int | None = None
    container_type: str | None = None
    # Stable, deterministic element identifier scoped to a single
    # extracted document. The format is human-readable so a finding that
    # cites it can be debugged at a glance: ``p<body_index>`` for body
    # paragraphs, ``t<table>r<row>`` for table-cell rows, ``s<n>h<i>`` /
    # ``s<n>f<i>`` for section header / footer paragraphs, ``tb<box>p<para>``
    # for text-box paragraphs, ``fn<id>p<para>`` / ``en<id>p<para>`` for
    # footnote / endnote paragraphs, and ``meta:hf`` / ``meta:tb`` /
    # ``meta:fn`` / ``meta:en`` for the synthetic delimiter that precedes
    # each supplemental block. The id is stable within a single extraction
    # run; for cross-run stability the document_id of the owning
    # ``ExtractedSpec`` should also be checked. Empty string for legacy
    # mappings constructed by tests that predate element ids.
    element_id: str = ""
    # Section heading text the element belongs to (best-effort). Surfacing
    # this lets a downstream applier disambiguate identical text in
    # different sections without re-scanning the paragraph map.
    section_id: str = ""


@dataclass
class ExtractedSpec:
    filename: str
    content: str
    word_count: int
    source_path: str = ""
    source_format: str = ""
    paragraph_map: list[ParagraphMapping] | None = None
    # Stable, human-debuggable document identifier. Defaults to
    # the filename without extension; when filenames could collide the
    # caller can override. Element ids are only unique inside a single
    # document, so a downstream applier pairs ``(document_id, element_id)``
    # when it disambiguates findings that cite an id.
    document_id: str = ""
    # Warnings emitted during text extraction
    # that the report banner surfaces so reviewers can spot specs where
    # text content may not have been fully captured (drawing-heavy
    # documents, embedded objects, etc.). Empty list by default; populated
    # only when the extractor's heuristics fire. Listed per spec — the
    # run-diagnostics banner counts the number of specs with any warnings,
    # not the total warning count, so a single spec with multiple
    # warnings still counts as one affected file.
    extraction_warnings: list[str] = field(default_factory=list)
    # True when the source document contained pending Word "Track Changes"
    # (revision) markup at extraction time. The extracted ``content`` is the
    # Accept-All-Changes view (insertions kept, deletions removed); this flag
    # lets the report advise reviewers that the spec was read as accept-all so
    # they can confirm that is the version they meant to review. Detected across
    # every surface the extractor reads — the body (incl. tables and text
    # boxes), section headers/footers, and footnote/endnote parts — so the
    # advisory fires even when a redline is confined to a header/footer or note.
    # Defaults False (the common case).
    tracked_changes_detected: bool = False


def _derive_document_id(filename: str) -> str:
    """Return a stable, human-readable document id for ``filename``.

    Ids stay debuggable: the filename without its extension is
    enough as a per-run identifier and reads cleanly in logs. Callers that
    expect cross-run stability across renames should override
    ``ExtractedSpec.document_id`` themselves.
    """
    if not filename:
        return ""
    return Path(filename).stem or filename


def _is_heading_paragraph(text: str) -> bool:
    """Heuristic match for a CSI / DSA spec heading paragraph.

    A cheap, deterministic section attribution lets the
    paragraph map can carry a ``section_id`` without re-walking the doc.
    The heuristic only has to be close enough that downstream prompts and
    reports can group paragraphs by section. False positives are harmless
    — they shift the section boundary by one paragraph.
    """
    stripped = (text or "").strip()
    if not stripped or len(stripped) > 80:
        return False
    # "PART 1 GENERAL" / "SECTION 23 05 23" — explicit headings.
    upper = stripped.upper()
    if upper.startswith("PART ") or upper.startswith("SECTION "):
        return True
    # "1.01 SUMMARY" / "2.3.A …" — numbered CSI subheadings.
    first_token = stripped.split(maxsplit=1)[0]
    if first_token and first_token[0].isdigit() and any(
        ch == "." for ch in first_token
    ):
        return True
    return False


# Threshold above which a spec is flagged as
# drawing-heavy. The pipeline writes the raw count and the proportion into
# the warning message so reviewers can see why the spec was flagged
# (drawings, embedded pictures, OLE objects). 20% is conservative; a
# typical drawing-supplemented spec carries figures inline at ~10% of body
# elements. Above that proportion the assumption that text extraction
# captures the reviewable content stops holding and the warning prompts
# a manual visual check.
_CONTENT_LOSS_WARNING_THRESHOLD = 0.20


def _detect_content_loss_warning(body) -> str | None:
    """Return a content-loss warning string for ``body`` or ``None`` if clean.

    Counts how many direct children of
    ``<w:body>`` (paragraphs and tables) contain at least one descendant
    ``<w:drawing>``, ``<w:pict>``, or ``<w:object>`` element. When that
    proportion exceeds :data:`_CONTENT_LOSS_WARNING_THRESHOLD`, the spec
    is likely drawing-heavy and text-only extraction cannot capture the
    reviewable content. The returned string is appended to
    ``ExtractedSpec.extraction_warnings`` so the report banner surfaces a
    visible count (per-spec, not per-drawing) and the operator knows to
    verify the spec visually.

    The ``<w:sectPr>`` body child (section properties) is metadata and
    not counted as a content element. Returns ``None`` when there are no
    body children, no embedded objects, or the proportion is below the
    threshold so the caller can keep ``extraction_warnings`` empty for
    the common case.
    """
    drawing_qn = qn("w:drawing")
    pict_qn = qn("w:pict")
    object_qn = qn("w:object")
    sect_pr_qn = qn("w:sectPr")

    total_body_elements = 0
    non_text_elements = 0
    drawings = 0
    pictures = 0
    objects = 0
    for child in body:
        if child.tag == sect_pr_qn:
            continue
        total_body_elements += 1
        child_drawings = len(child.findall(".//" + drawing_qn))
        child_pictures = len(child.findall(".//" + pict_qn))
        child_objects = len(child.findall(".//" + object_qn))
        if child_drawings or child_pictures or child_objects:
            non_text_elements += 1
        drawings += child_drawings
        pictures += child_pictures
        objects += child_objects

    if total_body_elements == 0:
        return None
    if non_text_elements == 0:
        return None
    proportion = non_text_elements / total_body_elements
    if proportion <= _CONTENT_LOSS_WARNING_THRESHOLD:
        return None

    percent = round(proportion * 100)
    return (
        f"Spec contains {percent}% non-text elements "
        f"({drawings} drawings, {pictures} pictures, {objects} OLE objects). "
        "Some content may not have been extracted for review. Verify visually."
    )


# Footnotes and endnotes live in their own package parts (``word/footnotes.xml``
# / ``word/endnotes.xml``), not under ``<w:body>``, so the body walk never
# reaches them. The parts are identified by their OOXML content type. Word
# seeds every document that has the part with structural notes (``separator``
# / ``continuationSeparator``) that carry no authored text; those are skipped
# by their ``w:type`` so an empty ``[Footnote -1]`` label never reaches review.
_FOOTNOTES_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
)
_ENDNOTES_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
)
_STRUCTURAL_NOTE_TYPES = {"separator", "continuationSeparator", "continuationNotice"}


def _collect_textbox_mappings(body) -> list[ParagraphMapping]:
    """Extract text authored inside drawing / VML text boxes.

    Text-box text is stored in ``<w:txbxContent>`` elements nested inside
    ``<w:drawing>`` (modern DrawingML) or ``<w:pict>`` (legacy VML) runs.
    ``Paragraph.text`` does not descend into them, so the plain body walk
    silently drops a requirement authored in a callout / sidebar text box —
    a "miss a real problem" gap (TRUST_AUDIT P0-6). This collects every text
    box in document order and emits one mapping per non-empty text-box
    paragraph. A nested text box is reached by the same descendant search
    and its parent's ``Paragraph.text`` does not include it, so each box is
    captured exactly once (no duplication, no miss).
    """
    txbx_qn = qn("w:txbxContent")
    p_qn = qn("w:p")
    mappings: list[ParagraphMapping] = []
    for box_index, txbx in enumerate(body.findall(".//" + txbx_qn)):
        for para_index, para_el in enumerate(txbx.findall(p_qn)):
            text = _accept_all_paragraph_text(para_el).strip()
            if not text:
                continue
            mappings.append(
                ParagraphMapping(
                    body_index=-1,
                    element_type="textbox",
                    text=f"[Text Box] {text}",
                    table_index=None,
                    row_index=None,
                    cell_index=None,
                    container_type="textbox",
                    element_id=f"tb{box_index}p{para_index}",
                    section_id="",
                )
            )
    return mappings


def _find_part_by_content_type(doc_part, content_type: str):
    """Return the related package part of ``content_type`` (or ``None``).

    Relationship ids are not stable across authoring tools, so footnote /
    endnote parts are located by their OOXML content type. Shared by the
    note-extraction path and the tracked-change detector.
    """
    for rel in doc_part.rels.values():
        if rel.is_external:
            continue
        target = rel.target_part
        if getattr(target, "content_type", None) == content_type:
            return target
    return None


def _collect_note_mappings(
    doc_part,
    *,
    content_type: str,
    note_tag: str,
    label: str,
    id_prefix: str,
) -> list[ParagraphMapping]:
    """Extract footnote / endnote text from the package part of ``content_type``.

    Footnotes and endnotes are not under ``<w:body>``; they hang off the
    document part by relationship. The part is located by content type
    (relationship ids are not stable), parsed defensively, and walked for
    ``<w:footnote>`` / ``<w:endnote>`` elements. Structural notes
    (``separator`` etc.) are skipped by ``w:type``. Returns one mapping per
    non-empty note paragraph, or an empty list when the part is absent (the
    common case) or unreadable — body text is the primary deliverable and a
    malformed notes part must never sink the whole extraction.
    """
    note_part = _find_part_by_content_type(doc_part, content_type)
    if note_part is None:
        return []
    try:
        root = parse_xml(note_part.blob)
    except Exception:
        return []

    element_type = label.lower()
    w_p = qn("w:p")
    w_id = qn("w:id")
    w_type = qn("w:type")
    mappings: list[ParagraphMapping] = []
    for note in root.findall(qn(note_tag)):
        if note.get(w_type) in _STRUCTURAL_NOTE_TYPES:
            continue
        note_id = note.get(w_id) or "?"
        for para_index, para_el in enumerate(note.findall(w_p)):
            text = _accept_all_paragraph_text(para_el).strip()
            if not text:
                continue
            mappings.append(
                ParagraphMapping(
                    body_index=-1,
                    element_type=element_type,
                    text=f"[{label} {note_id}] {text}",
                    table_index=None,
                    row_index=None,
                    cell_index=None,
                    container_type=element_type,
                    element_id=f"{id_prefix}{note_id}p{para_index}",
                    section_id="",
                )
            )
    return mappings


def _append_supplemental_block(
    paragraphs: list[str],
    paragraph_map: list[ParagraphMapping],
    *,
    delimiter: str,
    delimiter_id: str,
    container_type: str,
    entries: list[ParagraphMapping],
) -> None:
    """Append a labeled block (delimiter + its entries) to both the flat text
    list and the paragraph map, in lockstep.

    Supplemental content (text boxes, footnotes, endnotes, headers/footers)
    does not flow inline in ``<w:body>``, so each kind is rendered as its own
    labeled block after the body. Appending to ``paragraphs`` and
    ``paragraph_map`` together preserves the reconstruction invariant (the
    map's text must join back to ``content``). A no-op when ``entries`` is
    empty, so a spec with none of a given kind produces byte-identical output.
    """
    if not entries:
        return
    paragraph_map.append(
        ParagraphMapping(
            body_index=-1,
            element_type="meta",
            text=delimiter,
            table_index=None,
            row_index=None,
            cell_index=None,
            container_type=container_type,
            element_id=delimiter_id,
            section_id="",
        )
    )
    paragraphs.append(delimiter)
    paragraph_map.extend(entries)
    paragraphs.extend(entry.text for entry in entries)


# ---------------------------------------------------------------------------
# Tracked-changes (revision) handling
# ---------------------------------------------------------------------------
#
# When a reviewer leaves Word's "Track Changes" on, edits are stored as
# revision markup rather than applied to the text:
#   * <w:ins> wraps inserted runs (kept when changes are accepted),
#   * <w:del> wraps deleted runs whose text lives in <w:delText> (removed),
#   * <w:moveTo> / <w:moveFrom> wrap the destination / source of a move.
# python-docx's ``Paragraph.text`` selects only direct-child <w:r>/<w:hyperlink>
# runs and reads only <w:t> (never <w:delText>), so it silently drops BOTH
# inserted text (nested under <w:ins>) and deleted text — a hybrid that matches
# neither Word's "Accept All Changes" nor "Reject All Changes" view, and a
# combined edit (delete "2019", insert "2025") collapses to "Comply with  CBC.".
# We instead reconstruct the Accept-All view: keep insertions and move
# destinations, drop deletions and move sources. That is the text that will
# remain once the redline is accepted — i.e. what will actually be issued —
# computed in memory without modifying the source file.
_W_R = qn("w:r")
_W_HYPERLINK = qn("w:hyperlink")
_W_INS = qn("w:ins")
_W_DEL = qn("w:del")
_W_MOVE_FROM = qn("w:moveFrom")
_W_MOVE_TO = qn("w:moveTo")

# Revision wrappers whose content survives "Accept All Changes" — the walk
# descends through these to reach the runs they wrap.
_ACCEPTED_REVISION_WRAPPERS = frozenset({_W_INS, _W_MOVE_TO})

# Any of these anywhere in the document means a reviewer left tracked changes
# pending (used only for the report advisory, not for text extraction).
_REVISION_MARKER_TAGS = (_W_INS, _W_DEL, _W_MOVE_FROM, _W_MOVE_TO)


def _collect_accept_all_text(container, parts: list[str]) -> None:
    """Append the Accept-All run/hyperlink text under ``container`` to ``parts``.

    Mirrors python-docx ``CT_P.text`` (which concatenates the ``.text`` of each
    direct-child ``<w:r>`` / ``<w:hyperlink>``), with one addition: it descends
    through *accepted* revision wrappers (``<w:ins>`` / ``<w:moveTo>``) to reach
    the runs they wrap, and skips ``<w:del>`` / ``<w:moveFrom>`` entirely (those
    disappear on accept). It deliberately does **not** descend into any other
    container (``<w:smartTag>``, ``<w:sdt>``, ``<w:pPr>``, …), exactly as
    python-docx does not — so a document with no revision markup yields output
    byte-identical to ``Paragraph.text``. Run-level text translation (tabs,
    breaks, no-break hyphens; ``<w:delText>`` excluded) is handled by
    ``CT_R.text`` / ``CT_Hyperlink.text``.
    """
    for child in container:
        tag = child.tag
        if not isinstance(tag, str):
            continue  # comments / processing instructions carry no run text
        if tag == _W_R or tag == _W_HYPERLINK:
            parts.append(child.text or "")
        elif tag in _ACCEPTED_REVISION_WRAPPERS:
            _collect_accept_all_text(child, parts)


def _accept_all_paragraph_text(p_el) -> str:
    """Return a paragraph element's text as if all tracked changes were accepted.

    See :func:`_collect_accept_all_text`. For a paragraph with no revision
    markup this equals python-docx ``Paragraph.text``.
    """
    parts: list[str] = []
    _collect_accept_all_text(p_el, parts)
    return "".join(parts)


def _accept_all_cell_text(cell) -> str:
    """Accept-All text for a whole table cell.

    Matches python-docx ``_Cell.text`` (its paragraphs joined by newlines, not
    descending into nested tables) but resolves each paragraph through the
    revision-aware walk.
    """
    return "\n".join(_accept_all_paragraph_text(p._p) for p in cell.paragraphs)


def _element_has_tracked_changes(el) -> bool:
    """True when ``el`` contains any pending tracked-change markup."""
    return any(el.find(".//" + tag) is not None for tag in _REVISION_MARKER_TAGS)


def _document_has_tracked_changes(doc) -> bool:
    """True when the document carries pending tracked-change markup on any
    surface the extractor reads.

    The advisory must mirror extraction coverage so a reviewer is always told
    when *any* extracted text was resolved to the Accept-All view — not just
    body text. Revision markup can live in three places the extractor reads:
    the body (including tables and text boxes, all nested under ``<w:body>``),
    the section headers/footers, and the footnote/endnote package parts. Header/
    footer and note parts hang off the document by relationship, so the body
    scan alone misses a redline confined to them (e.g. a revision note in a
    page header). Parsing of note parts is defensive — an unreadable part never
    sinks detection, matching ``_collect_note_mappings``.
    """
    if _element_has_tracked_changes(doc.element.body):
        return True
    for section in doc.sections:
        for container in (section.header, section.footer):
            if any(_element_has_tracked_changes(p._p) for p in container.paragraphs):
                return True
    for content_type in (_FOOTNOTES_CONTENT_TYPE, _ENDNOTES_CONTENT_TYPE):
        note_part = _find_part_by_content_type(doc.part, content_type)
        if note_part is None:
            continue
        try:
            root = parse_xml(note_part.blob)
        except Exception:
            continue
        if _element_has_tracked_changes(root):
            return True
    return False


def extract_text_from_docx(filepath: Path) -> ExtractedSpec:
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")
    if filepath.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {filepath}")
    try:
        doc = Document(filepath)
    except PackageNotFoundError:
        raise ValueError(f"Invalid or corrupted .docx file: {filepath}")
    except Exception as e:
        raise ValueError(f"Could not read .docx file: {filepath} — {e}")

    paragraphs: list[str] = []
    paragraph_map: list[ParagraphMapping] = []
    table_counter = 0
    # Track the most recently seen heading paragraph so each
    # element below it can carry a ``section_id``. Reset to empty when the
    # extractor crosses a top-level "PART ..." boundary so subsequent
    # subheadings nest under the right ancestor.
    current_section: str = ""

    for body_index, child in enumerate(doc.element.body):
        if child.tag.endswith("}p"):
            text = _accept_all_paragraph_text(child).strip()
            if text:
                paragraphs.append(text)
                if _is_heading_paragraph(text):
                    current_section = text
                paragraph_map.append(
                    ParagraphMapping(
                        body_index=body_index,
                        element_type="paragraph",
                        text=text,
                        table_index=None,
                        row_index=None,
                        cell_index=None,
                        element_id=f"p{body_index}",
                        section_id=current_section,
                    )
                )
        elif child.tag.endswith("}tbl"):
            table = DocxTable(child, doc)
            for row_index, row in enumerate(table.rows):
                row_text = [
                    cell_text
                    for cell in row.cells
                    if (cell_text := _accept_all_cell_text(cell).strip())
                ]
                if row_text:
                    joined_text = " | ".join(row_text)
                    paragraphs.append(joined_text)
                    paragraph_map.append(
                        ParagraphMapping(
                            body_index=body_index,
                            element_type="table_cell",
                            text=joined_text,
                            table_index=table_counter,
                            row_index=row_index,
                            cell_index=None,
                            element_id=f"t{table_counter}r{row_index}",
                            section_id=current_section,
                        )
                    )
            table_counter += 1

    header_footer_entries: list[ParagraphMapping] = []
    for section_index, section in enumerate(doc.sections):
        for container_name, container in (("header", section.header), ("footer", section.footer)):
            for para_index, para in enumerate(container.paragraphs):
                text = _accept_all_paragraph_text(para._p).strip()
                if not text:
                    continue
                prefixed = f"[{container_name.title()}] {text}"
                container_tag = "h" if container_name == "header" else "f"
                header_footer_entries.append(
                    ParagraphMapping(
                        body_index=-1,
                        element_type=container_name,
                        text=prefixed,
                        table_index=None,
                        row_index=None,
                        cell_index=None,
                        section_index=section_index,
                        container_type=container_name,
                        element_id=f"s{section_index}{container_tag}{para_index}",
                        section_id="",
                    )
                )

    # Supplemental content that python-docx's body walk does not surface:
    # text boxes (text nested in <w:txbxContent> inside drawings / VML) and
    # footnotes / endnotes (separate package parts). A requirement authored
    # in any of these would otherwise be invisible to the reviewer — a
    # silent "miss a real problem" gap (TRUST_AUDIT P0-6). Each kind is
    # rendered as its own labeled block after the body (mirroring the
    # header/footer block); the blocks no-op when their source is absent, so
    # a spec with none of them produces byte-identical output to before.
    _append_supplemental_block(
        paragraphs,
        paragraph_map,
        delimiter="===== TEXT BOX CONTENT =====",
        delimiter_id="meta:tb",
        container_type="textbox",
        entries=_collect_textbox_mappings(doc.element.body),
    )
    _append_supplemental_block(
        paragraphs,
        paragraph_map,
        delimiter="===== FOOTNOTE CONTENT =====",
        delimiter_id="meta:fn",
        container_type="footnote",
        entries=_collect_note_mappings(
            doc.part,
            content_type=_FOOTNOTES_CONTENT_TYPE,
            note_tag="w:footnote",
            label="Footnote",
            id_prefix="fn",
        ),
    )
    _append_supplemental_block(
        paragraphs,
        paragraph_map,
        delimiter="===== ENDNOTE CONTENT =====",
        delimiter_id="meta:en",
        container_type="endnote",
        entries=_collect_note_mappings(
            doc.part,
            content_type=_ENDNOTES_CONTENT_TYPE,
            note_tag="w:endnote",
            label="Endnote",
            id_prefix="en",
        ),
    )
    _append_supplemental_block(
        paragraphs,
        paragraph_map,
        delimiter="===== HEADER/FOOTER CONTENT =====",
        delimiter_id="meta:hf",
        container_type="header_footer",
        entries=header_footer_entries,
    )

    content = "\n\n".join(paragraphs)
    reconstructed = "\n\n".join(m.text for m in paragraph_map)
    if reconstructed != content:
        # Controlled error preserves context (audit Issue 10). The raw assert
        # version was stripped under -O and produced an opaque AssertionError.
        raise ValueError(
            f"Paragraph map for '{filepath.name}' does not reconstruct extracted content "
            f"(map_chars={len(reconstructed)}, content_chars={len(content)})."
        )

    # Scan the body for embedded drawings /
    # pictures / objects. When the proportion of non-text elements
    # exceeds the threshold, the spec is likely drawing-heavy and text-
    # only extraction may have missed reviewable content. The warning
    # rides on ``extraction_warnings`` so the run-diagnostics banner can
    # count affected specs and surface the count to the reviewer.
    extraction_warnings: list[str] = []
    content_loss_warning = _detect_content_loss_warning(doc.element.body)
    if content_loss_warning is not None:
        extraction_warnings.append(content_loss_warning)

    # The extracted ``content`` above is the Accept-All-Changes view. Flag
    # whether any pending revision markup was present on any extracted surface
    # (body, headers/footers, footnote/endnote parts) so the report can advise
    # the reviewer the spec was read as accept-all (see ``tracked_changes_detected``).
    tracked_changes_detected = _document_has_tracked_changes(doc)

    return ExtractedSpec(
        filename=filepath.name,
        content=content,
        word_count=len(content.split()),
        source_path=str(filepath),
        source_format="docx",
        paragraph_map=paragraph_map,
        document_id=_derive_document_id(filepath.name),
        extraction_warnings=extraction_warnings,
        tracked_changes_detected=tracked_changes_detected,
    )


def extract_text(filepath: Path) -> ExtractedSpec:
    filepath = Path(filepath)
    ext = filepath.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: '{ext}'. Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return extract_text_from_docx(filepath)


def extract_multiple_specs(
    filepaths: list[Path],
    *,
    max_workers: int | None = None,
) -> list[ExtractedSpec]:
    """Extract a list of specs in parallel.

    Phase 5.2 (audit Section 9.2): bounded thread pool for I/O-bound DOCX
    parsing. Result order is preserved to match ``filepaths`` so downstream
    deterministic ordering (filenames, dedup keys, request maps) does not
    change. ``max_workers=1`` (or a single file) runs sequentially.
    """
    if not filepaths:
        return []
    paths = [Path(fp) for fp in filepaths]
    if len(paths) == 1:
        return [extract_text(paths[0])]
    workers = max_workers if max_workers is not None else min(8, len(paths))
    workers = max(1, workers)
    if workers == 1:
        return [extract_text(p) for p in paths]
    from concurrent.futures import ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=workers) as pool:
        return list(pool.map(extract_text, paths))


def _extract_pdf_text(filepath: Path) -> str:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError as exc:
        raise ValueError(
            "PDF support requires the 'pypdf' package. Install with: pip install pypdf"
        ) from exc
    try:
        reader = PdfReader(str(filepath))
    except PdfReadError as exc:
        raise ValueError(f"Invalid or corrupted PDF: {filepath} — {exc}")
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {filepath} — {exc}")
    if getattr(reader, "is_encrypted", False):
        raise ValueError(f"PDF is encrypted and cannot be read: {filepath}")
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = text.strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_plaintext(filepath: Path) -> str:
    """Read a UTF-8 Markdown / plain-text context attachment verbatim.

    Drawing digests saved by the standalone analyzer (``python -m src.drawings``)
    are Markdown, and reviewers may keep project notes as ``.txt``; both are
    reference material spliced into ``project_context`` as-is. Undecodable bytes
    are replaced rather than raising, so one stray byte never sinks the whole
    attachment (the result is reviewed by a human-readable model, not parsed).
    """
    return filepath.read_text(encoding="utf-8", errors="replace")


def extract_context_text(filepath: Path) -> str:
    """Extract plain text from a Project Context attachment.

    Accepts ``.docx`` / ``.pdf`` (text extracted) and ``.md`` / ``.txt`` (read
    verbatim). Returns a plain string suitable for splicing into the
    project_context prompt block. Unlike ``extract_text``, this does not build a
    paragraph map — the result is reference material, not an editable spec.
    """
    filepath = Path(filepath)
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")
    ext = filepath.suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(filepath).content
    if ext == ".pdf":
        return _extract_pdf_text(filepath)
    if ext in {".md", ".txt"}:
        return _extract_plaintext(filepath)
    raise ValueError(
        f"Unsupported context attachment format: '{ext}'. "
        f"Supported: {', '.join(sorted(CONTEXT_ATTACHMENT_EXTENSIONS))}"
    )
