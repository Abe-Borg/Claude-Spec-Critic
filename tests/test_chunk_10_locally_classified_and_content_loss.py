"""Chunk 10 tests — LOCALLY_CLASSIFIED tightening + DOCX content-loss warning.

Chunk 10 of the Trust Upgrade has two unrelated surfaces:

* **10a — Tighten local-skip keywords.** The router used to drop a
  GRIPES finding through the ``"formatting"`` keyword into ``local_skip``;
  that keyword is too broad (a real CMC formatting requirement could
  match) and is removed outright. ``"leed"`` and ``"internal
  contradiction"`` still route to ``local_skip`` (web search adds no
  signal for either) but are now tagged with ``requires_elevated_confidence``
  on the resulting :class:`VerificationResult` so a downstream applier
  can raise the bar for the residual-risk classes without paying for
  verification that wouldn't add evidence.

* **10b — DOCX content-loss warning.** A drawing-heavy spec (more than
  20% of body children carrying ``<w:drawing>`` / ``<w:pict>`` /
  ``<w:object>`` elements) produces a warning on
  ``ExtractedSpec.extraction_warnings`` so the run-diagnostics banner
  can count affected specs and prompt a manual visual review.

Both surfaces round-trip through resume state and feed into the report
exporter without requiring a cache schema bump (the elevated-confidence
flag is router-derived telemetry; local-skip results never reach the
verification cache because they aren't grounded).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.input.extractor import (
    ExtractedSpec,
    _CONTENT_LOSS_WARNING_THRESHOLD,
    _detect_content_loss_warning,
    extract_text_from_docx,
)
from src.orchestration.pipeline import PipelineResult
from src.output.report_exporter import _summarize_run_diagnostics, export_report
from src.output.report_status import (
    ReportStatus,
    classify_status,
    summarize_edit_actions,
    summarize_statuses,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verification_router import (
    classify_finding_for_verification,
    local_skip_requires_elevated_confidence,
)
from src.verification.verifier import _local_skip_result


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _gripe(*, issue: str, code_reference: str | None = None) -> Finding:
    return Finding(
        severity="GRIPES",
        fileName="Section_22_1000.docx",
        section="2.1",
        issue=issue,
        actionType="EDIT",
        existingText=None,
        replacementText=None,
        codeReference=code_reference,
    )


def _finding_with_local_skip_verification(
    *,
    issue: str,
    edit_confidence: float = 1.0,
) -> Finding:
    """Build a finding whose verification matches what the verifier would
    produce for a local-skip routing decision. Used to test the composite-
    confidence multiplier without spinning up the real verifier."""
    f = _gripe(issue=issue)
    f.existingText = "old text"
    f.replacementText = "new text"
    proposal = EditProposal(
        action_type="EDIT",
        existing_text="old text",
        replacement_text="new text",
        edit_confidence=edit_confidence,
    )
    f.edit_proposal = proposal
    elevated = local_skip_requires_elevated_confidence(f)
    f.verification = _local_skip_result(requires_elevated_confidence=elevated)
    return f


# ===========================================================================
# 10a — Tighten LOCALLY_CLASSIFIED
# ===========================================================================


class TestFormattingKeywordRemoved:
    """``"formatting"`` is too broad — a real CMC formatting requirement
    (e.g. "label valves per ASME A13.1 color formatting") could match
    and silently bypass verification. Chunk 10 removes it outright."""

    def test_formatting_keyword_routes_to_web_required(self):
        f = _gripe(issue="Color formatting label requirements may be wrong")
        assert classify_finding_for_verification(f) == "web_required"

    def test_formatting_keyword_does_not_set_elevated_confidence(self):
        f = _gripe(issue="Color formatting label requirements may be wrong")
        # Since the routing is web_required (not local_skip), the
        # elevated-confidence helper has nothing to flag.
        assert local_skip_requires_elevated_confidence(f) is False


class TestElevatedConfidenceKeywords:
    """``"leed"`` and ``"internal contradiction"`` still route to local_skip
    but get tagged with ``requires_elevated_confidence`` for a downstream
    applier to act on."""

    def test_leed_routes_to_local_skip(self):
        f = _gripe(issue="LEED reference is inappropriate for K-12 project")
        assert classify_finding_for_verification(f) == "local_skip"

    def test_leed_requires_elevated_confidence(self):
        f = _gripe(issue="LEED reference is inappropriate for K-12 project")
        assert local_skip_requires_elevated_confidence(f) is True

    def test_internal_contradiction_routes_to_local_skip(self):
        f = _gripe(issue="Internal contradiction: section 2.1 vs 2.3")
        assert classify_finding_for_verification(f) == "local_skip"

    def test_internal_contradiction_requires_elevated_confidence(self):
        f = _gripe(issue="Internal contradiction: section 2.1 vs 2.3")
        assert local_skip_requires_elevated_confidence(f) is True

    def test_code_reference_overrides_elevated_path_too(self):
        # A finding with a code reference always needs web grounding,
        # regardless of which keyword list it matches.
        f = _gripe(
            issue="LEED reference may conflict with CBC requirements",
            code_reference="CBC §1234",
        )
        assert classify_finding_for_verification(f) == "web_required"
        assert local_skip_requires_elevated_confidence(f) is False

    def test_high_severity_overrides_elevated_path_too(self):
        # Severity gate applies regardless of which keyword list matches.
        f = _gripe(issue="Internal contradiction: section 2.1 vs 2.3")
        f.severity = "HIGH"
        assert classify_finding_for_verification(f) == "web_required"
        assert local_skip_requires_elevated_confidence(f) is False


class TestRegularKeywordWinsOverElevated:
    """A finding that matches BOTH a regular keyword AND an elevated
    keyword takes the regular path (no elevated-confidence flag).
    The regular-list match is the stronger signal."""

    def test_placeholder_and_leed_takes_regular_path(self):
        # Both "placeholder" (regular) and "leed" (elevated) match — the
        # regular path wins and the flag stays False.
        f = _gripe(issue="LEED placeholder needs to be replaced")
        assert classify_finding_for_verification(f) == "local_skip"
        assert local_skip_requires_elevated_confidence(f) is False

    def test_todo_and_internal_contradiction_takes_regular_path(self):
        f = _gripe(issue="TODO: resolve internal contradiction in section 2")
        assert classify_finding_for_verification(f) == "local_skip"
        assert local_skip_requires_elevated_confidence(f) is False


class TestLocalSkipResultCarriesFlag:
    """``_local_skip_result`` propagates the elevated-confidence flag
    onto the resulting VerificationResult so downstream consumers
    (resume state, a future applier) can read it."""

    def test_default_flag_is_false(self):
        result = _local_skip_result()
        assert result.requires_elevated_confidence is False

    def test_explicit_true_propagates(self):
        result = _local_skip_result(requires_elevated_confidence=True)
        assert result.requires_elevated_confidence is True

    def test_default_classify_status_is_locally_classified(self):
        # Sanity: the local-skip result still classifies as
        # LOCALLY_CLASSIFIED regardless of the elevated flag.
        f = _gripe(issue="LEED reference is inappropriate")
        f.verification = _local_skip_result(requires_elevated_confidence=True)
        assert classify_status(f) is ReportStatus.LOCALLY_CLASSIFIED


# ===========================================================================
# 10b — DOCX content-loss warning
# ===========================================================================


def _add_run(paragraph_element, text: str) -> None:
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    paragraph_element.append(run)


def _make_drawing_run() -> OxmlElement:
    run = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    run.append(drawing)
    return run


def _make_pict_run() -> OxmlElement:
    run = OxmlElement("w:r")
    pict = OxmlElement("w:pict")
    run.append(pict)
    return run


def _make_object_run() -> OxmlElement:
    run = OxmlElement("w:r")
    obj = OxmlElement("w:object")
    run.append(obj)
    return run


def _make_drawing_heavy_doc(
    tmp_path: Path,
    *,
    num_text_paragraphs: int = 2,
    num_drawings: int = 5,
    num_pictures: int = 0,
    num_objects: int = 0,
    filename: str = "drawing_heavy.docx",
) -> Path:
    """Build a docx with the requested mix of text and non-text body children.

    Each drawing/picture/object is placed in its own paragraph so the
    body element count is ``num_text_paragraphs + num_drawings +
    num_pictures + num_objects``. The text paragraphs carry plain runs;
    the non-text paragraphs carry a single run with the embedded element
    (drawing / pict / object) and no text — mimicking how Word saves a
    bare inline figure or OLE object.
    """
    doc = Document()
    for i in range(num_text_paragraphs):
        doc.add_paragraph(f"Paragraph {i + 1} with reviewable text content.")
    for _ in range(num_drawings):
        para = doc.add_paragraph()
        para._element.append(_make_drawing_run())
    for _ in range(num_pictures):
        para = doc.add_paragraph()
        para._element.append(_make_pict_run())
    for _ in range(num_objects):
        para = doc.add_paragraph()
        para._element.append(_make_object_run())
    out = tmp_path / filename
    doc.save(out)
    return out


class TestExtractedSpecHasWarningsField:
    """``ExtractedSpec`` has a new ``extraction_warnings`` list field."""

    def test_default_empty_list(self):
        spec = ExtractedSpec(filename="x.docx", content="", word_count=0)
        assert spec.extraction_warnings == []

    def test_each_instance_gets_its_own_list(self):
        # Sanity: ``field(default_factory=list)`` produces a fresh list
        # per instance so mutating one spec's warnings doesn't leak into
        # another's.
        a = ExtractedSpec(filename="a.docx", content="", word_count=0)
        b = ExtractedSpec(filename="b.docx", content="", word_count=0)
        a.extraction_warnings.append("test")
        assert b.extraction_warnings == []


class TestDetectContentLossWarning:
    """``_detect_content_loss_warning`` is a pure helper over the body
    element. Exercised directly so the warning-generation logic can be
    tested independently of disk I/O."""

    def test_threshold_constant_is_documented_value(self):
        # The plan calls for a 20% threshold; pin the constant so a
        # future bump is intentional.
        assert _CONTENT_LOSS_WARNING_THRESHOLD == 0.20

    def test_clean_spec_returns_none(self, tmp_path: Path):
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=5,
            num_drawings=0,
        )
        doc = Document(str(doc_path))
        assert _detect_content_loss_warning(doc.element.body) is None

    def test_below_threshold_returns_none(self, tmp_path: Path):
        # 1 drawing out of 10 body children = 10%; below the 20% floor.
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=9,
            num_drawings=1,
        )
        doc = Document(str(doc_path))
        assert _detect_content_loss_warning(doc.element.body) is None

    def test_at_exact_threshold_returns_none(self, tmp_path: Path):
        # 2 drawings out of 10 = 20%; the gate is strict (>) so the
        # exact threshold value should NOT trigger a warning. Avoiding
        # this off-by-one keeps a borderline spec (exactly 20%) from
        # generating noisy warnings on every run.
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=8,
            num_drawings=2,
        )
        doc = Document(str(doc_path))
        assert _detect_content_loss_warning(doc.element.body) is None

    def test_above_threshold_returns_warning_string(self, tmp_path: Path):
        # 5 drawings out of 7 body children = ~71%; well above 20%.
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=2,
            num_drawings=5,
        )
        doc = Document(str(doc_path))
        warning = _detect_content_loss_warning(doc.element.body)
        assert warning is not None
        # The plan-mandated warning shape: percent, then the breakdown.
        assert "%" in warning
        assert "5 drawings" in warning
        assert "0 pictures" in warning
        assert "0 OLE objects" in warning
        assert "Verify visually" in warning

    def test_warning_message_breaks_down_by_type(self, tmp_path: Path):
        # Mixed embedded content: drawings + pictures + objects.
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=1,
            num_drawings=3,
            num_pictures=2,
            num_objects=1,
        )
        doc = Document(str(doc_path))
        warning = _detect_content_loss_warning(doc.element.body)
        assert warning is not None
        assert "3 drawings" in warning
        assert "2 pictures" in warning
        assert "1 OLE objects" in warning


class TestExtractTextFromDocxAttachesWarnings:
    """The full extraction path populates ``ExtractedSpec.extraction_warnings``
    when the helper fires."""

    def test_clean_spec_has_empty_warnings(self, tmp_path: Path):
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=5,
            num_drawings=0,
        )
        spec = extract_text_from_docx(doc_path)
        assert spec.extraction_warnings == []

    def test_drawing_heavy_spec_has_warning(self, tmp_path: Path):
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=2,
            num_drawings=5,
        )
        spec = extract_text_from_docx(doc_path)
        assert len(spec.extraction_warnings) == 1
        assert "drawings" in spec.extraction_warnings[0]

    def test_word_count_still_computed(self, tmp_path: Path):
        # Warning generation must not interfere with the rest of the
        # extraction pipeline.
        doc_path = _make_drawing_heavy_doc(
            tmp_path,
            num_text_paragraphs=2,
            num_drawings=5,
        )
        spec = extract_text_from_docx(doc_path)
        assert spec.word_count > 0
        assert "Paragraph 1" in spec.content


# ===========================================================================
# 10b — Pipeline integration: PipelineResult carries extracted_specs
# ===========================================================================


class TestPipelineResultCarriesExtractedSpecs:
    """``PipelineResult`` has a new ``extracted_specs`` field so the
    report exporter can read extraction warnings without re-extracting."""

    def test_default_empty_list(self):
        result = PipelineResult(review_result=None)
        assert result.extracted_specs == []

    def test_field_accepts_list_of_specs(self):
        spec = ExtractedSpec(
            filename="x.docx",
            content="",
            word_count=0,
            extraction_warnings=["warning"],
        )
        result = PipelineResult(review_result=None, extracted_specs=[spec])
        assert result.extracted_specs == [spec]
        assert result.extracted_specs[0].extraction_warnings == ["warning"]


class TestBannerSurfacesExtractionWarnings:
    """The Run Diagnostics banner (Chunk 6) counts specs with non-empty
    extraction warnings. The slot was reserved in Chunk 6; Chunk 10
    populates ``ExtractedSpec.extraction_warnings`` from the extractor
    and wires ``extracted_specs`` through the PipelineResult so the
    banner row reflects real data."""

    def _findings_for_banner(self) -> list[Finding]:
        # Need at least one finding so the report exporter has something
        # to render. The finding details don't matter for the banner —
        # only the spec-level warnings do.
        return [
            Finding(
                severity="HIGH",
                fileName="Section_22_1000.docx",
                section="2.1",
                issue="Stale reference",
                actionType="REPORT_ONLY",
                existingText=None,
                replacementText=None,
                codeReference="CBC §1234",
            )
        ]

    def test_warning_count_zero_when_no_warnings(self):
        # A clean run with no extraction warnings shows 0 in the slot.
        clean_spec = ExtractedSpec(
            filename="x.docx",
            content="text",
            word_count=1,
        )
        result = PipelineResult(
            review_result=ReviewResult(findings=[]),
            extracted_specs=[clean_spec],
        )
        summary = _summarize_run_diagnostics(
            findings=[],
            status_counts=summarize_statuses([]),
            edit_action_counts=summarize_edit_actions([]),
            cross_check_result=None,
            pipeline_result=result,
        )
        assert summary["extraction_warning_count"] == 0

    def test_warning_count_reflects_affected_spec_count(self):
        # Three specs: two with warnings, one clean. The banner reports
        # the number of specs affected (2), not the total warning count.
        specs = [
            ExtractedSpec(
                filename="a.docx", content="t", word_count=1,
                extraction_warnings=["w1"],
            ),
            ExtractedSpec(filename="b.docx", content="t", word_count=1),
            ExtractedSpec(
                filename="c.docx", content="t", word_count=1,
                extraction_warnings=["w2", "w3"],  # multiple warnings still counts once
            ),
        ]
        result = PipelineResult(
            review_result=ReviewResult(findings=[]),
            extracted_specs=specs,
        )
        summary = _summarize_run_diagnostics(
            findings=[],
            status_counts=summarize_statuses([]),
            edit_action_counts=summarize_edit_actions([]),
            cross_check_result=None,
            pipeline_result=result,
        )
        assert summary["extraction_warning_count"] == 2

    def test_banner_value_cell_is_red_when_warnings_present(self, tmp_path: Path):
        # End-to-end: a PipelineResult with extraction_warnings produces
        # a red-highlighted "Spec content extraction warnings" row in
        # the exported report's banner. The shading is the same
        # FFE5E5 light-red the verification-failure row uses (Chunk 6
        # established the highlight wiring).
        from docx.oxml.ns import qn as _qn

        f = self._findings_for_banner()[0]
        warning_spec = ExtractedSpec(
            filename=f.fileName,
            content="text",
            word_count=1,
            extraction_warnings=["Spec contains 50% non-text elements"],
        )

        class _Stub:
            review_result = ReviewResult(findings=[f])
            cross_check_result = None
            files_reviewed = [f.fileName]
            leed_alerts: list = []
            placeholder_alerts: list = []
            cycle_label = "2025"
            total_elapsed_seconds = 1.0
            extracted_specs = [warning_spec]

        out = tmp_path / "report.docx"
        export_report(_Stub(), out)
        doc = Document(str(out))

        # Find the row whose first cell label is "Spec content
        # extraction warnings" and read its value-cell shading.
        shading_hex: str | None = None
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                if row.cells[0].text.strip() != "Spec content extraction warnings":
                    continue
                value_cell = row.cells[1]
                tcPr = value_cell._tc.find(_qn("w:tcPr"))
                if tcPr is None:
                    continue
                shd = tcPr.find(_qn("w:shd"))
                if shd is not None:
                    shading_hex = shd.get(_qn("w:fill"))
        assert shading_hex is not None, (
            "Banner row 'Spec content extraction warnings' should be highlighted "
            "when at least one spec has extraction_warnings."
        )
        assert shading_hex.upper() == "FFE5E5"

    def test_banner_value_cell_unshaded_when_no_warnings(self, tmp_path: Path):
        from docx.oxml.ns import qn as _qn

        f = self._findings_for_banner()[0]
        clean_spec = ExtractedSpec(
            filename=f.fileName,
            content="text",
            word_count=1,
        )

        class _Stub:
            review_result = ReviewResult(findings=[f])
            cross_check_result = None
            files_reviewed = [f.fileName]
            leed_alerts: list = []
            placeholder_alerts: list = []
            cycle_label = "2025"
            total_elapsed_seconds = 1.0
            extracted_specs = [clean_spec]

        out = tmp_path / "report.docx"
        export_report(_Stub(), out)
        doc = Document(str(out))

        # No warnings → the row label is present but the value cell
        # should not carry the highlight shading.
        shading_hex: str | None = None
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                if row.cells[0].text.strip() != "Spec content extraction warnings":
                    continue
                value_cell = row.cells[1]
                tcPr = value_cell._tc.find(_qn("w:tcPr"))
                if tcPr is None:
                    continue
                shd = tcPr.find(_qn("w:shd"))
                if shd is not None:
                    shading_hex = shd.get(_qn("w:fill"))
        assert shading_hex is None
