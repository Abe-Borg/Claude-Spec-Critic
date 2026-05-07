import json
from datetime import datetime, timezone
from pathlib import Path

import pytest
from anthropic._utils import maybe_transform
from anthropic.types import TextBlock, ToolUseBlock
from anthropic.types.beta.messages.batch_create_params import BatchCreateParams
from docx import Document

from src.code_cycles import CALIFORNIA_2025
from src.extractor import (
    CONTEXT_ATTACHMENT_EXTENSIONS,
    SUPPORTED_EXTENSIONS,
    ExtractedSpec,
    extract_context_text,
    extract_text,
)
from src.prompts import get_system_prompt, get_single_spec_user_message
from src.pipeline import (
    _deduplicate_findings,
    BatchSubmission,
    CollectedBatchState,
    PipelineResult,
    collect_review_batch_results,
    finalize_batch_result,
    start_batch_review,
    run_cross_check_for_batch,
)
from src.reviewer import Finding, ReviewResult, _stream_review
from src.cross_checker import run_cross_check
from src.batch import BatchJob, BatchStatus, submit_verification_batch
from src import gui
from src import batch_state_store
from src.verifier import verify_finding, collect_verification_batch_results, VerificationResult
from src.diagnostics import DiagnosticsReport
from src.resume_state import (
    PHASE_REVIEW_POLL,
    PHASE_REVIEW_COLLECT,
    PHASE_VERIFICATION_POLL,
    PHASE_FINALIZE,
    build_resume_state,
    deserialize_resume_state,
    serialize_batch_job,
    deserialize_batch_job,
)


def test_extract_text_from_docx_builds_paragraph_map(tmp_path: Path):
    source = tmp_path / "mapping.docx"
    doc = Document()
    doc.add_paragraph("First paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "R1C1"
    table.rows[0].cells[1].text = "R1C2"
    table.rows[1].cells[0].text = "R2C1"
    table.rows[1].cells[1].text = ""
    doc.add_paragraph("Last paragraph")
    doc.save(source)

    spec = extract_text(source)

    assert spec.paragraph_map is not None
    assert "\n\n".join(m.text for m in spec.paragraph_map) == spec.content
    assert [m.body_index for m in spec.paragraph_map] == sorted(m.body_index for m in spec.paragraph_map)
    assert any(m.element_type == "table_cell" and m.table_index is not None for m in spec.paragraph_map)


def _make_submission(*, batch_id: str = "msgbatch_test", cross_check_enabled: bool = True, prepared_specs: list[ExtractedSpec] | None = None) -> BatchSubmission:
    return BatchSubmission(
        job=BatchJob(
            batch_id=batch_id,
            job_type="review",
            request_map={"review__spec__0": {"filename": "spec.docx", "index": 0, "type": "review"}},
            created_at=123.0,
        ),
        files_reviewed=["spec.docx"],
        review_request_ids=["review__spec__0"],
        cycle_label="2025",
        cross_check_enabled=cross_check_enabled,
        prepared_specs=prepared_specs,
    )


def _make_finding(issue: str = "Missing note") -> Finding:
    return Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1",
        issue=issue,
        actionType="EDIT",
        existingText="old",
        replacementText="new",
        codeReference="CBC",
        confidence=0.9,
    )


def _make_review_state(submission: BatchSubmission) -> CollectedBatchState:
    review = ReviewResult(findings=[_make_finding("Main review issue")], cross_check_status="completed")
    cross = ReviewResult(findings=[_make_finding("Cross-check issue")], cross_check_status="completed")
    return CollectedBatchState(
        submission=submission,
        review_result=review,
        files_reviewed=["spec.docx"],
        leed_alerts=[{"filename": "spec.docx", "type": "leed"}],
        placeholder_alerts=[{"filename": "spec.docx", "type": "placeholder"}],
        cross_check_result=cross,
        cross_check_skipped_due_to_missing_specs=True,
    )


def _make_dummy_app():
    class _Entry:
        def get(self):
            return "key"

    class _Var:
        def __init__(self):
            self.value = None

        def set(self, value):
            self.value = value

    class _Log:
        def log_error(self, _m):
            pass

        def log_warning(self, _m):
            pass

        def log(self, *_args, **_kwargs):
            pass

        def log_step(self, *_args, **_kwargs):
            pass

    class _Button:
        def set_processing(self):
            pass

        def configure(self, **_kwargs):
            pass

    class _Progress:
        def pack(self, **_kwargs):
            pass

        def set(self, _v):
            pass

        def configure(self, **_kwargs):
            pass

    dummy = type("Dummy", (), {})()
    dummy.api_key_entry = _Entry()
    dummy.log = _Log()
    dummy._cross_check_var = _Var()
    dummy.run_button = _Button()
    dummy.progress_bar = _Progress()
    dummy._poll_batch = lambda: None
    dummy._collect_batch_results = lambda: None
    dummy._resume_verification_poll = lambda _loaded: None
    dummy._on_review_complete = lambda _result: None
    dummy._reset_ui = lambda: None
    dummy._is_valid_verification_resume_state = lambda loaded: gui.SpecReviewApp._is_valid_verification_resume_state(dummy, loaded)
    dummy.after = lambda *_args, **_kwargs: None
    return dummy


def test_supported_extensions_docx_only(tmp_path: Path):
    assert SUPPORTED_EXTENSIONS == {".docx"}
    pdf = tmp_path / "sample.pdf"
    pdf.write_text("not a docx")
    with pytest.raises(ValueError):
        extract_text(pdf)


def test_extract_context_text_supports_docx_and_pdf(tmp_path: Path):
    assert CONTEXT_ATTACHMENT_EXTENSIONS == {".docx", ".pdf"}

    docx_path = tmp_path / "context.docx"
    doc = Document()
    doc.add_paragraph("Project alpha overview")
    doc.add_paragraph("Owner: Acme USD")
    doc.save(docx_path)
    docx_text = extract_context_text(docx_path)
    assert "Project alpha overview" in docx_text
    assert "Owner: Acme USD" in docx_text

    pypdf = pytest.importorskip("pypdf")
    pdf_path = tmp_path / "context.pdf"
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with pdf_path.open("wb") as fh:
        writer.write(fh)
    # A blank page produces no text; the helper must still succeed and return a string.
    assert isinstance(extract_context_text(pdf_path), str)

    bogus = tmp_path / "context.txt"
    bogus.write_text("hi")
    with pytest.raises(ValueError):
        extract_context_text(bogus)


def test_cycle_prompt_includes_current_editions():
    p2025 = get_system_prompt(CALIFORNIA_2025)
    assert "CBC 2025" in p2025
    assert "<findings_json>" in p2025
    assert "<FINDINGS_JSON>" not in p2025

    msg = get_single_spec_user_message("Body", "file.docx", cycle=CALIFORNIA_2025)
    assert "ASCE 7-22" in msg


def test_dedup_does_not_merge_different_edits():
    f1 = _make_finding("Same issue")
    f2 = Finding(severity="HIGH", fileName="b.docx", section="1", issue="Same issue", actionType="EDIT", existingText="different", replacementText="bar", codeReference="CBC", confidence=0.8)
    deduped = _deduplicate_findings([f1, f2])
    assert len(deduped) == 2


def test_cross_check_skip_status():
    result = run_cross_check([], [])
    assert result.cross_check_status == "skipped"


def test_stream_review_marks_incomplete_when_stop_reason_not_end_turn():
    class _FakeStream:
        def __init__(self):
            self.text_stream = iter(['[{"severity":"HIGH"'])

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def get_final_message(self):
            usage = type("Usage", (), {"input_tokens": 7, "output_tokens": 11})()
            return type("Resp", (), {"stop_reason": "max_tokens", "usage": usage})()

    class _FakeMessages:
        def stream(self, **_kwargs):
            return _FakeStream()

    class _FakeClient:
        messages = _FakeMessages()

    result = _stream_review(_FakeClient(), "sys", "user")

    assert result.parse_status == "incomplete"
    assert result.stop_reason == "max_tokens"
    assert result.error is not None
    assert "stop_reason: max_tokens" in result.error


def test_round_trip_review_poll_via_gui_save_load(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    state_path = tmp_path / "batch_state.json"
    monkeypatch.setattr(batch_state_store, "_batch_state_path", lambda: state_path)

    specs = [ExtractedSpec(filename="spec.docx", content="Section text", word_count=2, source_path="/tmp/spec.docx", source_format="docx")]
    submission = _make_submission(batch_id="msgbatch_test_roundtrip", cross_check_enabled=True, prepared_specs=specs)

    gui.save_batch_state(build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission))
    loaded = gui.load_batch_state()

    assert loaded is not None
    loaded_submission = loaded["submission"]
    assert loaded["phase"] == PHASE_REVIEW_POLL
    assert loaded_submission.job.batch_id == "msgbatch_test_roundtrip"
    assert loaded_submission.cycle_label == "2025"
    assert loaded_submission.cross_check_enabled is True
    assert loaded_submission.prepared_specs is not None
    assert loaded_submission.prepared_specs[0].filename == "spec.docx"


def test_verify_finding_accumulates_pause_turn_search_evidence(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")

    class _Block:
        def __init__(self, block_type, text=None, results=None):
            self.type = block_type
            if text is not None:
                self.text = text
            if results is not None:
                self.results = results

    class _UrlResult:
        def __init__(self, url: str):
            self.url = url

    class _Response:
        def __init__(self, stop_reason: str, content: list):
            self.stop_reason = stop_reason
            self.content = content

    class _StreamCtx:
        def __init__(self, response):
            self.response = response

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def get_final_message(self):
            return self.response

    class _Messages:
        def __init__(self):
            self.calls = 0

        def stream(self, **_kwargs):
            self.calls += 1
            if self.calls == 1:
                return _StreamCtx(
                    _Response(
                        "pause_turn",
                        [_Block("web_search_tool_result", results=[_UrlResult("https://iccsafe.org/a")])],
                    )
                )
            return _StreamCtx(
                _Response("end_turn", [_Block("text", text='{"verdict":"CONFIRMED","explanation":"ok","sources":[],"correction":null}')])
            )

    class _Client:
        messages = _Messages()

    monkeypatch.setattr("src.verifier._get_client", lambda: _Client())
    finding = _make_finding("Outdated reference")
    result = verify_finding(finding)

    assert result.verdict == "CONFIRMED"
    assert "https://iccsafe.org/a" in result.sources


def test_verify_finding_search_gate_passes_when_search_only_in_earlier_turn(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")

    class _Block:
        def __init__(self, block_type, text=None, results=None):
            self.type = block_type
            if text is not None:
                self.text = text
            if results is not None:
                self.results = results

    class _UrlResult:
        def __init__(self, url: str):
            self.url = url

    class _Response:
        def __init__(self, stop_reason: str, content: list):
            self.stop_reason = stop_reason
            self.content = content

    class _StreamCtx:
        def __init__(self, response):
            self.response = response

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def get_final_message(self):
            return self.response

    class _Messages:
        def __init__(self):
            self.calls = 0

        def stream(self, **_kwargs):
            self.calls += 1
            if self.calls == 1:
                return _StreamCtx(_Response("pause_turn", [_Block("web_search_tool_result", results=[_UrlResult("https://nfpa.org/b")])]))
            return _StreamCtx(_Response("end_turn", [_Block("text", text='{"verdict":"DISPUTED","explanation":"nope","sources":[],"correction":null}')]))

    class _Client:
        messages = _Messages()

    monkeypatch.setattr("src.verifier._get_client", lambda: _Client())
    finding = _make_finding("Claim that should be disputed")
    result = verify_finding(finding)

    assert result.verdict == "DISPUTED"
    assert "did not perform web search" not in (result.explanation or "").lower()


def test_verify_finding_single_response_still_works(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")

    class _Block:
        def __init__(self, block_type, text=None, results=None):
            self.type = block_type
            if text is not None:
                self.text = text
            if results is not None:
                self.results = results

    class _UrlResult:
        def __init__(self, url: str):
            self.url = url

    class _Response:
        def __init__(self):
            self.stop_reason = "end_turn"
            self.content = [
                _Block("web_search_tool_result", results=[_UrlResult("https://ashrae.org/c")]),
                _Block("text", text='{"verdict":"CORRECTED","explanation":"adjust text","sources":[],"correction":"new text"}'),
            ]

    class _StreamCtx:
        def __init__(self, response):
            self.response = response

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def get_final_message(self):
            return self.response

    class _Messages:
        def stream(self, **_kwargs):
            return _StreamCtx(_Response())

    class _Client:
        messages = _Messages()

    monkeypatch.setattr("src.verifier._get_client", lambda: _Client())
    finding = _make_finding("Minor wording issue")
    result = verify_finding(finding)

    assert result.verdict == "CORRECTED"
    assert "https://ashrae.org/c" in result.sources


def test_round_trip_review_collect_state_payload():
    submission = _make_submission(batch_id="msgbatch_collect")
    review_state = _make_review_state(submission)

    restored = deserialize_resume_state(build_resume_state(phase=PHASE_REVIEW_COLLECT, submission=submission, review_state=review_state))

    assert restored["phase"] == PHASE_REVIEW_COLLECT
    restored_review_state = restored["review_state"]
    assert restored_review_state.review_result.findings[0].issue == "Main review issue"
    assert restored_review_state.files_reviewed == ["spec.docx"]
    assert restored_review_state.leed_alerts[0]["type"] == "leed"
    assert restored_review_state.cross_check_skipped_due_to_missing_specs is True


def test_round_trip_verification_poll_state_payload():
    submission = _make_submission(batch_id="msgbatch_verify")
    review_state = _make_review_state(submission)
    verification_batch = BatchJob(batch_id="msgbatch_verify_inner", job_type="verify", request_map={"verify__0": {"finding_idx": 0}}, created_at=999.0)

    restored = deserialize_resume_state(
        build_resume_state(
            phase=PHASE_VERIFICATION_POLL,
            submission=submission,
            review_state=review_state,
            verification_batch=verification_batch,
            verification_started=True,
        )
    )

    assert restored["phase"] == PHASE_VERIFICATION_POLL
    assert restored["verification_batch"].batch_id == "msgbatch_verify_inner"
    assert restored["review_state"].review_result.findings[0].issue == "Main review issue"
    assert restored["resume_flags"]["verification_started"] is True


def test_round_trip_finalize_state_payload_uses_review_state():
    submission = _make_submission(batch_id="msgbatch_finalize")
    review_state = _make_review_state(submission)

    restored = deserialize_resume_state(
        build_resume_state(
            phase=PHASE_FINALIZE,
            submission=submission,
            review_state=review_state,
            verification_started=True,
            verification_completed=True,
        )
    )

    assert restored["phase"] == PHASE_FINALIZE
    assert restored["review_state"].review_result.findings[0].issue == "Main review issue"
    assert restored["review_state"].cross_check_result is not None
    assert restored["resume_flags"]["verification_completed"] is True


def test_verify_finding_non_end_turn_returns_unverified_with_stop_reason(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")

    class _StreamCtx:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def get_final_message(self):
            return type("Resp", (), {"stop_reason": "max_tokens", "content": []})()

    class _FakeMessages:
        def stream(self, **_kwargs):
            return _StreamCtx()

    class _FakeClient:
        messages = _FakeMessages()

    monkeypatch.setattr("src.verifier._get_client", lambda: _FakeClient())

    result = verify_finding(_make_finding("Issue"))
    assert result.verdict == "UNVERIFIED"
    assert "stop_reason: max_tokens" in result.explanation


def test_verification_batch_job_round_trip():
    job = BatchJob(batch_id="msgbatch_verify_roundtrip", job_type="verify", request_map={"verify__0": {"batch_idx": 0, "finding_idx": 3}}, created_at=10.0)
    payload = serialize_batch_job(job)
    restored = deserialize_batch_job(payload)
    assert restored.batch_id == job.batch_id
    assert restored.job_type == "verify"
    assert restored.request_map["verify__0"]["finding_idx"] == 3


def test_load_batch_state_legacy_phase_migrates_to_review_poll(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    state_path = tmp_path / "batch_state.json"
    monkeypatch.setattr(batch_state_store, "_batch_state_path", lambda: state_path)
    # Use a fresh ``saved_at`` so the fixture does not age past
    # BATCH_STATE_MAX_AGE_HOURS as wall-clock time advances.
    saved_at = datetime.now(timezone.utc).isoformat()
    state_path.write_text(
        json.dumps({
            "saved_at": saved_at,
            "phase": "review",
            "batch_id": "msgbatch_legacy",
            "job_type": "review",
            "request_map": {"review__spec__0": {"filename": "spec.docx", "index": 0, "type": "review"}},
            "created_at": 123.0,
            "files_reviewed": ["spec.docx"],
            "review_request_ids": ["review__spec__0"],
            "leed_alerts": [],
            "placeholder_alerts": [],
            "cross_check_enabled": False,
            "export_mode": False,
        }),
        encoding="utf-8",
    )
    loaded = gui.load_batch_state()
    assert loaded is not None
    assert loaded["phase"] == PHASE_REVIEW_POLL
    assert loaded["submission"].job.batch_id == "msgbatch_legacy"


def test_resume_batch_disables_cross_check_when_specs_missing():
    dummy = _make_dummy_app()
    submission = _make_submission(batch_id="msgbatch_crosscheck_missing", cross_check_enabled=True, prepared_specs=None)

    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_REVIEW_POLL, "submission": submission})

    assert dummy._cross_check_for_review is False


def test_resume_batch_phase_router_calls_expected_handler():
    called = {"poll": 0, "collect": 0, "verify": 0}
    dummy = _make_dummy_app()
    dummy._poll_batch = lambda: called.__setitem__("poll", called["poll"] + 1)
    dummy._collect_batch_results = lambda: called.__setitem__("collect", called["collect"] + 1)
    dummy._resume_verification_poll = lambda _loaded: called.__setitem__("verify", called["verify"] + 1)
    submission = _make_submission(batch_id="msgbatch_router")
    review_state = _make_review_state(submission)
    verification_batch = BatchJob(batch_id="msgbatch_verify_router", job_type="verify", request_map={"verify__0": {"finding_idx": 0}}, created_at=1.0)

    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_REVIEW_POLL, "submission": submission})
    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_REVIEW_COLLECT, "submission": submission})
    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_VERIFICATION_POLL, "submission": submission, "review_state": review_state, "verification_batch": verification_batch})
    assert called == {"poll": 1, "collect": 1, "verify": 1}


def test_run_cross_check_for_batch_skips_when_specs_missing():
    submission = _make_submission(batch_id="msgbatch_cross", cross_check_enabled=True)
    state = CollectedBatchState(submission=submission, review_result=ReviewResult(findings=[_make_finding("x")]))

    result = run_cross_check_for_batch(state, specs=None)

    assert result.cross_check_skipped_due_to_missing_specs is True
    assert result.cross_check_result is not None
    assert result.cross_check_result.cross_check_status == "skipped"


def test_resume_verification_state_rejected_when_batch_missing(monkeypatch: pytest.MonkeyPatch):
    dummy = _make_dummy_app()
    called = {"reset": 0, "resume": 0, "delete": 0}
    dummy._reset_ui = lambda: called.__setitem__("reset", called["reset"] + 1)
    dummy._resume_verification_poll = lambda _loaded: called.__setitem__("resume", called["resume"] + 1)
    monkeypatch.setattr(batch_state_store, "delete_batch_state", lambda: called.__setitem__("delete", called["delete"] + 1))
    submission = _make_submission(batch_id="msgbatch_invalid_verify_missing")
    review_state = _make_review_state(submission)

    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_VERIFICATION_POLL, "submission": submission, "review_state": review_state})

    assert called["resume"] == 0
    assert called["reset"] == 1
    assert called["delete"] == 1


def test_resume_verification_state_rejected_when_batch_id_malformed(monkeypatch: pytest.MonkeyPatch):
    dummy = _make_dummy_app()
    called = {"reset": 0, "resume": 0, "delete": 0}
    dummy._reset_ui = lambda: called.__setitem__("reset", called["reset"] + 1)
    dummy._resume_verification_poll = lambda _loaded: called.__setitem__("resume", called["resume"] + 1)
    monkeypatch.setattr(batch_state_store, "delete_batch_state", lambda: called.__setitem__("delete", called["delete"] + 1))
    submission = _make_submission(batch_id="msgbatch_invalid_verify_bad_id")
    review_state = _make_review_state(submission)
    bad_batch = BatchJob(batch_id="bad_id", job_type="verify", request_map={"verify__0": {"finding_idx": 0}}, created_at=1.0)

    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_VERIFICATION_POLL, "submission": submission, "review_state": review_state, "verification_batch": bad_batch})

    assert called["resume"] == 0
    assert called["reset"] == 1
    assert called["delete"] == 1


def test_resume_verification_state_rejected_when_review_state_missing(monkeypatch: pytest.MonkeyPatch):
    dummy = _make_dummy_app()
    called = {"reset": 0, "resume": 0, "delete": 0}
    dummy._reset_ui = lambda: called.__setitem__("reset", called["reset"] + 1)
    dummy._resume_verification_poll = lambda _loaded: called.__setitem__("resume", called["resume"] + 1)
    monkeypatch.setattr(batch_state_store, "delete_batch_state", lambda: called.__setitem__("delete", called["delete"] + 1))
    submission = _make_submission(batch_id="msgbatch_invalid_verify_missing_review")
    verification_batch = BatchJob(batch_id="msgbatch_verify_good", job_type="verify", request_map={"verify__0": {"finding_idx": 0}}, created_at=1.0)

    gui.SpecReviewApp._resume_batch(dummy, {"phase": PHASE_VERIFICATION_POLL, "submission": submission, "verification_batch": verification_batch})

    assert called["resume"] == 0
    assert called["reset"] == 1
    assert called["delete"] == 1


def test_discard_pending_batch_deletes_state(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    state_path = tmp_path / "batch_state.json"
    monkeypatch.setattr(batch_state_store, "_batch_state_path", lambda: state_path)
    submission = _make_submission(batch_id="msgbatch_discard")
    gui.save_batch_state(build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission))

    assert state_path.exists()
    gui.delete_batch_state()
    assert not state_path.exists()


def test_load_batch_state_invalid_submission_batch_id_deletes_file(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    state_path = tmp_path / "batch_state.json"
    monkeypatch.setattr(batch_state_store, "_batch_state_path", lambda: state_path)
    state_path.write_text(
        """{
  "version": "2.3.0",
  "saved_at": "2026-03-18T00:00:00+00:00",
  "phase": "review_poll",
  "submission": {
    "job": {
      "batch_id": "bad_batch",
      "job_type": "review",
      "request_map": {"review__spec__0": {"filename": "spec.docx", "index": 0, "type": "review"}},
      "created_at": 123.0,
      "status": "submitted"
    }
  }
}""",
        encoding="utf-8",
    )

    loaded = gui.load_batch_state()

    assert loaded is None
    assert not state_path.exists()


def test_verification_request_map_preserves_sorted_confidence_indices(monkeypatch: pytest.MonkeyPatch):
    class _FakeBatches:
        def create(self, requests):
            assert [r["custom_id"] for r in requests] == ["verify__0", "verify__1"]
            return type("B", (), {"id": "msgbatch_verify_idx"})()

    class _FakeMessages:
        batches = _FakeBatches()

    class _FakeClient:
        messages = _FakeMessages()

    monkeypatch.setattr("src.batch._get_client", lambda: _FakeClient())

    finding_a = Finding(severity="HIGH", fileName="a", section="1", issue="a", actionType="EDIT", existingText=None, replacementText=None, codeReference=None, confidence=0.9)
    finding_b = Finding(severity="HIGH", fileName="b", section="1", issue="b", actionType="EDIT", existingText=None, replacementText=None, codeReference=None, confidence=0.1)
    job = submit_verification_batch(
        [finding_a, finding_b],
        build_prompt_fn=lambda _: "prompt",
        system_prompt_fn=lambda _cycle: "system prompt",
    )
    # Lower confidence finding_b should be first in batch but still map to original index 1.
    assert job.request_map["verify__0"]["finding_idx"] == 1
    assert job.request_map["verify__1"]["finding_idx"] == 0


def test_submit_verification_batch_requires_system_prompt_fn():
    finding = Finding(severity="HIGH", fileName="a", section="1", issue="a", actionType="EDIT", existingText=None, replacementText=None, codeReference=None, confidence=0.9)
    with pytest.raises(TypeError):
        submit_verification_batch([finding], build_prompt_fn=lambda _: "prompt")


def test_start_batch_review_always_preserves_prepared_specs(monkeypatch: pytest.MonkeyPatch):
    specs = [ExtractedSpec(filename="spec.docx", content="Spec text", word_count=2, source_path="/tmp/spec.docx", source_format="docx")]
    prepared = type("Prepared", (), {"specs": specs, "leed_alerts": [], "placeholder_alerts": []})()
    monkeypatch.setattr("src.pipeline._prepare_specs", lambda **_kwargs: prepared)
    monkeypatch.setattr("src.pipeline.submit_review_batch", lambda _specs, **_kwargs: BatchJob(batch_id="msgbatch_review", job_type="review", request_map={"review__spec__0": {"filename": "spec.docx", "index": 0}}, created_at=1.0))

    submission = start_batch_review(input_dir=Path("."), cross_check_enabled=False)

    assert submission.prepared_specs is not None
    assert submission.prepared_specs[0].filename == "spec.docx"


def test_on_poll_result_terminal_status_collects_partial_results(monkeypatch: pytest.MonkeyPatch):
    dummy = _make_dummy_app()
    called = {"collect": 0, "save": 0}
    dummy._batch_submission = _make_submission(batch_id="msgbatch_poll_terminal")
    dummy._collect_batch_results = lambda: called.__setitem__("collect", called["collect"] + 1)
    dummy._diagnostics_report = None
    dummy.progress_bar = type("P", (), {"set": lambda self, _v: None})()
    dummy._schedule_next_poll = lambda _delay: None
    monkeypatch.setattr(batch_state_store, "save_batch_state", lambda _state: called.__setitem__("save", called["save"] + 1))

    status = BatchStatus(status="failed", processing=0, succeeded=1, errored=1, canceled=0, expired=0, total=2)
    gui.SpecReviewApp._on_poll_result(dummy, status)

    assert called["collect"] == 1
    assert called["save"] == 1


def test_collect_verification_batch_results_collects_after_terminal_status(monkeypatch: pytest.MonkeyPatch):
    finding = _make_finding("verification item")
    job = BatchJob(batch_id="msgbatch_verify_terminal", job_type="verify", request_map={"verify__0": {"finding_idx": 0}}, created_at=1.0)
    monkeypatch.setattr("src.verifier.poll_batch", lambda _batch_id: BatchStatus(status="failed", processing=0, succeeded=1, errored=0, canceled=0, expired=0, total=1))

    collect_verification_batch_results(job, [finding], log=lambda *_a, **_k: None, progress=lambda _p, _m: None, poll_interval=0)

    assert finding.verification is not None
    assert finding.verification.verdict == "UNVERIFIED"


def test_continuation_request_accepts_sdk_content_blocks():
    from src.verifier import _build_continuation_request

    content_blocks = [
        TextBlock(type="text", text="Searching code section 403...", citations=None),
        ToolUseBlock(type="tool_use", id="toolu_test123", name="web_search", input={"query": "2025 CMC outside air"}),
    ]
    request = _build_continuation_request("prompt text", content_blocks, cycle=CALIFORNIA_2025)

    transformed = maybe_transform({"requests": [{"custom_id": "verify__0", "params": request}]}, BatchCreateParams)

    assistant_message = transformed["requests"][0]["params"]["messages"][1]
    assert assistant_message["role"] == "assistant"
    assert assistant_message["content"][0]["type"] == "text"
    assert assistant_message["content"][1]["type"] == "tool_use"


def test_finalize_batch_result_sets_total_elapsed_seconds():
    submission = _make_submission(batch_id="msgbatch_elapsed")
    state = _make_review_state(submission)

    result = finalize_batch_result(state)

    assert isinstance(result, PipelineResult)
    assert result.total_elapsed_seconds is not None
    assert result.total_elapsed_seconds >= 0


def test_extract_text_includes_header_footer_content(tmp_path: Path):
    source = tmp_path / "header_footer.docx"
    doc = Document()
    doc.add_paragraph("Main body paragraph")
    doc.sections[0].header.paragraphs[0].text = "CBC (2019) EDITION"
    doc.sections[0].footer.paragraphs[0].text = "NORTHWOOD E. S. CLASSROOM BUILDING"
    doc.save(source)

    spec = extract_text(source)

    assert "===== HEADER/FOOTER CONTENT =====" in spec.content
    assert "[Header] CBC (2019) EDITION" in spec.content
    assert "[Footer] NORTHWOOD E. S. CLASSROOM BUILDING" in spec.content
    assert spec.paragraph_map is not None
    assert "\n\n".join(m.text for m in spec.paragraph_map) == spec.content
    assert any(m.element_type == "header" for m in spec.paragraph_map)
    assert any(m.element_type == "footer" for m in spec.paragraph_map)


def test_diagnostics_summary_counts_verification_phase_from_events():
    report = DiagnosticsReport()
    report.log("review", "info", "review start")
    report.log("review", "info", "verification event carried in review", {"verdict": "CONFIRMED", "confidence": 0.9})
    report.log("edit_application", "info", "edit applied")
    report.finish()

    summary = report.summary()
    assert "verification" in summary["phase_durations"]
