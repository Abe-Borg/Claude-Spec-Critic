"""The Cross-Spec Coordination section header must stay visible in Word.

Every per-finding "Sources" Heading 4 ships collapsed-by-default
(``<w15:collapsed/>``); Word folds subsequent paragraphs into that zone
until it reaches a paragraph whose outline level is at or above the
collapsed heading's. The Title-styled section headers ("Findings",
"Cross-Spec Coordination") carry no native outline level, so the
cross-check banner — and the standalone page-break paragraph in front of
it — used to be swallowed into the LAST finding's collapsed Sources
panel: hidden on open, absent from the Navigation Pane, with the
coordination findings dangling under the last severity group as if they
were extra per-spec findings.

These tests pin the fix: both Title headers carry an explicit
``outlineLvl=0`` (collapse-zone terminator + Navigation Pane entry)
while keeping the Title visual style, and the cross-check page break
rides on the heading itself instead of a swallowable standalone
paragraph.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.output.report_exporter import export_report
from src.review.reviewer import Finding, ReviewResult
from src.verification.verifier import VerificationResult

_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"


def _finding(*, file: str, issue: str, verification: VerificationResult | None = None) -> Finding:
    f = Finding(
        severity="MEDIUM",
        fileName=file,
        section="2.1",
        issue=issue,
        actionType="EDIT",
        existingText="old text",
        replacementText="new text",
        codeReference="CBC §1234",
        confidence=0.8,
    )
    f.verification = verification
    return f


def _verified() -> VerificationResult:
    return VerificationResult(
        verdict="CONFIRMED",
        explanation="Verified against CBC §1234.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
        cache_status="miss",
        model_used="claude-sonnet-4-6",
        verification_mode="standard_reasoning",
        web_search_requests=3,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(self, *, review_result: ReviewResult, cross_check_result=None):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = "2025"
        self.total_elapsed_seconds = 1.0


def _para_index(doc: Document, text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    return -1


def _outline_level(paragraph) -> int | None:
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    el = pPr.find(qn("w:outlineLvl"))
    return None if el is None else int(el.get(qn("w:val")))


def _is_collapsed(paragraph) -> bool:
    pPr = paragraph._p.pPr
    return pPr is not None and pPr.find(f"{{{_W15_NS}}}collapsed") is not None


class TestCrossCheckHeaderVisibility:
    def _doc(self, tmp_path: Path) -> Document:
        # Mirrors the bug shape: the last (only) review finding renders a
        # collapsed-by-default Sources panel immediately before the
        # cross-check section.
        review = ReviewResult(
            findings=[
                _finding(
                    file="23 25 00 - HVAC Water Treatment.docx",
                    issue="Stale ASHRAE 188 edition cited.",
                    verification=_verified(),
                )
            ]
        )
        cross = ReviewResult(
            findings=[
                _finding(
                    file="23 52 16 - Condensing Boilers.docx",
                    issue="Boiler protocol conflicts with DDC section.",
                )
            ],
            cross_check_status="completed",
            thinking="Boiler BAS integration protocol gap between sections.",
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=review, cross_check_result=cross), out
        )
        return Document(str(out))

    def test_header_carries_outline_level_zero(self, tmp_path: Path):
        doc = self._doc(tmp_path)
        idx = _para_index(doc, "Cross-Spec Coordination")
        assert idx != -1, "cross-check section header missing from report"
        header = doc.paragraphs[idx]
        # Outline level 0 terminates any open collapse zone and lists the
        # section in the Navigation Pane; the Title look is unchanged.
        assert _outline_level(header) == 0
        assert header.style.name == "Title"

    def test_fixture_renders_collapsed_sources_panel_before_header(self, tmp_path: Path):
        doc = self._doc(tmp_path)
        idx = _para_index(doc, "Cross-Spec Coordination")
        collapsed_sources = [
            i
            for i, p in enumerate(doc.paragraphs[:idx])
            if p.text.strip() == "Sources" and _is_collapsed(p)
        ]
        assert collapsed_sources, (
            "fixture must render a collapsed Sources panel before the "
            "section header — the exact shape that used to swallow it"
        )

    def test_page_break_rides_on_the_heading(self, tmp_path: Path):
        doc = self._doc(tmp_path)
        header = doc.paragraphs[_para_index(doc, "Cross-Spec Coordination")]
        assert header.paragraph_format.page_break_before is True
        # The old standalone page-break paragraph was unleveled body text
        # and got folded into the collapsed zone along with the header.
        for p in doc.paragraphs:
            for br in p._p.findall(".//" + qn("w:br")):
                assert br.get(qn("w:type")) != "page", (
                    "standalone page-break paragraph reintroduced before "
                    "the cross-check section"
                )

    def test_findings_header_carries_outline_level_zero(self, tmp_path: Path):
        doc = self._doc(tmp_path)
        idx = _para_index(doc, "Findings")
        assert idx != -1
        header = doc.paragraphs[idx]
        assert _outline_level(header) == 0
        assert header.style.name == "Title"
