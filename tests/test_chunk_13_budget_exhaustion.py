"""Chunk 13 tests — budget-exhaustion sentinel + banner counter.

Chunk 13 of the Trust Upgrade adds a per-finding sentinel that fires
when the verifier consumed its full mode-scaled web_search budget
without producing a grounded verdict. The contract has five surfaces:

* ``VerificationResult.budget_exhausted`` defaults to False, round-trips
  through resume state, and is set by both the real-time and batch
  paths when ``web_search_requests >= decision.web_search_max_uses``
  AND the final verdict is UNVERIFIED.
* ``VerificationCache.put`` refuses to persist
  ``budget_exhausted=True`` results — same transient-signal rationale
  as ``verification_failed`` (a re-run at a higher severity may
  succeed; freezing the shortfall as a durable verdict suppresses
  re-verification).
* ``report_status.is_budget_exhausted(finding)`` and
  :func:`summarize_budget_exhausted` expose the flag for renderer /
  banner consumption without touching private verifier state.
* The per-finding status line in the report appends a
  ``(search budget exhausted)`` sub-label when the flag is True. The
  trust-level classification stays INSUFFICIENT_EVIDENCE — no new
  top-level :class:`ReportStatus` value — but the sub-label
  distinguishes "verifier had no headroom" from "verifier ran out
  of evidence early".
* The Run Diagnostics banner has a new "Budget-exhausted findings"
  row (red-highlighted when > 0) and a recovery hint paragraph that
  points operators at the severity-tiered budget knob.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from src.core.code_cycles import DEFAULT_CYCLE
from src.output.report_exporter import (
    _summarize_run_diagnostics,
    _write_run_diagnostics_banner,
)
from src.output.report_status import (
    EditActionLabel,
    ReportStatus,
    classify_edit_action,
    classify_status,
    is_budget_exhausted,
    summarize_budget_exhausted,
    summarize_edit_actions,
    summarize_statuses,
)
from src.review.reviewer import EditProposal, Finding
from src.verification.verification_cache import VerificationCache
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    severity: str = "MEDIUM",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Cited model lookup failed",
    confidence: float = 0.6,
    action: str = "EDIT",
    existing: str | None = "old text",
    replacement: str | None = "new text",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
    )
    f.verification = verification
    return f


def _exhausted_verification(
    *,
    explanation: str = "Searched all 5 budgeted sources; no authoritative passage.",
    web_search_requests: int = 5,
) -> VerificationResult:
    """A canonical budget-exhausted result for tests."""
    return VerificationResult(
        verdict="UNVERIFIED",
        explanation=explanation,
        grounded=False,
        web_search_requests=web_search_requests,
        budget_exhausted=True,
        verification_mode="standard_reasoning",
    )


def _clean_unverified() -> VerificationResult:
    """A clean UNVERIFIED (no failure, no budget exhaustion)."""
    return VerificationResult(
        verdict="UNVERIFIED",
        explanation="Verifier searched 2 of 5 budgeted sources; nothing matched.",
        grounded=False,
        web_search_requests=2,
        budget_exhausted=False,
    )


# ---------------------------------------------------------------------------
# 1. VerificationResult field
# ---------------------------------------------------------------------------


class TestVerificationResultBudgetExhaustedField:
    def test_default_is_false(self):
        result = VerificationResult(verdict="UNVERIFIED")
        assert result.budget_exhausted is False

    def test_constructor_round_trip(self):
        result = VerificationResult(
            verdict="UNVERIFIED",
            explanation="budget used",
            budget_exhausted=True,
        )
        assert result.budget_exhausted is True


# ---------------------------------------------------------------------------
# 2. is_budget_exhausted helper
# ---------------------------------------------------------------------------


class TestIsBudgetExhausted:
    def test_returns_false_for_finding_without_verification(self):
        f = _finding(verification=None)
        assert is_budget_exhausted(f) is False

    def test_returns_false_when_flag_is_false(self):
        f = _finding(verification=_clean_unverified())
        assert is_budget_exhausted(f) is False

    def test_returns_true_when_flag_is_set(self):
        f = _finding(verification=_exhausted_verification())
        assert is_budget_exhausted(f) is True

    def test_defaults_to_false_on_legacy_result_without_attribute(self):
        # Strip the attribute to simulate a payload constructed before
        # the field existed. The helper must not crash.
        v = _clean_unverified()
        delattr(v, "budget_exhausted")
        f = _finding(verification=v)
        assert is_budget_exhausted(f) is False


# ---------------------------------------------------------------------------
# 3. classify_status — same trust level, no new enum
# ---------------------------------------------------------------------------


class TestClassifyStatusBudgetExhausted:
    def test_budget_exhausted_classifies_as_insufficient_evidence(self):
        # Chunk 13: no new top-level ReportStatus — exhausted findings
        # stay on INSUFFICIENT_EVIDENCE. The sub-label is a rendering
        # concern, not a trust-level one.
        f = _finding(verification=_exhausted_verification())
        assert classify_status(f) is ReportStatus.INSUFFICIENT_EVIDENCE

    def test_failure_sentinel_takes_priority_over_budget_exhausted(self):
        # If a result improbably carries both flags, the operational
        # failure surface wins (an operator needs to see the crash
        # before the budget knob). We don't change classify_status —
        # this asserts the existing ordering still holds.
        v = VerificationResult(
            verdict="UNVERIFIED",
            verification_failed=True,
            budget_exhausted=True,
            web_search_requests=5,
        )
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFICATION_FAILED

    def test_supportive_verdict_with_high_search_count_does_not_set_status(self):
        # The verifier never sets budget_exhausted on a grounded
        # CONFIRMED; this asserts the classification still works when
        # search_count is high but grounded is True (defensive).
        v = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://example.com/"],
            accepted_sources=["https://example.com/"],
            web_search_requests=5,
        )
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFIED_SUPPORTED


# ---------------------------------------------------------------------------
# 4. classify_edit_action — proposal presence drives the label
# ---------------------------------------------------------------------------


class TestClassifyEditActionBudgetExhausted:
    def test_exhausted_with_proposal_is_edit_suggested(self):
        # The app emits edit instructions but never applies them. A
        # budget-exhausted finding with a proposal is labeled
        # EDIT_SUGGESTED; its INSUFFICIENT_EVIDENCE status rides along in
        # the sidecar for a downstream applier to act on.
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.99,
        )
        f = _finding(
            verification=_exhausted_verification(),
            edit_proposal=proposal,
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_exhausted_without_proposal_routes_to_report_only(self):
        f = _finding(
            verification=_exhausted_verification(),
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            edit_proposal=None,
        )
        assert classify_edit_action(f) is EditActionLabel.REPORT_ONLY


# ---------------------------------------------------------------------------
# 5. Cache — refuses to persist budget_exhausted=True
# ---------------------------------------------------------------------------


class TestCacheRejectsBudgetExhausted:
    def _finding_for_cache(self) -> Finding:
        return Finding(
            severity="HIGH",
            fileName="Section_22_1000.docx",
            section="2.1",
            issue="claim about Acme model XYZ-42",
            actionType="EDIT",
            existingText="model XYZ-42",
            replacementText="model XYZ-42 (current)",
            codeReference=None,
            confidence=0.6,
        )

    def test_ungrounded_exhausted_result_is_not_cached(self):
        # The standard case: UNVERIFIED + grounded=False +
        # budget_exhausted=True. The grounded guard already drops
        # this, but the explicit guard is defense-in-depth.
        cache = VerificationCache()
        f = self._finding_for_cache()
        cache.put(
            f,
            cycle=DEFAULT_CYCLE,
            result=_exhausted_verification(),
        )
        assert cache.get(f, cycle=DEFAULT_CYCLE) is None
        assert cache.stats()["size"] == 0

    def test_grounded_exhausted_result_is_not_cached_either(self):
        # Contrived case: grounded=True + sources + budget_exhausted=True
        # (the production paths never produce this combination, but
        # the explicit guard prevents future call sites from caching it).
        cache = VerificationCache()
        f = self._finding_for_cache()
        cache.put(
            f,
            cycle=DEFAULT_CYCLE,
            result=VerificationResult(
                verdict="CONFIRMED",
                grounded=True,
                sources=["https://example.com/"],
                accepted_sources=["https://example.com/"],
                source_quote="snippet",
                budget_exhausted=True,
            ),
        )
        assert cache.get(f, cycle=DEFAULT_CYCLE) is None

    def test_clean_grounded_result_still_caches(self):
        # Sanity check: the new guard does not block a normal
        # grounded verdict (no budget_exhausted flag set).
        cache = VerificationCache()
        f = self._finding_for_cache()
        cache.put(
            f,
            cycle=DEFAULT_CYCLE,
            result=VerificationResult(
                verdict="CONFIRMED",
                grounded=True,
                sources=["https://example.com/"],
                accepted_sources=["https://example.com/"],
                source_quote="snippet",
            ),
        )
        hit = cache.get(f, cycle=DEFAULT_CYCLE)
        assert hit is not None
        assert hit.verdict == "CONFIRMED"
        # The hit clone must default budget_exhausted to False — the
        # cache layer never persists it.
        assert hit.budget_exhausted is False


# ---------------------------------------------------------------------------
# 7. summarize_budget_exhausted aggregation
# ---------------------------------------------------------------------------


class TestSummarizeBudgetExhausted:
    def test_empty_returns_zero(self):
        assert summarize_budget_exhausted([]) == 0

    def test_clean_finding_does_not_count(self):
        f = _finding(verification=_clean_unverified())
        assert summarize_budget_exhausted([f]) == 0

    def test_one_exhausted_counts_as_one(self):
        f = _finding(verification=_exhausted_verification())
        assert summarize_budget_exhausted([f]) == 1

    def test_multiple_findings_are_summed(self):
        findings = [
            _finding(file="a.docx", verification=_exhausted_verification()),
            _finding(file="b.docx", verification=_exhausted_verification()),
            _finding(file="c.docx", verification=_clean_unverified()),
        ]
        assert summarize_budget_exhausted(findings) == 2

    def test_finding_without_verification_does_not_count(self):
        f = _finding(verification=None)
        assert summarize_budget_exhausted([f]) == 0


# ---------------------------------------------------------------------------
# 8. _summarize_run_diagnostics includes the count
# ---------------------------------------------------------------------------


def _findings_to_summary(findings: list[Finding]) -> dict:
    status_counts = summarize_statuses(findings)
    edit_action_counts = summarize_edit_actions(findings)
    return _summarize_run_diagnostics(
        findings=findings,
        status_counts=status_counts,
        edit_action_counts=edit_action_counts,
        cross_check_result=None,
        pipeline_result=None,
    )


class TestSummaryIncludesBudgetExhausted:
    def test_clean_run_has_zero_budget_exhausted(self):
        summary = _findings_to_summary([])
        assert summary["budget_exhausted_count"] == 0

    def test_run_with_exhausted_findings_has_positive_count(self):
        f = _finding(verification=_exhausted_verification())
        summary = _findings_to_summary([f])
        assert summary["budget_exhausted_count"] == 1

    def test_summary_keys_are_stable(self):
        # The renderer reads ``budget_exhausted_count`` by key, so a
        # rename would silently break the banner. Pin the key here.
        summary = _findings_to_summary([])
        assert "budget_exhausted_count" in summary


# ---------------------------------------------------------------------------
# 9. _write_run_diagnostics_banner renders the row + hint
# ---------------------------------------------------------------------------


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestBannerRendersBudgetExhausted:
    def test_banner_row_appears_with_count(self):
        doc = Document()
        summary = {
            "edit_suggested": 1,
            "report_only": 0,
            "suppressed": 0,
            "verification_failed": 0,
            "cache_replay_count": 0,
            "oldest_cache_age_days": None,
            "demotion_count": 0,
            "extraction_warning_count": 0,
            "cross_check": None,
            "budget_exhausted_count": 2,
        }
        _write_run_diagnostics_banner(doc, summary)
        text = _all_text_from(doc)
        assert "Budget-exhausted findings" in text
        # The count must appear in the table.
        assert "2" in text

    def test_zero_count_renders_without_hint(self):
        doc = Document()
        summary = {
            "edit_suggested": 1,
            "report_only": 0,
            "suppressed": 0,
            "verification_failed": 0,
            "cache_replay_count": 0,
            "oldest_cache_age_days": None,
            "demotion_count": 0,
            "extraction_warning_count": 0,
            "cross_check": None,
            "budget_exhausted_count": 0,
        }
        _write_run_diagnostics_banner(doc, summary)
        text = _all_text_from(doc)
        # Row label still appears (the banner shape is stable).
        assert "Budget-exhausted findings" in text
        # But the recovery-hint sentence must NOT render at zero.
        assert "exhausted the verifier's web_search budget" not in text

    def test_nonzero_count_renders_recovery_hint(self):
        doc = Document()
        summary = {
            "edit_suggested": 1,
            "report_only": 0,
            "suppressed": 0,
            "verification_failed": 0,
            "cache_replay_count": 0,
            "oldest_cache_age_days": None,
            "demotion_count": 0,
            "extraction_warning_count": 0,
            "cross_check": None,
            "budget_exhausted_count": 3,
        }
        _write_run_diagnostics_banner(doc, summary)
        text = _all_text_from(doc)
        # The hint paragraph references the severity-budget knob so
        # the reviewer knows what action to take.
        assert "exhausted the verifier's web_search budget" in text
        assert "CRITICAL" in text and "GRIPES" in text

    def test_recovery_hint_numbers_match_severity_budget(self):
        # Regression guard: the hint used to hard-code "CRITICAL / HIGH
        # receive 7 searches", but _SEVERITY_MAX_USES has CRITICAL=8. The
        # numbers are now rendered from web_search_max_uses_for_severity so
        # they can't drift. Assert each severity's number matches policy.
        from src.core.api_config import (
            _SEVERITY_MAX_USES,
            web_search_max_uses_for_severity,
        )

        doc = Document()
        summary = {
            "edit_suggested": 1,
            "report_only": 0,
            "suppressed": 0,
            "verification_failed": 0,
            "cache_replay_count": 0,
            "oldest_cache_age_days": None,
            "demotion_count": 0,
            "extraction_warning_count": 0,
            "cross_check": None,
            "budget_exhausted_count": 2,
        }
        _write_run_diagnostics_banner(doc, summary)
        text = _all_text_from(doc)
        for sev, expected in _SEVERITY_MAX_USES.items():
            assert expected == web_search_max_uses_for_severity(sev)
            assert f"{sev} {expected}" in text, (
                f"Expected the hint to render '{sev} {expected}' from policy."
            )
        # CRITICAL is 8, not 7 — the old wrong text must be gone, and the
        # dead-end "raise CRITICAL severity" framing replaced.
        assert "CRITICAL 8" in text
        assert "receive 7 searches" not in text
        assert "CRITICAL findings already receive the maximum" in text


# ---------------------------------------------------------------------------
# 10. Verifier source inspection — the detection lives in the call site
# ---------------------------------------------------------------------------


class TestVerifierSourceInspection:
    """Belt-and-suspenders: read the verifier source and confirm both
    the real-time path and the batch wave path set ``budget_exhausted``
    on UNVERIFIED-with-budget-hit results. The end-to-end path can't
    be driven without a real API call; source inspection catches a
    future refactor that drops the flag.
    """

    def test_make_unverified_accepts_budget_exhausted_kwarg(self):
        source = Path("src/verification/verifier.py").read_text(encoding="utf-8")
        # The helper signature must accept the new kwarg so the
        # not-grounded early returns can flag exhausted budget.
        assert "budget_exhausted: bool = False," in source

    def test_real_time_path_sets_budget_exhausted(self):
        source = Path("src/verification/verifier.py").read_text(encoding="utf-8")
        # The success path must compute and stamp the flag after
        # _enforce_grounding_invariant so a downgraded verdict still
        # picks up the sub-label.
        assert "budget_was_exhausted" in source
        assert "parsed.budget_exhausted = True" in source

    def test_batch_wave_path_sets_budget_exhausted(self):
        source = Path("src/verification/verifier.py").read_text(encoding="utf-8")
        # The batch path must stamp the flag too — both paths must
        # apply the same condition.
        assert "parsed.budget_exhausted = True" in source
        # The batch path comparison should reference web_search_max_uses
        # so it tracks the routing decision the request was built with.
        assert "decision.web_search_max_uses" in source
