"""Chunk O tests: deterministic checks expansion.

Plan section "Chunk O — Deterministic Checks Expansion". The chunk:

* adds three new deterministic detectors to ``preprocessor.py`` —
  ``detect_unresolved_template_markers``, ``detect_invalid_code_cycle_strings``,
  and ``detect_duplicate_paragraphs``;
* stamps every alert dict with a stable ``deterministic_rule`` identifier so
  consumers can branch on a rule id instead of sniffing the human label;
* propagates every alert list (LEED, placeholder, code-cycle, structural,
  naming, plus the three new ones) through ``_PreparedSpecs`` →
  ``BatchSubmission`` → ``CollectedBatchState`` → ``PipelineResult``;
* renders every alert category in the report exporter's Alerts section
  with a ``(deterministic check)`` suffix per directive 2;
* extends the verification router's local-skip keyword list so a finding
  whose text mentions one of the new rules does not pay for a Sonnet+
  web_search call.

Coverage is organized into one class per rule + integration smoke checks
for the pipeline plumbing, report rendering, verification routing, and
resume-state round-trip.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.core.code_cycles import CALIFORNIA_2025
from src.input.preprocessor import (
    DETERMINISTIC_RULE_DUPLICATE_HEADING,
    DETERMINISTIC_RULE_DUPLICATE_PARAGRAPH,
    DETERMINISTIC_RULE_EMPTY_SECTION,
    DETERMINISTIC_RULE_INCONSISTENT_FILENAME,
    DETERMINISTIC_RULE_INVALID_CODE_CYCLE,
    DETERMINISTIC_RULE_LEED,
    DETERMINISTIC_RULE_PLACEHOLDER,
    DETERMINISTIC_RULE_STALE_ASCE7,
    DETERMINISTIC_RULE_STALE_CODE_CYCLE,
    DETERMINISTIC_RULE_TEMPLATE_MARKER,
    PreprocessResult,
    detect_duplicate_headings,
    detect_duplicate_paragraphs,
    detect_empty_sections,
    detect_inconsistent_file_naming,
    detect_invalid_code_cycle_strings,
    detect_leed_references,
    detect_placeholders,
    detect_stale_code_cycle_references,
    detect_unresolved_template_markers,
    preprocess_spec,
)
from src.review.reviewer import Finding
from src.verification_router import classify_finding_for_verification


# ---------------------------------------------------------------------------
# Rule id wiring
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# detect_unresolved_template_markers
# ---------------------------------------------------------------------------

class TestUnresolvedTemplateMarkers:
    def test_flags_todo_with_colon(self) -> None:
        alerts = detect_unresolved_template_markers("TODO: confirm with PM.", "s.docx")
        assert any("TODO" in a["match"].upper() for a in alerts)
        assert all(a["deterministic_rule"] == DETERMINISTIC_RULE_TEMPLATE_MARKER for a in alerts)

    def test_flags_todo_followed_by_uppercase_word(self) -> None:
        # Bare ``TODO Confirm`` is also flagged so capitalized continuations
        # don't slip past the colon-only rule.
        alerts = detect_unresolved_template_markers("TODO Confirm hanger spacing.", "s.docx")
        assert any("TODO" in a["match"].upper() for a in alerts)

    def test_does_not_flag_lowercase_to_do_phrase(self) -> None:
        # "things to do" in prose must not trigger.
        alerts = detect_unresolved_template_markers(
            "There are several things to do before submittal.", "s.docx"
        )
        assert alerts == []

    def test_flags_fixme(self) -> None:
        alerts = detect_unresolved_template_markers("FIXME before issue.", "s.docx")
        assert any("FIXME" in a["match"].upper() for a in alerts)

    def test_flags_xxx_marker(self) -> None:
        alerts = detect_unresolved_template_markers("Confirm XXX before issue.", "s.docx")
        assert any("XXX" in a["match"].upper() for a in alerts)

    def test_does_not_flag_model_number_like_xxx_dash(self) -> None:
        # Model numbers ("XXX-12") and digits ("XXX2") should not trigger.
        alerts = detect_unresolved_template_markers("Model XXX-12 specified.", "s.docx")
        assert alerts == []

    def test_flags_question_mark_placeholder(self) -> None:
        alerts = detect_unresolved_template_markers("Capacity: ??? gpm.", "s.docx")
        assert any("???" in a["match"] for a in alerts)

    def test_does_not_flag_double_question(self) -> None:
        # Two question marks ("Is it correct??") are not the placeholder
        # we want to flag.
        alerts = detect_unresolved_template_markers("Is this correct??", "s.docx")
        assert alerts == []

    def test_flags_lorem_ipsum(self) -> None:
        alerts = detect_unresolved_template_markers(
            "Lorem ipsum dolor sit amet.", "s.docx"
        )
        assert any("ipsum" in a["match"].lower() for a in alerts)

    def test_alert_dict_has_expected_keys(self) -> None:
        alerts = detect_unresolved_template_markers("TODO: fix.", "spec.docx")
        assert alerts
        a = alerts[0]
        assert a["filename"] == "spec.docx"
        assert a["deterministic_rule"] == DETERMINISTIC_RULE_TEMPLATE_MARKER
        assert a["position"] == 0
        assert "context" in a


# ---------------------------------------------------------------------------
# detect_invalid_code_cycle_strings
# ---------------------------------------------------------------------------

class TestInvalidCodeCycleStrings:
    def test_flags_2018_cbc_as_invalid(self) -> None:
        # California never published a 2018 cycle.
        alerts = detect_invalid_code_cycle_strings("Per the 2018 CBC.", "s.docx")
        assert any("2018" in a["match"] for a in alerts)
        assert all(a["deterministic_rule"] == DETERMINISTIC_RULE_INVALID_CODE_CYCLE for a in alerts)
        assert all(a["found_year"] == "2018" for a in alerts)

    def test_flags_2020_cmc_as_invalid(self) -> None:
        alerts = detect_invalid_code_cycle_strings("See 2020 CMC for venting.", "s.docx")
        assert any("2020" in a["match"] for a in alerts)

    def test_does_not_flag_real_cycle_years(self) -> None:
        # Each of these is a real California cycle year and must not trigger.
        for year in ("2010", "2013", "2016", "2019", "2022", "2025", "2028"):
            alerts = detect_invalid_code_cycle_strings(f"Per {year} CBC.", "s.docx")
            assert alerts == [], f"unexpectedly flagged real cycle year {year}"

    def test_does_not_flag_year_without_code_abbrev(self) -> None:
        # A bare year ("In 2018, the school...") with no code reference
        # must not trigger.
        alerts = detect_invalid_code_cycle_strings(
            "In 2018, the school was renovated.", "s.docx"
        )
        assert alerts == []

    def test_disjoint_from_stale_cycle_detector(self) -> None:
        # 2019 CBC is *stale* (real history). 2018 CBC is *invalid*. The two
        # detectors must not double-count the same span. Inputs are
        # constructed so each detector only sees the year it owns.
        content = "Per the 2019 CBC. Per the 2018 CBC."
        stale = detect_stale_code_cycle_references(content, "s.docx", CALIFORNIA_2025)
        invalid = detect_invalid_code_cycle_strings(content, "s.docx")
        stale_years = {a.get("found_year") for a in stale}
        invalid_years = {a.get("found_year") for a in invalid}
        assert "2019" in stale_years and "2018" not in stale_years
        assert "2018" in invalid_years and "2019" not in invalid_years

    def test_flags_california_code_full_name_invalid_year(self) -> None:
        # The third stale-cycle pattern ("2024 California Building Code")
        # should also surface as invalid.
        alerts = detect_invalid_code_cycle_strings(
            "Comply with 2024 California Building Code.", "s.docx"
        )
        assert any("2024" in a["match"] for a in alerts)


# ---------------------------------------------------------------------------
# detect_duplicate_paragraphs
# ---------------------------------------------------------------------------

class TestDuplicateParagraphs:
    def test_flags_repeated_long_paragraph(self) -> None:
        para = (
            "Submittals shall be provided for all piping accessories within "
            "10 days of award and shall include manufacturer cut sheets."
        )
        content = f"1.01 GENERAL\n\n{para}\n\n2.01 PRODUCTS\n\n{para}"
        alerts = detect_duplicate_paragraphs(content, "s.docx")
        assert len(alerts) == 1
        assert alerts[0]["deterministic_rule"] == DETERMINISTIC_RULE_DUPLICATE_PARAGRAPH
        assert alerts[0]["occurrence_count"] == 2

    def test_does_not_flag_short_repeats(self) -> None:
        # Short headings ("PART 1") repeat by design and must not trigger.
        content = "PART 1\n\nbody\n\nPART 1\n\nmore body"
        assert detect_duplicate_paragraphs(content, "s.docx") == []

    def test_normalizes_whitespace_and_case(self) -> None:
        para = (
            "Provide all anchorage hardware in stainless steel where exposed "
            "to weather, per the structural drawings and SSF-12."
        )
        content = (
            f"{para}\n\n"
            # Same paragraph but with extra whitespace and different case.
            f"  PROVIDE   ALL  anchorage  hardware  in  stainless  steel  "
            f"WHERE  exposed  to  weather,  per  the  structural  drawings  "
            f"and  SSF-12.   "
        )
        alerts = detect_duplicate_paragraphs(content, "s.docx")
        assert len(alerts) == 1

    def test_reports_each_occurrence_after_the_first(self) -> None:
        para = (
            "Cleaning shall be performed at the end of each shift and after "
            "all penetrations are sealed per the manufacturer instructions."
        )
        content = "\n\n".join([para, para, para])  # 3 copies → 2 alerts
        alerts = detect_duplicate_paragraphs(content, "s.docx")
        assert len(alerts) == 2

    def test_respects_min_length_kwarg(self) -> None:
        para = "Short clause but long enough."
        content = f"{para}\n\n{para}"
        # default min_length=80 → no alert
        assert detect_duplicate_paragraphs(content, "s.docx") == []
        # custom 20 → flagged
        alerts = detect_duplicate_paragraphs(content, "s.docx", min_length=20)
        assert len(alerts) == 1

    def test_empty_content_returns_empty_list(self) -> None:
        assert detect_duplicate_paragraphs("", "s.docx") == []


# ---------------------------------------------------------------------------
# preprocess_spec aggregator
# ---------------------------------------------------------------------------

class TestPreprocessSpecAggregator:
    def test_aggregates_all_chunk_o_alerts(self) -> None:
        long_dup = (
            "Provide a complete and operational system tested per the "
            "manufacturer's instructions before substantial completion."
        )
        content = (
            "1.01 GENERAL\n\n"
            f"{long_dup}\n\n"
            "TODO: confirm pipe sizing.\n\n"
            "Per the 2018 CBC.\n\n"
            "[INSERT PROJECT NAME]\n\n"
            f"{long_dup}\n\n"
            "LEED Silver targeted."
        )
        result = preprocess_spec(content, "23 21 13.docx", cycle=CALIFORNIA_2025)
        assert isinstance(result, PreprocessResult)
        assert result.template_marker_alerts, "template marker not detected"
        assert result.invalid_code_cycle_alerts, "invalid code cycle not detected"
        assert result.duplicate_paragraph_alerts, "duplicate paragraph not detected"
        # Existing detectors keep working alongside the new ones.
        assert result.placeholder_alerts
        assert result.leed_alerts

    def test_no_cycle_still_runs_chunk_o_detectors(self) -> None:
        # The new detectors do not require a cycle, so callers that pass
        # ``cycle=None`` still get template / duplicate / (no-op invalid)
        # results without crashing.
        content = "TODO: confirm.\n\n" + "x" * 200
        result = preprocess_spec(content, "s.docx")
        assert result.template_marker_alerts
        # invalid_code_cycle_alerts is fine to be empty when no code cite.
        assert result.invalid_code_cycle_alerts == []


# ---------------------------------------------------------------------------
# Pipeline plumbing — alerts flow from prepare → submission → result
# ---------------------------------------------------------------------------

class TestPipelinePlumbing:
    def test_finalize_batch_result_forwards_chunk_o_alerts(self) -> None:
        """finalize_batch_result copies every alert list onto the result."""
        from src.batch import BatchJob
        from src.pipeline import BatchSubmission, CollectedBatchState, finalize_batch_result
        from src.review.reviewer import ReviewResult

        sub = BatchSubmission(
            job=BatchJob(batch_id="msgbatch_test", job_type="review", request_map={}, created_at=0.0),
        )
        sentinel_codecycle = [{"filename": "s.docx", "match": "2019 CBC", "deterministic_rule": DETERMINISTIC_RULE_STALE_CODE_CYCLE}]
        sentinel_template = [{"filename": "s.docx", "match": "TODO: x", "deterministic_rule": DETERMINISTIC_RULE_TEMPLATE_MARKER}]
        state = CollectedBatchState(
            submission=sub,
            review_result=ReviewResult(findings=[]),
            code_cycle_alerts=sentinel_codecycle,
            template_marker_alerts=sentinel_template,
        )
        result = finalize_batch_result(state)
        assert result.code_cycle_alerts == sentinel_codecycle
        assert result.template_marker_alerts == sentinel_template

    def test_collect_review_batch_results_forwards_submission_alerts(self, monkeypatch) -> None:
        """collect_review_batch_results copies submission alerts onto state."""
        from src.batch import BatchJob
        from src.pipeline import BatchSubmission, collect_review_batch_results

        # Stub the network-facing retrieve_review_results so this test stays
        # hermetic. An empty result map means no findings are produced.
        monkeypatch.setattr("src.pipeline.retrieve_review_results", lambda job, model: {})
        monkeypatch.setattr(
            "src.pipeline._recover_retryable_review_batch_results",
            lambda submission, results, log: results,
        )

        sub = BatchSubmission(
            job=BatchJob(batch_id="msgbatch_test", job_type="review", request_map={}, created_at=0.0),
            code_cycle_alerts=[{"filename": "s.docx", "match": "2019 CBC"}],
            duplicate_paragraph_alerts=[{"filename": "s.docx", "match": "x" * 80}],
        )
        state = collect_review_batch_results(sub)
        assert state.code_cycle_alerts == sub.code_cycle_alerts
        assert state.duplicate_paragraph_alerts == sub.duplicate_paragraph_alerts


# ---------------------------------------------------------------------------
# Verification router — new keywords route to local_skip
# ---------------------------------------------------------------------------

class TestVerificationRouterChunkO:
    """The router treats GRIPES findings about the new rules as local_skip."""

    @pytest.fixture
    def gripe_finding(self) -> Finding:
        return Finding(
            severity="GRIPES",
            fileName="s.docx",
            section="2.1",
            issue="placeholder",
            actionType="EDIT",
            existingText=None,
            replacementText=None,
            codeReference=None,
        )

    def test_todo_finding_routes_to_local_skip(self, gripe_finding: Finding) -> None:
        gripe_finding.issue = "Unresolved TODO marker in section 2.1"
        assert classify_finding_for_verification(gripe_finding) == "local_skip"

    def test_fixme_finding_routes_to_local_skip(self, gripe_finding: Finding) -> None:
        gripe_finding.issue = "FIXME left in the spec"
        assert classify_finding_for_verification(gripe_finding) == "local_skip"

    def test_invalid_code_cycle_finding_routes_to_local_skip(self, gripe_finding: Finding) -> None:
        gripe_finding.issue = "Invalid code cycle year 2018"
        assert classify_finding_for_verification(gripe_finding) == "local_skip"

    def test_duplicate_paragraph_finding_routes_to_local_skip(self, gripe_finding: Finding) -> None:
        gripe_finding.issue = "Duplicate paragraph in submittals section"
        assert classify_finding_for_verification(gripe_finding) == "local_skip"

    def test_high_severity_overrides_local_skip(self, gripe_finding: Finding) -> None:
        # Severity gate: anything above GRIPES needs web verification even if
        # the issue text looks local.
        gripe_finding.severity = "HIGH"
        gripe_finding.issue = "Duplicate paragraph in submittals section"
        assert classify_finding_for_verification(gripe_finding) == "web_required"

    def test_code_reference_overrides_local_skip(self, gripe_finding: Finding) -> None:
        gripe_finding.issue = "FIXME — cycle reference may be wrong"
        gripe_finding.codeReference = "CBC 1605"
        assert classify_finding_for_verification(gripe_finding) == "web_required"


# ---------------------------------------------------------------------------
# Report exporter integration — every alert section renders
# ---------------------------------------------------------------------------

class _StubPipelineResult:
    """Duck-typed PipelineResult that exercises every alert section.

    Kept inline here so the chunk O snapshot tests don't share a fixture
    with the chunk N tests (each chunk owns its own surface).
    """

    def __init__(self, **kwargs) -> None:
        from src.review.reviewer import ReviewResult

        self.review_result = kwargs.get("review_result", ReviewResult(findings=[]))
        self.cross_check_result = None
        self.files_reviewed = kwargs.get("files_reviewed", ["s.docx"])
        self.cycle_label = "2025"
        self.total_elapsed_seconds = 1.0
        self.leed_alerts = kwargs.get("leed_alerts", [])
        self.placeholder_alerts = kwargs.get("placeholder_alerts", [])
        self.code_cycle_alerts = kwargs.get("code_cycle_alerts", [])
        self.structural_alerts = kwargs.get("structural_alerts", [])
        self.naming_alerts = kwargs.get("naming_alerts", [])
        self.template_marker_alerts = kwargs.get("template_marker_alerts", [])
        self.invalid_code_cycle_alerts = kwargs.get("invalid_code_cycle_alerts", [])
        self.duplicate_paragraph_alerts = kwargs.get("duplicate_paragraph_alerts", [])


def _alert(text: str, rule_id: str, filename: str = "s.docx") -> dict:
    return {
        "filename": filename,
        "type": "test alert",
        "match": text,
        "context": text,
        "position": 0,
        "deterministic_rule": rule_id,
    }


def _doc_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestReportExporterChunkOIntegration:
    def test_export_renders_template_marker_section(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult(
            template_marker_alerts=[_alert("TODO: confirm hanger spacing.", DETERMINISTIC_RULE_TEMPLATE_MARKER)],
        )
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        assert "Unresolved Template Markers" in text
        assert "(deterministic check)" in text
        assert "TODO: confirm hanger spacing." in text

    def test_export_renders_invalid_code_cycle_section(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult(
            invalid_code_cycle_alerts=[_alert("2018 CBC", DETERMINISTIC_RULE_INVALID_CODE_CYCLE)],
        )
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        assert "Invalid California Code Cycle Years" in text
        assert "2018 CBC" in text

    def test_export_renders_duplicate_paragraph_section(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult(
            duplicate_paragraph_alerts=[_alert("Provide cut sheets for all submittals.", DETERMINISTIC_RULE_DUPLICATE_PARAGRAPH)],
        )
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        assert "Duplicate Paragraphs" in text
        assert "Provide cut sheets" in text

    def test_export_renders_structural_section(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult(
            structural_alerts=[_alert("1.02 EMPTY SECTION", DETERMINISTIC_RULE_EMPTY_SECTION)],
        )
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        assert "Structural Issues" in text
        assert "1.02 EMPTY SECTION" in text

    def test_export_renders_stale_cycle_section(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult(
            code_cycle_alerts=[_alert("2019 CBC", DETERMINISTIC_RULE_STALE_CODE_CYCLE)],
        )
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        assert "Stale California Code Cycle References" in text
        assert "2019 CBC" in text

    def test_export_renders_naming_section(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult(
            naming_alerts=[_alert("23-23-13 - Refrigerant Piping.docx", DETERMINISTIC_RULE_INCONSISTENT_FILENAME, filename="23-23-13 - Refrigerant Piping.docx")],
        )
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        assert "Inconsistent Filenames" in text
        assert "23-23-13" in text

    def test_export_skips_alerts_heading_when_no_alerts(self, tmp_path: Path) -> None:
        from src.report_exporter import export_report

        result = _StubPipelineResult()  # every list empty
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)
        # No alerts → no top-level "Alerts" heading, no deterministic-check
        # banner. (Other report sections may still render.)
        assert "(deterministic check)" not in text


# ---------------------------------------------------------------------------
# Resume-state round-trip — new alert lists survive serialization
# ---------------------------------------------------------------------------

class TestResumeStateChunkO:
    def test_submission_round_trips_chunk_o_alerts(self) -> None:
        from src.batch import BatchJob
        from src.pipeline import BatchSubmission
        from src.resume_state import deserialize_submission, serialize_submission

        sub = BatchSubmission(
            job=BatchJob(batch_id="msgbatch_test", job_type="review", request_map={}, created_at=0.0),
            template_marker_alerts=[_alert("TODO: x", DETERMINISTIC_RULE_TEMPLATE_MARKER)],
            invalid_code_cycle_alerts=[_alert("2018 CBC", DETERMINISTIC_RULE_INVALID_CODE_CYCLE)],
            duplicate_paragraph_alerts=[_alert("p" * 100, DETERMINISTIC_RULE_DUPLICATE_PARAGRAPH)],
            code_cycle_alerts=[_alert("2019 CBC", DETERMINISTIC_RULE_STALE_CODE_CYCLE)],
            structural_alerts=[_alert("1.02 EMPTY", DETERMINISTIC_RULE_EMPTY_SECTION)],
            naming_alerts=[_alert("23-23-13.docx", DETERMINISTIC_RULE_INCONSISTENT_FILENAME)],
        )
        restored = deserialize_submission(serialize_submission(sub))
        assert restored.template_marker_alerts == sub.template_marker_alerts
        assert restored.invalid_code_cycle_alerts == sub.invalid_code_cycle_alerts
        assert restored.duplicate_paragraph_alerts == sub.duplicate_paragraph_alerts
        assert restored.code_cycle_alerts == sub.code_cycle_alerts
        assert restored.structural_alerts == sub.structural_alerts
        assert restored.naming_alerts == sub.naming_alerts

    def test_legacy_submission_payload_loads_with_empty_chunk_o_lists(self) -> None:
        """Older resume-state JSON omits the chunk O keys."""
        from src.resume_state import deserialize_submission

        legacy = {
            "job": {
                "batch_id": "msgbatch_test",
                "job_type": "review",
                "request_map": {},
                "created_at": 0.0,
            },
            "files_reviewed": [],
            "review_request_ids": [],
            "leed_alerts": [],
            "placeholder_alerts": [],
            "model": "claude-opus-4-7",
            "project_context": "",
            "code_cycle": "2025",
            "cross_check_enabled": False,
            "review_mode": "comprehensive",
            "prepared_specs": [],
        }
        restored = deserialize_submission(legacy)
        # Backward compatibility: new fields default to empty lists.
        assert restored.template_marker_alerts == []
        assert restored.invalid_code_cycle_alerts == []
        assert restored.duplicate_paragraph_alerts == []
        assert restored.code_cycle_alerts == []
        assert restored.structural_alerts == []
        assert restored.naming_alerts == []

    def test_collected_state_round_trips_chunk_o_alerts(self) -> None:
        from src.batch import BatchJob
        from src.pipeline import BatchSubmission, CollectedBatchState
        from src.resume_state import (
            deserialize_collected_batch_state,
            serialize_collected_batch_state,
        )
        from src.review.reviewer import ReviewResult

        sub = BatchSubmission(
            job=BatchJob(batch_id="msgbatch_test", job_type="review", request_map={}, created_at=0.0),
        )
        state = CollectedBatchState(
            submission=sub,
            review_result=ReviewResult(findings=[]),
            template_marker_alerts=[_alert("TODO: x", DETERMINISTIC_RULE_TEMPLATE_MARKER)],
            invalid_code_cycle_alerts=[_alert("2018 CBC", DETERMINISTIC_RULE_INVALID_CODE_CYCLE)],
            duplicate_paragraph_alerts=[_alert("p" * 100, DETERMINISTIC_RULE_DUPLICATE_PARAGRAPH)],
        )
        restored = deserialize_collected_batch_state(
            serialize_collected_batch_state(state), submission=sub
        )
        assert restored.template_marker_alerts == state.template_marker_alerts
        assert restored.invalid_code_cycle_alerts == state.invalid_code_cycle_alerts
        assert restored.duplicate_paragraph_alerts == state.duplicate_paragraph_alerts
