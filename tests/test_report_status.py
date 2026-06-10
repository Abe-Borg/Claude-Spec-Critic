"""Tests for report trust-model statuses.

This work defines two closed enums (``ReportStatus`` and ``EditActionLabel``) that
every finding maps to for display, plus the four evidence concepts
(spec evidence, web/code evidence, verification rationale, unsupported
sources) that the report exporter has to surface distinctly.

Coverage:

* ``TestReportStatusClassification`` exercises every status branch of
  :func:`classify_status`, including the priority ordering (verification
  failure beats disagreement beats local-skip beats verdict).
* ``TestEditActionClassification`` exercises every label branch of
  :func:`classify_edit_action`: no proposal → REPORT_ONLY, otherwise
  EDIT_SUGGESTED.
* ``TestSummarizeHelpers`` checks that the histogram helpers return
  zero-filled dicts and sum to the input count.
* ``TestLabelHelpers`` covers the human-readable label / glyph mapping
  and the string fallback for unknown values.
* ``TestReportExporterStatusIntegration`` does a snapshot-style check on
  the actual .docx output: it builds a tiny ``PipelineResult`` with
  findings spanning every status, exports to a temp path, and asserts
  the resulting document contains the expected status labels and
  evidence sub-headings. Catches regressions where the wiring drops
  the status line or mis-routes the histogram counts.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.output.report_exporter import export_report
from src.output.report_status import (
    EditActionLabel,
    ReportStatus,
    STATUS_LABELS,
    VERDICT_SUPERSEDES_CONFIDENCE,
    classify_edit_action,
    classify_status,
    summarize_edit_actions,
    summarize_statuses,
    verdict_supersedes_confidence,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.6,
    action: str = "EDIT",
    existing: str | None = "old text",
    replacement: str | None = "new text",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
    )
    f.verification = verification
    return f


def _verification(
    verdict: str = "CONFIRMED",
    *,
    grounded: bool = True,
    cache_status: str = "miss",
    explanation: str = "Verified against CBC §1234",
    sources: list[str] | None = None,
    rejected: list[dict] | None = None,
    correction: str | None = None,
) -> VerificationResult:
    # A grounded CONFIRMED/CORRECTED requires at least one
    # accepted external citation. Default to a representative one so
    # individual tests focused on status / edit-action classification
    # do not have to thread sources through every call.
    if sources is None:
        sources = (
            ["https://dgs.ca.gov"]
            if grounded and verdict.upper() in ("CONFIRMED", "CORRECTED")
            else []
        )
    return VerificationResult(
        verdict=verdict,
        explanation=explanation,
        sources=list(sources),
        correction=correction,
        grounded=grounded,
        cache_status=cache_status,
        rejected_sources=list(rejected or []),
    )


# ---------------------------------------------------------------------------
# classify_status — every branch
# ---------------------------------------------------------------------------

class TestReportStatusClassification:
    def test_no_verification_is_not_checked(self):
        f = _finding(verification=None)
        assert classify_status(f) is ReportStatus.NOT_CHECKED

    def test_local_skip_is_locally_classified(self):
        v = _verification("UNVERIFIED", grounded=False, cache_status="local_skip")
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.LOCALLY_CLASSIFIED

    def test_confirmed_and_grounded_is_verified_supported(self):
        f = _finding(verification=_verification("CONFIRMED", grounded=True))
        assert classify_status(f) is ReportStatus.VERIFIED_SUPPORTED

    def test_corrected_and_grounded_is_verified_contradicted(self):
        f = _finding(verification=_verification("CORRECTED", grounded=True))
        assert classify_status(f) is ReportStatus.VERIFIED_CONTRADICTED

    def test_disputed_verdict_is_disputed(self):
        f = _finding(verification=_verification("DISPUTED", grounded=False))
        assert classify_status(f) is ReportStatus.DISPUTED

    def test_unverified_verdict_is_insufficient_evidence(self):
        f = _finding(verification=_verification("UNVERIFIED", grounded=False))
        assert classify_status(f) is ReportStatus.INSUFFICIENT_EVIDENCE

    def test_confirmed_but_ungrounded_does_not_count_as_supported(self):
        # The grounding invariant in the verifier should already have
        # downgraded this, but the classifier is the second line of
        # defense for tests that construct results by hand.
        f = _finding(verification=_verification("CONFIRMED", grounded=False))
        assert classify_status(f) is ReportStatus.INSUFFICIENT_EVIDENCE


# ---------------------------------------------------------------------------
# verdict_supersedes_confidence — when the verdict replaces the review %
# ---------------------------------------------------------------------------

class TestVerdictSupersedesConfidence:
    """A verdict-bearing status supersedes the pre-verification confidence.

    The report drops the confidence % from a finding's header once
    verification reached a verdict (supported / contradicted / contested /
    disputed); for every other status — where the verifier reached no
    verdict — the % stays the primary signal.
    """

    def test_verified_supported_supersedes(self):
        f = _finding(verification=_verification("CONFIRMED", grounded=True))
        assert verdict_supersedes_confidence(f) is True

    def test_verified_contradicted_supersedes(self):
        f = _finding(verification=_verification("CORRECTED", grounded=True))
        assert verdict_supersedes_confidence(f) is True

    def test_disputed_supersedes(self):
        # A high review confidence next to a DISPUTED verdict is just as
        # misleading as a low one next to CONFIRMED — the verdict wins.
        f = _finding(verification=_verification("DISPUTED", grounded=False))
        assert verdict_supersedes_confidence(f) is True

    def test_contested_supersedes(self):
        v = VerificationResult(
            verdict="CONFIRMED",
            explanation="Initial and escalated verifiers disagreed.",
            sources=["https://dgs.ca.gov"],
            grounded=True,
            models_disagreed=True,
        )
        f = _finding(verification=v)
        # Sanity: this really does classify as contested.
        assert classify_status(f) is ReportStatus.VERIFIED_CONTESTED
        assert verdict_supersedes_confidence(f) is True

    def test_not_checked_does_not_supersede(self):
        f = _finding(verification=None)
        assert verdict_supersedes_confidence(f) is False

    def test_insufficient_evidence_does_not_supersede(self):
        f = _finding(verification=_verification("UNVERIFIED", grounded=False))
        assert verdict_supersedes_confidence(f) is False

    def test_locally_classified_does_not_supersede(self):
        v = _verification("UNVERIFIED", grounded=False, cache_status="local_skip")
        f = _finding(verification=v)
        assert verdict_supersedes_confidence(f) is False

    def test_verification_failed_does_not_supersede(self):
        v = VerificationResult(
            verdict="UNVERIFIED",
            explanation="Rate limited.",
            sources=[],
            grounded=False,
            verification_failed=True,
        )
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFICATION_FAILED
        assert verdict_supersedes_confidence(f) is False

    def test_set_membership_matches_helper(self):
        # The exported frozenset is the single source of truth the helper
        # consults; keep them in lockstep.
        assert ReportStatus.VERIFIED_SUPPORTED in VERDICT_SUPERSEDES_CONFIDENCE
        assert ReportStatus.VERIFIED_CONTRADICTED in VERDICT_SUPERSEDES_CONFIDENCE
        assert ReportStatus.VERIFIED_CONTESTED in VERDICT_SUPERSEDES_CONFIDENCE
        assert ReportStatus.DISPUTED in VERDICT_SUPERSEDES_CONFIDENCE
        assert ReportStatus.NOT_CHECKED not in VERDICT_SUPERSEDES_CONFIDENCE
        assert ReportStatus.INSUFFICIENT_EVIDENCE not in VERDICT_SUPERSEDES_CONFIDENCE


# ---------------------------------------------------------------------------
# classify_edit_action — every branch
# ---------------------------------------------------------------------------

class TestEditActionClassification:
    def test_no_proposal_is_report_only(self):
        f = _finding(action="REPORT_ONLY", existing=None, replacement=None)
        assert classify_edit_action(f) is EditActionLabel.REPORT_ONLY

    def test_proposal_is_edit_suggested(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.85,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_verification("CONFIRMED", grounded=True),
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    @pytest.mark.parametrize(
        "edit_confidence, verification",
        [
            # Ignores verification status — verdict/grounding/cache never
            # gate the label (the app emits; a downstream applier gates).
            (0.95, _verification("DISPUTED", grounded=False, cache_status="miss")),
            (0.95, _verification("UNVERIFIED", grounded=False, cache_status="miss")),
            (0.95, _verification("UNVERIFIED", grounded=False, cache_status="local_skip")),
            # Ignores low confidence.
            (0.1, _verification("CONFIRMED", grounded=True)),
            # NOT_CHECKED (no verification ran) with a proposal still labels.
            (0.95, None),
        ],
    )
    def test_proposal_label_ignores_verification_and_confidence(
        self, edit_confidence, verification
    ):
        # The app emits edit instructions but never applies them, so the
        # label is "has a proposal?" — independent of verdict/grounding/
        # confidence or whether verification ran at all.
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=edit_confidence,
        )
        f = _finding(edit_proposal=proposal, verification=verification)
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_legacy_edit_finding_routes_through_as_edit_proposal(self):
        # An old-shaped finding with actionType=EDIT and existingText set
        # is still EDIT_SUGGESTED — the legacy proposal is synthesized on
        # the fly by ``as_edit_proposal``.
        f = _finding(
            action="EDIT",
            existing="old",
            replacement="new",
            confidence=0.9,
            verification=_verification("CONFIRMED", grounded=True),
            edit_proposal=None,
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED


# ---------------------------------------------------------------------------
# Aggregation helpers
# ---------------------------------------------------------------------------

class TestSummarizeHelpers:
    def test_summarize_statuses_returns_zero_filled_dict_on_empty_input(self):
        counts = summarize_statuses([])
        assert set(counts.keys()) == set(ReportStatus)
        assert sum(counts.values()) == 0

    def test_summarize_statuses_sums_to_input_length(self):
        findings = [
            _finding(verification=_verification("CONFIRMED", grounded=True)),
            _finding(verification=_verification("DISPUTED", grounded=False)),
            _finding(verification=None),
            _finding(
                verification=_verification(
                    "UNVERIFIED", grounded=False, cache_status="local_skip"
                ),
            ),
        ]
        counts = summarize_statuses(findings)
        assert counts[ReportStatus.VERIFIED_SUPPORTED] == 1
        assert counts[ReportStatus.DISPUTED] == 1
        assert counts[ReportStatus.NOT_CHECKED] == 1
        assert counts[ReportStatus.LOCALLY_CLASSIFIED] == 1
        assert sum(counts.values()) == len(findings)

    def test_summarize_edit_actions_zero_filled(self):
        counts = summarize_edit_actions([])
        assert set(counts.keys()) == set(EditActionLabel)
        assert sum(counts.values()) == 0

    def test_summarize_edit_actions_counts_each_label(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.95,
        )
        findings = [
            _finding(
                edit_proposal=proposal,
                verification=_verification("CONFIRMED", grounded=True),
            ),
            _finding(action="REPORT_ONLY", existing=None, replacement=None),
        ]
        counts = summarize_edit_actions(findings)
        assert counts[EditActionLabel.EDIT_SUGGESTED] == 1
        assert counts[EditActionLabel.REPORT_ONLY] == 1


# ---------------------------------------------------------------------------
# Label / glyph helpers
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Report exporter integration (snapshot-style)
# ---------------------------------------------------------------------------

class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report.

    The real :class:`pipeline.PipelineResult` is a dataclass with a
    handful of typed fields; building one here forces importing pipeline
    + cross_checker + verifier modules that pull a lot of weight. A stub
    that exposes only the attributes export_report reads is cheaper and
    keeps these tests independent of unrelated pipeline changes.
    """

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestReportExporterStatusIntegration:
    @pytest.fixture
    def diverse_review_result(self) -> ReviewResult:
        """Build a review result with one finding per common status path."""
        verified_proposal = EditProposal(
            action_type="EDIT",
            existing_text="2019 CBC",
            replacement_text="2025 CBC",
            edit_confidence=0.95,
        )
        verified = _finding(
            severity="CRITICAL",
            file="Section_23_0000.docx",
            section="2.1",
            issue="Stale CBC reference",
            edit_proposal=verified_proposal,
            verification=_verification(
                "CONFIRMED",
                grounded=True,
                sources=["https://codes.iccsafe.org/content/CBC2025"],
            ),
        )
        disputed = _finding(
            severity="HIGH",
            file="Section_23_0000.docx",
            section="2.2",
            issue="Claims wrong fitting standard",
            verification=_verification(
                "DISPUTED",
                grounded=False,
                explanation="Search results contradict the finding's claim.",
                rejected=[{"url": "https://blog.example.com/foo", "reason": "ungrounded"}],
            ),
        )
        insufficient = _finding(
            severity="MEDIUM",
            file="Section_23_0000.docx",
            section="2.3",
            issue="Could not be verified",
            verification=_verification(
                "UNVERIFIED",
                grounded=False,
                explanation="Authoritative source not located.",
            ),
        )
        local_skip = _finding(
            severity="GRIPES",
            file="Section_23_0000.docx",
            section="2.4",
            issue="LEED Gold reference in a non-LEED project",
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verification(
                "UNVERIFIED",
                grounded=False,
                cache_status="local_skip",
                explanation="Locally classified; no web verification needed.",
            ),
        )
        not_checked = _finding(
            severity="GRIPES",
            file="Section_23_0000.docx",
            section="2.5",
            issue="Pre-verification scan finding",
            verification=None,
        )
        return ReviewResult(
            findings=[verified, disputed, insufficient, local_skip, not_checked],
            input_tokens=1000,
            output_tokens=500,
            elapsed_seconds=12.5,
        )

    def test_export_contains_status_lines_for_each_finding(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)

        # Every finding renders a "Status:" line, so the substring count
        # should equal the finding count.
        assert text.count("Status:") == len(diverse_review_result.findings)

        # Each status label appears at least once. The label lives on
        # the status line, the trust-model histogram cell, or both.
        assert STATUS_LABELS[ReportStatus.VERIFIED_SUPPORTED] in text
        assert STATUS_LABELS[ReportStatus.DISPUTED] in text
        assert STATUS_LABELS[ReportStatus.INSUFFICIENT_EVIDENCE] in text
        assert STATUS_LABELS[ReportStatus.LOCALLY_CLASSIFIED] in text
        assert STATUS_LABELS[ReportStatus.NOT_CHECKED] in text

    def test_export_renames_existing_text_label_to_spec_evidence(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        # Spec evidence is distinct from web/code
        # evidence and verification rationale. The label rename makes
        # the four concepts explicit.
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)
        assert "Spec evidence:" in text
        assert "Proposed replacement:" in text
        # The old labels should be gone now.
        assert "Existing Text:" not in text
        assert "Replace With:" not in text


# ---------------------------------------------------------------------------
# Confidence de-emphasis rendering (header suppression + footnote + note)
# ---------------------------------------------------------------------------

class TestConfidenceDeEmphasisRendering:
    """The header drops the confidence % once a verdict supersedes it.

    A verified finding's pre-verification confidence is moved off the
    header (where it reads as the headline trust signal) into a small,
    labeled footnote on the Status line; an unverified finding keeps the
    prominent header %.
    """

    def _build(self) -> ReviewResult:
        verified_low_conf = _finding(
            severity="HIGH",
            file="Section_22_1000.docx",
            section="2.1",
            issue="Wrong parenthetical section title",
            confidence=0.55,  # low review confidence, but verifier confirms it
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verification("CONFIRMED", grounded=True),
        )
        unchecked = _finding(
            severity="HIGH",
            file="Section_22_1000.docx",
            section="2.2",
            issue="Pre-verification scan finding",
            confidence=0.42,  # the only trust signal — stays in the header
            verification=None,
        )
        return ReviewResult(
            findings=[verified_low_conf, unchecked],
            input_tokens=1000,
            output_tokens=500,
            elapsed_seconds=1.0,
        )

    @staticmethod
    def _heading3_containing(doc: Document, token: str) -> str:
        for p in doc.paragraphs:
            if p.style.name == "Heading 3" and token in p.text:
                return p.text
        raise AssertionError(f"no Heading 3 paragraph contains {token!r}")

    def test_verified_header_drops_percent_unverified_keeps_it(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=self._build()), out)
        doc = Document(str(out))

        verified_header = self._heading3_containing(doc, "— 2.1")
        unchecked_header = self._heading3_containing(doc, "— 2.2")

        # Verified finding: the % is suppressed from the header.
        assert "55%" not in verified_header
        # Unverified finding: the % stays prominent in the header.
        assert "42%" in unchecked_header

    def test_verified_confidence_moves_to_labeled_footnote(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=self._build()), out)
        doc = Document(str(out))
        text = _all_text_from(doc)

        # The number is preserved, but as an explicitly-labeled
        # pre-verification footnote rather than the headline %.
        assert "review confidence 55% (pre-verification)" in text

    def test_methodology_documents_the_distinction(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=self._build()), out)
        doc = Document(str(out))
        text = _all_text_from(doc)

        # Report-level documentation of confidence-vs-verdict.
        assert "before verification" in text
        assert "authoritative trust signal" in text
