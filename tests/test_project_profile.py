"""Unit pins for :class:`ProjectProfile` and the module capability flag.

WS-2a of ``docs/hyperscale_datacenter_module_plan.md``. Hermetic — no API key,
no network, no tkinter.
"""
from __future__ import annotations

import dataclasses

import pytest

from src.core.project_profile import (
    CA_PROVINCES,
    COUNTRY_DISPLAY,
    US_STATES,
    ProjectProfile,
    normalize_country,
    states_for_country,
)


def _ashburn() -> ProjectProfile:
    return ProjectProfile(
        city="Ashburn", state_or_province="VA", country="US", client_name="ExampleCo"
    )


class TestNormalization:
    @pytest.mark.parametrize(
        "raw, expected",
        [
            ("US", "US"), ("usa", "US"), ("USA", "US"), ("United States", "US"),
            ("america", "US"), ("CA", "CA"), ("Canada", "CA"), ("canada", "CA"),
            ("  Canada  ", "CA"), ("Mexico", ""), ("", ""),
        ],
    )
    def test_normalize_country(self, raw, expected):
        assert normalize_country(raw) == expected

    def test_fields_are_trimmed_and_country_folded(self):
        p = ProjectProfile(
            city="  Markham ", state_or_province=" ON ", country=" Canada ",
            client_name="  ExampleCo ",
        )
        assert p.city == "Markham"
        assert p.state_or_province == "ON"
        assert p.country == "CA"
        assert p.client_name == "ExampleCo"

    def test_unknown_country_is_preserved_but_incomplete(self):
        p = ProjectProfile(city="X", state_or_province="ZZ", country="Narnia", client_name="C")
        # Unrecognized country is kept verbatim (trimmed) but fails completeness.
        assert p.country == "Narnia"
        assert p.is_complete() is False


class TestDisplayForms:
    def test_display_line(self):
        assert _ashburn().display_line() == "Ashburn, Virginia, USA — Client: ExampleCo"

    def test_project_meta_lines(self):
        assert _ashburn().project_meta_lines() == [
            "Project: Ashburn, Virginia, USA",
            "Client: ExampleCo",
        ]

    def test_country_and_state_display_fallbacks(self):
        # Unknown state code falls back to the code itself.
        p = ProjectProfile(city="X", state_or_province="ZZ", country="US", client_name="C")
        assert p.state_display == "ZZ"
        assert p.country_display == "USA"

    def test_canadian_province_display(self):
        p = ProjectProfile(city="Markham", state_or_province="ON", country="CA", client_name="C")
        assert p.state_display == "Ontario"
        assert p.country_display == "Canada"
        assert p.display_line() == "Markham, Ontario, Canada — Client: C"


class TestUserLocation:
    def test_shape_matches_engine_default(self):
        # Mirrors the engine's hardcoded {"country": "US", "region": "California"}
        # shape: ISO country code + full region name + city.
        loc = _ashburn().web_search_user_location()
        assert loc == {
            "type": "approximate",
            "country": "US",
            "region": "Virginia",
            "city": "Ashburn",
        }

    def test_canadian_user_location(self):
        loc = ProjectProfile(
            city="Markham", state_or_province="ON", country="CA", client_name="C"
        ).web_search_user_location()
        assert loc["country"] == "CA"
        assert loc["region"] == "Ontario"


class TestFingerprint:
    def test_is_16_hex(self):
        fp = _ashburn().jurisdiction_fingerprint()
        assert len(fp) == 16
        int(fp, 16)  # hex

    def test_stable_across_case_and_whitespace(self):
        a = ProjectProfile(city="Markham", state_or_province="ON", country="CA", client_name="C1")
        b = ProjectProfile(city=" markham ", state_or_province="ON", country="Canada", client_name="C2")
        # Client differs but the jurisdiction (country|state|city) is the same.
        assert a.jurisdiction_fingerprint() == b.jurisdiction_fingerprint()

    def test_different_city_differs(self):
        a = ProjectProfile(city="Markham", state_or_province="ON", country="CA", client_name="C")
        b = ProjectProfile(city="Toronto", state_or_province="ON", country="CA", client_name="C")
        assert a.jurisdiction_fingerprint() != b.jurisdiction_fingerprint()

    def test_typo_city_misroutes_deliberately(self):
        # The plan's field example: "Marham" vs "Markham" must fingerprint
        # differently (the reason input echo-back matters).
        good = ProjectProfile(city="Markham", state_or_province="ON", country="CA", client_name="C")
        typo = ProjectProfile(city="Marham", state_or_province="ON", country="CA", client_name="C")
        assert good.jurisdiction_fingerprint() != typo.jurisdiction_fingerprint()


class TestCompleteness:
    def test_complete(self):
        assert _ashburn().is_complete() is True

    @pytest.mark.parametrize(
        "overrides",
        [
            dict(city=""), dict(state_or_province=""), dict(client_name=""),
            dict(country=""), dict(country="Mexico"),
        ],
    )
    def test_incomplete(self, overrides):
        base = dict(city="Ashburn", state_or_province="VA", country="US", client_name="ExampleCo")
        base.update(overrides)
        assert ProjectProfile(**base).is_complete() is False


class TestSerialization:
    def test_round_trip(self):
        p = _ashburn()
        assert ProjectProfile.from_dict(p.to_dict()) == p

    def test_from_dict_defensive(self):
        assert ProjectProfile.from_dict(None) is None
        assert ProjectProfile.from_dict("nope") is None
        assert ProjectProfile.from_dict({}) is None
        # All-empty payload degrades to None (profile-less), not a hollow object.
        assert ProjectProfile.from_dict(
            {"city": "", "state_or_province": "", "country": "", "client_name": ""}
        ) is None

    def test_from_dict_normalizes(self):
        p = ProjectProfile.from_dict(
            {"city": " Markham ", "state_or_province": "ON", "country": "Canada", "client_name": "C"}
        )
        assert p is not None
        assert p.country == "CA"
        assert p.city == "Markham"

    def test_from_dict_partial_is_not_none_but_incomplete(self):
        p = ProjectProfile.from_dict({"city": "Ashburn"})
        assert p is not None
        assert p.is_complete() is False


class TestStateTables:
    def test_counts(self):
        assert len(US_STATES) == 51  # 50 states + DC
        assert len(CA_PROVINCES) == 13
        assert set(COUNTRY_DISPLAY) == {"US", "CA"}

    def test_states_for_country(self):
        assert states_for_country("US") is US_STATES
        assert states_for_country("Canada") is CA_PROVINCES
        assert states_for_country("Mexico") == {}


class TestProfilePersistence:
    """Profile threads submission -> pending state -> submission, additively."""

    def _submission(self, **overrides):
        from src.batch.batch import BatchJob
        from src.orchestration.pipeline import BatchSubmission

        job = BatchJob(
            batch_id="msgbatch_PROF",
            job_type="review",
            request_map={"review__a__0": {"filename": "a.docx", "index": 0, "type": "review"}},
            created_at=1700000000.0,
        )
        base = dict(
            job=job,
            files_reviewed=["a.docx"],
            review_request_ids=["review__a__0"],
            model="claude-opus-4-8",
        )
        base.update(overrides)
        return BatchSubmission(**base)

    def test_pending_round_trip_with_profile(self, tmp_path):
        from src.orchestration.batch_resume import (
            PendingBatch,
            load_pending_batch,
            save_pending_batch,
        )

        prof = _ashburn().to_dict()
        pending = PendingBatch.from_submission(self._submission(project_profile=prof))
        assert pending.project_profile == prof

        path = tmp_path / "pb.json"
        save_pending_batch(pending, path=path)
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.project_profile == prof
        # to_submission reconstructs the profile onto the rebuilt submission.
        sub = loaded.to_submission(log=lambda *a, **k: None)
        assert sub.project_profile == prof

    def test_pending_round_trip_without_profile(self, tmp_path):
        from src.orchestration.batch_resume import (
            PendingBatch,
            load_pending_batch,
            save_pending_batch,
        )

        pending = PendingBatch.from_submission(self._submission())
        assert pending.project_profile is None
        path = tmp_path / "pb.json"
        save_pending_batch(pending, path=path)
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.project_profile is None

    def test_legacy_state_file_without_profile_key_defaults_none(self, tmp_path):
        import json

        from src.orchestration.batch_resume import load_pending_batch

        # A pending_batch.json written before WS-2 has no project_profile key;
        # it must load (additive field, no schema bump) with the profile None.
        legacy = {
            "batch_id": "msgbatch_LEGACY",
            "model": "claude-opus-4-8",
            "request_map": {},
            "review_request_ids": [],
            "files_reviewed": [],
            "input_dir": "",
            "files": [],
            "cycle_label": "2025",
            "module_id": "california_k12_mep",
            "project_context": "",
            "cross_check_enabled": False,
            "submitted_at": 1700000000.0,
            "run_id": "",
            "app_version": "3.0.0",
            "schema_version": 1,
        }
        path = tmp_path / "pb.json"
        path.write_text(json.dumps(legacy), encoding="utf-8")
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.project_profile is None

    def test_new_state_file_keeps_schema_version_1(self, tmp_path):
        import json

        from src.orchestration.batch_resume import (
            PendingBatch,
            save_pending_batch,
        )

        # Additive field, defensive loader: no schema bump.
        path = tmp_path / "pb.json"
        save_pending_batch(
            PendingBatch.from_submission(self._submission(project_profile=_ashburn().to_dict())),
            path=path,
        )
        data = json.loads(path.read_text(encoding="utf-8"))
        assert data["schema_version"] == 1
        assert data["project_profile"]["city"] == "Ashburn"

    def test_finalize_stamps_profile_onto_pipeline_result(self):
        from src.orchestration.pipeline import (
            CollectedBatchState,
            finalize_batch_result,
        )
        from src.review.reviewer import ReviewResult

        prof = _ashburn().to_dict()
        state = CollectedBatchState(
            submission=self._submission(project_profile=prof),
            review_result=ReviewResult(findings=[]),
        )
        assert finalize_batch_result(state).project_profile == prof

    def test_finalize_profile_none_on_profileless_run(self):
        from src.orchestration.pipeline import (
            CollectedBatchState,
            finalize_batch_result,
        )
        from src.review.reviewer import ReviewResult

        state = CollectedBatchState(
            submission=self._submission(),
            review_result=ReviewResult(findings=[]),
        )
        assert finalize_batch_result(state).project_profile is None


class TestReportTitleBlock:
    def test_title_block_appends_project_and_client_lines(self):
        from docx import Document

        from src.output.report_exporter import _write_title_block

        class _StubReview:
            model = "claude-opus-4-8"

        doc = Document()
        _write_title_block(
            doc, _StubReview(), ["a.docx"], cycle_label="dc-ibc-2024",
            profile=_ashburn(),
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Project: Ashburn, Virginia, USA" in text
        assert "Client: ExampleCo" in text

    def test_title_block_without_profile_has_no_project_lines(self):
        from docx import Document

        from src.output.report_exporter import _write_title_block

        class _StubReview:
            model = "claude-opus-4-8"

        doc = Document()
        _write_title_block(doc, _StubReview(), ["a.docx"], cycle_label="2025")
        text = "\n".join(p.text for p in doc.paragraphs)
        # Profile-less run: no Project:/Client: lines at all (byte-identical
        # to a pre-WS-2 title block).
        assert "Project:" not in text
        assert "Client:" not in text


class TestDiagnosticsProfileSummary:
    def test_default_empty(self):
        from src.orchestration.diagnostics import DiagnosticsReport

        assert DiagnosticsReport().project_profile_summary == ""

    def test_carries_summary(self):
        from src.orchestration.diagnostics import DiagnosticsReport

        summary = _ashburn().display_line()
        assert (
            DiagnosticsReport(project_profile_summary=summary).project_profile_summary
            == summary
        )


class TestModuleCapabilityFlag:
    def test_default_off_for_existing_modules(self):
        from src.modules import CALIFORNIA_K12_MEP, DATACENTER_FIRE

        # Both current modules keep the flag off (no location-aware behavior
        # until a later workstream flips it for the DC module).
        assert CALIFORNIA_K12_MEP.project_profile_enabled is False
        assert DATACENTER_FIRE.project_profile_enabled is False

    def test_flag_is_additive_and_validates(self):
        from src.modules import CALIFORNIA_K12_MEP, validate_module_registry

        # Flipping the flag on a module must not break registry validation in
        # WS-2 (the conditional content-slot rules arrive with the slots).
        enabled = dataclasses.replace(CALIFORNIA_K12_MEP, project_profile_enabled=True)
        validate_module_registry([enabled])
