"""Chunk 12 tests — VERIFIED_CONTESTED status + models_disagreed sentinel.

Chunk 12 of the Trust Upgrade surfaces escalation disagreements as a
dedicated trust-model status. When the initial Sonnet verifier and the
escalated Opus verifier reach different verdicts on a finding AND both
verdicts were grounded (each had at least one accepted citation), the
finding renders as VERIFIED_CONTESTED rather than under the regular
verdict-based classification.

The contract has these surfaces:

* ``ReportStatus.VERIFIED_CONTESTED`` exists with a registered display
  label ("Verified — but models disagreed"), glyph (⚡), color
  (purple), and shading entry. It appears in ``STATUS_DISPLAY_ORDER``
  between VERIFIED_CONTRADICTED and LOCALLY_CLASSIFIED.
* ``VerificationResult.models_disagreed`` defaults to False and
  round-trips through serialize/deserialize (resume state) and through
  the cache persist path / clone helpers (so cache replays preserve
  the contested status).
* ``VerificationResult.initial_sources`` defaults to ``[]`` and
  round-trips the same way as ``models_disagreed``. Populated during
  escalation with the initial verifier's accepted citations so the
  evidence panel can render both citation sets side-by-side.
* ``classify_status`` returns VERIFIED_CONTESTED when
  ``models_disagreed`` is True, prioritized BEFORE the verdict-based
  branches so a final CONFIRMED+grounded verdict that disagreed with a
  DISPUTED initial does not slip through as VERIFIED_SUPPORTED.
* ``classify_edit_action`` labels a VERIFIED_CONTESTED finding
  EDIT_SUGGESTED when it carries a proposal and REPORT_ONLY otherwise;
  the contested status rides along for a downstream applier to act on.
* ``verify_finding`` sets ``models_disagreed`` only when BOTH passes
  were grounded AND the verdicts differ. Initial-UNVERIFIED-then-
  CONFIRMED escalations do NOT trigger it — the initial pass was not
  grounded, so the disagreement is "Sonnet found nothing, Opus
  did", not "two verifiers reading the same sources disagreed".
* The exported report's evidence panel renders the disagreement
  sentence inline ("manual review recommended") and adds an "Initial
  verifier sources:" sub-section listing the initial verifier's URLs
  for findings with ``models_disagreed=True``.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from src.core.code_cycles import DEFAULT_CYCLE
from src.output.report_exporter import STATUS_COLORS, STATUS_SHADING, export_report
from src.output.report_status import (
    EditActionLabel,
    ReportStatus,
    STATUS_DISPLAY_ORDER,
    STATUS_GLYPHS,
    STATUS_LABELS,
    classify_edit_action,
    classify_status,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verification_cache import (
    VerificationCache,
    _clone_for_hit,
    _clone_for_store,
    _CacheEntry,
    _result_to_dict,
)
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.6,
    action: str = "EDIT",
    existing: str | None = "old text",
    replacement: str | None = "new text",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
    )
    f.verification = verification
    return f


def _contested_verification(
    *,
    final_verdict: str = "CONFIRMED",
    final_model: str = "claude-opus-4-7",
    initial_verdict: str = "DISPUTED",
    initial_model: str = "claude-sonnet-4-6",
    final_sources: list[str] | None = None,
    initial_sources: list[str] | None = None,
) -> VerificationResult:
    """Build a VerificationResult representing a real model disagreement.

    Both passes grounded (each with accepted citations) but reached
    different verdicts. The final ``verdict`` field is the Opus result
    (CONFIRMED by default) because the escalation logic swapped to it;
    ``initial_verdict`` captures the Sonnet result (DISPUTED by default).
    """
    if final_sources is None:
        final_sources = ["https://codes.iccsafe.org/content/CBC2025"]
    if initial_sources is None:
        initial_sources = ["https://nfpa.org/codes/13/2022"]
    return VerificationResult(
        verdict=final_verdict,
        explanation="Opus verified; Sonnet disagreed.",
        sources=list(final_sources),
        accepted_sources=list(final_sources),
        grounded=True,
        model_used=final_model,
        escalated=True,
        cache_status="miss",
        web_search_requests=5,
        source_quote="Section 4.2 requires a minimum of 6 ft clearance.",
        escalation_attempted=True,
        initial_model=initial_model,
        initial_verdict=initial_verdict,
        escalation_changed_verdict=True,
        escalation_reason="ungrounded_critical_high",
        models_disagreed=True,
        initial_sources=list(initial_sources),
    )


# ---------------------------------------------------------------------------
# 1. ReportStatus enum + display metadata
# ---------------------------------------------------------------------------


class TestVerifiedContestedReportStatus:
    def test_enum_value_exists(self):
        assert ReportStatus.VERIFIED_CONTESTED == "VERIFIED_CONTESTED"

    def test_label_is_registered(self):
        assert ReportStatus.VERIFIED_CONTESTED in STATUS_LABELS
        # The label must explicitly mention disagreement so a quick
        # scroll of the summary table tells the reviewer what the status
        # means without expanding.
        label = STATUS_LABELS[ReportStatus.VERIFIED_CONTESTED]
        assert "disagreed" in label.lower()

    def test_glyph_is_lightning(self):
        # Lightning bolt distinguishes from ⚠ (operational failure) and
        # the verdict glyphs (✓ / ✎ / ✗).
        assert STATUS_GLYPHS[ReportStatus.VERIFIED_CONTESTED] == "⚡"

    def test_glyph_distinct_from_other_statuses(self):
        # Sanity check: no other status uses the same glyph.
        glyphs_minus_contested = {
            g for s, g in STATUS_GLYPHS.items() if s is not ReportStatus.VERIFIED_CONTESTED
        }
        assert "⚡" not in glyphs_minus_contested

    def test_display_order_includes_contested(self):
        assert ReportStatus.VERIFIED_CONTESTED in STATUS_DISPLAY_ORDER

    def test_display_order_places_contested_between_verified_and_uncertain(self):
        # The plan calls for VERIFIED_CONTESTED between VERIFIED_CONTRADICTED
        # and DISPUTED. Concretely: after the supportive verified bucket,
        # before LOCALLY_CLASSIFIED / INSUFFICIENT_EVIDENCE / DISPUTED.
        order = list(STATUS_DISPLAY_ORDER)
        i_contested = order.index(ReportStatus.VERIFIED_CONTESTED)
        i_contradicted = order.index(ReportStatus.VERIFIED_CONTRADICTED)
        i_disputed = order.index(ReportStatus.DISPUTED)
        assert i_contested > i_contradicted
        assert i_contested < i_disputed

    def test_color_is_registered(self):
        assert ReportStatus.VERIFIED_CONTESTED in STATUS_COLORS
        assert ReportStatus.VERIFIED_CONTESTED in STATUS_SHADING

    def test_color_is_purple_and_distinct(self):
        # Plan: purple, distinct from amber/red/green. Spot-check the
        # shading hex value is the purple we picked, and that it differs
        # from the other status colors.
        contested = STATUS_SHADING[ReportStatus.VERIFIED_CONTESTED]
        assert contested == "800080"  # Purple
        for status in ReportStatus:
            if status is ReportStatus.VERIFIED_CONTESTED:
                continue
            assert STATUS_SHADING[status] != contested


# ---------------------------------------------------------------------------
# 2. VerificationResult fields
# ---------------------------------------------------------------------------


class TestVerificationResultContestedFields:
    def test_default_models_disagreed_is_false(self):
        result = VerificationResult(verdict="UNVERIFIED")
        assert result.models_disagreed is False

    def test_default_initial_sources_is_empty_list(self):
        result = VerificationResult(verdict="UNVERIFIED")
        assert result.initial_sources == []
        # Must be a fresh list per instance — not a shared mutable default.
        result.initial_sources.append("https://x")
        other = VerificationResult(verdict="UNVERIFIED")
        assert other.initial_sources == []

    def test_field_round_trips_through_constructor(self):
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://a"],
            models_disagreed=True,
            initial_sources=["https://b", "https://c"],
        )
        assert result.models_disagreed is True
        assert result.initial_sources == ["https://b", "https://c"]


# ---------------------------------------------------------------------------
# 3. classify_status — VERIFIED_CONTESTED branch
# ---------------------------------------------------------------------------


class TestClassifyStatusContested:
    def test_models_disagreed_overrides_confirmed_supported(self):
        # The scenario from the plan: Sonnet says DISPUTED, Opus says
        # CONFIRMED. After the escalation swap, ``result.verdict`` is
        # CONFIRMED and ``result.grounded`` is True. Without the sentinel
        # check the status would render as VERIFIED_SUPPORTED — exactly
        # the bug Chunk 12 fixes.
        f = _finding(verification=_contested_verification())
        assert classify_status(f) is ReportStatus.VERIFIED_CONTESTED

    def test_models_disagreed_overrides_corrected(self):
        v = _contested_verification(final_verdict="CORRECTED")
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFIED_CONTESTED

    def test_models_disagreed_overrides_disputed(self):
        v = _contested_verification(
            final_verdict="DISPUTED",
            initial_verdict="CONFIRMED",
        )
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFIED_CONTESTED

    def test_verification_failed_still_wins_over_disagreement(self):
        # An operational failure beats a model disagreement: if the
        # verifier broke, the reported "disagreement" is unreliable
        # signal at best.
        v = _contested_verification()
        v.verification_failed = True
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFICATION_FAILED

    def test_no_disagreement_flag_falls_through_to_verdict(self):
        # Sanity: a CONFIRMED+grounded verdict without the disagreement
        # flag classifies as VERIFIED_SUPPORTED.
        v = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://x"],
            accepted_sources=["https://x"],
        )
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.VERIFIED_SUPPORTED


# ---------------------------------------------------------------------------
# 4. classify_edit_action — proposal presence drives the label
# ---------------------------------------------------------------------------


class TestClassifyEditActionContested:
    def test_contested_with_proposal_is_edit_suggested(self):
        # The app emits edit instructions but never applies them. A
        # contested finding with a proposal is labeled EDIT_SUGGESTED;
        # the VERIFIED_CONTESTED status rides along in the sidecar so a
        # downstream applier sees the disagreement and can skip it.
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.95,
        )
        f = _finding(
            verification=_contested_verification(),
            edit_proposal=proposal,
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_contested_without_proposal_routes_to_report_only(self):
        # No proposal short-circuits to REPORT_ONLY regardless of status.
        f = _finding(
            verification=_contested_verification(),
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            edit_proposal=None,
        )
        assert classify_edit_action(f) is EditActionLabel.REPORT_ONLY


# ---------------------------------------------------------------------------
# 5. Detection — verifier sets models_disagreed only when both grounded
# ---------------------------------------------------------------------------


class TestModelsDisagreedDetection:
    """Per the plan: the verifier sets ``models_disagreed=True`` only
    when both the initial and escalated passes produced grounded
    verdicts AND those verdicts differ. The "both grounded" condition is
    stricter than ``escalation_changed_verdict`` — an
    initial-UNVERIFIED-then-CONFIRMED escalation should NOT register as
    a disagreement, because the initial pass didn't actually ground
    anything to disagree about.

    We can't easily drive the live verifier (would require an API key
    and network access), but the contract can be enforced via source
    inspection. The escalation block in verify_finding is the single
    place that sets the flag.
    """

    def test_verify_finding_source_uses_strict_both_grounded_condition(self):
        # The merge logic that sets ``models_disagreed`` now lives in the
        # shared ``_apply_escalation_outcome`` helper (reused by the
        # real-time and batch escalation paths so they cannot drift).
        # Assert (a) verify_finding still snapshots the initial grounded
        # flag BEFORE the escalation call and passes it into the helper,
        # and (b) the helper expresses the strict "both grounded AND
        # verdicts differ" condition rather than the weaker "verdicts
        # differ" alone.
        source = Path("src/verification/verifier.py").read_text(encoding="utf-8")
        # The snapshot of the initial grounded flag must happen BEFORE
        # the escalated call (otherwise the swap below would clobber it).
        snapshot_idx = source.find("initial_grounded_snapshot = bool(result.grounded)")
        assert snapshot_idx > 0, (
            "Expected an `initial_grounded_snapshot` capture before the "
            "escalation call so models_disagreed can reference the "
            "pre-swap grounded state."
        )
        # verify_finding must hand the snapshot to the shared merge helper.
        passes_idx = source.find("initial_grounded=initial_grounded_snapshot")
        assert passes_idx > snapshot_idx, (
            "Expected verify_finding to pass `initial_grounded="
            "initial_grounded_snapshot` into `_apply_escalation_outcome` "
            "AFTER taking the snapshot."
        )
        # The helper is the single source of truth for the disagreement
        # condition. Isolate its body and assert the strict three-part
        # condition.
        helper_idx = source.find("def _apply_escalation_outcome(")
        assert helper_idx > 0, "Expected the shared _apply_escalation_outcome helper."
        models_disagreed_idx = source.find("result.models_disagreed = (", helper_idx)
        assert models_disagreed_idx > helper_idx, (
            "Expected `result.models_disagreed = (...)` inside the helper."
        )
        block = source[models_disagreed_idx : models_disagreed_idx + 400]
        assert "initial_grounded" in block
        assert "esc_result.grounded" in block
        assert "esc_result.verdict != initial_verdict" in block

    def test_verify_finding_records_initial_sources_snapshot(self):
        source = Path("src/verification/verifier.py").read_text(encoding="utf-8")
        snapshot_idx = source.find("initial_sources_snapshot = list(result.sources or [])")
        assert snapshot_idx > 0, (
            "Expected an `initial_sources_snapshot = list(result.sources or [])` "
            "capture before the escalation call so initial_sources persists "
            "after the potential swap to esc_result."
        )
        # verify_finding hands the snapshot to the shared helper, which sets
        # ``result.initial_sources`` unconditionally.
        passes_idx = source.find("initial_sources=initial_sources_snapshot")
        assert passes_idx > snapshot_idx, (
            "Expected verify_finding to pass `initial_sources="
            "initial_sources_snapshot` into `_apply_escalation_outcome` "
            "after taking the snapshot."
        )
        helper_idx = source.find("def _apply_escalation_outcome(")
        assert source.find("result.initial_sources = list(initial_sources)", helper_idx) > helper_idx, (
            "Expected the helper to set result.initial_sources from the "
            "passed-in snapshot."
        )


# ---------------------------------------------------------------------------
# 7. Cache persist + replay — contested fields round-trip
# ---------------------------------------------------------------------------


class TestCacheContested:
    def _finding_for_cache(self) -> Finding:
        return Finding(
            severity="HIGH",
            fileName="Section_22_1000.docx",
            section="2.1",
            issue="claim about clearance",
            actionType="EDIT",
            existingText="5 ft clearance",
            replacementText="6 ft clearance",
            codeReference="NFPA 13 §10",
            confidence=0.6,
        )

    def test_result_to_dict_includes_contested_fields(self):
        # The cache persist path must serialize the disagreement
        # telemetry so a cache replay surfaces VERIFIED_CONTESTED for
        # the same finding the original run flagged.
        result = _contested_verification()
        d = _result_to_dict(result)
        assert "models_disagreed" in d
        assert d["models_disagreed"] is True
        assert "initial_sources" in d
        assert d["initial_sources"] == result.initial_sources

    def test_clone_for_store_preserves_contested_state(self):
        original = _contested_verification()
        clone = _clone_for_store(original)
        assert clone.models_disagreed is True
        assert clone.initial_sources == original.initial_sources
        # The clone's initial_sources must be an independent list so
        # mutations to one do not leak into the other.
        clone.initial_sources.append("https://new")
        assert "https://new" not in original.initial_sources

    def test_clone_for_hit_preserves_contested_state(self):
        stored = _contested_verification()
        # Strip the cache_status off the stored result so we can verify
        # _clone_for_hit stamps it back to "hit".
        stored.cache_status = "miss"
        entry = _CacheEntry(result=stored, created_ts=1_700_000_000.0)
        hit = _clone_for_hit(entry)
        assert hit.cache_status == "hit"
        assert hit.models_disagreed is True
        assert hit.initial_sources == stored.initial_sources

    def test_cache_persist_and_load_preserves_contested(self, tmp_path: Path):
        # End-to-end: persist a contested grounded verdict to disk, load
        # it back, verify the replay still classifies as VERIFIED_CONTESTED.
        cache = VerificationCache()
        f = self._finding_for_cache()
        cache.put(f, cycle=DEFAULT_CYCLE, result=_contested_verification())
        path = tmp_path / "verification_cache.json"
        cache.save_to_disk(path)
        # Sanity check: the JSON on disk actually contains the new keys.
        raw = json.loads(path.read_text(encoding="utf-8"))
        entries = list(raw["entries"].values())
        assert entries, "Expected at least one entry persisted."
        assert entries[0]["result"]["models_disagreed"] is True
        assert entries[0]["result"]["initial_sources"]
        # Reload into a fresh cache and verify replay.
        fresh = VerificationCache()
        fresh.load_from_disk(path)
        replayed = fresh.get(f, cycle=DEFAULT_CYCLE)
        assert replayed is not None
        assert replayed.models_disagreed is True
        assert replayed.initial_sources

    def test_cache_load_legacy_entry_defaults_neutral(self, tmp_path: Path):
        # A v3 cache entry written before Chunk 12 lacks both new keys.
        # Loading it must not crash and must default the fields to
        # neutral values so the legacy entry classifies via the
        # verdict-based branches.
        import time as _time

        recent_ts = _time.time() - 86400  # 1 day old — well inside the 60d TTL
        path = tmp_path / "verification_cache.json"
        path.write_text(
            json.dumps(
                {
                    "version": 3,
                    "saved_at": recent_ts,
                    "entries": {
                        "test_key|EDIT|NFPA 13 §10|abc123": {
                            "created_ts": recent_ts,
                            "result": {
                                "verdict": "CONFIRMED",
                                "explanation": "Verified.",
                                "sources": ["https://x"],
                                "correction": None,
                                "grounded": True,
                                "model_used": "claude-sonnet-4-6",
                                "escalated": False,
                                "web_search_requests": 1,
                                "successful_source_count": 1,
                                "search_error_count": 0,
                                "searched_sources": ["https://x"],
                                "cited_sources": ["https://x"],
                                "accepted_sources": ["https://x"],
                                "rejected_sources": [],
                                "verification_profile": "code_standard",
                                "verification_mode": "standard_reasoning",
                                "source_quote": "snippet",
                                "web_fetch_requests": 0,
                                "fetched_sources": [],
                                # models_disagreed + initial_sources
                                # intentionally omitted (legacy entry).
                            },
                        }
                    },
                }
            ),
            encoding="utf-8",
        )
        cache = VerificationCache()
        loaded = cache.load_from_disk(path)
        assert loaded == 1
        # The cache key is computed from the finding, not the raw key
        # string in the JSON. Instead of fishing the result out by key
        # we inspect the entries map directly.
        with cache._lock:
            (entry,) = list(cache._entries.values())
        assert entry.result.models_disagreed is False
        assert entry.result.initial_sources == []


# ---------------------------------------------------------------------------
# 8. Report rendering — evidence panel surfaces disagreement
# ---------------------------------------------------------------------------


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestContestedEvidencePanelRendering:
    def test_renders_contested_status_label(self, tmp_path: Path):
        contested = _finding(verification=_contested_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[contested])), out
        )
        text = _all_text_from(Document(str(out)))
        # The top-level status badge must show the contested label.
        assert "Verified — but models disagreed" in text or "disagreed" in text

    def test_renders_disagreement_sentence_in_escalation_history(
        self, tmp_path: Path
    ):
        contested = _finding(verification=_contested_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[contested])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Escalation history:" in text
        # The Chunk 12 expanded sentence must appear when both passes
        # were grounded and disagreed. The shorter "(models disagreed)"
        # parenthetical from the existing escalation block belongs to
        # the weaker ``escalation_changed_verdict`` path — for the
        # strict ``models_disagreed`` case we render the full sentence
        # with "manual review recommended".
        assert "manual review recommended" in text.lower()

    def test_renders_both_models_in_escalation_line(self, tmp_path: Path):
        contested = _finding(verification=_contested_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[contested])), out
        )
        text = _all_text_from(Document(str(out)))
        # Both verifier model identifiers must appear inline so a
        # reviewer sees "Sonnet 4.6: DISPUTED → Opus 4.7: CONFIRMED."
        assert "claude-sonnet-4-6" in text
        assert "claude-opus-4-7" in text
        assert "DISPUTED" in text  # initial verdict
        assert "CONFIRMED" in text  # final verdict

    def test_renders_initial_verifier_sources_section(self, tmp_path: Path):
        contested = _finding(verification=_contested_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[contested])), out
        )
        text = _all_text_from(Document(str(out)))
        # The dedicated sub-section must appear, listing the initial
        # verifier's citations. Below it the existing "Web/code
        # evidence" sub-section will show the final verifier's
        # citations, so a reviewer can compare side-by-side.
        assert "Initial verifier sources" in text
        assert "https://nfpa.org/codes/13/2022" in text  # initial source
        assert "https://codes.iccsafe.org/content/CBC2025" in text  # final source

    def test_no_initial_sources_section_when_not_contested(self, tmp_path: Path):
        # A finding that escalated but did NOT trigger models_disagreed
        # (e.g. initial UNVERIFIED, escalated CONFIRMED) should not get
        # the "Initial verifier sources" sub-section. The existing
        # "Escalation history" line is still rendered (escalation
        # happened); only the Chunk 12 sub-section is conditional.
        v = VerificationResult(
            verdict="CONFIRMED",
            explanation="Verified after escalation.",
            sources=["https://opus-source"],
            accepted_sources=["https://opus-source"],
            grounded=True,
            model_used="claude-opus-4-7",
            escalated=True,
            cache_status="miss",
            web_search_requests=5,
            source_quote="snippet",
            escalation_attempted=True,
            initial_model="claude-sonnet-4-6",
            initial_verdict="UNVERIFIED",  # Sonnet didn't ground anything
            escalation_changed_verdict=True,
            escalation_reason="initial_unverified",
            models_disagreed=False,  # Sonnet wasn't grounded, so no real disagreement
            initial_sources=[],  # nothing for Sonnet to cite
        )
        finding = _finding(verification=v)
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[finding])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Initial verifier sources" not in text


# ---------------------------------------------------------------------------
# 9. Resume retry-failed-only env var stub
# ---------------------------------------------------------------------------
