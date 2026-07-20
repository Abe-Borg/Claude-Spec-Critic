"""End-to-end: the location-aware pipeline under the ARCHITECTURAL module.

Sibling of ``tests/test_datacenter_e2e.py`` (which drives the fire module and
proves CA-neutrality through the same path — not duplicated here). Drives
``run_batch_collection_headless`` for ``datacenter_architectural``
(``project_profile_enabled=True``) with every network-touching stage faked
(review collect, verification rounds, cross-check, compliance) and asserts the
same artifact contract the fire e2e pins, re-themed to architectural content:

- the report's "Jurisdiction & Client Requirements" section, the project /
  client title lines, the coverage matrix + advisories, the ``[Compliance]``-
  labeled findings, and the two Run-Diagnostics banner rows — under the
  ARCHITECTURAL report title;
- the sidecar (schema v4) with the ``project`` block, the
  ``requirements_coverage`` matrix, and a compliance edit entry (``lc-`` id);
- the standalone ``<report-stem>.profile.json`` export stamped with this
  module's id;
- verification round 2 actually receiving the compliance findings.

All client/location identifiers are anonymized fixtures.
"""
from __future__ import annotations

import json

from docx import Document

from src.batch.batch import BatchJob
from src.core.project_profile import ProjectProfile
from src.input.extractor import ExtractedSpec
from src.modules import DATACENTER_ARCHITECTURAL
from src.orchestration import pipeline as pl
from src.orchestration.pipeline import BatchSubmission, run_batch_collection_headless
from src.output.edit_sidecar import (
    write_edit_instructions_sidecar,
    write_requirements_profile_sidecar,
)
from src.output.report_exporter import export_report
from src.research import DimensionStatus, RequirementsProfile, ResearchItem
from src.review.reviewer import Finding, ReviewResult
from src.verification.verification_cache import VerificationCache


# ---------------------------------------------------------------------------
# Fixtures — a two-spec architectural package + a researched profile.
# ---------------------------------------------------------------------------

_DOORS = "08 11 13 - Hollow Metal Doors and Frames.docx"
_FIRESTOP = "07 84 13 - Penetration Firestopping.docx"

_DOORS_CONTENT = (
    "PART 1 - GENERAL\n\n"
    "1.01 SUMMARY\n\n"
    "A. Comply with 2015 IBC Chapter 7.\n\n"
    "B. Provide fire-rated hollow metal doors and frames per NFPA 80."
)
_FIRESTOP_CONTENT = "1.01 SUMMARY\n\nA. Provide penetration firestopping as indicated."


def _specs() -> list[ExtractedSpec]:
    return [
        ExtractedSpec(filename=_DOORS, content=_DOORS_CONTENT, word_count=22),
        ExtractedSpec(filename=_FIRESTOP, content=_FIRESTOP_CONTENT, word_count=8),
    ]


def _profile() -> ProjectProfile:
    return ProjectProfile(
        city="Ashburn", state_or_province="VA", country="US", client_name="ExampleCo"
    )


def _requirements_profile() -> RequirementsProfile:
    """Grounded + ungrounded + process-advisory items across the sections."""
    items = [
        ResearchItem(
            item_id="r-000000000001",
            dimension_id="governing_codes",
            topic="Building code edition",
            category="governing_code",
            requirement="The 2021 Virginia USBC (2021 IBC basis) governs.",
            authority="Virginia DHCD",
            code_reference="13VAC5-63",
            accepted_sources=["https://law.lis.virginia.gov/admincode/"],
            grounded=True,
            confidence=0.92,
        ),
        ResearchItem(
            item_id="r-000000000002",
            dimension_id="governing_codes",
            topic="Referenced fire-door standard",
            category="referenced_standard",
            requirement="The adopted code references NFPA 80-2019.",
            authority="Virginia DHCD",
            code_reference="IBC 716",
            accepted_sources=["https://law.lis.virginia.gov/admincode/"],
            grounded=True,
            confidence=0.8,
        ),
        ResearchItem(
            item_id="r-000000000003",
            dimension_id="ahj_requirements",
            topic="Air-barrier testing window",
            category="ahj_requirement",
            requirement="Whole-building airtightness tests are witnessed April through October only.",
            authority="Loudoun County Building Department",
            accepted_sources=["https://www.loudoun.gov/building"],
            grounded=True,
            confidence=0.7,
            actionability="process_advisory",
        ),
        ResearchItem(
            item_id="r-000000000004",
            dimension_id="client_standards",
            topic="Owner roofing preference",
            category="client_standard",
            requirement="ExampleCo prefers mechanically attached single-ply roofing on data halls.",
            authority="ExampleCo",
            grounded=False,
            confidence=0.4,
        ),
    ]
    statuses = [
        DimensionStatus("governing_codes", "completed", item_count=2, grounded_count=2,
                        web_search_requests=18),
        DimensionStatus("ahj_requirements", "completed", item_count=1, grounded_count=1,
                        web_search_requests=12),
        DimensionStatus("client_standards", "failed",
                        error="owner standards confidential / unretrievable"),
        DimensionStatus("site_environment", "completed", item_count=0, grounded_count=0,
                        web_search_requests=6),
    ]
    return RequirementsProfile(
        items=items,
        dimension_statuses=statuses,
        research_date="2026-07-15",
        project=_profile().to_dict(),
    )


def _review_finding() -> Finding:
    # EDIT whose existingText is verbatim in _DOORS_CONTENT so it survives the
    # deterministic anchor validation and reaches the edit sidecar (rf- id).
    return Finding(
        severity="MEDIUM",
        fileName=_DOORS,
        section="1.01",
        issue="Spec cites a stale IBC edition rather than the adopted edition.",
        actionType="EDIT",
        existingText="Comply with 2015 IBC Chapter 7.",
        replacementText="Comply with the current IBC edition adopted for this project location.",
        codeReference="IBC (current adopted edition)",
        confidence=0.9,
    )


def _cross_finding() -> Finding:
    return Finding(
        severity="HIGH",
        fileName=_FIRESTOP,
        section="3.05",
        issue="Firestop system schedule references wall types the Division 08 door schedule does not protect.",
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference=None,
    )


def _compliance_result() -> ReviewResult:
    # An ADD anchored verbatim on _DOORS_CONTENT (survives anchor validation)
    # so a compliance edit lands in the sidecar; coverage exercises the
    # report's represented / missing rows.
    add = Finding(
        severity="HIGH",
        fileName=_DOORS,
        section="1.01",
        issue="The package does not reference the adopted NFPA 80-2019 edition "
        "(profile requirement r-000000000002).",
        actionType="ADD",
        existingText=None,
        replacementText="C. Comply with NFPA 80-2019 as referenced by the adopted code.",
        codeReference="IBC 716",
        anchorText="PART 1 - GENERAL",
        insertPosition="after",
        confidence=0.85,
    )
    return ReviewResult(
        findings=[add],
        cross_check_status="completed",
        coverage=[
            {"requirement_id": "r-000000000001", "status": "represented",
             "evidence": "Comply with 2015 IBC Chapter 7.", "fileName": _DOORS},
            {"requirement_id": "r-000000000002", "status": "missing",
             "evidence": None, "fileName": None},
        ],
        thinking="One adopted standard edition is missing from the package.",
    )


def _submission(module, *, profile, requirements_profile, cross_check) -> BatchSubmission:
    cid0, cid1 = "review__08_11_13__0", "review__07_84_13__1"
    job = BatchJob(
        batch_id="msgbatch_DCARCH_E2E",
        job_type="review",
        request_map={
            cid0: {"filename": _DOORS, "index": 0, "type": "review"},
            cid1: {"filename": _FIRESTOP, "index": 1, "type": "review"},
        },
        created_at=1_700_000_000.0,
    )
    return BatchSubmission(
        job=job,
        files_reviewed=[_DOORS, _FIRESTOP],
        review_request_ids=[cid0, cid1],
        model="claude-opus-4-8",
        project_context="Hyperscale data-center new-build program.",
        prepared_specs=_specs(),
        cycle_label=module.cycle.label,
        module_id=module.module_id,
        project_profile=profile.to_dict() if profile is not None else None,
        requirements_profile=requirements_profile.to_dict()
        if requirements_profile is not None else None,
        cross_check_enabled=cross_check,
    )


class _VerificationSpy:
    """Records every findings list handed to start_batch_verification.

    Returns ``None`` (all findings resolved locally) so no batch is submitted
    and no network is touched — both verification rounds route through here.
    """

    def __init__(self):
        self.calls: list[list[Finding]] = []

    def __call__(self, findings, **kwargs):
        self.calls.append(list(findings))
        return None

    @property
    def all_findings(self) -> list[Finding]:
        return [f for call in self.calls for f in call]


def _install_fakes(monkeypatch, *, review_map, cross_result, verify_spy,
                   compliance_result=None):
    monkeypatch.setattr(pl, "retrieve_review_results", lambda job, *, model: review_map)
    monkeypatch.setattr(pl, "start_batch_verification", verify_spy)
    if cross_result is not None:
        monkeypatch.setattr(
            pl, "run_chunked_cross_check", lambda specs, existing, **kw: cross_result
        )
    if compliance_result is not None:
        # run_compliance_for_batch imports this lazily from the package.
        import src.compliance as compliance_pkg

        monkeypatch.setattr(
            compliance_pkg,
            "run_chunked_compliance_check",
            lambda specs, profile, existing, **kw: compliance_result,
        )


def _doc_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# The architectural (flag ON) end-to-end run.
# ---------------------------------------------------------------------------


class TestDatacenterArchitecturalEndToEnd:
    def _run(self, monkeypatch):
        verify_spy = _VerificationSpy()
        review_map = {
            "review__08_11_13__0": ReviewResult(
                findings=[_review_finding()], parse_status="ok"
            ),
            "review__07_84_13__1": ReviewResult(findings=[], parse_status="ok"),
        }
        _install_fakes(
            monkeypatch,
            review_map=review_map,
            cross_result=ReviewResult(
                findings=[_cross_finding()], cross_check_status="completed"
            ),
            verify_spy=verify_spy,
            compliance_result=_compliance_result(),
        )
        submission = _submission(
            DATACENTER_ARCHITECTURAL,
            profile=_profile(),
            requirements_profile=_requirements_profile(),
            cross_check=True,
        )
        result = run_batch_collection_headless(
            submission, cache=VerificationCache(), log=lambda *a, **k: None
        )
        return result, verify_spy

    def test_pipeline_result_carries_all_phases(self, monkeypatch):
        result, verify_spy = self._run(monkeypatch)

        # Review + cross-check + compliance all present and labeled/id-stamped.
        assert [f.issue for f in result.review_result.findings]
        assert result.cross_check_result.findings
        comp = result.compliance_result
        assert comp is not None and comp.cross_check_status == "completed"
        assert len(comp.findings) == 1
        lc = comp.findings[0]
        assert lc.section.startswith("[Compliance]")
        assert lc.finding_id.startswith("lc-")
        # The compliance edit survived anchor validation (verbatim anchor).
        assert lc.actionType == "ADD"

        # Identity + profile ride through to the result.
        assert result.module_id == "datacenter_architectural"
        assert result.project_profile["city"] == "Ashburn"
        assert result.requirements_profile["research_date"] == "2026-07-15"

        # Verification round 2 actually received the compliance finding.
        submitted_compliance = [
            f for f in verify_spy.all_findings
            if (f.section or "").startswith("[Compliance]")
        ]
        assert submitted_compliance, "compliance findings must reach verification round 2"

    def test_report_has_location_aware_surfaces_under_arch_title(self, monkeypatch, tmp_path):
        result, _ = self._run(monkeypatch)
        out = tmp_path / "report.docx"
        export_report(result, out)
        text = _doc_text(out)

        # The ARCHITECTURAL report identity (never the fire module's).
        assert "Spec Critic — Architectural Specification Review Report" in text
        assert "Code Cycle: dc-arch-ibc-2024" in text
        assert "Fire Protection Specification Review Report" not in text

        # Requirements section + title lines.
        assert "Jurisdiction & Client Requirements" in text
        assert "Ashburn, Virginia, USA" in text
        assert "ExampleCo" in text
        # Coverage matrix (represented + missing) and advisories.
        assert "Requirements Coverage" in text
        assert "Represented" in text and "MISSING" in text
        assert "Process & Schedule Advisories" in text
        assert "[UNVERIFIED]" in text
        # Compliance findings render (labeled) and the compliance section shows.
        assert "Local-Code Compliance" in text
        assert "NFPA 80-2019" in text
        # Both conditional Run-Diagnostics banner rows.
        assert "Location/client research" in text
        assert "Local-code compliance" in text
        # standards=() -> no pinned-editions methodology paragraph.
        assert "per the dc-arch-ibc-2024 cycle:" not in text

    def test_sidecar_v4_and_profile_json(self, monkeypatch, tmp_path):
        result, _ = self._run(monkeypatch)
        out = tmp_path / "report.docx"
        export_report(result, out)
        sidecar_path = write_edit_instructions_sidecar(result, out)
        profile_path = write_requirements_profile_sidecar(result, out)

        sidecar = json.loads(sidecar_path.read_text())
        assert sidecar["schema_version"] == 4
        assert sidecar["project"]["city"] == "Ashburn"
        coverage_ids = {c["requirement_id"] for c in sidecar["requirements_coverage"]}
        assert {"r-000000000001", "r-000000000002"} <= coverage_ids
        # A compliance edit entry (lc- id) AND the review edit (rf- id) are present.
        prefixes = {e["finding_id"][:3] for e in sidecar["edits"]}
        assert "lc-" in prefixes
        assert "rf-" in prefixes

        assert profile_path is not None
        profile_json = json.loads(profile_path.read_text())
        assert profile_json["research_date"] == "2026-07-15"
        assert profile_json["module_id"] == "datacenter_architectural"
        assert profile_json["requirements_profile"]["items"]
