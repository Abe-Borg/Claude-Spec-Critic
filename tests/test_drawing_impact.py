"""Tests for the drawing-impact synthesis pass (WS-5).

Covers, hermetically (no network, no API key):

* Digest extraction from Project Context (the gate).
* The tool schema stays inside the strict-mode supported subset.
* Prompt builders (system-prompt stability, user-message shape).
* Payload parsing: level/relationship coercion, hallucinated-id drop,
  dedup, markdown sanitization.
* ``run_drawing_impact`` end-to-end against a scripted fake client
  (tool path, text fallback, incomplete, exception, empty findings).
* ``run_drawing_impact_for_batch`` gating (no digest -> None; digest ->
  runs) and ``finalize_batch_result`` carry-through.
* Report rendering: the section + banner row appear only when a result is
  present, so a drawing-less report is byte-identical.
"""
from __future__ import annotations

import json
import types
from pathlib import Path

import pytest
from docx import Document

import src.drawing_impact as di_pkg
from src.drawing_impact import (
    DrawingImpactResult,
    build_impact_system_prompt,
    build_impact_user_message,
    extract_drawing_digest,
    render_findings_block,
    run_drawing_impact,
)
from src.drawing_impact.impact_synthesizer import (
    _extract_impact_object,
    _parse_impact_payload,
)
from src.core.api_config import (
    MODEL_SONNET_5,
    PHASE_DRAWING_IMPACT,
    cache_policy_for,
    drawing_impact_max_tokens,
    effort_config_for,
)
from src.gui.context_attachment import wrap_attachment
from src.input.drawing_digest import DIGEST_ATTACHMENT_LABEL
from src.orchestration import pipeline
from src.orchestration.pipeline import (
    BatchSubmission,
    CollectedBatchState,
    PipelineResult,
    finalize_batch_result,
    run_drawing_impact_for_batch,
)
from src.review.reviewer import Finding, ReviewResult
from src.review.structured_schemas import (
    DRAWING_IMPACT_SCHEMA,
    DRAWING_IMPACT_TOOL_NAME,
    drawing_impact_tool,
    drawing_impact_tool_choice,
)
from tests.fixtures.fake_anthropic import (
    FakeMessage,
    FakeTextBlock,
    FakeToolUseBlock,
    FakeUsage,
)


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    finding_id: str = "rf-aaaaaaaaaaaa",
    severity: str = "HIGH",
    file: str = "230553.docx",
    section: str = "2.1",
    issue: str = "Valve tag does not match the mechanical schedule.",
) -> Finding:
    return Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference=None,
        confidence=0.7,
        finding_id=finding_id,
    )


def _impact_payload(**overrides) -> dict:
    payload = {
        "impact_level": "moderate",
        "narrative": "The drawings clarified the valve schedule.",
        "finding_links": [
            {
                "finding_id": "rf-aaaaaaaaaaaa",
                "relationship": "corroborated",
                "explanation": "The M-501 schedule confirms the tag mismatch.",
                "sheet_references": ["plans.pdf p.5"],
            }
        ],
    }
    payload.update(overrides)
    return payload


class _FakeStream:
    def __init__(self, message, text: str = ""):
        self._message = message
        self._text = text

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False

    @property
    def text_stream(self):
        return iter([self._text]) if self._text else iter(())

    def get_final_message(self):
        return self._message


class _FakeMessages:
    def __init__(self, responder):
        self._responder = responder
        self.calls: list[dict] = []

    def stream(self, **kwargs):
        self.calls.append(kwargs)
        result = self._responder(kwargs)
        if isinstance(result, Exception):
            raise result
        message, text = result if isinstance(result, tuple) else (result, "")
        return _FakeStream(message, text)


class FakeClient:
    def __init__(self, responder):
        self.messages = _FakeMessages(responder)

    @property
    def calls(self):
        return self.messages.calls


def _tool_message(payload: dict, *, stop_reason: str = "tool_use") -> FakeMessage:
    return FakeMessage(
        content=[FakeToolUseBlock(name=DRAWING_IMPACT_TOOL_NAME, input=dict(payload))],
        stop_reason=stop_reason,
        usage=FakeUsage(input_tokens=1200, output_tokens=300),
    )


# ---------------------------------------------------------------------------
# 1. Digest extraction (the gate)
# ---------------------------------------------------------------------------


class TestDigestExtraction:
    def test_empty_and_none(self):
        assert extract_drawing_digest("") == ""
        assert extract_drawing_digest(None) == ""

    def test_no_digest_block(self):
        assert extract_drawing_digest("Just some free-text project notes.") == ""

    def test_single_block(self):
        digest = "SHEET INDEX\nA-101 Overall Plan [plans.pdf p.1]"
        ctx = "Project notes.\n\n" + wrap_attachment(DIGEST_ATTACHMENT_LABEL, digest)
        assert extract_drawing_digest(ctx) == digest

    def test_multiple_blocks_joined(self):
        a = wrap_attachment(DIGEST_ATTACHMENT_LABEL, "First set [a.pdf p.1]")
        b = wrap_attachment(DIGEST_ATTACHMENT_LABEL, "Second set [b.pdf p.1]")
        ctx = a + "\n\n" + b
        got = extract_drawing_digest(ctx)
        assert "First set [a.pdf p.1]" in got
        assert "Second set [b.pdf p.1]" in got

    def test_similarly_named_attachment_not_matched(self):
        # A context file the user happened to name like the digest carries the
        # extension in its label, so its BEGIN line differs and is not matched.
        ctx = wrap_attachment("Construction Drawing Digest.docx", "not a digest")
        assert extract_drawing_digest(ctx) == ""


# ---------------------------------------------------------------------------
# 2. Schema + phase registration
# ---------------------------------------------------------------------------


class TestSchemaAndRegistration:
    def test_strict_subset(self):
        def _walk(schema):
            if schema.get("type") == "object" or "properties" in schema:
                assert schema.get("additionalProperties") is False
                assert sorted(schema["required"]) == sorted(schema["properties"])
            for value in schema.get("properties", {}).values():
                _walk(value)
                if isinstance(value.get("items"), dict):
                    _walk(value["items"])

        _walk(DRAWING_IMPACT_SCHEMA)
        text = json.dumps(DRAWING_IMPACT_SCHEMA)
        for forbidden in ('"minimum"', '"maximum"', '"minLength"'):
            assert forbidden not in text

    def test_tool_builder_and_choice(self):
        tool = drawing_impact_tool(model=MODEL_SONNET_5)
        assert tool["name"] == DRAWING_IMPACT_TOOL_NAME == "submit_drawing_impact"
        assert tool["strict"] is True
        assert drawing_impact_tool_choice() == {
            "type": "auto",
            "disable_parallel_tool_use": True,
        }

    def test_phase_registered(self):
        # An unregistered phase silently caps at the verification 16k default
        # and drops effort/cache — pin the explicit registration.
        assert drawing_impact_max_tokens(model=MODEL_SONNET_5) == 16_000
        assert effort_config_for(model=MODEL_SONNET_5, phase=PHASE_DRAWING_IMPACT) == {
            "effort": "high"
        }
        assert cache_policy_for(PHASE_DRAWING_IMPACT).caches_anything is True


# ---------------------------------------------------------------------------
# 3. Prompt builders
# ---------------------------------------------------------------------------


class TestPromptBuilders:
    def test_system_prompt_is_stable_and_grounded(self):
        p1 = build_impact_system_prompt()
        p2 = build_impact_system_prompt()
        assert p1 == p2  # byte-identical across calls (cacheable)
        assert "submit_drawing_impact" in p1
        # Honesty / grounding language must be present.
        assert "Never invent a page reference" in p1
        assert "manufactured" in p1

    def test_user_message_wraps_digest_and_findings(self):
        f = _finding()
        msg = build_impact_user_message(
            "DIGEST BODY [plans.pdf p.2]", [f], module_display_name="California K-12"
        )
        assert "<drawing_digest>" in msg and "</drawing_digest>" in msg
        assert "<review_findings>" in msg
        assert f.finding_id in msg
        assert "California K-12" in msg
        assert "1 finding(s)" in msg

    def test_findings_block_escapes_and_tags(self):
        f = _finding(issue="Angle < 30 deg & tag mismatch")
        block = render_findings_block([f])
        # Reserved characters in the body are escaped, not raw.
        assert "&lt;" in block and "&amp;" in block
        assert f'id="{f.finding_id}"' in block
        assert 'severity="HIGH"' in block


# ---------------------------------------------------------------------------
# 4. Payload parsing
# ---------------------------------------------------------------------------


class TestParsePayload:
    def test_happy_path(self):
        level, narrative, links = _parse_impact_payload(
            _impact_payload(), {"rf-aaaaaaaaaaaa"}
        )
        assert level == "moderate"
        assert narrative == "The drawings clarified the valve schedule."
        assert len(links) == 1
        assert links[0].finding_id == "rf-aaaaaaaaaaaa"
        assert links[0].relationship == "corroborated"
        assert links[0].sheet_references == ["plans.pdf p.5"]

    def test_unknown_level_coerces_to_minimal(self):
        level, _, _ = _parse_impact_payload(
            _impact_payload(impact_level="enormous"), {"rf-aaaaaaaaaaaa"}
        )
        assert level == "minimal"

    def test_unknown_relationship_coerces_to_contextualized(self):
        payload = _impact_payload(
            finding_links=[
                {
                    "finding_id": "rf-aaaaaaaaaaaa",
                    "relationship": "obliterated",
                    "explanation": "x",
                    "sheet_references": [],
                }
            ]
        )
        _, _, links = _parse_impact_payload(payload, {"rf-aaaaaaaaaaaa"})
        assert links[0].relationship == "contextualized"

    def test_hallucinated_id_dropped(self):
        payload = _impact_payload(
            finding_links=[
                {
                    "finding_id": "rf-GHOSTGHOST",
                    "relationship": "contradicted",
                    "explanation": "invented",
                    "sheet_references": [],
                }
            ]
        )
        _, _, links = _parse_impact_payload(payload, {"rf-aaaaaaaaaaaa"})
        assert links == []

    def test_duplicate_id_deduped_first_wins(self):
        payload = _impact_payload(
            finding_links=[
                {
                    "finding_id": "rf-aaaaaaaaaaaa",
                    "relationship": "corroborated",
                    "explanation": "first",
                    "sheet_references": [],
                },
                {
                    "finding_id": "rf-aaaaaaaaaaaa",
                    "relationship": "contradicted",
                    "explanation": "second",
                    "sheet_references": [],
                },
            ]
        )
        _, _, links = _parse_impact_payload(payload, {"rf-aaaaaaaaaaaa"})
        assert len(links) == 1
        assert links[0].explanation == "first"

    def test_markdown_narrative_sanitized(self):
        _, narrative, _ = _parse_impact_payload(
            _impact_payload(narrative="## Overview\nThe drawings helped."),
            {"rf-aaaaaaaaaaaa"},
        )
        assert not narrative.startswith("#")
        assert "The drawings helped." in narrative

    def test_non_list_links_tolerated(self):
        _, _, links = _parse_impact_payload(
            _impact_payload(finding_links="not a list"), {"rf-aaaaaaaaaaaa"}
        )
        assert links == []

    def test_text_fallback_object_extraction(self):
        body = "<drawing_impact_json>" + json.dumps(_impact_payload()) + "</drawing_impact_json>"
        obj = _extract_impact_object("prose... " + body + " ...more")
        assert obj is not None and obj["impact_level"] == "moderate"
        assert _extract_impact_object("no json here") is None


# ---------------------------------------------------------------------------
# 5. run_drawing_impact end-to-end (scripted fake client)
# ---------------------------------------------------------------------------


class TestRunDrawingImpact:
    def test_tool_path_completes(self):
        client = FakeClient(lambda kw: _tool_message(_impact_payload()))
        result = run_drawing_impact(
            digest_text="DIGEST [plans.pdf p.5]",
            findings=[_finding()],
            client=client,
        )
        assert result.status == "completed"
        assert result.impact_level == "moderate"
        assert result.linked_finding_count == 1
        assert result.input_tokens == 1200 and result.output_tokens == 300
        # The request carried the tool + a real max_tokens cap.
        assert client.calls[0]["max_tokens"] == 16_000
        assert client.calls[0]["tools"][0]["name"] == DRAWING_IMPACT_TOOL_NAME

    def test_text_fallback_path_completes(self):
        body = "<drawing_impact_json>" + json.dumps(_impact_payload()) + "</drawing_impact_json>"
        msg = FakeMessage(content=[FakeTextBlock(text=body)], stop_reason="end_turn")
        client = FakeClient(lambda kw: (msg, body))
        result = run_drawing_impact(
            digest_text="DIGEST", findings=[_finding()], client=client
        )
        assert result.status == "completed"
        assert result.linked_finding_count == 1

    def test_incomplete_stop_reason_fails(self):
        msg = FakeMessage(content=[FakeTextBlock(text="...")], stop_reason="max_tokens")
        client = FakeClient(lambda kw: (msg, "..."))
        result = run_drawing_impact(
            digest_text="DIGEST", findings=[_finding()], client=client
        )
        assert result.status == "failed"
        assert "max_tokens" in result.error

    def test_api_exception_fails_without_raising(self):
        client = FakeClient(lambda kw: ValueError("boom"))
        result = run_drawing_impact(
            digest_text="DIGEST", findings=[_finding()], client=client, max_retries=1
        )
        assert result.status == "failed"
        assert "boom" in result.error

    def test_empty_findings_still_runs(self):
        # No findings -> narrative-only; the model can still speak to the
        # drawings' overall contribution.
        client = FakeClient(
            lambda kw: _tool_message(_impact_payload(finding_links=[]))
        )
        result = run_drawing_impact(digest_text="DIGEST", findings=[], client=client)
        assert result.status == "completed"
        assert result.linked_finding_count == 0
        # A finding with no id is not linkable and never reaches the prompt.
        assert "<review_findings>" in client.calls[0]["messages"][0]["content"]

    def test_hallucinated_link_dropped_end_to_end(self):
        payload = _impact_payload(
            finding_links=[
                {
                    "finding_id": "rf-NOTREAL0000",
                    "relationship": "corroborated",
                    "explanation": "ghost",
                    "sheet_references": [],
                }
            ]
        )
        client = FakeClient(lambda kw: _tool_message(payload))
        result = run_drawing_impact(
            digest_text="DIGEST", findings=[_finding()], client=client
        )
        assert result.status == "completed"
        assert result.finding_links == []


# ---------------------------------------------------------------------------
# 6. Pipeline gating + carry-through
# ---------------------------------------------------------------------------


def _submission(project_context: str) -> BatchSubmission:
    job = types.SimpleNamespace(created_at=0.0, batch_id="batch_fake")
    return BatchSubmission(job=job, project_context=project_context)


def _state(project_context: str, findings: list[Finding]) -> CollectedBatchState:
    return CollectedBatchState(
        submission=_submission(project_context),
        review_result=ReviewResult(findings=findings),
    )


class TestRunDrawingImpactForBatch:
    def test_no_digest_is_noop(self, monkeypatch):
        called = {"n": 0}

        def _spy(**kwargs):
            called["n"] += 1
            return DrawingImpactResult(status="completed")

        monkeypatch.setattr(di_pkg, "run_drawing_impact", _spy)
        state = _state("Just project notes, no drawings.", [_finding()])
        out = run_drawing_impact_for_batch(state)
        assert out.drawing_impact_result is None
        assert called["n"] == 0  # the API pass never ran

    def test_digest_present_runs_and_sets_result(self, monkeypatch):
        seen = {}

        def _spy(*, digest_text, findings, **kwargs):
            seen["digest"] = digest_text
            seen["ids"] = [f.finding_id for f in findings]
            return DrawingImpactResult(status="completed", impact_level="substantial")

        monkeypatch.setattr(di_pkg, "run_drawing_impact", _spy)
        digest = "SHEET INDEX\nM-501 [plans.pdf p.5]"
        ctx = wrap_attachment(DIGEST_ATTACHMENT_LABEL, digest)
        state = _state(ctx, [_finding()])
        out = run_drawing_impact_for_batch(state)
        assert out.drawing_impact_result is not None
        assert out.drawing_impact_result.impact_level == "substantial"
        assert seen["digest"] == digest
        assert seen["ids"] == ["rf-aaaaaaaaaaaa"]

    def test_failed_result_is_recorded_not_raised(self, monkeypatch):
        monkeypatch.setattr(
            di_pkg,
            "run_drawing_impact",
            lambda **kw: DrawingImpactResult(status="failed", error="network down"),
        )
        ctx = wrap_attachment(DIGEST_ATTACHMENT_LABEL, "DIGEST [p.1]")
        state = _state(ctx, [_finding()])
        out = run_drawing_impact_for_batch(state)
        assert out.drawing_impact_result.status == "failed"

    def test_finalize_carries_result(self):
        state = _state(
            wrap_attachment(DIGEST_ATTACHMENT_LABEL, "DIGEST"), [_finding()]
        )
        state.drawing_impact_result = DrawingImpactResult(
            status="completed", impact_level="moderate"
        )
        result = finalize_batch_result(state)
        assert result.drawing_impact_result is not None
        assert result.drawing_impact_result.impact_level == "moderate"


# ---------------------------------------------------------------------------
# 7. Report rendering
# ---------------------------------------------------------------------------


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _pipeline_result(drawing_impact=None, findings=None) -> PipelineResult:
    findings = findings if findings is not None else [_finding()]
    return PipelineResult(
        review_result=ReviewResult(findings=findings),
        files_reviewed=["230553.docx"],
        drawing_impact_result=drawing_impact,
    )


class TestReportRendering:
    def test_no_result_no_section(self, tmp_path: Path):
        out = tmp_path / "r.docx"
        export = _import_export_report()
        export(_pipeline_result(drawing_impact=None), out)
        text = _all_text(Document(str(out)))
        assert "How the Drawings Informed This Review" not in text
        assert "Drawing analysis impact" not in text  # banner row absent

    def test_completed_result_renders_section(self, tmp_path: Path):
        out = tmp_path / "r.docx"
        impact = DrawingImpactResult(
            status="completed",
            impact_level="substantial",
            narrative="The drawings surfaced a schedule the specs omitted.",
            finding_links=[
                di_pkg.DrawingFindingLink(
                    finding_id="rf-aaaaaaaaaaaa",
                    relationship="corroborated",
                    explanation="M-501 confirms the tag mismatch.",
                    sheet_references=["plans.pdf p.5"],
                )
            ],
        )
        export = _import_export_report()
        export(_pipeline_result(drawing_impact=impact), out)
        text = _all_text(Document(str(out)))
        assert "How the Drawings Informed This Review" in text
        assert "Substantial" in text
        assert "The drawings surfaced a schedule the specs omitted." in text
        assert "Corroborated by drawings" in text
        assert "rf-aaaaaaaaaaaa" in text
        assert "plans.pdf p.5" in text
        # Banner row present.
        assert "Drawing analysis impact" in text

    def test_no_links_renders_background_note(self, tmp_path: Path):
        out = tmp_path / "r.docx"
        impact = DrawingImpactResult(
            status="completed",
            impact_level="minimal",
            narrative="The drawings mostly restated the specs.",
            finding_links=[],
        )
        export = _import_export_report()
        export(_pipeline_result(drawing_impact=impact), out)
        text = _all_text(Document(str(out)))
        assert "How the Drawings Informed This Review" in text
        assert "No individual finding turned on drawing content" in text

    def test_failed_result_renders_honest_note(self, tmp_path: Path):
        out = tmp_path / "r.docx"
        impact = DrawingImpactResult(status="failed", error="rate limited")
        export = _import_export_report()
        export(_pipeline_result(drawing_impact=impact), out)
        text = _all_text(Document(str(out)))
        assert "How the Drawings Informed This Review" in text
        assert "did not complete" in text
        assert "still provided" in text  # honest: drawings still informed the review
        assert "analysis failed" in text  # banner row


def _import_export_report():
    # Imported lazily so a docx-less collection still imports the module-level
    # symbols above (export_report needs python-docx, always present here).
    from src.output.report_exporter import export_report

    return export_report
