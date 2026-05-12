"""Chunk 9 — DOCX edit-safety hardening and transactional apply.

Covers the safety contracts introduced in Chunk 9 of the repair plan:

- The unsafe-markup detector (:func:`detect_unsafe_markup`) flags every
  WordprocessingML construct in :data:`_UNSAFE_TAGS` and returns a
  human-readable refusal reason via :attr:`UnsafeMarkupResult.detail`.
- ``apply_edits_to_spec`` refuses to mutate paragraphs that carry
  hyperlinks, field codes, drawings, comments, tracked changes, bookmarks,
  or content controls. The corresponding :class:`EditOutcome` carries
  ``status="skipped"`` and ``refused_unsafe_markup=True``.
- Table-cell edits are also gated by unsafe-markup detection, and a global
  ``SPEC_CRITIC_TABLE_CELL_AUTO_EDIT=0`` switch refuses every table-cell
  auto-edit.
- ADD anchors and whole-paragraph DELETE targets are gated the same way.
- The transactional all-or-none policy suppresses the output write when at
  least one auto-edit failed; ``aborted_transactional=True`` is the
  visible signal.
- Successful edits still preserve the document text the locator pointed
  at (regression guard).
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.edit_locator import EditLocation, LocatorResult
from src.extractor import ParagraphMapping, extract_text_from_docx
from src.reviewer import Finding
from src.spec_editor import (
    UnsafeMarkupResult,
    apply_edits_to_spec,
    build_edit_actions,
    detect_unsafe_markup,
)
from tests.fixtures.docx_fixtures import (
    make_paragraph_spec,
    make_paragraph_with_bookmark,
    make_paragraph_with_comment_range,
    make_paragraph_with_drawing,
    make_paragraph_with_field_code,
    make_paragraph_with_hyperlink,
    make_paragraph_with_tracked_change,
    make_table_with_unsafe_cell,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    action: str = "EDIT",
    existing: str = "",
    replacement: str | None = "",
    severity: str = "HIGH",
    file_name: str = "spec.docx",
    anchor: str | None = None,
    insert_position: str | None = None,
) -> Finding:
    return Finding(
        severity=severity,
        fileName=file_name,
        section="1.0",
        issue="Issue",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="Code",
        confidence=0.9,
        anchorText=anchor,
        insertPosition=insert_position,
    )


def _build_locator_result(
    *,
    mapping: ParagraphMapping,
    match_start: int,
    match_end: int,
    matched_text: str,
    replacement_text: str | None,
    action: str = "EDIT",
    severity: str = "HIGH",
    confidence: float = 1.0,
    match_method: str = "exact",
) -> LocatorResult:
    location = EditLocation(
        mapping=mapping,
        match_start=match_start,
        match_end=match_end,
        matched_text=matched_text,
        match_confidence=confidence,
        match_method=match_method,
    )
    finding = _finding(
        action=action,
        existing=matched_text,
        replacement=replacement_text,
        severity=severity,
    )
    return LocatorResult(
        finding=finding,
        status="matched",
        locations=[location],
        replacement_text=replacement_text,
        action_type=action,
        warning=None,
    )


def _locate_substring(mapping: ParagraphMapping, needle: str) -> tuple[int, int]:
    idx = mapping.text.find(needle)
    assert idx >= 0, f"needle {needle!r} not in mapping text {mapping.text!r}"
    return idx, idx + len(needle)


# ---------------------------------------------------------------------------
# detect_unsafe_markup() unit tests
# ---------------------------------------------------------------------------


def test_detect_unsafe_markup_plain_paragraph_is_safe(tmp_path: Path):
    source = make_paragraph_spec(tmp_path, ["Plain content with no rich structure."])
    doc = Document(source)
    para = doc.paragraphs[0]
    result = detect_unsafe_markup(para._element)
    assert isinstance(result, UnsafeMarkupResult)
    assert result.unsafe is False
    assert result.reasons == ()
    assert result.detail == ""


def test_detect_unsafe_markup_hyperlink(tmp_path: Path):
    source = make_paragraph_with_hyperlink(tmp_path)
    doc = Document(source)
    # The second paragraph is the one that carries the hyperlink.
    para = doc.paragraphs[1]
    result = detect_unsafe_markup(para._element)
    assert result.unsafe is True
    assert "hyperlink" in result.reasons
    assert "hyperlink" in result.detail
    assert "manual review" in result.detail.lower()


def test_detect_unsafe_markup_field_code(tmp_path: Path):
    source = make_paragraph_with_field_code(tmp_path)
    doc = Document(source)
    para = doc.paragraphs[1]
    result = detect_unsafe_markup(para._element)
    assert result.unsafe is True
    assert "field character" in result.reasons
    # The instr text is also detected as a distinct unsafe reason.
    assert "field instruction text" in result.reasons


def test_detect_unsafe_markup_drawing(tmp_path: Path):
    source = make_paragraph_with_drawing(tmp_path)
    doc = Document(source)
    para = doc.paragraphs[1]
    result = detect_unsafe_markup(para._element)
    assert result.unsafe is True
    assert "drawing/image" in result.reasons


def test_detect_unsafe_markup_tracked_change(tmp_path: Path):
    source = make_paragraph_with_tracked_change(tmp_path)
    doc = Document(source)
    para = doc.paragraphs[1]
    result = detect_unsafe_markup(para._element)
    assert result.unsafe is True
    assert "tracked insertion" in result.reasons


def test_detect_unsafe_markup_comment_range(tmp_path: Path):
    source = make_paragraph_with_comment_range(tmp_path)
    doc = Document(source)
    para = doc.paragraphs[1]
    result = detect_unsafe_markup(para._element)
    assert result.unsafe is True
    assert "comment range" in result.reasons or "comment reference" in result.reasons


def test_detect_unsafe_markup_bookmark(tmp_path: Path):
    source = make_paragraph_with_bookmark(tmp_path)
    doc = Document(source)
    para = doc.paragraphs[1]
    result = detect_unsafe_markup(para._element)
    assert result.unsafe is True
    assert "bookmark range" in result.reasons


def test_detect_unsafe_markup_handles_none():
    """Defensive: None / objects without iter() return safe.

    The detector might be called on a body element that turned out not to
    be a paragraph (degenerate mapping); failing safe rather than raising
    keeps the apply pipeline robust.
    """
    assert detect_unsafe_markup(None).unsafe is False


# ---------------------------------------------------------------------------
# End-to-end refusal tests — apply_edits_to_spec()
# ---------------------------------------------------------------------------


def _apply_to_paragraph_with_text(
    tmp_path: Path,
    source: Path,
    paragraph_index: int,
    needle: str,
    replacement: str,
) -> tuple[Document, list]:
    """Re-extract the spec, build a single-edit action against the indicated
    paragraph, run apply_edits_to_spec, and return ``(saved_or_source_doc,
    report.outcomes)``. The output document is loaded from the saved file
    when one was written; otherwise from the source so callers can verify
    the original is untouched.
    """
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[paragraph_index]
    start, end = _locate_substring(mapping, needle)
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text=needle,
        replacement_text=replacement,
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    if output.exists():
        saved = Document(output)
    else:
        saved = Document(source)
    return saved, report


def test_normal_paragraph_edit_still_succeeds(tmp_path: Path):
    """Regression guard: Chunk 9 must not block the legitimate auto-edit path."""
    source = make_paragraph_spec(tmp_path, ["Provide ASCE 7-16 seismic bracing."])
    saved, _ = _apply_to_paragraph_with_text(
        tmp_path, source, 0, "ASCE 7-16", "ASCE 7-22"
    )
    assert saved.paragraphs[0].text == "Provide ASCE 7-22 seismic bracing."


def test_hyperlink_paragraph_refuses(tmp_path: Path):
    source = make_paragraph_with_hyperlink(tmp_path)
    spec = extract_text_from_docx(source)
    # Second paragraph in the map carries the hyperlink.
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "manufacturer datasheet")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="manufacturer datasheet",
        replacement_text="manufacturer cut sheet",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.edits_skipped == 1
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "hyperlink" in report.outcomes[0].detail.lower()
    # Source paragraph still contains the original text.
    saved = Document(output)
    assert "manufacturer datasheet" in saved.paragraphs[1].text


def test_field_code_paragraph_refuses(tmp_path: Path):
    source = make_paragraph_with_field_code(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "page number")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="page number",
        replacement_text="page reference",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "field" in report.outcomes[0].detail.lower()


def test_drawing_paragraph_refuses(tmp_path: Path):
    source = make_paragraph_with_drawing(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "Figure caption")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="Figure caption",
        replacement_text="Schematic caption",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "drawing" in report.outcomes[0].detail.lower()


def test_tracked_change_paragraph_refuses(tmp_path: Path):
    source = make_paragraph_with_tracked_change(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "proposed insertion")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="proposed insertion",
        replacement_text="confirmed addition",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "tracked" in report.outcomes[0].detail.lower()


def test_comment_range_paragraph_refuses(tmp_path: Path):
    source = make_paragraph_with_comment_range(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "reviewer comment range")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="reviewer comment range",
        replacement_text="updated comment range",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "comment" in report.outcomes[0].detail.lower()


def test_bookmark_paragraph_refuses(tmp_path: Path):
    source = make_paragraph_with_bookmark(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "bookmark range marker")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="bookmark range marker",
        replacement_text="navigation marker",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "bookmark" in report.outcomes[0].detail.lower()


# ---------------------------------------------------------------------------
# Table-cell unsafe-markup refusal
# ---------------------------------------------------------------------------


def test_unsafe_table_cell_refuses(tmp_path: Path):
    """A table cell whose paragraph carries a hyperlink is refused."""
    source = make_table_with_unsafe_cell(tmp_path)
    spec = extract_text_from_docx(source)
    # Find the table-cell mapping.
    cell_mapping = next(
        (m for m in spec.paragraph_map if m.element_type == "table_cell"),
        None,
    )
    assert cell_mapping is not None
    start, end = _locate_substring(cell_mapping, "manufacturer datasheet")
    result = _build_locator_result(
        mapping=cell_mapping,
        match_start=start,
        match_end=end,
        matched_text="manufacturer datasheet",
        replacement_text="manufacturer cut sheet",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "hyperlink" in report.outcomes[0].detail.lower()


def test_table_cell_auto_edit_disabled_refuses_even_safe_cells(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """Operator-set SPEC_CRITIC_TABLE_CELL_AUTO_EDIT=0 refuses any table cell."""
    # Build a plain-text table (no hyperlinks); without the override the
    # edit would succeed.
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "208V power"
    table.cell(0, 1).text = "Notes"
    source = tmp_path / "plain_table.docx"
    doc.save(source)
    spec = extract_text_from_docx(source)
    cell_mapping = next(
        (m for m in spec.paragraph_map if m.element_type == "table_cell"),
        None,
    )
    assert cell_mapping is not None
    start, end = _locate_substring(cell_mapping, "208V")
    result = _build_locator_result(
        mapping=cell_mapping,
        match_start=start,
        match_end=end,
        matched_text="208V",
        replacement_text="240V",
    )

    monkeypatch.setenv("SPEC_CRITIC_TABLE_CELL_AUTO_EDIT", "0")
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    assert "table-cell auto-edit is disabled" in report.outcomes[0].detail.lower()


# ---------------------------------------------------------------------------
# ADD anchor and DELETE refusal
# ---------------------------------------------------------------------------


def test_add_anchor_with_unsafe_markup_refuses(tmp_path: Path):
    """ADD beside a hyperlink-carrying anchor refuses to insert."""
    source = make_paragraph_with_hyperlink(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    location = EditLocation(
        mapping=mapping,
        match_start=0,
        match_end=len(mapping.text),
        matched_text=mapping.text,
        match_confidence=1.0,
        match_method="exact",
    )
    finding = _finding(
        action="ADD",
        existing=mapping.text,
        replacement="NEW PARAGRAPH AFTER ANCHOR.",
        anchor=mapping.text,
        insert_position="after",
    )
    result = LocatorResult(
        finding=finding,
        status="matched",
        locations=[location],
        replacement_text="NEW PARAGRAPH AFTER ANCHOR.",
        action_type="ADD",
        warning=None,
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    saved = Document(output)
    # No new paragraph was inserted next to the unsafe anchor.
    paragraph_texts = [p.text for p in saved.paragraphs]
    assert "NEW PARAGRAPH AFTER ANCHOR." not in paragraph_texts


def test_whole_paragraph_delete_with_unsafe_markup_refuses(tmp_path: Path):
    """Whole-paragraph DELETE on a hyperlink-carrying paragraph refuses."""
    source = make_paragraph_with_hyperlink(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    location = EditLocation(
        mapping=mapping,
        match_start=0,
        match_end=len(mapping.text),
        matched_text=mapping.text,
        match_confidence=1.0,
        match_method="exact",
    )
    finding = _finding(action="DELETE", existing=mapping.text, replacement=None)
    result = LocatorResult(
        finding=finding,
        status="matched",
        locations=[location],
        replacement_text=None,
        action_type="DELETE",
        warning=None,
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 0
    assert report.outcomes[0].refused_unsafe_markup is True
    saved = Document(output)
    # Paragraph is still there.
    assert any(
        "manufacturer datasheet" in p.text for p in saved.paragraphs
    )


# ---------------------------------------------------------------------------
# Transactional all-or-none output
# ---------------------------------------------------------------------------


def test_failed_edit_does_not_produce_partially_mutated_output(tmp_path: Path):
    """Failed edit in batch → output is suppressed; applied edits are demoted.

    Build two edits against the same source:
      * one valid (would succeed in isolation),
      * one invalid (offsets out of range → ``failed``).

    With the default all-or-none policy in place, the output file must
    contain none of the changes from either edit. The valid edit's outcome
    is demoted to ``skipped`` with a clear "Output suppressed" detail; the
    invalid edit retains its ``failed`` status. The report carries
    ``aborted_transactional=True``.
    """
    source = make_paragraph_spec(
        tmp_path,
        ["Provide ASCE 7-16 seismic bracing.", "Comply with CBC 2019."],
    )
    spec = extract_text_from_docx(source)

    good_mapping = spec.paragraph_map[0]
    good_start, good_end = _locate_substring(good_mapping, "ASCE 7-16")
    good = _build_locator_result(
        mapping=good_mapping,
        match_start=good_start,
        match_end=good_end,
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )

    bad_mapping = spec.paragraph_map[1]
    # Pick an out-of-range body_index so the apply path takes the
    # explicit ``"failed"`` branch ("Body index is out of range") instead
    # of the precondition-skip branch.
    bad_mapping_out_of_range = ParagraphMapping(
        body_index=9999,
        element_type=bad_mapping.element_type,
        text=bad_mapping.text,
        table_index=bad_mapping.table_index,
        row_index=bad_mapping.row_index,
        cell_index=bad_mapping.cell_index,
    )
    location_bad = EditLocation(
        mapping=bad_mapping_out_of_range,
        match_start=0,
        match_end=len(bad_mapping.text),
        matched_text=bad_mapping.text,
        match_confidence=1.0,
        match_method="exact",
    )
    bad = LocatorResult(
        finding=_finding(
            existing=bad_mapping.text,
            replacement="X",
        ),
        status="matched",
        locations=[location_bad],
        replacement_text="X",
        action_type="EDIT",
        warning=None,
    )

    output = tmp_path / "output.docx"
    # The locator-safety classifier might downgrade ``bad`` to manual
    # review because its offsets are nonsensical; force-feed the action
    # plan directly so build_edit_actions cannot drop it.
    from src.spec_editor import EditAction

    actions = [
        EditAction(
            locator_result=good,
            location=good.locations[0],
            replacement_text=good.replacement_text,
            action_type=good.action_type,
            finding_index=0,
        ),
        EditAction(
            locator_result=bad,
            location=bad.locations[0],
            replacement_text=bad.replacement_text,
            action_type=bad.action_type,
            finding_index=1,
        ),
    ]
    report = apply_edits_to_spec(source, output, actions)

    assert report.aborted_transactional is True
    # No edits considered "applied" because the policy demoted them.
    assert report.edits_applied == 0
    # The output file is either not written or, if written for the reopen
    # validation step, the disk file content must NOT contain the would-be
    # mutation. (The implementation suppresses the disk write entirely on
    # all-or-none abort.)
    if output.exists():
        saved = Document(output)
        # If a previous test created the file in this tmpdir we are still
        # safe: the all-or-none abort guarantees the source content was
        # never written to ``output`` in this run, so we re-open the source
        # instead to verify the intent.
        # The all-or-none branch never calls write_bytes, so the file
        # should not exist in a clean tmp_path.
        # Be lenient — just verify the mutation is absent if a file exists.
        first_para = saved.paragraphs[0].text
        assert "ASCE 7-22" not in first_para
    else:
        # Standard case: output never written.
        pass

    # The originally-valid edit is reported as skipped with the suppress
    # detail.
    detail_blob = " ".join(o.detail for o in report.outcomes)
    assert "Output suppressed" in detail_blob
    # At least one outcome retained the failed status from the bad edit.
    assert any(o.status == "failed" for o in report.outcomes)


def test_best_effort_mode_writes_partial_output(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """SPEC_CRITIC_EDIT_TRANSACTIONAL=0 reverts to legacy best-effort writes.

    Operators who explicitly opt out of the all-or-none policy get the
    pre-Chunk-9 behavior: the output is written even when some edits
    failed. This guards the opt-out path itself.
    """
    monkeypatch.setenv("SPEC_CRITIC_EDIT_TRANSACTIONAL", "0")

    source = make_paragraph_spec(tmp_path, ["Provide ASCE 7-16 seismic bracing."])
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[0]
    start, end = _locate_substring(mapping, "ASCE 7-16")
    good = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )

    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([good]))
    assert report.aborted_transactional is False
    assert report.edits_applied == 1
    assert output.exists()
    saved = Document(output)
    assert "ASCE 7-22" in saved.paragraphs[0].text


def test_successful_edit_preserves_unrelated_formatting(tmp_path: Path):
    """Regression: Chunk 9 changes do not corrupt unrelated runs.

    Build a paragraph with two runs (the second carrying bold formatting),
    edit the first run, and verify the second run's run-text is still
    present in the saved file.
    """
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Provide ASCE 7-16 ")
    bold_run = para.add_run("seismic bracing")
    bold_run.bold = True
    para.add_run(" per code.")
    source = tmp_path / "formatted.docx"
    doc.save(source)

    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[0]
    start, end = _locate_substring(mapping, "ASCE 7-16")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.edits_applied == 1
    saved = Document(output)
    # Replacement landed and the bold run survives.
    paragraph = saved.paragraphs[0]
    assert "ASCE 7-22" in paragraph.text
    assert "seismic bracing" in paragraph.text
    # Find the run that still carries bold formatting.
    bold_runs = [r for r in paragraph.runs if r.bold]
    assert bold_runs, "expected at least one bold run to survive the edit"
    assert any("seismic" in r.text for r in bold_runs)


def test_edit_report_warning_lists_unsafe_refusals(tmp_path: Path):
    """Refusal details surface in EditReport.warnings for the report layer."""
    source = make_paragraph_with_hyperlink(tmp_path)
    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[1]
    start, end = _locate_substring(mapping, "manufacturer datasheet")
    result = _build_locator_result(
        mapping=mapping,
        match_start=start,
        match_end=end,
        matched_text="manufacturer datasheet",
        replacement_text="manufacturer cut sheet",
    )
    output = tmp_path / "output.docx"
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.warnings, "expected an unsafe-markup warning to surface"
    assert any("hyperlink" in w.lower() for w in report.warnings)
