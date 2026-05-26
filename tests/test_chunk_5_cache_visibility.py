"""Chunk 5 tests — cache-replay visibility (age badge + TTL default + hint).

Chunk 5 of the Trust Upgrade surfaces cache-replay information so a
reviewer can spot stale verdicts at a glance and force re-verification
when needed. The contract has four surfaces:

* ``VerificationResult.cache_entry_created_ts`` is populated by
  :func:`verification_cache._clone_for_hit` (sourced from the sidecar
  ``_CacheEntry.created_ts``). Default 0.0 means "not from a cache hit."
* ``cache_ttl_days()`` defaults to 60 days when unset. Explicit ``0``
  preserves the legacy "no expiry" behavior; malformed values fall back
  to the 60-day default.
* ``_write_finding_entry`` renders an inline "Cache replay — Nd old"
  badge for cache-hit findings, color-tiered (amber <30d, orange
  30-90d, red >90d).
* The evidence panel surfaces a one-line hint pointing at the cache
  file path so reviewers know exactly where to delete an entry to
  trigger re-verification.

Resume state round-trip is verified separately because a resumed report
must render the same badge the original run would have shown.
"""
from __future__ import annotations

import time
from pathlib import Path

import pytest
from docx import Document

from src.core.code_cycles import DEFAULT_CYCLE
from src.orchestration.resume_state import (
    deserialize_verification_result,
    serialize_verification_result,
)
from src.output.report_exporter import (
    CACHE_AGE_COLORS,
    _cache_age_tier,
    _cache_entry_age_days,
    export_report,
)
from src.review.reviewer import Finding, ReviewResult
from src.verification.verification_cache import (
    VerificationCache,
    _DEFAULT_CACHE_TTL_DAYS,
    cache_ttl_days,
    default_cache_path,
)
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.8,
    verification: VerificationResult | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType="EDIT",
        existingText="2019 CBC",
        replacementText="2025 CBC",
        codeReference="CBC §1234",
        confidence=confidence,
    )
    f.verification = verification
    return f


def _cache_hit_result(
    *,
    age_days: int = 5,
    verdict: str = "CONFIRMED",
) -> VerificationResult:
    """Return a VerificationResult that looks like a cache hit of the given age."""
    created_ts = time.time() - (age_days * 86400)
    return VerificationResult(
        verdict=verdict,
        explanation="Cached verdict from prior run.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        accepted_sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
        model_used="claude-sonnet-4-6",
        cache_status="hit",
        source_quote="The 2025 California Building Code adopts the 2024 IBC.",
        cache_entry_created_ts=created_ts,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = None
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 1. VerificationResult.cache_entry_created_ts field
# ---------------------------------------------------------------------------


class TestVerificationResultCacheTimestamp:
    def test_default_is_zero(self):
        result = VerificationResult(verdict="CONFIRMED")
        assert result.cache_entry_created_ts == 0.0

    def test_field_round_trips_through_constructor(self):
        ts = time.time()
        result = VerificationResult(
            verdict="CONFIRMED",
            cache_entry_created_ts=ts,
        )
        assert result.cache_entry_created_ts == ts


# ---------------------------------------------------------------------------
# 2. Cache _clone_for_hit stamps created_ts on the result
# ---------------------------------------------------------------------------


class TestCloneForHitStampsTimestamp:
    def _finding_for_cache(self) -> Finding:
        return Finding(
            severity="HIGH",
            fileName="Section_22_1000.docx",
            section="2.1",
            issue="claim about NFPA 13",
            actionType="EDIT",
            existingText="per NFPA 13 (2019)",
            replacementText="per NFPA 13 (2022)",
            codeReference="NFPA 13 §10",
            confidence=0.7,
        )

    def test_cache_hit_carries_entry_creation_timestamp(self):
        # Put a grounded verdict, then immediately fetch it. The hit's
        # cache_entry_created_ts must equal the entry's stored created_ts
        # (which is approximately time.time() at put-time).
        cache = VerificationCache()
        f = self._finding_for_cache()
        before_put = time.time()
        cache.put(
            f,
            cycle=DEFAULT_CYCLE,
            result=VerificationResult(
                verdict="CONFIRMED",
                grounded=True,
                sources=["https://nfpa.org/"],
                accepted_sources=["https://nfpa.org/"],
                source_quote="snippet",
            ),
        )
        after_put = time.time()
        hit = cache.get(f, cycle=DEFAULT_CYCLE)
        assert hit is not None
        assert hit.cache_status == "hit"
        # The stamped timestamp must fall in the [before_put, after_put]
        # window — proves the value is the entry's created_ts, not the
        # default 0.0 or the time of the get() call.
        assert before_put <= hit.cache_entry_created_ts <= after_put

    def test_cache_miss_does_not_set_timestamp(self):
        # A miss returns None; the caller constructs a fresh result on
        # the miss path which keeps the 0.0 default. This test confirms
        # the cache itself doesn't fabricate a non-zero timestamp on
        # miss.
        cache = VerificationCache()
        f = self._finding_for_cache()
        assert cache.get(f, cycle=DEFAULT_CYCLE) is None

    def test_timestamp_survives_save_and_load(self, tmp_path: Path):
        # When the cache persists to disk and reloads, the entry's
        # created_ts is preserved. The next hit reflects that older
        # timestamp, which is the contract that lets the badge color
        # an entry as stale.
        cache_path = tmp_path / "cache.json"
        cache = VerificationCache()
        f = self._finding_for_cache()
        cache.put(
            f,
            cycle=DEFAULT_CYCLE,
            result=VerificationResult(
                verdict="CONFIRMED",
                grounded=True,
                sources=["https://nfpa.org/"],
                accepted_sources=["https://nfpa.org/"],
                source_quote="snippet",
            ),
        )
        cache.save_to_disk(cache_path)
        first_hit = cache.get(f, cycle=DEFAULT_CYCLE)
        assert first_hit is not None
        stored_ts = first_hit.cache_entry_created_ts

        reloaded = VerificationCache()
        # Use a TTL of 0 so the stored entry is not pruned on load.
        import os
        os.environ["SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS"] = "0"
        try:
            reloaded.load_from_disk(cache_path)
        finally:
            os.environ.pop("SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS", None)
        second_hit = reloaded.get(f, cycle=DEFAULT_CYCLE)
        assert second_hit is not None
        assert second_hit.cache_status == "hit"
        # The reloaded entry's timestamp should equal (within float
        # precision) the original entry's timestamp — disk round-trip
        # is lossless.
        assert abs(second_hit.cache_entry_created_ts - stored_ts) < 1.0


# ---------------------------------------------------------------------------
# 3. cache_ttl_days default + explicit overrides
# ---------------------------------------------------------------------------


class TestCacheTtlDefault:
    def test_default_is_sixty(self, monkeypatch: pytest.MonkeyPatch):
        monkeypatch.delenv("SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS", raising=False)
        assert cache_ttl_days() == _DEFAULT_CACHE_TTL_DAYS == 60

    def test_explicit_zero_means_no_expiry(self, monkeypatch: pytest.MonkeyPatch):
        # The plan calls this out explicitly: operators can opt out of
        # the new default by setting the env var to "0".
        monkeypatch.setenv("SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS", "0")
        assert cache_ttl_days() == 0

    def test_positive_override(self, monkeypatch: pytest.MonkeyPatch):
        monkeypatch.setenv("SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS", "14")
        assert cache_ttl_days() == 14

    def test_malformed_falls_back_to_default(self, monkeypatch: pytest.MonkeyPatch):
        # Chunk 5 / Trust Upgrade: previous behavior fell back to 0 for
        # malformed values, which silently disabled expiry on any typo.
        # New behavior falls back to the 60-day default so a typo never
        # turns the cache into a permanent database by accident.
        monkeypatch.setenv("SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS", "garbage")
        assert cache_ttl_days() == 60

    def test_negative_falls_back_to_default(self, monkeypatch: pytest.MonkeyPatch):
        monkeypatch.setenv("SPEC_CRITIC_VERIFICATION_CACHE_TTL_DAYS", "-1")
        assert cache_ttl_days() == 60


# ---------------------------------------------------------------------------
# 4. Age-tier classifier
# ---------------------------------------------------------------------------


class TestCacheAgeTier:
    def test_fresh_below_thirty_days(self):
        assert _cache_age_tier(0) == "fresh"
        assert _cache_age_tier(15) == "fresh"
        assert _cache_age_tier(29) == "fresh"

    def test_stale_thirty_to_ninety_days(self):
        assert _cache_age_tier(30) == "stale"
        assert _cache_age_tier(60) == "stale"
        assert _cache_age_tier(90) == "stale"

    def test_very_stale_above_ninety_days(self):
        assert _cache_age_tier(91) == "very_stale"
        assert _cache_age_tier(180) == "very_stale"
        assert _cache_age_tier(365) == "very_stale"

    def test_each_tier_has_distinct_color(self):
        # All three tiers must map to distinct RGB values so a reader
        # can tell them apart visually.
        colors = {CACHE_AGE_COLORS[t] for t in ("fresh", "stale", "very_stale")}
        assert len(colors) == 3


# ---------------------------------------------------------------------------
# 5. _cache_entry_age_days helper
# ---------------------------------------------------------------------------


class TestCacheEntryAgeDays:
    def test_returns_age_for_cache_hit(self):
        vr = _cache_hit_result(age_days=10)
        assert _cache_entry_age_days(vr) == 10

    def test_returns_none_for_non_hit(self):
        # A miss-status result must never produce a badge, even if the
        # timestamp field somehow has a value.
        vr = VerificationResult(
            verdict="CONFIRMED",
            cache_status="miss",
            cache_entry_created_ts=time.time() - 86400,
        )
        assert _cache_entry_age_days(vr) is None

    def test_returns_none_for_local_skip(self):
        vr = VerificationResult(
            verdict="UNVERIFIED",
            cache_status="local_skip",
            cache_entry_created_ts=0.0,
        )
        assert _cache_entry_age_days(vr) is None

    def test_returns_none_for_legacy_zero_timestamp(self):
        # Legacy resume payloads predating Chunk 5 have
        # cache_entry_created_ts=0.0. The badge is suppressed (rather
        # than rendered as a nonsense epoch-1970 age).
        vr = VerificationResult(
            verdict="CONFIRMED",
            cache_status="hit",
            cache_entry_created_ts=0.0,
        )
        assert _cache_entry_age_days(vr) is None

    def test_returns_none_for_future_timestamp(self):
        # Clock skew anomaly: a timestamp in the future must not
        # produce a negative-age badge. Suppress instead.
        vr = VerificationResult(
            verdict="CONFIRMED",
            cache_status="hit",
            cache_entry_created_ts=time.time() + 86400,
        )
        assert _cache_entry_age_days(vr) is None

    def test_returns_none_for_missing_verification(self):
        # Defensive: a finding without a verification object renders
        # no badge.
        assert _cache_entry_age_days(None) is None


# ---------------------------------------------------------------------------
# 6. Resume state round-trips cache_entry_created_ts
# ---------------------------------------------------------------------------


class TestResumeStateRoundTrip:
    def test_serialize_includes_timestamp(self):
        ts = time.time() - 86400
        result = VerificationResult(
            verdict="CONFIRMED",
            cache_status="hit",
            cache_entry_created_ts=ts,
        )
        payload = serialize_verification_result(result)
        assert payload is not None
        assert payload["cache_entry_created_ts"] == pytest.approx(ts)

    def test_deserialize_restores_timestamp(self):
        ts = time.time() - 86400
        payload = {
            "verdict": "CONFIRMED",
            "explanation": "Cached.",
            "sources": ["https://x"],
            "correction": None,
            "grounded": True,
            "model_used": "claude-sonnet-4-6",
            "escalated": False,
            "cache_status": "hit",
            "web_search_requests": 0,
            "successful_source_count": 0,
            "search_error_count": 0,
            "searched_sources": [],
            "cited_sources": [],
            "accepted_sources": ["https://x"],
            "rejected_sources": [],
            "verification_profile": "",
            "verification_mode": "",
            "source_quote": "",
            "verification_failed": False,
            "cache_entry_created_ts": ts,
        }
        result = deserialize_verification_result(payload)
        assert result is not None
        assert result.cache_entry_created_ts == pytest.approx(ts)

    def test_deserialize_legacy_payload_defaults_zero(self):
        # State files written before Chunk 5 don't carry the field.
        # Missing → 0.0 (badge suppressed; safe fallback when the
        # original timestamp was never recorded).
        payload = {
            "verdict": "CONFIRMED",
            "explanation": "Cached.",
            "sources": ["https://x"],
            "correction": None,
            "grounded": True,
            "model_used": "claude-sonnet-4-6",
            "escalated": False,
            "cache_status": "hit",
            "web_search_requests": 0,
            "successful_source_count": 0,
            "search_error_count": 0,
            "searched_sources": [],
            "cited_sources": [],
            "accepted_sources": ["https://x"],
            "rejected_sources": [],
            "verification_profile": "",
            "verification_mode": "",
            "source_quote": "",
            "verification_failed": False,
            # cache_entry_created_ts intentionally omitted
        }
        result = deserialize_verification_result(payload)
        assert result is not None
        assert result.cache_entry_created_ts == 0.0


# ---------------------------------------------------------------------------
# 7. Report rendering — badge appears on status line for cache hits
# ---------------------------------------------------------------------------


class TestReportBadgeRendering:
    # Chunk 6 / Trust Upgrade: the badge format is "Cache replay — Nd
    # old" (em-dash, age suffix). The Run Diagnostics banner introduced
    # in Chunk 6 also surfaces a "Cache replays" row, so assertions key
    # off the em-dash + age form rather than the bare "Cache replay"
    # substring to avoid colliding with the banner label.
    _BADGE_PREFIX = "Cache replay — "

    def test_fresh_cache_hit_renders_badge(self, tmp_path: Path):
        f = _finding(verification=_cache_hit_result(age_days=10))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert self._BADGE_PREFIX in text
        assert "10d old" in text

    def test_stale_cache_hit_renders_badge_with_age(self, tmp_path: Path):
        f = _finding(verification=_cache_hit_result(age_days=45))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert self._BADGE_PREFIX in text
        assert "45d old" in text

    def test_very_stale_cache_hit_renders_badge_with_age(self, tmp_path: Path):
        f = _finding(verification=_cache_hit_result(age_days=120))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert self._BADGE_PREFIX in text
        assert "120d old" in text

    def test_miss_does_not_render_badge(self, tmp_path: Path):
        vr = VerificationResult(
            verdict="CONFIRMED",
            explanation="Verified.",
            sources=["https://x"],
            accepted_sources=["https://x"],
            grounded=True,
            cache_status="miss",
            source_quote="snippet",
        )
        f = _finding(verification=vr)
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert self._BADGE_PREFIX not in text

    def test_local_skip_does_not_render_badge(self, tmp_path: Path):
        vr = VerificationResult(
            verdict="UNVERIFIED",
            explanation="Locally classified.",
            cache_status="local_skip",
        )
        f = _finding(verification=vr)
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert self._BADGE_PREFIX not in text

    def test_legacy_cache_hit_without_timestamp_does_not_render_badge(
        self, tmp_path: Path
    ):
        # A legacy resume payload may carry cache_status="hit" but
        # cache_entry_created_ts=0.0. The badge is suppressed rather
        # than displaying a nonsense age.
        vr = VerificationResult(
            verdict="CONFIRMED",
            explanation="Cached (legacy).",
            sources=["https://x"],
            accepted_sources=["https://x"],
            grounded=True,
            cache_status="hit",
            source_quote="snippet",
            cache_entry_created_ts=0.0,
        )
        f = _finding(verification=vr)
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert self._BADGE_PREFIX not in text


# ---------------------------------------------------------------------------
# 8. Report rendering — force-refresh hint appears for cache hits
# ---------------------------------------------------------------------------


class TestForceRefreshHint:
    def test_cache_hit_renders_force_refresh_hint(self, tmp_path: Path):
        # The hint must include the cache file path so a reviewer can
        # locate and delete the entry without consulting docs.
        f = _finding(verification=_cache_hit_result(age_days=5))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "To force re-verification" in text
        # The actual cache path is configured via env; default points to
        # ~/.spec_critic/verification_cache.json. The hint must contain
        # the resolved path so a reviewer doesn't have to guess.
        assert str(default_cache_path()) in text

    def test_miss_does_not_render_force_refresh_hint(self, tmp_path: Path):
        vr = VerificationResult(
            verdict="CONFIRMED",
            explanation="Verified.",
            sources=["https://x"],
            accepted_sources=["https://x"],
            grounded=True,
            cache_status="miss",
            source_quote="snippet",
        )
        f = _finding(verification=vr)
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "To force re-verification" not in text

    def test_hint_uses_configured_cache_path(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        # When the operator sets SPEC_CRITIC_CACHE_PATH, the hint must
        # point at the configured location, not the hardcoded default.
        custom = tmp_path / "custom_cache.json"
        monkeypatch.setenv("SPEC_CRITIC_CACHE_PATH", str(custom))
        f = _finding(verification=_cache_hit_result(age_days=5))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert str(custom) in text


# ---------------------------------------------------------------------------
# 9. Diagnostics — cache_hits / cache_misses counter (already wired)
# ---------------------------------------------------------------------------


class TestDiagnosticsCacheCounter:
    """Plan section 5d says ``DiagnosticsReport`` already tracks cache
    hits/misses via ``verification_evidence`` and they make it into the
    run summary. This test guards against an accidental drop in the
    rollup output."""

    def test_cache_hits_and_misses_present_in_summary(self):
        from src.orchestration.diagnostics import DiagnosticsReport

        report = DiagnosticsReport()
        report.log(
            "verification",
            "info",
            "cache hit",
            data={"verdict": "CONFIRMED", "cache_status": "hit"},
        )
        report.log(
            "verification",
            "info",
            "cache miss",
            data={"verdict": "CONFIRMED", "cache_status": "miss"},
        )
        summary = report.summary()
        evidence = summary["verification_evidence"]
        assert evidence["cache_hits"] == 1
        assert evidence["cache_misses"] == 1
