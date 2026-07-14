"""Behavior pins for the ``datacenter_fire`` module's routing + report surfaces.

WS-1 of ``docs/hyperscale_datacenter_module_plan.md``. Complements the byte-exact
prompt goldens (``test_golden_datacenter_surfaces.py``) with decision pins:
verification routing (a CRITICAL fire-marshal / FM-Global finding rides the
deep-reasoning path under the DC cycle, and the *same* finding routes
differently under the default California cycle), cross-check chunk assignment,
and the domain-worded report surfaces (DC title / phrase / pinned editions,
generic jurisdiction-free cycle wording — never California's).

Hermetic: no API key, no network.
"""
from __future__ import annotations

import pytest

from src.modules import DATACENTER_FIRE
from src.review.reviewer import Finding


def _finding(issue: str, *, severity: str = "CRITICAL", filename: str = "21 13 13 - Wet.docx") -> Finding:
    return Finding(
        severity=severity,
        fileName=filename,
        section="1.01",
        issue=issue,
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference=None,
    )


# ---------------------------------------------------------------------------
# Verification routing under the DC cycle
# ---------------------------------------------------------------------------


class TestDatacenterRouting:
    def test_fire_marshal_critical_routes_jurisdictional_deep(self):
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        finding = _finding("The fire marshal requires a witnessed acceptance test before occupancy.")
        routed = select_routing(finding, local_skip=False, cycle=DATACENTER_FIRE.cycle)
        assert routed.profile is VerificationProfile.JURISDICTIONAL
        assert routed.mode is VerificationMode.DEEP_REASONING

    def test_fm_global_critical_routes_jurisdictional_deep(self):
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        finding = _finding("The FM Global data sheet requirement conflicts with the specified sprinkler density.")
        routed = select_routing(finding, local_skip=False, cycle=DATACENTER_FIRE.cycle)
        assert routed.profile is VerificationProfile.JURISDICTIONAL
        assert routed.mode is VerificationMode.DEEP_REASONING

    def test_same_finding_routes_differently_under_california(self):
        # The DC jurisdictional vocabulary knows "fire marshal"; the default
        # California vocabulary does not, so the same finding is a plain
        # constructability claim there (standard reasoning, not deep).
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        finding = _finding("The fire marshal requires a witnessed acceptance test before occupancy.")
        default_routed = select_routing(finding, local_skip=False)
        assert default_routed.profile is VerificationProfile.CONSTRUCTABILITY
        assert default_routed.mode is VerificationMode.STANDARD_REASONING

    def test_manufacturer_and_code_standard_keywords_classify(self):
        from src.verification.verification_profiles import (
            VerificationProfile,
            classify_finding_profile,
        )

        kw = DATACENTER_FIRE.profile_keywords
        assert (
            classify_finding_profile(
                _finding("Provide a Viking model VK-100 per the datasheet.", severity="MEDIUM"),
                keywords=kw,
            )
            is VerificationProfile.MANUFACTURER
        )
        assert (
            classify_finding_profile(
                _finding("Cites an NFPA standard section that does not exist.", severity="HIGH"),
                keywords=kw,
            )
            is VerificationProfile.CODE_STANDARD
        )


# ---------------------------------------------------------------------------
# Cross-check chunk assignment
# ---------------------------------------------------------------------------


class TestDatacenterChunkAssignment:
    def test_module_groups_drive_assignment(self):
        from src.cross_check.cross_checker import _assign_chunk, _chunk_label

        groups = DATACENTER_FIRE.cross_check_chunk_groups
        assert _assign_chunk("21 13 13 - Wet-Pipe Sprinkler Systems.docx", groups) == "div_21"
        assert _assign_chunk("28 31 00 - Fire Detection and Alarm.docx", groups) == "div_28"
        assert _assign_chunk("22 11 16 - Fire Water Service.docx", groups) == "div_22"
        # Division 23 has no DC chunk group -> pools into the reserved general.
        assert _assign_chunk("23 05 00 - Common HVAC.docx", groups) == "general"
        assert _chunk_label("div_28", groups) == "Division 28 — Fire Detection & Alarm"


# ---------------------------------------------------------------------------
# Report surfaces
# ---------------------------------------------------------------------------


class TestDatacenterReportSurfaces:
    def test_methodology_note_renders_dc_phrase_and_pinned_editions(self):
        from docx import Document

        from src.output.report_exporter import _write_methodology_note

        doc = Document()
        _write_methodology_note(
            doc, cycle_label=DATACENTER_FIRE.cycle.label, module=DATACENTER_FIRE
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "relevant to hyperscale data-center fire protection projects." in text
        # No jurisdiction label -> generic cycle sentence (no California).
        assert "This review used dc-ibc-2024 code cycle references." in text
        # The pinned-editions paragraph renders the DC cycle's own standards,
        # never California's.
        assert "per the dc-ibc-2024 cycle:" in text
        assert "NFPA 13 2022" in text
        assert "NFPA 855 2023" in text
        assert "California" not in text

    def test_title_block_renders_fire_protection_title_and_bare_cycle(self):
        from docx import Document

        from src.output.report_exporter import _write_title_block

        class _StubReview:
            model = "claude-opus-4-8"

        doc = Document()
        _write_title_block(
            doc,
            _StubReview(),
            ["21 13 13 - Wet.docx"],
            cycle_label=DATACENTER_FIRE.cycle.label,
            module=DATACENTER_FIRE,
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Spec Critic — Fire Protection Specification Review Report" in text
        # jurisdiction_label is empty -> bare "Code Cycle: dc-ibc-2024".
        assert "Code Cycle: dc-ibc-2024" in text
        assert "California" not in text

    def test_alerts_render_generic_jurisdiction_free_headings(self):
        from docx import Document

        from src.output.report_exporter import _write_alerts

        doc = Document()
        _write_alerts(
            doc,
            [],
            [],
            code_cycle_alerts=[{"filename": "21 13 13 - Wet.docx", "context": "2018 IBC"}],
            invalid_code_cycle_alerts=[{"filename": "21 13 13 - Wet.docx", "context": "2019 IBC"}],
            module=DATACENTER_FIRE,
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        # No jurisdiction word in the headings.
        assert "Stale Code Cycle References" in text
        assert "Invalid Code Cycle Years" in text
        assert "California" not in text
        # Published I-code years come from plausible_cycle_years.
        assert "2009, 2012, 2015, 2018, 2021, 2024" in text
