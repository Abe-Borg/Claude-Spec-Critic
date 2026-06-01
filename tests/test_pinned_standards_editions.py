"""Tests for pinned standards editions.

Pins the NFPA, ASHRAE, IAPMO, and UL editions adopted by California for the
current cycle. The contract spans four surfaces:

* ``CodeCycle`` carries the pinned editions in a single ``standards`` tuple of
  :class:`StandardEdition` (replacing the legacy flat ``nfpa13`` … fields), with
  ``CALIFORNIA_2025`` populated from the adoption matrix.
* The reviewer system prompt and user message reference the pinned editions so
  the model knows which editions to compare against.
* The verifier system prompt renders a "Pinned standards editions" block before
  the search budget, with explicit instructions to flag edition drift.
* The methodology note in the exported report enumerates the pinned editions so
  reviewers see which editions drove the verdicts.

The editions for the fire-protection standards (NFPA 13/14/20/24/25/72) are
locked to the California Fire Code 2025, Chapter 80 adoption table.
"""
from __future__ import annotations

from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document

from src.core.code_cycles import (
    AVAILABLE_CYCLES,
    CALIFORNIA_2025,
    CodeCycle,
    DEFAULT_CYCLE,
    StandardEdition,
)
from src.output.report_exporter import (
    _render_pinned_editions_note,
    export_report,
)
from src.review.prompts import get_single_spec_user_message, get_system_prompt
from src.review.reviewer import Finding, ReviewResult
from src.verification.verifier import (
    _get_verification_system_prompt,
    _pinned_standards_lines,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(file: str = "test.docx") -> Finding:
    return Finding(
        severity="HIGH",
        fileName=file,
        section="2.1",
        issue="Stale code reference",
        actionType="EDIT",
        existingText="2019 CBC",
        replacementText="2025 CBC",
        codeReference="CBC §1234",
        confidence=0.8,
    )


def _bare_cycle(label: str = "bare") -> CodeCycle:
    """A cycle with the code years populated but no pinned standards."""
    return CodeCycle(
        label=label,
        cbc="2025",
        cmc="2025",
        cpc="2025",
        energy_code="2025",
        calgreen="2025",
        asce7="7-22",
        asce7_previous="7-16",
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = None
        self.files_reviewed = (
            [review_result.findings[0].fileName] if review_result.findings else ["test.docx"]
        )
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ===========================================================================
# 1. StandardEdition rendering
# ===========================================================================


class TestStandardEditionRender:
    def test_plain_edition_renders_year_only(self):
        std = StandardEdition("NFPA 14", "2024")
        assert std.edition_phrase == "2024"
        assert std.description == "NFPA 14 2024"

    def test_ca_amended_renders_one_directional_phrase(self):
        # The phrasing must make clear California amends the standard, not the
        # reverse — "<edition>, as amended by California".
        std = StandardEdition("NFPA 13", "2025", ca_amended=True)
        assert std.edition_phrase == "2025, as amended by California"
        assert std.description == "NFPA 13 2025, as amended by California"

    def test_ca_edition_note_renders_inline(self):
        std = StandardEdition("NFPA 25", "2013", ca_amended=True, note="California Edition")
        assert std.edition_phrase == "2013 California Edition"

    def test_note_without_amendment_renders_parenthetical(self):
        std = StandardEdition("UL 300", "2005", note="revised")
        assert std.edition_phrase == "2005 (revised)"

    def test_is_verified_reflects_source(self):
        assert StandardEdition("NFPA 13", "2025", source="CFC 2025, Ch. 80").is_verified
        assert not StandardEdition("ASHRAE 90.1", "2022", source="UNVERIFIED: confirm").is_verified
        # No source at all is treated as not verified.
        assert not StandardEdition("X", "1").is_verified


# ===========================================================================
# 2. CodeCycle carries the standards collection
# ===========================================================================


class TestCodeCycleFields:
    def test_codecycle_has_standards_for_each_org(self):
        names = {std.name for std in CALIFORNIA_2025.standards}
        for required in (
            "NFPA 13", "NFPA 14", "NFPA 20", "NFPA 24", "NFPA 25", "NFPA 72",
            "ASHRAE 62.1", "ASHRAE 90.1", "ASHRAE 15",
            "IAPMO Uniform Plumbing TSC",
            "UL 300", "UL 555", "UL 555S", "UL 268", "UL 1479",
        ):
            assert required in names, f"missing pinned standard: {required}"

    def test_nfpa_editions_match_cfc_2025_chapter_80(self):
        # Locks in the fact-checked corrections against the California Fire
        # Code 2025, Chapter 80 referenced-standards table.
        expected = {
            "NFPA 13": "2025",
            "NFPA 14": "2024",
            "NFPA 20": "2025",
            "NFPA 24": "2025",
            "NFPA 25": "2013",
            "NFPA 72": "2025",
        }
        for name, edition in expected.items():
            std = CALIFORNIA_2025.standard(name)
            assert std is not None, f"missing {name}"
            assert std.edition == edition, f"{name} expected {edition}, got {std.edition}"

    def test_standards_is_hashable_tuple(self):
        # Frozen dataclass + tuple of frozen StandardEdition stays hashable.
        assert isinstance(CALIFORNIA_2025.standards, tuple)
        assert hash(CALIFORNIA_2025) == hash(CALIFORNIA_2025)
        assert CALIFORNIA_2025 in {CALIFORNIA_2025}

    def test_california_2025_has_nonempty_pinned_editions(self):
        assert CALIFORNIA_2025.standards
        assert CALIFORNIA_2025.standard("NFPA 13") is not None
        assert CALIFORNIA_2025.standard("NFPA 72") is not None

    def test_edition_phrase_lookup(self):
        assert CALIFORNIA_2025.edition_phrase("NFPA 13") == "2025, as amended by California"
        assert CALIFORNIA_2025.edition_phrase("NFPA 14") == "2024"
        assert CALIFORNIA_2025.edition_phrase("NFPA 25") == "2013 California Edition"
        # Missing standard yields an empty string so callers can fall back.
        assert CALIFORNIA_2025.edition_phrase("NFPA 9999") == ""

    def test_default_cycle_has_no_pinned_standards(self):
        # A future cycle that doesn't populate standards still constructs fine.
        c = _bare_cycle("future")
        assert c.standards == ()
        assert c.edition_phrase("NFPA 13") == ""

    def test_unverified_standards_surface(self):
        # The provenance flag is queryable; the ASHRAE energy editions and the
        # UL listings are deliberately marked UNVERIFIED pending code confirmation.
        unverified = {std.name for std in CALIFORNIA_2025.unverified_standards()}
        assert "ASHRAE 90.1" in unverified
        assert "NFPA 13" not in unverified


# ===========================================================================
# 3. Reviewer system prompt + user message reference pinned editions
# ===========================================================================


class TestReviewerPromptPinnedEditions:
    @pytest.mark.parametrize(
        "label",
        ["NFPA 13", "NFPA 72", "ASHRAE 62.1", "ASHRAE 90.1"],
    )
    def test_system_prompt_mentions_nfpa_ashrae_editions(self, label):
        sp = get_system_prompt(CALIFORNIA_2025)
        assert f"{label} {CALIFORNIA_2025.edition_phrase(label)}" in sp

    def test_user_message_mentions_pinned_editions(self):
        um = get_single_spec_user_message(
            "spec content", "23 21 13 - Hydronic.docx", cycle=CALIFORNIA_2025
        )
        assert "NFPA 13" in um
        assert "NFPA 72" in um
        assert "ASHRAE 62.1" in um
        assert "ASHRAE 90.1" in um

    def test_system_prompt_stable_across_calls(self):
        # Cache breakpoint invariant — the prompt must be byte-stable across calls.
        assert get_system_prompt(CALIFORNIA_2025) == get_system_prompt(CALIFORNIA_2025)

    def test_system_prompt_changes_when_pinned_edition_changes(self):
        # Two cycles that differ only in the pinned NFPA 13 edition should
        # produce different system prompts.
        alt = replace(
            CALIFORNIA_2025,
            standards=(StandardEdition("NFPA 13", "2099"),),
        )
        assert get_system_prompt(CALIFORNIA_2025) != get_system_prompt(alt)


# ===========================================================================
# 4. Verifier system prompt block
# ===========================================================================


class TestVerifierPinnedEditionsBlock:
    def test_pinned_lines_render_each_populated_standard(self):
        joined = "\n".join(_pinned_standards_lines(CALIFORNIA_2025))
        assert "Pinned standards editions" in joined
        assert "NFPA 13: 2025, as amended by California" in joined
        assert "NFPA 72: 2025, as amended by California" in joined
        assert "ASHRAE 62.1: 2019" in joined
        assert "ASHRAE 90.1: 2022" in joined
        assert "IAPMO Uniform Plumbing TSC" in joined
        assert "UL 300:" in joined
        assert "UL 555:" in joined

    def test_pinned_lines_include_drift_instruction(self):
        joined = "\n".join(_pinned_standards_lines(CALIFORNIA_2025))
        assert "flag" in joined.lower()
        assert "edition" in joined.lower()

    def test_pinned_lines_empty_when_no_editions(self):
        assert _pinned_standards_lines(_bare_cycle()) == []

    def test_pinned_lines_skip_unpinned_standards(self):
        partial = replace(
            _bare_cycle("partial"),
            standards=(StandardEdition("NFPA 13", "2022"),),
        )
        joined = "\n".join(_pinned_standards_lines(partial))
        assert "NFPA 13: 2022" in joined
        assert "NFPA 72:" not in joined
        assert "ASHRAE" not in joined

    def test_verifier_pinned_block_renders_for_text_fallback_prompt(self):
        prompt = _get_verification_system_prompt(
            CALIFORNIA_2025, include_verdict_tool=False
        )
        assert "Pinned standards editions" in prompt
        assert "NFPA 13" in prompt

    def test_verifier_pinned_block_before_search_budget(self):
        prompt = _get_verification_system_prompt(
            CALIFORNIA_2025, include_verdict_tool=True
        )
        pinned_idx = prompt.find("Pinned standards editions")
        budget_idx = prompt.find("Search budget")
        assert pinned_idx >= 0
        assert budget_idx >= 0
        assert pinned_idx < budget_idx

    def test_verifier_prompt_stable_across_calls(self):
        a = _get_verification_system_prompt(CALIFORNIA_2025, include_verdict_tool=True)
        b = _get_verification_system_prompt(CALIFORNIA_2025, include_verdict_tool=True)
        assert a == b


# ===========================================================================
# 5. Methodology note enumerates pinned editions
# ===========================================================================


class TestMethodologyNotePinnedEditions:
    def test_render_note_includes_all_populated_editions(self):
        note = _render_pinned_editions_note("2025")
        assert "NFPA 13 2025, as amended by California" in note
        assert "ASHRAE 62.1 2019" in note
        assert "UL 300" in note

    def test_render_note_empty_when_cycle_has_no_pinning(self):
        AVAILABLE_CYCLES["__test_bare__"] = _bare_cycle()
        try:
            assert _render_pinned_editions_note("__test_bare__") == ""
        finally:
            AVAILABLE_CYCLES.pop("__test_bare__", None)

    def test_render_note_unknown_label_falls_back_to_default(self):
        note = _render_pinned_editions_note("not-a-real-cycle")
        assert "NFPA 13" in note


# ===========================================================================
# 6. End-to-end report export carries pinned editions
# ===========================================================================


class TestExportedReportPinnedEditions:
    def test_exported_report_contains_pinned_editions_paragraph(self, tmp_path: Path):
        f = _finding()
        result = ReviewResult(findings=[f])
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=result), out)
        text = _all_text_from(Document(str(out)))
        assert "pinned the following standards editions" in text
        assert "NFPA 13" in text


# ===========================================================================
# 7. DEFAULT_CYCLE remains California 2025 (regression guard)
# ===========================================================================


class TestDefaultCycleInvariant:
    def test_default_cycle_is_california_2025(self):
        assert DEFAULT_CYCLE is CALIFORNIA_2025
        assert list(AVAILABLE_CYCLES.keys()) == ["2025"]
