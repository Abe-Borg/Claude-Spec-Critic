"""WS-3 requirements-research fan-out: module contract, schema, runner, splice.

Hermetic throughout — the fake streaming client below scripts every
per-dimension response, and the fan-out is exercised end-to-end (parallel
dimensions, pause_turn continuation, grounding, partial failure, all-fail
abort) without touching the network.
"""
from __future__ import annotations

import dataclasses
import json
import time

import pytest

from src.core.api_config import (
    DEFAULT_VERIFICATION_MAX_USES,
    MODEL_SONNET_46,
    RESEARCH_DEFAULT_MAX_FETCHES,
    RESEARCH_DEFAULT_MAX_SEARCHES,
    build_web_search_tool,
)
from src.core.project_profile import ProjectProfile
from src.modules import DEFAULT_MODULE, ResearchDimension, validate_module_registry
from src.research import (
    CorpusSignals,
    DimensionStatus,
    RequirementsProfile,
    ResearchFanoutError,
    ResearchItem,
    run_requirements_research,
    scrape_corpus_signals,
    splice_profile_into_context,
)
from src.research import requirements_research as rr
from src.review.structured_schemas import (
    REQUIREMENTS_RESEARCH_SCHEMA,
    RESEARCH_TOOL_NAME,
    requirements_research_tool,
)
from tests.fixtures.fake_anthropic import (
    FakeMessage,
    FakeServerToolUsage,
    FakeTextBlock,
    FakeUsage,
    pause_turn_response,
    research_tool_use_response,
    sample_research_profile_payload,
)


# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------


def _complete_profile() -> ProjectProfile:
    return ProjectProfile(
        city="Markham",
        state_or_province="ON",
        country="Canada",
        client_name="ExampleCo",
    )


def _dimension(dimension_id: str = "alpha", marker: str | None = None, **overrides) -> ResearchDimension:
    marker = marker or dimension_id.upper()
    return ResearchDimension(
        dimension_id=dimension_id,
        title=dimension_id.title(),
        prompt_template=f"{marker} research brief for {{city}}, {{state_or_province}}.",
        **overrides,
    )


def _enabled_module(**overrides):
    """A profile-enabled module built off the CA default (not registered).

    Carries every profile-gated slot the D-2 conditional validation demands
    (research persona/dimensions + the WS-4 compliance persona/severities).
    """
    defaults = dict(
        project_profile_enabled=True,
        research_persona="You are a test research assistant.",
        research_dimensions=(_dimension("alpha"),),
        compliance_persona="You are a test compliance reviewer.",
        compliance_severity_definitions="- CRITICAL — permit-blocking omission.",
    )
    defaults.update(overrides)
    return dataclasses.replace(DEFAULT_MODULE, **defaults)


class _FakeStream:
    def __init__(self, message):
        self._message = message

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False

    def get_final_message(self):
        return self._message


class _FakeMessagesAPI:
    def __init__(self, route):
        self._route = route
        self.calls: list[dict] = []

    def stream(self, **kwargs):
        self.calls.append(kwargs)
        result = self._route(kwargs)
        if isinstance(result, Exception):
            raise result
        return _FakeStream(result)


class FakeResearchClient:
    """Scripted ``client.messages.stream`` stand-in.

    ``route`` receives the request kwargs and returns a FakeMessage (or an
    Exception to raise). Dimension fan-out runs on worker threads, so routes
    key off the user-message marker text rather than call order.
    """

    def __init__(self, route):
        self.messages = _FakeMessagesAPI(route)

    @property
    def calls(self) -> list[dict]:
        return self.messages.calls


def _route_by_marker(script: dict[str, list]):
    """Route requests by dimension marker; each marker consumes its own list."""
    remaining = {marker: list(items) for marker, items in script.items()}

    def route(kwargs):
        user_content = kwargs["messages"][0]["content"]
        for marker, items in remaining.items():
            if marker in user_content:
                if not items:
                    raise AssertionError(f"script exhausted for marker {marker!r}")
                return items.pop(0)
        raise AssertionError(f"no route for user message: {user_content[:120]!r}")

    return route


class _LogCollector:
    def __init__(self):
        self.entries: list[tuple[str, str]] = []

    def __call__(self, msg, **kwargs):
        self.entries.append((str(msg), str(kwargs.get("level", "info"))))

    def messages(self, level: str | None = None) -> list[str]:
        return [m for m, lvl in self.entries if level is None or lvl == level]


# ---------------------------------------------------------------------------
# Module contract (D-2 conditional validation)
# ---------------------------------------------------------------------------


class TestResearchSlotValidation:
    def test_enabled_module_with_research_slots_validates(self):
        validate_module_registry([_enabled_module()])

    def test_enabled_requires_persona(self):
        with pytest.raises(ValueError, match="research_persona"):
            validate_module_registry([_enabled_module(research_persona="")])

    def test_enabled_requires_at_least_one_dimension(self):
        with pytest.raises(ValueError, match="research dimension"):
            validate_module_registry([_enabled_module(research_dimensions=())])

    def test_disabled_requires_empty_persona(self):
        with pytest.raises(ValueError, match="research_persona must be empty"):
            validate_module_registry(
                [
                    dataclasses.replace(
                        DEFAULT_MODULE, research_persona="dead content"
                    )
                ]
            )

    def test_disabled_requires_empty_dimensions(self):
        with pytest.raises(ValueError, match="research_dimensions must be empty"):
            validate_module_registry(
                [
                    dataclasses.replace(
                        DEFAULT_MODULE, research_dimensions=(_dimension(),)
                    )
                ]
            )

    def test_disabled_requires_empty_corpus_patterns(self):
        with pytest.raises(ValueError, match="corpus_signal_patterns must be empty"):
            validate_module_registry(
                [
                    dataclasses.replace(
                        DEFAULT_MODULE, corpus_signal_patterns=("Basis of Design",)
                    )
                ]
            )

    def test_duplicate_dimension_ids_rejected(self):
        with pytest.raises(ValueError, match="duplicate research dimension_id"):
            validate_module_registry(
                [
                    _enabled_module(
                        research_dimensions=(_dimension("alpha"), _dimension("alpha"))
                    )
                ]
            )

    def test_unknown_placeholder_rejected(self):
        bad = ResearchDimension(
            dimension_id="alpha",
            title="Alpha",
            prompt_template="Research {nonexistent_placeholder} now.",
        )
        with pytest.raises(ValueError, match="does not format"):
            validate_module_registry([_enabled_module(research_dimensions=(bad,))])

    def test_profile_placeholders_accepted(self):
        dim = ResearchDimension(
            dimension_id="alpha",
            title="Alpha",
            prompt_template=(
                "Codes for {city}, {state_or_province}, {country} for "
                "{client_name} under the {cbc} CBC."
            ),
        )
        validate_module_registry([_enabled_module(research_dimensions=(dim,))])

    def test_negative_budget_rejected(self):
        bad = _dimension("alpha", max_searches=-1)
        with pytest.raises(ValueError, match="non-negative"):
            validate_module_registry([_enabled_module(research_dimensions=(bad,))])

    def test_noncompiling_corpus_pattern_rejected(self):
        with pytest.raises(ValueError, match="does not compile"):
            validate_module_registry(
                [_enabled_module(corpus_signal_patterns=("[unclosed",))]
            )

    def test_default_modules_pass_unchanged(self):
        # Both registered modules have the flag off + empty research slots;
        # the new conditional validation must not disturb them.
        from src.modules import AVAILABLE_MODULES

        validate_module_registry(AVAILABLE_MODULES.values())


# ---------------------------------------------------------------------------
# Schema (strict-mode subset) + tool builder
# ---------------------------------------------------------------------------


def _walk_schema_objects(schema: dict):
    if schema.get("type") == "object" or "properties" in schema:
        yield schema
    for value in schema.get("properties", {}).values():
        yield from _walk_schema_objects(value)
        if value.get("type") == "array" and isinstance(value.get("items"), dict):
            yield from _walk_schema_objects(value["items"])


class TestResearchSchema:
    def test_every_object_level_is_strict_subset(self):
        for obj in _walk_schema_objects(REQUIREMENTS_RESEARCH_SCHEMA):
            assert obj.get("additionalProperties") is False
            assert sorted(obj["required"]) == sorted(obj["properties"].keys())

    def test_no_numeric_or_length_constraints(self):
        text = json.dumps(REQUIREMENTS_RESEARCH_SCHEMA)
        for forbidden in ('"minimum"', '"maximum"', '"minLength"', '"maxLength"'):
            assert forbidden not in text

    def test_enums_only_on_non_nullable_strings(self):
        # Strict-mode rejects an enum on a union type with a null member.
        def _check(schema: dict):
            if "enum" in schema:
                assert schema.get("type") == "string"
            for value in schema.get("properties", {}).values():
                _check(value)
                if isinstance(value.get("items"), dict):
                    _check(value["items"])

        _check(REQUIREMENTS_RESEARCH_SCHEMA)

    def test_tool_builder_strict_for_known_model(self):
        tool = requirements_research_tool(model=MODEL_SONNET_46)
        assert tool["name"] == RESEARCH_TOOL_NAME == "submit_requirements_research"
        assert tool["strict"] is True
        assert tool["input_schema"] is REQUIREMENTS_RESEARCH_SCHEMA

    def test_tool_builder_lenient_for_unknown_model(self):
        assert "strict" not in requirements_research_tool(model="claude-mystery-9")

    def test_no_tool_choice_helper_exists(self):
        # Research must send NO tool_choice: the _20260209 web server tools
        # run programmatic tool calling under the hood, and the API 400s on
        # disable_parallel_tool_use combined with it. The helper was removed
        # so a future call site can't reintroduce the rejected shape.
        import src.review.structured_schemas as schemas

        assert not hasattr(schemas, "research_tool_choice")


class TestWebSearchUserLocation:
    def test_default_is_byte_identical_to_legacy_shape(self):
        # CA-neutrality pin (invariant 2): with no user_location the tool
        # dict must be exactly today's hardcoded-California shape.
        tool = build_web_search_tool(max_uses=DEFAULT_VERIFICATION_MAX_USES)
        assert tool["user_location"] == {
            "type": "approximate",
            "country": "US",
            "region": "California",
        }

    def test_profile_location_threads_through(self):
        loc = _complete_profile().web_search_user_location()
        tool = build_web_search_tool(max_uses=3, user_location=loc)
        assert tool["user_location"] == {
            "type": "approximate",
            "country": "CA",
            "region": "Ontario",
            "city": "Markham",
        }
        # The builder copies — mutating the tool must not reach the profile dict.
        tool["user_location"]["city"] = "Elsewhere"
        assert loc["city"] == "Markham"


# ---------------------------------------------------------------------------
# Corpus-signal scrape
# ---------------------------------------------------------------------------


def _spec(content: str, filename: str = "21 13 13 Wet-Pipe.docx"):
    from src.input.extractor import ExtractedSpec

    return ExtractedSpec(filename=filename, content=content, word_count=len(content.split()))


@pytest.fixture
def stub_corpus_tokens(monkeypatch):
    """Word-count stand-in for ``count_tokens`` in the corpus-signals module.

    Real ``count_tokens`` downloads the cl100k_base encoding on first use;
    stubbing it keeps the suite network-free — the same pattern as
    ``test_context_attachments``.
    """
    from src.research import corpus_signals as cs

    monkeypatch.setattr(cs, "count_tokens", lambda text: len(text.split()))


class TestCorpusSignals:
    def test_empty_corpus_has_no_signals(self):
        signals = scrape_corpus_signals([_spec("Plain text only.")], module=_enabled_module())
        assert not signals.has_signals
        assert signals.render_block() == ""

    def test_module_patterns_match_document_names(self):
        module = _enabled_module(corpus_signal_patterns=(r"Basis of Design",))
        spec = _spec("Refer to the Owner's Basis of Design document, Rev 4.")
        signals = scrape_corpus_signals([spec], module=module)
        assert len(signals.document_names) == 1
        assert "Basis of Design" in signals.document_names[0]

    def test_consultant_and_edition_and_standard_signals(self):
        content = (
            "The Owner's risk consultant shall review all shop drawings. "
            "The NBC-referenced edition governs where editions conflict. "
            "Sprinkler systems shall comply with NFPA 13-2022 throughout."
        )
        signals = scrape_corpus_signals([_spec(content)], module=_enabled_module())
        assert any("risk consultant" in m for m in signals.consultant_insurer_mentions)
        assert any("edition governs" in s for s in signals.edition_governance_sentences)
        assert "NFPA 13 (2022)" in signals.standards_with_editions

    def test_signals_dedupe_across_specs(self):
        module = _enabled_module()
        content = "Comply with NFPA 13-2022."
        signals = scrape_corpus_signals(
            [_spec(content, "a.docx"), _spec(content, "b.docx")], module=module
        )
        assert signals.standards_with_editions == ["NFPA 13 (2022)"]

    def test_render_block_marks_empty_categories(self, stub_corpus_tokens):
        signals = scrape_corpus_signals(
            [_spec("Comply with NFPA 13-2022.")], module=_enabled_module()
        )
        block = signals.render_block()
        assert "Standards cited with edition years:\n- NFPA 13 (2022)" in block
        assert block.count("(none detected)") == 3


# ---------------------------------------------------------------------------
# The fan-out runner
# ---------------------------------------------------------------------------


class TestResearchFanout:
    def test_happy_path_two_dimensions_merge_in_module_order(self):
        module = _enabled_module(
            research_dimensions=(_dimension("alpha"), _dimension("beta"))
        )
        alpha_payload = sample_research_profile_payload()
        beta_payload = {
            "summary": "",
            "items": [
                {
                    "topic": "Seismic",
                    "category": "site_environment",
                    "requirement": "NBC seismic category applies to restraint design.",
                    "actionability": "spec_requirement",
                    "authority": None,
                    "code_reference": None,
                    "source_urls": ["https://nrc.example.ca/seismic"],
                    "confidence": 0.8,
                    "notes": None,
                }
            ],
        }
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [research_tool_use_response(payload=alpha_payload)],
                    "BETA": [
                        research_tool_use_response(
                            payload=beta_payload,
                            searched_urls=["https://nrc.example.ca/seismic"],
                        )
                    ],
                }
            )
        )
        log = _LogCollector()
        profile = run_requirements_research(
            module, _complete_profile(), client=client, log=log
        )
        assert [s.dimension_id for s in profile.dimension_statuses] == ["alpha", "beta"]
        assert all(s.status == "completed" for s in profile.dimension_statuses)
        # Items merged in module dimension order regardless of completion order.
        assert [i.dimension_id for i in profile.items] == ["alpha", "alpha", "beta"]
        assert profile.completed_dimensions == 2 and profile.failed_dimensions == 0
        assert profile.research_date == time.strftime("%Y-%m-%d")
        assert profile.project == _complete_profile().to_dict()
        # Echo-back line fired before any spend (D-1).
        assert any("Researching requirements for Markham, Ontario, Canada" in m
                   for m in log.messages("step"))

    def test_item_ids_are_stable_and_prefixed(self):
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        ids = [i.item_id for i in profile.items]
        assert all(i.startswith("r-") and len(i) == 14 for i in ids)
        assert len(set(ids)) == len(ids)

    def test_partial_failure_continues_flagged(self):
        module = _enabled_module(
            research_dimensions=(_dimension("alpha"), _dimension("beta"))
        )
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [research_tool_use_response()],
                    "BETA": [RuntimeError("beta exploded")],
                }
            )
        )
        log = _LogCollector()
        profile = run_requirements_research(
            module, _complete_profile(), client=client, log=log
        )
        statuses = {s.dimension_id: s for s in profile.dimension_statuses}
        assert statuses["alpha"].status == "completed"
        assert statuses["beta"].status == "failed"
        assert "beta exploded" in statuses["beta"].error
        assert {i.dimension_id for i in profile.items} == {"alpha"}
        assert any("PARTIALLY" in m for m in log.messages("warning"))

    def test_all_dimensions_failed_raises(self):
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [RuntimeError("no network")]})
        )
        with pytest.raises(ResearchFanoutError, match="no network"):
            run_requirements_research(
                _enabled_module(), _complete_profile(), client=client
            )

    def test_grounding_partitions_accepted_and_rejected(self):
        payload = {
            "summary": "",
            "items": [
                {
                    "topic": "Grounded",
                    "category": "governing_code",
                    "requirement": "Grounded requirement.",
                    "actionability": "spec_requirement",
                    "authority": None,
                    "code_reference": None,
                    "source_urls": ["https://codes.example.gov/adoption"],
                    "confidence": 0.9,
                    "notes": None,
                },
                {
                    "topic": "Invented",
                    "category": "governing_code",
                    "requirement": "Ungrounded requirement.",
                    "actionability": "spec_requirement",
                    "authority": None,
                    "code_reference": None,
                    "source_urls": ["https://invented.example.com/nowhere"],
                    "confidence": 0.9,
                    "notes": None,
                },
            ],
        }
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        research_tool_use_response(
                            payload=payload,
                            searched_urls=["https://codes.example.gov/adoption"],
                        )
                    ]
                }
            )
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        grounded, ungrounded = profile.items
        assert grounded.grounded is True
        assert grounded.accepted_sources == ["https://codes.example.gov/adoption"]
        assert ungrounded.grounded is False
        assert ungrounded.accepted_sources == []
        # Kept, never dropped — it renders with the [UNVERIFIED] marker.
        assert "[UNVERIFIED]" in profile.render_text()
        assert profile.dimension_statuses[0].grounded_count == 1

    def test_pause_turn_resumes_and_pools_grounding(self):
        payload = {
            "summary": "",
            "items": [
                {
                    "topic": "From first turn",
                    "category": "ahj_requirement",
                    "requirement": "Requirement grounded by the pre-pause search.",
                    "actionability": "spec_requirement",
                    "authority": None,
                    "code_reference": None,
                    # Cited URL was retrieved BEFORE the pause — pooling
                    # across continuation responses must accept it.
                    "source_urls": ["https://ahj.example.gov/bulletin"],
                    "confidence": 0.8,
                    "notes": None,
                }
            ],
        }
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        pause_turn_response(
                            searched_urls=["https://ahj.example.gov/bulletin"],
                            web_search_requests=2,
                        ),
                        research_tool_use_response(payload=payload, searched_urls=[]),
                    ]
                }
            )
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        assert profile.dimension_statuses[0].status == "completed"
        assert profile.items[0].grounded is True
        assert len(client.calls) == 2
        # The resume re-sends the assistant content with no synthetic user turn.
        resumed = client.calls[1]["messages"]
        assert [m["role"] for m in resumed] == ["user", "assistant"]
        assert profile.dimension_statuses[0].web_search_requests == 2

    def test_pause_turn_resume_elides_oversized_fetched_pdf(self):
        """A fetched >600-page PDF must not ride the continuation resume.

        The API enforces its per-request PDF page limit on re-sent
        ``web_fetch`` documents (observed live: ``messages.1.content.22.
        pdf.source.base64.data: A maximum of 600 PDF pages may be
        provided`` killed a research dimension mid-continuation), so the
        resume path elides the oversized payload instead of 400ing.
        """
        import base64 as _b64
        import io as _io

        from pypdf import PdfWriter

        from src.core.resend_sanitizer import MAX_RESEND_PDF_PAGES

        writer = PdfWriter()
        for _ in range(MAX_RESEND_PDF_PAGES + 1):
            writer.add_blank_page(width=72, height=72)
        buf = _io.BytesIO()
        writer.write(buf)
        oversized_pdf_b64 = _b64.b64encode(buf.getvalue()).decode("ascii")

        pause_with_pdf = FakeMessage(
            content=[
                {
                    "type": "web_fetch_tool_result",
                    "tool_use_id": "srvtoolu_fetch_1",
                    "content": {
                        "type": "web_fetch_result",
                        "url": "https://codes.example.gov/full-code.pdf",
                        "content": {
                            "type": "document",
                            "title": "Full building code",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": oversized_pdf_b64,
                            },
                        },
                    },
                }
            ],
            stop_reason="pause_turn",
            usage=FakeUsage(
                server_tool_use=FakeServerToolUsage(
                    web_search_requests=1, web_fetch_requests=1
                )
            ),
        )
        client = FakeResearchClient(
            _route_by_marker(
                {"ALPHA": [pause_with_pdf, research_tool_use_response()]}
            )
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        assert profile.dimension_statuses[0].status == "completed"
        assert len(client.calls) == 2
        resumed = client.calls[1]["messages"]
        assert [m["role"] for m in resumed] == ["user", "assistant"]
        serialized = json.dumps(resumed[1]["content"])
        assert "application/pdf" not in serialized
        assert "elided" in serialized
        # The result block itself (and its URL) survives — only the PDF
        # payload is swapped for the plain-text note.
        assert "https://codes.example.gov/full-code.pdf" in serialized

    def test_failed_dimension_preserves_telemetry_from_completed_calls(self):
        """A dimension that fails mid-continuation still reports the searches
        (and tokens) its completed calls already billed — a failure must not
        read as zero-cost in diagnostics."""
        module = _enabled_module(
            research_dimensions=(_dimension("alpha"), _dimension("beta"))
        )
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        pause_turn_response(web_search_requests=2),
                        RuntimeError("boom"),
                    ],
                    "BETA": [research_tool_use_response()],
                }
            )
        )
        profile = run_requirements_research(module, _complete_profile(), client=client)
        by_id = {s.dimension_id: s for s in profile.dimension_statuses}
        assert by_id["alpha"].status == "failed"
        assert "RuntimeError: boom" in by_id["alpha"].error
        # The pre-failure pause_turn call's searches are preserved.
        assert by_id["alpha"].web_search_requests == 2
        assert by_id["beta"].status == "completed"

    def test_failed_dimension_telemetry_spans_retried_attempts(self, monkeypatch):
        """Billed usage from a retried (abandoned) attempt still reaches the
        terminal failure status — the aggregate must span attempts, not just
        the last one."""
        monkeypatch.setattr(rr.time, "sleep", lambda _s: None)
        module = _enabled_module(
            research_dimensions=(_dimension("alpha"), _dimension("beta"))
        )
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        # Attempt 1: a pause that billed 2 searches, then a
                        # retryable transport error on the continuation.
                        pause_turn_response(web_search_requests=2),
                        RuntimeError("connection reset by peer"),
                        # Attempt 2: another billed pause, then a
                        # non-retryable terminal error.
                        pause_turn_response(web_search_requests=1),
                        RuntimeError("boom"),
                    ],
                    "BETA": [research_tool_use_response()],
                }
            )
        )
        profile = run_requirements_research(module, _complete_profile(), client=client)
        by_id = {s.dimension_id: s for s in profile.dimension_statuses}
        assert by_id["alpha"].status == "failed"
        assert "RuntimeError: boom" in by_id["alpha"].error
        # 2 searches from the retried attempt + 1 from the terminal attempt.
        assert by_id["alpha"].web_search_requests == 3
        assert by_id["beta"].status == "completed"

    def test_pause_turn_over_budget_ceiling_fails_dimension(self):
        module = _enabled_module(
            research_dimensions=(_dimension("alpha", max_searches=1),)
        )
        client = FakeResearchClient(
            _route_by_marker(
                {"ALPHA": [pause_turn_response(web_search_requests=3)]}
            )
        )
        with pytest.raises(ResearchFanoutError, match="budget ceiling"):
            run_requirements_research(module, _complete_profile(), client=client)

    def test_incomplete_stop_reason_fails_dimension(self):
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        FakeMessage(
                            content=[FakeTextBlock(text="truncat")],
                            stop_reason="max_tokens",
                        )
                    ]
                }
            )
        )
        with pytest.raises(ResearchFanoutError, match="incomplete"):
            run_requirements_research(
                _enabled_module(), _complete_profile(), client=client
            )

    def test_tagged_json_fallback_parses(self):
        payload = sample_research_profile_payload()
        text = f"<research_json>{json.dumps(payload)}</research_json>"
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        FakeMessage(
                            content=[FakeTextBlock(text=text)],
                            stop_reason="end_turn",
                        )
                    ]
                }
            )
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        assert len(profile.items) == 2
        # No searches ran, so nothing can ground.
        assert all(not i.grounded for i in profile.items)

    def test_no_payload_fails_dimension(self):
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [
                        FakeMessage(
                            content=[FakeTextBlock(text="I could not research.")],
                            stop_reason="end_turn",
                        )
                    ]
                }
            )
        )
        with pytest.raises(ResearchFanoutError, match="no parseable payload"):
            run_requirements_research(
                _enabled_module(), _complete_profile(), client=client
            )

    def test_actionability_and_confidence_clamped_at_parse(self):
        payload = {
            "summary": "",
            "items": [
                {
                    "topic": "Weird",
                    "category": "governing_code",
                    "requirement": "Requirement with junk metadata.",
                    "actionability": "banana",
                    "authority": None,
                    "code_reference": None,
                    "source_urls": [],
                    "confidence": 7.5,
                    "notes": None,
                }
            ],
        }
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response(payload=payload)]})
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        item = profile.items[0]
        # Unknown actionability coerces to spec_requirement (over-check, never skip).
        assert item.actionability == "spec_requirement"
        assert item.confidence == 1.0

    def test_corpus_signals_ride_user_message_as_data(self, stub_corpus_tokens):
        signals = CorpusSignals(standards_with_editions=["NFPA 13 (2022)"])
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        run_requirements_research(
            _enabled_module(),
            _complete_profile(),
            corpus_signals=signals,
            client=client,
        )
        user_message = client.calls[0]["messages"][0]["content"]
        assert "<corpus_signals>" in user_message
        assert "NFPA 13 (2022)" in user_message

    def test_request_shape_carries_location_budgets_and_tools(self):
        module = _enabled_module(
            research_dimensions=(_dimension("alpha", max_searches=24, max_fetches=8),)
        )
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        run_requirements_research(module, _complete_profile(), client=client)
        kwargs = client.calls[0]
        tools = kwargs["tools"]
        names = [t.get("name") for t in tools]
        assert names == ["web_search", "web_fetch", "submit_requirements_research"]
        assert tools[0]["max_uses"] == 24
        assert tools[0]["user_location"]["city"] == "Markham"
        assert tools[1]["max_uses"] == 8
        # web_fetch has no location parameter — the key must never appear.
        assert "user_location" not in tools[1]
        # Trailing cache breakpoint lands on the output tool.
        assert tools[-1].get("cache_control") == {"type": "ephemeral", "ttl": "1h"}
        # NO tool_choice key: disable_parallel_tool_use is rejected (400)
        # alongside the _20260209 web tools' programmatic tool calling.
        assert "tool_choice" not in kwargs
        assert kwargs["max_tokens"] == 24_000

    def test_engine_default_budgets_apply_when_dimension_says_zero(self):
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        tools = client.calls[0]["tools"]
        assert tools[0]["max_uses"] == RESEARCH_DEFAULT_MAX_SEARCHES
        assert tools[1]["max_uses"] == RESEARCH_DEFAULT_MAX_FETCHES

    def test_diag_rollup_records_per_dimension(self):
        class _FakeDiag:
            def __init__(self):
                self.calls = []

            def record_api_call(self, **kwargs):
                self.calls.append(kwargs)

        diag = _FakeDiag()
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        run_requirements_research(
            _enabled_module(), _complete_profile(), client=client, diag=diag
        )
        assert len(diag.calls) == 1
        call = diag.calls[0]
        assert call["phase"] == "location_research"
        assert call["extra"]["dimension_id"] == "alpha"
        assert call["extra"]["dimension_status"] == "completed"
        assert call["level"] == "info"
        # B3: the row records the real request cap, not 0.
        from src.core.api_config import research_max_tokens

        assert call["max_output_tokens"] == research_max_tokens(
            model=call["model"]
        )
        assert call["max_output_tokens"] > 0
        # B2: parse-time drops are counted (0 on a clean payload).
        assert call["extra"]["dropped_item_count"] == 0

    def test_zero_item_completion_surfaces_as_warning(self):
        """B1: a completed dimension with 0 items is an unresearched area,
        not an unqualified success — warning log, warning diag row, and an
        honesty line in the rendered profile block."""

        class _FakeDiag:
            def __init__(self):
                self.calls = []

            def record_api_call(self, **kwargs):
                self.calls.append(kwargs)

        diag = _FakeDiag()
        log = _LogCollector()
        client = FakeResearchClient(
            _route_by_marker(
                {"ALPHA": [research_tool_use_response(payload={"items": []})]}
            )
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client, diag=diag, log=log
        )
        by_id = {s.dimension_id: s for s in profile.dimension_statuses}
        assert by_id["alpha"].status == "completed"
        assert by_id["alpha"].item_count == 0
        assert [s.dimension_id for s in profile.empty_completed_dimensions] == ["alpha"]
        # Coordinator log at warning, naming the honesty framing.
        warnings = log.messages("warning")
        assert any(
            "completed with 0 items" in m and "unresearched" in m for m in warnings
        )
        # Diag row at warning level.
        assert diag.calls[0]["level"] == "warning"
        assert diag.calls[0]["extra"]["item_count"] == 0
        # Rendered block tells the review model too.
        rendered = profile.render_text()
        assert "completed without finding any requirements (alpha)" in rendered
        assert "unresearched, not confirmed-clean" in rendered

    def test_parse_time_item_drops_are_counted(self):
        payload = {
            "items": [
                "not-a-dict",
                {"requirement": "   ", "category": "governing_code"},
                {
                    "requirement": "Real requirement.",
                    "category": "governing_code",
                    "topic": "Code",
                },
            ]
        }
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response(payload=payload)]})
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        by_id = {s.dimension_id: s for s in profile.dimension_statuses}
        assert by_id["alpha"].item_count == 1
        assert by_id["alpha"].dropped_item_count == 2

    def test_dropped_item_count_round_trips(self):
        status = DimensionStatus(
            dimension_id="alpha",
            status="completed",
            item_count=1,
            dropped_item_count=2,
        )
        profile = RequirementsProfile(
            items=[], dimension_statuses=[status], research_date="2026-07-21"
        )
        restored = RequirementsProfile.from_dict(profile.to_dict())
        assert restored is not None
        assert restored.dimension_statuses[0].dropped_item_count == 2
        # Legacy rows without the key load with the 0 default.
        legacy = profile.to_dict()
        del legacy["dimension_statuses"][0]["dropped_item_count"]
        restored_legacy = RequirementsProfile.from_dict(legacy)
        assert restored_legacy is not None
        assert restored_legacy.dimension_statuses[0].dropped_item_count == 0

    def test_profile_without_empty_dimensions_renders_no_note(self):
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        profile = run_requirements_research(
            _enabled_module(), _complete_profile(), client=client
        )
        assert profile.empty_completed_dimensions == []
        assert "unresearched, not confirmed-clean" not in profile.render_text()


# ---------------------------------------------------------------------------
# Rendered profile block (deterministic golden) + round-trip
# ---------------------------------------------------------------------------


def _fixed_profile() -> RequirementsProfile:
    return RequirementsProfile(
        items=[
            ResearchItem(
                item_id="r-aaaaaaaaaaaa",
                dimension_id="governing_codes",
                topic="Building code",
                category="governing_code",
                requirement="The 2024 IBC as amended governs.",
                authority="State Agency",
                code_reference="IBC 2024",
                source_urls=["https://codes.example.gov/adoption"],
                accepted_sources=["https://codes.example.gov/adoption"],
                grounded=True,
                confidence=0.9,
            ),
            ResearchItem(
                item_id="r-bbbbbbbbbbbb",
                dimension_id="ahj_requirements",
                topic="Flow test window",
                category="ahj_requirement",
                requirement="Hydrant flow tests witnessed April through October only.",
                authority="City Fire Marshal",
                source_urls=["https://city.example.gov/flow"],
                accepted_sources=["https://city.example.gov/flow"],
                grounded=True,
                confidence=0.7,
                actionability="process_advisory",
            ),
            ResearchItem(
                item_id="r-cccccccccccc",
                dimension_id="governing_codes",
                topic="Local amendment",
                category="local_amendment",
                requirement="Municipal amendment requires supervised valves.",
                confidence=0.5,
            ),
        ],
        dimension_statuses=[
            DimensionStatus(dimension_id="governing_codes", status="completed", item_count=2),
            DimensionStatus(dimension_id="ahj_requirements", status="completed", item_count=1),
            DimensionStatus(dimension_id="client_standards", status="failed", error="boom"),
        ],
        research_date="2026-07-14",
        project=ProjectProfile("Markham", "ON", "CA", "ExampleCo").to_dict(),
    )


_EXPECTED_RENDER = """PROJECT REQUIREMENTS PROFILE
Project: Markham, Ontario, Canada | Client: ExampleCo
Generated by location/client research (2 of 3 dimensions completed), researched 2026-07-14. Edition and process facts are as-of that date.
Items marked [UNVERIFIED] could not be grounded in retrieved sources.
Items marked [PROCESS] are project-team process/schedule advisories, not specification content.

GOVERNING CODES & AMENDMENTS
- [r-aaaaaaaaaaaa] The 2024 IBC as amended governs. (Authority: State Agency; Ref: IBC 2024; Sources: https://codes.example.gov/adoption; confidence 90%)
- [r-cccccccccccc] Municipal amendment requires supervised valves. (Sources: [UNVERIFIED]; confidence 50%)

AHJ REQUIREMENTS
- [r-bbbbbbbbbbbb] [PROCESS] Hydrant flow tests witnessed April through October only. (Authority: City Fire Marshal; Sources: https://city.example.gov/flow; confidence 70%)"""


class TestRenderedProfile:
    def test_render_text_is_byte_pinned(self):
        assert _fixed_profile().render_text() == _EXPECTED_RENDER

    def test_render_is_deterministic(self):
        assert _fixed_profile().render_text() == _fixed_profile().render_text()

    def test_round_trip_preserves_everything(self):
        original = _fixed_profile()
        restored = RequirementsProfile.from_dict(original.to_dict())
        assert restored is not None
        assert restored.to_dict() == original.to_dict()
        assert restored.render_text() == original.render_text()

    def test_from_dict_rejects_garbage(self):
        assert RequirementsProfile.from_dict(None) is None
        assert RequirementsProfile.from_dict("nope") is None
        assert RequirementsProfile.from_dict({}) is None


# ---------------------------------------------------------------------------
# Context splice + lowest-confidence-first trim
# ---------------------------------------------------------------------------


class TestSpliceProfileIntoContext:
    def test_within_cap_appends_wrapped_attachment(self, monkeypatch):
        monkeypatch.setattr(
            rr, "context_within_token_cap", lambda text: (len(text.split()), True)
        )
        profile = _fixed_profile()
        effective, dropped = splice_profile_into_context("User context.", profile)
        assert dropped == 0
        assert effective.startswith("User context.\n\n")
        assert "--- BEGIN ATTACHMENT: Project Requirements Profile ---" in effective
        assert "--- END ATTACHMENT: Project Requirements Profile ---" in effective
        assert profile.render_text() in effective

    def test_over_cap_drops_lowest_confidence_first(self, monkeypatch):
        profile = _fixed_profile()
        # Fake cap: only a render without the 0.5-confidence item fits.
        full_marker = "r-cccccccccccc"

        def fake_cap(text):
            fits = full_marker not in text
            return len(text), fits

        monkeypatch.setattr(rr, "context_within_token_cap", fake_cap)
        log = _LogCollector()
        effective, dropped = splice_profile_into_context(
            "User context.", profile, log=log
        )
        assert dropped == 1
        assert full_marker not in effective
        # Higher-confidence items survive.
        assert "r-aaaaaaaaaaaa" in effective and "r-bbbbbbbbbbbb" in effective
        # The STRUCTURED profile keeps every item — only the render trims.
        assert len(profile.items) == 3
        assert any("dropped 1 lowest-confidence" in m for m in log.messages("warning"))

    def test_degenerate_cap_drops_block_entirely(self, monkeypatch):
        profile = _fixed_profile()
        monkeypatch.setattr(rr, "context_within_token_cap", lambda text: (len(text), False))
        log = _LogCollector()
        effective, dropped = splice_profile_into_context(
            "User context.", profile, log=log
        )
        assert effective == "User context."
        assert dropped == len(profile.items)
        assert any("dropped from review context entirely" in m for m in log.messages("warning"))


# ---------------------------------------------------------------------------
# Pipeline gate + persistence (research never re-runs on resume)
# ---------------------------------------------------------------------------


class TestResearchPhaseGate:
    def test_gate_conditions(self):
        from src.orchestration.pipeline import _research_phase_applies

        complete = _complete_profile()
        incomplete = ProjectProfile("", "ON", "CA", "ExampleCo")
        assert not _research_phase_applies(DEFAULT_MODULE, complete)  # flag off
        assert not _research_phase_applies(_enabled_module(), None)  # no profile
        assert not _research_phase_applies(_enabled_module(), incomplete)
        assert _research_phase_applies(_enabled_module(), complete)
        assert not _research_phase_applies(
            _enabled_module(
                project_profile_enabled=True,
                research_persona="p",
                research_dimensions=(),
            ),
            complete,
        )

    def _submit(self, monkeypatch, tmp_path, *, module, profile):
        """Drive start_batch_review hermetically; capture what it submits."""
        from src.input.extractor import ExtractedSpec
        from src.batch.batch import BatchJob
        from src.orchestration import pipeline

        captured: dict = {}

        def fake_prepare(**kwargs):
            captured["prepare_context"] = kwargs["project_context"]
            spec = ExtractedSpec(filename="a.docx", content="body", word_count=1)
            return pipeline._PreparedSpecs(
                specs=[spec], leed_alerts=[], placeholder_alerts=[]
            )

        def fake_submit(specs, *, project_context, model, cycle, pre_detected_alerts):
            captured["submit_context"] = project_context
            return BatchJob(
                batch_id="batch_1",
                job_type="review",
                request_map={
                    "review__a__0": {"filename": "a.docx", "index": 0, "type": "review"}
                },
                created_at=time.time(),
            )

        monkeypatch.setattr(pipeline, "_prepare_specs", fake_prepare)
        monkeypatch.setattr(pipeline, "submit_review_batch", fake_submit)
        submission = pipeline.start_batch_review(
            input_dir=tmp_path,
            files=None,
            project_context="User context.",
            module=module,
            project_profile=profile,
        )
        return submission, captured

    def test_profile_less_run_never_calls_runner(self, monkeypatch, tmp_path):
        from src.orchestration import pipeline

        def _explode(**kwargs):
            raise AssertionError("research phase must not run for a profile-less run")

        monkeypatch.setattr(pipeline, "_run_research_phase", _explode)
        submission, captured = self._submit(
            monkeypatch, tmp_path, module=DEFAULT_MODULE, profile=None
        )
        assert submission.requirements_profile is None
        assert captured["prepare_context"] == "User context."
        assert captured["submit_context"] == "User context."

    def test_enabled_run_splices_effective_context(self, monkeypatch, tmp_path):
        from src.orchestration import pipeline

        def fake_research(**kwargs):
            assert kwargs["user_context"] == "User context."
            return "EFFECTIVE CONTEXT", {"items": [], "research_date": "2026-07-14"}

        monkeypatch.setattr(pipeline, "_run_research_phase", fake_research)
        submission, captured = self._submit(
            monkeypatch, tmp_path, module=_enabled_module(), profile=_complete_profile()
        )
        # Preflight counts and the batch submit both see the spliced context.
        assert captured["prepare_context"] == "EFFECTIVE CONTEXT"
        assert captured["submit_context"] == "EFFECTIVE CONTEXT"
        assert submission.project_context == "EFFECTIVE CONTEXT"
        assert submission.requirements_profile == {
            "items": [],
            "research_date": "2026-07-14",
        }


class TestRunResearchPhaseIntegration:
    def test_end_to_end_scrape_research_splice(
        self, tmp_path, monkeypatch, stub_corpus_tokens
    ):
        """The full pre-submit glue: real DOCX → cached extraction → corpus
        scrape → fan-out (fake client) → splice into the effective context."""
        from docx import Document

        from src.orchestration import pipeline

        doc = Document()
        doc.add_paragraph("Sprinkler systems shall comply with NFPA 13-2022.")
        doc.add_paragraph("The Owner's risk consultant shall review submittals.")
        spec_path = tmp_path / "21 13 13 Wet-Pipe Sprinkler Systems.docx"
        doc.save(str(spec_path))

        monkeypatch.setattr(rr, "context_within_token_cap", lambda text: (0, True))
        client = FakeResearchClient(
            _route_by_marker({"ALPHA": [research_tool_use_response()]})
        )
        monkeypatch.setattr(rr, "_get_client", lambda: client)

        log = _LogCollector()
        effective, profile_dict = pipeline._run_research_phase(
            module=_enabled_module(),
            profile=_complete_profile(),
            input_dir=tmp_path,
            files=[spec_path],
            user_context="Operator context.",
            log=log,
            progress=lambda *a, **k: None,
        )
        assert effective.startswith("Operator context.")
        assert "PROJECT REQUIREMENTS PROFILE" in effective
        restored = RequirementsProfile.from_dict(profile_dict)
        assert restored is not None and len(restored.items) == 2
        # The corpus scrape fed the spec's own vocabulary to the researcher.
        user_message = client.calls[0]["messages"][0]["content"]
        assert "<corpus_signals>" in user_message
        assert "NFPA 13 (2022)" in user_message
        assert any("risk consultant" in m for m in (user_message,))

    def _forbid_runner(self, monkeypatch):
        """Fail the test if the API-backed fan-out is ever reached."""

        def _boom(*args, **kwargs):
            raise AssertionError(
                "research fan-out must not run when the spec set is unusable"
            )

        # ``_run_research_phase`` binds the runner via a deferred
        # ``from ..research import ...`` at call time, so patching the
        # package attribute intercepts it.
        monkeypatch.setattr("src.research.run_requirements_research", _boom)

    def test_empty_input_dir_aborts_before_research_spend(
        self, tmp_path, monkeypatch
    ):
        # PR #299 review (Codex P2): a run that cannot submit must fail
        # BEFORE the research budget is spent, not after.
        from src.orchestration import pipeline

        self._forbid_runner(monkeypatch)
        with pytest.raises(FileNotFoundError, match="No specification files found"):
            pipeline._run_research_phase(
                module=_enabled_module(),
                profile=_complete_profile(),
                input_dir=tmp_path,
                files=None,
                user_context="",
                log=_LogCollector(),
                progress=lambda *a, **k: None,
            )

    def test_extraction_failure_aborts_before_research_spend(
        self, tmp_path, monkeypatch
    ):
        from src.orchestration import pipeline

        self._forbid_runner(monkeypatch)

        def _corrupt(_files):
            raise ValueError("corrupt DOCX: not a zip archive")

        monkeypatch.setattr(pipeline, "extract_multiple_specs_cached", _corrupt)
        spec_path = tmp_path / "21 13 13 Broken.docx"
        spec_path.write_bytes(b"not a real docx")
        with pytest.raises(ValueError, match="corrupt DOCX"):
            pipeline._run_research_phase(
                module=_enabled_module(),
                profile=_complete_profile(),
                input_dir=tmp_path,
                files=[spec_path],
                user_context="",
                log=_LogCollector(),
                progress=lambda *a, **k: None,
            )

    def test_all_empty_specs_abort_before_research_spend(
        self, tmp_path, monkeypatch
    ):
        from src.input.extractor import ExtractedSpec
        from src.orchestration import pipeline

        self._forbid_runner(monkeypatch)
        empty = ExtractedSpec(filename="a.docx", content="   ", word_count=0)
        monkeypatch.setattr(
            pipeline, "extract_multiple_specs_cached", lambda _files: [empty]
        )
        spec_path = tmp_path / "a.docx"
        spec_path.write_bytes(b"placeholder")
        with pytest.raises(FileNotFoundError, match="All files failed extraction"):
            pipeline._run_research_phase(
                module=_enabled_module(),
                profile=_complete_profile(),
                input_dir=tmp_path,
                files=[spec_path],
                user_context="",
                log=_LogCollector(),
                progress=lambda *a, **k: None,
            )


class TestRequirementsProfilePersistence:
    def _submission(self, requirements_profile):
        from src.batch.batch import BatchJob
        from src.orchestration.pipeline import BatchSubmission

        return BatchSubmission(
            job=BatchJob(
                batch_id="batch_1",
                job_type="review",
                request_map={},
                created_at=1000.0,
            ),
            model="claude-opus-4-8",
            project_context="ctx",
            requirements_profile=requirements_profile,
        )

    def test_pending_batch_round_trips_requirements_profile(self, tmp_path):
        from src.orchestration.batch_resume import (
            PendingBatch,
            load_pending_batch,
            save_pending_batch,
        )

        profile_dict = _fixed_profile().to_dict()
        pending = PendingBatch.from_submission(
            self._submission(profile_dict), run_id="r1", app_version="3.0.0"
        )
        path = tmp_path / "pending.json"
        save_pending_batch(pending, path=path)
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.requirements_profile == profile_dict
        restored = RequirementsProfile.from_dict(loaded.requirements_profile)
        assert restored is not None and len(restored.items) == 3

    def test_legacy_state_without_key_loads_none(self, tmp_path):
        import json as _json

        from src.orchestration.batch_resume import load_pending_batch

        legacy = {
            "schema_version": 1,
            "batch_id": "batch_legacy",
            "model": "claude-opus-4-8",
        }
        path = tmp_path / "pending.json"
        path.write_text(_json.dumps(legacy), encoding="utf-8")
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.requirements_profile is None

    def test_finalize_carries_requirements_profile(self):
        from src.orchestration.pipeline import CollectedBatchState, finalize_batch_result
        from src.review.reviewer import ReviewResult

        profile_dict = {"items": [], "research_date": "2026-07-14"}
        state = CollectedBatchState(
            submission=self._submission(profile_dict),
            review_result=ReviewResult(findings=[]),
        )
        result = finalize_batch_result(state)
        assert result.requirements_profile == profile_dict


# ---------------------------------------------------------------------------
# Fan-out progress emissions (WS2 / B5)
# ---------------------------------------------------------------------------


class TestResearchProgress:
    def test_fanout_emits_real_completion_fractions(self):
        """Each completed dimension advances progress by its real fraction —
        the legacy behavior emitted a flat 0.0 for the whole fan-out, which
        froze the run bar for the entire multi-minute research phase."""
        module = _enabled_module(
            research_dimensions=(_dimension("alpha"), _dimension("beta"))
        )
        client = FakeResearchClient(
            _route_by_marker(
                {
                    "ALPHA": [research_tool_use_response()],
                    "BETA": [research_tool_use_response()],
                }
            )
        )
        emissions: list[float] = []

        def progress(pct, _msg, **_kwargs):
            emissions.append(round(float(pct), 1))

        run_requirements_research(
            module, _complete_profile(), client=client, progress=progress
        )
        # One 0.0 start emission, then one real fraction per completed
        # dimension (completion order is thread-dependent, values aren't).
        assert emissions[0] == 0.0
        assert sorted(emissions[1:]) == [50.0, 100.0]
