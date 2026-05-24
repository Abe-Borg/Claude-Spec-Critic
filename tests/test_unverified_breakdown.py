"""Report 'Verification:' line reconciles with the trust-status breakdown.

The raw UNVERIFIED *verdict* count lumps locally-classified findings
(resolved without a web search) together with genuinely-unverifiable ones,
so a reader saw a large "N unverified" in the summary line that disagreed
with the much smaller "Insufficient evidence" count in the Trust Model
Summary. The summary line now appends a status breakdown of the
unverified-verdict findings using the same labels as the trust model.
"""
from __future__ import annotations

from types import SimpleNamespace

from docx import Document

from src.output.report_exporter import _write_summary_table
from src.review.reviewer import Finding
from src.verification.verifier import VerificationResult


def _finding(verdict, *, cache_status="miss", grounded=False, sources=None):
    f = Finding(
        severity="HIGH",
        fileName="22 11 00.docx",
        section="2.1",
        issue="i",
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference="",
        confidence=0.5,
    )
    f.verification = VerificationResult(
        verdict=verdict,
        grounded=grounded,
        cache_status=cache_status,
        sources=list(sources or []),
        accepted_sources=list(sources or []),
    )
    return f


def _review(findings):
    return SimpleNamespace(
        critical_count=0,
        high_count=len(findings),
        medium_count=0,
        gripe_count=0,
        total_count=len(findings),
        input_tokens=0,
        output_tokens=0,
        elapsed_seconds=1.0,
        findings=findings,
    )


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_unverified_breakdown_reconciles_with_trust_status():
    findings = [
        _finding("UNVERIFIED", cache_status="local_skip"),  # locally classified
        _finding("UNVERIFIED", cache_status="local_skip"),  # locally classified
        _finding("UNVERIFIED", cache_status="local_skip"),  # locally classified
        _finding("UNVERIFIED", cache_status="miss"),        # insufficient evidence
        _finding("UNVERIFIED", cache_status="miss"),        # insufficient evidence
        _finding("CONFIRMED", cache_status="miss", grounded=True, sources=["https://x"]),
    ]
    doc = Document()
    _write_summary_table(doc, _review(findings), None)
    text = _all_text(doc)

    # Raw verdict line still present (5 of 6 are UNVERIFIED verdict).
    assert "5 unverified" in text
    assert "1 confirmed" in text
    # The breakdown note explains the 5 — 3 are merely locally classified,
    # only 2 are genuinely insufficient — so the line reconciles with the
    # Trust Model Summary instead of overstating "unverified".
    assert "Of the unverified-verdict findings:" in text
    assert "3 locally classified (deterministic)" in text
    assert "2 insufficient evidence" in text


def test_no_breakdown_when_no_unverified():
    findings = [
        _finding("CONFIRMED", grounded=True, sources=["https://x"]),
        _finding("CORRECTED", grounded=True, sources=["https://y"]),
    ]
    doc = Document()
    _write_summary_table(doc, _review(findings), None)
    text = _all_text(doc)
    assert "unverified" not in text.lower()
    assert "Of the unverified-verdict findings:" not in text


def test_breakdown_counts_sum_to_unverified_verdict_count():
    # Property: the breakdown partitions exactly the UNVERIFIED-verdict
    # findings, so it can't drift from the raw count.
    findings = [
        _finding("UNVERIFIED", cache_status="local_skip"),
        _finding("UNVERIFIED", cache_status="miss"),
        _finding("UNVERIFIED", cache_status="miss"),
        _finding("CONFIRMED", grounded=True, sources=["https://x"]),
    ]
    from src.output.report_status import (
        STATUS_DISPLAY_ORDER,
        classify_status,
        status_label,
    )

    unv = [f for f in findings if f.verification.verdict == "UNVERIFIED"]
    counts: dict = {}
    for f in unv:
        st = classify_status(f)
        counts[st] = counts.get(st, 0) + 1
    assert sum(counts.values()) == len(unv) == 3

    doc = Document()
    _write_summary_table(doc, _review(findings), None)
    text = _all_text(doc)
    for s in STATUS_DISPLAY_ORDER:
        if counts.get(s, 0) > 0:
            assert f"{counts[s]} {status_label(s).lower()}" in text
