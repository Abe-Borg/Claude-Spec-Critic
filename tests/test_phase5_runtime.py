"""Phase 5 regression tests: batch runtime, extraction, resume validation.

Covers audit Sections 8.5 (rich-formatting downgrade), 9.1 (polling
backoff), 9.2 (parallel extraction), 9.3 (cross-check overlap dependency
filter), and 9.5 (resume-state validation).
"""

from __future__ import annotations

import time
from pathlib import Path

import pytest
from docx import Document

from src.batch import BatchJob
from src.batch_runtime import (
    PollPolicy,
    _progressive_poll_interval,
    poll_batch_bounded,
)
from src.edit_candidates import (
    SAFETY_AUTO_SAFE,
    SAFETY_AUTO_WITH_CAUTION,
    SAFETY_MANUAL_REVIEW,
)
from src.edit_locator import locate_edit
from src.extractor import (
    ParagraphMapping,
    extract_multiple_specs,
    extract_text_from_docx,
)
from src.pipeline import (
    BatchSubmission,
    CollectedBatchState,
    classify_cross_check_dependencies,
)
from src.resume_state import (
    PHASE_REVIEW_POLL,
    build_resume_state,
    deserialize_resume_state,
)
from src.reviewer import Finding, ReviewResult
from src.verifier import VerificationResult


# --- Section 8.5: rich-formatting downgrade ---------------------------------


def _finding(
    *,
    action: str = "EDIT",
    existing: str = "Provide ASCE 7-16 bracing.",
    replacement: str | None = "Provide ASCE 7-22 bracing.",
    verdict: str | None = "CONFIRMED",
) -> Finding:
    f = Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1.0",
        issue="Issue",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC",
        confidence=0.9,
    )
    if verdict is not None:
        f.verification = VerificationResult(verdict=verdict)
    return f


def test_extractor_records_run_formatting(tmp_path: Path):
    source = tmp_path / "fmt.docx"
    doc = Document()
    plain = doc.add_paragraph()
    plain.add_run("All plain text here.")
    rich = doc.add_paragraph()
    rich.add_run("Plain prefix ")
    bold_run = rich.add_run("BOLD")
    bold_run.bold = True
    rich.add_run(" suffix.")
    doc.save(source)

    spec = extract_text_from_docx(source)
    plain_map = spec.paragraph_map[0]
    rich_map = spec.paragraph_map[1]

    assert plain_map.run_count == 1
    assert plain_map.distinct_formatting_runs == 1
    assert rich_map.run_count == 3
    assert rich_map.distinct_formatting_runs >= 2


def test_locator_downgrades_whole_paragraph_replace_on_rich_paragraph():
    rich_mapping = ParagraphMapping(
        body_index=0,
        element_type="paragraph",
        text="Provide ASCE 7-16 bracing.",
        table_index=None,
        row_index=None,
        cell_index=None,
        run_count=2,
        distinct_formatting_runs=2,
    )
    result = locate_edit(_finding(existing=rich_mapping.text), [rich_mapping])
    assert result.status == "matched"
    # Whole-paragraph replace on rich-formatted paragraph -> manual review.
    assert result.safety_category == SAFETY_MANUAL_REVIEW


def test_locator_downgrades_partial_replace_on_rich_paragraph():
    rich_mapping = ParagraphMapping(
        body_index=0,
        element_type="paragraph",
        text="Provide ASCE 7-16 bracing detail per drawings.",
        table_index=None,
        row_index=None,
        cell_index=None,
        run_count=3,
        distinct_formatting_runs=2,
    )
    result = locate_edit(_finding(existing="ASCE 7-16"), [rich_mapping])
    assert result.status == "matched"
    # Partial replace on rich-formatted paragraph -> caution, not auto-safe.
    assert result.safety_category == SAFETY_AUTO_WITH_CAUTION


def test_locator_keeps_auto_safe_for_plain_paragraph():
    plain_mapping = ParagraphMapping(
        body_index=0,
        element_type="paragraph",
        text="Provide ASCE 7-16 bracing.",
        table_index=None,
        row_index=None,
        cell_index=None,
        run_count=1,
        distinct_formatting_runs=1,
    )
    result = locate_edit(_finding(existing=plain_mapping.text), [plain_mapping])
    assert result.status == "matched"
    assert result.safety_category == SAFETY_AUTO_SAFE


# --- Section 9.1: progressive polling backoff -------------------------------


def test_progressive_poll_interval_starts_at_base():
    policy = PollPolicy(
        poll_interval_seconds=10,
        backoff_after_seconds=300,
        max_poll_interval_seconds=120,
    )
    assert _progressive_poll_interval(elapsed_seconds=0, policy=policy) == 10
    assert _progressive_poll_interval(elapsed_seconds=300, policy=policy) == 10


def test_progressive_poll_interval_ramps_up():
    policy = PollPolicy(
        poll_interval_seconds=10,
        backoff_after_seconds=300,
        max_poll_interval_seconds=120,
    )
    midway = _progressive_poll_interval(elapsed_seconds=450, policy=policy)
    assert 10 < midway < 120


def test_progressive_poll_interval_caps_at_max():
    policy = PollPolicy(
        poll_interval_seconds=10,
        backoff_after_seconds=300,
        max_poll_interval_seconds=120,
    )
    assert _progressive_poll_interval(elapsed_seconds=10_000, policy=policy) == 120


def test_poll_batch_bounded_uses_progressive_interval(monkeypatch):
    """Long-elapsed polling should sleep longer than short-elapsed polling."""
    sleep_calls: list[float] = []

    def fake_sleep(seconds: float) -> None:
        sleep_calls.append(seconds)

    poll_count = {"n": 0}

    from src.batch import BatchStatus

    def fake_poll(_batch_id: str):
        poll_count["n"] += 1
        # First call: in_progress; second call: terminal.
        if poll_count["n"] >= 2:
            return BatchStatus(
                status="ended", processing=0,
                succeeded=1, errored=0, canceled=0, expired=0, total=1,
            )
        return BatchStatus(
            status="in_progress", processing=1,
            succeeded=0, errored=0, canceled=0, expired=0, total=1,
        )

    monkeypatch.setattr("src.batch_runtime.time.sleep", fake_sleep)
    monkeypatch.setattr("src.batch_runtime.poll_batch", fake_poll)

    policy = PollPolicy(
        poll_interval_seconds=5,
        backoff_after_seconds=60,
        max_poll_interval_seconds=60,
    )
    outcome = poll_batch_bounded(
        "msgbatch_x", policy=policy, log=lambda *_: None,
        progress_cb=lambda *_: None,
    )
    assert outcome.terminal is True
    assert sleep_calls and sleep_calls[0] == 5  # base interval immediately


# --- Section 9.2: parallel extraction --------------------------------------


def test_extract_multiple_specs_preserves_order(tmp_path: Path):
    paths = []
    for letter in ["a", "b", "c", "d"]:
        p = tmp_path / f"{letter}.docx"
        doc = Document()
        doc.add_paragraph(f"content {letter}")
        doc.save(p)
        paths.append(p)
    extracted = extract_multiple_specs(paths)
    assert [s.filename for s in extracted] == ["a.docx", "b.docx", "c.docx", "d.docx"]


def test_extract_multiple_specs_single_file_runs_inline(tmp_path: Path):
    p = tmp_path / "only.docx"
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(p)
    extracted = extract_multiple_specs([p])
    assert len(extracted) == 1 and extracted[0].filename == "only.docx"


def test_extract_multiple_specs_propagates_errors(tmp_path: Path):
    good = tmp_path / "good.docx"
    Document().add_paragraph("hi")
    d = Document()
    d.add_paragraph("hi")
    d.save(good)
    bad = tmp_path / "missing.docx"  # does not exist
    with pytest.raises((FileNotFoundError, ValueError)):
        extract_multiple_specs([good, bad])


# --- Section 9.3: cross-check dependency filter -----------------------------


def _cross_finding(
    *, file: str, section: str, issue: str = "coordination issue"
) -> Finding:
    return Finding(
        severity="HIGH",
        fileName=file,
        section=section,
        issue=issue,
        actionType="EDIT",
        existingText="x",
        replacementText="y",
        codeReference=None,
        confidence=0.5,
    )


def _review_finding_disputed(*, file: str, section: str) -> Finding:
    f = Finding(
        severity="HIGH",
        fileName=file,
        section=section,
        issue="upstream",
        actionType="EDIT",
        existingText="a",
        replacementText="b",
        codeReference=None,
        confidence=0.6,
    )
    f.verification = VerificationResult(verdict="DISPUTED")
    return f


def test_drop_cross_check_findings_with_disputed_upstream_drops_match():
    cross = [
        _cross_finding(file="A.docx", section="2.1"),
        _cross_finding(file="B.docx", section="3.0"),
    ]
    review = [_review_finding_disputed(file="A.docx", section="2.1")]
    log_messages: list[str] = []
    kept, _suppressed = classify_cross_check_dependencies(
        cross, review, log=lambda m, **_k: log_messages.append(m)
    )
    assert [f.fileName for f in kept] == ["B.docx"]
    assert any("DISPUTED" in m for m in log_messages)


def test_drop_cross_check_findings_passthrough_without_disputed():
    cross = [_cross_finding(file="A.docx", section="2.1")]
    review = [
        Finding(
            severity="HIGH", fileName="A.docx", section="2.1", issue="ok",
            actionType="EDIT", existingText="a", replacementText="b",
            codeReference=None, confidence=0.6,
            verification=VerificationResult(verdict="CONFIRMED"),
        )
    ]
    kept, _suppressed = classify_cross_check_dependencies(cross, review)
    assert kept == cross


def test_drop_cross_check_findings_uses_affected_files():
    cross = [_cross_finding(file="A.docx", section="2.1")]
    cross[0].affected_files = ["A.docx", "C.docx"]
    review = [_review_finding_disputed(file="C.docx", section="2.1")]
    kept, _suppressed = classify_cross_check_dependencies(cross, review)
    # Cross finding lists C.docx in affected_files, which matches a DISPUTED
    # review finding -> drop.
    assert kept == []


# --- Section 9.5: resume-state validation -----------------------------------


def _build_minimal_submission(batch_id: str = "msgbatch_test123") -> BatchSubmission:
    job = BatchJob(
        batch_id=batch_id,
        job_type="review",
        request_map={"req1": {"index": 0, "filename": "x.docx"}},
        created_at=time.time(),
        status="submitted",
    )
    return BatchSubmission(job=job, files_reviewed=["x.docx"], review_request_ids=["req1"])


def test_resume_state_rejects_invalid_batch_id():
    submission = _build_minimal_submission(batch_id="not_a_batch_id")
    payload = build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission)
    with pytest.raises(ValueError, match="Invalid batch_id"):
        deserialize_resume_state(payload)


def test_resume_state_rejects_unknown_phase():
    submission = _build_minimal_submission()
    payload = build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission)
    payload["phase"] = "what_is_this_phase"
    with pytest.raises(ValueError, match="Unsupported resume phase"):
        deserialize_resume_state(payload)


def test_resume_state_rejects_empty_request_map():
    submission = _build_minimal_submission()
    payload = build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission)
    payload["submission"]["job"]["request_map"] = {}
    with pytest.raises(ValueError, match="non-empty"):
        deserialize_resume_state(payload)


def test_resume_state_rejects_malformed_request_map_entry():
    submission = _build_minimal_submission()
    payload = build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission)
    payload["submission"]["job"]["request_map"] = {"req1": "not-a-dict"}
    with pytest.raises(ValueError, match="must be a dict"):
        deserialize_resume_state(payload)


def test_resume_state_validates_verification_batch_id():
    submission = _build_minimal_submission()
    payload = build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission)
    payload["verification_batch"] = {
        "batch_id": "wrong-prefix",
        "request_map": {"v1": {"finding_idx": 0}},
        "created_at": time.time(),
    }
    with pytest.raises(ValueError, match="Invalid batch_id"):
        deserialize_resume_state(payload)


def test_resume_state_round_trips_with_schema_stamp():
    submission = _build_minimal_submission()
    payload = build_resume_state(phase=PHASE_REVIEW_POLL, submission=submission)
    assert payload["schema"] == "v2"
    restored = deserialize_resume_state(payload)
    assert restored["phase"] == PHASE_REVIEW_POLL
    assert restored["submission"].job.batch_id == submission.job.batch_id
    assert restored["schema"] == "v2"
