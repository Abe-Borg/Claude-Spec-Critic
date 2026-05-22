"""Phase 1 / Step 1.1 — Replacement-text typographic normalization.

Unit tests for the new ``src/editing/replacement_style.py`` module
plus a small integration test that pins the rollup wiring from
``apply_edits.execute_edit_plan`` into ``DiagnosticsReport``.
The end-to-end "applied edit keeps the doc's style" tests live in
``tests/test_spec_editor.py`` next to the other apply-edits tests.


The auto-apply pipeline used to land the model's replacement text in the
source document verbatim. Claude routinely emits curly quotes, em-dashes,
and Unicode apostrophes — typography that does not match most CSI
templates, which use straight quotes and ASCII hyphens consistently.
After the edit landed, the new sentence looked visibly different from
its neighbors.

These tests pin the new ``replacement_style`` module's contract:

1. ``profile_document_style`` reads a sample of the source document and
   majority-votes per dimension (quotes, dashes, apostrophes, NBSP in
   measurements). Empty samples default to ASCII/straight to preserve
   the legacy passthrough behavior for documents the profiler cannot
   classify.
2. ``normalize_replacement_text`` rewrites a single replacement string
   to match the profile. The function is idempotent and a no-op when
   the profile is ``None`` (the kill-switch path) or already matches.
3. ``normalize_replacement_style_enabled`` honors the
   ``SPEC_CRITIC_NORMALIZE_REPLACEMENT_STYLE`` env-var kill switch and
   defaults enabled.

The integration tests in ``tests/test_spec_editor.py`` cover end-to-end
behavior — an applied edit on a curly-quote document keeps curly, and
the inverse holds.
"""
from __future__ import annotations

import pytest

from src.editing.replacement_style import (
    DocumentStyleProfile,
    normalize_replacement_style_enabled,
    normalize_replacement_text,
    profile_document_style,
)


# ---------------------------------------------------------------------------
# Profile
# ---------------------------------------------------------------------------


class TestProfileDocumentStyle:
    def test_pure_straight_quote_doc(self):
        texts = [
            'Provide "schedule 40" piping per ASCE 7-22.',
            "Don't substitute without engineer's approval.",
        ]
        profile = profile_document_style(texts)
        assert profile.prefers_straight_quotes is True
        assert profile.prefers_ascii_apostrophe is True

    def test_pure_curly_quote_doc(self):
        texts = [
            "Provide “schedule 40” piping per ASCE 7-22.",
            "Don’t substitute without engineer’s approval.",
        ]
        profile = profile_document_style(texts)
        assert profile.prefers_straight_quotes is False
        assert profile.prefers_ascii_apostrophe is False

    def test_mixed_doc_straight_majority_wins(self):
        texts = [
            'Provide "schedule 40" piping.',
            'Provide "schedule 80" fittings.',
            "Use “red” color coding.",
        ]
        profile = profile_document_style(texts)
        # 4 straight " vs 2 curly. Majority straight.
        assert profile.prefers_straight_quotes is True

    def test_tie_defaults_to_straight(self):
        texts = ['Provide "x" and “y”.']
        # 2 straight " vs 2 curly. Tie -> straight per spec.
        profile = profile_document_style(texts)
        assert profile.prefers_straight_quotes is True

    def test_empty_doc_defaults_to_straight(self):
        profile = profile_document_style([])
        assert profile.prefers_straight_quotes is True
        assert profile.prefers_ascii_apostrophe is True
        assert profile.prefers_hyphen_dash is True
        assert profile.uses_nbsp_in_measurements is False

    def test_empty_strings_dont_skew_profile(self):
        texts = ["", "", "  ", ""]
        profile = profile_document_style(texts)
        assert profile.prefers_straight_quotes is True

    def test_hyphen_dash_preference(self):
        texts = [
            "Provide R-454B refrigerant per ASHRAE 34.",
            "Install at 60-80 psi operating pressure.",
            "Use Class B fire-rated assemblies.",
        ]
        # Hyphens between word chars dominate; no em-dashes.
        profile = profile_document_style(texts)
        assert profile.prefers_hyphen_dash is True

    def test_em_dash_preference(self):
        texts = [
            "Provide R—rated assemblies—see schedule.",
            "Install 60—80 psi operating pressure.",
            "Use Class B—fire-rated.",
        ]
        # Em-dashes between word chars dominate.
        profile = profile_document_style(texts)
        assert profile.prefers_hyphen_dash is False

    def test_nbsp_in_measurements_detected(self):
        texts = [
            "Install at 60 psi operating pressure.",
            "Provide 12 in. clearance.",
            "Pipe at 5 ft elevation.",
        ]
        profile = profile_document_style(texts)
        assert profile.uses_nbsp_in_measurements is True

    def test_nbsp_default_when_no_measurements(self):
        texts = ["Provide piping.", "Install per code."]
        profile = profile_document_style(texts)
        assert profile.uses_nbsp_in_measurements is False

    def test_nbsp_minority_does_not_flip(self):
        texts = [
            "Install at 60 psi.",
            "Provide 12 in. clearance.",
            "Pipe at 5 ft elevation.",
            "Loading at 100 psi.",
        ]
        # 3 plain space vs 1 NBSP. Profiler should not declare NBSP preference.
        profile = profile_document_style(texts)
        assert profile.uses_nbsp_in_measurements is False


# ---------------------------------------------------------------------------
# Normalization
# ---------------------------------------------------------------------------


_STRAIGHT_PROFILE = DocumentStyleProfile(
    prefers_straight_quotes=True,
    prefers_hyphen_dash=True,
    prefers_ascii_apostrophe=True,
    uses_nbsp_in_measurements=False,
)
_CURLY_PROFILE = DocumentStyleProfile(
    prefers_straight_quotes=False,
    prefers_hyphen_dash=False,
    prefers_ascii_apostrophe=False,
    uses_nbsp_in_measurements=False,
)


class TestNormalizeReplacementText:
    def test_curly_to_straight_double_quotes(self):
        text = "Provide “schedule 40” piping."
        out, changed = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        assert out == 'Provide "schedule 40" piping.'
        assert changed is True

    def test_unicode_apostrophe_to_ascii(self):
        text = "Don’t substitute without engineer’s approval."
        out, changed = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        assert out == "Don't substitute without engineer's approval."
        assert changed is True

    def test_em_dash_with_spaces_to_hyphen(self):
        text = "Provide R-454B — see schedule."
        out, changed = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        assert out == "Provide R-454B - see schedule."
        assert changed is True

    def test_en_dash_with_spaces_to_hyphen(self):
        text = "Install – per code.".replace("–", "–")
        text = "Item A – Item B"
        out, changed = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        assert out == "Item A - Item B"
        assert changed is True

    def test_em_dash_between_words_left_alone(self):
        """Bare em-dash inside a token (e.g., page ranges) is intentional."""
        text = "Pages 12—15 of the standard."
        out, changed = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        assert out == "Pages 12—15 of the standard."
        assert changed is False

    def test_idempotent_double_normalize_equals_single(self):
        text = "Provide “schedule 40” — don’t substitute."
        once, _ = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        twice, second_changed = normalize_replacement_text(once, _STRAIGHT_PROFILE)
        assert once == twice
        assert second_changed is False

    def test_no_op_when_profile_matches_straight(self):
        text = 'Already "straight" with don\'t apostrophe.'
        out, changed = normalize_replacement_text(text, _STRAIGHT_PROFILE)
        assert out == text
        assert changed is False

    def test_no_op_when_profile_is_none(self):
        text = "Provide “schedule 40” piping."
        out, changed = normalize_replacement_text(text, None)
        assert out == text
        assert changed is False

    def test_no_op_when_text_is_empty(self):
        out, changed = normalize_replacement_text("", _STRAIGHT_PROFILE)
        assert out == ""
        assert changed is False

    def test_inverse_straight_to_curly_double_quote(self):
        text = 'Provide "schedule 40" piping.'
        out, changed = normalize_replacement_text(text, _CURLY_PROFILE)
        assert out == "Provide “schedule 40” piping."
        assert changed is True

    def test_inverse_straight_to_curly_apostrophe_internal(self):
        text = "Don't substitute without engineer's approval."
        out, changed = normalize_replacement_text(text, _CURLY_PROFILE)
        assert out == "Don’t substitute without engineer’s approval."
        assert changed is True

    def test_inverse_idempotent(self):
        text = 'Provide "schedule 40" piping with "don\'t" issue.'
        once, _ = normalize_replacement_text(text, _CURLY_PROFILE)
        twice, second_changed = normalize_replacement_text(once, _CURLY_PROFILE)
        assert once == twice
        assert second_changed is False

    def test_nbsp_inserted_in_measurements(self):
        profile = DocumentStyleProfile(
            prefers_straight_quotes=True,
            prefers_hyphen_dash=True,
            prefers_ascii_apostrophe=True,
            uses_nbsp_in_measurements=True,
        )
        text = "Install at 60 psi and 12 in. clearance."
        out, changed = normalize_replacement_text(text, profile)
        assert out == "Install at 60 psi and 12 in. clearance."
        assert changed is True


# ---------------------------------------------------------------------------
# Env-var kill switch
# ---------------------------------------------------------------------------


class TestNormalizeReplacementStyleEnabled:
    def test_default_enabled(self, monkeypatch: pytest.MonkeyPatch):
        monkeypatch.delenv("SPEC_CRITIC_NORMALIZE_REPLACEMENT_STYLE", raising=False)
        assert normalize_replacement_style_enabled() is True

    @pytest.mark.parametrize("value", ["0", "false", "FALSE", "No", "off", " 0 "])
    def test_disabled_via_env(
        self, monkeypatch: pytest.MonkeyPatch, value: str
    ):
        monkeypatch.setenv("SPEC_CRITIC_NORMALIZE_REPLACEMENT_STYLE", value)
        assert normalize_replacement_style_enabled() is False

    @pytest.mark.parametrize("value", ["1", "true", "yes", "on", "anything-else", ""])
    def test_other_values_keep_default(
        self, monkeypatch: pytest.MonkeyPatch, value: str
    ):
        monkeypatch.setenv("SPEC_CRITIC_NORMALIZE_REPLACEMENT_STYLE", value)
        assert normalize_replacement_style_enabled() is True


# ---------------------------------------------------------------------------
# Rollup wiring: apply_edits.execute_edit_plan → DiagnosticsReport
# ---------------------------------------------------------------------------


def _make_simple_finding(*, file_name: str, existing: str, replacement: str):
    """Minimal Finding for the execute_edit_plan rollup test."""
    from src.review.reviewer import Finding

    return Finding(
        severity="HIGH",
        fileName=file_name,
        section="2.1",
        issue="Style mismatch.",
        actionType="EDIT",
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC 2025",
        confidence=0.9,
    )


def test_execute_edit_plan_aggregates_normalize_count_into_diagnostics(
    tmp_path,
):
    """``execute_edit_plan`` rolls the per-spec counter into the diagnostics report.

    Pins the wiring from ``EditReport.replacement_normalized_count`` to
    ``DiagnosticsReport.replacement_text_normalized_count`` so a later
    refactor cannot silently break the user-visible counter.
    """
    from docx import Document

    from src.editing.apply_edits import execute_edit_plan
    from src.input.extractor import extract_text_from_docx
    from src.orchestration.diagnostics import DiagnosticsReport

    source = tmp_path / "spec.docx"
    doc = Document()
    # Pure curly-quote doc so the profile votes curly with no ambiguity.
    doc.add_paragraph("Provide “schedule 40” steel piping.")
    doc.add_paragraph("Confirm “seismic” bracing per ASCE 7-22.")
    doc.add_paragraph("Don’t substitute without engineer’s approval.")
    doc.save(source)

    spec = extract_text_from_docx(source)
    spec.filename = "spec.docx"

    finding = _make_simple_finding(
        file_name="spec.docx",
        existing="ASCE 7-22",
        replacement='Per "ASCE 7-22" with don\'t substitute',
    )

    diagnostics = DiagnosticsReport()
    reports = execute_edit_plan(
        selected_finding_indices=[0],
        all_findings=[finding],
        cross_check_findings=[],
        extracted_specs=[spec],
        source_paths=[source],
        output_dir=tmp_path / "out",
        diagnostics=diagnostics,
    )

    assert len(reports) == 1
    assert reports[0].edits_applied == 1
    assert reports[0].replacement_normalized_count == 1
    assert diagnostics.replacement_text_normalized_count == 1
    # And the to_text() rollup surfaces the counter under the new section.
    text = diagnostics.to_text()
    assert "AUTO-APPLY QUALITY" in text
    assert "Replacement text normalized" in text


# ---------------------------------------------------------------------------
# Phase 5 / Step 5.1 — Verifier correction replaceability sanity check.
#
# The verifier prompt asks for "1-2 sentences explaining the verdict and the
# corrected reference text" — that is NOT the same as "clean replacement text
# suitable for direct substitution into a spec paragraph." Corrections often
# contain explanatory parentheticals, source citations, URLs, or restated
# context. ``edit_locator._resolve_replacement_text`` previously used
# ``verification.correction`` verbatim as the replacement text whenever the
# verdict was ``CORRECTED``; the sanity check below decides whether the
# correction looks replaceable, and the locator falls back to the model's
# original ``replacement_text`` when it does not.
# ---------------------------------------------------------------------------


from src.editing.replacement_style import (
    correction_looks_replaceable,
    use_verifier_correction_as_replacement_enabled,
)


class TestCorrectionLooksReplaceable:
    def test_clean_short_correction_passes(self):
        """Short, prose-only correction matching the original's length is fine."""
        assert (
            correction_looks_replaceable(
                "Comply with ASCE 7-22 for seismic design.",
                "Comply with ASCE 7-16 for seismic design.",
            )
            is True
        )

    def test_clean_correction_passes_when_original_replacement_is_none(self):
        """A clean correction is replaceable even when no original was supplied."""
        # An EDIT proposal would always have a replacement; in legacy paths a
        # bare CORRECTED verdict can still arrive without one. Treat that as
        # the corrected-text path having no length baseline — only the
        # structural checks (URLs, parentheticals) apply.
        assert (
            correction_looks_replaceable(
                "Comply with ASCE 7-22 for seismic design.", None
            )
            is True
        )

    def test_paragraph_length_correction_against_short_original_fails(self):
        """When the correction is ~3× longer than the original, reject it.

        A clean replacement should be roughly the same shape as the model's
        original attempt. A 5-word original answered by a paragraph signals
        the verifier emitted explanation rather than substitution text.
        """
        original = "Use ASCE 7-22."
        correction = (
            "The applicable standard is ASCE 7-22 because the 2025 California "
            "Building Code adopted that revision through Title 24 Part 2, and "
            "the 7-16 reference is obsolete for projects permitted after "
            "January 1 2026."
        )
        assert correction_looks_replaceable(correction, original) is False

    def test_parenthetical_citation_rejected_when_original_had_none(self):
        """A correction adding ``(per Section X.Y)`` is explanatory, not replacement text."""
        assert (
            correction_looks_replaceable(
                "Comply with ASCE 7-22 (per CBC § 1613.1) for seismic design.",
                "Comply with ASCE 7-16 for seismic design.",
            )
            is False
        )

    def test_parenthetical_citation_allowed_when_original_had_one(self):
        """If the original already used a parenthetical, the correction can keep one."""
        assert (
            correction_looks_replaceable(
                "Comply with ASCE 7-22 (latest revision) for seismic design.",
                "Comply with ASCE 7-16 (current revision) for seismic design.",
            )
            is True
        )

    def test_url_rejected(self):
        """URLs in the correction are a smoking gun — never replacement text."""
        assert (
            correction_looks_replaceable(
                "See https://www.iccsafe.org/asce-7-22 for details.",
                "See ASCE 7-22 for details.",
            )
            is False
        )

    def test_url_rejected_even_when_original_had_one(self):
        """URLs in body text don't belong in spec paragraphs; reject either way."""
        # Conservative: even if the original somehow had a URL, we'd rather
        # use the original than risk landing a verifier-cited URL.
        assert (
            correction_looks_replaceable(
                "See https://example.com for details.",
                "See https://other.example.com for details.",
            )
            is False
        )

    def test_current_qualifier_rejected_when_original_had_none(self):
        """``current`` / ``latest`` / ``as of [year]`` are explanatory qualifiers."""
        assert (
            correction_looks_replaceable(
                "Comply with the current ASCE 7-22 for seismic design.",
                "Comply with ASCE 7-16 for seismic design.",
            )
            is False
        )

    def test_latest_qualifier_rejected(self):
        assert (
            correction_looks_replaceable(
                "Comply with the latest ASCE 7-22 for seismic design.",
                "Comply with ASCE 7-16 for seismic design.",
            )
            is False
        )

    def test_as_of_year_qualifier_rejected(self):
        assert (
            correction_looks_replaceable(
                "Comply with ASCE 7-22 as of 2024 for seismic design.",
                "Comply with ASCE 7-16 for seismic design.",
            )
            is False
        )

    def test_current_qualifier_allowed_when_original_had_one(self):
        """If the model's own replacement already used the qualifier, it's fine."""
        assert (
            correction_looks_replaceable(
                "Comply with the current ASCE 7-22 standard.",
                "Comply with the current ASCE 7-16 standard.",
            )
            is True
        )

    def test_empty_correction_is_not_replaceable(self):
        """An empty correction is by definition not replaceable."""
        assert correction_looks_replaceable("", "Comply with ASCE 7-16.") is False

    def test_whitespace_only_correction_is_not_replaceable(self):
        assert (
            correction_looks_replaceable("   \n  ", "Comply with ASCE 7-16.")
            is False
        )


class TestUseVerifierCorrectionAsReplacementEnabled:
    def test_default_disabled(self, monkeypatch: pytest.MonkeyPatch):
        """Default behavior is the new path: do NOT trust the correction verbatim."""
        monkeypatch.delenv(
            "SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT", raising=False
        )
        assert use_verifier_correction_as_replacement_enabled() is False

    @pytest.mark.parametrize("value", ["1", "true", "yes", "on", "TRUE", "On"])
    def test_enabled_via_env(
        self, monkeypatch: pytest.MonkeyPatch, value: str
    ):
        monkeypatch.setenv(
            "SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT", value
        )
        assert use_verifier_correction_as_replacement_enabled() is True

    @pytest.mark.parametrize(
        "value", ["0", "false", "no", "off", "FALSE", " 0 ", ""]
    )
    def test_disabled_via_env(
        self, monkeypatch: pytest.MonkeyPatch, value: str
    ):
        monkeypatch.setenv(
            "SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT", value
        )
        assert use_verifier_correction_as_replacement_enabled() is False


# ---------------------------------------------------------------------------
# edit_candidates wiring: the candidate UI uses the same sanity check.
#
# The candidate UI surfaces the resolved replacement text so the user can
# preview what the auto-apply path will land in the document. If the UI
# kept showing the verifier's parenthetical-citation correction while the
# locator silently swapped in ``replacement_text``, the user would see a
# different preview than the actual applied edit — defeating the point of
# the preview.
# ---------------------------------------------------------------------------


class TestEditCandidateRespectsCorrectionSanityCheck:
    def test_clean_correction_still_surfaces_in_candidate(self):
        """A clean CORRECTED.correction surfaces to the candidate UI verbatim."""
        from src.editing.edit_candidates import classify_edit_candidates
        from src.review.reviewer import Finding
        from src.verification.verifier import VerificationResult

        finding = Finding(
            severity="HIGH",
            fileName="spec.docx",
            section="2.1",
            issue="Cite the correct edition.",
            actionType="EDIT",
            existingText="ASCE 7-16",
            replacementText="ASCE 7-22",
            codeReference="ASCE 7",
            confidence=0.9,
        )
        finding.verification = VerificationResult(
            verdict="CORRECTED",
            correction="ASCE 7-22",
            explanation="Verified",
            sources=["https://example.com/asce-7-22"],
        )

        candidates = classify_edit_candidates([finding])
        # The candidate UI should preview the verifier's clean correction.
        assert candidates[0].replacement_text == "ASCE 7-22"

    def test_parenthetical_correction_falls_back_in_candidate(self):
        """An explanatory parenthetical correction falls back to replacement_text."""
        from src.editing.edit_candidates import classify_edit_candidates
        from src.review.reviewer import Finding
        from src.verification.verifier import VerificationResult

        finding = Finding(
            severity="HIGH",
            fileName="spec.docx",
            section="2.1",
            issue="Cite the correct edition.",
            actionType="EDIT",
            existingText="ASCE 7-16",
            replacementText="ASCE 7-22",
            codeReference="ASCE 7",
            confidence=0.9,
        )
        finding.verification = VerificationResult(
            verdict="CORRECTED",
            correction="ASCE 7-22 (per CBC § 1613.1).",
            explanation="Verified",
            sources=["https://example.com/asce-7-22"],
        )

        candidates = classify_edit_candidates([finding])
        # The UI should preview the model's clean replacement, not the
        # verifier's parenthetical citation. The locator will do the
        # same swap at apply time, so preview matches reality.
        assert candidates[0].replacement_text == "ASCE 7-22"

    def test_env_var_restores_legacy_candidate_preview(
        self, monkeypatch: pytest.MonkeyPatch
    ):
        """``SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT=1`` skips the check."""
        monkeypatch.setenv(
            "SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT", "1"
        )
        from src.editing.edit_candidates import classify_edit_candidates
        from src.review.reviewer import Finding
        from src.verification.verifier import VerificationResult

        finding = Finding(
            severity="HIGH",
            fileName="spec.docx",
            section="2.1",
            issue="Cite the correct edition.",
            actionType="EDIT",
            existingText="ASCE 7-16",
            replacementText="ASCE 7-22",
            codeReference="ASCE 7",
            confidence=0.9,
        )
        finding.verification = VerificationResult(
            verdict="CORRECTED",
            correction="ASCE 7-22 (per CBC § 1613.1).",
            explanation="Verified",
            sources=["https://example.com/asce-7-22"],
        )

        candidates = classify_edit_candidates([finding])
        # Env-on → legacy verbatim preview.
        assert (
            candidates[0].replacement_text == "ASCE 7-22 (per CBC § 1613.1)."
        )
