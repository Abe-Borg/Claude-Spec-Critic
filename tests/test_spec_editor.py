from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.editing.edit_locator import EditLocation, LocatorResult
from src.input.extractor import ParagraphMapping
from src.review.reviewer import Finding
from src.editing.spec_editor import apply_edits_to_spec, build_edit_actions


def _finding(*, action: str = "EDIT", existing: str = "", replacement: str | None = "") -> Finding:
    return Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1.0",
        issue="Issue",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="Code",
        confidence=0.9,
    )


def _locator_result(
    *,
    action: str = "EDIT",
    status: str = "matched",
    body_index: int = 0,
    element_type: str = "paragraph",
    text: str,
    match_start: int,
    match_end: int,
    matched_text: str,
    replacement_text: str | None,
    confidence: float = 1.0,
    row_index: int | None = None,
    severity: str = "HIGH",
    anchor_text: str | None = None,
    insert_position: str | None = None,
) -> LocatorResult:
    mapping = ParagraphMapping(
        body_index=body_index,
        element_type=element_type,
        text=text,
        table_index=0 if element_type == "table_cell" else None,
        row_index=row_index,
        cell_index=None,
    )
    location = EditLocation(
        mapping=mapping,
        match_start=match_start,
        match_end=match_end,
        matched_text=matched_text,
        match_confidence=confidence,
        match_method="exact",
    )
    return LocatorResult(
        finding=Finding(
            severity=severity,
            fileName="spec.docx",
            section="1.0",
            issue="Issue",
            actionType=action,
            existingText=matched_text,
            replacementText=replacement_text,
            codeReference="Code",
            confidence=0.9,
            anchorText=anchor_text,
            insertPosition=insert_position,
        ),
        status=status,
        locations=[location],
        replacement_text=replacement_text,
        action_type=action,
        warning=None,
    )


def test_apply_edits_simple_paragraph_replacement(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide seismic bracing per ASCE 7-16.")
    doc.save(source)

    result = _locator_result(
        text="Provide seismic bracing per ASCE 7-16.",
        match_start=28,
        match_end=37,
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))

    saved = Document(output)
    assert saved.paragraphs[0].text == "Provide seismic bracing per ASCE 7-22."
    assert report.edits_applied == 1
    assert report.edits_failed == 0


def test_apply_edits_multi_run_replacement(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Provide seismic bracing per ")
    b1 = p.add_run("ASCE")
    b1.bold = True
    p.add_run(" ")
    b2 = p.add_run("7-16")
    b2.bold = True
    p.add_run(".")
    doc.save(source)

    result = _locator_result(
        text="Provide seismic bracing per ASCE 7-16.",
        match_start=28,
        match_end=37,
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )
    apply_edits_to_spec(source, output, build_edit_actions([result]))

    saved = Document(output)
    para = saved.paragraphs[0]
    assert para.text == "Provide seismic bracing per ASCE 7-22."
    assert para.runs[0].text == "Provide seismic bracing per "
    assert para.runs[-1].text == "."


def test_apply_edits_delete_entire_paragraph(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Keep this paragraph")
    doc.add_paragraph("Delete this paragraph")
    doc.add_paragraph("Keep this too")
    doc.save(source)

    full = "Delete this paragraph"
    result = _locator_result(
        action="DELETE",
        text=full,
        body_index=1,
        match_start=0,
        match_end=len(full),
        matched_text=full,
        replacement_text=None,
    )
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))

    saved = Document(output)
    assert [p.text for p in saved.paragraphs] == ["Keep this paragraph", "Keep this too"]
    assert report.edits_applied == 1


def test_conflict_resolution_applies_non_overlapping_in_reverse_order(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    text = "Alpha Beta Gamma Delta"
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    first = _locator_result(
        text=text,
        match_start=6,
        match_end=10,
        matched_text="Beta",
        replacement_text="BETA",
    )
    second = _locator_result(
        text=text,
        match_start=17,
        match_end=22,
        matched_text="Delta",
        replacement_text="DELTA",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([first, second]))
    saved = Document(output)

    assert saved.paragraphs[0].text == "Alpha BETA Gamma DELTA"
    assert report.edits_applied == 2


def test_conflict_resolution_skips_both_on_ambiguous_partial_overlap(tmp_path: Path):
    """Chunk D3.1: partial overlap (no containment) is ambiguous — skip both.

    Previously the higher-confidence edit silently won and was applied. Per
    the delta plan, ambiguous overlapping edits must be flagged for manual
    review rather than auto-applied, since picking either intent silently
    discards the other.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    text = "abc def ghi"
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    high = _locator_result(
        text=text,
        match_start=4,
        match_end=7,
        matched_text="def",
        replacement_text="XYZ",
        confidence=0.95,
    )
    low = _locator_result(
        text=text,
        match_start=2,
        match_end=6,
        matched_text="c de",
        replacement_text="1234",
        confidence=0.70,
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([low, high]))
    saved = Document(output)

    # Paragraph unchanged; both edits routed to manual review.
    assert saved.paragraphs[0].text == "abc def ghi"
    assert report.edits_applied == 0
    assert report.edits_skipped == 2
    for outcome in report.outcomes:
        assert outcome.status == "skipped"
        assert "ambiguous" in outcome.detail.lower()
        assert "manual review" in outcome.detail.lower()


def test_conflict_resolution_prefers_broader_subsuming_edit(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    text = "Pipe markers include refrigerant piping and condensate piping using R454B notation."
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    gripe = _locator_result(
        text=text,
        match_start=text.index("R454B"),
        match_end=text.index("R454B") + len("R454B"),
        matched_text="R454B",
        replacement_text="R-454B",
        confidence=1.0,
        severity="GRIPES",
    )
    medium = _locator_result(
        text=text,
        match_start=0,
        match_end=len(text),
        matched_text=text,
        replacement_text="Pipe markers shall separate refrigerant piping from condensate piping and use R-454B notation.",
        confidence=1.0,
        severity="MEDIUM",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([gripe, medium]))
    saved = Document(output)

    assert saved.paragraphs[0].text.startswith("Pipe markers shall separate refrigerant piping")
    assert report.edits_applied == 1
    assert report.edits_skipped == 1
    # Step 4.1 refined the skipped detail to make the containment relation
    # explicit. The narrower's "R-454B" is in the broader's replacement, so
    # the new detail reads "intent preserved" instead of the previous
    # generic "broader/higher-priority" wording.
    skipped_details = [
        o.detail for o in report.outcomes if o.status == "skipped"
    ]
    assert any("intent preserved" in d.lower() for d in skipped_details)


def test_strict_containment_preserves_narrower_intent_in_broader_replacement(tmp_path: Path):
    """Step 4.1: narrower edit's replacement appears in broader's → 'intent preserved'.

    A GRIPES typo fix nested inside a MEDIUM paragraph rewrite. The broader
    edit's replacement text contains the narrower's correction verbatim, so
    the narrower's intent is preserved by the broader's application and the
    skipped outcome should note that. ``contained_edit_lost_intent`` stays
    False.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    text = "Pipe markers include refrigerant piping and condensate piping using R454B notation."
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    gripe = _locator_result(
        text=text,
        match_start=text.index("R454B"),
        match_end=text.index("R454B") + len("R454B"),
        matched_text="R454B",
        replacement_text="R-454B",
        confidence=1.0,
        severity="GRIPES",
    )
    medium = _locator_result(
        text=text,
        match_start=0,
        match_end=len(text),
        matched_text=text,
        replacement_text="Pipe markers shall separate refrigerant piping from condensate piping and use R-454B notation.",
        confidence=1.0,
        severity="MEDIUM",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([gripe, medium]))

    # Broader wins, narrower is skipped. The narrower's correction
    # ("R-454B") appears in the broader's replacement, so intent is
    # preserved. No diagnostics escalation.
    assert report.edits_applied == 1
    assert report.edits_skipped == 1
    assert report.contained_edits_lost_intent_count == 0
    skipped = [o for o in report.outcomes if o.status == "skipped"]
    assert len(skipped) == 1
    assert skipped[0].contained_edit_lost_intent is False
    assert "intent preserved" in skipped[0].detail.lower()


def test_strict_containment_loses_narrower_intent_when_broader_replaces_it(tmp_path: Path):
    """Step 4.1: narrower edit's replacement NOT in broader's → 'manual review recommended'.

    A GRIPES typo fix nested inside a MEDIUM paragraph rewrite where the
    broader edit's replacement text discards the narrower's correction. The
    narrower's intent is lost — the broader still wins (more agency to the
    user) but the skipped outcome must flag the loss so it shows up in the
    report and diagnostics.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    text = "Pipe markers include refrigerant piping and condensate piping using R454B notation."
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    # Narrower: typo fix R454B → R-454B
    gripe = _locator_result(
        text=text,
        match_start=text.index("R454B"),
        match_end=text.index("R454B") + len("R454B"),
        matched_text="R454B",
        replacement_text="R-454B",
        confidence=1.0,
        severity="GRIPES",
    )
    # Broader: rewrites the whole sentence but uses a different refrigerant.
    # The narrower's "R-454B" is not preserved here — broader picked
    # R-32 instead, so the GRIPES typo fix is silently discarded by
    # the broader edit.
    medium = _locator_result(
        text=text,
        match_start=0,
        match_end=len(text),
        matched_text=text,
        replacement_text="Pipe markers shall separate refrigerant piping from condensate piping and use R-32 notation.",
        confidence=1.0,
        severity="MEDIUM",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([gripe, medium]))

    assert report.edits_applied == 1
    assert report.edits_skipped == 1
    assert report.contained_edits_lost_intent_count == 1
    skipped = [o for o in report.outcomes if o.status == "skipped"]
    assert len(skipped) == 1
    assert skipped[0].contained_edit_lost_intent is True
    detail = skipped[0].detail.lower()
    assert "not preserved" in detail or "manual review" in detail


def test_strict_containment_narrower_action_loses_when_broader_processed_later(tmp_path: Path):
    """Step 4.1: same logic when the narrower edit is the one in `accepted`.

    The conflict resolver processes actions in descending-start-offset
    order. When the narrower edit is processed first and the broader
    arrives second, the broader still wins; the narrower already in
    `accepted` is the one that gets skipped. This test exercises the other
    branch (winner is `action`, `overlap` is the loser) of the resolver.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    # Construct the text so the narrower edit has a higher start offset than
    # the broader edit. Reverse-sort processes the narrower first.
    text = "Comply with NFPA 13 and other applicable standards in Section 21 13 13."
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    # Narrower edit (higher start offset, processed first by reverse-sort):
    # typo fix on the section number.
    gripe = _locator_result(
        text=text,
        match_start=text.index("Section 21 13 13"),
        match_end=text.index("Section 21 13 13") + len("Section 21 13 13"),
        matched_text="Section 21 13 13",
        replacement_text="Section 21 13 16",
        confidence=1.0,
        severity="GRIPES",
    )
    # Broader edit (covers the whole paragraph, processed second).
    # Replacement removes the section reference entirely.
    medium = _locator_result(
        text=text,
        match_start=0,
        match_end=len(text),
        matched_text=text,
        replacement_text="Comply with NFPA 13 and other applicable industry standards.",
        confidence=1.0,
        severity="MEDIUM",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([gripe, medium]))

    assert report.edits_applied == 1
    assert report.edits_skipped == 1
    assert report.contained_edits_lost_intent_count == 1
    skipped = [o for o in report.outcomes if o.status == "skipped"]
    assert len(skipped) == 1
    assert skipped[0].contained_edit_lost_intent is True


def test_table_cell_edit_updates_target_only(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "R1C1"
    table.cell(0, 1).text = "Allowance Amount"
    table.cell(0, 2).text = "$10,000"
    doc.save(source)

    joined = "R1C1 | Allowance Amount | $10,000"
    start = joined.find("Allowance Amount")
    result = _locator_result(
        text=joined,
        element_type="table_cell",
        row_index=0,
        match_start=start,
        match_end=start + len("Allowance Amount"),
        matched_text="Allowance Amount",
        replacement_text="Allowance Value",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    table = saved.tables[0]

    assert table.cell(0, 0).text == "R1C1"
    assert table.cell(0, 1).text == "Allowance Value"
    assert table.cell(0, 2).text == "$10,000"
    assert report.edits_applied == 1


def test_same_source_and_output_raises_value_error(tmp_path: Path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("Hello")
    doc.save(source)

    result = _locator_result(
        text="Hello",
        match_start=0,
        match_end=5,
        matched_text="Hello",
        replacement_text="Hi",
    )

    with pytest.raises(ValueError):
        apply_edits_to_spec(source, source, build_edit_actions([result]))


# ---------------------------------------------------------------------------
# Phase 1 / Step 1.1 — Replacement-text typographic normalization
#
# Integration tests verifying that ``apply_edits_to_spec`` profiles the
# source document's typography and normalizes the model's replacement text
# to match before writing it into the file. The ``replacement_style``
# unit tests in ``test_replacement_style.py`` cover the pure functions;
# these tests cover the wiring through the edit pipeline.
# ---------------------------------------------------------------------------


def test_replacement_normalized_to_curly_when_doc_uses_curly_quotes(tmp_path: Path):
    """Model emits straight quotes; source doc is curly. Applied edit keeps curly."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    # Pure curly-quote document to make the profiler's vote unambiguous.
    doc.add_paragraph("Provide “schedule 40” steel piping per project standards.")
    doc.add_paragraph("Don’t substitute without engineer’s approval.")
    doc.add_paragraph("Confirm “seismic” bracing per ASCE 7-22.")
    target_text = "ASCE 7-22"
    paragraph_index = 2
    doc.save(source)

    full_text = "Confirm “seismic” bracing per ASCE 7-22."
    start = full_text.index(target_text)
    result = _locator_result(
        text=full_text,
        body_index=paragraph_index,
        match_start=start,
        match_end=start + len(target_text),
        matched_text=target_text,
        # Model's replacement uses straight ASCII quotes.
        replacement_text='Per "ASCE 7-22" with don\'t substitute',
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    # The applied edit should have rewritten straight " to curly " and
    # straight ' to curly ' so the new sentence matches the doc's style.
    applied_text = saved.paragraphs[paragraph_index].text
    assert "“ASCE 7-22”" in applied_text
    assert "don’t" in applied_text
    assert report.replacement_normalized_count == 1


def test_replacement_normalized_to_straight_when_doc_uses_straight_quotes(tmp_path: Path):
    """Inverse of the above: doc uses straight, model emits curly."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph('Provide "schedule 40" steel piping per project standards.')
    doc.add_paragraph("Don't substitute without engineer's approval.")
    doc.add_paragraph('Confirm "seismic" bracing per ASCE 7-22.')
    target_text = "ASCE 7-22"
    paragraph_index = 2
    doc.save(source)

    full_text = 'Confirm "seismic" bracing per ASCE 7-22.'
    start = full_text.index(target_text)
    result = _locator_result(
        text=full_text,
        body_index=paragraph_index,
        match_start=start,
        match_end=start + len(target_text),
        matched_text=target_text,
        # Claude likes curly — should land as straight in this doc.
        replacement_text="Per “ASCE 7-22” with don’t substitute",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    applied_text = saved.paragraphs[paragraph_index].text
    assert '"ASCE 7-22"' in applied_text
    assert "don't" in applied_text
    assert report.replacement_normalized_count == 1


def test_replacement_unchanged_when_already_matches(tmp_path: Path):
    """When the replacement already matches the doc style, no normalize counter."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph('Provide "schedule 40" steel piping per project standards.')
    full_text = 'Provide "schedule 40" steel piping per project standards.'
    doc.save(source)

    start = full_text.index("schedule 40")
    result = _locator_result(
        text=full_text,
        match_start=start,
        match_end=start + len("schedule 40"),
        matched_text="schedule 40",
        replacement_text="schedule 80",  # plain ASCII; matches profile.
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))

    assert report.edits_applied == 1
    assert report.replacement_normalized_count == 0


def test_replacement_normalization_disabled_via_env_var(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """SPEC_CRITIC_NORMALIZE_REPLACEMENT_STYLE=0 keeps the model's text verbatim."""
    monkeypatch.setenv("SPEC_CRITIC_NORMALIZE_REPLACEMENT_STYLE", "0")
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide “schedule 40” steel piping per project standards.")
    full_text = "Provide “schedule 40” steel piping per project standards."
    doc.save(source)

    start = full_text.index("schedule 40")
    result = _locator_result(
        text=full_text,
        match_start=start,
        match_end=start + len("schedule 40"),
        matched_text="schedule 40",
        replacement_text="schedule 80",
    )
    # Override style profile defaults via an explicit passthrough to confirm
    # the env var short-circuits the profile creation path too.
    report = apply_edits_to_spec(source, output, build_edit_actions([result]))

    assert report.edits_applied == 1
    assert report.replacement_normalized_count == 0


# ---------------------------------------------------------------------------
# Phase 1 / Step 1.2 — Punctuation boundary preservation
#
# When the model's existingText includes terminal punctuation but the
# replacement_text does not (or vice versa), the applied edit used to
# silently drop or double the punctuation. The fix is a deterministic
# pass that compares the trailing character of existing vs replacement
# and inspects the character immediately after the match in the live
# paragraph to decide whether to add or strip a terminating punctuation
# mark on the replacement.
# ---------------------------------------------------------------------------


def test_punctuation_boundary_preserves_trailing_period(tmp_path: Path):
    """existing ends with '.', replacement does not -> add the period back."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide seismic bracing per ASCE 7-16.")
    doc.save(source)

    full = "Provide seismic bracing per ASCE 7-16."
    start = full.index("per ASCE 7-16.")
    result = _locator_result(
        text=full,
        match_start=start,
        match_end=start + len("per ASCE 7-16."),
        matched_text="per ASCE 7-16.",
        replacement_text="per ASCE 7-22",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert saved.paragraphs[0].text == "Provide seismic bracing per ASCE 7-22."
    assert report.punctuation_boundary_fixed_count == 1


def test_punctuation_boundary_prevents_doubled_period(tmp_path: Path):
    """existing and replacement both end with '.' next char is also '.' -> strip."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    # Construct a sentence whose terminal period is the char after the match.
    doc.add_paragraph("Confirm sec 5.")
    doc.save(source)

    full = "Confirm sec 5."
    start = full.index("sec 5")
    result = _locator_result(
        text=full,
        match_start=start,
        match_end=start + len("sec 5"),
        matched_text="sec 5",
        replacement_text="sec 6.",  # already ends with period, next char is also period.
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert saved.paragraphs[0].text == "Confirm sec 6."  # not "Confirm sec 6.."
    assert report.punctuation_boundary_fixed_count == 1


def test_punctuation_boundary_noop_when_already_correct(tmp_path: Path):
    """No boundary issue: replacement ends as expected, no counter bump."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide ASCE 7-16 references.")
    doc.save(source)

    full = "Provide ASCE 7-16 references."
    start = full.index("ASCE 7-16")
    result = _locator_result(
        text=full,
        match_start=start,
        match_end=start + len("ASCE 7-16"),
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert saved.paragraphs[0].text == "Provide ASCE 7-22 references."
    assert report.punctuation_boundary_fixed_count == 0


def test_punctuation_boundary_preserves_trailing_comma(tmp_path: Path):
    """existing ends with ',', replacement does not -> add the comma back."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide ASCE 7-16, ASHRAE 90.1 and CBC 2025.")
    doc.save(source)

    full = "Provide ASCE 7-16, ASHRAE 90.1 and CBC 2025."
    start = full.index("ASCE 7-16,")
    result = _locator_result(
        text=full,
        match_start=start,
        match_end=start + len("ASCE 7-16,"),
        matched_text="ASCE 7-16,",
        replacement_text="ASCE 7-22",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert saved.paragraphs[0].text == "Provide ASCE 7-22, ASHRAE 90.1 and CBC 2025."
    assert report.punctuation_boundary_fixed_count == 1


def test_punctuation_boundary_does_not_add_period_when_next_is_word(
    tmp_path: Path,
):
    """The fix only adds punctuation back when the next char is whitespace or end-of-paragraph.

    If the match was mid-word (unusual but possible with normalized
    matching), don't fabricate a period that interrupts the next token.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide foo.bar baseline.")
    doc.save(source)

    full = "Provide foo.bar baseline."
    start = full.index("foo.")
    result = _locator_result(
        text=full,
        match_start=start,
        match_end=start + len("foo."),
        matched_text="foo.",
        replacement_text="qux",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    # Period is *not* re-added because the char after the match is 'b', not
    # whitespace or end-of-paragraph. The fix is conservative.
    assert saved.paragraphs[0].text == "Provide quxbar baseline."
    assert report.punctuation_boundary_fixed_count == 0


def test_punctuation_boundary_fix_disabled_via_env(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """SPEC_CRITIC_PUNCTUATION_BOUNDARY_FIX=0 disables the fix entirely."""
    monkeypatch.setenv("SPEC_CRITIC_PUNCTUATION_BOUNDARY_FIX", "0")
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide seismic bracing per ASCE 7-16.")
    doc.save(source)

    full = "Provide seismic bracing per ASCE 7-16."
    start = full.index("per ASCE 7-16.")
    result = _locator_result(
        text=full,
        match_start=start,
        match_end=start + len("per ASCE 7-16."),
        matched_text="per ASCE 7-16.",
        replacement_text="per ASCE 7-22",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    # With the fix off, the terminal period is lost.
    assert saved.paragraphs[0].text == "Provide seismic bracing per ASCE 7-22"
    assert report.punctuation_boundary_fixed_count == 0


# ---------------------------------------------------------------------------
# Phase 1 / Step 1.3 — Whole-paragraph DELETE for table cells
#
# _is_whole_paragraph_delete used to require element_type=="paragraph",
# so a DELETE covering the entire matched paragraph inside a table cell
# fell through to substring deletion. The substring deletion clears the
# paragraph's text but leaves the empty <w:p> in the cell, which Word
# renders as a blank line. The fix removes the paragraph element from
# its cell when (a) the cell has more than one paragraph, or leaves it
# empty when removing it would violate Word's "every cell needs at
# least one paragraph" rule.
# ---------------------------------------------------------------------------


def _make_table_with_two_para_cell(tmp_path: Path) -> Path:
    """One-row table whose only cell contains two paragraphs."""
    source = tmp_path / "two_para_cell.docx"
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Use python-docx's first auto-paragraph for the first line so we
    # control the paragraph ordering deterministically.
    cell.paragraphs[0].text = "Delete this header line"
    cell.add_paragraph("Keep this body line")
    doc.save(source)
    return source


def test_table_cell_whole_paragraph_delete_removes_paragraph_element(
    tmp_path: Path,
):
    """Whole-paragraph DELETE on a cell with multiple paragraphs removes the element."""
    source = _make_table_with_two_para_cell(tmp_path)
    output = tmp_path / "output.docx"

    target = "Delete this header line"
    # Match position 0 of the cell paragraph; the locator-supplied
    # match_start/match_end span the whole paragraph text.
    result = _locator_result(
        action="DELETE",
        element_type="table_cell",
        body_index=1,
        row_index=0,
        text=f"{target} | ",  # unused for the resolver, see _resolve_cell_and_offsets
        match_start=0,
        match_end=len(target),
        matched_text=target,
        replacement_text=None,
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    cell_paragraphs = saved.tables[0].cell(0, 0).paragraphs
    cell_texts = [p.text for p in cell_paragraphs]
    # Cell now has exactly ONE paragraph (the body line). No empty
    # placeholder paragraph above it.
    assert cell_texts == ["Keep this body line"]


def test_table_cell_whole_paragraph_delete_when_only_paragraph_keeps_empty_para(
    tmp_path: Path,
):
    """Whole-paragraph DELETE on a cell with one paragraph leaves an empty one.

    Word requires every cell to contain at least one paragraph; removing
    the only one would produce an invalid <w:tc> element. The fix is to
    clear the paragraph's text via the existing substring path when the
    cell has just one paragraph left.
    """
    source = tmp_path / "single_para_cell.docx"
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "Delete me entirely"
    doc.save(source)

    output = tmp_path / "output.docx"
    target = "Delete me entirely"
    result = _locator_result(
        action="DELETE",
        element_type="table_cell",
        body_index=1,
        row_index=0,
        text=f"{target}",
        match_start=0,
        match_end=len(target),
        matched_text=target,
        replacement_text=None,
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    cell_paragraphs = saved.tables[0].cell(0, 0).paragraphs
    # Exactly one paragraph remains; its text is empty. Word treats this
    # as a valid (empty) cell, not as a missing structural element.
    assert len(cell_paragraphs) == 1
    assert cell_paragraphs[0].text == ""


def test_table_cell_partial_delete_still_uses_substring_path(tmp_path: Path):
    """Partial DELETE (not whole paragraph) is unchanged — substring removal still applies."""
    source = tmp_path / "partial_delete.docx"
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "Keep prefix and delete this suffix"
    doc.save(source)

    output = tmp_path / "output.docx"
    target = " and delete this suffix"
    full = "Keep prefix and delete this suffix"
    start = full.index(target)
    result = _locator_result(
        action="DELETE",
        element_type="table_cell",
        body_index=1,
        row_index=0,
        text=full,
        match_start=start,
        match_end=start + len(target),
        matched_text=target,
        replacement_text=None,
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    cell_paragraphs = saved.tables[0].cell(0, 0).paragraphs
    # Paragraph survives, just with the suffix gone.
    assert len(cell_paragraphs) == 1
    assert cell_paragraphs[0].text == "Keep prefix"


# ---------------------------------------------------------------------------
# Phase 2 / Step 2.1 — Strip list / numbering properties from inherited pPr
#
# When an ADD action's anchor paragraph is part of a numbered list, the
# legacy code deep-copied the anchor's <w:pPr> verbatim, so the inserted
# paragraph joined the list. Word auto-renumbered the items after it,
# the new paragraph inherited indentation/outline level it should not
# have, and the visual result was a list item that wasn't supposed to
# exist. The fix strips <w:numPr>, <w:outlineLvl>, and <w:pBdr> from
# the cloned pPr unconditionally and strips <w:ind> when the inserted
# text doesn't itself read as list-shaped.
# ---------------------------------------------------------------------------


def _make_numbered_list_anchor_spec(tmp_path: Path) -> Path:
    """Three-item numbered list using w:numPr so the inserted paragraph
    can be checked for list-membership inheritance."""
    source = tmp_path / "numbered_anchor.docx"
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    for text in ("First list item.", "Second list item.", "Third list item."):
        para = doc.add_paragraph(text)
        ppr = para._element.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        ppr.append(num_pr)
        # Outline level + paragraph border to exercise the broader strip
        # rules. <w:ind> is added without a list prefix in the insert
        # text so the inserted paragraph should drop it.
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "2")
        ppr.append(outline)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ppr.append(ind)
        border = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        border.append(top)
        ppr.append(border)
    doc.save(source)
    return source


def _paragraph_ppr_children(paragraph) -> set[str]:
    """Return the local-name set of every <w:pPr> child element."""
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return set()
    return {child.tag.split("}", 1)[-1] for child in ppr}


def test_add_inserted_paragraph_strips_numbering_from_inherited_ppr(
    tmp_path: Path,
):
    """ADD next to a numbered-list anchor inherits style but NOT numPr."""
    source = _make_numbered_list_anchor_spec(tmp_path)
    output = tmp_path / "output.docx"

    # Anchor is "Second list item." at body_index=2 (PART line at 0,
    # first item at 1, second at 2). ADD "after" inserts a sibling
    # paragraph that should NOT pick up <w:numPr> or <w:outlineLvl>.
    anchor_text = "Second list item."
    result = _locator_result(
        action="ADD",
        body_index=2,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="Inserted commentary on second item.",
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    assert report.edits_applied == 1

    # The inserted paragraph sits immediately after the original anchor
    # paragraph index. Body indices: PART(0), item1(1), item2(2),
    # INSERTED(3), item3(4).
    inserted = saved.paragraphs[3]
    assert inserted.text == "Inserted commentary on second item."

    children = _paragraph_ppr_children(inserted)
    # Hard guarantee: list-numbering and outline-level dropped.
    assert "numPr" not in children
    assert "outlineLvl" not in children
    # Paragraph border is also stripped — see the rationale in
    # _clean_inherited_ppr.
    assert "pBdr" not in children
    # Indent is stripped because the insert text does NOT look
    # list-shaped (no leading "A.", "1.", "•", etc.).
    assert "ind" not in children

    # The neighbors still have their list structure.
    item_one = saved.paragraphs[1]
    item_three = saved.paragraphs[4]
    assert "numPr" in _paragraph_ppr_children(item_one)
    assert "numPr" in _paragraph_ppr_children(item_three)


def test_add_inserted_list_shaped_text_keeps_indent(tmp_path: Path):
    """When the inserted text reads as a list item, <w:ind> survives.

    The numbering/outline strip still runs (we never want to silently
    extend a numbered list), but indentation is preserved because the
    visual intent of the inserted item matches the anchor's indented
    list context.
    """
    source = _make_numbered_list_anchor_spec(tmp_path)
    output = tmp_path / "output.docx"

    anchor_text = "Second list item."
    result = _locator_result(
        action="ADD",
        body_index=2,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        # Literal "A. " prefix tells the cleaner the inserted paragraph
        # is itself list-shaped; the inherited <w:ind> stays so it
        # visually aligns with the surrounding items.
        replacement_text="A. Inserted list item for the second item context.",
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    assert report.edits_applied == 1

    inserted = saved.paragraphs[3]
    children = _paragraph_ppr_children(inserted)
    # Numbering and outline level are always stripped — we trust the
    # literal prefix, not the list semantics.
    assert "numPr" not in children
    assert "outlineLvl" not in children
    # Indent inherited because the insert text reads list-shaped.
    assert "ind" in children


def test_add_inherits_list_numbering_env_flag_reverts_to_legacy(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """SPEC_CRITIC_ADD_INHERITS_LIST_NUMBERING=1 preserves legacy pPr deepcopy."""
    monkeypatch.setenv("SPEC_CRITIC_ADD_INHERITS_LIST_NUMBERING", "1")
    source = _make_numbered_list_anchor_spec(tmp_path)
    output = tmp_path / "output.docx"

    anchor_text = "Second list item."
    result = _locator_result(
        action="ADD",
        body_index=2,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="Inserted commentary on second item.",
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    assert report.edits_applied == 1

    inserted = saved.paragraphs[3]
    children = _paragraph_ppr_children(inserted)
    # Legacy behavior: inserted paragraph inherits the full pPr.
    assert "numPr" in children
    assert "outlineLvl" in children


def test_add_inserted_paragraph_inherits_style_id(tmp_path: Path):
    """When the anchor's pPr has a <w:pStyle>, the inserted paragraph keeps it.

    Step 2.1 only strips list/numbering machinery; the paragraph style id
    is what binds the inserted paragraph to the document's font/size
    conventions and must survive.
    """
    source = tmp_path / "styled_anchor.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    para = doc.add_paragraph("Body paragraph with explicit style.")
    ppr = para._element.get_or_add_pPr()
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), "BodyText")
    ppr.append(pstyle)
    # Add justification too — that is preserved as well.
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    ppr.append(jc)
    doc.save(source)

    anchor_text = "Body paragraph with explicit style."
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="Sibling paragraph that should inherit style.",
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    assert report.edits_applied == 1

    inserted = saved.paragraphs[2]
    # In the styled-anchor case the anchor has a pStyle, so
    # ``_reference_style_for_text`` short-circuits at the styled
    # build_paragraph_element path which writes a fresh pPr that only
    # carries the pStyle. Verify the inserted paragraph carries that
    # style id rather than picking up the legacy unbounded inheritance.
    children = _paragraph_ppr_children(inserted)
    assert "pStyle" in children
    pstyle_el = inserted._element.find(qn("w:pPr")).find(qn("w:pStyle"))
    assert pstyle_el.get(qn("w:val")) == "BodyText"


# ---------------------------------------------------------------------------
# Phase 2 / Step 2.2 — Refuse to guess ADD position
#
# When an ADD's ``insertPosition`` is not explicitly "before" or "after",
# the legacy heuristic compared normalized text but sliced raw bytes,
# producing inserted paragraphs that contained a chopped fragment of
# the anchor at their start when anchor/replacement differed in
# whitespace, dash style, or case. The parser already demotes ADD
# findings without a usable insertPosition at parse time (Chunk 7), so
# reaching the apply layer implies a legacy resume payload or a
# directly-constructed Finding bypassing the parser. The defensive
# refusal in ``_apply_add_action`` keeps the visual bug out of the
# output document either way.
# ---------------------------------------------------------------------------


def test_add_without_explicit_insert_position_is_skipped(tmp_path: Path):
    """ADD finding reaching apply layer without insertPosition is refused."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    doc.add_paragraph("Anchor paragraph for ADD.")
    doc.save(source)

    anchor_text = "Anchor paragraph for ADD."
    # No insert_position passed — defaults to None in the helper, which
    # mirrors a legacy resume payload that bypassed the parser.
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="Anchor paragraph for ADD. Plus appended text.",
        anchor_text=anchor_text,
        insert_position=None,
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    # No edit applied; finding routed to manual review with an
    # explanatory skip detail.
    assert report.edits_applied == 0
    assert report.edits_skipped == 1
    assert len(saved.paragraphs) == 2  # nothing inserted
    outcome = report.outcomes[0]
    assert outcome.status == "skipped"
    assert "insertPosition" in outcome.detail
    assert "manual review" in outcome.detail.lower()
    assert outcome.add_demoted_missing_position is True
    assert report.add_demoted_missing_position_count == 1


def test_add_with_explicit_after_inserts_correctly(tmp_path: Path):
    """ADD with explicit insert_position='after' inserts new paragraph after anchor."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    doc.add_paragraph("Anchor paragraph.")
    doc.add_paragraph("Trailing paragraph.")
    doc.save(source)

    anchor_text = "Anchor paragraph."
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="New paragraph between anchor and trailing.",
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert [p.text for p in saved.paragraphs] == [
        "PART 2 PRODUCTS",
        "Anchor paragraph.",
        "New paragraph between anchor and trailing.",
        "Trailing paragraph.",
    ]
    assert report.add_demoted_missing_position_count == 0


def test_add_with_explicit_before_inserts_correctly(tmp_path: Path):
    """ADD with explicit insert_position='before' inserts new paragraph before anchor."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    doc.add_paragraph("Anchor paragraph.")
    doc.save(source)

    anchor_text = "Anchor paragraph."
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="New paragraph before anchor.",
        anchor_text=anchor_text,
        insert_position="before",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert [p.text for p in saved.paragraphs] == [
        "PART 2 PRODUCTS",
        "New paragraph before anchor.",
        "Anchor paragraph.",
    ]


def test_add_skipped_count_aggregates_into_diagnostics(tmp_path: Path):
    """The per-spec EditReport counter rolls up into DiagnosticsReport.

    The defensive refusal in ``_apply_add_action`` is only reachable
    when a LocatorResult is constructed without going through
    ``locate_edit`` (which short-circuits at ``Finding.as_edit_proposal``
    via parse-time validation). Build the EditAction by hand, run the
    per-spec edit pass, then simulate the
    ``apply_edits.execute_edit_plan`` aggregation pattern so we know
    the rollup line landed.
    """
    from src.orchestration.diagnostics import DiagnosticsReport

    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    doc.add_paragraph("Anchor paragraph for ADD.")
    doc.save(source)

    anchor_text = "Anchor paragraph for ADD."
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text="Anchor paragraph for ADD. Plus appended text.",
        anchor_text=anchor_text,
        insert_position=None,
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    assert report.add_demoted_missing_position_count == 1

    diag = DiagnosticsReport()
    diag.add_demoted_missing_position_count += (
        report.add_demoted_missing_position_count
    )
    assert diag.add_demoted_missing_position_count == 1


# ---------------------------------------------------------------------------
# Phase 2 / Step 2.3 — Smarter paragraph splitting for inserted content
#
# Legacy ``_split_insert_paragraphs`` split only on blank-line separators
# (``\n\s*\n+``). Content with single newlines between items collapsed
# into one paragraph, which Word rendered as a single paragraph with
# embedded soft breaks. The fix distinguishes three shapes:
#   1. Double-newline separators -> definitely separate paragraphs.
#   2. Single-newline-separated lines that all read as list items
#      (A./1./•/–/-) -> separate paragraphs.
#   3. Otherwise -> soft breaks inside one paragraph; collapse internal
#      whitespace to a single space.
# ---------------------------------------------------------------------------


def test_split_insert_paragraphs_unchanged_on_double_newline():
    from src.editing.spec_editor import _split_insert_paragraphs

    text = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
    assert _split_insert_paragraphs(text) == [
        "First paragraph.",
        "Second paragraph.",
        "Third paragraph.",
    ]


def test_split_insert_paragraphs_splits_list_with_single_newlines():
    """Single newlines between items that all carry list prefixes -> split."""
    from src.editing.spec_editor import _split_insert_paragraphs

    text = "A. First item of the list.\nB. Second item.\nC. Third item."
    assert _split_insert_paragraphs(text) == [
        "A. First item of the list.",
        "B. Second item.",
        "C. Third item.",
    ]


def test_split_insert_paragraphs_joins_prose_with_single_newlines():
    """Single newlines inside prose collapse to a single space."""
    from src.editing.spec_editor import _split_insert_paragraphs

    text = (
        "This is a long sentence that\n"
        "the model split across\n"
        "three soft breaks."
    )
    assert _split_insert_paragraphs(text) == [
        "This is a long sentence that the model split across three soft breaks."
    ]


def test_split_insert_paragraphs_mixed_chunks():
    """Double-newline chunks delimit paragraphs; each chunk is then
    classified (list vs prose) independently."""
    from src.editing.spec_editor import _split_insert_paragraphs

    text = (
        "Intro prose split\nacross two soft lines.\n\n"
        "1. First numbered item.\n2. Second numbered item.\n\n"
        "Trailing prose paragraph."
    )
    assert _split_insert_paragraphs(text) == [
        "Intro prose split across two soft lines.",
        "1. First numbered item.",
        "2. Second numbered item.",
        "Trailing prose paragraph.",
    ]


def test_split_insert_paragraphs_single_line_unchanged():
    from src.editing.spec_editor import _split_insert_paragraphs

    assert _split_insert_paragraphs("Lone paragraph text.") == [
        "Lone paragraph text."
    ]


def test_split_insert_paragraphs_bullet_list_with_single_newlines():
    from src.editing.spec_editor import _split_insert_paragraphs

    text = "• First bullet.\n• Second bullet.\n• Third bullet."
    assert _split_insert_paragraphs(text) == [
        "• First bullet.",
        "• Second bullet.",
        "• Third bullet.",
    ]


def test_add_inserts_three_list_items_from_single_newline_text(tmp_path: Path):
    """End-to-end: ADD with single-newline-separated list items writes
    three paragraphs, not one with soft breaks."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    doc.add_paragraph("Anchor before list.")
    doc.save(source)

    anchor_text = "Anchor before list."
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text=(
            "A. First inserted item.\n"
            "B. Second inserted item.\n"
            "C. Third inserted item."
        ),
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert [p.text for p in saved.paragraphs] == [
        "PART 2 PRODUCTS",
        "Anchor before list.",
        "A. First inserted item.",
        "B. Second inserted item.",
        "C. Third inserted item.",
    ]


def test_add_collapses_soft_break_prose_to_single_paragraph(tmp_path: Path):
    """Prose with single newlines lands as one paragraph with single spaces."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    doc.add_paragraph("Anchor before prose.")
    doc.save(source)

    anchor_text = "Anchor before prose."
    result = _locator_result(
        action="ADD",
        body_index=1,
        text=anchor_text,
        match_start=0,
        match_end=len(anchor_text),
        matched_text=anchor_text,
        replacement_text=(
            "This explanatory paragraph spans\n"
            "what the model wrote as three\n"
            "soft-broken lines."
        ),
        anchor_text=anchor_text,
        insert_position="after",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert report.edits_applied == 1
    assert [p.text for p in saved.paragraphs] == [
        "PART 2 PRODUCTS",
        "Anchor before prose.",
        "This explanatory paragraph spans what the model wrote as three soft-broken lines.",
    ]
