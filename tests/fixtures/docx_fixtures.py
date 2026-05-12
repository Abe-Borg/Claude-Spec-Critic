"""Tiny in-memory DOCX builders for edit-safety / locator tests.

Chunk A baseline: Chunk F (edit precondition offset safety) and later
edit-related work need cheap DOCX fixtures so they can exercise the real
``ExtractedSpec`` paragraph_map without parsing large real specs.

These helpers are intentionally minimal:
- ``make_paragraph_spec`` builds a single-section spec from a list of strings.
- ``make_table_spec`` builds a one-table spec with the given cell values.
- ``make_real_world_section_spec`` returns a spec that looks like a
  realistic CSI-section snippet (PART / 1.01 / paragraph), good enough
  for testing locator behavior under anchor matching.

Chunk 9 adds unsafe-markup builders for the new edit-refusal contract:
- ``make_paragraph_with_hyperlink`` injects a w:hyperlink element.
- ``make_paragraph_with_field_code`` injects begin/instr/sep/end field characters.
- ``make_paragraph_with_drawing`` injects a w:drawing element.
- ``make_paragraph_with_tracked_change`` injects a w:ins / w:del element.
- ``make_paragraph_with_comment_range`` injects w:commentRangeStart / End / Reference.
- ``make_paragraph_with_bookmark`` injects w:bookmarkStart / End.

Each helper attaches the unsafe element to an otherwise normal paragraph
whose plain text contains the substring callers pass in; the precondition
revalidation logic can still find the substring, so tests can demonstrate
that the unsafe-markup detector — not a missing precondition — is what
refuses the edit.

Callers receive the on-disk ``Path`` so they can feed it to
``extract_text_from_docx`` and get a populated ``paragraph_map`` back.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def make_paragraph_spec(tmp_path: Path, paragraphs: list[str], *, filename: str = "spec.docx") -> Path:
    """Save a docx with one paragraph per element of ``paragraphs``."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_table_spec(
    tmp_path: Path,
    rows: list[list[str]],
    *,
    filename: str = "table_spec.docx",
    leading_paragraph: str = "PART 2 PRODUCTS",
) -> Path:
    """Save a docx with a single table built from ``rows``."""
    doc = Document()
    doc.add_paragraph(leading_paragraph)
    if not rows:
        raise ValueError("rows must not be empty")
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            table.cell(r_idx, c_idx).text = cell_text
    out = tmp_path / filename
    doc.save(out)
    return out


def make_real_world_section_spec(
    tmp_path: Path,
    *,
    filename: str = "23 21 13 - Hydronic.docx",
    code_year: str = "2022",
) -> Path:
    """Return a small spec that mirrors common CSI-section structure."""
    paragraphs = [
        f"SECTION 23 21 13 - HYDRONIC PIPING",
        "PART 1 GENERAL",
        "1.01 SUMMARY",
        f"A. Comply with California Plumbing Code {code_year} requirements.",
        "1.02 REFERENCES",
        f"A. CBC {code_year} - California Building Code.",
        "PART 2 PRODUCTS",
        "2.01 GENERAL",
        "A. Provide hydronic piping as scheduled.",
        "PART 3 EXECUTION",
        "3.01 INSTALLATION",
        "A. Install per manufacturer's written instructions.",
    ]
    return make_paragraph_spec(tmp_path, paragraphs, filename=filename)


# ---------------------------------------------------------------------------
# Chunk 9 — unsafe-markup paragraph builders.
#
# Each helper crafts a paragraph whose run-text reads as the supplied
# ``text`` string for python-docx purposes (so the locator can still find
# the substring), but whose XML carries the unsafe WordprocessingML
# construct named after the helper. The detect_unsafe_markup() scan in
# spec_editor.py should refuse to mutate these paragraphs.
# ---------------------------------------------------------------------------


def _add_plain_run(paragraph_element, text: str) -> None:
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    run.append(text_el)
    paragraph_element.append(run)


def make_paragraph_with_hyperlink(
    tmp_path: Path,
    *,
    text: str = "See manufacturer datasheet for equipment specifications.",
    filename: str = "hyperlink_spec.docx",
) -> Path:
    """Paragraph whose text contains a w:hyperlink element."""
    doc = Document()
    doc.add_paragraph(f"PART 1 GENERAL")
    para = doc.add_paragraph()
    p_elem = para._element
    # Plain prefix.
    _add_plain_run(p_elem, text + " ")
    # Hyperlink with anchor attribute (no relationship needed for the unsafe-markup test).
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "ref1")
    h_run = OxmlElement("w:r")
    h_text = OxmlElement("w:t")
    h_text.text = "see link"
    h_run.append(h_text)
    hyperlink.append(h_run)
    p_elem.append(hyperlink)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_paragraph_with_field_code(
    tmp_path: Path,
    *,
    text: str = "Reference page number in this paragraph.",
    filename: str = "field_spec.docx",
) -> Path:
    """Paragraph containing a w:fldChar begin / instr / end sequence."""
    doc = Document()
    doc.add_paragraph(f"PART 1 GENERAL")
    para = doc.add_paragraph()
    p_elem = para._element
    _add_plain_run(p_elem, text + " ")
    # Field begin.
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    p_elem.append(begin_run)
    # Field instruction.
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    instr_run.append(instr)
    p_elem.append(instr_run)
    # Field end.
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    p_elem.append(end_run)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_paragraph_with_drawing(
    tmp_path: Path,
    *,
    text: str = "Figure caption associated with embedded drawing.",
    filename: str = "drawing_spec.docx",
) -> Path:
    """Paragraph containing a (minimal) w:drawing element."""
    doc = Document()
    doc.add_paragraph(f"PART 1 GENERAL")
    para = doc.add_paragraph()
    p_elem = para._element
    _add_plain_run(p_elem, text + " ")
    drawing_run = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    drawing_run.append(drawing)
    p_elem.append(drawing_run)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_paragraph_with_tracked_change(
    tmp_path: Path,
    *,
    text: str = "Tracked-change paragraph with proposed insertion.",
    filename: str = "tracked_spec.docx",
) -> Path:
    """Paragraph carrying a w:ins (tracked insertion) element."""
    doc = Document()
    doc.add_paragraph(f"PART 1 GENERAL")
    para = doc.add_paragraph()
    p_elem = para._element
    _add_plain_run(p_elem, text + " ")
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Reviewer")
    ins.set(qn("w:date"), "2026-01-01T00:00:00Z")
    ins_run = OxmlElement("w:r")
    ins_text = OxmlElement("w:t")
    ins_text.text = "tracked addition"
    ins_run.append(ins_text)
    ins.append(ins_run)
    p_elem.append(ins)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_paragraph_with_comment_range(
    tmp_path: Path,
    *,
    text: str = "Paragraph with reviewer comment range attached.",
    filename: str = "comment_spec.docx",
) -> Path:
    """Paragraph wrapped by w:commentRangeStart / End and a w:commentReference."""
    doc = Document()
    doc.add_paragraph(f"PART 1 GENERAL")
    para = doc.add_paragraph()
    p_elem = para._element
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), "0")
    p_elem.append(start)
    _add_plain_run(p_elem, text)
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), "0")
    p_elem.append(end)
    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), "0")
    ref_run.append(ref)
    p_elem.append(ref_run)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_paragraph_with_bookmark(
    tmp_path: Path,
    *,
    text: str = "Paragraph contains a bookmark range marker.",
    filename: str = "bookmark_spec.docx",
) -> Path:
    """Paragraph wrapped by w:bookmarkStart / End markers."""
    doc = Document()
    doc.add_paragraph(f"PART 1 GENERAL")
    para = doc.add_paragraph()
    p_elem = para._element
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), "ref_bookmark")
    p_elem.append(start)
    _add_plain_run(p_elem, text)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    p_elem.append(end)
    out = tmp_path / filename
    doc.save(out)
    return out


def make_table_with_unsafe_cell(
    tmp_path: Path,
    *,
    cell_text: str = "Linked manufacturer datasheet 12345",
    filename: str = "unsafe_table_spec.docx",
) -> Path:
    """Single-row, single-column table whose only cell carries a hyperlink."""
    doc = Document()
    doc.add_paragraph("PART 2 PRODUCTS")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Clear the default paragraph the cell ships with so we control the runs.
    para = cell.paragraphs[0]
    p_elem = para._element
    for child in list(p_elem):
        p_elem.remove(child)
    _add_plain_run(p_elem, cell_text + " ")
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "ref1")
    h_run = OxmlElement("w:r")
    h_text = OxmlElement("w:t")
    h_text.text = "(link)"
    h_run.append(h_text)
    hyperlink.append(h_run)
    p_elem.append(hyperlink)
    out = tmp_path / filename
    doc.save(out)
    return out
