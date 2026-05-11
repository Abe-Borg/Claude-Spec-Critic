"""Chunk F regression tests: edit precondition offset safety.

The recorded ``EditLocation.match_start`` / ``match_end`` can become stale if
an earlier edit in the same apply pass shifts paragraph text. The Chunk F
contract is:

* If the recorded offsets still bracket the expected text, mutate them.
* If the recorded offsets are stale but the expected text is uniquely
  present elsewhere in the live paragraph, mutate the corrected offsets
  (never the stale ones).
* If the expected text is missing or appears more than once in the live
  paragraph, skip the edit — guessing risks replacing the wrong span.

These tests assert each of those branches and prove that stale offsets can
no longer corrupt the wrong slice of a paragraph.
"""

from pathlib import Path

import pytest
from docx import Document

from src.edit_locator import EditLocation, LocatorResult
from src.extractor import ParagraphMapping, extract_text_from_docx
from src.reviewer import Finding
from src.spec_editor import (
    PreconditionResult,
    _precondition_holds_for_paragraph,
    apply_edits_to_spec,
    build_edit_actions,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    action: str = "EDIT",
    existing: str = "",
    replacement: str | None = "",
) -> Finding:
    return Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1.0",
        issue="Issue",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="Code",
        confidence=0.9,
    )


def _locator_result(
    *,
    text: str,
    match_start: int,
    match_end: int,
    matched_text: str,
    replacement_text: str | None,
    body_index: int = 0,
    element_type: str = "paragraph",
    row_index: int | None = None,
    confidence: float = 1.0,
    action: str = "EDIT",
) -> LocatorResult:
    mapping = ParagraphMapping(
        body_index=body_index,
        element_type=element_type,
        text=text,
        table_index=0 if element_type == "table_cell" else None,
        row_index=row_index,
        cell_index=None,
    )
    location = EditLocation(
        mapping=mapping,
        match_start=match_start,
        match_end=match_end,
        matched_text=matched_text,
        match_confidence=confidence,
        match_method="exact",
    )
    return LocatorResult(
        finding=_finding(
            action=action,
            existing=matched_text,
            replacement=replacement_text,
        ),
        status="matched",
        locations=[location],
        replacement_text=replacement_text,
        action_type=action,
        warning=None,
    )


def _live_paragraph(text: str, tmp_path: Path):
    """Materialize a single-paragraph docx and return the live Paragraph object."""
    source = tmp_path / "live.docx"
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)
    return Document(source).paragraphs[0]


# ---------------------------------------------------------------------------
# Direct precondition unit tests
# ---------------------------------------------------------------------------


def test_precondition_returns_recorded_offsets_when_unchanged(tmp_path: Path):
    paragraph = _live_paragraph("Install ASCE 7-16 bracing on every chiller.", tmp_path)
    text = paragraph.text
    start = text.index("ASCE 7-16")
    end = start + len("ASCE 7-16")

    result = _precondition_holds_for_paragraph(paragraph, start, end, "ASCE 7-16")

    assert isinstance(result, PreconditionResult)
    assert result.ok is True
    assert (result.match_start, result.match_end) == (start, end)
    assert "recorded offsets" in result.detail


def test_precondition_corrects_offsets_when_text_shifted_earlier(tmp_path: Path):
    """Recorded offsets point past the expected text after an upstream deletion."""
    paragraph = _live_paragraph("foo and bar.", tmp_path)
    # Pretend the recorded offsets came from a longer paragraph where "bar"
    # used to sit later. The expected text is uniquely present at a different
    # offset now.
    stale_start = 9
    stale_end = 12  # stale slice would be "ar." not "bar"

    result = _precondition_holds_for_paragraph(paragraph, stale_start, stale_end, "bar")

    assert result.ok is True
    actual_start = paragraph.text.index("bar")
    assert (result.match_start, result.match_end) == (actual_start, actual_start + 3)
    assert "offsets corrected" in result.detail


def test_precondition_corrects_offsets_when_text_shifted_later(tmp_path: Path):
    """Recorded offsets fall before the expected text after upstream insertion."""
    paragraph = _live_paragraph("Lead-in text. Replace TARGET token.", tmp_path)
    actual_start = paragraph.text.index("TARGET")
    # Pretend earlier "Lead-in text. " was not present when offsets were
    # recorded; the recorded offsets point to a span before the actual TARGET.
    stale_start = 8
    stale_end = stale_start + len("TARGET")

    result = _precondition_holds_for_paragraph(paragraph, stale_start, stale_end, "TARGET")

    assert result.ok is True
    assert (result.match_start, result.match_end) == (actual_start, actual_start + len("TARGET"))


def test_precondition_skips_when_text_duplicated(tmp_path: Path):
    """Two occurrences in the live paragraph → refuse to guess."""
    paragraph = _live_paragraph("foo and another foo.", tmp_path)
    result = _precondition_holds_for_paragraph(paragraph, 0, 3, "foo")
    # Recorded offsets [0, 3] still match the first "foo" exactly, so the
    # primary branch accepts it. Verify the recorded-offsets path returns the
    # ORIGINAL offsets — it is not allowed to shift to the second occurrence.
    assert result.ok is True
    assert (result.match_start, result.match_end) == (0, 3)

    # But if the recorded offsets are stale (do not match), the duplicate
    # presence should make the precondition fail.
    stale = _precondition_holds_for_paragraph(paragraph, 5, 8, "foo")
    assert stale.ok is False
    assert "appears 2 times" in stale.detail
    assert "manual review" in stale.detail.lower()


def test_precondition_skips_when_text_missing(tmp_path: Path):
    paragraph = _live_paragraph("Only real content.", tmp_path)
    result = _precondition_holds_for_paragraph(paragraph, 0, 5, "ghost")
    assert result.ok is False
    assert "no longer present" in result.detail


def test_precondition_skips_when_expected_text_empty(tmp_path: Path):
    paragraph = _live_paragraph("anything", tmp_path)
    # Stale offsets and no expected text — must not silently accept.
    result = _precondition_holds_for_paragraph(paragraph, 99, 102, "")
    assert result.ok is False


# ---------------------------------------------------------------------------
# End-to-end safety tests through apply_edits_to_spec
# ---------------------------------------------------------------------------


def test_replace_uses_corrected_offsets_when_recorded_offsets_stale(tmp_path: Path):
    """Stale recorded offsets + uniquely-present target → replace at the corrected span.

    Without the Chunk F fix, the precondition fallback would accept the
    edit (because the unique substring still exists) and then
    `_replace_in_paragraph` would happily overwrite whatever now sits at
    [match_start, match_end) — which is no longer the target text. The
    fix returns corrected offsets and the replacement lands on the
    correct span.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    live_text = "Lead-in clause inserted later. Replace TARGET token here."
    doc = Document()
    doc.add_paragraph(live_text)
    doc.save(source)

    # Stale offsets: the locator believed TARGET was at the start of the
    # paragraph (e.g. extracted from an older version that didn't have the
    # lead-in clause). The current text has TARGET much further to the right.
    real_start = live_text.index("TARGET")
    stale_start = 8
    stale_end = stale_start + len("TARGET")
    assert live_text[stale_start:stale_end] != "TARGET"
    assert real_start != stale_start

    result = _locator_result(
        text=live_text,
        match_start=stale_start,
        match_end=stale_end,
        matched_text="TARGET",
        replacement_text="REPLACED",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    final = saved.paragraphs[0].text

    # Replacement landed on the live TARGET, not on whatever sat at the
    # stale offsets. Lead-in clause survives unmodified.
    assert "TARGET" not in final
    assert final == "Lead-in clause inserted later. Replace REPLACED token here."
    assert report.edits_applied == 1


def test_replace_skipped_when_target_text_is_duplicated_in_live_paragraph(tmp_path: Path):
    """Stale offsets + duplicated target → skip; never guess between occurrences.

    Demonstrates the core safety guarantee: if the recorded slice doesn't
    match and the expected text appears more than once in the live
    paragraph, the editor refuses to apply the edit. Picking either
    occurrence at random would risk corrupting the wrong span.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    live_text = "TWIN here and TWIN there."
    doc = Document()
    doc.add_paragraph(live_text)
    doc.save(source)

    # Stale offsets that don't slice "TWIN" — typical of a locator computed
    # against an older version of the paragraph or a different normalization.
    stale_start = 5
    stale_end = stale_start + len("TWIN")
    assert live_text[stale_start:stale_end] != "TWIN"

    result = _locator_result(
        text=live_text,
        match_start=stale_start,
        match_end=stale_end,
        matched_text="TWIN",
        replacement_text="SOLE",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    # Paragraph is unchanged; nothing was guessed at.
    assert saved.paragraphs[0].text == live_text
    assert report.edits_applied == 0
    skipped = [o for o in report.outcomes if o.status == "skipped"]
    assert len(skipped) == 1
    assert "appears 2 times" in skipped[0].detail
    assert "manual review" in skipped[0].detail.lower()


def test_replace_skipped_when_target_text_missing_from_live_paragraph(tmp_path: Path):
    """Stale offsets + missing target → skip with a clear diagnostic."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    live_text = "Real paragraph with no ghost text inside."
    doc = Document()
    doc.add_paragraph(live_text)
    doc.save(source)

    result = _locator_result(
        text=live_text,
        match_start=0,
        match_end=len("ghost phrase"),
        matched_text="ghost phrase",
        replacement_text="REPLACED",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)

    assert saved.paragraphs[0].text == live_text
    assert report.edits_applied == 0
    skipped = [o for o in report.outcomes if o.status == "skipped"]
    assert len(skipped) == 1
    assert "no longer present" in skipped[0].detail


def test_sequential_edits_with_offset_safety_in_one_paragraph(tmp_path: Path):
    """Two non-conflicting edits in one paragraph still both apply.

    Guards the realistic multi-edit case: descending-start ordering means
    the higher-offset edit runs first and the lower-offset edit's recorded
    offsets are still valid when it runs. Even if `LONGREPLACEMENT` shifts
    text after position 0, the second edit at position 0 is unaffected.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    text = "Replace foo and then bar with baz."
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)

    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[0]

    foo_start = text.index("foo")
    bar_start = text.index("bar")

    edit_foo = LocatorResult(
        finding=_finding(existing="foo", replacement="LONGREPLACEMENT"),
        status="matched",
        locations=[
            EditLocation(
                mapping=mapping,
                match_start=foo_start,
                match_end=foo_start + 3,
                matched_text="foo",
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="LONGREPLACEMENT",
        action_type="EDIT",
        warning=None,
    )
    edit_bar = LocatorResult(
        finding=_finding(existing="bar", replacement="QUX"),
        status="matched",
        locations=[
            EditLocation(
                mapping=mapping,
                match_start=bar_start,
                match_end=bar_start + 3,
                matched_text="bar",
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="QUX",
        action_type="EDIT",
        warning=None,
    )

    report = apply_edits_to_spec(
        source, output, build_edit_actions([edit_foo, edit_bar])
    )
    saved = Document(output)
    assert saved.paragraphs[0].text == "Replace LONGREPLACEMENT and then QUX with baz."
    assert report.edits_applied == 2
    assert report.edits_failed == 0


def test_table_cell_edit_uses_corrected_offsets(tmp_path: Path):
    """The table-cell path must apply the same offset-correction contract."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Pad cell"
    table.cell(0, 1).text = "GHOST text and TARGET phrase"
    doc.save(source)

    # Pretend the locator recorded the row's joined text. We deliberately
    # supply STALE recorded offsets for "TARGET" (off by a few characters)
    # so the precondition falls back to the unique-substring branch and
    # corrects them. Without Chunk F, the replacement would land on
    # whatever sits at the stale offsets.
    joined = "Pad cell | GHOST text and TARGET phrase"
    real_start = joined.index("TARGET")
    real_end = real_start + len("TARGET")

    # Use stale offsets that are inside the cell text but pointing at
    # the wrong span ("OST t" instead of "TARGET").
    stale_start = real_start - 10
    stale_end = stale_start + len("TARGET")
    assert joined[stale_start:stale_end] != "TARGET"

    result = _locator_result(
        text=joined,
        element_type="table_cell",
        row_index=0,
        match_start=stale_start,
        match_end=stale_end,
        matched_text="TARGET",
        replacement_text="REPLACED",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    cell_text = saved.tables[0].cell(0, 1).text
    # The replacement landed on the correct "TARGET" span, not the stale
    # slice at [stale_start, stale_end). "GHOST text" is still intact.
    assert "TARGET" not in cell_text
    assert "REPLACED" in cell_text
    assert "GHOST text and" in cell_text
    assert report.edits_applied == 1


def test_table_cell_edit_skipped_when_target_duplicated(tmp_path: Path):
    """Duplicated target inside a cell with stale offsets → skip, no guessing."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "TWIN here and TWIN there"
    doc.save(source)

    cell_text = "TWIN here and TWIN there"
    # Stale offsets inside the cell that don't slice "TWIN".
    stale_start = 5
    stale_end = stale_start + len("TWIN")
    assert cell_text[stale_start:stale_end] != "TWIN"

    result = _locator_result(
        text=cell_text,
        element_type="table_cell",
        row_index=0,
        match_start=stale_start,
        match_end=stale_end,
        matched_text="TWIN",
        replacement_text="SOLE",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    final_cell = saved.tables[0].cell(0, 0).text
    # Duplicated target + stale offsets → skip; the cell text is unchanged.
    assert final_cell == "TWIN here and TWIN there"
    assert report.edits_applied == 0
    skipped = [o for o in report.outcomes if o.status == "skipped"]
    assert len(skipped) == 1
    assert "appears 2 times" in skipped[0].detail


def test_existing_safe_replacement_still_applies(tmp_path: Path):
    """Sanity: a fresh, unambiguous edit with valid offsets still works.

    Guards against the Chunk F refactor accidentally tightening the
    happy-path behavior.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide seismic bracing per ASCE 7-16.")
    doc.save(source)

    text = "Provide seismic bracing per ASCE 7-16."
    start = text.index("ASCE 7-16")
    end = start + len("ASCE 7-16")

    result = _locator_result(
        text=text,
        match_start=start,
        match_end=end,
        matched_text="ASCE 7-16",
        replacement_text="ASCE 7-22",
    )

    report = apply_edits_to_spec(source, output, build_edit_actions([result]))
    saved = Document(output)
    assert saved.paragraphs[0].text == "Provide seismic bracing per ASCE 7-22."
    assert report.edits_applied == 1
