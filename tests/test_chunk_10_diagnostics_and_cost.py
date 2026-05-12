"""Chunk 10 — bounded diagnostics + cost-estimator regression tests.

Covers the four acceptance criteria the plan asks for:

1. Diagnostics truncate oversized payloads (byte caps + global cap).
2. Secret-looking values are not emitted (key + value scrubbing).
3. Estimated costs compute correctly for representative model / phase /
   token cases (Opus / Sonnet / Haiku, batch discount, cache rates,
   web-search add-on, unknown-model fallback).
4. Report renders the cost summary when available and gracefully omits
   it when pricing is unavailable.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from src.api_config import (
    MODEL_HAIKU_45,
    MODEL_OPUS_47,
    MODEL_SONNET_46,
)
from src.cost_estimator import (
    BATCH_DISCOUNT,
    PRICING_AS_OF,
    WEB_SEARCH_PRICE_PER_1K,
    estimate_event_cost,
    estimate_run_cost,
    format_usd,
    model_pricing,
)
from src.diagnostics import (
    _DEFAULT_MAX_EVENT_DATA_BYTES,
    DiagnosticsReport,
    _scrub_and_bound,
)


# ===========================================================================
# 1) Diagnostics byte caps
# ===========================================================================


class TestEventByteCaps:
    def test_oversized_event_data_is_truncated_in_place(self):
        report = DiagnosticsReport(max_event_data_bytes=512)
        huge_text = "x" * 5000  # ~5 KB
        report.log("review", "info", "big payload", {"raw_response": huge_text})

        event = report.events[0]
        # The huge string was truncated to a smaller form.
        assert "raw_response" in (event.data or {})
        # Either the string itself was capped, or the field was elided.
        rendered = event.data["raw_response"]
        assert isinstance(rendered, str)
        assert len(rendered.encode("utf-8")) <= 5000  # never grew
        # Counter incremented.
        assert report.events_truncated_by_size >= 1

    def test_global_cap_drops_oldest_events_when_exceeded(self):
        report = DiagnosticsReport(
            max_events=1000,
            max_event_data_bytes=64 * 1024,
            max_total_data_bytes=4 * 1024,  # 4 KB total
        )
        # Each event carries ~1 KB of data — after 5 events we exceed 4 KB.
        for i in range(8):
            report.log(
                "review",
                "info",
                f"event {i}",
                {"payload": "p" * 1000, "i": i},
            )
        assert report.total_data_bytes <= 4 * 1024
        assert report.events_dropped >= 1
        assert report.bytes_dropped > 0
        # The newest event is still present.
        assert report.events[-1].message == "event 7"

    def test_default_caps_are_applied(self):
        """The defaults exist and look sane."""
        report = DiagnosticsReport()
        assert report.max_event_data_bytes == _DEFAULT_MAX_EVENT_DATA_BYTES
        assert report.max_total_data_bytes > 0

    def test_truncation_visible_in_summary(self):
        report = DiagnosticsReport(max_event_data_bytes=256)
        report.log("review", "info", "x", {"raw_response": "a" * 5000})
        report.finish()
        summary = report.summary()
        assert summary["events_truncated_by_size"] >= 1


# ===========================================================================
# 2) Secret scrubbing
# ===========================================================================


class TestSecretScrubbing:
    def test_secret_shaped_keys_are_redacted(self):
        scrubbed = _scrub_and_bound(
            {
                "api_key": "sk-ant-PROD-abcdef1234567890",
                "AUTHORIZATION": "Bearer abcd1234",
                "password": "hunter2",
                "client_secret": "topsecret",
                "ok_field": "safe value",
            }
        )
        assert scrubbed["api_key"] == "<redacted>"
        assert scrubbed["AUTHORIZATION"] == "<redacted>"
        assert scrubbed["password"] == "<redacted>"
        assert scrubbed["client_secret"] == "<redacted>"
        assert scrubbed["ok_field"] == "safe value"

    def test_secret_shaped_values_are_redacted_even_under_safe_keys(self):
        scrubbed = _scrub_and_bound(
            {
                "message": "Found sk-ant-PROD-abcdef1234567890 in payload",
                "akia": "AKIA0123456789ABCDEF",
                "auth": "Bearer aaaaaaaaaaaaaaaaaaaa",
            }
        )
        # Either the secret-key match or the secret-value match redacts.
        for key, val in scrubbed.items():
            assert "sk-ant-" not in str(val)
            assert "AKIA0123456789ABCDEF" not in str(val)
            assert "Bearer aaaaaaaaaaaa" not in str(val)

    def test_nested_secrets_are_scrubbed(self):
        scrubbed = _scrub_and_bound(
            {
                "config": {"api_key": "sk-ant-abc12345"},
                "list_of_things": [
                    {"password": "p"},
                    "sk-ant-DEEPdeepdeep12345",
                ],
            }
        )
        assert scrubbed["config"]["api_key"] == "<redacted>"
        assert scrubbed["list_of_things"][0]["password"] == "<redacted>"
        assert scrubbed["list_of_things"][1] == "<redacted>"

    def test_record_api_call_does_not_emit_secrets(self):
        report = DiagnosticsReport()
        report.record_api_call(
            phase="review",
            model="opus",
            input_tokens=10,
            output_tokens=10,
            extra={
                "api_key": "sk-ant-LIVE-1234567890abcdef",
                "user_query": "Hello",
            },
        )
        event = report.events[0]
        # The api_key field is scrubbed.
        assert event.data["api_key"] == "<redacted>"
        # The safe field round-trips.
        assert event.data["user_query"] == "Hello"
        # Counter incremented at least once.
        assert report.secrets_redacted >= 1

    def test_full_report_json_contains_no_secret_values(self):
        report = DiagnosticsReport()
        report.log(
            "review",
            "info",
            "leaky log",
            {"context": "trace contains sk-ant-PROD-aaaaaaaaaaaaaaaa here"},
        )
        report.finish()
        rendered = json.dumps(report.to_dict(), default=str)
        assert "sk-ant-PROD-aaaaaaaaaaaaaaaa" not in rendered


# ===========================================================================
# 3) Cost estimator
# ===========================================================================


class TestModelPricing:
    def test_known_models_have_pricing(self):
        assert model_pricing(MODEL_OPUS_47) is not None
        assert model_pricing(MODEL_SONNET_46) is not None
        assert model_pricing(MODEL_HAIKU_45) is not None

    def test_unknown_model_returns_none(self):
        assert model_pricing("claude-opus-9000") is None
        assert model_pricing("") is None
        assert model_pricing(None) is None  # type: ignore[arg-type]


class TestPerEventCost:
    def test_opus_realtime_basic_input_output(self):
        """1M input + 100K output on Opus realtime = $15 + $7.50 = $22.50."""
        cost = estimate_event_cost(
            {
                "model": MODEL_OPUS_47,
                "input_tokens": 1_000_000,
                "output_tokens": 100_000,
                "call_mode": "realtime",
            }
        )
        assert cost is not None
        assert cost["input_usd"] == 15.0
        assert cost["output_usd"] == 7.5
        assert cost["total_usd"] == 22.5

    def test_batch_discount_halves_input_output(self):
        cost = estimate_event_cost(
            {
                "model": MODEL_SONNET_46,
                "input_tokens": 1_000_000,
                "output_tokens": 100_000,
                "call_mode": "batch",
            }
        )
        assert cost is not None
        # Sonnet: $3 input * 0.5 = $1.50; $15 output * 0.1 * 0.5 = $0.75
        assert cost["input_usd"] == pytest.approx(1.5)
        assert cost["output_usd"] == pytest.approx(0.75)
        assert cost["total_usd"] == pytest.approx(2.25)

    def test_cache_writes_and_reads_priced_separately_from_input(self):
        cost = estimate_event_cost(
            {
                "model": MODEL_SONNET_46,
                "input_tokens": 0,
                "output_tokens": 0,
                "cache_creation_input_tokens": 1_000_000,
                "cache_read_input_tokens": 1_000_000,
            }
        )
        assert cost is not None
        # Sonnet cache write 1h: $6/MTok, read: $0.30/MTok
        assert cost["cache_write_usd"] == 6.0
        assert cost["cache_read_usd"] == 0.3
        assert cost["total_usd"] == pytest.approx(6.3)

    def test_cache_pricing_is_not_discounted_on_batch(self):
        cost = estimate_event_cost(
            {
                "model": MODEL_SONNET_46,
                "cache_creation_input_tokens": 1_000_000,
                "cache_read_input_tokens": 1_000_000,
                "call_mode": "batch",
            }
        )
        assert cost is not None
        # Cache pricing identical regardless of mode.
        assert cost["cache_write_usd"] == 6.0
        assert cost["cache_read_usd"] == 0.3

    def test_web_search_priced_per_thousand(self):
        cost = estimate_event_cost(
            {
                "model": MODEL_HAIKU_45,
                "input_tokens": 0,
                "output_tokens": 0,
                "web_search_requests": 100,
            }
        )
        assert cost is not None
        # $10 per 1k → 100 requests = $1
        assert cost["web_search_usd"] == pytest.approx(1.0)
        assert cost["total_usd"] == pytest.approx(1.0)

    def test_unknown_model_returns_none(self):
        cost = estimate_event_cost(
            {
                "model": "claude-future-9000",
                "input_tokens": 1_000_000,
                "output_tokens": 1_000_000,
            }
        )
        assert cost is None

    def test_empty_data_returns_none(self):
        assert estimate_event_cost({}) is None
        assert estimate_event_cost(None) is None  # type: ignore[arg-type]


class TestRunCostAggregation:
    def _record(self, report, phase, model, input_t, output_t, **kw):
        report.record_api_call(
            phase=phase,
            model=model,
            input_tokens=input_t,
            output_tokens=output_t,
            **kw,
        )

    def test_aggregates_across_phases_and_models(self):
        report = DiagnosticsReport()
        self._record(report, "review", MODEL_OPUS_47, 1_000_000, 100_000)
        self._record(report, "verification", MODEL_SONNET_46, 200_000, 20_000)
        self._record(report, "triage", MODEL_HAIKU_45, 50_000, 5_000)

        ec = estimate_run_cost(report.events)
        assert ec["available"] is True
        assert set(ec["by_phase"].keys()) == {"review", "verification", "triage"}
        # By-model rollup has all three.
        assert set(ec["by_model"].keys()) == {
            MODEL_OPUS_47,
            MODEL_SONNET_46,
            MODEL_HAIKU_45,
        }
        # Pricing snapshot is included.
        assert ec["pricing_as_of"] == PRICING_AS_OF
        # Disclaimer notes are non-empty.
        assert any("Estimated" in n or "Anthropic" in n for n in ec["notes"])
        # Currency carried through.
        assert ec["currency"] == "USD"

    def test_unknown_model_routes_to_missing_pricing_calls(self):
        report = DiagnosticsReport()
        self._record(report, "review", MODEL_OPUS_47, 100_000, 10_000)
        self._record(report, "review", "claude-mystery", 100_000, 10_000)

        ec = estimate_run_cost(report.events)
        assert ec["available"] is True  # at least one priced call
        assert ec["missing_pricing_calls"] == 1
        assert ec["missing_pricing_models"] == ["claude-mystery"]
        # Note explicitly cites the unknown model.
        assert any("claude-mystery" in n for n in ec["notes"])
        # The phase bucket records the missing call too.
        assert ec["by_phase"]["review"]["missing_pricing_calls"] == 1

    def test_no_api_calls_reports_unavailable(self):
        report = DiagnosticsReport()
        report.log("review", "info", "just a status message")
        ec = estimate_run_cost(report.events)
        assert ec["available"] is False
        assert ec["total_usd"] == 0.0

    def test_every_call_unknown_model_reports_unavailable(self):
        report = DiagnosticsReport()
        self._record(report, "review", "claude-future", 100, 100)
        ec = estimate_run_cost(report.events)
        assert ec["available"] is False
        assert ec["missing_pricing_calls"] == 1
        assert "claude-future" in ec["missing_pricing_models"]


class TestFormatUSD:
    def test_dollars_use_two_decimals(self):
        assert format_usd(12.345) == "$12.35"
        assert format_usd(1.0) == "$1.00"
        assert format_usd(1234.5) == "$1,234.50"

    def test_cents_show_three_decimals(self):
        # 0.01 .. 0.999
        assert format_usd(0.123) == "$0.123"
        assert format_usd(0.01) == "$0.010"

    def test_sub_penny_uses_four_decimals(self):
        assert format_usd(0.001234) == "$0.0012"

    def test_zero_is_clean(self):
        assert format_usd(0.0) == "$0.00"

    def test_negative_clamps_to_zero(self):
        assert format_usd(-5.0) == "$0.00"


# ===========================================================================
# 4) Diagnostics summary integration + to_text rendering
# ===========================================================================


class TestSummaryIntegration:
    def test_summary_includes_estimated_cost(self):
        report = DiagnosticsReport()
        report.record_api_call(
            phase="review",
            model=MODEL_OPUS_47,
            input_tokens=1_000_000,
            output_tokens=100_000,
        )
        s = report.summary()
        assert "estimated_cost" in s
        ec = s["estimated_cost"]
        assert ec["available"] is True
        assert ec["total_usd"] == 22.5

    def test_summary_marks_unavailable_when_no_api_calls(self):
        report = DiagnosticsReport()
        report.log("review", "info", "just text")
        s = report.summary()
        ec = s["estimated_cost"]
        assert ec["available"] is False

    def test_to_text_renders_cost_section_when_available(self):
        report = DiagnosticsReport()
        report.record_api_call(
            phase="review",
            model=MODEL_OPUS_47,
            input_tokens=1_000_000,
            output_tokens=100_000,
        )
        report.finish()
        text = report.to_text()
        assert "ESTIMATED API COST" in text
        assert "$22.50" in text
        assert "Pricing As Of:" in text
        assert PRICING_AS_OF in text

    def test_to_text_says_unavailable_when_no_pricing(self):
        report = DiagnosticsReport()
        report.record_api_call(
            phase="review",
            model="claude-future-9000",
            input_tokens=100,
            output_tokens=100,
        )
        report.finish()
        text = report.to_text()
        assert "ESTIMATED API COST" in text
        assert "Cost unavailable" in text

    def test_summary_records_truncation_and_redaction_counters(self):
        report = DiagnosticsReport(max_event_data_bytes=256)
        report.record_api_call(
            phase="review",
            model="opus",
            input_tokens=10,
            output_tokens=10,
            extra={
                "api_key": "sk-ant-LIVE-aaaaaaaaaaaaaaaa",
                "raw_response": "x" * 5000,
            },
        )
        s = report.summary()
        assert s["secrets_redacted"] >= 1
        assert s["events_truncated_by_size"] >= 1


# ===========================================================================
# 5) Word report integration
# ===========================================================================


class TestReportIntegration:
    """Make sure the Word exporter renders the cost section without crashing.

    Uses a stub PipelineResult so the test stays fully offline.
    """

    def _stub_result(self):
        from src.reviewer import ReviewResult

        class _StubResult:
            review_result = ReviewResult(findings=[], model=MODEL_OPUS_47)
            cross_check_result = None
            files_reviewed = ["test.docx"]
            leed_alerts: list = []
            placeholder_alerts: list = []
            total_elapsed_seconds = 1.0
            cycle_label = "2025"

        return _StubResult()

    def test_export_with_cost_summary_renders(self, tmp_path: Path):
        from src.report_exporter import export_report

        ec = {
            "available": True,
            "total_usd": 1.2345,
            "currency": "USD",
            "pricing_as_of": PRICING_AS_OF,
            "by_phase": {
                "review": {
                    "calls": 1,
                    "input_usd": 0.5,
                    "output_usd": 0.5,
                    "cache_write_usd": 0.1,
                    "cache_read_usd": 0.1,
                    "web_search_usd": 0.0,
                    "total_usd": 1.2,
                    "missing_pricing_calls": 0,
                }
            },
            "by_model": {
                MODEL_OPUS_47: {
                    "calls": 1,
                    "total_usd": 1.2,
                    "input_tokens": 100,
                    "output_tokens": 100,
                    "cache_creation_input_tokens": 0,
                    "cache_read_input_tokens": 0,
                    "web_search_requests": 0,
                }
            },
            "missing_pricing_models": [],
            "missing_pricing_calls": 0,
            "priced_calls": 1,
            "web_search_requests": 0,
            "notes": ["Estimated only."],
        }
        out = tmp_path / "report.docx"
        export_report(self._stub_result(), out, estimated_cost=ec)
        assert out.exists()

        # Re-open to confirm "Estimated API Cost" heading is present.
        from docx import Document

        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Estimated API Cost" in all_text
        # Dollar figure rendered.
        assert "$1.2" in all_text  # tolerant of formatting

    def test_export_without_cost_summary_still_works(self, tmp_path: Path):
        from src.report_exporter import export_report

        out = tmp_path / "report.docx"
        export_report(self._stub_result(), out, estimated_cost=None)
        assert out.exists()

        from docx import Document

        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Estimated API Cost" not in all_text

    def test_export_with_unavailable_cost_renders_disclaimer(self, tmp_path: Path):
        from src.report_exporter import export_report

        ec = {
            "available": False,
            "total_usd": 0.0,
            "currency": "USD",
            "pricing_as_of": PRICING_AS_OF,
            "by_phase": {},
            "by_model": {},
            "missing_pricing_models": ["claude-future-9000"],
            "missing_pricing_calls": 1,
            "priced_calls": 0,
            "web_search_requests": 0,
            "notes": [],
        }
        out = tmp_path / "report.docx"
        export_report(self._stub_result(), out, estimated_cost=ec)

        from docx import Document

        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Estimated API Cost" in all_text
        assert "unavailable" in all_text.lower()
        assert "claude-future-9000" in all_text


# ===========================================================================
# 6) Constants sanity
# ===========================================================================


class TestPricingTableSanity:
    def test_batch_discount_is_50_percent(self):
        assert BATCH_DISCOUNT == 0.5

    def test_web_search_price_matches_documented_rate(self):
        assert WEB_SEARCH_PRICE_PER_1K == 10.0

    def test_cache_pricing_ratios_match_anthropic_multipliers(self):
        """1h cache write = 2x input; cache read = 0.1x input."""
        for model in (MODEL_OPUS_47, MODEL_SONNET_46, MODEL_HAIKU_45):
            p = model_pricing(model)
            assert p is not None
            assert p.cache_write_1h_rate == pytest.approx(p.input_rate * 2.0)
            assert p.cache_read_rate == pytest.approx(p.input_rate * 0.1)

    def test_opus_more_expensive_than_sonnet_than_haiku(self):
        opus = model_pricing(MODEL_OPUS_47)
        sonnet = model_pricing(MODEL_SONNET_46)
        haiku = model_pricing(MODEL_HAIKU_45)
        assert opus and sonnet and haiku
        assert opus.input_rate > sonnet.input_rate > haiku.input_rate
        assert opus.output_rate > sonnet.output_rate > haiku.output_rate
