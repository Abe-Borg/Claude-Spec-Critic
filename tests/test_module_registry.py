"""Module registry, module-identity threading, and prompt-slot contracts.

Phase 1 introduced the registry + identity threading; Phase 2 moved the
prompt content (personas, severity anchors, categories, few-shot examples,
verifier source tiers) onto the module as validated slots. Covers:

1. **Registry** — ``ReviewModule`` is frozen; ``get_module`` resolves known
   ids and degrades unknown / missing ids to the default California module
   (mirroring ``AVAILABLE_CYCLES.get(label, DEFAULT_CYCLE)``); registry
   validation fails fast on duplicate ids, duplicate cycle labels (the
   verification-cache namespace rule), and empty fields.
2. **Persistence round-trip** — ``module_id`` rides
   ``BatchSubmission -> PendingBatch -> pending_batch.json -> to_submission``
   verbatim, and a LEGACY state file written before the field existed loads
   with the default module (no schema bump).
3. **Downstream stamping** — ``finalize_batch_result`` carries the
   submission's module id onto ``PipelineResult``; the trace recorder writes
   it into ``run.json``.
4. **The ``module_for_cycle`` bridge** — content layers still keyed by
   ``cycle=`` resolve their module through the registry's unique-label
   reverse lookup.
5. **Prompt-slot content validation** — registration rejects empty slots,
   templates that don't format, few-shot examples that the real parser
   would demote, and per-request element-id leakage into the cached prefix.
6. **Slots drive output** — a synthetic module's strings actually appear in
   the assembled reviewer / cross-check / verifier prompts (the byte-exact
   CA rendering is separately pinned by ``test_golden_domain_surfaces``).
"""
from __future__ import annotations

import dataclasses
import json

import pytest

from src.batch.batch import BatchJob
from src.core.code_cycles import BaseCode, CALIFORNIA_2025, CodeCycle
from src.modules import (
    AVAILABLE_MODULES,
    CALIFORNIA_K12_MEP,
    ChunkGroup,
    DEFAULT_MODULE,
    DetectorVocabulary,
    ProfileKeywords,
    ReviewModule,
    code_basis_format_kwargs,
    get_module,
    module_for_cycle,
    validate_module_registry,
)
from src.modules.base import _iter_json_objects
from src.orchestration.batch_resume import (
    PendingBatch,
    load_pending_batch,
    save_pending_batch,
)
from src.orchestration.pipeline import (
    BatchSubmission,
    CollectedBatchState,
    PipelineResult,
    finalize_batch_result,
)
from src.review.reviewer import ReviewResult
from src.tracing.recorder import TraceRecorder


def _cycle(label: str) -> CodeCycle:
    return CodeCycle(
        label=label,
        base_codes=(BaseCode("code", "Test Code", label),) if label else (),
        asce7="7-22",
        asce7_previous="7-16",
    )


_TEST_VOCABULARY = DetectorVocabulary(
    code_abbreviations=("TBC",),
    plausible_cycle_years=("2019", "2022"),
    valid_cycle_years=("2019", "2022", "2025"),
    asce7_plausible_editions=("16", "22"),
    jurisdiction_label="Test",
)

_TEST_PROFILE_KEYWORDS = ProfileKeywords(
    jurisdictional=("testville", "test marshal"),
    manufacturer=("testcorp", "model number"),
    code_standard=("tbc", "test code"),
    internal_coordination=("placeholder", "typo", "internal contradiction"),
)

_TEST_CHUNK_GROUPS = (
    ChunkGroup("div_21", "Division 21 — Fire Suppression", ("21",)),
    ChunkGroup("div_22", "Division 22 — Plumbing", ("22",)),
)


_VALID_EXAMPLE_BLOCK = """\
Example 1 — REPORT_ONLY:
{
  "severity": "MEDIUM",
  "fileName": "a.docx",
  "section": "1.01",
  "issue": "Test issue.",
  "actionType": "REPORT_ONLY",
  "existingText": null,
  "replacementText": null,
  "codeReference": null,
  "confidence": 0.5
}
"""


def _module(module_id: str = "test_module", label: str = "9999", **overrides) -> ReviewModule:
    slots = dict(
        module_id=module_id,
        display_name="Test Module",
        description="test",
        cycle=_cycle(label),
        reviewer_persona="You are a test reviewer.",
        review_user_intro="Review the following test document.",
        review_severity_definitions=(
            "CRITICAL — test.\nHIGH — test.\nMEDIUM — test.\nGRIPES — test."
        ),
        review_confidence_high_example="a test high-confidence example",
        review_categories_template="1. Internal contradictions within the spec.",
        review_examples=_VALID_EXAMPLE_BLOCK,
        cross_check_persona="You are a test coordination reviewer.",
        cross_check_severity_definitions=(
            "CRITICAL — test.\nHIGH — test.\nMEDIUM — test.\nGRIPES — test."
        ),
        verifier_persona="You are a test verification assistant.",
        verifier_source_priorities="1. Test sources:\n   example.gov",
        review_user_code_basis_line="Current code cycle: Test Code {code}.",
        cross_check_code_basis_line="Current cycle: Test Code {code}.",
        verifier_system_code_basis_lines="Current code cycle: Test Code {code}.",
        verifier_user_code_basis_lines="Current cycle: Test Code {code}.",
        detector_vocabulary=_TEST_VOCABULARY,
        profile_keywords=_TEST_PROFILE_KEYWORDS,
        cross_check_chunk_groups=_TEST_CHUNK_GROUPS,
        report_context_phrase="test projects",
        report_title="Spec Critic — Test Specification Review Report",
    )
    slots.update(overrides)
    return ReviewModule(**slots)


class _StubReview:
    """Minimal duck-typed ReviewResult for the title-block writer."""

    model = "test-model"


def _submission(**overrides) -> BatchSubmission:
    job = BatchJob(
        batch_id="msgbatch_MODTEST",
        job_type="review",
        request_map={"review__a__0": {"filename": "a.docx", "index": 0, "type": "review"}},
        created_at=1700000000.0,
    )
    base = dict(
        job=job,
        files_reviewed=["a.docx"],
        review_request_ids=["review__a__0"],
        model="claude-opus-4-8",
    )
    base.update(overrides)
    return BatchSubmission(**base)


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


class TestRegistry:
    def test_default_module_is_california_k12_mep(self):
        assert DEFAULT_MODULE is CALIFORNIA_K12_MEP
        assert DEFAULT_MODULE.module_id == "california_k12_mep"
        assert DEFAULT_MODULE.cycle is CALIFORNIA_2025
        assert AVAILABLE_MODULES == {"california_k12_mep": CALIFORNIA_K12_MEP}

    def test_get_module_resolves_known_id(self):
        assert get_module("california_k12_mep") is CALIFORNIA_K12_MEP

    @pytest.mark.parametrize("bad_id", [None, "", "  ", "not_a_module"])
    def test_get_module_degrades_to_default(self, bad_id):
        assert get_module(bad_id) is DEFAULT_MODULE

    def test_get_module_strips_whitespace(self):
        assert get_module("  california_k12_mep  ") is CALIFORNIA_K12_MEP

    def test_review_module_is_frozen(self):
        with pytest.raises(dataclasses.FrozenInstanceError):
            CALIFORNIA_K12_MEP.module_id = "other"  # type: ignore[misc]


class TestRegistryValidation:
    def test_current_registry_validates(self):
        validate_module_registry(AVAILABLE_MODULES.values())

    def test_duplicate_module_id_rejected(self):
        with pytest.raises(ValueError, match="Duplicate module_id"):
            validate_module_registry(
                [_module("dup", "9998"), _module("dup", "9999")]
            )

    def test_duplicate_cycle_label_rejected(self):
        # Cycle labels namespace the verification cache AND back the
        # module_for_cycle bridge; two modules sharing one label would let
        # cached verdicts collide and make the bridge ambiguous.
        with pytest.raises(ValueError, match="Duplicate cycle label"):
            validate_module_registry(
                [_module("first", "9999"), _module("second", "9999")]
            )

    def test_empty_module_id_rejected(self):
        with pytest.raises(ValueError, match="module_id"):
            validate_module_registry([_module("", "9999")])

    def test_unstripped_module_id_rejected(self):
        with pytest.raises(ValueError, match="module_id"):
            validate_module_registry([_module(" padded ", "9999")])

    def test_empty_display_name_rejected(self):
        bad = dataclasses.replace(_module(), display_name="  ")
        with pytest.raises(ValueError, match="display_name"):
            validate_module_registry([bad])

    def test_missing_cycle_label_rejected(self):
        with pytest.raises(ValueError, match="cycle label"):
            validate_module_registry([_module("ok", "")])


# ---------------------------------------------------------------------------
# module_for_cycle bridge
# ---------------------------------------------------------------------------


class TestModuleForCycleBridge:
    def test_resolves_california_cycle(self):
        assert module_for_cycle(CALIFORNIA_2025) is CALIFORNIA_K12_MEP

    def test_unknown_label_degrades_to_default(self):
        assert module_for_cycle(_cycle("0000")) is DEFAULT_MODULE

    def test_label_less_input_degrades_to_default(self):
        assert module_for_cycle(None) is DEFAULT_MODULE
        assert module_for_cycle(object()) is DEFAULT_MODULE


# ---------------------------------------------------------------------------
# Prompt-slot content validation
# ---------------------------------------------------------------------------


class TestContentValidation:
    def test_empty_prompt_slot_rejected(self):
        bad = dataclasses.replace(_module(), reviewer_persona="   ")
        with pytest.raises(ValueError, match="reviewer_persona"):
            validate_module_registry([bad])

    def test_unknown_categories_placeholder_rejected(self):
        bad = dataclasses.replace(
            _module(), review_categories_template="1. Cycle {not_a_placeholder}."
        )
        with pytest.raises(ValueError, match="review_categories_template"):
            validate_module_registry([bad])

    def test_examples_without_json_rejected(self):
        bad = dataclasses.replace(_module(), review_examples="Prose only, no JSON.")
        with pytest.raises(ValueError, match="no parseable JSON"):
            validate_module_registry([bad])

    @pytest.mark.parametrize(
        "leak, match",
        [
            ("Also set evidenceElementId.", "evidenceElementId"),
            ('Quote the <para id="p7"> element.', "<para"),
            ('Cite the <row id="t0r0"> entry.', "<row"),
            ('Reference the <heading id="p0"> line.', "<heading"),
        ],
    )
    def test_example_element_id_leak_rejected(self, leak, match):
        # Per-request concepts inside the cached system-prompt prefix: the
        # evidenceElementId field and EVERY element-id wrapper tag the
        # id-tagged rendering emits (<para>, <row>, <heading>).
        bad = dataclasses.replace(
            _module(),
            review_examples=_VALID_EXAMPLE_BLOCK + "\n" + leak,
        )
        with pytest.raises(ValueError, match=match):
            validate_module_registry([bad])

    def test_example_teaching_demotable_edit_rejected(self):
        # An EDIT without replacementText is demoted by the parser — an
        # example teaching that shape would waste model output every run.
        block = """\
Example — bad EDIT:
{
  "severity": "MEDIUM",
  "fileName": "a.docx",
  "section": "1.01",
  "issue": "Bad shape.",
  "actionType": "EDIT",
  "existingText": "old text",
  "replacementText": null,
  "codeReference": null,
  "confidence": 0.5
}
"""
        bad = dataclasses.replace(_module(), review_examples=block)
        with pytest.raises(ValueError, match="demoted"):
            validate_module_registry([bad])

    def test_example_noop_edit_rejected(self):
        # existingText byte-identical to replacementText is the no-op EDIT
        # shape the parser demotes (TRUST_AUDIT P1-1).
        block = """\
Example — no-op EDIT:
{
  "severity": "MEDIUM",
  "fileName": "a.docx",
  "section": "1.01",
  "issue": "No-op.",
  "actionType": "EDIT",
  "existingText": "same text",
  "replacementText": "same text",
  "codeReference": null,
  "confidence": 0.5
}
"""
        bad = dataclasses.replace(_module(), review_examples=block)
        with pytest.raises(ValueError, match="demoted"):
            validate_module_registry([bad])

    def test_example_bad_severity_rejected(self):
        block = _VALID_EXAMPLE_BLOCK.replace('"MEDIUM"', '"BLOCKER"')
        bad = dataclasses.replace(_module(), review_examples=block)
        with pytest.raises(ValueError, match="severity"):
            validate_module_registry([bad])

    def test_example_confidence_out_of_range_rejected(self):
        block = _VALID_EXAMPLE_BLOCK.replace('"confidence": 0.5', '"confidence": 1.5')
        bad = dataclasses.replace(_module(), review_examples=block)
        with pytest.raises(ValueError, match="confidence"):
            validate_module_registry([bad])

    def test_california_examples_carry_multiple_json_shapes(self):
        # The CA module registered successfully at import (so its examples
        # already passed the real contract); pin that the block actually
        # contains the three JSON teaching shapes.
        examples = list(_iter_json_objects(CALIFORNIA_K12_MEP.review_examples))
        assert len(examples) == 3
        assert [e["actionType"] for e in examples] == ["EDIT", "ADD", "REPORT_ONLY"]


# ---------------------------------------------------------------------------
# Code basis (Phase 3)
# ---------------------------------------------------------------------------


class TestCodeBasisAccessors:
    def test_california_base_codes_and_accessors(self):
        cycle = CALIFORNIA_2025
        assert [c.key for c in cycle.base_codes] == [
            "cbc", "cmc", "cpc", "energy", "calgreen",
        ]
        assert cycle.code_year("cbc") == "2025"
        assert cycle.code_year("nope") == ""
        # First base code is the primary — the stale-detector target.
        assert cycle.primary_code_year == "2025"
        assert CodeCycle(label="empty").primary_code_year == ""

    def test_code_basis_format_kwargs_shape(self):
        kwargs = code_basis_format_kwargs(CALIFORNIA_2025)
        assert kwargs["cbc"] == "2025"
        assert kwargs["energy"] == "2025"
        assert kwargs["asce7"] == "7-22"
        assert kwargs["asce7_prev"] == "7-16"
        assert kwargs["pinned_standards"].startswith("NFPA 13 2025")


class TestCodeBasisValidation:
    def test_missing_base_codes_rejected(self):
        cycle = CodeCycle(label="9999", asce7="7-22", asce7_previous="7-16")
        bad = dataclasses.replace(_module(), cycle=cycle)
        with pytest.raises(ValueError, match="base_codes"):
            validate_module_registry([bad])

    def test_duplicate_base_code_keys_rejected(self):
        cycle = CodeCycle(
            label="9999",
            base_codes=(
                BaseCode("code", "A", "9999"),
                BaseCode("code", "B", "9999"),
            ),
            asce7="7-22",
            asce7_previous="7-16",
        )
        bad = dataclasses.replace(_module(), cycle=cycle)
        with pytest.raises(ValueError, match="unique"):
            validate_module_registry([bad])

    def test_code_basis_line_with_unknown_placeholder_rejected(self):
        bad = dataclasses.replace(
            _module(), cross_check_code_basis_line="Current cycle: {missing}."
        )
        with pytest.raises(ValueError, match="cross_check_code_basis_line"):
            validate_module_registry([bad])


class TestDetectorVocabularyValidation:
    def test_plausible_years_must_be_subset_of_valid(self):
        # The subset relationship is what keeps the stale and invalid
        # detectors disjoint by construction.
        bad_vocab = dataclasses.replace(
            _TEST_VOCABULARY, plausible_cycle_years=("2019", "1999")
        )
        bad = dataclasses.replace(_module(), detector_vocabulary=bad_vocab)
        with pytest.raises(ValueError, match="subset"):
            validate_module_registry([bad])

    def test_uncompilable_extra_pattern_rejected(self):
        bad_vocab = dataclasses.replace(
            _TEST_VOCABULARY, stale_cycle_extra_patterns=(r"\b(20\d{2}",)
        )
        bad = dataclasses.replace(_module(), detector_vocabulary=bad_vocab)
        with pytest.raises(ValueError, match="does not compile"):
            validate_module_registry([bad])

    def test_extra_pattern_without_year_group_rejected(self):
        bad_vocab = dataclasses.replace(
            _TEST_VOCABULARY, stale_cycle_extra_patterns=(r"\b20\d{2}\s+Code\b",)
        )
        bad = dataclasses.replace(_module(), detector_vocabulary=bad_vocab)
        with pytest.raises(ValueError, match="group 1"):
            validate_module_registry([bad])

    def test_empty_abbreviations_rejected(self):
        bad_vocab = dataclasses.replace(_TEST_VOCABULARY, code_abbreviations=())
        bad = dataclasses.replace(_module(), detector_vocabulary=bad_vocab)
        with pytest.raises(ValueError, match="abbreviation"):
            validate_module_registry([bad])

    def test_list_valued_fields_are_coerced_to_tuples(self):
        # Config-loaded vocabularies arrive with lists; the dataclass must
        # coerce them so the preprocessor's pattern cache (keyed by the
        # hashable vocabulary) never sees an unhashable field mid-run.
        vocab = DetectorVocabulary(
            code_abbreviations=["TBC"],  # type: ignore[arg-type]
            plausible_cycle_years=["2019"],  # type: ignore[arg-type]
            valid_cycle_years=["2019", "2025"],  # type: ignore[arg-type]
            asce7_plausible_editions=["22"],  # type: ignore[arg-type]
            stale_cycle_extra_patterns=[],  # type: ignore[arg-type]
        )
        assert vocab.code_abbreviations == ("TBC",)
        assert isinstance(vocab.valid_cycle_years, tuple)
        hash(vocab)  # must be cache-key safe
        validate_module_registry([_module(detector_vocabulary=vocab)])

    def test_bare_string_sequence_field_rejected_at_construction(self):
        # tuple("CBC") would silently become ("C", "B", "C") — reject instead.
        with pytest.raises(TypeError, match="single string"):
            DetectorVocabulary(
                code_abbreviations="TBC",  # type: ignore[arg-type]
                plausible_cycle_years=("2019",),
                valid_cycle_years=("2019",),
                asce7_plausible_editions=("22",),
            )

    def test_unhashable_element_rejected_at_registration(self):
        # A list nested INSIDE a tuple survives coercion but would blow up
        # the pattern cache — the hashability gate catches it at startup.
        bad_vocab = dataclasses.replace(
            _TEST_VOCABULARY, code_abbreviations=(["TBC"],)  # type: ignore[arg-type]
        )
        bad = dataclasses.replace(_module(), detector_vocabulary=bad_vocab)
        with pytest.raises(ValueError, match="hashable"):
            validate_module_registry([bad])


class TestVocabularyDrivenDetection:
    def test_synthetic_vocabulary_drives_stale_detection(self, monkeypatch):
        from src.input.preprocessor import detect_stale_code_cycle_references
        from src.modules import registry as registry_mod

        mod = _module("syn_detect", "7777")
        monkeypatch.setitem(registry_mod._MODULES_BY_CYCLE_LABEL, "7777", mod)

        content = "Comply with 2019 TBC Chapter 6 and with 2019 CBC Chapter 6."
        alerts = detect_stale_code_cycle_references(content, "a.docx", mod.cycle)
        # "TBC" is this module's abbreviation; "CBC" is NOT in its
        # vocabulary, so only one citation alerts.
        assert len(alerts) == 1
        assert alerts[0]["match"] == "2019 TBC"
        assert alerts[0]["expected_year"] == "7777"  # primary_code_year

    def test_invalid_detector_uses_vocabulary_and_label(self):
        from src.input.preprocessor import detect_invalid_code_cycle_strings

        alerts = detect_invalid_code_cycle_strings(
            "Per 2018 TBC requirements.", "a.docx", vocabulary=_TEST_VOCABULARY
        )
        assert len(alerts) == 1
        assert alerts[0]["type"] == "Invalid Test code cycle year (2018)"
        # 2019 is a valid year for this vocabulary — not invalid.
        assert not detect_invalid_code_cycle_strings(
            "Per 2019 TBC requirements.", "a.docx", vocabulary=_TEST_VOCABULARY
        )

    def test_leed_flag_gates_the_leed_detector(self, monkeypatch):
        from src.input.preprocessor import preprocess_spec
        from src.modules import registry as registry_mod

        no_leed_vocab = dataclasses.replace(
            _TEST_VOCABULARY, flag_leed_references=False
        )
        mod = _module("syn_leed", "7778", detector_vocabulary=no_leed_vocab)
        monkeypatch.setitem(registry_mod._MODULES_BY_CYCLE_LABEL, "7778", mod)

        content = "Project shall achieve LEED-NC Silver certification."
        assert preprocess_spec(content, "a.docx", cycle=mod.cycle).leed_alerts == []
        # The default (California) module keeps flagging LEED references.
        assert preprocess_spec(content, "a.docx", cycle=None).leed_alerts


# ---------------------------------------------------------------------------
# Verification-profile keywords + chunk groups (Phase 4)
# ---------------------------------------------------------------------------


class TestProfileParsing:
    def test_legacy_california_ahj_maps_to_jurisdictional(self):
        from src.verification.verification_profiles import (
            VerificationProfile,
            parse_verification_profile,
        )

        assert (
            parse_verification_profile("california_ahj")
            is VerificationProfile.JURISDICTIONAL
        )
        assert (
            parse_verification_profile("jurisdictional")
            is VerificationProfile.JURISDICTIONAL
        )
        assert (
            parse_verification_profile("not-a-profile")
            is VerificationProfile.CONSTRUCTABILITY
        )
        assert (
            parse_verification_profile(None)
            is VerificationProfile.CONSTRUCTABILITY
        )
        assert (
            parse_verification_profile(VerificationProfile.MANUFACTURER)
            is VerificationProfile.MANUFACTURER
        )


class TestModuleKeywordsDriveRouting:
    def test_keywords_param_drives_classification(self):
        from src.review.reviewer import Finding
        from src.verification.verification_profiles import (
            VerificationProfile,
            classify_finding_profile,
        )

        finding = Finding(
            severity="CRITICAL",
            fileName="21 13 13 - Sprinklers.docx",
            section="1.01",
            issue="Ruling by the Testville marshal conflicts with this clause.",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference=None,
        )
        # The synthetic module's vocabulary knows "testville"; the default
        # (California) vocabulary does not.
        assert (
            classify_finding_profile(finding, keywords=_TEST_PROFILE_KEYWORDS)
            is VerificationProfile.JURISDICTIONAL
        )
        assert (
            classify_finding_profile(finding)
            is VerificationProfile.CONSTRUCTABILITY
        )

    def test_select_routing_resolves_keywords_from_cycle(self, monkeypatch):
        from src.modules import registry as registry_mod
        from src.review.reviewer import Finding
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        mod = _module("syn_routing", "6666")
        monkeypatch.setitem(registry_mod._MODULES_BY_CYCLE_LABEL, "6666", mod)
        finding = Finding(
            severity="CRITICAL",
            fileName="21 13 13 - Sprinklers.docx",
            section="1.01",
            issue="Ruling by the Testville marshal conflicts with this clause.",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference=None,
        )
        # With the owning module's cycle: jurisdictional -> CRITICAL goes deep.
        routed = select_routing(finding, local_skip=False, cycle=mod.cycle)
        assert routed.profile is VerificationProfile.JURISDICTIONAL
        assert routed.mode is VerificationMode.DEEP_REASONING
        # Without module context the default (CA) vocabulary sees nothing.
        default_routed = select_routing(finding, local_skip=False)
        assert default_routed.profile is VerificationProfile.CONSTRUCTABILITY
        assert default_routed.mode is VerificationMode.STANDARD_REASONING

    def test_empty_keyword_field_rejected(self):
        bad_keywords = dataclasses.replace(_TEST_PROFILE_KEYWORDS, jurisdictional=())
        bad = dataclasses.replace(_module(), profile_keywords=bad_keywords)
        with pytest.raises(ValueError, match="jurisdictional"):
            validate_module_registry([bad])


class TestChunkGroupValidation:
    def test_duplicate_prefix_across_groups_rejected(self):
        groups = (
            ChunkGroup("a", "A", ("21",)),
            ChunkGroup("b", "B", ("21", "22")),
        )
        bad = dataclasses.replace(_module(), cross_check_chunk_groups=groups)
        with pytest.raises(ValueError, match="more than one chunk group"):
            validate_module_registry([bad])

    def test_reserved_general_id_rejected(self):
        groups = (ChunkGroup("general", "General", ("21",)),)
        bad = dataclasses.replace(_module(), cross_check_chunk_groups=groups)
        with pytest.raises(ValueError, match="reserved"):
            validate_module_registry([bad])

    def test_duplicate_chunk_id_rejected(self):
        groups = (
            ChunkGroup("dup", "A", ("21",)),
            ChunkGroup("dup", "B", ("22",)),
        )
        bad = dataclasses.replace(_module(), cross_check_chunk_groups=groups)
        with pytest.raises(ValueError, match="duplicate chunk_id"):
            validate_module_registry([bad])

    def test_module_groups_drive_chunk_assignment(self):
        from src.cross_check.cross_checker import _assign_chunk, _chunk_label

        # The synthetic module maps 21/22 only — 23 pools into "general".
        assert _assign_chunk("21 13 13 - Wet.docx", _TEST_CHUNK_GROUPS) == "div_21"
        assert _assign_chunk("23 05 00 - HVAC.docx", _TEST_CHUNK_GROUPS) == "general"
        assert _chunk_label("div_22", _TEST_CHUNK_GROUPS) == "Division 22 — Plumbing"
        # Default (no groups passed) resolves the California module's map.
        assert _assign_chunk("23 05 00 - HVAC.docx") == "div_23"


# ---------------------------------------------------------------------------
# Slots drive the assembled prompts
# ---------------------------------------------------------------------------


class TestSlotsDriveOutput:
    def test_synthetic_module_slots_appear_in_prompts(self, monkeypatch):
        from src.cross_check.cross_checker import _cross_system_prompt
        from src.modules import registry as registry_mod
        from src.review.prompts import get_single_spec_user_message, get_system_prompt
        from src.verification.verifier import _get_verification_system_prompt

        mod = _module(
            "synthetic",
            "8888",
            reviewer_persona="You are a synthetic persona for slot testing.",
            review_user_intro="Synthetic intro line.",
            review_confidence_high_example="a synthetic high example",
            review_categories_template="1. Synthetic category for cycle {code}.",
            cross_check_persona="Synthetic cross-check persona.",
            verifier_persona="Synthetic verifier persona.",
            verifier_source_priorities="1. Synthetic sources:\n   synthetic.example",
        )
        monkeypatch.setitem(registry_mod._MODULES_BY_CYCLE_LABEL, "8888", mod)

        sp = get_system_prompt(mod.cycle)
        assert sp.startswith("You are a synthetic persona for slot testing.\n\n<task>")
        assert "1. Synthetic category for cycle 8888." in sp
        assert "Current code cycle: Test Code 8888." in get_single_spec_user_message(
            "body", "a.docx", "", cycle=mod.cycle,
            paragraph_map=None, pre_detected_alerts=None,
        )
        assert "(e.g., a synthetic high example)" in sp
        assert "Example 1 — REPORT_ONLY:" in sp

        um = get_single_spec_user_message(
            "body", "a.docx", "", cycle=mod.cycle,
            paragraph_map=None, pre_detected_alerts=None,
        )
        assert um.startswith("Synthetic intro line.\n\n")

        cp = _cross_system_prompt(mod.cycle)
        assert cp.startswith("Synthetic cross-check persona.\n\n")

        vp = _get_verification_system_prompt(mod.cycle, include_verdict_tool=True)
        assert "Synthetic verifier persona." in vp
        assert "   synthetic.example" in vp
        # Engine protocol stays regardless of module.
        assert "Prefer authoritative sources in this priority order:" in vp


# ---------------------------------------------------------------------------
# GUI selection persistence + report surfaces (Phase 5)
# ---------------------------------------------------------------------------


class TestUiStatePersistence:
    def test_round_trip(self, tmp_path):
        from src.core.ui_state import load_selected_module_id, save_selected_module_id

        state_path = tmp_path / "ui_state.json"
        assert load_selected_module_id(path=state_path) == ""
        save_selected_module_id("california_k12_mep", path=state_path)
        assert load_selected_module_id(path=state_path) == "california_k12_mep"

    def test_corrupt_file_reads_as_no_selection(self, tmp_path):
        from src.core.ui_state import load_selected_module_id

        state_path = tmp_path / "ui_state.json"
        state_path.write_text("{not json", encoding="utf-8")
        assert load_selected_module_id(path=state_path) == ""

    def test_env_override_path(self, tmp_path, monkeypatch):
        from src.core.ui_state import ui_state_path

        monkeypatch.setenv("SPEC_CRITIC_UI_STATE_PATH", str(tmp_path / "s.json"))
        assert ui_state_path() == tmp_path / "s.json"

    def test_stale_saved_id_degrades_to_default(self, tmp_path):
        # The GUI resolves the saved id through get_module, so an id from
        # an uninstalled module falls back to the default module.
        from src.core.ui_state import load_selected_module_id, save_selected_module_id

        state_path = tmp_path / "ui_state.json"
        save_selected_module_id("module_that_no_longer_exists", path=state_path)
        resolved = get_module(load_selected_module_id(path=state_path))
        assert resolved is DEFAULT_MODULE


class TestReportSurfaces:
    def test_methodology_note_renders_module_phrase_and_jurisdiction(self):
        from docx import Document

        from src.output.report_exporter import _write_methodology_note

        doc = Document()
        _write_methodology_note(
            doc,
            cycle_label="9999",
            module=_module(report_context_phrase="hyperscale data-center projects"),
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "relevant to hyperscale data-center projects." in text
        # Cycle-references sentence takes the module vocabulary's
        # jurisdiction wording; the test cycle pins no standards, so no
        # pinned-editions paragraph renders.
        assert "This review used Test 9999 code cycle references." in text
        assert "California" not in text

    def test_methodology_note_default_keeps_california_sentences(self):
        from docx import Document

        from src.output.report_exporter import _write_methodology_note

        doc = Document()
        _write_methodology_note(doc)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "relevant to California K-12 DSA projects." in text
        assert "This review used California 2025 code cycle references." in text
        # Pinned-editions paragraph renders the default module's own cycle.
        assert "per the 2025 California cycle:" in text

    def test_title_block_renders_module_title_and_jurisdiction(self):
        from docx import Document

        from src.output.report_exporter import _write_title_block

        doc = Document()
        _write_title_block(
            doc,
            _StubReview(),
            ["a.docx"],
            cycle_label="9999",
            module=_module(
                report_title="Spec Critic — Fire Protection Specification Review Report"
            ),
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Spec Critic — Fire Protection Specification Review Report" in text
        assert "Code Cycle: Test 9999" in text
        assert "California" not in text

    def test_title_block_without_jurisdiction_renders_bare_cycle_line(self):
        from dataclasses import replace as dc_replace

        from docx import Document

        from src.output.report_exporter import _write_title_block

        doc = Document()
        _write_title_block(
            doc,
            _StubReview(),
            ["a.docx"],
            cycle_label="9999",
            module=_module(
                detector_vocabulary=dc_replace(
                    _TEST_VOCABULARY, jurisdiction_label=""
                )
            ),
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Code Cycle: 9999" in text

    def test_title_block_default_keeps_california_lines(self):
        from docx import Document

        from src.output.report_exporter import _write_title_block

        doc = Document()
        _write_title_block(doc, _StubReview(), ["a.docx"])
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Spec Critic — M&P Specification Review Report" in text
        assert "Code Cycle: California 2025" in text

    def test_alert_sections_render_module_jurisdiction_and_years(self):
        from docx import Document

        from src.output.report_exporter import _write_alerts

        doc = Document()
        _write_alerts(
            doc,
            [],
            [],
            code_cycle_alerts=[{"filename": "a.docx", "context": "2019 TBC"}],
            invalid_code_cycle_alerts=[{"filename": "a.docx", "context": "2018 TBC"}],
            module=_module(),
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Stale Test Code Cycle References" in text
        assert "Invalid Test Code Cycle Years" in text
        # Published years come from plausible_cycle_years — NOT
        # valid_cycle_years, which also admits anticipated future cycles.
        assert "(Test publishes cycles every 3 years: 2019, 2022)" in text
        assert "California" not in text

    def test_alert_sections_generic_without_jurisdiction_or_cadence(self):
        from dataclasses import replace as dc_replace

        from docx import Document

        from src.output.report_exporter import _write_alerts

        vocab = dc_replace(
            _TEST_VOCABULARY,
            jurisdiction_label="",
            plausible_cycle_years=("2018", "2021", "2027"),
            valid_cycle_years=("2018", "2021", "2027"),
        )
        doc = Document()
        _write_alerts(
            doc,
            [],
            [],
            code_cycle_alerts=[{"filename": "a.docx", "context": "2018 TBC"}],
            invalid_code_cycle_alerts=[{"filename": "a.docx", "context": "2019 TBC"}],
            module=_module(detector_vocabulary=vocab),
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Stale Code Cycle References" in text
        assert "Invalid Code Cycle Years" in text
        # Irregular gaps (3 then 6) drop the cadence claim.
        assert "(known cycle years: 2018, 2021, 2027)" in text
        assert "publishes cycles every" not in text

    def test_alert_sections_default_keep_california_wording(self):
        from docx import Document

        from src.output.report_exporter import _write_alerts

        doc = Document()
        _write_alerts(
            doc,
            [],
            [],
            code_cycle_alerts=[{"filename": "a.docx", "context": "2019 CBC"}],
            invalid_code_cycle_alerts=[{"filename": "a.docx", "context": "2018 CBC"}],
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Stale California Code Cycle References" in text
        assert (
            "The following references cite a historical California code cycle "
            "rather than the one selected for this review:"
        ) in text
        assert "Invalid California Code Cycle Years" in text
        # 2028 is in valid_cycle_years (anticipated, detector-admissible)
        # but is not a published cycle, so it must not appear here.
        assert (
            "The following references cite a year that is not a real California "
            "code cycle (California publishes cycles every 3 years: 2010, 2013, "
            "2016, 2019, 2022, 2025). These are likely typos:"
        ) in text

    def test_default_module_report_phrase_and_title(self):
        assert DEFAULT_MODULE.report_context_phrase == "California K-12 DSA projects"
        assert DEFAULT_MODULE.report_title == (
            "Spec Critic — M&P Specification Review Report"
        )

    def test_diagnostics_report_carries_module_id(self):
        from src.orchestration.diagnostics import DiagnosticsReport

        assert DiagnosticsReport().module_id == ""
        assert DiagnosticsReport(module_id="california_k12_mep").module_id == (
            "california_k12_mep"
        )


# ---------------------------------------------------------------------------
# Persistence round-trip (submission -> pending state -> submission)
# ---------------------------------------------------------------------------


class TestModuleIdRoundTrip:
    def test_submission_defaults_carry_the_default_module(self):
        sub = _submission()
        assert sub.module_id == DEFAULT_MODULE.module_id
        assert PipelineResult(review_result=None).module_id == DEFAULT_MODULE.module_id

    def test_pending_batch_round_trips_module_id(self, tmp_path):
        state_path = tmp_path / "pending_batch.json"
        pending = PendingBatch.from_submission(_submission())
        assert pending.module_id == DEFAULT_MODULE.module_id

        save_pending_batch(pending, path=state_path)
        loaded = load_pending_batch(path=state_path)
        assert loaded is not None
        assert loaded.module_id == DEFAULT_MODULE.module_id

        # to_submission resolves the module and re-stamps identity. Files are
        # absent, so reconstruction is the findings-only path (no re-extract).
        sub = loaded.to_submission(log=lambda *a, **k: None)
        assert sub.module_id == DEFAULT_MODULE.module_id
        assert sub.cycle_label == DEFAULT_MODULE.cycle.label
        assert sub.job.request_map == {
            "review__a__0": {"filename": "a.docx", "index": 0, "type": "review"}
        }

    def test_legacy_state_file_without_module_id_defaults(self, tmp_path):
        """A pending_batch.json written before Phase 1 has no module_id key.

        It must load (same schema version — the field was additive) and
        resolve to the default California module, the only configuration a
        legacy file could have been written by.
        """
        state_path = tmp_path / "pending_batch.json"
        legacy = {
            "batch_id": "msgbatch_LEGACY",
            "model": "claude-opus-4-8",
            "request_map": {},
            "review_request_ids": [],
            "files_reviewed": [],
            "input_dir": "",
            "files": [],
            "cycle_label": "2025",
            "project_context": "",
            "cross_check_enabled": False,
            "submitted_at": 1700000000.0,
            "run_id": "",
            "app_version": "2.9.0",
            "schema_version": 1,
        }
        state_path.write_text(json.dumps(legacy), encoding="utf-8")
        loaded = load_pending_batch(path=state_path)
        assert loaded is not None
        assert loaded.module_id == DEFAULT_MODULE.module_id

    def test_new_state_file_keeps_schema_version_1(self, tmp_path):
        # Additive field, defensive loader: old readers ignore the key, so
        # the schema version intentionally does NOT bump.
        state_path = tmp_path / "pending_batch.json"
        save_pending_batch(PendingBatch.from_submission(_submission()), path=state_path)
        data = json.loads(state_path.read_text(encoding="utf-8"))
        assert data["schema_version"] == 1
        assert data["module_id"] == DEFAULT_MODULE.module_id


# ---------------------------------------------------------------------------
# Downstream stamping
# ---------------------------------------------------------------------------


class TestDownstreamStamping:
    def test_finalize_carries_module_id_onto_pipeline_result(self):
        state = CollectedBatchState(
            submission=_submission(),
            review_result=ReviewResult(findings=[]),
        )
        result = finalize_batch_result(state)
        assert result.module_id == DEFAULT_MODULE.module_id
        assert result.cycle_label == DEFAULT_MODULE.cycle.label

    def test_finalize_defaults_when_submission_lacks_the_field(self):
        class _BareJob:
            created_at = 1700000000.0

        class _BareSubmission:
            job = _BareJob()
            cycle_label = "2025"
            cross_check_enabled = False
            prepared_specs = None
            files_reviewed: list = []

        state = CollectedBatchState(
            submission=_BareSubmission(),  # type: ignore[arg-type]
            review_result=ReviewResult(findings=[]),
        )
        assert finalize_batch_result(state).module_id == DEFAULT_MODULE.module_id

    def test_trace_run_meta_records_module_id(self, tmp_path):
        trace_dir = tmp_path / "trace"
        rec = TraceRecorder(run_id="mod1", trace_dir=trace_dir, capture_level="default")
        rec.start(
            mode="batch",
            model="claude-opus-4-8",
            cycle_label="2025",
            module_id="california_k12_mep",
        )
        rec.stop()
        meta = json.loads((trace_dir / "run.json").read_text(encoding="utf-8"))
        assert meta["module_id"] == "california_k12_mep"
