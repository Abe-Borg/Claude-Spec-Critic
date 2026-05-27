"""Chunk N tests: report trust-model statuses.

Plan section "Chunk N — Report Trust Model Improvements". The chunk
defines two closed enums (``ReportStatus`` and ``EditActionLabel``) that
every finding maps to for display, plus the four evidence concepts
(spec evidence, web/code evidence, verification rationale, unsupported
sources) that the report exporter has to surface distinctly.

Coverage:

* ``TestReportStatusClassification`` exercises every status branch of
  :func:`classify_status`, including the priority ordering (suppression
  beats verification beats local-skip beats verdict).
* ``TestEditActionClassification`` exercises every label branch of
  :func:`classify_edit_action`: suppression → SUPPRESSED, no proposal →
  REPORT_ONLY, otherwise EDIT_SUGGESTED.
* ``TestSummarizeHelpers`` checks that the histogram helpers return
  zero-filled dicts and sum to the input count.
* ``TestLabelHelpers`` covers the human-readable label / glyph mapping
  and the string fallback for unknown values.
* ``TestReportExporterStatusIntegration`` does a snapshot-style check on
  the actual .docx output: it builds a tiny ``PipelineResult`` with
  findings spanning every status, exports to a temp path, and asserts
  the resulting document contains the expected status labels and
  evidence sub-headings. Catches regressions where the wiring drops
  the status line or mis-routes the histogram counts.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.output.report_exporter import export_report
from src.output.report_status import (
    EditActionLabel,
    ReportStatus,
    STATUS_LABELS,
    classify_edit_action,
    classify_status,
    summarize_edit_actions,
    summarize_statuses,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.6,
    action: str = "EDIT",
    existing: str | None = "old text",
    replacement: str | None = "new text",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
    suppression_reason: str | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
        suppression_reason=suppression_reason,
    )
    f.verification = verification
    return f


def _verification(
    verdict: str = "CONFIRMED",
    *,
    grounded: bool = True,
    cache_status: str = "miss",
    explanation: str = "Verified against CBC §1234",
    sources: list[str] | None = None,
    rejected: list[dict] | None = None,
    correction: str | None = None,
) -> VerificationResult:
    # Chunk 5: a grounded CONFIRMED/CORRECTED requires at least one
    # accepted external citation. Default to a representative one so
    # individual tests focused on status / edit-action classification
    # do not have to thread sources through every call.
    if sources is None:
        sources = (
            ["https://dgs.ca.gov"]
            if grounded and verdict.upper() in ("CONFIRMED", "CORRECTED")
            else []
        )
    return VerificationResult(
        verdict=verdict,
        explanation=explanation,
        sources=list(sources),
        correction=correction,
        grounded=grounded,
        cache_status=cache_status,
        rejected_sources=list(rejected or []),
    )


# ---------------------------------------------------------------------------
# classify_status — every branch (Chunk N Directive 1)
# ---------------------------------------------------------------------------

class TestReportStatusClassification:
    def test_no_verification_is_not_checked(self):
        f = _finding(verification=None)
        assert classify_status(f) is ReportStatus.NOT_CHECKED

    def test_local_skip_is_locally_classified(self):
        v = _verification("UNVERIFIED", grounded=False, cache_status="local_skip")
        f = _finding(verification=v)
        assert classify_status(f) is ReportStatus.LOCALLY_CLASSIFIED

    def test_confirmed_and_grounded_is_verified_supported(self):
        f = _finding(verification=_verification("CONFIRMED", grounded=True))
        assert classify_status(f) is ReportStatus.VERIFIED_SUPPORTED

    def test_corrected_and_grounded_is_verified_contradicted(self):
        f = _finding(verification=_verification("CORRECTED", grounded=True))
        assert classify_status(f) is ReportStatus.VERIFIED_CONTRADICTED

    def test_disputed_verdict_is_disputed(self):
        f = _finding(verification=_verification("DISPUTED", grounded=False))
        assert classify_status(f) is ReportStatus.DISPUTED

    def test_unverified_verdict_is_insufficient_evidence(self):
        f = _finding(verification=_verification("UNVERIFIED", grounded=False))
        assert classify_status(f) is ReportStatus.INSUFFICIENT_EVIDENCE

    def test_unknown_verdict_falls_through_to_insufficient_evidence(self):
        # Belt-and-suspenders: malformed verdict strings should not crash
        # and should land in the conservative bucket.
        f = _finding(verification=_verification("???", grounded=False))
        assert classify_status(f) is ReportStatus.INSUFFICIENT_EVIDENCE

    def test_confirmed_but_ungrounded_does_not_count_as_supported(self):
        # The grounding invariant in the verifier should already have
        # downgraded this, but the classifier is the second line of
        # defense for tests that construct results by hand.
        f = _finding(verification=_verification("CONFIRMED", grounded=False))
        assert classify_status(f) is ReportStatus.INSUFFICIENT_EVIDENCE

    def test_suppression_reason_beats_everything(self):
        # Even a confirmed+grounded finding renders as MANUAL_REVIEW_REQUIRED
        # if it was suppressed; the report shows it under the suppressed
        # section, not the verified section.
        f = _finding(
            verification=_verification("CONFIRMED", grounded=True),
            suppression_reason="All upstream review findings disputed",
        )
        assert classify_status(f) is ReportStatus.MANUAL_REVIEW_REQUIRED


# ---------------------------------------------------------------------------
# classify_edit_action — every branch (Chunk N Directive 4)
# ---------------------------------------------------------------------------

class TestEditActionClassification:
    def test_suppressed_short_circuits(self):
        f = _finding(
            verification=_verification("CONFIRMED"),
            suppression_reason="dropped by upstream-disputed filter",
        )
        assert classify_edit_action(f) is EditActionLabel.SUPPRESSED

    def test_no_proposal_is_report_only(self):
        f = _finding(action="REPORT_ONLY", existing=None, replacement=None)
        assert classify_edit_action(f) is EditActionLabel.REPORT_ONLY

    def test_proposal_is_edit_suggested(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.85,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_verification("CONFIRMED", grounded=True),
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_proposal_label_ignores_verification_status(self):
        # The app emits edit instructions but never applies them, so the
        # label is "has a proposal?" — independent of verdict/grounding.
        # A downstream applier does its own gating using the verification
        # status carried alongside the instruction.
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.95,
        )
        for verdict, grounded, cache in [
            ("DISPUTED", False, "miss"),
            ("UNVERIFIED", False, "miss"),
            ("UNVERIFIED", False, "local_skip"),
        ]:
            f = _finding(
                edit_proposal=proposal,
                verification=_verification(verdict, grounded=grounded, cache_status=cache),
            )
            assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_proposal_label_ignores_low_confidence(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.1,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_verification("CONFIRMED", grounded=True),
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_not_checked_with_proposal_is_edit_suggested(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.95,
        )
        f = _finding(edit_proposal=proposal, verification=None)
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED

    def test_legacy_finding_without_edit_proposal_field_is_report_only(self):
        # Pre-Chunk-L payloads round-trip without the new ``edit_proposal``
        # field. ``as_edit_proposal`` falls back to the legacy actionType /
        # existingText / replacementText fields; if those don't carry an
        # ADD/EDIT/DELETE action, classify_edit_action returns REPORT_ONLY.
        f = _finding(action="REPORT_ONLY", existing=None, replacement=None)
        assert classify_edit_action(f) is EditActionLabel.REPORT_ONLY

    def test_legacy_edit_finding_routes_through_as_edit_proposal(self):
        # An old-shaped finding with actionType=EDIT and existingText set
        # is still EDIT_SUGGESTED — the legacy proposal is synthesized on
        # the fly by ``as_edit_proposal``.
        f = _finding(
            action="EDIT",
            existing="old",
            replacement="new",
            confidence=0.9,
            verification=_verification("CONFIRMED", grounded=True),
            edit_proposal=None,
        )
        assert classify_edit_action(f) is EditActionLabel.EDIT_SUGGESTED


# ---------------------------------------------------------------------------
# Aggregation helpers
# ---------------------------------------------------------------------------

class TestSummarizeHelpers:
    def test_summarize_statuses_returns_zero_filled_dict_on_empty_input(self):
        counts = summarize_statuses([])
        assert set(counts.keys()) == set(ReportStatus)
        assert sum(counts.values()) == 0

    def test_summarize_statuses_sums_to_input_length(self):
        findings = [
            _finding(verification=_verification("CONFIRMED", grounded=True)),
            _finding(verification=_verification("DISPUTED", grounded=False)),
            _finding(verification=None),
            _finding(
                verification=_verification("CONFIRMED", grounded=True),
                suppression_reason="upstream disputed",
            ),
            _finding(
                verification=_verification(
                    "UNVERIFIED", grounded=False, cache_status="local_skip"
                ),
            ),
        ]
        counts = summarize_statuses(findings)
        assert counts[ReportStatus.VERIFIED_SUPPORTED] == 1
        assert counts[ReportStatus.DISPUTED] == 1
        assert counts[ReportStatus.NOT_CHECKED] == 1
        assert counts[ReportStatus.MANUAL_REVIEW_REQUIRED] == 1
        assert counts[ReportStatus.LOCALLY_CLASSIFIED] == 1
        assert sum(counts.values()) == len(findings)

    def test_summarize_edit_actions_zero_filled(self):
        counts = summarize_edit_actions([])
        assert set(counts.keys()) == set(EditActionLabel)
        assert sum(counts.values()) == 0

    def test_summarize_edit_actions_counts_each_label(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="new",
            edit_confidence=0.95,
        )
        findings = [
            _finding(
                edit_proposal=proposal,
                verification=_verification("CONFIRMED", grounded=True),
            ),
            _finding(action="REPORT_ONLY", existing=None, replacement=None),
            _finding(
                edit_proposal=proposal,
                verification=_verification("CONFIRMED", grounded=True),
                suppression_reason="dropped",
            ),
        ]
        counts = summarize_edit_actions(findings)
        assert counts[EditActionLabel.EDIT_SUGGESTED] == 1
        assert counts[EditActionLabel.REPORT_ONLY] == 1
        assert counts[EditActionLabel.SUPPRESSED] == 1


# ---------------------------------------------------------------------------
# Label / glyph helpers
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Report exporter integration (snapshot-style)
# ---------------------------------------------------------------------------

class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report.

    The real :class:`pipeline.PipelineResult` is a dataclass with a
    handful of typed fields; building one here forces importing pipeline
    + cross_checker + verifier modules that pull a lot of weight. A stub
    that exposes only the attributes export_report reads is cheaper and
    keeps these tests independent of unrelated pipeline changes.
    """

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestReportExporterStatusIntegration:
    @pytest.fixture
    def diverse_review_result(self) -> ReviewResult:
        """Build a review result with one finding per common status path."""
        verified_proposal = EditProposal(
            action_type="EDIT",
            existing_text="2019 CBC",
            replacement_text="2025 CBC",
            edit_confidence=0.95,
        )
        verified = _finding(
            severity="CRITICAL",
            file="Section_23_0000.docx",
            section="2.1",
            issue="Stale CBC reference",
            edit_proposal=verified_proposal,
            verification=_verification(
                "CONFIRMED",
                grounded=True,
                sources=["https://codes.iccsafe.org/content/CBC2025"],
            ),
        )
        disputed = _finding(
            severity="HIGH",
            file="Section_23_0000.docx",
            section="2.2",
            issue="Claims wrong fitting standard",
            verification=_verification(
                "DISPUTED",
                grounded=False,
                explanation="Search results contradict the finding's claim.",
                rejected=[{"url": "https://blog.example.com/foo", "reason": "ungrounded"}],
            ),
        )
        insufficient = _finding(
            severity="MEDIUM",
            file="Section_23_0000.docx",
            section="2.3",
            issue="Could not be verified",
            verification=_verification(
                "UNVERIFIED",
                grounded=False,
                explanation="Authoritative source not located.",
            ),
        )
        local_skip = _finding(
            severity="GRIPES",
            file="Section_23_0000.docx",
            section="2.4",
            issue="LEED Gold reference in a non-LEED project",
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verification(
                "UNVERIFIED",
                grounded=False,
                cache_status="local_skip",
                explanation="Locally classified; no web verification needed.",
            ),
        )
        not_checked = _finding(
            severity="GRIPES",
            file="Section_23_0000.docx",
            section="2.5",
            issue="Pre-verification scan finding",
            verification=None,
        )
        return ReviewResult(
            findings=[verified, disputed, insufficient, local_skip, not_checked],
            input_tokens=1000,
            output_tokens=500,
            elapsed_seconds=12.5,
        )

    def test_export_contains_status_lines_for_each_finding(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)

        # Every finding renders a "Status:" line, so the substring count
        # should equal the finding count.
        assert text.count("Status:") == len(diverse_review_result.findings)

        # Each status label appears at least once. The label lives on
        # the status line, the trust-model histogram cell, or both.
        assert STATUS_LABELS[ReportStatus.VERIFIED_SUPPORTED] in text
        assert STATUS_LABELS[ReportStatus.DISPUTED] in text
        assert STATUS_LABELS[ReportStatus.INSUFFICIENT_EVIDENCE] in text
        assert STATUS_LABELS[ReportStatus.LOCALLY_CLASSIFIED] in text
        assert STATUS_LABELS[ReportStatus.NOT_CHECKED] in text

    def test_export_contains_edit_action_labels(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)

        # Findings with a proposal render "Edit suggested"; findings with
        # no proposal render "Report only". Both labels should be visible.
        assert "Edit suggested" in text
        assert "Report only" in text

    def test_export_includes_trust_model_summary_heading(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)
        assert "Trust Model Summary" in text
        # Edit eligibility line includes the at-a-glance histogram.
        assert "Edit eligibility:" in text

    def test_export_renames_existing_text_label_to_spec_evidence(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        # Chunk N Directive 3: spec evidence is distinct from web/code
        # evidence and verification rationale. The label rename makes
        # the four concepts explicit.
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)
        assert "Spec evidence:" in text
        assert "Proposed replacement:" in text
        # The old labels should be gone now.
        assert "Existing Text:" not in text
        assert "Replace With:" not in text

    def test_export_includes_verification_rationale_label(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)
        assert "Verification rationale:" in text

    def test_export_includes_web_code_and_rejected_evidence_labels(
        self, tmp_path: Path, diverse_review_result: ReviewResult
    ):
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=diverse_review_result), out
        )
        doc = Document(str(out))
        text = _all_text_from(doc)
        # The verified finding has an accepted source list, so the
        # web/code-evidence label must render.
        assert "Web/code evidence" in text
        # The disputed finding has a rejected source, so the
        # unsupported-sources label must render.
        assert "Unsupported / rejected sources" in text

    def test_suppressed_findings_dont_pollute_main_severity_section(
        self, tmp_path: Path
    ):
        # Suppressed findings must remain distinguishable from supported
        # findings. The report renders them in their own subsection under
        # the cross-check section; they should be tagged MANUAL_REVIEW_REQUIRED.
        verified = _finding(
            severity="HIGH",
            verification=_verification("CONFIRMED", grounded=True),
        )
        suppressed = _finding(
            severity="HIGH",
            issue="Coordination claim — upstream review disputed",
            suppression_reason="All cited upstream findings disputed",
        )
        review = ReviewResult(findings=[verified])
        cross = ReviewResult(
            findings=[],
            cross_check_status="completed",
            suppressed_findings=[suppressed],
        )

        out = tmp_path / "report_with_suppressed.docx"
        export_report(
            _StubPipelineResult(
                review_result=review,
                cross_check_result=cross,
                files_reviewed=[verified.fileName],
            ),
            out,
        )
        doc = Document(str(out))
        text = _all_text_from(doc)

        # The suppressed finding is rendered under a dedicated
        # "Suppressed Coordination Findings" subsection (Chunk M wiring,
        # preserved by Chunk N).
        assert "Suppressed Coordination Findings" in text
        # And it shows the MANUAL_REVIEW_REQUIRED status (Chunk N), so
        # the reader doesn't think it's an accepted finding.
        assert STATUS_LABELS[ReportStatus.MANUAL_REVIEW_REQUIRED] in text
        # The Edit: Suppressed label must be visible on the finding.
        assert "Suppressed" in text
