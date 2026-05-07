"""Regression tests for audit Sprint 1 (P0) fixes.

Each block here corresponds to a numbered audit issue from the
implementation plan. Tests are written so they would have failed
before the corresponding fix landed.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest
from docx import Document

from src.apply_edits import execute_edit_plan
from src.batch import BatchJob
from src.edit_locator import EditLocation, LocatorResult, locate_edit
from src.extractor import ExtractedSpec, ParagraphMapping, extract_text
from src.pipeline import (
    BatchSubmission,
    CollectedBatchState,
    _deduplicate_findings,
    collect_review_batch_results,
    run_cross_check_for_batch,
)
from src.reviewer import Finding, ReviewResult, _parse_findings
from src.spec_editor import build_edit_actions
from src.verifier import (
    _classify_wave_results,
    _collect_search_evidence,
    _content_block_to_plain,
    _search_gate_failure,
)
from src.resume_state import deserialize_finding, serialize_finding


# ---------------------------------------------------------------------------
# Test scaffolding
# ---------------------------------------------------------------------------


class _FakeBlock:
    def __init__(self, block_type: str, **kwargs):
        self.type = block_type
        for k, v in kwargs.items():
            setattr(self, k, v)


class _FakeUrlResult:
    def __init__(self, url: str):
        self.type = "web_search_result"
        self.url = url


class _FakeError:
    def __init__(self):
        self.type = "web_search_tool_result_error"


class _FakeUsage:
    def __init__(self, web_search_requests: int):
        class _ServerToolUse:
            pass

        self.server_tool_use = _ServerToolUse()
        self.server_tool_use.web_search_requests = web_search_requests


class _FakeMessage:
    def __init__(self, content, *, usage=None, stop_reason="end_turn"):
        self.content = content
        self.usage = usage
        self.stop_reason = stop_reason


def _make_finding(
    *,
    file_name: str = "spec.docx",
    issue: str = "Issue",
    existing: str | None = "old text",
    replacement: str | None = "new text",
    action: str = "EDIT",
) -> Finding:
    return Finding(
        severity="HIGH",
        fileName=file_name,
        section="1",
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC",
        confidence=0.9,
    )


def _mapping(text: str, *, idx: int) -> ParagraphMapping:
    return ParagraphMapping(
        body_index=idx,
        element_type="paragraph",
        text=text,
        table_index=None,
        row_index=None,
        cell_index=None,
    )


# ---------------------------------------------------------------------------
# Issue 1 — Verification search-success counting
# ---------------------------------------------------------------------------


def test_search_evidence_does_not_count_error_only_block_as_success():
    message = _FakeMessage(
        [_FakeBlock("web_search_tool_result", content=[_FakeError(), _FakeError()])],
        usage=_FakeUsage(web_search_requests=2),
    )
    urls, success_count, error_count = _collect_search_evidence(message)
    assert urls == []
    assert success_count == 0
    assert error_count == 2


def test_search_gate_fails_when_only_errors_returned():
    message = _FakeMessage(
        [_FakeBlock("web_search_tool_result", content=[_FakeError()])],
        usage=_FakeUsage(web_search_requests=1),
    )
    failure = _search_gate_failure(message)
    assert failure is not None
    assert "all 1 search requests failed" in failure.lower()


def test_search_gate_passes_when_at_least_one_real_result():
    message = _FakeMessage(
        [
            _FakeBlock(
                "web_search_tool_result",
                content=[_FakeError(), _FakeUrlResult("https://example.org/a")],
            )
        ],
        usage=_FakeUsage(web_search_requests=1),
    )
    assert _search_gate_failure(message) is None
    urls, success, errors = _collect_search_evidence(message)
    assert urls == ["https://example.org/a"]
    assert success == 1
    assert errors == 1


# ---------------------------------------------------------------------------
# Issue 2 — Deduplication truncation collision
# ---------------------------------------------------------------------------


def test_dedup_does_not_collide_when_only_first_200_chars_match():
    base_prefix = "x" * 250
    f1 = _make_finding(existing=base_prefix + "ENDING_A")
    f2 = _make_finding(file_name="other.docx", existing=base_prefix + "ENDING_B")
    deduped = _deduplicate_findings([f1, f2])
    assert len(deduped) == 2


def test_dedup_still_merges_truly_identical_findings_across_files():
    f1 = _make_finding(existing="identical text", replacement="x")
    f2 = _make_finding(file_name="other.docx", existing="identical text", replacement="x")
    deduped = _deduplicate_findings([f1, f2])
    assert len(deduped) == 1
    assert sorted(deduped[0].affected_files) == ["other.docx", "spec.docx"]


# ---------------------------------------------------------------------------
# Issue 3 — affected_files fan-out at edit application
# ---------------------------------------------------------------------------


def test_execute_edit_plan_applies_to_every_affected_file(tmp_path: Path):
    file_a = tmp_path / "a.docx"
    file_b = tmp_path / "b.docx"
    output_dir = tmp_path / "out"

    for path in (file_a, file_b):
        doc = Document()
        doc.add_paragraph("Provide ASCE 7-16 bracing.")
        doc.save(path)

    spec_a = extract_text(file_a)
    spec_a.filename = "a.docx"
    spec_b = extract_text(file_b)
    spec_b.filename = "b.docx"

    grouped = _make_finding(
        file_name="a.docx",
        existing="ASCE 7-16",
        replacement="ASCE 7-22",
    )
    grouped.affected_files = ["a.docx", "b.docx"]
    from src.verifier import VerificationResult as _VR
    grouped.verification = _VR(verdict="CONFIRMED", explanation="", sources=[])

    reports = execute_edit_plan(
        selected_finding_indices=[0],
        all_findings=[grouped],
        cross_check_findings=[],
        extracted_specs=[spec_a, spec_b],
        source_paths=[file_a, file_b],
        output_dir=output_dir,
    )

    assert len(reports) == 2
    edited_paths = {r.output_path for r in reports}
    assert any(p.name.startswith("a_edited") for p in edited_paths)
    assert any(p.name.startswith("b_edited") for p in edited_paths)
    for r in reports:
        assert r.edits_applied == 1
        saved = Document(r.output_path)
        assert saved.paragraphs[0].text == "Provide ASCE 7-22 bracing."


# ---------------------------------------------------------------------------
# Issue 4 — Block ambiguous edit auto-apply
# ---------------------------------------------------------------------------


def test_ambiguous_locator_result_does_not_produce_edit_action():
    paragraph_map = [
        _mapping("Use non-shrink grout at equipment bases.", idx=0),
        _mapping("Use non-shrink grout at equipment bases for support pads.", idx=1),
    ]
    finding = _make_finding(existing="Use non-shrink grout at equipment bases")

    result = locate_edit(finding, paragraph_map)
    assert result.status == "ambiguous"

    actions = build_edit_actions([result])
    assert actions == []
    assert result.warning is not None
    assert "manually" in result.warning.lower() or "manual" in result.warning.lower()


# ---------------------------------------------------------------------------
# Issue 5 — Explicit ADD anchor model
# ---------------------------------------------------------------------------


def test_parse_findings_extracts_anchor_text_and_position():
    items = [
        {
            "severity": "HIGH",
            "fileName": "spec.docx",
            "section": "1.0",
            "issue": "Missing seismic note",
            "actionType": "ADD",
            "existingText": None,
            "replacementText": "Provide bracing per ASCE 7-22.",
            "codeReference": "ASCE 7-22",
            "confidence": 0.9,
            "anchorText": "PART 2 - PRODUCTS",
            "insertPosition": "AFTER",
        }
    ]
    findings = _parse_findings(items)
    assert len(findings) == 1
    f = findings[0]
    assert f.anchorText == "PART 2 - PRODUCTS"
    assert f.insertPosition == "after"


def test_parse_findings_drops_invalid_insert_position():
    items = [
        {
            "severity": "HIGH",
            "fileName": "spec.docx",
            "section": "1.0",
            "issue": "issue",
            "actionType": "ADD",
            "existingText": None,
            "replacementText": "x",
            "codeReference": "x",
            "confidence": 0.5,
            "anchorText": "anchor",
            "insertPosition": "sideways",
        }
    ]
    findings = _parse_findings(items)
    assert findings[0].insertPosition is None


def test_add_action_with_anchor_text_locates_anchor(tmp_path: Path):
    paragraph_map = [
        _mapping("PART 2 - PRODUCTS", idx=0),
        _mapping("Existing requirement.", idx=1),
    ]
    finding = Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="2",
        issue="Add seismic clause",
        actionType="ADD",
        existingText=None,
        replacementText="Provide bracing per ASCE 7-22.",
        codeReference="ASCE 7-22",
        confidence=0.9,
        anchorText="PART 2 - PRODUCTS",
        insertPosition="after",
    )
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.locations[0].mapping.body_index == 0


def test_finding_anchor_fields_round_trip_through_resume_state():
    finding = Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="2",
        issue="x",
        actionType="ADD",
        existingText=None,
        replacementText="new",
        codeReference="x",
        confidence=0.9,
        anchorText="ANCHOR PARAGRAPH",
        insertPosition="before",
    )
    restored = deserialize_finding(serialize_finding(finding))
    assert restored.anchorText == "ANCHOR PARAGRAPH"
    assert restored.insertPosition == "before"


# ---------------------------------------------------------------------------
# Issue 6 — DELETE earlier in document does not corrupt later ADD insertion
# ---------------------------------------------------------------------------


def test_delete_followed_by_add_inserts_at_correct_anchor(tmp_path: Path):
    source = tmp_path / "src.docx"
    output = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("Paragraph A — keep")
    doc.add_paragraph("Paragraph B — delete me")
    doc.add_paragraph("Paragraph C — anchor for ADD")
    doc.save(source)

    spec = extract_text(source)
    paragraph_map = spec.paragraph_map
    assert paragraph_map is not None
    anchor_mapping = next(m for m in paragraph_map if m.text.startswith("Paragraph C"))
    delete_mapping = next(m for m in paragraph_map if m.text.startswith("Paragraph B"))

    delete_finding = Finding(
        severity="HIGH",
        fileName="src.docx",
        section="1",
        issue="drop B",
        actionType="DELETE",
        existingText=delete_mapping.text,
        replacementText=None,
        codeReference="x",
        confidence=0.95,
    )
    delete_locator = LocatorResult(
        finding=delete_finding,
        status="matched",
        locations=[
            EditLocation(
                mapping=delete_mapping,
                match_start=0,
                match_end=len(delete_mapping.text),
                matched_text=delete_mapping.text,
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text=None,
        action_type="DELETE",
        warning=None,
    )

    add_finding = Finding(
        severity="HIGH",
        fileName="src.docx",
        section="1",
        issue="add D",
        actionType="ADD",
        existingText=None,
        replacementText="Paragraph D — inserted after C",
        codeReference="x",
        confidence=0.95,
        anchorText=anchor_mapping.text,
        insertPosition="after",
    )
    add_locator = LocatorResult(
        finding=add_finding,
        status="matched",
        locations=[
            EditLocation(
                mapping=anchor_mapping,
                match_start=0,
                match_end=len(anchor_mapping.text),
                matched_text=anchor_mapping.text,
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="Paragraph D — inserted after C",
        action_type="ADD",
        warning=None,
    )

    actions = build_edit_actions([delete_locator, add_locator])
    assert len(actions) == 2

    from src.spec_editor import apply_edits_to_spec

    report = apply_edits_to_spec(source, output, actions)
    assert report.edits_applied == 2
    assert report.edits_failed == 0

    saved = Document(output)
    texts = [p.text for p in saved.paragraphs]
    # Paragraph B should be gone; D should immediately follow C.
    assert "Paragraph B — delete me" not in texts
    c_idx = texts.index("Paragraph C — anchor for ADD")
    assert texts[c_idx + 1] == "Paragraph D — inserted after C"


# ---------------------------------------------------------------------------
# Issue 7 — Errored batch results surface as truncated specs and skip cross-check
# ---------------------------------------------------------------------------


def _submission_with_two_specs(specs: list[ExtractedSpec]) -> BatchSubmission:
    return BatchSubmission(
        job=BatchJob(
            batch_id="msgbatch_test_audit",
            job_type="review",
            request_map={
                "review__spec_a__0": {"filename": "a.docx", "index": 0, "type": "review"},
                "review__spec_b__1": {"filename": "b.docx", "index": 1, "type": "review"},
            },
            created_at=1.0,
        ),
        files_reviewed=["a.docx", "b.docx"],
        review_request_ids=["review__spec_a__0", "review__spec_b__1"],
        cycle_label="2025",
        cross_check_enabled=True,
        prepared_specs=specs,
    )


def test_batch_errored_request_is_recorded_as_truncated_spec(monkeypatch: pytest.MonkeyPatch):
    spec_a = ExtractedSpec(filename="a.docx", content="A", word_count=1, source_path="/tmp/a.docx", source_format="docx")
    spec_b = ExtractedSpec(filename="b.docx", content="B", word_count=1, source_path="/tmp/b.docx", source_format="docx")
    submission = _submission_with_two_specs([spec_a, spec_b])

    review_results = {
        "review__spec_a__0": ReviewResult(findings=[_make_finding(file_name="a.docx", issue="OK")], parse_status="ok"),
        "review__spec_b__1": ReviewResult(findings=[], error="Batch request errored: server overload"),
    }
    monkeypatch.setattr(
        "src.pipeline.retrieve_review_results", lambda _job, *, model: review_results
    )
    monkeypatch.setattr(
        "src.pipeline._recover_retryable_review_batch_results",
        lambda _submission, results, *, log: results,
    )

    state = collect_review_batch_results(submission)
    assert "b.docx" in state.truncated_specs
    assert state.review_result.error is not None


def test_run_cross_check_for_batch_skips_failed_specs(monkeypatch: pytest.MonkeyPatch):
    spec_a = ExtractedSpec(filename="a.docx", content="A", word_count=1, source_path="/tmp/a.docx", source_format="docx")
    spec_b = ExtractedSpec(filename="b.docx", content="B", word_count=1, source_path="/tmp/b.docx", source_format="docx")
    submission = _submission_with_two_specs([spec_a, spec_b])
    state = CollectedBatchState(
        submission=submission,
        review_result=ReviewResult(findings=[_make_finding(file_name="a.docx")]),
        truncated_specs=["b.docx"],
    )

    captured: dict[str, Any] = {}

    def fake_run_cross_check(specs, findings, **kwargs):
        captured["spec_filenames"] = [s.filename for s in specs]
        return ReviewResult(findings=[], cross_check_status="completed")

    # Pipeline now invokes the chunked wrapper, which delegates to
    # ``run_cross_check`` when the input fits. Patch the chunked entry
    # point so the test still observes which specs were forwarded.
    monkeypatch.setattr("src.pipeline.run_chunked_cross_check", fake_run_cross_check)

    run_cross_check_for_batch(state, specs=[spec_a, spec_b])

    assert captured["spec_filenames"] == ["a.docx"]


# ---------------------------------------------------------------------------
# Issue 8 — Continuation content stored as plain dicts
# ---------------------------------------------------------------------------


def test_content_block_to_plain_converts_text_block_to_dict():
    from anthropic.types import TextBlock

    block = TextBlock(type="text", text="hello", citations=None)
    plain = _content_block_to_plain(block)
    assert isinstance(plain, dict)
    assert plain["type"] == "text"
    assert plain["text"] == "hello"


def test_classify_wave_results_stores_plain_dicts_for_pause_turn(monkeypatch: pytest.MonkeyPatch):
    finding = _make_finding()

    class _PauseMessage:
        stop_reason = "pause_turn"

        def __init__(self):
            from anthropic.types import TextBlock
            self.content = [TextBlock(type="text", text="thinking...", citations=None)]

    class _BatchEntry:
        class _Result:
            type = "succeeded"
            message = _PauseMessage()

        result = _Result()

    monkeypatch.setattr(
        "src.verifier.retrieve_verification_results_detailed",
        lambda _job: {"verify__0": _BatchEntry()},
    )

    job = BatchJob(
        batch_id="msgbatch_pause_audit",
        job_type="verify",
        request_map={"verify__0": {"finding_idx": 0}},
        created_at=1.0,
    )
    contexts = {"verify__0": {"finding_idx": 0, "original_prompt": "p"}}
    outcomes = _classify_wave_results(job=job, findings=[finding], request_contexts=contexts)
    assert len(outcomes) == 1
    blocks = outcomes[0].assistant_content_blocks
    assert blocks is not None
    assert all(isinstance(b, dict) for b in blocks)
    assert blocks[0]["type"] == "text"


# ---------------------------------------------------------------------------
# Issue 9 — Final retry does not sleep and exception detail is preserved
# ---------------------------------------------------------------------------


def test_stream_review_final_retry_does_not_sleep_and_preserves_error(monkeypatch: pytest.MonkeyPatch):
    from src.reviewer import _stream_review

    sleeps: list[float] = []
    monkeypatch.setattr("src.reviewer.time.sleep", lambda s: sleeps.append(s))

    class _ConnReset(Exception):
        pass

    def _raise(*_args, **_kwargs):
        # The retry loop's generic-exception branch treats messages
        # containing "connection reset" as transient and retryable.
        raise _ConnReset("connection reset by peer (test)")

    class _FakeMessages:
        def stream(self, **kwargs):
            _raise()

    class _FakeClient:
        messages = _FakeMessages()

    result = _stream_review(_FakeClient(), "sys", "user", max_retries=2)
    # Two attempts: one sleep between them, none after the final attempt.
    assert len(sleeps) == 1
    assert result.error is not None
    # Exception detail must survive instead of being swallowed by a generic
    # "Failed after N attempts." message (audit Issue 9).
    assert "connection reset" in result.error.lower()


# ---------------------------------------------------------------------------
# Issue 10 — Extractor invariant raises ValueError, not AssertionError
# ---------------------------------------------------------------------------


def test_extractor_invariant_violation_raises_value_error(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    source = tmp_path / "broken.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(source)

    real_mapping_class = ParagraphMapping
    call_state = {"count": 0}

    def _broken_paragraph_mapping(*args, **kwargs):
        call_state["count"] += 1
        if call_state["count"] == 1:
            kwargs.setdefault("text", "")
            kwargs["text"] = "INTENTIONALLY MISMATCHED"
        return real_mapping_class(*args, **kwargs)

    monkeypatch.setattr("src.extractor.ParagraphMapping", _broken_paragraph_mapping)

    with pytest.raises(ValueError) as exc_info:
        extract_text(source)
    assert "does not reconstruct" in str(exc_info.value)
