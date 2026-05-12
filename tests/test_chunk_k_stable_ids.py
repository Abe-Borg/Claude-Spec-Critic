"""Chunk K tests: stable document IDs for deterministic finding/edit targeting.

Plan section: "Chunk K — Stable Document IDs for Deterministic Finding/Edit
Targeting". The chunk has five sub-phases:

* K1 — extraction adds stable element ids + document ids on every spec.
* K2 — prompt builders render spec content with the ids visible to the model.
* K3 — structured finding schema accepts an optional ``evidenceElementId``.
* K4 — the locator prefers ids when supplied, validates exact-text quotes
  against the live element, and falls back to text-only matching only when
  no id is available.
* K5 — diagnostics surface how often each locator method was used.

These tests assert each sub-phase end-to-end and prove the listed
acceptance criteria from the plan (id generation, prompt includes ids,
model-like output with id locates the correct paragraph, duplicate text
is disambiguated, invalid id requires manual review, IDs are
deterministic / debuggable, backward compatibility, resume round-trip).
"""

from __future__ import annotations

import json
import os
from pathlib import Path

import pytest

from src.code_cycles import DEFAULT_CYCLE
from src.diagnostics import DiagnosticsReport
from src.edit_locator import _id_anchored_match, locate_edit
from src.extractor import (
    ParagraphMapping,
    _derive_document_id,
    _is_heading_paragraph,
    extract_text_from_docx,
)
from src.prompt_serialization import (
    TAG_HEADING,
    TAG_PARA,
    TAG_ROW,
    element_ids_enabled,
    render_spec_with_ids,
)
from src.prompts import get_single_spec_user_message
from src.resume_state import deserialize_finding, serialize_finding
from src.reviewer import Finding, _parse_findings
from src.structured_schemas import _FINDING_OBJECT_SCHEMA
from tests.fixtures.docx_fixtures import make_paragraph_spec, make_table_spec


# ---------------------------------------------------------------------------
# K1 — Extraction ID Model
# ---------------------------------------------------------------------------


class TestK1ExtractionIdModel:
    def test_paragraph_mappings_get_stable_element_ids(self, tmp_path: Path):
        path = make_paragraph_spec(
            tmp_path,
            [
                "SECTION 23 05 23 - HVAC",
                "PART 1 GENERAL",
                "1.01 SUMMARY",
                "A. Install copper piping with Type L wall thickness.",
            ],
        )
        spec = extract_text_from_docx(path)

        ids = [m.element_id for m in spec.paragraph_map or []]
        # Body paragraphs use ``p<body_index>``. The first four body
        # elements in the docx are paragraphs, so ids are p0..p3.
        assert ids == ["p0", "p1", "p2", "p3"]

    def test_table_cells_get_table_row_ids(self, tmp_path: Path):
        path = make_table_spec(
            tmp_path,
            rows=[["Tag", "Value"], ["P-1", "100 gpm"]],
        )
        spec = extract_text_from_docx(path)

        cell_ids = [
            m.element_id for m in spec.paragraph_map or [] if m.element_type == "table_cell"
        ]
        assert cell_ids == ["t0r0", "t0r1"]

    def test_section_id_attribution_follows_heading_paragraphs(self, tmp_path: Path):
        path = make_paragraph_spec(
            tmp_path,
            [
                "PART 1 GENERAL",
                "1.01 SUMMARY",
                "A. Body paragraph one.",
                "1.02 SCOPE",
                "A. Body paragraph two.",
            ],
        )
        spec = extract_text_from_docx(path)
        by_id = {m.element_id: m for m in spec.paragraph_map or []}
        # The body paragraphs nest under the nearest heading. The heading
        # paragraph itself owns its own section_id (it is its own heading).
        assert by_id["p2"].section_id == "1.01 SUMMARY"
        assert by_id["p4"].section_id == "1.02 SCOPE"

    def test_document_id_is_filename_stem(self, tmp_path: Path):
        path = make_paragraph_spec(
            tmp_path,
            ["PART 1 GENERAL", "Body."],
            filename="23 05 23 - HVAC.docx",
        )
        spec = extract_text_from_docx(path)
        assert spec.document_id == "23 05 23 - HVAC"

    def test_element_ids_are_deterministic_across_runs(self, tmp_path: Path):
        # K1 directive 4: ids should not be Python object memory ids. The
        # same document extracted twice should produce byte-identical ids.
        path = make_paragraph_spec(tmp_path, ["A.", "B.", "C."])
        spec_a = extract_text_from_docx(path)
        spec_b = extract_text_from_docx(path)
        ids_a = [m.element_id for m in spec_a.paragraph_map or []]
        ids_b = [m.element_id for m in spec_b.paragraph_map or []]
        assert ids_a == ids_b

    def test_derive_document_id_handles_empty(self):
        assert _derive_document_id("") == ""
        assert _derive_document_id("noext") == "noext"
        assert _derive_document_id("hvac.docx") == "hvac"

    def test_heading_detection_heuristic(self):
        assert _is_heading_paragraph("PART 1 GENERAL")
        assert _is_heading_paragraph("SECTION 23 05 23")
        assert _is_heading_paragraph("1.01 SUMMARY")
        assert _is_heading_paragraph("2.3.A Hot water")
        assert not _is_heading_paragraph(
            "A. Body paragraph longer than the heading length limit "
            "definitely no longer a heading."
        )
        assert not _is_heading_paragraph("")

    def test_header_footer_paragraphs_get_distinct_ids(self, tmp_path: Path):
        from docx import Document

        path = tmp_path / "with_header.docx"
        doc = Document()
        doc.add_paragraph("Body paragraph.")
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Header Top"
        section.footer.paragraphs[0].text = "Footer Bottom"
        doc.save(path)

        spec = extract_text_from_docx(path)
        ids = {m.element_id for m in spec.paragraph_map or [] if m.element_id}
        # The synthetic meta delimiter plus the header/footer entries
        # all get their own ids that don't collide with body paragraphs.
        assert "meta:hf" in ids
        assert any(eid.startswith("s0h") for eid in ids)
        assert any(eid.startswith("s0f") for eid in ids)


# ---------------------------------------------------------------------------
# K2 — Prompt Serialization With IDs
# ---------------------------------------------------------------------------


class TestK2PromptSerializationWithIds:
    def test_id_tagged_rendering_emits_one_wrapper_per_element(self, tmp_path: Path):
        path = make_paragraph_spec(
            tmp_path,
            ["PART 1 GENERAL", "1.01 SUMMARY", "A. Body line."],
        )
        spec = extract_text_from_docx(path)
        rendered = render_spec_with_ids(
            spec.content, spec.paragraph_map, filename=spec.filename,
        )
        assert f'<{TAG_HEADING} id="p0">PART 1 GENERAL</{TAG_HEADING}>' in rendered
        assert f'<{TAG_HEADING} id="p1">1.01 SUMMARY</{TAG_HEADING}>' in rendered
        # Body paragraphs carry the owning section attribute.
        assert (
            f'<{TAG_PARA} id="p2" section="1.01 SUMMARY">A. Body line.</{TAG_PARA}>'
            in rendered
        )

    def test_id_tagged_rendering_uses_row_tag_for_tables(self, tmp_path: Path):
        path = make_table_spec(
            tmp_path,
            rows=[["Tag", "Value"], ["P-1", "100 gpm"]],
        )
        spec = extract_text_from_docx(path)
        rendered = render_spec_with_ids(spec.content, spec.paragraph_map, filename=spec.filename)
        assert f'<{TAG_ROW} id="t0r0"' in rendered
        assert f'<{TAG_ROW} id="t0r1"' in rendered

    def test_user_message_includes_id_hint_when_map_supplied(self, tmp_path: Path):
        path = make_paragraph_spec(tmp_path, ["PART 1 GENERAL", "A. Body."])
        spec = extract_text_from_docx(path)
        msg = get_single_spec_user_message(
            spec.content,
            spec.filename,
            cycle=DEFAULT_CYCLE,
            paragraph_map=spec.paragraph_map,
        )
        # The id hint line is the K2 signal that the model should cite ids.
        assert "evidenceElementId" in msg
        assert "<para id=" in msg or "<heading id=" in msg

    def test_user_message_legacy_path_omits_id_hint_when_no_map(self, tmp_path: Path):
        # Legacy callers that hand the function a raw string keep the
        # pre-Chunk-K rendering — no id wrappers, no id hint.
        msg = get_single_spec_user_message(
            "PART 1 GENERAL\n\nA. Body.",
            "spec.docx",
            cycle=DEFAULT_CYCLE,
        )
        assert "evidenceElementId" not in msg
        assert "<para id=" not in msg

    def test_env_toggle_reverts_to_legacy_rendering(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setenv("SPEC_CRITIC_ELEMENT_IDS", "0")
        assert element_ids_enabled() is False
        path = make_paragraph_spec(tmp_path, ["PART 1 GENERAL", "A. Body."])
        spec = extract_text_from_docx(path)
        msg = get_single_spec_user_message(
            spec.content,
            spec.filename,
            cycle=DEFAULT_CYCLE,
            paragraph_map=spec.paragraph_map,
        )
        # With the env flag off, the wrapper falls back to <spec> only.
        assert "evidenceElementId" not in msg
        assert "<para id=" not in msg

    def test_render_spec_with_ids_falls_back_when_no_map(self):
        # Defensive path: a caller that has the body text but no map (e.g.
        # legacy resume payloads) must still get a syntactically valid
        # ``<spec>`` block instead of an empty body.
        rendered = render_spec_with_ids("Body only.", None, filename="x.docx")
        assert rendered.startswith("<spec filename=\"x.docx\">")
        assert "Body only." in rendered
        assert rendered.endswith("</spec>")

    def test_render_spec_escapes_evil_filename(self, tmp_path: Path):
        # Chunk G regression — filename with reserved chars must not break
        # the wrapper when the K2 path renders it.
        spec_map = [
            ParagraphMapping(
                body_index=0, element_type="paragraph", text="Body.",
                table_index=None, row_index=None, cell_index=None,
                element_id="p0", section_id="",
            )
        ]
        rendered = render_spec_with_ids(
            "Body.", spec_map, filename='evil".docx',
        )
        # Quote characters must be escaped so the opening tag stays intact.
        assert "&quot;" in rendered or "\\\"" in rendered or 'filename="evil"' not in rendered

    def test_system_prompt_unchanged_after_chunk_k(self, tmp_path: Path):
        # Cache-prefix safety: the system prompt is the prompt-cache
        # breakpoint. K2 must not move bytes here, or every run pays the
        # cache-write cost on first call without recouping it later.
        from src.prompts import get_system_prompt

        prompt = get_system_prompt(DEFAULT_CYCLE)
        # The system prompt instructs the model about the review tool
        # but should NOT mention the K2 id wrappers — those live in
        # the per-request user message so the cached prefix stays
        # stable across runs.
        assert "evidenceElementId" not in prompt
        assert "<para id=" not in prompt


# ---------------------------------------------------------------------------
# K3 — Schema Update for Evidence IDs
# ---------------------------------------------------------------------------


class TestK3SchemaUpdateForEvidenceIds:
    def test_schema_declares_evidence_field(self):
        assert "evidenceElementId" in _FINDING_OBJECT_SCHEMA["required"]
        prop = _FINDING_OBJECT_SCHEMA["properties"]["evidenceElementId"]
        # Nullable string so strict-mode constrained sampling still has a
        # deterministic shape but legacy "no id" findings remain valid.
        assert prop["type"] == ["string", "null"]

    def test_parser_accepts_evidence_id(self):
        findings = _parse_findings(
            [
                {
                    "severity": "HIGH",
                    "fileName": "a.docx",
                    "section": "1.01",
                    "issue": "X",
                    "actionType": "EDIT",
                    "existingText": "old text",
                    "replacementText": "new text",
                    "codeReference": None,
                    "confidence": 0.9,
                    "anchorText": None,
                    "insertPosition": None,
                    "evidenceElementId": "p17",
                }
            ]
        )
        assert findings[0].evidenceElementId == "p17"

    def test_parser_normalizes_empty_string_to_none(self):
        findings = _parse_findings(
            [
                {
                    "severity": "HIGH",
                    "fileName": "a.docx",
                    "section": "1.01",
                    "issue": "X",
                    "actionType": "EDIT",
                    "existingText": "old",
                    "replacementText": "new",
                    "codeReference": None,
                    "confidence": 0.9,
                    "evidenceElementId": "  ",
                }
            ]
        )
        assert findings[0].evidenceElementId is None

    def test_parser_legacy_payload_without_field_loads_as_none(self):
        findings = _parse_findings(
            [
                {
                    "severity": "HIGH",
                    "fileName": "a.docx",
                    "section": "1.01",
                    "issue": "X",
                    "actionType": "EDIT",
                    "existingText": "old",
                    "replacementText": "new",
                    "codeReference": None,
                    "confidence": 0.9,
                }
            ]
        )
        assert findings[0].evidenceElementId is None

    def test_resume_serializer_roundtrips_evidence_id(self):
        f = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1.01",
            issue="X",
            actionType="EDIT",
            existingText="old",
            replacementText="new",
            codeReference=None,
            confidence=0.9,
            evidenceElementId="t2r3",
        )
        payload = serialize_finding(f)
        restored = deserialize_finding(payload)
        assert restored.evidenceElementId == "t2r3"

    def test_resume_serializer_handles_legacy_payload(self):
        legacy = {
            "severity": "HIGH",
            "fileName": "a.docx",
            "section": "1.01",
            "issue": "X",
            "actionType": "EDIT",
            "existingText": None,
            "replacementText": None,
            "codeReference": None,
            "confidence": 0.9,
        }
        restored = deserialize_finding(legacy)
        assert restored.evidenceElementId is None

    def test_resume_payload_is_json_round_trippable(self):
        # Resume payloads are written to disk as JSON; the new field must
        # survive the round trip without type drift.
        f = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1.01",
            issue="X",
            actionType="EDIT",
            existingText="old",
            replacementText="new",
            codeReference=None,
            confidence=0.9,
            evidenceElementId="p17",
        )
        payload = serialize_finding(f)
        text = json.dumps(payload)
        restored = deserialize_finding(json.loads(text))
        assert restored.evidenceElementId == "p17"


# ---------------------------------------------------------------------------
# K4 — Locator Prefers IDs
# ---------------------------------------------------------------------------


def _pm(text: str, *, idx: int, eid: str, section: str = "") -> ParagraphMapping:
    return ParagraphMapping(
        body_index=idx,
        element_type="paragraph",
        text=text,
        table_index=None,
        row_index=None,
        cell_index=None,
        element_id=eid,
        section_id=section,
    )


def _table_pm(text: str, *, idx: int, eid: str, row: int = 0) -> ParagraphMapping:
    return ParagraphMapping(
        body_index=idx,
        element_type="table_cell",
        text=text,
        table_index=0,
        row_index=row,
        cell_index=None,
        element_id=eid,
        section_id="",
    )


def _finding(
    *,
    existing: str | None = "old",
    replacement: str | None = "new",
    evidence_id: str | None = None,
    action: str = "EDIT",
    anchor: str | None = None,
) -> Finding:
    return Finding(
        severity="HIGH",
        fileName="a.docx",
        section="1.01",
        issue="x",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference=None,
        confidence=0.9,
        evidenceElementId=evidence_id,
        anchorText=anchor,
    )


class TestK4LocatorPrefersIds:
    def test_id_locates_unique_paragraph(self):
        pm = [_pm("Body line.", idx=0, eid="p0")]
        f = _finding(existing="Body line.", evidence_id="p0")
        result = locate_edit(f, pm)
        assert result.status == "matched"
        assert result.locations[0].match_method == "id"
        assert result.locations[0].mapping.element_id == "p0"

    def test_id_disambiguates_duplicate_exact_text(self):
        pm = [
            _pm("Same text here.", idx=0, eid="p0", section="1.0 SUMMARY"),
            _pm("Same text here.", idx=1, eid="p1", section="2.0 PRODUCTS"),
        ]
        # Without id this is ambiguous (verified separately below). With
        # id="p1" the locator binds to the second paragraph deterministically.
        f = _finding(existing="Same text here.", evidence_id="p1")
        result = locate_edit(f, pm)
        assert result.status == "matched"
        assert result.locations[0].mapping.body_index == 1
        assert result.locations[0].match_method == "id"
        assert result.safety_category == "AUTO_SAFE"

    def test_without_id_duplicate_text_remains_ambiguous(self):
        # Negative control for the above: prove the id is what fixes
        # the ambiguity, not some other change in the locator.
        pm = [
            _pm("Same text here.", idx=0, eid="p0"),
            _pm("Same text here.", idx=1, eid="p1"),
        ]
        f = _finding(existing="Same text here.", evidence_id=None)
        result = locate_edit(f, pm)
        assert result.status == "ambiguous"

    def test_id_falls_back_to_normalized_quote(self):
        # Real prompts sometimes pad whitespace inconsistently. The id
        # path should still match if the live element contains the
        # normalized form of the recorded quote.
        pm = [_pm("Install  COPPER  piping.", idx=0, eid="p0")]
        f = _finding(existing="install copper piping.", evidence_id="p0")
        result = locate_edit(f, pm)
        assert result.status == "matched"
        assert result.locations[0].match_method == "id"

    def test_invalid_id_does_not_silently_fall_back(self):
        # K4 contract: if the id is set but unusable, do not text-match
        # against the rest of the document. The model named a specific
        # element; a "similar text" match somewhere else is almost
        # certainly the wrong target.
        pm = [_pm("Body line.", idx=0, eid="p0")]
        f = _finding(existing="Body line.", evidence_id="p99")
        result = locate_edit(f, pm)
        assert result.status == "not_found"
        assert result.safety_category == "MANUAL_REVIEW"
        assert "p99" in (result.warning or "")

    def test_id_set_but_quote_missing_requires_manual_review(self):
        pm = [_pm("Live text has shifted.", idx=0, eid="p0")]
        f = _finding(existing="Text that is not present.", evidence_id="p0")
        result = locate_edit(f, pm)
        assert result.status == "not_found"
        assert result.safety_category == "MANUAL_REVIEW"

    def test_id_for_table_cell_is_auto_with_caution(self):
        pm = [_table_pm("P-1 | 100 gpm", idx=0, eid="t0r0", row=0)]
        f = _finding(existing="100 gpm", evidence_id="t0r0")
        result = locate_edit(f, pm)
        assert result.status == "matched"
        assert result.locations[0].match_method == "id"
        # Table cells stay in caution territory so the table-cell
        # precondition revalidation in spec_editor still gates the
        # actual mutation.
        assert result.safety_category == "AUTO_WITH_CAUTION"

    def test_legacy_finding_without_id_uses_text_matching(self):
        # Backward compatibility: a finding with no evidenceElementId
        # should match exactly the same way it did pre-Chunk K.
        pm = [_pm("Body line.", idx=0, eid="p0")]
        f = _finding(existing="Body line.", evidence_id=None)
        result = locate_edit(f, pm)
        assert result.status == "matched"
        assert result.locations[0].match_method == "exact"

    def test_id_anchored_match_returns_warning_for_missing_id(self):
        # Direct test of the helper so the public locate_edit doesn't
        # need to be the only path that exercises this branch.
        pm = [_pm("Body.", idx=0, eid="p0")]
        f = _finding(existing="Body.", evidence_id="p_does_not_exist")
        locations, warning = _id_anchored_match(f, "Body.", pm)
        assert locations == []
        assert warning is not None
        assert "p_does_not_exist" in warning

    def test_id_with_no_existing_text_takes_whole_paragraph(self):
        # ADD actions can carry an id but no existingText — the id alone
        # identifies the anchor paragraph and the whole paragraph span
        # is returned so spec_editor can place new content adjacent to it.
        pm = [_pm("Anchor paragraph text.", idx=0, eid="p0")]
        f = _finding(existing=None, evidence_id="p0", action="ADD", anchor=None)
        locations, warning = _id_anchored_match(f, "", pm)
        assert warning is None
        assert len(locations) == 1
        assert locations[0].match_start == 0
        assert locations[0].match_end == len(pm[0].text)


# ---------------------------------------------------------------------------
# K5 — Reports Include IDs Internally
# ---------------------------------------------------------------------------


class TestK5ReportsIncludeIds:
    def test_diagnostics_records_locator_methods(self):
        d = DiagnosticsReport()
        d.record_locator_method("id")
        d.record_locator_method("id")
        d.record_locator_method("exact")
        d.record_locator_method("fuzzy")
        summary = d.summary()
        assert summary["locator_methods"] == {"id": 2, "exact": 1, "fuzzy": 1}

    def test_diagnostics_text_omits_locator_methods_when_empty(self):
        d = DiagnosticsReport()
        text = d.to_text()
        assert "Locator Methods" not in text

    def test_diagnostics_text_surfaces_locator_methods_when_present(self):
        d = DiagnosticsReport()
        d.record_locator_method("id")
        text = d.to_text()
        assert "Locator Methods" in text

    def test_execute_edit_plan_threads_diagnostics(self, tmp_path: Path):
        # Light integration test: build a one-paragraph spec, ask for an
        # id-anchored edit, and confirm the diagnostics counter is
        # incremented under the ``id`` bucket.
        from src.apply_edits import execute_edit_plan

        source = make_paragraph_spec(
            tmp_path,
            ["PART 1 GENERAL", "A. Hot water at 110 degrees."],
            filename="hvac.docx",
        )
        spec = extract_text_from_docx(source)
        # The body paragraph for "A. Hot water…" is p1 — heading is p0.
        body_pm = next(m for m in spec.paragraph_map if "Hot water" in m.text)

        f = Finding(
            severity="HIGH",
            fileName="hvac.docx",
            section="1.01",
            issue="Wrong temperature",
            actionType="EDIT",
            existingText="110 degrees",
            replacementText="120 degrees",
            codeReference=None,
            confidence=0.9,
            evidenceElementId=body_pm.element_id,
        )
        diagnostics = DiagnosticsReport()
        execute_edit_plan(
            selected_finding_indices=[0],
            all_findings=[f],
            cross_check_findings=[],
            extracted_specs=[spec],
            source_paths=[source],
            output_dir=tmp_path / "out",
            diagnostics=diagnostics,
        )
        assert diagnostics.locator_methods.get("id", 0) == 1


# ---------------------------------------------------------------------------
# End-to-end integration check
# ---------------------------------------------------------------------------


class TestChunkKEndToEnd:
    def test_extract_render_parse_locate(self, tmp_path: Path):
        # The full Chunk K loop: extract → render with ids → simulate a
        # model output that cites an id → parse → locate. Proves K1..K4
        # compose without integration glue.
        path = make_paragraph_spec(
            tmp_path,
            [
                "SECTION 23 05 23 - HVAC",
                "PART 1 GENERAL",
                "1.01 SUMMARY",
                "A. Hot water at 110 degrees.",
                "1.02 SCOPE",
                "A. Hot water at 110 degrees.",  # duplicate text on purpose
            ],
            filename="hvac.docx",
        )
        spec = extract_text_from_docx(path)
        rendered = render_spec_with_ids(
            spec.content, spec.paragraph_map, filename=spec.filename,
        )

        # Find the id of the duplicate body paragraph in section 1.02.
        target_mapping = next(
            m for m in spec.paragraph_map
            if m.element_type == "paragraph"
            and m.section_id == "1.02 SCOPE"
            and "Hot water" in m.text
        )
        assert target_mapping.element_id

        # Confirm the rendered spec body actually includes that id and
        # the exact body text the model would quote.
        assert f'id="{target_mapping.element_id}"' in rendered
        assert "Hot water at 110 degrees." in rendered

        # Simulate the model emitting a structured finding that cites
        # the id alongside the exact quote.
        findings = _parse_findings(
            [
                {
                    "severity": "HIGH",
                    "fileName": spec.filename,
                    "section": "1.02 SCOPE",
                    "issue": "Temperature too low.",
                    "actionType": "EDIT",
                    "existingText": "110 degrees",
                    "replacementText": "120 degrees",
                    "codeReference": None,
                    "confidence": 0.95,
                    "anchorText": None,
                    "insertPosition": None,
                    "evidenceElementId": target_mapping.element_id,
                }
            ]
        )
        assert findings[0].evidenceElementId == target_mapping.element_id

        # The locator should bind to the section-1.02 paragraph
        # specifically — without the id this would be ambiguous.
        result = locate_edit(findings[0], spec.paragraph_map)
        assert result.status == "matched"
        assert result.locations[0].match_method == "id"
        assert result.locations[0].mapping.element_id == target_mapping.element_id
