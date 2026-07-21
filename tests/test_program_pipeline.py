"""Composite hyperscale program orchestration and output contracts."""
from __future__ import annotations

import threading
import time
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from src.batch.batch import BatchJob, BatchStatus
from src.gui import batch_controller as gui_batch
from src.modules import require_module
from src.orchestration import program_pipeline as pp
from src.orchestration.batch_resume import (
    PendingBatch,
    PendingProgramRun,
    load_pending_run,
    save_pending_program_run,
)
from src.orchestration.pipeline import BatchSubmission, PipelineResult
from src.output.edit_sidecar import build_edit_instructions
from src.output.report_exporter import export_report
from src.programs import (
    HYPERSCALE_DATACENTER_PROGRAM,
    RoutingState,
    SpecAssignment,
    SpecRoutingDecision,
    resolve_saved_program,
)
from src.review.reviewer import Finding, ReviewResult
from src.tracing import current_span


def _assignment(name: str, module_ids: tuple[str, ...]) -> SpecAssignment:
    state = RoutingState.SUPPORTED if module_ids else RoutingState.UNSUPPORTED
    return SpecAssignment(
        source_path=str(Path("C:/specs") / name),
        decision=SpecRoutingDecision(
            spec_id=name,
            program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
            automatic_state=state,
            automatic_module_ids=module_ids,
            confidence=0.95,
            evidence=(),
        ),
    )


def _submission(
    module_id: str,
    name: str,
    *,
    project_context: str = "",
) -> BatchSubmission:
    module = require_module(module_id)
    request_id = f"review__{module_id}__0"
    return BatchSubmission(
        job=BatchJob(
            batch_id=f"msgbatch_{module_id}",
            job_type="review",
            request_map={
                request_id: {"filename": name, "index": 0, "type": "review"}
            },
            created_at=1_700_000_000.0,
        ),
        files_reviewed=[name],
        review_request_ids=[request_id],
        model="test-model",
        project_context=project_context,
        cycle_label=module.cycle.label,
        module_id=module_id,
    )


def _finding(name: str, finding_id: str = "rf-shared") -> Finding:
    return Finding(
        severity="HIGH",
        fileName=name,
        section="1.01",
        issue="A concrete specification conflict.",
        actionType="EDIT",
        existingText="Old requirement.",
        replacementText="Corrected requirement.",
        codeReference=None,
        finding_id=finding_id,
    )


def _result(module_id: str, name: str, *, with_finding: bool = True) -> PipelineResult:
    module = require_module(module_id)
    findings = [_finding(name)] if with_finding else []
    return PipelineResult(
        review_result=ReviewResult(findings=findings, model="test-model"),
        files_reviewed=[name],
        cycle_label=module.cycle.label,
        module_id=module_id,
    )


def test_program_prepares_every_partition_before_any_submission(monkeypatch):
    assignments = (
        _assignment("21 13 13 Fire Sprinklers.docx", ("datacenter_fire",)),
        _assignment("07 27 26 Air Barriers.docx", ("datacenter_architecture",)),
        _assignment("26 24 13 Switchboards.docx", ("datacenter_electrical",)),
        _assignment(
            "28 46 00 Fire Detection and Alarm.docx",
            ("datacenter_electronic_safety_security",),
        ),
    )
    events: list[str] = []
    research_semaphores: list[object] = []
    trace_parent_settings: list[tuple[object | None, bool]] = []
    submit_worker_settings: list[int | None] = []
    event_lock = threading.Lock()
    prepare_barrier = threading.Barrier(4)

    def fake_prepare(*, files, module, **_kwargs):
        with event_lock:
            events.append(f"prepare:{module.module_id}")
            research_semaphores.append(_kwargs["research_call_semaphore"])
            trace_parent_settings.append(
                (
                    _kwargs["_trace_parent"],
                    _kwargs["_trace_inherit_current_parent"],
                )
            )
        # A serial prepare loop cannot pass this barrier.  It proves research /
        # preflight work from distinct modules is genuinely allowed to overlap.
        prepare_barrier.wait(timeout=2)
        return SimpleNamespace(
            module=module,
            prepared=SimpleNamespace(
                specs=[SimpleNamespace(filename=Path(path).name) for path in files]
            ),
        )

    def fake_submit(prepared, **_kwargs):
        with event_lock:
            events.append(f"submit:{prepared.module.module_id}")
            submit_worker_settings.append(_kwargs.get("realtime_review_workers"))
        return _submission(
            prepared.module.module_id,
            prepared.prepared.specs[0].filename,
        )

    monkeypatch.setattr(pp, "prepare_batch_review", fake_prepare)
    monkeypatch.setattr(pp, "submit_prepared_batch_review", fake_submit)
    monkeypatch.setattr(pp, "program_prepare_max_workers", lambda: 4)

    submission = pp.start_program_review(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        input_dir=Path("C:/specs"),
        model="test-model",
        realtime_review_workers=6,
    )

    expected_prepares = {
        "prepare:datacenter_fire",
        "prepare:datacenter_architecture",
        "prepare:datacenter_electrical",
        "prepare:datacenter_electronic_safety_security",
    }
    first_submit = next(i for i, event in enumerate(events) if event.startswith("submit:"))
    assert set(events[:first_submit]) == expected_prepares
    assert len({id(semaphore) for semaphore in research_semaphores}) == 1
    assert trace_parent_settings == [(None, False)] * 4
    assert submit_worker_settings == [6, 6, 6, 6]
    assert events[first_submit:] == [
        "submit:datacenter_fire",
        "submit:datacenter_architecture",
        "submit:datacenter_electrical",
        "submit:datacenter_electronic_safety_security",
    ]
    assert tuple(submission.partitions) == (
        "datacenter_fire",
        "datacenter_architecture",
        "datacenter_electrical",
        "datacenter_electronic_safety_security",
    )


def test_parallel_preparation_failure_closes_successful_abandoned_sibling_trace(
    monkeypatch,
):
    assignments = (
        _assignment("21 13 13 Fire Sprinklers.docx", ("datacenter_fire",)),
        _assignment("07 27 26 Air Barriers.docx", ("datacenter_architecture",)),
    )
    closed: list[tuple[str, bool, dict]] = []

    def fake_prepare(*, module, **_kwargs):
        if module.module_id == "datacenter_architecture":
            raise RuntimeError("architecture preflight failed")
        return SimpleNamespace(
            module=module,
            trace_pipeline=SimpleNamespace(span_id="trace-fire-prepared"),
        )

    monkeypatch.setattr(pp, "prepare_batch_review", fake_prepare)
    monkeypatch.setattr(
        pp._trace,
        "capture_pipeline_end_by_id",
        lambda span_id, *, success, summary=None: closed.append(
            (span_id, success, dict(summary or {}))
        ),
    )

    with pytest.raises(RuntimeError, match="architecture preflight failed"):
        pp.prepare_program_review(
            program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
            assignments=assignments,
            input_dir=Path("C:/specs"),
            model="test-model",
        )

    assert closed == [
        (
            "trace-fire-prepared",
            False,
            {
                "module_id": "datacenter_fire",
                "phase": "preparation_aborted",
                "error": (
                    "Sibling preparation failed: datacenter_architecture: "
                    "architecture preflight failed"
                ),
            },
        )
    ]


def test_global_realtime_submission_failure_closes_every_prepared_trace(monkeypatch):
    name = "00 00 00 Shared.docx"
    module_ids = ("datacenter_fire", "datacenter_architecture")
    assignments = (_assignment(name, module_ids),)
    partitions = {}
    for module_id in module_ids:
        module = require_module(module_id)
        partitions[module_id] = SimpleNamespace(
            module=module,
            prepared=SimpleNamespace(
                specs=[
                    SimpleNamespace(
                        filename=name,
                        content=f"Prepared content for {module_id}",
                        paragraph_map=[],
                    )
                ],
                pre_detected_by_filename={},
            ),
            effective_context="",
            model="test-model",
            review_transport="realtime",
            trace_pipeline=SimpleNamespace(span_id=f"trace-{module_id}"),
            diagnostics=None,
        )
    prepared = pp.PreparedProgramReview(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        partitions=partitions,
        review_transport="realtime",
    )
    closed: list[tuple[str, bool, dict]] = []
    monkeypatch.setattr(
        pp,
        "run_realtime_review_jobs",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            RuntimeError("global review preflight failed")
        ),
    )
    monkeypatch.setattr(
        pp._trace,
        "capture_pipeline_end_by_id",
        lambda span_id, *, success, summary=None: closed.append(
            (span_id, success, dict(summary or {}))
        ),
    )

    with pytest.raises(pp.ProgramSubmissionError, match="global review preflight failed"):
        pp.submit_prepared_program_review(prepared)

    assert closed == [
        (
            f"trace-{module_id}",
            False,
            {
                "module_id": module_id,
                "phase": "submission",
                "error": "global review preflight failed",
            },
        )
        for module_id in module_ids
    ]


def test_batch_submission_failure_closes_failed_and_unsubmitted_traces(monkeypatch):
    module_files = {
        "datacenter_fire": "21 13 13 Fire Sprinklers.docx",
        "datacenter_architecture": "07 27 26 Air Barriers.docx",
        "datacenter_electrical": "26 24 13 Switchboards.docx",
    }
    assignments = tuple(
        _assignment(filename, (module_id,))
        for module_id, filename in module_files.items()
    )
    partitions = {
        module_id: SimpleNamespace(
            module=require_module(module_id),
            prepared=SimpleNamespace(
                specs=[SimpleNamespace(filename=filename)]
            ),
            trace_pipeline=SimpleNamespace(span_id=f"trace-{module_id}"),
        )
        for module_id, filename in module_files.items()
    }
    prepared = pp.PreparedProgramReview(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        partitions=partitions,
        review_transport="batch",
    )
    closed: list[tuple[str, bool, dict]] = []

    def fake_submit(child, **_kwargs):
        module_id = child.module.module_id
        if module_id == "datacenter_architecture":
            raise RuntimeError("architecture submission failed")
        return _submission(module_id, child.prepared.specs[0].filename)

    monkeypatch.setattr(pp, "submit_prepared_batch_review", fake_submit)
    monkeypatch.setattr(
        pp._trace,
        "capture_pipeline_end_by_id",
        lambda span_id, *, success, summary=None: closed.append(
            (span_id, success, dict(summary or {}))
        ),
    )

    with pytest.raises(pp.ProgramSubmissionError) as caught:
        pp.submit_prepared_program_review(prepared)

    assert tuple(caught.value.partial_submission.partitions) == (
        "datacenter_fire",
    )
    assert closed == [
        (
            "trace-datacenter_architecture",
            False,
            {
                "module_id": "datacenter_architecture",
                "phase": "submission",
                "error": "architecture submission failed",
            },
        ),
        (
            "trace-datacenter_electrical",
            False,
            {
                "module_id": "datacenter_electrical",
                "phase": "submission_aborted",
                "error": "architecture submission failed",
            },
        ),
    ]


def test_program_realtime_review_uses_one_global_pool_and_local_child_ids(monkeypatch):
    name = "00 00 00 Shared.docx"
    assignments = (
        _assignment(
            name,
            tuple(HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids),
        ),
    )
    diag = object()
    partitions = {}
    for module_id in HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids:
        module = require_module(module_id)
        spec = SimpleNamespace(
            filename=name,
            content=f"Prepared content for {module_id}",
            paragraph_map=[],
        )
        partitions[module_id] = SimpleNamespace(
            module=module,
            prepared=SimpleNamespace(
                specs=[spec],
                pre_detected_by_filename={},
            ),
            effective_context=f"Context for {module_id}",
            model="test-model",
            review_transport="realtime",
            trace_pipeline=None,
            diagnostics=diag,
        )
    prepared = pp.PreparedProgramReview(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        partitions=partitions,
        review_transport="realtime",
    )

    global_runs: list[list[object]] = []
    configured_workers: list[int | None] = []
    built_children: dict[str, dict] = {}
    callbacks: list[tuple[str, ...]] = []
    progress_values: list[float] = []

    def fake_global_run(jobs, *, progress, diagnostics, max_workers=None, **_kwargs):
        assert diagnostics is diag
        global_runs.append(list(jobs))
        configured_workers.append(max_workers)
        progress(50.0, "half complete")
        progress(100.0, "complete")
        return {
            job.job_key: ReviewResult(findings=[], model=job.request_spec.model)
            for job in jobs
        }

    def fake_child_submission(
        child, *, realtime_results, request_map, started_at
    ):
        module_id = child.module.module_id
        built_children[module_id] = {
            "results": dict(realtime_results),
            "request_map": dict(request_map),
            "started_at": started_at,
        }
        submission = _submission(module_id, name)
        submission.review_transport = "realtime"
        submission.realtime_results = dict(realtime_results)
        submission.job.request_map = dict(request_map)
        submission.review_request_ids = list(request_map)
        submission.job.created_at = started_at
        return submission

    monkeypatch.setattr(pp, "run_realtime_review_jobs", fake_global_run)
    monkeypatch.setattr(pp, "build_realtime_batch_submission", fake_child_submission)

    submission = pp.submit_prepared_program_review(
        prepared,
        realtime_review_workers=6,
        progress=lambda value, _message, **_kwargs: progress_values.append(value),
        on_partition_submitted=lambda current: callbacks.append(
            tuple(current.partitions)
        ),
    )

    assert len(global_runs) == 1
    assert configured_workers == [6]
    jobs = global_runs[0]
    assert len(jobs) == 4
    # The same routed spec legitimately has the same child custom_id in each
    # module; only the opaque program job key is global.
    assert {job.custom_id for job in jobs} == {"review__00_00_00_Shared__0"}
    assert len({job.job_key for job in jobs}) == 4
    assert {job.job_key[0] for job in jobs} == set(
        HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids
    )
    assert tuple(submission.partitions) == tuple(
        HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids
    )
    assert callbacks == [
        tuple(HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids[:index])
        for index in range(1, 5)
    ]
    assert all(
        set(payload["results"]) == set(payload["request_map"])
        for payload in built_children.values()
    )
    assert progress_values == sorted(progress_values)
    assert progress_values[-1] == 55.0


def test_program_result_marks_skips_and_missing_partitions_partial():
    assignments = (
        _assignment(
            "00 00 00 Combined.docx",
            (
                "datacenter_fire",
                "datacenter_architecture",
                "datacenter_electrical",
                "datacenter_electronic_safety_security",
            ),
        ),
        _assignment("27 10 00 Structured Cabling.docx", ()),
    )
    result = pp.ProgramPipelineResult(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        module_results={
            "datacenter_fire": _result(
                "datacenter_fire", "00 00 00 Combined.docx", with_finding=False
            )
        },
    )

    assert result.status == "partial"
    assert result.skipped_files == ["27 10 00 Structured Cabling.docx"]
    assert result.missing_module_ids == [
        "datacenter_architecture",
        "datacenter_electrical",
        "datacenter_electronic_safety_security",
    ]
    assert result.routed_request_count == 1
    assert result.expected_routed_request_count == 4


def test_partial_submission_counts_only_partitions_actually_submitted():
    assignments = (
        _assignment("21 13 13 Fire Sprinklers.docx", ("datacenter_fire",)),
        _assignment("07 27 26 Air Barriers.docx", ("datacenter_architecture",)),
        _assignment(
            "00 00 00 Combined.docx",
            ("datacenter_fire", "datacenter_architecture"),
        ),
    )
    fire_submission = _submission(
        "datacenter_fire", "21 13 13 Fire Sprinklers.docx"
    )
    fire_submission.files_reviewed.append("00 00 00 Combined.docx")
    fire_submission.review_request_ids.append("review__datacenter_fire__1")
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        partitions={"datacenter_fire": fire_submission},
    )

    assert submission.files_reviewed == [
        "21 13 13 Fire Sprinklers.docx",
        "00 00 00 Combined.docx",
    ]
    assert submission.expected_files_reviewed == [
        "21 13 13 Fire Sprinklers.docx",
        "07 27 26 Air Barriers.docx",
        "00 00 00 Combined.docx",
    ]
    assert submission.routed_request_count == 2
    assert submission.expected_routed_request_count == 4
    assert submission.missing_module_ids == ("datacenter_architecture",)


def test_partial_realtime_submission_skips_remote_batch_polling(monkeypatch):
    name = "21 13 13 Fire Sprinklers.docx"
    assignment = _assignment(name, ("datacenter_fire",))
    partial = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(assignment,),
        partitions={"datacenter_fire": _submission("datacenter_fire", name)},
        review_transport="realtime",
    )
    calls: list[str] = []
    monkeypatch.setattr(
        gui_batch,
        "on_realtime_reviewed",
        lambda _app, _submission: calls.append("collect"),
    )
    monkeypatch.setattr(
        gui_batch,
        "on_batch_submitted",
        lambda _app, _submission: calls.append("poll"),
    )

    gui_batch._continue_partial_program(object(), partial)

    assert calls == ["collect"]


def test_partial_batch_submission_log_reports_actual_and_expected_counts():
    fire_name = "21 13 13 Fire Sprinklers.docx"
    architecture_name = "07 27 26 Air Barriers.docx"
    partial = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(
            _assignment(fire_name, ("datacenter_fire",)),
            _assignment(architecture_name, ("datacenter_architecture",)),
        ),
        partitions={
            "datacenter_fire": _submission("datacenter_fire", fire_name)
        },
    )
    warnings: list[str] = []
    successes: list[str] = []
    app = SimpleNamespace(
        progress_bar=SimpleNamespace(set=lambda _value: None),
        log=SimpleNamespace(
            log_warning=warnings.append,
            log_success=successes.append,
            log=lambda *_args, **_kwargs: None,
            log_step=lambda *_args, **_kwargs: None,
        ),
        run_button=SimpleNamespace(configure=lambda **_kwargs: None),
        _poll_batch=lambda: None,
    )

    gui_batch.on_batch_submitted(app, partial)

    assert successes == []
    assert any("Partial batch submission retained" in item for item in warnings)
    assert any("1 of 2 routed spec(s)" in item for item in warnings)
    assert any("1 of 2 module request(s)" in item for item in warnings)


def test_batch_poll_progress_stops_at_collection_boundary():
    values: list[float] = []
    app = SimpleNamespace(
        _diagnostics_report=None,
        progress_bar=SimpleNamespace(set=values.append),
        log=SimpleNamespace(log=lambda *_args, **_kwargs: None),
    )
    status = BatchStatus(
        status="ended",
        processing=0,
        succeeded=10,
        errored=0,
        canceled=0,
        expired=0,
        total=10,
    )

    gui_batch.update_poll_progress(app, status)

    assert values == [0.55]


def test_verification_progress_is_mapped_into_its_collection_stage():
    values: list[float] = []
    mapped = gui_batch._verification_stage_progress(
        lambda value, _message, **_kwargs: values.append(value),
        stage_start=55.0,
        stage_end=72.0,
    )

    mapped(60.0, "start")
    mapped(77.5, "middle")
    mapped(95.0, "end")

    assert values == [55.0, 63.5, 72.0]


def test_bare_batch_recovery_resolves_hyperscale_module_short_name():
    program = HYPERSCALE_DATACENTER_PROGRAM
    assert gui_batch._resolve_recovery_module(program, "1").module_id == (
        program.implemented_module_ids[0]
    )
    assert gui_batch._resolve_recovery_module(
        program, "datacenter_architecture"
    ).module_id == "datacenter_architecture"
    assert gui_batch._resolve_recovery_module(
        program, "electrical"
    ).module_id == "datacenter_electrical"
    assert gui_batch._resolve_recovery_module(program, "4").module_id == (
        "datacenter_electronic_safety_security"
    )
    assert gui_batch._resolve_recovery_module(
        program, "datacenter_electronic_safety_security"
    ).module_id == "datacenter_electronic_safety_security"
    assert gui_batch._resolve_recovery_module(program, "security").module_id == (
        "datacenter_electronic_safety_security"
    )


def test_stale_saved_program_falls_back_to_valid_legacy_module():
    resolved = resolve_saved_program("removed_program", "datacenter_architecture")
    assert resolved is HYPERSCALE_DATACENTER_PROGRAM


def test_program_collection_runs_one_qualified_drawing_pass(monkeypatch):
    name = "00 00 00 Combined.docx"
    assignment = _assignment(
        name,
        (
            "datacenter_fire",
            "datacenter_architecture",
            "datacenter_electrical",
            "datacenter_electronic_safety_security",
        ),
    )
    digest_context = (
        "--- BEGIN ATTACHMENT: Construction Drawing Digest ---\n"
        "A101 shows the data hall enclosure. [set.pdf p.1]\n"
        "--- END ATTACHMENT: Construction Drawing Digest ---"
    )
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(assignment,),
        partitions={
            module_id: _submission(
                module_id, name, project_context=digest_context
            )
            for module_id in HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids
        },
    )
    child_calls: list[tuple[str, bool]] = []
    active = 0
    max_active = 0
    active_lock = threading.Lock()
    overlap = threading.Barrier(2)
    sentinel = object()

    def fake_collect(child, *, include_drawing_impact, **_kwargs):
        nonlocal active, max_active
        with active_lock:
            active += 1
            max_active = max(max_active, active)
        child_calls.append((child.module_id, include_drawing_impact))
        overlap.wait(timeout=2)
        # Reverse each pair's completion order so insertion-by-as_completed
        # would visibly perturb the declared program order.
        if child.module_id in {"datacenter_fire", "datacenter_electrical"}:
            time.sleep(0.02)
        with active_lock:
            active -= 1
        return _result(child.module_id, name)

    def fake_drawing(*, findings, module, **_kwargs):
        assert module.display_name == HYPERSCALE_DATACENTER_PROGRAM.display_name
        assert {finding.finding_id for finding in findings} == {
            "datacenter_fire::rf-shared",
            "datacenter_architecture::rf-shared",
            "datacenter_electrical::rf-shared",
            "datacenter_electronic_safety_security::rf-shared",
        }
        return sentinel

    monkeypatch.setattr(pp, "run_batch_collection_headless", fake_collect)
    monkeypatch.setattr(pp, "_make_verification_cache", lambda **_kwargs: object())
    monkeypatch.setattr(pp, "_persist_verification_cache", lambda *_a, **_k: None)
    monkeypatch.setattr("src.drawing_impact.run_drawing_impact", fake_drawing)
    monkeypatch.setattr(pp, "program_collection_max_workers", lambda: 2)

    result = pp.collect_program_results(submission)

    assert set(child_calls) == {
        ("datacenter_fire", False),
        ("datacenter_architecture", False),
        ("datacenter_electrical", False),
        ("datacenter_electronic_safety_security", False),
    }
    assert max_active == 2
    assert tuple(result.module_results) == (
        "datacenter_fire",
        "datacenter_architecture",
        "datacenter_electrical",
        "datacenter_electronic_safety_security",
    )
    assert result.drawing_impact_result is sentinel


def test_later_collection_failure_retains_completed_module_result(monkeypatch):
    fire_name = "21 13 13 Fire Sprinklers.docx"
    architecture_name = "07 27 26 Air Barriers.docx"
    assignments = (
        _assignment(fire_name, ("datacenter_fire",)),
        _assignment(architecture_name, ("datacenter_architecture",)),
    )
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=assignments,
        partitions={
            "datacenter_fire": _submission("datacenter_fire", fire_name),
            "datacenter_architecture": _submission(
                "datacenter_architecture", architecture_name
            ),
        },
    )
    submission.partitions["datacenter_architecture"].trace_span_id = (
        "trace-architecture-failed"
    )
    persisted: list[object] = []
    finished: list[str] = []
    progress_values: list[float] = []
    trace_ends: list[tuple[str, bool, dict]] = []
    overlap = threading.Barrier(2)

    def fake_collect(child, **_kwargs):
        overlap.wait(timeout=2)
        if child.module_id == "datacenter_architecture":
            finished.append(child.module_id)
            raise RuntimeError("architecture result endpoint timed out")
        time.sleep(0.02)
        finished.append(child.module_id)
        return _result(child.module_id, fire_name, with_finding=False)

    monkeypatch.setattr(pp, "run_batch_collection_headless", fake_collect)
    monkeypatch.setattr(pp, "_make_verification_cache", lambda **_kwargs: object())
    monkeypatch.setattr(
        pp,
        "_persist_verification_cache",
        lambda cache, **_kwargs: persisted.append(cache),
    )
    monkeypatch.setattr(pp, "program_collection_max_workers", lambda: 2)
    monkeypatch.setattr(
        pp._trace,
        "capture_pipeline_end_by_id",
        lambda span_id, *, success, summary=None: trace_ends.append(
            (span_id, success, dict(summary or {}))
        ),
    )

    result = pp.collect_program_results(
        submission,
        progress=lambda value, _message, **_kwargs: progress_values.append(value),
    )

    assert set(finished) == {"datacenter_fire", "datacenter_architecture"}
    assert tuple(result.module_results) == ("datacenter_fire",)
    assert result.module_errors == {
        "datacenter_architecture": "architecture result endpoint timed out"
    }
    assert result.missing_module_ids == ["datacenter_architecture"]
    assert result.status == "partial"
    assert len(persisted) == 1
    assert progress_values == sorted(progress_values)
    assert progress_values[-1] == 95.0
    assert trace_ends == [
        (
            "trace-architecture-failed",
            False,
            {
                "module_id": "datacenter_architecture",
                "phase": "collection",
                "error": "architecture result endpoint timed out",
            },
        )
    ]


def test_concurrent_collection_activates_each_module_trace_parent(monkeypatch):
    fire_name = "21 13 13 Fire Sprinklers.docx"
    architecture_name = "07 27 26 Air Barriers.docx"
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(
            _assignment(fire_name, ("datacenter_fire",)),
            _assignment(architecture_name, ("datacenter_architecture",)),
        ),
        partitions={
            "datacenter_fire": _submission("datacenter_fire", fire_name),
            "datacenter_architecture": _submission(
                "datacenter_architecture", architecture_name
            ),
        },
    )
    submission.partitions["datacenter_fire"].trace_span_id = "trace-fire"
    submission.partitions["datacenter_architecture"].trace_span_id = "trace-arch"
    seen: dict[str, str | None] = {}
    overlap = threading.Barrier(2)

    def fake_collect(child, **_kwargs):
        active = current_span()
        seen[child.module_id] = active.span_id if active is not None else None
        overlap.wait(timeout=2)
        return _result(
            child.module_id,
            child.files_reviewed[0],
            with_finding=False,
        )

    monkeypatch.setattr(pp, "run_batch_collection_headless", fake_collect)
    monkeypatch.setattr(pp, "_make_verification_cache", lambda **_kwargs: object())
    monkeypatch.setattr(pp, "_persist_verification_cache", lambda *_a, **_k: None)
    monkeypatch.setattr(pp, "program_collection_max_workers", lambda: 2)

    pp.collect_program_results(submission)

    assert seen == {
        "datacenter_fire": "trace-fire",
        "datacenter_architecture": "trace-arch",
    }
    assert current_span() is None


def test_pending_program_manifest_round_trip_and_strict_membership(tmp_path):
    name = "00 00 00 Combined.docx"
    assignment = _assignment(
        name,
        (
            "datacenter_fire",
            "datacenter_architecture",
            "datacenter_electrical",
            "datacenter_electronic_safety_security",
        ),
    )
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(assignment,),
        partitions={
            module_id: _submission(module_id, name)
            for module_id in HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids
        },
    )
    path = tmp_path / "pending.json"
    save_pending_program_run(
        PendingProgramRun.from_submission(submission), path=path
    )

    loaded = load_pending_run(path=path)
    assert isinstance(loaded, PendingProgramRun)
    assert loaded.program_id == HYPERSCALE_DATACENTER_PROGRAM.program_id
    assert set(loaded.partitions) == {
        "datacenter_fire",
        "datacenter_architecture",
        "datacenter_electrical",
        "datacenter_electronic_safety_security",
    }

    loaded.partitions["california_k12_mep"] = loaded.partitions.pop(
        "datacenter_fire"
    )
    try:
        loaded.to_submission()
    except ValueError as exc:
        assert "not a member" in str(exc)
    else:  # pragma: no cover - explicit honesty contract
        raise AssertionError("out-of-program resume partition was accepted")


def test_pending_program_manifest_accepts_legacy_two_module_run(tmp_path):
    """A pre-electrical fire/architecture manifest remains resumable."""
    name = "00 00 00 Combined.docx"
    assignment = _assignment(
        name, ("datacenter_fire", "datacenter_architecture")
    )
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(assignment,),
        partitions={
            module_id: _submission(module_id, name)
            for module_id in ("datacenter_fire", "datacenter_architecture")
        },
    )
    path = tmp_path / "pending-two-module.json"
    save_pending_program_run(
        PendingProgramRun.from_submission(submission), path=path
    )

    loaded = load_pending_run(path=path)
    assert isinstance(loaded, PendingProgramRun)
    resumed = loaded.to_submission()
    assert tuple(resumed.partitions) == (
        "datacenter_fire",
        "datacenter_architecture",
    )
    assert resumed.missing_module_ids == ()


def test_pending_program_manifest_accepts_pre_security_three_module_run(tmp_path):
    """A pre-security fire/architecture/electrical manifest remains resumable."""
    name = "00 00 00 Combined.docx"
    legacy_module_ids = (
        "datacenter_fire",
        "datacenter_architecture",
        "datacenter_electrical",
    )
    assignment = _assignment(name, legacy_module_ids)
    submission = pp.ProgramSubmission(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(assignment,),
        partitions={
            module_id: _submission(module_id, name)
            for module_id in legacy_module_ids
        },
    )
    path = tmp_path / "pending-three-module.json"
    save_pending_program_run(
        PendingProgramRun.from_submission(submission), path=path
    )

    loaded = load_pending_run(path=path)
    assert isinstance(loaded, PendingProgramRun)
    resumed = loaded.to_submission()
    assert tuple(resumed.partitions) == legacy_module_ids
    assert resumed.missing_module_ids == ()


def test_single_batch_resume_never_falls_back_or_changes_cycle():
    child = _submission(
        "datacenter_architecture", "07 27 26 Air Barriers.docx"
    )
    pending = PendingBatch.from_submission(child)
    pending.module_id = "removed_architecture_module"
    try:
        pending.to_submission()
    except KeyError:
        pass
    else:  # pragma: no cover - explicit honesty contract
        raise AssertionError("unknown resume module silently fell back")

    pending.module_id = "datacenter_architecture"
    pending.cycle_label = "obsolete-architecture-cycle"
    try:
        pending.to_submission()
    except ValueError as exc:
        assert "different code basis" in str(exc)
    else:  # pragma: no cover
        raise AssertionError("cycle drift was accepted")


def test_program_report_and_sidecar_preserve_module_provenance(tmp_path):
    name = "00 00 00 Combined.docx"
    assignment = _assignment(
        name,
        (
            "datacenter_fire",
            "datacenter_architecture",
            "datacenter_electrical",
            "datacenter_electronic_safety_security",
        ),
    )
    result = pp.ProgramPipelineResult(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(assignment,),
        module_results={
            module_id: _result(module_id, name)
            for module_id in HYPERSCALE_DATACENTER_PROGRAM.implemented_module_ids
        },
    )

    report_path = export_report(result, tmp_path / "program.docx")
    doc = Document(report_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Hyperscale Data Center Specification Review Report" in text
    assert require_module("datacenter_fire").display_name in text
    assert require_module("datacenter_architecture").display_name in text
    assert require_module("datacenter_electrical").display_name in text
    assert (
        require_module("datacenter_electronic_safety_security").display_name
        in text
    )

    sidecar = build_edit_instructions(result, report_path=report_path)
    assert sidecar["schema_version"] == 5
    assert sidecar["program_id"] == HYPERSCALE_DATACENTER_PROGRAM.program_id
    assert sidecar["module_errors"] == {}
    assert sidecar["submission_coverage"] == {
        "submitted_files": [name],
        "expected_files": [name],
        "submitted_requests": 4,
        "expected_requests": 4,
    }
    assert sidecar["edit_count"] == 4
    assert {entry["module_id"] for entry in sidecar["edits"]} == {
        "datacenter_fire",
        "datacenter_architecture",
        "datacenter_electrical",
        "datacenter_electronic_safety_security",
    }


def test_partial_program_report_surfaces_module_collection_error(tmp_path):
    name = "00 00 00 Combined.docx"
    result = pp.ProgramPipelineResult(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(
            _assignment(name, ("datacenter_fire", "datacenter_architecture")),
        ),
        module_results={
            "datacenter_fire": _result(
                "datacenter_fire", name, with_finding=False
            )
        },
        module_errors={
            "datacenter_architecture": "architecture result endpoint timed out"
        },
    )

    report_path = export_report(result, tmp_path / "partial-program.docx")
    doc = Document(report_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Coverage status: PARTIAL" in text
    assert "architecture result endpoint timed out" in text

    sidecar = build_edit_instructions(result, report_path=report_path)
    assert sidecar["module_errors"] == result.module_errors


def test_partial_submission_report_distinguishes_submitted_from_expected(tmp_path):
    fire_name = "21 13 13 Fire Sprinklers.docx"
    architecture_name = "07 27 26 Air Barriers.docx"
    result = pp.ProgramPipelineResult(
        program_id=HYPERSCALE_DATACENTER_PROGRAM.program_id,
        assignments=(
            _assignment(fire_name, ("datacenter_fire",)),
            _assignment(architecture_name, ("datacenter_architecture",)),
        ),
        module_results={
            "datacenter_fire": _result(
                "datacenter_fire", fire_name, with_finding=False
            )
        },
    )

    assert result.files_reviewed == [fire_name]
    assert result.expected_files_reviewed == [fire_name, architecture_name]
    assert result.routed_request_count == 1
    assert result.expected_routed_request_count == 2

    report_path = export_report(result, tmp_path / "partial-submission.docx")
    doc = Document(report_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Routed Specifications Submitted: 1 of 2" in text
    assert "Routed Review Requests Submitted: 1 of 2" in text

    sidecar = build_edit_instructions(result, report_path=report_path)
    assert sidecar["submission_coverage"] == {
        "submitted_files": [fire_name],
        "expected_files": [fire_name, architecture_name],
        "submitted_requests": 1,
        "expected_requests": 2,
    }
