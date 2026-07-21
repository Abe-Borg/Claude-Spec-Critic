"""Verification reconciliation onto the requirements profile (WS6 / E5).

CORRECTED / DISPUTED / contested findings must flow back onto the profile
surfaces: matched items get an inline amber marker, unattributable
corrections surface as one profile-level caution, and the ``.profile.json``
export carries a machine-readable ``verification_corrections`` list.
"""
from __future__ import annotations

from docx import Document

from src.output.verification_reconciliation import (
    VerificationCorrection,
    collect_verification_corrections,
    correction_marker_text,
    corrections_by_item_id,
    serialize_corrections,
    unattributed_corrections,
)
from src.review.reviewer import Finding, ReviewResult
from src.verification.verifier import VerificationResult


def _finding(
    *,
    issue: str = "Cited climate zone appears wrong",
    verdict: str | None = "CORRECTED",
    explanation: str = "",
    correction: str | None = None,
    finding_id: str = "rf-0123456789ab",
    grounded: bool = True,
    models_disagreed: bool = False,
) -> Finding:
    f = Finding(
        severity="HIGH",
        fileName="08 80 00 GLAZING.docx",
        section="1.4",
        issue=issue,
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference="",
        finding_id=finding_id,
    )
    if verdict is not None:
        f.verification = VerificationResult(
            verdict=verdict,
            explanation=explanation,
            correction=correction,
            grounded=grounded,
            sources=["https://example.gov/x"] if grounded else [],
            accepted_sources=["https://example.gov/x"] if grounded else [],
            models_disagreed=models_disagreed,
        )
    return f


class TestCollect:
    def test_corrected_finding_selected_with_item_ids_from_issue(self):
        f = _finding(
            issue="Profile item r-aaaaaaaaaaaa cites Climate Zone 7; site is Zone 6.",
        )
        corrections = collect_verification_corrections([f])
        assert len(corrections) == 1
        c = corrections[0]
        assert c.verdict_label == "corrected"
        assert c.item_ids == ("r-aaaaaaaaaaaa",)
        assert c.finding_id == "rf-0123456789ab"

    def test_ids_extracted_from_rationale_and_correction(self):
        f = _finding(
            issue="Wrong climate zone cited.",
            explanation="The researched item r-bbbbbbbbbbbb asserted Zone 7.",
            correction="Zone 6 per NECB mapping; supersedes r-cccccccccccc.",
        )
        (c,) = collect_verification_corrections([f])
        assert set(c.item_ids) == {"r-bbbbbbbbbbbb", "r-cccccccccccc"}

    def test_disputed_selected(self):
        (c,) = collect_verification_corrections([_finding(verdict="DISPUTED")])
        assert c.verdict_label == "disputed"

    def test_contested_via_models_disagreed(self):
        f = _finding(verdict="CONFIRMED", models_disagreed=True)
        (c,) = collect_verification_corrections([f])
        assert c.verdict_label == "contested"

    def test_confirmed_and_unverified_not_selected(self):
        confirmed = _finding(verdict="CONFIRMED")
        unverified = _finding(verdict="UNVERIFIED", grounded=False)
        unchecked = _finding(verdict=None)
        assert collect_verification_corrections([confirmed, unverified, unchecked]) == []

    def test_duplicate_ids_collapse(self):
        f = _finding(
            issue="r-aaaaaaaaaaaa is wrong",
            explanation="confirming r-aaaaaaaaaaaa is wrong",
        )
        (c,) = collect_verification_corrections([f])
        assert c.item_ids == ("r-aaaaaaaaaaaa",)

    def test_malformed_ids_ignored(self):
        f = _finding(issue="r-XYZ and r-12345 are not ids; r-abcdefabcdef is.")
        (c,) = collect_verification_corrections([f])
        assert c.item_ids == ("r-abcdefabcdef",)


class TestHelpers:
    def test_index_and_unattributed_split(self):
        with_id = VerificationCorrection("rf-1", "a.docx", "corrected", ("r-aaaaaaaaaaaa",))
        without = VerificationCorrection("rf-2", "b.docx", "disputed", ())
        by_id = corrections_by_item_id([with_id, without])
        assert by_id == {"r-aaaaaaaaaaaa": [with_id]}
        assert unattributed_corrections([with_id, without]) == [without]

    def test_marker_text(self):
        c = VerificationCorrection("rf-1", "a.docx", "corrected", ())
        assert correction_marker_text(c) == (
            "⚠ corrected by verification — see finding rf-1"
        )

    def test_marker_falls_back_to_file_name(self):
        c = VerificationCorrection("", "a.docx", "disputed", ())
        assert "see finding a.docx" in correction_marker_text(c)

    def test_serialize(self):
        c = VerificationCorrection("rf-1", "a.docx", "corrected", ("r-aaaaaaaaaaaa",))
        assert serialize_corrections([c]) == [
            {
                "finding_id": "rf-1",
                "file_name": "a.docx",
                "verdict": "corrected",
                "requirement_item_ids": ["r-aaaaaaaaaaaa"],
            }
        ]


def _doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestRenderAnnotations:
    def _profile(self):
        from src.research.requirements_research import (
            DimensionStatus,
            RequirementsProfile,
            ResearchItem,
        )

        return RequirementsProfile(
            items=[
                ResearchItem(
                    item_id="r-aaaaaaaaaaaa",
                    dimension_id="governing_codes",
                    topic="Climate zone",
                    category="governing_code",
                    requirement="The site is in Climate Zone 7.",
                    grounded=True,
                    accepted_sources=["https://example.gov/x"],
                    confidence=0.8,
                )
            ],
            dimension_statuses=[
                DimensionStatus(
                    dimension_id="governing_codes", status="completed", item_count=1
                )
            ],
            research_date="2026-07-14",
            project={"city": "Markham", "state_or_province": "ON", "country": "CA",
                     "client_name": "ExampleCo"},
        )

    def _module(self):
        import dataclasses

        from src.modules import DEFAULT_MODULE

        return dataclasses.replace(
            DEFAULT_MODULE,
            project_profile_enabled=True,
            research_persona="p",
            research_dimensions=(),
            compliance_persona="p",
            compliance_severity_definitions="- CRITICAL — x.",
        )

    def test_matched_item_gets_marker(self):
        from src.output.report_exporter import _write_requirements_section

        doc = Document()
        corrections = [
            VerificationCorrection(
                "rf-0123456789ab", "a.docx", "corrected", ("r-aaaaaaaaaaaa",)
            )
        ]
        _write_requirements_section(
            doc, self._profile(), None, self._module(), corrections=corrections
        )
        text = _doc_text(doc)
        assert "corrected by verification — see finding rf-0123456789ab" in text
        # Attributed corrections do NOT trigger the profile-wide caution.
        assert "individual requirement items may be affected" not in text

    def test_unattributed_correction_renders_profile_caution(self):
        from src.output.report_exporter import _write_requirements_section

        doc = Document()
        corrections = [
            VerificationCorrection("rf-0123456789ab", "a.docx", "corrected", ())
        ]
        _write_requirements_section(
            doc, self._profile(), None, self._module(), corrections=corrections
        )
        text = _doc_text(doc)
        assert "Verification corrected or disputed one or more claims" in text
        assert "corrected by verification — see finding" not in text

    def test_no_corrections_renders_neither(self):
        from src.output.report_exporter import _write_requirements_section

        doc = Document()
        _write_requirements_section(doc, self._profile(), None, self._module())
        text = _doc_text(doc)
        assert "by verification" not in text
        assert "Verification corrected or disputed" not in text


class TestProfileJsonExport:
    def test_export_carries_verification_corrections(self):
        from src.modules import require_module
        from src.orchestration.pipeline import PipelineResult
        from src.output.edit_sidecar import build_requirements_profile_export

        module = require_module("datacenter_fire")
        corrected = _finding(
            issue="Item r-aaaaaaaaaaaa cites Climate Zone 7; site is Zone 6.",
        )
        result = PipelineResult(
            review_result=ReviewResult(findings=[corrected]),
            files_reviewed=["a.docx"],
            cycle_label=module.cycle.label,
            module_id=module.module_id,
            requirements_profile={"items": [], "research_date": "2026-07-14"},
        )
        export = build_requirements_profile_export(result)
        assert export is not None
        assert export["verification_corrections"] == [
            {
                "finding_id": "rf-0123456789ab",
                "file_name": "08 80 00 GLAZING.docx",
                "verdict": "corrected",
                "requirement_item_ids": ["r-aaaaaaaaaaaa"],
            }
        ]
