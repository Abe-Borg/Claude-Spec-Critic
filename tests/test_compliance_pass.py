"""WS-4a compliance pass: checker, coverage merge, pipeline stage, report.

Hermetic throughout. The fake streaming client scripts every response; the
chunked tests stub ``count_tokens`` (network-free, same convention as the
rest of the suite) and shrink the input ceiling to force chunking.
"""
from __future__ import annotations

import dataclasses
import json

import pytest

from src.compliance import run_chunked_compliance_check, run_compliance_check
from src.compliance import compliance_checker as cc
from src.core.api_config import MODEL_SONNET_46
from src.input.extractor import ExtractedSpec
from src.modules import DEFAULT_MODULE, ResearchDimension
from src.research import DimensionStatus, RequirementsProfile, ResearchItem
from src.review.reviewer import Finding, ReviewResult
from src.review.structured_schemas import (
    COMPLIANCE_FINDINGS_SCHEMA,
    COMPLIANCE_TOOL_NAME,
    compliance_findings_tool,
    compliance_tool_choice,
)
from src.verification.verifier import VerificationResult
from tests.fixtures.fake_anthropic import (
    FakeMessage,
    FakeTextBlock,
    compliance_tool_use_response,
    sample_compliance_payload,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _spec(content: str, filename: str = "21 13 13 Wet-Pipe.docx") -> ExtractedSpec:
    return ExtractedSpec(
        filename=filename, content=content, word_count=len(content.split())
    )


def _item(item_id: str, requirement: str, *, grounded: bool = True, **overrides):
    defaults = dict(
        item_id=item_id,
        dimension_id="governing_codes",
        topic="Topic",
        category="governing_code",
        requirement=requirement,
        grounded=grounded,
        accepted_sources=["https://codes.example.gov/x"] if grounded else [],
        confidence=0.8,
    )
    defaults.update(overrides)
    return ResearchItem(**defaults)


def _profile(items=None) -> RequirementsProfile:
    if items is None:
        items = [
            _item("r-aaaaaaaaaaaa", "The 2024 IBC as amended governs."),
            _item("r-bbbbbbbbbbbb", "Municipal Amendment 12-2024 applies."),
        ]
    return RequirementsProfile(
        items=list(items),
        dimension_statuses=[
            DimensionStatus(dimension_id="governing_codes", status="completed")
        ],
        research_date="2026-07-14",
        project={"city": "Markham", "state_or_province": "ON", "country": "CA",
                 "client_name": "ExampleCo"},
    )


def _enabled_module(**overrides):
    defaults = dict(
        project_profile_enabled=True,
        research_persona="You are a test research assistant.",
        research_dimensions=(
            ResearchDimension(
                dimension_id="governing_codes",
                title="Governing codes",
                prompt_template="Codes for {city}.",
            ),
        ),
        compliance_persona="You are a test compliance reviewer.",
        compliance_severity_definitions="- CRITICAL — permit-blocking omission.",
    )
    defaults.update(overrides)
    return dataclasses.replace(DEFAULT_MODULE, **defaults)


class _FakeStream:
    def __init__(self, message):
        self._message = message

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False

    @property
    def text_stream(self):
        return iter(())

    def get_final_message(self):
        return self._message


class _FakeMessagesAPI:
    def __init__(self, route):
        self._route = route
        self.calls: list[dict] = []

    def stream(self, **kwargs):
        self.calls.append(kwargs)
        result = self._route(kwargs)
        if isinstance(result, Exception):
            raise result
        return _FakeStream(result)


class FakeComplianceClient:
    def __init__(self, route):
        self.messages = _FakeMessagesAPI(route)

    @property
    def calls(self):
        return self.messages.calls


def _route_single(response):
    return lambda kwargs: response


def _route_by_marker(script: dict[str, list]):
    remaining = {marker: list(items) for marker, items in script.items()}

    def route(kwargs):
        user_content = kwargs["messages"][0]["content"]
        for marker, items in remaining.items():
            if marker in user_content:
                if not items:
                    raise AssertionError(f"script exhausted for {marker!r}")
                return items.pop(0)
        raise AssertionError(f"no route for message: {user_content[:120]!r}")

    return route


@pytest.fixture(autouse=True)
def stub_compliance_tokens(monkeypatch):
    """Word-count stand-in for ``count_tokens`` in the compliance module.

    Real ``count_tokens`` downloads the cl100k_base encoding on first use;
    stubbing keeps the suite network-free (the established convention).
    """
    monkeypatch.setattr(cc, "count_tokens", lambda text: len(text.split()))


@pytest.fixture
def fake_client(monkeypatch):
    """Install a scripted client; tests set ``holder['route']``."""
    holder: dict = {"route": None}
    client = FakeComplianceClient(lambda kwargs: holder["route"](kwargs))
    monkeypatch.setattr(cc, "_get_client", lambda: client)
    holder["client"] = client
    return holder


# ---------------------------------------------------------------------------
# Schema + module contract
# ---------------------------------------------------------------------------


class TestComplianceSchema:
    def test_strict_subset(self):
        def _walk(schema):
            if schema.get("type") == "object" or "properties" in schema:
                assert schema.get("additionalProperties") is False
                assert sorted(schema["required"]) == sorted(schema["properties"])
            for value in schema.get("properties", {}).values():
                _walk(value)
                if isinstance(value.get("items"), dict):
                    _walk(value["items"])

        _walk(COMPLIANCE_FINDINGS_SCHEMA)
        text = json.dumps(COMPLIANCE_FINDINGS_SCHEMA)
        for forbidden in ('"minimum"', '"maximum"', '"minLength"'):
            assert forbidden not in text

    def test_tool_builder_and_choice(self):
        tool = compliance_findings_tool(model=MODEL_SONNET_46)
        assert tool["name"] == COMPLIANCE_TOOL_NAME == "submit_compliance_findings"
        assert tool["strict"] is True
        assert compliance_tool_choice() == {
            "type": "auto",
            "disable_parallel_tool_use": True,
        }

    def test_disabled_module_rejects_compliance_slots(self):
        from src.modules import validate_module_registry

        with pytest.raises(ValueError, match="compliance_persona must be empty"):
            validate_module_registry(
                [dataclasses.replace(DEFAULT_MODULE, compliance_persona="dead")]
            )

    def test_enabled_module_requires_compliance_slots(self):
        from src.modules import validate_module_registry

        with pytest.raises(ValueError, match="compliance_persona"):
            validate_module_registry([_enabled_module(compliance_persona="")])


# ---------------------------------------------------------------------------
# Single-pass checker
# ---------------------------------------------------------------------------


class TestRunComplianceCheck:
    def test_completed_path_parses_coverage_and_findings(self, fake_client):
        fake_client["route"] = _route_single(compliance_tool_use_response())
        module = _enabled_module()
        result = run_compliance_check(
            [_spec("Comply with the 2024 IBC as amended.")],
            _profile(),
            [],
            cycle=module.cycle,
        )
        assert result.cross_check_status == "completed"
        assert [c["status"] for c in result.coverage] == ["represented", "missing"]
        assert len(result.findings) == 1
        assert result.findings[0].actionType == "ADD"
        assert "missing from the package" in result.thinking

    def test_skips_without_grounded_items(self, fake_client):
        fake_client["route"] = _route_single(compliance_tool_use_response())
        ungrounded = _profile(
            [_item("r-aaaaaaaaaaaa", "Ungrounded requirement.", grounded=False)]
        )
        result = run_compliance_check(
            [_spec("body")], ungrounded, [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "skipped"
        assert "no grounded requirement items" in result.thinking
        assert fake_client["client"].calls == []

    def test_skips_without_specs(self, fake_client):
        fake_client["route"] = _route_single(compliance_tool_use_response())
        result = run_compliance_check([], _profile(), [], cycle=_enabled_module().cycle)
        assert result.cross_check_status == "skipped"
        assert fake_client["client"].calls == []

    def test_nonretryable_error_fails(self, fake_client):
        fake_client["route"] = _route_single(RuntimeError("boom"))
        result = run_compliance_check(
            [_spec("body")], _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "failed"
        assert "boom" in (result.error or "")

    def test_incomplete_stop_reason_fails(self, fake_client):
        fake_client["route"] = _route_single(
            FakeMessage(content=[FakeTextBlock(text="trunc")], stop_reason="max_tokens")
        )
        result = run_compliance_check(
            [_spec("body")], _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "failed"
        assert result.parse_status == "incomplete"

    def test_no_payload_fails(self, fake_client):
        fake_client["route"] = _route_single(
            FakeMessage(content=[FakeTextBlock(text="prose")], stop_reason="end_turn")
        )
        result = run_compliance_check(
            [_spec("body")], _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "failed"
        assert result.parse_status == "parse_error"

    def test_tagged_json_fallback(self, fake_client):
        payload = sample_compliance_payload()
        text = f"<compliance_json>{json.dumps(payload)}</compliance_json>"
        fake_client["route"] = _route_single(
            FakeMessage(content=[FakeTextBlock(text=text)], stop_reason="end_turn")
        )
        result = run_compliance_check(
            [_spec("body")], _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "completed"
        assert len(result.coverage) == 2

    def test_coverage_normalization_clamps_and_filters(self, fake_client):
        payload = {
            "compliance_summary": "",
            "coverage": [
                {"requirement_id": "r-aaaaaaaaaaaa", "status": "BANANA",
                 "evidence": None, "fileName": None},
                {"requirement_id": "r-000000000000", "status": "missing",
                 "evidence": None, "fileName": None},  # unknown id → dropped
                {"requirement_id": "", "status": "missing",
                 "evidence": None, "fileName": None},  # no id → dropped
            ],
            "findings": [],
        }
        fake_client["route"] = _route_single(
            compliance_tool_use_response(payload=payload)
        )
        result = run_compliance_check(
            [_spec("body")], _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.coverage == [
            {
                "requirement_id": "r-aaaaaaaaaaaa",
                "status": "unclear",
                "evidence": None,
                "fileName": None,
            }
        ]

    def test_user_message_separates_controlling_and_unverified(self, fake_client):
        fake_client["route"] = _route_single(compliance_tool_use_response())
        profile = _profile(
            [
                _item("r-aaaaaaaaaaaa", "Grounded requirement."),
                _item("r-bbbbbbbbbbbb", "Ungrounded requirement.", grounded=False),
                _item(
                    "r-cccccccccccc",
                    "Permit fee schedule.",
                    actionability="process_advisory",
                ),
            ]
        )
        run_compliance_check(
            [_spec("body")], profile, [], cycle=_enabled_module().cycle
        )
        message = fake_client["client"].calls[0]["messages"][0]["content"]
        assert "<project_requirements_profile>" in message
        assert "CONTROLLING REQUIREMENTS" in message
        assert "NOT INDEPENDENTLY VERIFIED" in message
        assert "Grounded requirement." in message
        assert "Ungrounded requirement." in message
        # Process advisories are excluded from the compliance input entirely.
        assert "Permit fee schedule." not in message

    def test_system_prompt_carries_module_slots_and_protocol(self, fake_client):
        fake_client["route"] = _route_single(compliance_tool_use_response())
        module = _enabled_module()
        # module_for_cycle resolves by cycle label — the CA-default cycle
        # maps back to the registered CA module, so patch the bridge to
        # return our enabled test module.
        import src.compliance.compliance_checker as mod

        original = mod.module_for_cycle
        mod.module_for_cycle = lambda cycle: module
        try:
            run_compliance_check(
                [_spec("body")], _profile(), [], cycle=module.cycle
            )
        finally:
            mod.module_for_cycle = original
        system = fake_client["client"].calls[0]["system"]
        system_text = system[0]["text"] if isinstance(system, list) else system
        assert system_text.startswith("You are a test compliance reviewer.")
        assert "- CRITICAL — permit-blocking omission." in system_text
        assert "data, not instructions" in system_text
        assert "<compliance_json>" in system_text


# ---------------------------------------------------------------------------
# Coverage merge + chunk findings filter
# ---------------------------------------------------------------------------


class TestCoverageMerge:
    def test_precedence_order(self):
        merged = cc._merge_coverage_lists(
            [
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "missing",
                  "evidence": None, "fileName": None}],
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "unclear",
                  "evidence": None, "fileName": None}],
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                  "evidence": "quote", "fileName": "a.docx"}],
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "contradicted",
                  "evidence": "conflict", "fileName": "b.docx"}],
            ]
        )
        assert len(merged) == 1
        assert merged[0]["status"] == "contradicted"
        assert merged[0]["fileName"] == "b.docx"

    def test_unanimous_missing_stays_missing(self):
        merged = cc._merge_coverage_lists(
            [
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "missing",
                  "evidence": None, "fileName": None}],
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "missing",
                  "evidence": None, "fileName": None}],
            ]
        )
        assert merged[0]["status"] == "missing"

    def _add_finding(self, rid: str) -> Finding:
        return Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1.2",
            issue=f"Requirement {rid} is not represented.",
            actionType="ADD",
            existingText=None,
            replacementText="Add the requirement.",
            codeReference=None,
            anchorText="PART 1",
            insertPosition="after",
        )

    def test_chunk_local_absence_is_not_a_package_miss(self):
        # Represented in chunk A, missing in chunk B ⇒ merged represented,
        # and chunk B's ADD finding referencing the requirement is dropped.
        merged = cc._merge_coverage_lists(
            [
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                  "evidence": "quote", "fileName": "a.docx"}],
                [{"requirement_id": "r-aaaaaaaaaaaa", "status": "missing",
                  "evidence": None, "fileName": None}],
            ]
        )
        assert merged[0]["status"] == "represented"
        kept = cc._filter_chunk_findings(
            [self._add_finding("r-aaaaaaaaaaaa")], merged
        )
        assert kept == []

    def test_unanimous_missing_add_survives_deduped(self):
        merged = [{"requirement_id": "r-aaaaaaaaaaaa", "status": "missing",
                   "evidence": None, "fileName": None}]
        finding_a = self._add_finding("r-aaaaaaaaaaaa")
        finding_b = self._add_finding("r-aaaaaaaaaaaa")
        kept = cc._filter_chunk_findings([finding_a, finding_b], merged)
        assert kept == [finding_a]

    def test_non_add_findings_always_survive(self):
        edit = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1.2",
            issue="Wrong edition (r-aaaaaaaaaaaa).",
            actionType="EDIT",
            existingText="2015 IBC",
            replacementText="2024 IBC",
            codeReference=None,
        )
        merged = [{"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                   "evidence": None, "fileName": None}]
        assert cc._filter_chunk_findings([edit], merged) == [edit]

    def test_add_without_requirement_reference_survives(self):
        orphan = self._add_finding("")
        orphan.issue = "No requirement id referenced here."
        assert cc._filter_chunk_findings([orphan], []) == [orphan]


class TestChunkedCompliance:
    def test_partial_chunk_failure_preserved(self, fake_client, monkeypatch):
        monkeypatch.setattr(cc, "count_tokens", lambda text: len(text.split()))
        monkeypatch.setattr(cc, "COMPLIANCE_RECOMMENDED_MAX", 5_000)
        pad = "word " * 2_000
        specs = [
            _spec(pad + " DIV21SPEC alpha", "21 13 13 Wet.docx"),
            _spec(pad + " DIV21SPEC beta", "21 13 16 Dry.docx"),
            _spec(pad + " DIV22SPEC alpha", "22 11 13 Water.docx"),
            _spec(pad + " DIV22SPEC beta", "22 11 16 Piping.docx"),
        ]
        div21_payload = {
            "compliance_summary": "Div 21 fine.",
            "coverage": [
                {"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                 "evidence": "quote", "fileName": "21 13 13 Wet.docx"},
            ],
            "findings": [],
        }
        fake_client["route"] = _route_by_marker(
            {
                "DIV21SPEC": [compliance_tool_use_response(payload=div21_payload)],
                "DIV22SPEC": [RuntimeError("chunk exploded")],
            }
        )
        result = run_chunked_compliance_check(
            specs, _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "completed"
        assert result.chunk_failures == 1
        assert [c["requirement_id"] for c in result.coverage] == ["r-aaaaaaaaaaaa"]

    def test_cross_chunk_merge_drops_disproven_add(self, fake_client, monkeypatch):
        monkeypatch.setattr(cc, "count_tokens", lambda text: len(text.split()))
        monkeypatch.setattr(cc, "COMPLIANCE_RECOMMENDED_MAX", 5_000)
        pad = "word " * 2_000
        specs = [
            _spec(pad + " DIV21SPEC alpha", "21 13 13 Wet.docx"),
            _spec(pad + " DIV21SPEC beta", "21 13 16 Dry.docx"),
            _spec(pad + " DIV22SPEC alpha", "22 11 13 Water.docx"),
            _spec(pad + " DIV22SPEC beta", "22 11 16 Piping.docx"),
        ]
        represented_payload = {
            "compliance_summary": "",
            "coverage": [
                {"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                 "evidence": "quote", "fileName": "21 13 13 Wet.docx"},
            ],
            "findings": [],
        }
        missing_payload = {
            "compliance_summary": "",
            "coverage": [
                {"requirement_id": "r-aaaaaaaaaaaa", "status": "missing",
                 "evidence": None, "fileName": None},
            ],
            "findings": [
                {
                    "severity": "HIGH",
                    "fileName": "22 11 13 Water.docx",
                    "section": "1.2",
                    "issue": "Requirement r-aaaaaaaaaaaa missing from this subset.",
                    "actionType": "ADD",
                    "existingText": None,
                    "replacementText": "Add it.",
                    "codeReference": None,
                    "confidence": 0.8,
                    "anchorText": "PART 1",
                    "insertPosition": "after",
                    "evidenceElementId": None,
                }
            ],
        }
        fake_client["route"] = _route_by_marker(
            {
                "DIV21SPEC": [compliance_tool_use_response(payload=represented_payload)],
                "DIV22SPEC": [compliance_tool_use_response(payload=missing_payload)],
            }
        )
        result = run_chunked_compliance_check(
            specs, _profile(), [], cycle=_enabled_module().cycle
        )
        assert result.cross_check_status == "completed"
        merged = {c["requirement_id"]: c["status"] for c in result.coverage}
        assert merged["r-aaaaaaaaaaaa"] == "represented"
        # The chunk-local ADD was disproven by the merge — no missing finding.
        assert result.findings == []

    def test_chunked_user_message_carries_subset_note(self, fake_client, monkeypatch):
        monkeypatch.setattr(cc, "count_tokens", lambda text: len(text.split()))
        monkeypatch.setattr(cc, "COMPLIANCE_RECOMMENDED_MAX", 5_000)
        pad = "word " * 2_000
        specs = [
            _spec(pad + " DIV21SPEC alpha", "21 13 13 Wet.docx"),
            _spec(pad + " DIV21SPEC beta", "21 13 16 Dry.docx"),
            _spec(pad + " DIV22SPEC alpha", "22 11 13 Water.docx"),
            _spec(pad + " DIV22SPEC beta", "22 11 16 Piping.docx"),
        ]
        fake_client["route"] = _route_by_marker(
            {
                "DIV21SPEC": [compliance_tool_use_response()],
                "DIV22SPEC": [compliance_tool_use_response()],
            }
        )
        run_chunked_compliance_check(
            specs, _profile(), [], cycle=_enabled_module().cycle
        )
        for call in fake_client["client"].calls:
            assert "one subset of a larger specification package" in (
                call["messages"][0]["content"]
            )


# ---------------------------------------------------------------------------
# Pipeline stage
# ---------------------------------------------------------------------------


def _state(module, requirements_profile=None, findings=None, cross_findings=None):
    import time as _time

    from src.batch.batch import BatchJob
    from src.orchestration.pipeline import BatchSubmission, CollectedBatchState

    submission = BatchSubmission(
        job=BatchJob(batch_id="b1", job_type="review", request_map={}, created_at=_time.time()),
        model="claude-opus-4-8",
        project_context="ctx",
        prepared_specs=[_spec("Comply with the 2024 IBC as amended.")],
        cycle_label=module.cycle.label,
        module_id=module.module_id,
        requirements_profile=requirements_profile,
    )
    state = CollectedBatchState(
        submission=submission,
        review_result=ReviewResult(findings=list(findings or [])),
    )
    if cross_findings is not None:
        state.cross_check_result = ReviewResult(
            findings=list(cross_findings), cross_check_status="completed"
        )
    return state


class TestRunComplianceForBatch:
    def test_flag_off_module_is_untouched(self, monkeypatch):
        from src.orchestration import pipeline

        state = _state(DEFAULT_MODULE)
        out = pipeline.run_compliance_for_batch(state)
        assert out.compliance_result is None

    def test_flag_on_without_profile_reports_skipped(self, monkeypatch):
        from src.orchestration import pipeline

        module = _enabled_module()
        monkeypatch.setattr(pipeline, "get_module", lambda _mid: module)
        state = _state(module, requirements_profile=None)
        out = pipeline.run_compliance_for_batch(state)
        assert out.compliance_result is not None
        assert out.compliance_result.cross_check_status == "skipped"
        assert "profile unavailable" in out.compliance_result.thinking

    def test_runs_and_stamps_labels_and_lc_ids(self, monkeypatch, fake_client):
        from src.orchestration import pipeline

        module = _enabled_module()
        monkeypatch.setattr(pipeline, "get_module", lambda _mid: module)
        fake_client["route"] = _route_single(compliance_tool_use_response())
        state = _state(module, requirements_profile=_profile().to_dict())
        out = pipeline.run_compliance_for_batch(state)
        comp = out.compliance_result
        assert comp is not None and comp.cross_check_status == "completed"
        assert len(comp.findings) == 1
        finding = comp.findings[0]
        assert finding.section.startswith("[Compliance]")
        assert finding.finding_id.startswith("lc-") and len(finding.finding_id) == 15

    def test_lc_ids_never_collide_with_rf_cf_on_identical_content(self):
        from src.orchestration.pipeline import (
            assign_compliance_finding_ids,
            assign_cross_check_finding_ids,
            compute_finding_id,
        )

        def _mk():
            return Finding(
                severity="HIGH",
                fileName="a.docx",
                section="1.2",
                issue="Identical content.",
                actionType="REPORT_ONLY",
                existingText=None,
                replacementText=None,
                codeReference=None,
            )

        rf = compute_finding_id(_mk())
        cf = assign_cross_check_finding_ids([_mk()])[0].finding_id
        lc = assign_compliance_finding_ids([_mk()])[0].finding_id
        assert rf[3:] == cf[3:] == lc[3:]  # same content digest
        assert len({rf, cf, lc}) == 3  # prefixes keep them distinct

    def test_disputed_findings_excluded_from_already_identified(
        self, monkeypatch, fake_client
    ):
        from src.orchestration import pipeline

        module = _enabled_module()
        monkeypatch.setattr(pipeline, "get_module", lambda _mid: module)
        fake_client["route"] = _route_single(compliance_tool_use_response())
        disputed = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1.1",
            issue="DISPUTED-REVIEW-FINDING marker.",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference=None,
            verification=VerificationResult(verdict="DISPUTED", explanation=""),
        )
        clean = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1.1",
            issue="CLEAN-REVIEW-FINDING marker.",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference=None,
        )
        state = _state(
            module,
            requirements_profile=_profile().to_dict(),
            findings=[disputed, clean],
        )
        pipeline.run_compliance_for_batch(state)
        message = fake_client["client"].calls[0]["messages"][0]["content"]
        assert "CLEAN-REVIEW-FINDING" in message
        assert "DISPUTED-REVIEW-FINDING" not in message

    def test_finalize_carries_compliance_result(self):
        from src.orchestration.pipeline import finalize_batch_result

        module = _enabled_module()
        state = _state(module, requirements_profile=_profile().to_dict())
        state.compliance_result = ReviewResult(
            findings=[], cross_check_status="completed", coverage=[]
        )
        result = finalize_batch_result(state)
        assert result.compliance_result is state.compliance_result


# ---------------------------------------------------------------------------
# Report surfaces
# ---------------------------------------------------------------------------


def _doc_text(path) -> str:
    from docx import Document as _Document

    doc = _Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _pipeline_result(module, *, requirements_profile=None, compliance_result=None):
    from src.orchestration.pipeline import PipelineResult

    return PipelineResult(
        review_result=ReviewResult(findings=[]),
        files_reviewed=["21 13 13 Wet-Pipe.docx"],
        cycle_label=module.cycle.label,
        module_id=module.module_id,
        requirements_profile=requirements_profile,
        compliance_result=compliance_result,
    )


class TestReportSurfaces:
    def test_requirements_section_renders(self, tmp_path):
        from src.output.report_exporter import export_report

        module = _enabled_module()
        profile = _profile(
            [
                _item("r-aaaaaaaaaaaa", "The 2024 IBC as amended governs."),
                _item("r-bbbbbbbbbbbb", "Ungrounded requirement.", grounded=False),
                _item(
                    "r-cccccccccccc",
                    "Flow tests witnessed April-October only.",
                    actionability="process_advisory",
                    category="ahj_requirement",
                ),
            ]
        )
        compliance = ReviewResult(
            findings=[],
            cross_check_status="completed",
            coverage=[
                {"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                 "evidence": "Comply with the 2024 IBC.", "fileName": "21 13 13 Wet-Pipe.docx"},
                {"requirement_id": "r-bbbbbbbbbbbb", "status": "missing",
                 "evidence": None, "fileName": None},
            ],
        )
        out = tmp_path / "report.docx"
        export_report(
            _pipeline_result(
                module,
                requirements_profile=profile.to_dict(),
                compliance_result=compliance,
            ),
            out,
        )
        text = _doc_text(out)
        assert "Jurisdiction & Client Requirements" in text
        assert "Markham, Ontario, Canada" in text
        assert "[UNVERIFIED]" in text
        assert "Requirements Coverage" in text
        assert "Represented" in text and "MISSING" in text
        assert "Process & Schedule Advisories" in text
        assert "Adopted & Referenced Editions" in text
        assert "Local-Code Compliance" in text
        # Banner rows for both phases.
        assert "Location/client research" in text
        assert "Local-code compliance" in text

    def test_profile_less_report_renders_no_ws4_surfaces(self, tmp_path):
        from src.output.report_exporter import export_report

        out = tmp_path / "report.docx"
        export_report(_pipeline_result(DEFAULT_MODULE), out)
        text = _doc_text(out)
        assert "Jurisdiction & Client Requirements" not in text
        assert "Local-Code Compliance" not in text
        assert "Location/client research" not in text
        assert "Local-code compliance" not in text

    def test_skipped_compliance_renders_red_banner_and_section(self, tmp_path):
        from src.output.report_exporter import export_report

        module = _enabled_module()
        skipped = ReviewResult(
            findings=[],
            cross_check_status="skipped",
            thinking="requirements profile unavailable",
        )
        out = tmp_path / "report.docx"
        export_report(
            _pipeline_result(
                module,
                requirements_profile=_profile().to_dict(),
                compliance_result=skipped,
            ),
            out,
        )
        text = _doc_text(out)
        assert "Local-code compliance" in text
        assert "skipped" in text
        assert "NOT evaluated against the researched" in text


class TestEmptyDimensionReportSurface:
    """WS3 / B1: a 0-item completed research dimension renders a red honesty
    paragraph in the Jurisdiction & Client Requirements section."""

    def _profile_with_empty_dimension(self) -> RequirementsProfile:
        return RequirementsProfile(
            items=[_item("r-aaaaaaaaaaaa", "The 2024 IBC as amended governs.")],
            dimension_statuses=[
                DimensionStatus(
                    dimension_id="governing_codes", status="completed", item_count=1
                ),
                DimensionStatus(
                    dimension_id="accessibility", status="completed", item_count=0
                ),
            ],
            research_date="2026-07-14",
            project={"city": "Markham", "state_or_province": "ON", "country": "CA",
                     "client_name": "ExampleCo"},
        )

    def test_empty_dimension_renders_red_paragraph(self, tmp_path):
        from src.output.report_exporter import export_report

        out = tmp_path / "report.docx"
        export_report(
            _pipeline_result(
                _enabled_module(),
                requirements_profile=self._profile_with_empty_dimension().to_dict(),
            ),
            out,
        )
        text = _doc_text(out)
        assert (
            "1 research dimension(s) completed without finding any requirements "
            "(accessibility)" in text
        )
        assert "unverified, not confirmed-clean" in text
        assert "1 dimension(s) returned no items." in text

    def test_no_empty_dimensions_renders_no_note(self, tmp_path):
        from src.output.report_exporter import export_report

        profile = RequirementsProfile(
            items=[_item("r-aaaaaaaaaaaa", "The 2024 IBC as amended governs.")],
            dimension_statuses=[
                DimensionStatus(
                    dimension_id="governing_codes", status="completed", item_count=1
                ),
            ],
            research_date="2026-07-14",
            project={"city": "Markham", "state_or_province": "ON", "country": "CA",
                     "client_name": "ExampleCo"},
        )
        out = tmp_path / "report.docx"
        export_report(
            _pipeline_result(
                _enabled_module(), requirements_profile=profile.to_dict()
            ),
            out,
        )
        text = _doc_text(out)
        assert "completed without finding any requirements" not in text
        assert "returned no items" not in text


class TestWs6ReportSurfaces:
    """WS6 fixes on the requirements/compliance report surfaces."""

    def test_represented_evidence_cell_keeps_quote_and_file(self, tmp_path):
        from src.output.report_exporter import export_report

        compliance = ReviewResult(
            findings=[],
            cross_check_status="completed",
            coverage=[
                {"requirement_id": "r-aaaaaaaaaaaa", "status": "represented",
                 "evidence": "Comply with the 2024 IBC.",
                 "fileName": "21 13 13 Wet-Pipe.docx"},
                {"requirement_id": "r-bbbbbbbbbbbb", "status": "missing",
                 "evidence": None, "fileName": "21 13 13 Wet-Pipe.docx"},
            ],
        )
        out = tmp_path / "report.docx"
        export_report(
            _pipeline_result(
                _enabled_module(),
                requirements_profile=_profile().to_dict(),
                compliance_result=compliance,
            ),
            out,
        )
        text = _doc_text(out)
        assert "Comply with the 2024 IBC. — 21 13 13 Wet-Pipe.docx" in text
        # E2: a MISSING row never renders a bare filename as pseudo-evidence.
        assert "(not found in reviewed package)" in text

    def test_profile_section_headings_preserve_initialisms(self, tmp_path):
        from src.output.report_exporter import export_report
        from src.research.requirements_research import ResearchItem

        profile = _profile(
            [
                _item(
                    "r-aaaaaaaaaaaa",
                    "AHJ requires witnessed flow tests.",
                    category="ahj_requirement",
                ),
            ]
        )
        out = tmp_path / "report.docx"
        export_report(
            _pipeline_result(
                _enabled_module(), requirements_profile=profile.to_dict()
            ),
            out,
        )
        text = _doc_text(out)
        # E9: explicit display map, not str.title().
        assert "AHJ Requirements" in text
        assert "Ahj Requirements" not in text

    def test_profile_less_pinned_editions_note_is_byte_identical(self):
        """E4/E10 must not disturb the CA wording: exact string pin."""
        from src.core.code_cycles import CALIFORNIA_2025
        from src.output.report_exporter import _render_pinned_editions_note

        note = _render_pinned_editions_note(CALIFORNIA_2025, "California")
        assert note.startswith(
            "This review pinned the following standards editions per the "
            "2025 California cycle: NFPA 13 2025, as amended by California; "
        )
        assert note.endswith(
            "Findings referencing other editions should be reviewed for "
            "relevance to the current cycle."
        )
        assert "fallback" not in note

    def test_profile_governed_pinned_editions_note_uses_fallback_framing(self):
        from src.core.code_cycles import CALIFORNIA_2025
        from src.output.report_exporter import _render_pinned_editions_note

        note = _render_pinned_editions_note(
            CALIFORNIA_2025, "California", profile_governed=True
        )
        assert note.startswith(
            "This review carried the following model-code fallback editions "
            "as its baseline: "
        )
        assert "Jurisdiction & Client Requirements section above" in note
        assert "pinned the following standards editions per the" not in note

    def test_verification_scope_note_only_with_compliance(self, tmp_path):
        """E7: the clarifying sentence renders only when a compliance result
        exists, so the CA summary block stays byte-identical."""
        from docx import Document as _Document

        from src.output.report_exporter import _write_summary_table
        from src.verification.verifier import VerificationResult

        f = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="1",
            issue="x",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference="",
        )
        f.verification = VerificationResult(
            verdict="CONFIRMED", explanation="ok", grounded=True
        )
        comp_finding = Finding(
            severity="HIGH",
            fileName="a.docx",
            section="[Compliance]",
            issue="missing requirement",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference="",
        )
        comp_finding.verification = VerificationResult(
            verdict="CONFIRMED", explanation="ok", grounded=True
        )
        compliance = ReviewResult(
            findings=[comp_finding], cross_check_status="completed"
        )
        review = ReviewResult(findings=[f])

        with_comp = _Document()
        _write_summary_table(with_comp, review, None, compliance_result=compliance)
        with_text = "\n".join(p.text for p in with_comp.paragraphs)
        # Tally includes the compliance finding: 2 confirmed.
        assert "2 confirmed" in with_text
        assert (
            "Verification counts include review, cross-check, and compliance "
            "findings." in with_text
        )

        without = _Document()
        _write_summary_table(without, review, None)
        without_text = "\n".join(p.text for p in without.paragraphs)
        assert "1 confirmed" in without_text
        assert "Verification counts include" not in without_text


class TestPackageSubsetNote:
    """E1 prompt side: a routed-subset run carries the subset note even on
    the non-chunked path, so 'missing' is classified relative to the subset."""

    def test_package_subset_adds_note_on_non_chunked_path(
        self, fake_client, stub_compliance_tokens
    ):
        fake_client["route"] = _route_by_marker(
            {"Comply": [compliance_tool_use_response()]}
        )
        run_chunked_compliance_check(
            [_spec("Comply with the 2024 IBC as amended.")],
            _profile(),
            [],
            cycle=_enabled_module().cycle,
            package_subset=True,
        )
        (call,) = fake_client["client"].calls
        assert "one subset of a larger specification package" in (
            call["messages"][0]["content"]
        )

    def test_default_non_chunked_path_has_no_note(
        self, fake_client, stub_compliance_tokens
    ):
        fake_client["route"] = _route_by_marker(
            {"Comply": [compliance_tool_use_response()]}
        )
        run_chunked_compliance_check(
            [_spec("Comply with the 2024 IBC as amended.")],
            _profile(),
            [],
            cycle=_enabled_module().cycle,
        )
        (call,) = fake_client["client"].calls
        assert "one subset of a larger specification package" not in (
            call["messages"][0]["content"]
        )

    def test_sizing_counts_the_subset_note_near_the_limit(
        self, fake_client, monkeypatch
    ):
        """Codex review (PR #320): the size check must measure the SAME
        message the non-chunked path sends. A package_subset corpus whose
        noteless prompt is just under the cap must take the CHUNKED path
        (the with-note prompt is over), never the non-chunked branch that
        then hits run_compliance_check's own over-limit guard and skips."""
        monkeypatch.setattr(cc, "count_tokens", lambda text: len(text.split()))
        # Two specs per division: singleton-division chunks pool into
        # ``general`` (the ≥2-specs-per-chunk invariant), which would rebuild
        # the full corpus in one chunk and defeat the fixture.
        pad = "word " * 300
        specs = [
            _spec(pad + " DIV21SPEC alpha", "21 13 13 Wet.docx"),
            _spec(pad + " DIV21SPEC beta", "21 13 16 Dry.docx"),
            _spec(pad + " DIV22SPEC alpha", "22 11 13 Water.docx"),
            _spec(pad + " DIV22SPEC beta", "22 11 16 Piping.docx"),
        ]
        cycle = _enabled_module().cycle
        # Cap = exactly the noteless full-corpus size, so adding the ~20-word
        # subset note pushes the full corpus over while each single-spec
        # chunk (roughly half the corpus) still fits comfortably.
        noteless_full = cc._build_compliance_user_message(specs, _profile(), [])
        system_words = len(cc._compliance_system_prompt(cycle).split())
        monkeypatch.setattr(
            cc,
            "COMPLIANCE_RECOMMENDED_MAX",
            system_words + len(noteless_full.split()),
        )
        fake_client["route"] = _route_by_marker(
            {
                "DIV21SPEC": [compliance_tool_use_response()],
                "DIV22SPEC": [compliance_tool_use_response()],
            }
        )
        result = run_chunked_compliance_check(
            specs,
            _profile(),
            [],
            cycle=cycle,
            package_subset=True,
        )
        # Chunked (two per-division calls), completed — not the pre-fix
        # skipped outcome from the inner over-limit guard.
        assert result.cross_check_status == "completed"
        assert len(fake_client["client"].calls) == 2
