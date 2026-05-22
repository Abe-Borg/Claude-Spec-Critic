"""Chunk 8 tests: separate report deduplication from executable edit identity.

Plan section "Chunk 8 — Separate report deduplication from executable edit
identity". The chunk preserves the per-file pre-merge edit fields when
:func:`pipeline._deduplicate_findings` collapses findings across files, so
edit execution does not fan one representative's ``existingText`` /
``replacementText`` / ``anchorText`` / ``evidenceElementId`` /
``edit_proposal`` across files that may have differed.

Acceptance scenarios from the plan:

1. Two files have the same semantic issue but different exact text. Grouping
   produces one display group, but the edit planner uses different original
   edit text per file.
2. The same file has two similar findings in different paragraphs. Edit
   planner keeps the locations separate.
3. Representative-only legacy payload (no per-file originals recorded) does
   not auto-apply across files — auto-edit fires only on the representative's
   own file; other affected files are routed to manual review.
4. Non-edit REPORT_ONLY grouped findings still display correctly.

Plus the back-compat surface: resume payloads round-trip
``occurrence_originals``; pre-Chunk-8 payloads load with the field empty.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.editing.apply_edits import execute_edit_plan
from src.input.extractor import ParagraphMapping, extract_text
from src.orchestration.pipeline import (
    FindingGroup,
    FindingOccurrence,
    _deduplicate_findings,
    group_findings,
)
from src.orchestration.resume_state import deserialize_finding, serialize_finding
from src.review.reviewer import EditProposal, Finding, REPORT_ONLY_ACTION


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_finding(
    *,
    file_name: str = "spec.docx",
    issue: str = "Code reference uses outdated CBC edition.",
    section: str = "2.1",
    severity: str = "HIGH",
    confidence: float = 0.8,
    action: str = "EDIT",
    existing: str | None = "per CBC 2019",
    replacement: str | None = "per CBC 2025",
    code_ref: str | None = "CBC 2025",
    anchor: str | None = None,
    insert_pos: str | None = None,
    evidence_id: str | None = None,
) -> Finding:
    return Finding(
        severity=severity,
        fileName=file_name,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference=code_ref,
        confidence=confidence,
        anchorText=anchor,
        insertPosition=insert_pos,
        evidenceElementId=evidence_id,
    )


def _write_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


# ---------------------------------------------------------------------------
# _deduplicate_findings — occurrence_originals capture
# ---------------------------------------------------------------------------


class TestDedupCapturesOriginals:
    def test_singleton_finding_has_empty_occurrence_originals(self):
        f = _make_finding(file_name="a.docx")
        out = _deduplicate_findings([f])
        assert len(out) == 1
        # Singleton: the finding IS its own original; no separate list needed.
        assert out[0].occurrence_originals == []

    def test_merged_multifile_finding_retains_one_original_per_file(self):
        # Two findings with identical dedup identity but different filenames
        # collapse into one merged representative — the originals are kept.
        a = _make_finding(file_name="a.docx", existing="per CBC 2019", replacement="per CBC 2025")
        b = _make_finding(file_name="b.docx", existing="per CBC 2019", replacement="per CBC 2025")
        out = _deduplicate_findings([a, b])
        assert len(out) == 1
        merged = out[0]
        assert sorted(merged.affected_files) == ["a.docx", "b.docx"]
        assert len(merged.occurrence_originals) == 2
        files_in_originals = sorted(o.fileName for o in merged.occurrence_originals)
        assert files_in_originals == ["a.docx", "b.docx"]

    def test_originals_preserve_per_file_edit_text(self):
        # Same dedup key, but each original carries its own text. The merged
        # representative's text is the rep's; per-file text lives on originals.
        a = _make_finding(file_name="a.docx", existing="per CBC 2019", replacement="per CBC 2025")
        b = _make_finding(file_name="b.docx", existing="per CBC 2019", replacement="per CBC 2025")
        # Tweak b's anchor to verify per-file fields are not overwritten.
        b.anchorText = "near top of section"
        out = _deduplicate_findings([a, b])
        merged = out[0]
        by_file = {o.fileName: o for o in merged.occurrence_originals}
        assert by_file["a.docx"].anchorText is None
        assert by_file["b.docx"].anchorText == "near top of section"

    def test_originals_themselves_have_empty_occurrence_originals(self):
        # No recursive nesting — members are leaves.
        a = _make_finding(file_name="a.docx")
        b = _make_finding(file_name="b.docx")
        out = _deduplicate_findings([a, b])
        for orig in out[0].occurrence_originals:
            assert orig.occurrence_originals == []

    def test_case_only_existing_text_difference_preserves_per_file_originals(self):
        """Phase 4 / Step 4.2 regression: case/whitespace-only differences in
        ``existingText`` collapse into one dedup group, but each per-file
        original retains its UNMODIFIED text.

        The dedup key normalizes ``existingText`` and ``replacementText``
        via ``_normalized_text_digest`` (lowercase + strip), so two
        findings whose existing text differs only in case or trailing
        whitespace land in the same group. The premise that triggered
        this test (P9 from the auto-apply quality plan) was that the
        per-file text for the dedup loser would be lost; in practice
        ``_deduplicate_findings`` keeps every group member in
        ``occurrence_originals`` as the original ``Finding`` object,
        which still carries its pre-normalization text. This regression
        test locks in that behavior so a future refactor cannot silently
        re-introduce the bug the plan worried about.
        """
        # Case-only difference in existing text.
        a = _make_finding(
            file_name="a.docx",
            existing="Per CBC 2019",
            replacement="Per CBC 2025",
        )
        b = _make_finding(
            file_name="b.docx",
            existing="per cbc 2019",
            replacement="per cbc 2025",
        )
        # Trailing-whitespace-only difference layered on top.
        c = _make_finding(
            file_name="c.docx",
            existing="per CBC 2019   ",
            replacement="per CBC 2025\n",
        )

        out = _deduplicate_findings([a, b, c])

        # All three collapse to one representative (same dedup key after
        # normalization).
        assert len(out) == 1
        merged = out[0]
        assert sorted(merged.affected_files) == ["a.docx", "b.docx", "c.docx"]

        # Every per-file original is preserved with its ORIGINAL text
        # intact — not the normalized form, not the representative's
        # form. Edit execution can rely on each file's own text.
        by_file = {o.fileName: o for o in merged.occurrence_originals}
        assert by_file["a.docx"].existingText == "Per CBC 2019"
        assert by_file["a.docx"].replacementText == "Per CBC 2025"
        assert by_file["b.docx"].existingText == "per cbc 2019"
        assert by_file["b.docx"].replacementText == "per cbc 2025"
        assert by_file["c.docx"].existingText == "per CBC 2019   "
        assert by_file["c.docx"].replacementText == "per CBC 2025\n"


# ---------------------------------------------------------------------------
# group_findings — FindingOccurrence binds per-file original
# ---------------------------------------------------------------------------


class TestGroupFindingsPerFileOriginal:
    def test_singleton_occurrence_binds_representative_as_original(self):
        f = _make_finding(file_name="A.docx")
        groups = group_findings([f])
        assert len(groups) == 1
        occ = groups[0].occurrences[0]
        # Singleton: the representative is the only original (Chunk 8
        # backfill so executable_finding() returns the right thing).
        assert occ.has_original()
        assert occ.original_finding is f
        assert occ.executable_finding() is f

    def test_merged_finding_binds_each_file_to_its_own_original(self):
        a = _make_finding(file_name="a.docx", existing="per CBC 2019", replacement="per CBC 2025")
        b = _make_finding(file_name="b.docx", existing="per CBC 2019", replacement="per CBC 2025")
        merged = _deduplicate_findings([a, b])[0]
        groups = group_findings([merged])
        occs_by_file = {o.file_name: o for o in groups[0].occurrences}
        # Each file's occurrence binds to the original that came from that file.
        assert occs_by_file["a.docx"].executable_finding().fileName == "a.docx"
        assert occs_by_file["b.docx"].executable_finding().fileName == "b.docx"
        # The representative is shared (display layer) — but the executable
        # finding is the per-file original, not the representative.
        assert occs_by_file["a.docx"].finding is merged
        assert occs_by_file["b.docx"].finding is merged
        assert occs_by_file["a.docx"].executable_finding() is not merged
        assert occs_by_file["b.docx"].executable_finding() is not merged

    def test_legacy_multifile_finding_without_originals_marks_missing(self):
        # Pre-Chunk-8 payload: a finding with multiple affected files but no
        # ``occurrence_originals``. The representative's own file binds to the
        # representative; other affected files have no original (manual
        # review path in apply_edits).
        f = _make_finding(file_name="a.docx")
        f.affected_files = ["a.docx", "b.docx"]
        # ``occurrence_originals`` left empty.
        groups = group_findings([f])
        occs_by_file = {o.file_name: o for o in groups[0].occurrences}
        assert occs_by_file["a.docx"].has_original()
        assert occs_by_file["a.docx"].original_finding is f
        assert not occs_by_file["b.docx"].has_original()
        assert occs_by_file["b.docx"].executable_finding() is f  # fallback only

# ---------------------------------------------------------------------------
# apply_edits.execute_edit_plan — uses per-file originals
# ---------------------------------------------------------------------------


class TestExecuteEditPlanUsesPerFileOriginals:
    def test_case_only_dedup_executor_uses_each_files_original_text(self, tmp_path: Path):
        """Phase 4 / Step 4.2 regression: end-to-end coverage for the P9
        scenario. Two files have ``existingText`` that differs only in
        case. They share a dedup key (normalization collapses them) but
        the executor must use each file's original text — which is the
        only thing that will actually match the source paragraph.
        """
        file_a = tmp_path / "a.docx"
        file_b = tmp_path / "b.docx"
        # Source paragraphs in mixed case. The executor needs each
        # file's per-original existingText to locate the right span.
        _write_docx(file_a, "Per CBC 2019 fire-rated assemblies.")
        _write_docx(file_b, "per cbc 2019 fire-rated assemblies.")
        spec_a = extract_text(file_a)
        spec_a.filename = "a.docx"
        spec_b = extract_text(file_b)
        spec_b.filename = "b.docx"

        # Two findings, same dedup key after normalization, different
        # case in existingText. The dedup helper should collapse them
        # AND preserve both per-file originals with their original text
        # intact so the executor's locator can find the match.
        a = _make_finding(
            file_name="a.docx",
            existing="Per CBC 2019",
            replacement="Per CBC 2025",
            code_ref="CBC 2025",
        )
        b = _make_finding(
            file_name="b.docx",
            existing="per cbc 2019",
            replacement="per cbc 2025",
            code_ref="CBC 2025",
        )

        deduped = _deduplicate_findings([a, b])
        assert len(deduped) == 1
        merged = deduped[0]
        # Both files captured; each per-file original retained its
        # original (case-different) existingText.
        assert sorted(merged.affected_files) == ["a.docx", "b.docx"]
        by_file = {o.fileName: o for o in merged.occurrence_originals}
        assert by_file["a.docx"].existingText == "Per CBC 2019"
        assert by_file["b.docx"].existingText == "per cbc 2019"

        reports = execute_edit_plan(
            selected_finding_indices=[0],
            all_findings=deduped,
            cross_check_findings=[],
            extracted_specs=[spec_a, spec_b],
            source_paths=[file_a, file_b],
            output_dir=tmp_path / "out",
        )
        assert len(reports) == 2
        # Both files edited cleanly via their own per-file originals.
        for r in reports:
            assert r.edits_applied == 1
        outputs = {r.source_path.name: r.output_path for r in reports}
        assert Document(outputs["a.docx"]).paragraphs[0].text == "Per CBC 2025 fire-rated assemblies."
        assert Document(outputs["b.docx"]).paragraphs[0].text == "per cbc 2025 fire-rated assemblies."

    def test_two_files_different_exact_text_use_their_own_text(self, tmp_path: Path):
        # Acceptance scenario 1 from the chunk: the merged finding's display
        # is one row, but the executor uses each file's own existingText.
        file_a = tmp_path / "a.docx"
        file_b = tmp_path / "b.docx"
        # Different exact text per file — would mis-apply if the executor
        # blindly used the representative's text.
        _write_docx(file_a, "Comply with ASCE 7-16 wind requirements.")
        _write_docx(file_b, "Comply with ASCE 7-22 wind requirements.")
        spec_a = extract_text(file_a)
        spec_a.filename = "a.docx"
        spec_b = extract_text(file_b)
        spec_b.filename = "b.docx"

        # Two findings with the same dedup key (same normalized issue,
        # section, code ref, action) but different per-file existingText.
        # Use the same replacement text and existingText hash so the dedup
        # key collapses both. To trigger that we use the same exact text;
        # test the executor's per-file routing by varying the *original's*
        # ``existingText`` after dedup populates ``occurrence_originals``.
        merged = _make_finding(
            file_name="a.docx",
            existing="ASCE 7-16",
            replacement="ASCE 7-25",
            code_ref="ASCE 7",
        )
        merged.affected_files = ["a.docx", "b.docx"]
        # Hand-built originals carrying file-specific edit text. This mirrors
        # what _deduplicate_findings would produce when the originals were
        # close in dedup key but had divergent existingText (e.g., a
        # normalization that squashed the version difference).
        orig_a = _make_finding(
            file_name="a.docx",
            existing="ASCE 7-16",
            replacement="ASCE 7-25",
        )
        orig_b = _make_finding(
            file_name="b.docx",
            existing="ASCE 7-22",
            replacement="ASCE 7-25",
        )
        merged.occurrence_originals = [orig_a, orig_b]

        reports = execute_edit_plan(
            selected_finding_indices=[0],
            all_findings=[merged],
            cross_check_findings=[],
            extracted_specs=[spec_a, spec_b],
            source_paths=[file_a, file_b],
            output_dir=tmp_path / "out",
        )
        assert len(reports) == 2
        # Both files were edited (each using its own existingText).
        for r in reports:
            assert r.edits_applied == 1
        # Verify the actual replacements landed correctly.
        outputs = {r.source_path.name: r.output_path for r in reports}
        assert Document(outputs["a.docx"]).paragraphs[0].text == "Comply with ASCE 7-25 wind requirements."
        assert Document(outputs["b.docx"]).paragraphs[0].text == "Comply with ASCE 7-25 wind requirements."

    def test_legacy_payload_only_edits_representative_file_others_manual_review(self, tmp_path: Path):
        # Acceptance scenario 3: representative-only legacy payload does NOT
        # auto-apply across files. The representative's own file edits; the
        # other affected file is routed to manual review.
        file_a = tmp_path / "a.docx"
        file_b = tmp_path / "b.docx"
        # Both happen to contain the representative's text — a real-world
        # legacy payload could not distinguish, so executor falls back to
        # safety routing.
        _write_docx(file_a, "Comply with ASCE 7-16 wind requirements.")
        _write_docx(file_b, "Comply with ASCE 7-16 wind requirements.")
        spec_a = extract_text(file_a)
        spec_a.filename = "a.docx"
        spec_b = extract_text(file_b)
        spec_b.filename = "b.docx"

        legacy = _make_finding(
            file_name="a.docx",
            existing="ASCE 7-16",
            replacement="ASCE 7-25",
        )
        legacy.affected_files = ["a.docx", "b.docx"]
        # NOTE: occurrence_originals deliberately NOT populated — the legacy
        # payload predates Chunk 8.
        assert legacy.occurrence_originals == []

        warnings: list[str] = []
        reports = execute_edit_plan(
            selected_finding_indices=[0],
            all_findings=[legacy],
            cross_check_findings=[],
            extracted_specs=[spec_a, spec_b],
            source_paths=[file_a, file_b],
            output_dir=tmp_path / "out",
            log=lambda msg: warnings.append(msg),
        )
        # Two reports: a.docx with auto-edit applied, b.docx as manual review.
        assert len(reports) == 2
        by_name = {r.source_path.name: r for r in reports}
        assert by_name["a.docx"].edits_applied == 1
        assert by_name["b.docx"].edits_applied == 0
        assert by_name["b.docx"].edits_skipped == 1
        assert any("manual review" in w.lower() for w in by_name["b.docx"].warnings)
        # Log surfaces the per-file safety routing decision.
        assert any("manual review" in m.lower() and "b.docx" in m for m in warnings)

    def test_singleton_finding_still_auto_edits(self, tmp_path: Path):
        # Backward-compat: a singleton finding (one file) goes through the
        # representative-IS-original fallback and edits as before.
        file_a = tmp_path / "a.docx"
        _write_docx(file_a, "Use CBC 2019 references.")
        spec_a = extract_text(file_a)
        spec_a.filename = "a.docx"

        f = _make_finding(
            file_name="a.docx",
            existing="CBC 2019",
            replacement="CBC 2025",
        )
        # No multi-file fan-out here: the executor binds the rep as its own
        # original because file_a is the rep's own file.
        reports = execute_edit_plan(
            selected_finding_indices=[0],
            all_findings=[f],
            cross_check_findings=[],
            extracted_specs=[spec_a],
            source_paths=[file_a],
            output_dir=tmp_path / "out",
        )
        assert len(reports) == 1
        assert reports[0].edits_applied == 1


# ---------------------------------------------------------------------------
# Same-file findings in different paragraphs stay separate
# ---------------------------------------------------------------------------


class TestSameFileDifferentParagraphsStaySeparate:
    def test_two_findings_same_file_different_text_not_merged(self):
        # Acceptance scenario 2: same file, similar issue text, different
        # exact existingText — must not be collapsed by dedup (so the edit
        # planner keeps the locations separate).
        f1 = _make_finding(
            file_name="a.docx",
            issue="Outdated code edition reference.",
            existing="per CBC 2019",
            replacement="per CBC 2025",
        )
        f2 = _make_finding(
            file_name="a.docx",
            issue="Outdated code edition reference.",
            existing="per CMC 2019",
            replacement="per CMC 2025",
        )
        out = _deduplicate_findings([f1, f2])
        # Different existingText → different dedup key → not merged.
        assert len(out) == 2
        assert {f.existingText for f in out} == {"per CBC 2019", "per CMC 2019"}


# ---------------------------------------------------------------------------
# REPORT_ONLY grouped findings still display correctly
# ---------------------------------------------------------------------------


class TestReportOnlyGroupedFindings:
    def test_report_only_finding_merges_and_stays_report_only(self):
        # Acceptance scenario 4: REPORT_ONLY findings can group across files
        # without rehydrating an edit slot. The merged representative stays
        # REPORT_ONLY, the per-file originals stay REPORT_ONLY, no edit
        # proposal materializes.
        a = _make_finding(
            file_name="a.docx",
            issue="Coordination conflict — sleeve sizing not coordinated with structural.",
            action=REPORT_ONLY_ACTION,
            existing=None,
            replacement=None,
        )
        b = _make_finding(
            file_name="b.docx",
            issue="Coordination conflict — sleeve sizing not coordinated with structural.",
            action=REPORT_ONLY_ACTION,
            existing=None,
            replacement=None,
        )
        out = _deduplicate_findings([a, b])
        assert len(out) == 1
        merged = out[0]
        assert merged.actionType == REPORT_ONLY_ACTION
        assert merged.as_edit_proposal() is None
        assert sorted(merged.affected_files) == ["a.docx", "b.docx"]
        assert len(merged.occurrence_originals) == 2
        for orig in merged.occurrence_originals:
            assert orig.actionType == REPORT_ONLY_ACTION
            assert orig.as_edit_proposal() is None

    def test_report_only_grouped_finding_renders_through_group_findings(self):
        # group_findings should still produce occurrences for REPORT_ONLY
        # findings so the report can list affected files even though no
        # edit will be executed.
        a = _make_finding(
            file_name="a.docx",
            action=REPORT_ONLY_ACTION,
            existing=None,
            replacement=None,
        )
        b = _make_finding(
            file_name="b.docx",
            action=REPORT_ONLY_ACTION,
            existing=None,
            replacement=None,
        )
        merged = _deduplicate_findings([a, b])[0]
        groups = group_findings([merged])
        names = [o.file_name for o in groups[0].occurrences]
        assert sorted(names) == ["a.docx", "b.docx"]


# ---------------------------------------------------------------------------
# Resume state round-trip
# ---------------------------------------------------------------------------


class TestResumeStateRoundTrip:
    def test_round_trip_preserves_occurrence_originals(self):
        # Same existingText/replacementText (so the dedup key collapses) but
        # different per-file anchor — anchor isn't part of the dedup key, so
        # this is a realistic shape for what occurrence_originals captures.
        a = _make_finding(file_name="a.docx", existing="ASCE 7-16", replacement="ASCE 7-25")
        a.anchorText = "near top of section 1"
        b = _make_finding(file_name="b.docx", existing="ASCE 7-16", replacement="ASCE 7-25")
        b.anchorText = "near top of section 2"
        merged = _deduplicate_findings([a, b])[0]
        payload = serialize_finding(merged)
        # The serialized form contains the originals.
        assert "occurrence_originals" in payload
        assert len(payload["occurrence_originals"]) == 2
        # Round-trip preserves per-file edit metadata.
        loaded = deserialize_finding(payload)
        by_file = {o.fileName: o for o in loaded.occurrence_originals}
        assert by_file["a.docx"].existingText == "ASCE 7-16"
        assert by_file["b.docx"].existingText == "ASCE 7-16"
        assert by_file["a.docx"].anchorText == "near top of section 1"
        assert by_file["b.docx"].anchorText == "near top of section 2"

    def test_legacy_payload_without_originals_loads_with_empty_list(self):
        # Pre-Chunk-8 payload: no ``occurrence_originals`` key at all.
        legacy_payload = {
            "severity": "HIGH",
            "fileName": "a.docx",
            "section": "2.1",
            "issue": "outdated reference",
            "actionType": "EDIT",
            "existingText": "old",
            "replacementText": "new",
            "codeReference": "CBC",
            "confidence": 0.7,
            "affected_files": ["a.docx", "b.docx"],
        }
        loaded = deserialize_finding(legacy_payload)
        assert loaded.occurrence_originals == []
        # Backward compatibility: the legacy fields still load.
        assert loaded.affected_files == ["a.docx", "b.docx"]

    def test_nested_originals_are_not_serialized_recursively(self):
        # Defense in depth: even if a member somehow carried its own
        # occurrence_originals, the second-level serialization clears it so
        # the JSON cannot grow without bound.
        member_with_nested = _make_finding(file_name="a.docx")
        member_with_nested.occurrence_originals = [_make_finding(file_name="a.docx")]
        merged = _make_finding(file_name="a.docx")
        merged.occurrence_originals = [member_with_nested]
        payload = serialize_finding(merged)
        nested = payload["occurrence_originals"][0]
        # The nested member's own occurrence_originals is forced empty.
        assert nested["occurrence_originals"] == []
