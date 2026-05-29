"""Tests for the Run Diagnostics banner.

A styled-table banner right after the title block surfaces operational
health at-a-glance:

* Edit-suggested / Report-only counts (from the
  edit-action histogram already computed for the trust-model summary).
* Cache replays with the oldest entry age (using
  ``cache_entry_created_ts``).
* Verification failures (the ``VERIFICATION_FAILED`` status),
  highlighted red when > 0.
* REPORT_ONLY demotions at parse time (the ``demotion_reason``).
* Spec content extraction warnings (slot reserved for the content-loss
  warning; renders 0 on every run until that lands).
* Cross-spec coordination status — skipped / failed / completed.

A failure recovery hint paragraph appears below the table whenever the
verification-failure count is non-zero so a reviewer can re-run only
those findings.

The plan's success criteria:

* A clean run shows the banner with all counts at expected values.
* A run with simulated verification failures shows red callouts.
* Cross-spec coordination skipped/failed counts appear in the banner.
"""
from __future__ import annotations

import time
from pathlib import Path

import pytest
from docx import Document

from src.output.report_exporter import (
    _summarize_run_diagnostics,
    export_report,
)
from src.output.report_status import (
    summarize_edit_actions,
    summarize_statuses,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.8,
    action: str = "EDIT",
    existing: str | None = "2019 CBC",
    replacement: str | None = "2025 CBC",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
    demotion_reason: str | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
        demotion_reason=demotion_reason,
    )
    f.verification = verification
    return f


def _verified_supported() -> VerificationResult:
    return VerificationResult(
        verdict="CONFIRMED",
        explanation="Verified against CBC §1234.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        accepted_sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
        cache_status="miss",
        source_quote="The 2025 CBC adopts the 2024 IBC.",
        model_used="claude-sonnet-4-6",
        verification_mode="standard_reasoning",
        web_search_requests=3,
    )


def _failed_verification() -> VerificationResult:
    return VerificationResult(
        verdict="UNVERIFIED",
        explanation="Server overloaded during verification: 529",
        grounded=False,
        verification_failed=True,
    )


def _cache_hit_result(*, age_days: int = 5) -> VerificationResult:
    created_ts = time.time() - (age_days * 86400)
    return VerificationResult(
        verdict="CONFIRMED",
        explanation="Cached verdict from prior run.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        accepted_sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
        model_used="claude-sonnet-4-6",
        cache_status="hit",
        source_quote="The 2025 CBC adopts the 2024 IBC.",
        cache_entry_created_ts=created_ts,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
        cross_check_result: ReviewResult | None = None,
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or (
            [review_result.findings[0].fileName] if review_result.findings else ["test.docx"]
        )
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _findings_to_summary(
    findings: list[Finding],
    *,
    cross_check_result: ReviewResult | None = None,
    pipeline_result=None,
) -> dict:
    """Convenience: drive _summarize_run_diagnostics with the same
    inputs export_report would feed it."""
    status_counts = summarize_statuses(findings)
    edit_action_counts = summarize_edit_actions(findings)
    return _summarize_run_diagnostics(
        findings=findings,
        status_counts=status_counts,
        edit_action_counts=edit_action_counts,
        cross_check_result=cross_check_result,
        pipeline_result=pipeline_result,
    )


# ---------------------------------------------------------------------------
# 1. _summarize_run_diagnostics — pure function over findings + counts
# ---------------------------------------------------------------------------


class TestSummarizeRunDiagnostics:
    def test_empty_run_returns_zero_counts(self):
        summary = _findings_to_summary([])
        assert summary["edit_suggested"] == 0
        assert summary["report_only"] == 0
        assert summary["verification_failed"] == 0
        assert summary["cache_replay_count"] == 0
        assert summary["oldest_cache_age_days"] is None
        assert summary["demotion_count"] == 0
        assert summary["extraction_warning_count"] == 0
        assert summary["cross_check"] is None

    def test_edit_suggested_count_reflects_proposal(self):
        # Any finding carrying an edit proposal is EDIT_SUGGESTED — the
        # app emits the instruction without gating on confidence.
        f = _finding(verification=_verified_supported(), confidence=0.9)
        summary = _findings_to_summary([f])
        assert summary["edit_suggested"] == 1

    def test_report_only_count_reflects_no_proposal(self):
        f = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verified_supported(),
        )
        summary = _findings_to_summary([f])
        assert summary["report_only"] == 1

    def test_verification_failed_count_uses_status_histogram(self):
        f = _finding(verification=_failed_verification())
        summary = _findings_to_summary([f])
        assert summary["verification_failed"] == 1

    def test_cache_replay_counts_only_hit_findings(self):
        hit = _finding(verification=_cache_hit_result(age_days=5))
        miss = _finding(
            file="Section_22_2000.docx",
            verification=_verified_supported(),
        )
        summary = _findings_to_summary([hit, miss])
        assert summary["cache_replay_count"] == 1
        assert summary["oldest_cache_age_days"] == 5

    def test_oldest_cache_age_is_maximum(self):
        # Among several cache hits, the oldest age wins so a reviewer
        # sees the worst-case staleness in the banner.
        ages = [5, 45, 120]
        findings = [
            _finding(
                file=f"Section_22_{i * 1000}.docx",
                verification=_cache_hit_result(age_days=age),
            )
            for i, age in enumerate(ages, start=1)
        ]
        summary = _findings_to_summary(findings)
        assert summary["cache_replay_count"] == 3
        assert summary["oldest_cache_age_days"] == 120

    def test_legacy_cache_hit_counts_but_no_age(self):
        # A cache_status="hit" with cache_entry_created_ts=0.0 (legacy
        # resume payload predating cache-age tracking) counts toward the cache
        # replay total but cannot contribute to the oldest-age display.
        # Note: _enforce_grounding_invariant downgrades verdicts where
        # accepted citation chain is missing, but for a legacy payload
        # the cache_status stays "hit" on the verification result.
        legacy = VerificationResult(
            verdict="CONFIRMED",
            sources=["https://x"],
            accepted_sources=["https://x"],
            grounded=True,
            cache_status="hit",
            source_quote="snippet",
            cache_entry_created_ts=0.0,
        )
        f = _finding(verification=legacy)
        summary = _findings_to_summary([f])
        assert summary["cache_replay_count"] == 1
        assert summary["oldest_cache_age_days"] is None

    def test_demotion_count_reflects_demotion_reason_field(self):
        # Findings with a parse-time demotion are the ones surfaced in
        # this row — they signal model-output shape issues, not a
        # deliberate REPORT_ONLY.
        f = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            demotion_reason="EDIT requested but existingText was empty",
        )
        summary = _findings_to_summary([f])
        assert summary["demotion_count"] == 1

    def test_demotion_count_ignores_empty_or_whitespace_reasons(self):
        # An empty string or pure whitespace must not be counted as a
        # demotion — those represent the "no reason recorded" default.
        f1 = _finding(action="REPORT_ONLY", existing=None, replacement=None, demotion_reason="")
        f2 = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            file="Other.docx",
            demotion_reason="   ",
        )
        summary = _findings_to_summary([f1, f2])
        assert summary["demotion_count"] == 0

    def test_cross_check_state_completed(self):
        cc = ReviewResult(findings=[], cross_check_status="completed")
        summary = _findings_to_summary([], cross_check_result=cc)
        assert summary["cross_check"] is not None
        assert summary["cross_check"]["status"] == "completed"
        assert summary["cross_check"]["finding_count"] == 0

    def test_cross_check_state_skipped(self):
        cc = ReviewResult(
            findings=[],
            cross_check_status="skipped",
            thinking="not enough specs",
        )
        summary = _findings_to_summary([], cross_check_result=cc)
        assert summary["cross_check"]["status"] == "skipped"
        assert summary["cross_check"]["reason"] == "not enough specs"

    def test_cross_check_state_failed(self):
        cc = ReviewResult(
            findings=[],
            cross_check_status="failed",
            error="API timeout",
        )
        summary = _findings_to_summary([], cross_check_result=cc)
        assert summary["cross_check"]["status"] == "failed"
        assert summary["cross_check"]["reason"] == "API timeout"

    def test_cross_check_none_means_not_run(self):
        # No cross-check requested → no cross-check row in the banner.
        # The summary records None so the renderer can suppress the row.
        summary = _findings_to_summary([], cross_check_result=None)
        assert summary["cross_check"] is None


# ---------------------------------------------------------------------------
# 2. Banner rendering — section heading + labels
# ---------------------------------------------------------------------------


class TestBannerRendering:
    def test_banner_heading_appears_in_report(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Run Diagnostics" in text


# ---------------------------------------------------------------------------
# 3. Banner content — counts reflect the input findings
# ---------------------------------------------------------------------------


class TestBannerCounts:
    @pytest.mark.parametrize(
        "n_failures, expected_phrase",
        [
            (1, "1 finding failed verification"),
            (2, "2 findings failed verification"),
        ],
    )
    def test_failure_hint_count_and_pluralization(
        self, tmp_path: Path, n_failures: int, expected_phrase: str
    ):
        # The failure-hint paragraph copies the failure count into its
        # prose and pluralizes "finding" vs "findings" accordingly. One
        # supported finding rides along for the singular case so the
        # banner also reports a non-failure.
        findings = [
            _finding(file=f"Failed_{i}.docx", verification=_failed_verification())
            for i in range(n_failures)
        ]
        if n_failures == 1:
            findings.append(_finding(file="OK.docx", verification=_verified_supported()))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=findings)),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Verification failures (operational)" in text
        assert expected_phrase in text


# ---------------------------------------------------------------------------
# 4. Cross-check status in the banner
# ---------------------------------------------------------------------------


class TestBannerCrossCheckStatus:
    def test_cross_check_none_omits_row(self, tmp_path: Path):
        # When cross-check wasn't run, no row is rendered (the
        # _StubPipelineResult passes None by default). The renderer's
        # row-suppression branch — the helper class covers the populated
        # states.
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Cross-spec coordination" not in text


# ---------------------------------------------------------------------------
# 5. Visual highlight — verification failures + extraction warnings
# ---------------------------------------------------------------------------


class TestBannerHighlights:
    """The plan calls for verification failures to highlight in red
    when > 0 — we verify by inspecting cell shading on the value
    column. Same treatment applies to extraction warnings."""

    def _value_cell_shading_for_label(self, doc: Document, label: str) -> str | None:
        """Walk every table in the doc; return the value-cell shading
        for the row whose first cell text matches ``label``.

        Returns the hex string (e.g., "FFE5E5") or ``None`` if the row
        has no shading element on its value cell.
        """
        from docx.oxml.ns import qn

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                if row.cells[0].text.strip() != label:
                    continue
                value_cell = row.cells[1]
                tcPr = value_cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    return None
                shd = tcPr.find(qn("w:shd"))
                if shd is None:
                    return None
                return shd.get(qn("w:fill"))
        return None

    def test_verification_failure_value_cell_is_red_when_nonzero(
        self, tmp_path: Path
    ):
        f = _finding(verification=_failed_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        doc = Document(str(out))
        shading = self._value_cell_shading_for_label(
            doc, "Verification failures (operational)"
        )
        assert shading is not None
        # The light-red shading we apply is FFE5E5; the exact value is
        # an implementation detail but it must be a red-family hex.
        assert shading.upper() == "FFE5E5"

    def test_verification_failure_value_cell_unshaded_when_zero(
        self, tmp_path: Path
    ):
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        doc = Document(str(out))
        shading = self._value_cell_shading_for_label(
            doc, "Verification failures (operational)"
        )
        # No shading element means we did not apply the highlight.
        assert shading is None
