"""Phase 4 regression tests: edit-safety categories, revalidation, and ordering.

Covers audit Sections 8.1, 8.3, 8.4 plus the existing 8.7 invariant
(source documents are never overwritten).
"""

from pathlib import Path

import pytest
from docx import Document

from src.edit_candidates import (
    SAFETY_AUTO_SAFE,
    SAFETY_AUTO_WITH_CAUTION,
    SAFETY_MANUAL_REVIEW,
    SAFETY_REPORT_ONLY,
    classify_edit_candidates,
)
from src.edit_locator import EditLocation, LocatorResult, locate_edit
from src.extractor import ParagraphMapping, extract_text_from_docx
from src.reviewer import Finding
from src.spec_editor import (
    EditAction,
    apply_edits_to_spec,
    build_edit_actions,
)
from src.verifier import VerificationResult


def _mapping(text: str, *, idx: int, element_type: str = "paragraph") -> ParagraphMapping:
    return ParagraphMapping(
        body_index=idx,
        element_type=element_type,
        text=text,
        table_index=0 if element_type == "table_cell" else None,
        row_index=0 if element_type == "table_cell" else None,
        cell_index=None,
    )


def _finding(
    *,
    action: str = "EDIT",
    existing: str = "Provide seismic bracing per ASCE 7-16.",
    replacement: str | None = "Provide seismic bracing per ASCE 7-22.",
    verdict: str | None = "CONFIRMED",
) -> Finding:
    f = Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1.0",
        issue="Issue",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC",
        confidence=0.9,
    )
    if verdict is not None:
        f.verification = VerificationResult(
            verdict=verdict, explanation="", sources=[], correction=None
        )
    return f


# -- Section 8.1: candidate-level safety categories ---------------------------


def test_candidate_confirmed_classified_auto_safe():
    candidates = classify_edit_candidates([_finding(verdict="CONFIRMED")])
    assert candidates[0].safety_category == SAFETY_AUTO_SAFE


def test_candidate_unverified_classified_auto_with_caution():
    candidates = classify_edit_candidates([_finding(verdict="UNVERIFIED")])
    cand = candidates[0]
    assert cand.safety_category == SAFETY_AUTO_WITH_CAUTION
    assert cand.eligible is True
    assert cand.default_selected is False


def test_candidate_disputed_classified_report_only():
    candidates = classify_edit_candidates([_finding(verdict="DISPUTED")])
    assert candidates[0].safety_category == SAFETY_REPORT_ONLY


def test_candidate_no_anchor_classified_report_only():
    candidates = classify_edit_candidates([_finding(existing="  ", verdict="CONFIRMED")])
    assert candidates[0].safety_category == SAFETY_REPORT_ONLY


# -- Section 8.1: locator-level safety categories ----------------------------


def test_locator_exact_paragraph_match_is_auto_safe():
    paragraph_map = [_mapping("Install copper piping with Type L wall thickness.", idx=0)]
    result = locate_edit(_finding(existing="Type L wall thickness"), paragraph_map)
    assert result.status == "matched"
    assert result.safety_category == SAFETY_AUTO_SAFE


def test_locator_section_anchored_match_downgrades_to_caution():
    paragraph_map = [
        _mapping("1.0 GENERAL", idx=0),
        _mapping("Voltage shall be 208V", idx=1),
        _mapping("2.0 PRODUCTS", idx=2),
        _mapping("Voltage shall be 480V", idx=3),
    ]
    finding = _finding(existing="208V")
    finding.section = "1.0"
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.safety_category == SAFETY_AUTO_WITH_CAUTION


def test_locator_fuzzy_match_classified_manual_review():
    paragraph_map = [_mapping("Provide 2-hour fire-resistance rated shaft enclosures.", idx=0)]
    result = locate_edit(_finding(existing="Provide two hour fire resistance shaft enclosure."), paragraph_map)
    assert result.status == "matched"
    assert result.safety_category == SAFETY_MANUAL_REVIEW


def test_locator_ambiguous_match_classified_manual_review():
    paragraph_map = [
        _mapping("Use non-shrink grout at equipment bases.", idx=0),
        _mapping("Use non-shrink grout at equipment bases and supports.", idx=1),
    ]
    result = locate_edit(_finding(existing="Use non-shrink grout at equipment bases"), paragraph_map)
    assert result.status == "ambiguous"
    assert result.safety_category == SAFETY_MANUAL_REVIEW


def test_locator_table_cell_match_downgrades_to_caution():
    paragraph_map = [_mapping("R1C1 | Allowance Amount | $10,000", idx=0, element_type="table_cell")]
    result = locate_edit(_finding(existing="Allowance Amount"), paragraph_map)
    assert result.status == "matched"
    assert result.safety_category == SAFETY_AUTO_WITH_CAUTION


def test_locator_cross_paragraph_match_downgrades_to_caution():
    paragraph_map = [
        _mapping("PART 1 - GENERAL", idx=0),
        _mapping("Provide submittals within ten days.", idx=1),
        _mapping("Submit operation and maintenance manuals.", idx=2),
    ]
    finding = _finding(
        existing="Provide submittals within ten days.\n\nSubmit operation and maintenance manuals."
    )
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.safety_category == SAFETY_AUTO_WITH_CAUTION


def test_locator_not_found_classified_report_only():
    paragraph_map = [_mapping("Unrelated text", idx=0)]
    result = locate_edit(_finding(existing="Completely absent sentence"), paragraph_map)
    assert result.status == "not_found"
    assert result.safety_category == SAFETY_REPORT_ONLY


# -- Section 8.1: build_edit_actions gating ----------------------------------


def _locator_with_category(
    *,
    text: str,
    category: str,
    action: str = "EDIT",
    matched: str | None = None,
    replacement: str = "new",
    status: str = "matched",
) -> LocatorResult:
    matched_text = matched or text
    mapping = _mapping(text, idx=0)
    location = EditLocation(
        mapping=mapping,
        match_start=0,
        match_end=len(matched_text),
        matched_text=matched_text,
        match_confidence=1.0,
        match_method="exact",
    )
    return LocatorResult(
        finding=_finding(action=action, existing=matched_text, replacement=replacement),
        status=status,
        locations=[location],
        replacement_text=replacement,
        action_type=action,
        warning=None,
        safety_category=category,
    )


def test_build_edit_actions_skips_manual_review_category():
    result = _locator_with_category(text="Hello world", category=SAFETY_MANUAL_REVIEW)
    assert build_edit_actions([result]) == []
    assert result.warning is not None
    assert "manual review" in result.warning.lower()


def test_build_edit_actions_skips_report_only_category():
    result = _locator_with_category(text="Hello world", category=SAFETY_REPORT_ONLY)
    assert build_edit_actions([result]) == []


def test_build_edit_actions_allows_caution_when_flag_true():
    result = _locator_with_category(text="Hello world", category=SAFETY_AUTO_WITH_CAUTION)
    assert len(build_edit_actions([result], allow_caution=True)) == 1


def test_build_edit_actions_blocks_caution_when_flag_false():
    result = _locator_with_category(text="Hello world", category=SAFETY_AUTO_WITH_CAUTION)
    assert build_edit_actions([result], allow_caution=False) == []
    assert result.warning is not None


def test_build_edit_actions_includes_auto_safe():
    result = _locator_with_category(text="Hello world", category=SAFETY_AUTO_SAFE)
    assert len(build_edit_actions([result])) == 1


# -- Section 8.4: structural ordering ----------------------------------------


def test_whole_paragraph_delete_runs_after_add(tmp_path: Path):
    """ADD anchored after a paragraph that will also be DELETEd in same pass."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Anchor paragraph for ADD")
    doc.add_paragraph("Delete me paragraph")
    doc.save(source)

    spec = extract_text_from_docx(source)
    paragraph_map = spec.paragraph_map

    anchor_mapping = paragraph_map[0]
    delete_mapping = paragraph_map[1]

    delete_locator = LocatorResult(
        finding=_finding(action="DELETE", existing=delete_mapping.text, replacement=None),
        status="matched",
        locations=[
            EditLocation(
                mapping=delete_mapping,
                match_start=0,
                match_end=len(delete_mapping.text),
                matched_text=delete_mapping.text,
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text=None,
        action_type="DELETE",
        warning=None,
    )

    add_finding = _finding(action="ADD", existing="", replacement="Added paragraph after anchor")
    add_finding.anchorText = anchor_mapping.text
    add_finding.insertPosition = "after"
    add_locator = LocatorResult(
        finding=add_finding,
        status="matched",
        locations=[
            EditLocation(
                mapping=anchor_mapping,
                match_start=0,
                match_end=len(anchor_mapping.text),
                matched_text=anchor_mapping.text,
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="Added paragraph after anchor",
        action_type="ADD",
        warning=None,
    )

    actions = build_edit_actions([delete_locator, add_locator])
    assert len(actions) == 2

    report = apply_edits_to_spec(source, output, actions)
    assert report.edits_applied == 2

    saved = Document(output)
    texts = [p.text for p in saved.paragraphs]
    assert texts == ["Anchor paragraph for ADD", "Added paragraph after anchor"]


def test_multiple_whole_paragraph_deletes_apply_in_descending_order(tmp_path: Path):
    """Two DELETEs at adjacent body indices both apply correctly."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Keep A")
    doc.add_paragraph("Drop B")
    doc.add_paragraph("Drop C")
    doc.add_paragraph("Keep D")
    doc.save(source)

    spec = extract_text_from_docx(source)
    paragraph_map = spec.paragraph_map

    drop_b = paragraph_map[1]
    drop_c = paragraph_map[2]

    def _delete_locator(mapping: ParagraphMapping) -> LocatorResult:
        return LocatorResult(
            finding=_finding(action="DELETE", existing=mapping.text, replacement=None),
            status="matched",
            locations=[
                EditLocation(
                    mapping=mapping,
                    match_start=0,
                    match_end=len(mapping.text),
                    matched_text=mapping.text,
                    match_confidence=1.0,
                    match_method="exact",
                )
            ],
            replacement_text=None,
            action_type="DELETE",
            warning=None,
        )

    actions = build_edit_actions([_delete_locator(drop_b), _delete_locator(drop_c)])
    assert len(actions) == 2

    report = apply_edits_to_spec(source, output, actions)
    assert report.edits_applied == 2

    saved = Document(output)
    assert [p.text for p in saved.paragraphs] == ["Keep A", "Keep D"]


# -- Section 8.3: revalidation before mutation -------------------------------


def test_replacement_skipped_when_recorded_offset_no_longer_matches(tmp_path: Path):
    """Two overlapping replacements: first runs, second's offsets become stale."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Provide ASCE 7-16 bracing per ASCE 7-16 code.")
    doc.save(source)

    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[0]
    text = mapping.text

    first_start = text.find("ASCE 7-16")
    first_end = first_start + len("ASCE 7-16")
    second_start = text.find("ASCE 7-16", first_end)
    second_end = second_start + len("ASCE 7-16")

    def _loc(start: int, end: int) -> LocatorResult:
        return LocatorResult(
            finding=_finding(existing=text[start:end], replacement="ASCE 7-22"),
            status="matched",
            locations=[
                EditLocation(
                    mapping=mapping,
                    match_start=start,
                    match_end=end,
                    matched_text=text[start:end],
                    match_confidence=1.0,
                    match_method="exact",
                )
            ],
            replacement_text="ASCE 7-22",
            action_type="EDIT",
            warning=None,
        )

    actions = build_edit_actions([_loc(first_start, first_end), _loc(second_start, second_end)])
    report = apply_edits_to_spec(source, output, actions)

    saved = Document(output)
    # Both occurrences updated thanks to revalidation falling back to unique
    # substring presence after the first replacement shifted character offsets.
    assert "ASCE 7-16" not in saved.paragraphs[0].text
    assert saved.paragraphs[0].text.count("ASCE 7-22") == 2
    assert report.edits_applied == 2


def test_revalidation_skips_when_target_text_absent(tmp_path: Path):
    """A locator result whose recorded matched_text isn't in the doc should skip."""
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Real paragraph contents only.")
    doc.save(source)

    spec = extract_text_from_docx(source)
    mapping = spec.paragraph_map[0]

    # Recorded matched_text is fabricated — it does not appear in the doc.
    stale = LocatorResult(
        finding=_finding(existing="ghost phrase", replacement="replacement"),
        status="matched",
        locations=[
            EditLocation(
                mapping=mapping,
                match_start=0,
                match_end=len("ghost phrase"),
                matched_text="ghost phrase",
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="replacement",
        action_type="EDIT",
        warning=None,
    )

    actions = build_edit_actions([stale])
    report = apply_edits_to_spec(source, output, actions)

    saved = Document(output)
    assert saved.paragraphs[0].text == "Real paragraph contents only."
    assert report.edits_applied == 0
    skip_outcomes = [outcome for outcome in report.outcomes if outcome.status == "skipped"]
    assert any("Precondition" in outcome.detail for outcome in skip_outcomes)


# -- Section 8.7: source-never-overwritten invariant -------------------------


def test_source_path_never_overwritten(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(source)

    original_bytes = source.read_bytes()

    result = LocatorResult(
        finding=_finding(existing="Hello world", replacement="Hi"),
        status="matched",
        locations=[
            EditLocation(
                mapping=_mapping("Hello world", idx=0),
                match_start=0,
                match_end=len("Hello world"),
                matched_text="Hello world",
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="Hi",
        action_type="EDIT",
        warning=None,
    )

    apply_edits_to_spec(source, output, build_edit_actions([result]))

    assert source.exists()
    assert source.read_bytes() == original_bytes
    assert output.exists()
    assert output.read_bytes() != original_bytes


def test_apply_edits_rejects_source_equal_to_output(tmp_path: Path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("Hello")
    doc.save(source)

    result = LocatorResult(
        finding=_finding(existing="Hello", replacement="Hi"),
        status="matched",
        locations=[
            EditLocation(
                mapping=_mapping("Hello", idx=0),
                match_start=0,
                match_end=len("Hello"),
                matched_text="Hello",
                match_confidence=1.0,
                match_method="exact",
            )
        ],
        replacement_text="Hi",
        action_type="EDIT",
        warning=None,
    )

    with pytest.raises(ValueError):
        apply_edits_to_spec(source, source, build_edit_actions([result]))
