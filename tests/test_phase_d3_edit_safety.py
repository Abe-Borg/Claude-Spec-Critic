"""Phase D3 regression tests: multi-edit-per-paragraph ordering + fuzzy gating.

Covers two chunks from the second-agent delta plan:

* **D3.1** — Multiple non-overlapping edits in the same paragraph must apply
  in descending-start order so a growth in an upstream edit does not shift
  the offsets of a downstream edit. Ambiguous partial overlaps must be
  routed to manual review instead of silently picking a winner.
* **D3.2** — Fuzzy locator matches (whole-document and section-anchored
  variants) must be classified as MANUAL_REVIEW so they never auto-apply.
"""

from pathlib import Path

from docx import Document

from src.edit_candidates import (
    SAFETY_AUTO_SAFE,
    SAFETY_AUTO_WITH_CAUTION,
    SAFETY_MANUAL_REVIEW,
)
from src.edit_locator import EditLocation, LocatorResult, locate_edit
from src.extractor import ParagraphMapping, extract_text_from_docx
from src.reviewer import Finding
from src.spec_editor import (
    apply_edits_to_spec,
    build_edit_actions,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    action: str = "EDIT",
    existing: str = "",
    replacement: str | None = "",
    severity: str = "HIGH",
) -> Finding:
    return Finding(
        severity=severity,
        fileName="spec.docx",
        section="1.0",
        issue="Issue",
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="Code",
        confidence=0.9,
    )


def _locator_for_paragraph(
    *,
    mapping: ParagraphMapping,
    match_start: int,
    match_end: int,
    matched_text: str,
    replacement_text: str | None,
    confidence: float = 1.0,
    action: str = "EDIT",
    severity: str = "HIGH",
) -> LocatorResult:
    location = EditLocation(
        mapping=mapping,
        match_start=match_start,
        match_end=match_end,
        matched_text=matched_text,
        match_confidence=confidence,
        match_method="exact",
    )
    return LocatorResult(
        finding=_finding(
            action=action,
            existing=matched_text,
            replacement=replacement_text,
            severity=severity,
        ),
        status="matched",
        locations=[location],
        replacement_text=replacement_text,
        action_type=action,
        warning=None,
    )


def _mapping(text: str, *, idx: int, element_type: str = "paragraph") -> ParagraphMapping:
    return ParagraphMapping(
        body_index=idx,
        element_type=element_type,
        text=text,
        table_index=0 if element_type == "table_cell" else None,
        row_index=0 if element_type == "table_cell" else None,
        cell_index=None,
    )


def _live_mapping(tmp_path: Path, text: str) -> tuple[Path, ParagraphMapping]:
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph(text)
    doc.save(source)
    spec = extract_text_from_docx(source)
    return source, spec.paragraph_map[0]


# ---------------------------------------------------------------------------
# Chunk D3.1 — multi-edit-per-paragraph ordering and overlap safety
# ---------------------------------------------------------------------------


def test_same_length_replacements_in_one_paragraph_both_apply(tmp_path: Path):
    """Two equal-length, non-overlapping replacements both apply cleanly.

    Baseline of the multi-edit case: no offset shifting at all (same lengths),
    so the descending order of application is irrelevant — but the test
    guards against a regression that drops one of two non-overlapping edits.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(tmp_path, "alpha bravo charlie delta echo")
    text = mapping.text

    bravo_start = text.index("bravo")
    delta_start = text.index("delta")
    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=bravo_start,
            match_end=bravo_start + len("bravo"),
            matched_text="bravo",
            replacement_text="BRAVO",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=delta_start,
            match_end=delta_start + len("delta"),
            matched_text="delta",
            replacement_text="DELTA",
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    assert saved.paragraphs[0].text == "alpha BRAVO charlie DELTA echo"
    assert report.edits_applied == 2
    assert report.edits_skipped == 0
    assert report.edits_failed == 0


def test_growth_before_downstream_edit_in_one_paragraph(tmp_path: Path):
    """Upstream edit grows the paragraph; downstream edit still applies.

    Descending-start ordering means the downstream (higher-offset) edit runs
    first, before the upstream (lower-offset) growth shifts anything. The
    downstream offsets stay valid. Even when the upstream edit then grows
    the paragraph, the upstream's own offsets at the low end are unaffected.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(tmp_path, "alpha bravo charlie delta echo")
    text = mapping.text

    bravo_start = text.index("bravo")
    delta_start = text.index("delta")
    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=bravo_start,
            match_end=bravo_start + len("bravo"),
            matched_text="bravo",
            replacement_text="BRAVO-MUCH-LONGER-REPLACEMENT",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=delta_start,
            match_end=delta_start + len("delta"),
            matched_text="delta",
            replacement_text="DELTA",
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    assert saved.paragraphs[0].text == (
        "alpha BRAVO-MUCH-LONGER-REPLACEMENT charlie DELTA echo"
    )
    assert report.edits_applied == 2
    assert report.edits_failed == 0


def test_shrink_before_downstream_edit_in_one_paragraph(tmp_path: Path):
    """Upstream edit shrinks the paragraph; downstream edit still applies.

    Mirrors the growth case but with the upstream replacement shorter than
    the original text. Descending-start ordering keeps the downstream
    offsets correct.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(
        tmp_path, "alpha REALLY-LONG-BRAVO charlie delta echo"
    )
    text = mapping.text

    big_start = text.index("REALLY-LONG-BRAVO")
    delta_start = text.index("delta")
    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=big_start,
            match_end=big_start + len("REALLY-LONG-BRAVO"),
            matched_text="REALLY-LONG-BRAVO",
            replacement_text="b",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=delta_start,
            match_end=delta_start + len("delta"),
            matched_text="delta",
            replacement_text="DELTA",
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    assert saved.paragraphs[0].text == "alpha b charlie DELTA echo"
    assert report.edits_applied == 2
    assert report.edits_failed == 0


def test_three_non_overlapping_edits_in_one_paragraph(tmp_path: Path):
    """Three non-overlapping edits in the same paragraph all apply.

    Stress-tests the descending-start ordering: even with two upstream
    growths, the highest-offset edit runs first and is unaffected by either
    growth.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(
        tmp_path, "alpha bravo charlie delta echo foxtrot"
    )
    text = mapping.text

    bravo_start = text.index("bravo")
    delta_start = text.index("delta")
    foxtrot_start = text.index("foxtrot")
    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=bravo_start,
            match_end=bravo_start + len("bravo"),
            matched_text="bravo",
            replacement_text="BRAVO-EXPANDED",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=delta_start,
            match_end=delta_start + len("delta"),
            matched_text="delta",
            replacement_text="DELTA-EXPANDED",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=foxtrot_start,
            match_end=foxtrot_start + len("foxtrot"),
            matched_text="foxtrot",
            replacement_text="FOXTROT",
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    assert saved.paragraphs[0].text == (
        "alpha BRAVO-EXPANDED charlie DELTA-EXPANDED echo FOXTROT"
    )
    assert report.edits_applied == 3
    assert report.edits_failed == 0


def test_partial_overlap_in_one_paragraph_is_manual_review(tmp_path: Path):
    """Partial overlap (neither contains the other) → both skipped.

    The previous behavior picked a winner via severity / confidence / span
    heuristics. Per the delta plan, ambiguous overlapping edits in the same
    paragraph must not be auto-applied; the paragraph stays untouched and
    both outcomes carry a "manual review" detail.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(tmp_path, "alpha bravo charlie delta echo")
    text = mapping.text

    # Edit A: replace "bravo charlie" with "X"
    # Edit B: replace "charlie delta" with "Y"
    # Spans overlap at " charlie " — neither contains the other.
    a_start = text.index("bravo charlie")
    a_end = a_start + len("bravo charlie")
    b_start = text.index("charlie delta")
    b_end = b_start + len("charlie delta")
    assert a_start < b_end and a_end > b_start  # overlap
    assert not (a_start <= b_start and a_end >= b_end)  # A doesn't contain B
    assert not (b_start <= a_start and b_end >= a_end)  # B doesn't contain A

    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=a_start,
            match_end=a_end,
            matched_text="bravo charlie",
            replacement_text="X",
            confidence=0.95,
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=b_start,
            match_end=b_end,
            matched_text="charlie delta",
            replacement_text="Y",
            confidence=0.85,
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    assert saved.paragraphs[0].text == "alpha bravo charlie delta echo"
    assert report.edits_applied == 0
    assert report.edits_skipped == 2
    for outcome in report.outcomes:
        assert outcome.status == "skipped"
        assert "ambiguous" in outcome.detail.lower()
        assert "manual review" in outcome.detail.lower()


def test_three_way_overlap_chain_taints_third_edit(tmp_path: Path):
    """A third edit overlapping a tainted region is also routed to manual review.

    The first two edits partial-overlap and are both removed from the
    accepted set. Without taint tracking, a third edit overlapping with
    either of those removed spans would slip through as "no overlap" and
    get auto-applied. Verify the taint range catches it.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(
        tmp_path, "alpha bravo charlie delta echo foxtrot golf"
    )
    text = mapping.text

    # A and B partial-overlap (no containment).
    a_start = text.index("bravo charlie")
    a_end = a_start + len("bravo charlie")
    b_start = text.index("charlie delta")
    b_end = b_start + len("charlie delta")
    # C overlaps with B's range (specifically the "delta" portion).
    c_start = text.index("delta echo")
    c_end = c_start + len("delta echo")

    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=a_start,
            match_end=a_end,
            matched_text="bravo charlie",
            replacement_text="X",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=b_start,
            match_end=b_end,
            matched_text="charlie delta",
            replacement_text="Y",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=c_start,
            match_end=c_end,
            matched_text="delta echo",
            replacement_text="Z",
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    # Paragraph fully unchanged: A and B are ambiguous, C overlaps the
    # tainted region from the A+B union and is also skipped.
    assert saved.paragraphs[0].text == "alpha bravo charlie delta echo foxtrot golf"
    assert report.edits_applied == 0
    assert report.edits_skipped == 3
    # At least one skipped outcome should mention the tainted region.
    tainted_details = [o.detail for o in report.outcomes if "already flagged" in o.detail]
    assert tainted_details, "expected a tainted-region detail for the third edit"


def test_strict_containment_still_keeps_broader_edit(tmp_path: Path):
    """Regression guard: strict containment is not "ambiguous".

    The delta plan only routes ambiguous partial overlaps to manual review.
    When one edit strictly contains the other, the broader edit's
    replacement subsumes the narrower edit's intent and is safe to apply.
    """
    output = tmp_path / "output.docx"
    source, mapping = _live_mapping(
        tmp_path,
        "Pipe markers include refrigerant piping and condensate piping using R454B notation.",
    )
    text = mapping.text

    narrow_start = text.index("R454B")
    narrow_end = narrow_start + len("R454B")
    broad_start = 0
    broad_end = len(text)
    # Broad strictly contains narrow.
    assert broad_start <= narrow_start and broad_end >= narrow_end

    edits = [
        _locator_for_paragraph(
            mapping=mapping,
            match_start=narrow_start,
            match_end=narrow_end,
            matched_text="R454B",
            replacement_text="R-454B",
            confidence=1.0,
            severity="GRIPES",
        ),
        _locator_for_paragraph(
            mapping=mapping,
            match_start=broad_start,
            match_end=broad_end,
            matched_text=text,
            replacement_text=(
                "Pipe markers shall separate refrigerant piping from condensate "
                "piping and use R-454B notation."
            ),
            confidence=1.0,
            severity="MEDIUM",
        ),
    ]

    report = apply_edits_to_spec(source, output, build_edit_actions(edits))
    saved = Document(output)
    assert saved.paragraphs[0].text.startswith("Pipe markers shall separate")
    assert "R-454B" in saved.paragraphs[0].text
    assert report.edits_applied == 1
    assert report.edits_skipped == 1


# ---------------------------------------------------------------------------
# Chunk D3.2 — fuzzy locator (whole-document and section-anchored) → manual review
# ---------------------------------------------------------------------------


def test_whole_document_fuzzy_match_classified_manual_review():
    """Whole-document fuzzy match falls into MANUAL_REVIEW.

    Re-asserts the existing fuzzy → manual-review rule so D3.2 does not
    accidentally weaken it while adding the section-anchored variant.
    """
    paragraph_map = [
        _mapping("Provide 2-hour fire-resistance rated shaft enclosures.", idx=0),
    ]
    finding = _finding(
        existing="Provide two hour fire resistance shaft enclosure.",
        replacement="Provide 2-hour fire-resistance rated shaft enclosures.",
    )
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.locations[0].match_method == "fuzzy"
    assert result.safety_category == SAFETY_MANUAL_REVIEW


def test_section_anchored_fuzzy_match_classified_manual_review():
    """Section-anchored fuzzy match must not be auto-applied.

    Narrowing the search window to a section heading does not make a
    paraphrase identification safe enough for silent document mutation.
    Previously the section-anchored path relabeled fuzzy hits as
    "section_anchored" and they slipped through as AUTO_WITH_CAUTION,
    which auto-applies under the default ``allow_caution=True`` policy.
    """
    paragraph_map = [
        _mapping("1.0 GENERAL", idx=0),
        _mapping("Provide 2-hour fire-resistance rated shaft enclosures.", idx=1),
        _mapping("2.0 PRODUCTS", idx=2),
    ]
    finding = _finding(
        existing="Provide two hour fire resistance shaft enclosure.",
        replacement="Provide 2-hour fire-resistance rated shaft enclosures.",
    )
    finding.section = "1.0"
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.locations[0].match_method == "section_anchored_fuzzy"
    assert result.safety_category == SAFETY_MANUAL_REVIEW


def test_section_anchored_exact_match_still_auto_with_caution():
    """Regression guard: section-anchored EXACT matches are not affected.

    D3.2 only downgrades the section-anchored FUZZY variant. Section-anchored
    exact matches keep their existing AUTO_WITH_CAUTION classification so
    short-text disambiguation paths continue to work.
    """
    paragraph_map = [
        _mapping("1.0 GENERAL", idx=0),
        _mapping("Voltage shall be 208V", idx=1),
        _mapping("2.0 PRODUCTS", idx=2),
        _mapping("Voltage shall be 480V", idx=3),
    ]
    finding = _finding(existing="208V", replacement="240V")
    finding.section = "1.0"
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.locations[0].match_method == "section_anchored"
    assert result.safety_category == SAFETY_AUTO_WITH_CAUTION


def test_build_edit_actions_excludes_section_anchored_fuzzy():
    """``build_edit_actions`` must not produce an auto-apply action for fuzzy.

    Even with the default ``allow_caution=True`` (which exists so that
    AUTO_WITH_CAUTION can still flow through to mutation), the
    section-anchored fuzzy locator result must fall into MANUAL_REVIEW and
    produce zero actions.
    """
    paragraph_map = [
        _mapping("1.0 GENERAL", idx=0),
        _mapping("Provide 2-hour fire-resistance rated shaft enclosures.", idx=1),
    ]
    finding = _finding(
        existing="Provide two hour fire resistance shaft enclosure.",
        replacement="Provide 2-hour fire-resistance rated shaft enclosures.",
    )
    finding.section = "1.0"
    result = locate_edit(finding, paragraph_map)
    assert result.safety_category == SAFETY_MANUAL_REVIEW

    actions = build_edit_actions([result], allow_caution=True)
    assert actions == []
    assert result.warning is not None
    assert "manual review" in result.warning.lower()


def test_build_edit_actions_excludes_whole_document_fuzzy():
    """Whole-document fuzzy results were already MANUAL_REVIEW; assert no actions."""
    paragraph_map = [
        _mapping("Provide 2-hour fire-resistance rated shaft enclosures.", idx=0),
    ]
    finding = _finding(
        existing="Provide two hour fire resistance shaft enclosure.",
        replacement="Provide 2-hour fire-resistance rated shaft enclosures.",
    )
    result = locate_edit(finding, paragraph_map)
    actions = build_edit_actions([result], allow_caution=True)
    assert actions == []
    assert result.safety_category == SAFETY_MANUAL_REVIEW


def test_exact_paragraph_match_still_auto_safe():
    """Regression guard: D3.2 leaves the auto-safe exact path untouched."""
    paragraph_map = [
        _mapping("Install copper piping with Type L wall thickness.", idx=0),
    ]
    finding = _finding(
        existing="Type L wall thickness",
        replacement="Type M wall thickness",
    )
    result = locate_edit(finding, paragraph_map)
    assert result.status == "matched"
    assert result.locations[0].match_method == "exact"
    assert result.safety_category == SAFETY_AUTO_SAFE
