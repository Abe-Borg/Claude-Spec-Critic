"""Batch resume / recovery: persistence + reconstruction + headless collection.

Covers the machinery that lets a detached review batch be reconnected and
finished without re-submitting:

- ``batch_resume`` persistence (save/load/clear) is a round-trip and is
  defensive against every malformed-file axis (the reliability gap the prior
  "you will be prompted to resume" copy promised but never implemented).
- ``reconstruct_batch_submission`` restores the request map verbatim, so a
  batch's findings come back even when the local source files are gone, and
  best-effort re-extraction repopulates specs when the files are present.
- ``thin_submission_from_batch_results`` rebuilds the request map from a bare
  batch id (no saved state) by reading the remote results' custom ids — the
  path that recovers a batch submitted before resume state existed.
- ``run_batch_collection_headless`` drives collect → verify → cross-check →
  finalize off-GUI and preserves findings end-to-end.
"""
from __future__ import annotations

import json
import time
import types
from pathlib import Path

from docx import Document

from src.batch.batch import BatchJob
from src.batch.batch_runtime import PollOutcome
from src.orchestration import batch_resume as br
from src.orchestration import pipeline as pl
from src.orchestration.batch_resume import (
    PendingBatch,
    _parse_review_custom_id,
    clear_pending_batch,
    load_pending_batch,
    pending_batch_path,
    save_pending_batch,
    thin_submission_from_batch_results,
)
from src.orchestration.pipeline import (
    BatchSubmission,
    reconstruct_batch_submission,
    run_batch_collection_headless,
)
from src.core.code_cycles import DEFAULT_CYCLE
from src.modules import DEFAULT_MODULE
from src.review.reviewer import Finding, ReviewResult
from src.verification.verification_cache import VerificationCache
from tests.fixtures.fake_anthropic import FakeBatchResult, FakeBatchResultEnvelope


def _submission(batch_id: str = "msgbatch_TEST") -> BatchSubmission:
    job = BatchJob(
        batch_id=batch_id,
        job_type="review",
        request_map={"review__a__0": {"filename": "22 11 16 - Water.docx", "index": 0, "type": "review"}},
        created_at=1700000000.0,
    )
    return BatchSubmission(
        job=job,
        files_reviewed=["22 11 16 - Water.docx"],
        review_request_ids=["review__a__0"],
        model="claude-opus-4-8",
        project_context="K-12 DSA project",
        cycle_label=DEFAULT_CYCLE.label,
        cross_check_enabled=True,
        prepared_specs=None,
    )


# ===========================================================================
# Persistence: save / load / clear
# ===========================================================================


class TestPersistence:
    def test_round_trip(self, tmp_path):
        path = tmp_path / "pending_batch.json"
        pending = PendingBatch.from_submission(
            _submission(),
            input_dir="/specs",
            files=["/specs/22 11 16 - Water.docx"],
            run_id="run123",
            app_version="3.0.0",
        )
        save_pending_batch(pending, path=path)
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.batch_id == "msgbatch_TEST"
        assert loaded.review_request_ids == ["review__a__0"]
        assert loaded.request_map == {"review__a__0": {"filename": "22 11 16 - Water.docx", "index": 0, "type": "review"}}
        assert loaded.files == ["/specs/22 11 16 - Water.docx"]
        assert loaded.cross_check_enabled is True
        assert loaded.run_id == "run123"
        assert loaded.submitted_at == 1700000000.0

    def test_missing_file_returns_none(self, tmp_path):
        assert load_pending_batch(path=tmp_path / "nope.json") is None

    def test_malformed_json_returns_none(self, tmp_path):
        path = tmp_path / "p.json"
        path.write_text("{not json", encoding="utf-8")
        assert load_pending_batch(path=path) is None

    def test_wrong_schema_version_returns_none(self, tmp_path):
        path = tmp_path / "p.json"
        path.write_text(json.dumps({"batch_id": "x", "schema_version": 999}), encoding="utf-8")
        assert load_pending_batch(path=path) is None

    def test_empty_batch_id_returns_none(self, tmp_path):
        path = tmp_path / "p.json"
        path.write_text(json.dumps({"batch_id": "  ", "schema_version": br._SCHEMA_VERSION}), encoding="utf-8")
        assert load_pending_batch(path=path) is None

    def test_non_numeric_submitted_at_is_coerced(self, tmp_path):
        path = tmp_path / "p.json"
        path.write_text(
            json.dumps({"batch_id": "msgbatch_X", "schema_version": br._SCHEMA_VERSION, "submitted_at": "yesterday"}),
            encoding="utf-8",
        )
        loaded = load_pending_batch(path=path)
        assert loaded is not None
        assert loaded.submitted_at == 0.0

    def test_clear_removes_file_and_is_idempotent(self, tmp_path):
        path = tmp_path / "p.json"
        save_pending_batch(PendingBatch.from_submission(_submission()), path=path)
        assert path.exists()
        clear_pending_batch(path=path)
        assert not path.exists()
        clear_pending_batch(path=path)  # second call must not raise

    def test_path_honors_env_override(self, tmp_path, monkeypatch):
        target = tmp_path / "custom" / "pending.json"
        monkeypatch.setenv("SPEC_CRITIC_PENDING_BATCH_PATH", str(target))
        assert pending_batch_path() == target


# ===========================================================================
# custom-id parsing
# ===========================================================================


class TestParseCustomId:
    def test_well_formed(self):
        assert _parse_review_custom_id("review__22_11_16___Water__3") == ("22_11_16___Water", 3)

    def test_simple(self):
        assert _parse_review_custom_id("review__SPEC__0") == ("SPEC", 0)

    def test_rejects_non_review(self):
        assert _parse_review_custom_id("verify__0") is None

    def test_rejects_non_integer_index(self):
        assert _parse_review_custom_id("review__SPEC__x") is None


# ===========================================================================
# reconstruct_batch_submission
# ===========================================================================


class TestReconstruct:
    def test_request_map_preserved_when_files_missing(self):
        """The core reliability property: a detached batch's findings are
        recoverable even when every source file is gone."""
        logs: list[tuple[str, str]] = []
        sub = reconstruct_batch_submission(
            batch_id="msgbatch_X",
            request_map={"review__a__0": {"filename": "a.docx", "index": 0, "type": "review"}},
            review_request_ids=["review__a__0"],
            files_reviewed=["a.docx"],
            input_dir="/gone",
            files=["/gone/a.docx"],
            model="claude-opus-4-8",
            project_context="",
            module=DEFAULT_MODULE,
            cross_check_enabled=True,
            created_at=123.0,
            log=lambda msg, level="info": logs.append((level, msg)),
        )
        assert sub.job.batch_id == "msgbatch_X"
        assert sub.job.request_map == {"review__a__0": {"filename": "a.docx", "index": 0, "type": "review"}}
        assert sub.review_request_ids == ["review__a__0"]
        assert sub.prepared_specs is None  # files unavailable → no re-extraction
        assert sub.job.created_at == 123.0
        assert any("not found" in msg for _lvl, msg in logs)

    def test_reextracts_when_files_present(self, tmp_path):
        path = tmp_path / "22 11 16 - Domestic Water Piping.docx"
        doc = Document()
        doc.add_paragraph("PART 1 - GENERAL")
        doc.add_paragraph("Provide copper tubing per applicable code.")
        doc.save(str(path))

        sub = reconstruct_batch_submission(
            batch_id="msgbatch_X",
            request_map={"review__a__0": {"filename": path.name, "index": 0, "type": "review"}},
            review_request_ids=["review__a__0"],
            files_reviewed=[path.name],
            input_dir=str(tmp_path),
            files=[str(path)],
            model="claude-opus-4-8",
            project_context="",
            module=DEFAULT_MODULE,
            cross_check_enabled=True,
            created_at=123.0,
        )
        assert sub.prepared_specs is not None
        assert len(sub.prepared_specs) == 1
        assert sub.prepared_specs[0].filename == path.name
        # request map is still the persisted one, not recomputed
        assert sub.review_request_ids == ["review__a__0"]

    def test_drift_warning_when_reviewed_set_differs(self, tmp_path):
        path = tmp_path / "actual.docx"
        doc = Document()
        doc.add_paragraph("PART 1 - GENERAL")
        doc.add_paragraph("Some requirement text here.")
        doc.save(str(path))
        logs: list[str] = []
        reconstruct_batch_submission(
            batch_id="msgbatch_X",
            request_map={},
            review_request_ids=[],
            files_reviewed=["a-different-name.docx"],  # ≠ extracted "actual.docx"
            input_dir=str(tmp_path),
            files=[str(path)],
            model="claude-opus-4-8",
            project_context="",
            module=DEFAULT_MODULE,
            cross_check_enabled=True,
            created_at=1.0,
            log=lambda msg, level="info": logs.append(msg),
        )
        assert any("differs from the originally reviewed" in m for m in logs)


# ===========================================================================
# thin_submission_from_batch_results
# ===========================================================================


class _FakeBatchesAPI:
    def __init__(self, results):
        self._results = results

    def results(self, _batch_id):
        return iter(self._results)


class _FakeMessages:
    def __init__(self, results):
        self.batches = _FakeBatchesAPI(results)


class _FakeClient:
    def __init__(self, results):
        self.messages = _FakeMessages(results)


class TestThinSubmission:
    def test_rebuilds_request_map_in_index_order(self, monkeypatch):
        # Two items returned out of order; the rebuilt map must order them by
        # the index encoded in the custom id.
        results = [
            FakeBatchResult(custom_id="review__BBB__1", result=FakeBatchResultEnvelope(type="succeeded")),
            FakeBatchResult(custom_id="review__AAA__0", result=FakeBatchResultEnvelope(type="succeeded")),
            FakeBatchResult(custom_id="garbage__9", result=FakeBatchResultEnvelope(type="succeeded")),
        ]
        # ``_get_client`` is imported inside the function from ``batch.batch``;
        # patch it there so the function picks up the fake.
        import src.batch.batch as batch_mod
        monkeypatch.setattr(batch_mod, "_get_client", lambda: _FakeClient(results))

        sub = thin_submission_from_batch_results("msgbatch_X", model="claude-opus-4-8")
        assert sub.review_request_ids == ["review__AAA__0", "review__BBB__1"]
        assert sub.job.request_map["review__AAA__0"]["index"] == 0
        assert sub.job.request_map["review__BBB__1"]["index"] == 1
        # unrecognized custom id is ignored, not mapped
        assert "garbage__9" not in sub.job.request_map
        assert sub.prepared_specs is None  # no files → findings-only

    def test_recovers_real_filename_from_supplied_files(self, monkeypatch, tmp_path):
        # The custom id carries the SANITIZED stem; with the real file supplied,
        # the thin path must map it back to the real filename so the failed-spec
        # cross-check exclusion and the filename-keyed repair fallback match.
        real_name = "22 11 16 - Water.docx"
        path = tmp_path / real_name
        doc = Document()
        doc.add_paragraph("PART 1 - GENERAL")
        doc.add_paragraph("Provide copper tubing per applicable code.")
        doc.save(str(path))
        # custom id for this file at index 0 (sanitizer keeps hyphens, maps
        # other non-alphanumerics to "_"): "review__22_11_16_-_Water__0".
        results = [
            FakeBatchResult(custom_id="review__22_11_16_-_Water__0", result=FakeBatchResultEnvelope(type="succeeded")),
        ]
        import src.batch.batch as batch_mod
        monkeypatch.setattr(batch_mod, "_get_client", lambda: _FakeClient(results))

        sub = thin_submission_from_batch_results(
            "msgbatch_X",
            model="claude-opus-4-8",
            input_dir=str(tmp_path),
            files=[str(path)],
            cross_check_enabled=True,
        )
        cid = "review__22_11_16_-_Water__0"
        assert sub.job.request_map[cid]["filename"] == real_name  # real name, not the stem
        assert sub.files_reviewed == [real_name]
        # Files exist → re-extraction populated specs with the matching name, so
        # no drift and cross-check can run.
        assert sub.prepared_specs is not None
        assert sub.prepared_specs[0].filename == real_name


class TestRepairFallbackResolvesByFilename:
    """Fix for the resume misalignment: the review-repair fallback must select
    the retryable spec by FILENAME, so a re-extracted prepared_specs in a
    different order than the persisted request_map indices never repairs (and
    mis-attributes) the wrong spec."""

    def _spec(self, name):
        return types.SimpleNamespace(filename=name, content=f"body of {name}")

    def test_repair_picks_spec_by_filename_not_position(self, monkeypatch):
        specA, specB, specC = self._spec("A.docx"), self._spec("B.docx"), self._spec("C.docx")
        # prepared_specs RE-EXTRACTED in a DIFFERENT order than the submit
        # indices below (index 1 == B, but prepared[1] == A).
        prepared = [specC, specA, specB]
        job = BatchJob(
            batch_id="X",
            job_type="review",
            request_map={
                "review__a__0": {"filename": "A.docx", "index": 0, "type": "review"},
                "review__b__1": {"filename": "B.docx", "index": 1, "type": "review"},
                "review__c__2": {"filename": "C.docx", "index": 2, "type": "review"},
            },
            created_at=0.0,
        )
        sub = BatchSubmission(
            job=job,
            files_reviewed=["A.docx", "B.docx", "C.docx"],
            review_request_ids=["review__a__0", "review__b__1", "review__c__2"],
            model="m",
            prepared_specs=prepared,
        )
        ok = ReviewResult(findings=[], parse_status="ok")
        bad = ReviewResult(findings=[], parse_status="parse_error", error="truncated")
        results = {"review__a__0": ok, "review__b__1": bad, "review__c__2": ok}

        captured: dict = {}

        def fake_submit(repair_specs, **_kw):
            captured["specs"] = list(repair_specs)
            return BatchJob(
                batch_id="repair",
                job_type="review",
                request_map={"review__r__0": {"filename": repair_specs[0].filename, "index": 0, "type": "review"}},
                created_at=0.0,
            )

        monkeypatch.setattr(pl, "submit_review_batch", fake_submit)
        monkeypatch.setattr(pl, "poll_batch_bounded", lambda *a, **k: PollOutcome(terminal=True, terminal_status="ended"))
        monkeypatch.setattr(
            pl, "preprocess_spec",
            lambda content, filename, *, cycle: types.SimpleNamespace(
                leed_alerts=[], placeholder_alerts=[], code_cycle_alerts=[],
                structural_alerts=[], template_marker_alerts=[],
                invalid_code_cycle_alerts=[], duplicate_paragraph_alerts=[],
            ),
        )
        monkeypatch.setattr(
            pl, "retrieve_review_results",
            lambda job, *, model: {"review__r__0": ReviewResult(findings=[], parse_status="ok")},
        )

        from src.orchestration.pipeline import _recover_retryable_review_batch_results
        _recover_retryable_review_batch_results(sub, dict(results), log=lambda *a, **k: None)

        # The repaired spec is B (filename match), NOT prepared[1] == A.
        assert [s.filename for s in captured["specs"]] == ["B.docx"]


# ===========================================================================
# run_batch_collection_headless
# ===========================================================================


class TestHeadlessCollection:
    def test_collects_and_finalizes_preserving_findings(self, monkeypatch):
        finding = Finding(
            severity="HIGH",
            fileName="22 11 16 - Water.docx",
            section="2.1",
            issue="pressure rating mismatch",
            actionType="EDIT",
            existingText="150 psi",
            replacementText="125 psi",
            codeReference="CPC 604",
        )
        rr = ReviewResult(findings=[finding], parse_status="ok")
        monkeypatch.setattr(pl, "retrieve_review_results", lambda job, *, model: {"review__a__0": rr})

        # Keep verification hermetic: pretend everything resolved locally so no
        # batch is submitted. Assert the verification stage was still invoked.
        called = {"verify": False}

        def _fake_start_verification(findings, **kwargs):
            called["verify"] = True
            return None  # all resolved locally → no batch

        monkeypatch.setattr(pl, "start_batch_verification", _fake_start_verification)

        sub = _submission()
        sub.cross_check_enabled = False  # no specs available → skip cross-check
        result = run_batch_collection_headless(sub, cache=VerificationCache(), log=lambda *a, **k: None)

        assert called["verify"] is True
        assert result.review_result is not None
        issues = [f.issue for f in result.review_result.findings]
        assert issues == ["pressure rating mismatch"]
        assert result.cross_check_result is None  # cross-check disabled
        assert result.cycle_label == DEFAULT_CYCLE.label

    def test_no_findings_skips_verification(self, monkeypatch):
        rr = ReviewResult(findings=[], parse_status="ok")
        monkeypatch.setattr(pl, "retrieve_review_results", lambda job, *, model: {"review__a__0": rr})
        called = {"verify": False}
        monkeypatch.setattr(pl, "start_batch_verification", lambda *a, **k: called.__setitem__("verify", True))

        sub = _submission()
        sub.cross_check_enabled = False
        result = run_batch_collection_headless(sub, cache=VerificationCache(), log=lambda *a, **k: None)
        assert called["verify"] is False  # no findings → verification not attempted
        assert result.review_result.findings == []
