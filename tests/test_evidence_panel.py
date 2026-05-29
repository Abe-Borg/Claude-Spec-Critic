"""Tests for the per-finding evidence panel in the report.

Plan section "Per-Finding Evidence Panel in Report".
Every finding with a verification result renders an evidence panel under
the Sources Heading 4 containing:

- Verifier model / verification mode / search budget
- Source quote (verbatim from web_search snippet)
- Verifier rationale (moved here from above the heading)
- Escalation history (when applicable)
- Web/code evidence + rejected source URLs
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.output.report_exporter import export_report
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers (mirror style of test_chunk_n_report_status)
# ---------------------------------------------------------------------------

def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.8,
    action: str = "EDIT",
    existing: str | None = "2019 CBC",
    replacement: str | None = "2025 CBC",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
    )
    f.verification = verification
    return f


def _verification(
    verdict: str = "CONFIRMED",
    *,
    grounded: bool = True,
    cache_status: str = "miss",
    explanation: str = "Verified against CBC §1234.",
    sources: list[str] | None = None,
    rejected: list[dict] | None = None,
    correction: str | None = None,
    model_used: str = "claude-sonnet-4-6",
    verification_mode: str = "standard_reasoning",
    web_search_requests: int = 3,
    source_quote: str = (
        "The 2025 California Building Code adopts the 2024 International "
        "Building Code with California amendments."
    ),
    escalation_attempted: bool = False,
    initial_verdict: str = "",
    initial_model: str = "",
    escalation_reason: str = "",
    escalation_changed_verdict: bool = False,
) -> VerificationResult:
    if sources is None:
        sources = (
            ["https://codes.iccsafe.org/content/CBC2025"]
            if grounded and verdict.upper() in ("CONFIRMED", "CORRECTED")
            else []
        )
    return VerificationResult(
        verdict=verdict,
        explanation=explanation,
        sources=list(sources),
        correction=correction,
        grounded=grounded,
        cache_status=cache_status,
        rejected_sources=list(rejected or []),
        model_used=model_used,
        verification_mode=verification_mode,
        web_search_requests=web_search_requests,
        source_quote=source_quote,
        escalation_attempted=escalation_attempted,
        initial_verdict=initial_verdict,
        initial_model=initial_model,
        escalation_reason=escalation_reason,
        escalation_changed_verdict=escalation_changed_verdict,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ===========================================================================
# Report rendering — evidence panel under Sources
# ===========================================================================

class TestEvidencePanelRendering:
    @pytest.fixture
    def review_with_verification(self) -> ReviewResult:
        verified = _finding(
            severity="CRITICAL",
            issue="Stale CBC reference",
            verification=_verification(
                verdict="CONFIRMED",
                model_used="claude-sonnet-4-6",
                verification_mode="standard_reasoning",
                web_search_requests=4,
                source_quote=(
                    "The 2025 California Building Code requires fire "
                    "sprinklers in all K-12 facilities."
                ),
                explanation="The 2025 CBC supersedes the 2019 cycle.",
                sources=["https://codes.iccsafe.org/content/CBC2025"],
            ),
        )
        return ReviewResult(findings=[verified])

    def test_renders_verification_mode(self, tmp_path: Path, review_with_verification):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        assert "Verification mode:" in text
        # Human-readable form, not the raw underscore string.
        assert "Standard reasoning" in text

    def test_renders_search_budget(self, tmp_path: Path, review_with_verification):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        # CRITICAL severity has an 8-call budget; the verification used
        # 4 of them. The exact "N of M" string should appear.
        assert "Search budget used:" in text
        assert "4 of 8 searches used" in text

    def test_renders_rationale_inside_evidence_panel(
        self, tmp_path: Path, review_with_verification
    ):
        # Rationale moves from above the Sources heading to under it
        # (next to the source quote). The label is preserved so the
        # report trust-model tests still pass.
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        assert "Verification rationale:" in text
        # Rationale text must follow source quote in the document
        # body — the panel order matters for readability.
        quote_idx = text.find("fire sprinklers in all K-12 facilities")
        rationale_idx = text.find("The 2025 CBC supersedes")
        assert quote_idx >= 0
        assert rationale_idx >= 0
        assert rationale_idx > quote_idx, (
            "rationale should render after the source quote inside the panel"
        )

    def test_local_skip_panel_suppresses_search_budget(self, tmp_path: Path):
        # A LOCAL_SKIP verification has nothing to do with the web
        # search budget; rendering "0 of N searches" would be
        # misleading. The line is suppressed.
        local = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verification(
                verdict="UNVERIFIED",
                grounded=False,
                cache_status="local_skip",
                verification_mode="local_skip",
                model_used="local",
                web_search_requests=0,
                source_quote="",
                explanation="Locally classified — placeholder GRIPES.",
                sources=[],
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[local])), out
        )
        text = _all_text_from(Document(str(out)))
        # Mode line still renders; budget line suppressed.
        assert "Local skip" in text
        assert "searches used" not in text


# ===========================================================================
# Escalation history rendering
# ===========================================================================

class TestEscalationHistoryRendering:
    def test_renders_when_escalation_attempted(self, tmp_path: Path):
        # Escalation history must surface when escalation_attempted is
        # set, even if the verdict did not change — the audit trail is
        # the point.
        escalated = _finding(
            severity="CRITICAL",
            verification=_verification(
                verdict="CONFIRMED",
                model_used="claude-opus-4-7",
                verification_mode="deep_reasoning",
                escalation_attempted=True,
                initial_verdict="UNVERIFIED",
                initial_model="claude-sonnet-4-6",
                escalation_reason="ungrounded_critical_high",
                escalation_changed_verdict=True,
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[escalated])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Escalation history:" in text
        assert "UNVERIFIED" in text
        assert "claude-sonnet-4-6" in text
        assert "claude-opus-4-7" in text
        assert "ungrounded_critical_high" in text
        # When the escalation changed the verdict, the panel must flag
        # the disagreement.
        assert "models disagreed" in text

    def test_omitted_when_no_escalation(self, tmp_path: Path):
        verified = _finding(
            verification=_verification(
                verdict="CONFIRMED",
                escalation_attempted=False,
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[verified])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Escalation history:" not in text


# ===========================================================================
# Integration — every CONFIRMED / CORRECTED finding gets a source quote
# ===========================================================================

class TestEvidencePanelSuccessCriteria:
    """Mirrors the success criteria from the chunk plan."""

    @pytest.mark.parametrize(
        "verdict, correction, source_quote, model_used",
        [
            (
                "CONFIRMED",
                None,
                "CBC §1234 explicitly requires fire sprinklers.",
                "claude-sonnet-4-6",
            ),
            (
                "CORRECTED",
                "Replace with 2025 CBC",
                "The 2025 CBC supersedes the 2019 cycle.",
                "claude-opus-4-7",
            ),
        ],
    )
    def test_grounded_finding_shows_source_quote_and_verifier_model(
        self, tmp_path: Path, verdict, correction, source_quote, model_used
    ):
        # Both source quote and verifier model must appear in the output
        # for any grounded CONFIRMED / CORRECTED finding.
        f = _finding(
            verification=_verification(
                verdict=verdict,
                correction=correction,
                source_quote=source_quote,
                model_used=model_used,
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert source_quote in text
        assert model_used in text
