"""Chunk 4 tests: per-finding evidence panel in the report.

Trust Upgrade plan section "Chunk 4 — Per-Finding Evidence Panel in Report".
Every finding with a verification result renders an evidence panel under
the Sources Heading 4 containing:

- Verifier model / verification mode / search budget
- Source quote (verbatim from web_search snippet, Chunk 2 schema)
- Verifier rationale (moved here from above the heading)
- Escalation history (when applicable)
- Web/code evidence + rejected source URLs

Findings with edit proposals also render a separate "Edit Target
Evidence" Heading 4 containing:

- Locator status, match method, match confidence
- Safety category (AUTO_SAFE / AUTO_WITH_CAUTION / MANUAL_REVIEW / REPORT_ONLY)
- Element id (when present)

The locator-evidence dict is populated during
``pipeline.finalize_batch_result`` / ``pipeline.run_review`` AND again
during ``apply_edits.execute_edit_plan`` so first-time exports and
resumed apply runs both have the data. It round-trips through resume
state for resume-after-apply workflows.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.editing.apply_edits import populate_locator_evidence
from src.editing.edit_locator import LocatorResult, locator_evidence_from_result
from src.input.extractor import ExtractedSpec, ParagraphMapping
from src.orchestration.resume_state import (
    deserialize_finding,
    serialize_finding,
)
from src.output.report_exporter import export_report
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers (mirror style of test_chunk_n_report_status)
# ---------------------------------------------------------------------------

def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.8,
    action: str = "EDIT",
    existing: str | None = "2019 CBC",
    replacement: str | None = "2025 CBC",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
    locator_evidence: dict | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
    )
    f.verification = verification
    f.locator_evidence = locator_evidence
    return f


def _verification(
    verdict: str = "CONFIRMED",
    *,
    grounded: bool = True,
    cache_status: str = "miss",
    explanation: str = "Verified against CBC §1234.",
    sources: list[str] | None = None,
    rejected: list[dict] | None = None,
    correction: str | None = None,
    model_used: str = "claude-sonnet-4-6",
    verification_mode: str = "standard_reasoning",
    web_search_requests: int = 3,
    source_quote: str = (
        "The 2025 California Building Code adopts the 2024 International "
        "Building Code with California amendments."
    ),
    escalation_attempted: bool = False,
    initial_verdict: str = "",
    initial_model: str = "",
    escalation_reason: str = "",
    escalation_changed_verdict: bool = False,
) -> VerificationResult:
    if sources is None:
        sources = (
            ["https://codes.iccsafe.org/content/CBC2025"]
            if grounded and verdict.upper() in ("CONFIRMED", "CORRECTED")
            else []
        )
    return VerificationResult(
        verdict=verdict,
        explanation=explanation,
        sources=list(sources),
        correction=correction,
        grounded=grounded,
        cache_status=cache_status,
        rejected_sources=list(rejected or []),
        model_used=model_used,
        verification_mode=verification_mode,
        web_search_requests=web_search_requests,
        source_quote=source_quote,
        escalation_attempted=escalation_attempted,
        initial_verdict=initial_verdict,
        initial_model=initial_model,
        escalation_reason=escalation_reason,
        escalation_changed_verdict=escalation_changed_verdict,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Helpers — paragraph map fixture (mirrors what extract_text_from_docx emits)
# ---------------------------------------------------------------------------

def _paragraph_map(text_blocks: list[tuple[str, str | None]]) -> list[ParagraphMapping]:
    """Build a minimal paragraph_map for locator testing.

    Each tuple is ``(text, element_id)``. element_id is the optional id
    the locator prefers when ``Finding.evidenceElementId`` matches.
    """
    out: list[ParagraphMapping] = []
    for idx, (text, element_id) in enumerate(text_blocks):
        out.append(
            ParagraphMapping(
                body_index=idx,
                element_type="paragraph",
                text=text,
                table_index=None,
                row_index=None,
                cell_index=None,
                element_id=element_id or "",
            )
        )
    return out


# ===========================================================================
# locator_evidence_from_result — wire-format helper
# ===========================================================================

class TestLocatorEvidenceFromResult:
    def test_extracts_best_location_fields(self):
        paragraph_map = _paragraph_map([("2019 CBC requires", "p-1")])
        from src.editing.edit_locator import locate_edit
        finding = _finding(
            existing="2019 CBC",
            replacement="2025 CBC",
            edit_proposal=EditProposal(
                action_type="EDIT",
                existing_text="2019 CBC",
                replacement_text="2025 CBC",
                edit_confidence=0.9,
            ),
        )
        result = locate_edit(finding, paragraph_map)
        evidence = locator_evidence_from_result(result)
        # The locator found the existing text exactly, so every field
        # should populate. Status is the locator's "matched" /
        # "not_found" / "ambiguous" enum.
        assert evidence["status"] == "matched"
        assert evidence["match_method"] in ("exact", "normalized")
        assert 0.0 < evidence["match_confidence"] <= 1.0
        # Safety category is one of the four closed-set values.
        assert evidence["safety_category"] in (
            "AUTO_SAFE", "AUTO_WITH_CAUTION", "MANUAL_REVIEW", "REPORT_ONLY",
        )
        # Element id is preserved when the matched paragraph has one.
        assert evidence["element_id"] == "p-1"

    def test_no_locations_returns_empty_fields(self):
        # An empty locations list (not_found / ambiguous outcomes) must
        # not invent confidence or element id values.
        result = LocatorResult(
            finding=_finding(),
            status="not_found",
            locations=[],
            replacement_text=None,
            action_type="EDIT",
        )
        evidence = locator_evidence_from_result(result)
        assert evidence["status"] == "not_found"
        assert evidence["match_method"] == ""
        assert evidence["match_confidence"] == 0.0
        assert evidence["element_id"] == ""


# ===========================================================================
# populate_locator_evidence — pipeline-side helper
# ===========================================================================

class TestPopulateLocatorEvidence:
    def test_stamps_evidence_on_finding_with_proposal(self):
        paragraph_map = _paragraph_map([("2019 CBC §1234", "p-1")])
        spec = ExtractedSpec(
            filename="Section_23.docx",
            content="2019 CBC §1234",
            word_count=3,
            paragraph_map=paragraph_map,
        )
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="2019 CBC",
            replacement_text="2025 CBC",
            edit_confidence=0.85,
        )
        finding = _finding(
            file="Section_23.docx",
            existing="2019 CBC",
            replacement="2025 CBC",
            edit_proposal=proposal,
        )
        populate_locator_evidence([finding], [spec])
        assert finding.locator_evidence is not None
        assert finding.locator_evidence["status"] == "matched"
        assert finding.locator_evidence["match_method"] in ("exact", "normalized")

    def test_no_proposal_leaves_evidence_none(self):
        # REPORT_ONLY findings never produce a proposal -> no locator
        # evidence should be stamped.
        spec = ExtractedSpec(
            filename="Section_23.docx",
            content="2019 CBC §1234",
            word_count=3,
            paragraph_map=_paragraph_map([("2019 CBC §1234", None)]),
        )
        finding = Finding(
            severity="MEDIUM",
            fileName="Section_23.docx",
            section="2.1",
            issue="Coordination-only finding",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference=None,
        )
        populate_locator_evidence([finding], [spec])
        assert finding.locator_evidence is None

    def test_file_unavailable_leaves_evidence_none(self):
        # When the spec for a finding's file is not present in
        # extracted_specs, the helper must leave locator_evidence
        # None rather than inventing one or crashing.
        spec = ExtractedSpec(
            filename="Other_File.docx",
            content="unrelated content",
            word_count=2,
            paragraph_map=_paragraph_map([("unrelated content", None)]),
        )
        finding = _finding(
            file="Section_23.docx",
            edit_proposal=EditProposal(
                action_type="EDIT",
                existing_text="2019 CBC",
                replacement_text="2025 CBC",
                edit_confidence=0.85,
            ),
        )
        populate_locator_evidence([finding], [spec])
        assert finding.locator_evidence is None


# ===========================================================================
# Resume-state serialization
# ===========================================================================

class TestResumeStateRoundTrip:
    def test_locator_evidence_round_trip(self):
        evidence = {
            "status": "matched",
            "match_method": "id",
            "match_confidence": 0.97,
            "safety_category": "AUTO_SAFE",
            "element_id": "p-42",
        }
        finding = _finding(
            edit_proposal=EditProposal(
                action_type="EDIT",
                existing_text="x",
                replacement_text="y",
                edit_confidence=0.9,
            ),
            locator_evidence=evidence,
        )
        payload = serialize_finding(finding)
        restored = deserialize_finding(payload)
        assert restored.locator_evidence == evidence

    def test_legacy_payload_has_no_locator_evidence(self):
        # Resume payloads written before Chunk 4 simply lack the key —
        # the deserializer must default to None instead of crashing.
        finding = _finding(
            edit_proposal=EditProposal(
                action_type="EDIT",
                existing_text="x",
                replacement_text="y",
                edit_confidence=0.9,
            ),
        )
        payload = serialize_finding(finding)
        payload.pop("locator_evidence", None)
        restored = deserialize_finding(payload)
        assert restored.locator_evidence is None

    def test_malformed_evidence_falls_back_to_none(self):
        # A malformed locator_evidence value (not a dict) should not
        # break the deserializer.
        finding = _finding()
        payload = serialize_finding(finding)
        payload["locator_evidence"] = "not a dict"
        restored = deserialize_finding(payload)
        assert restored.locator_evidence is None


# ===========================================================================
# Report rendering — evidence panel under Sources
# ===========================================================================

class TestEvidencePanelRendering:
    @pytest.fixture
    def review_with_verification(self) -> ReviewResult:
        verified = _finding(
            severity="CRITICAL",
            issue="Stale CBC reference",
            verification=_verification(
                verdict="CONFIRMED",
                model_used="claude-sonnet-4-6",
                verification_mode="standard_reasoning",
                web_search_requests=4,
                source_quote=(
                    "The 2025 California Building Code requires fire "
                    "sprinklers in all K-12 facilities."
                ),
                explanation="The 2025 CBC supersedes the 2019 cycle.",
                sources=["https://codes.iccsafe.org/content/CBC2025"],
            ),
        )
        return ReviewResult(findings=[verified])

    def test_renders_verifier_model(self, tmp_path: Path, review_with_verification):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        assert "Verifier model:" in text
        assert "claude-sonnet-4-6" in text

    def test_renders_verification_mode(self, tmp_path: Path, review_with_verification):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        assert "Verification mode:" in text
        # Human-readable form, not the raw underscore string.
        assert "Standard reasoning" in text

    def test_renders_search_budget(self, tmp_path: Path, review_with_verification):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        # CRITICAL severity has an 8-call budget; the verification used
        # 4 of them. The exact "N of M" string should appear.
        assert "Search budget used:" in text
        assert "4 of 8 searches used" in text

    def test_renders_source_quote(self, tmp_path: Path, review_with_verification):
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        assert "Source quote" in text
        assert "fire sprinklers in all K-12 facilities" in text

    def test_renders_rationale_inside_evidence_panel(
        self, tmp_path: Path, review_with_verification
    ):
        # Chunk 4: rationale moves from above the Sources heading to
        # under it (next to the source quote). The label is preserved
        # so Chunk N tests still pass.
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review_with_verification), out)
        text = _all_text_from(Document(str(out)))
        assert "Verification rationale:" in text
        # Rationale text must follow source quote in the document
        # body — the panel order matters for readability.
        quote_idx = text.find("fire sprinklers in all K-12 facilities")
        rationale_idx = text.find("The 2025 CBC supersedes")
        assert quote_idx >= 0
        assert rationale_idx >= 0
        assert rationale_idx > quote_idx, (
            "rationale should render after the source quote inside the panel"
        )

    def test_local_skip_panel_suppresses_search_budget(self, tmp_path: Path):
        # A LOCAL_SKIP verification has nothing to do with the web
        # search budget; rendering "0 of N searches" would be
        # misleading. The line is suppressed.
        local = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verification(
                verdict="UNVERIFIED",
                grounded=False,
                cache_status="local_skip",
                verification_mode="local_skip",
                model_used="local",
                web_search_requests=0,
                source_quote="",
                explanation="Locally classified — placeholder GRIPES.",
                sources=[],
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[local])), out
        )
        text = _all_text_from(Document(str(out)))
        # Mode line still renders; budget line suppressed.
        assert "Local skip" in text
        assert "searches used" not in text


# ===========================================================================
# Escalation history rendering
# ===========================================================================

class TestEscalationHistoryRendering:
    def test_renders_when_escalation_attempted(self, tmp_path: Path):
        # Escalation history must surface when escalation_attempted is
        # set, even if the verdict did not change — the audit trail is
        # the point.
        escalated = _finding(
            severity="CRITICAL",
            verification=_verification(
                verdict="CONFIRMED",
                model_used="claude-opus-4-7",
                verification_mode="deep_reasoning",
                escalation_attempted=True,
                initial_verdict="UNVERIFIED",
                initial_model="claude-sonnet-4-6",
                escalation_reason="ungrounded_critical_high",
                escalation_changed_verdict=True,
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[escalated])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Escalation history:" in text
        assert "UNVERIFIED" in text
        assert "claude-sonnet-4-6" in text
        assert "claude-opus-4-7" in text
        assert "ungrounded_critical_high" in text
        # When the escalation changed the verdict, the panel must flag
        # the disagreement.
        assert "models disagreed" in text

    def test_omitted_when_no_escalation(self, tmp_path: Path):
        verified = _finding(
            verification=_verification(
                verdict="CONFIRMED",
                escalation_attempted=False,
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[verified])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Escalation history:" not in text


# ===========================================================================
# Edit Target Evidence panel
# ===========================================================================

class TestEditTargetEvidencePanel:
    def test_renders_when_locator_evidence_populated(self, tmp_path: Path):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="2019 CBC",
            replacement_text="2025 CBC",
            edit_confidence=0.85,
        )
        finding = _finding(
            edit_proposal=proposal,
            verification=_verification(),
            locator_evidence={
                "status": "matched",
                "match_method": "id",
                "match_confidence": 0.97,
                "safety_category": "AUTO_SAFE",
                "element_id": "p-42",
            },
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[finding])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Edit Target Evidence" in text
        assert "Locator status:" in text
        assert "Match method:" in text
        # Human-readable label for the id match method.
        assert "Id-anchored" in text
        assert "Match confidence:" in text
        assert "97%" in text
        assert "Safety category:" in text
        assert "AUTO_SAFE" in text
        assert "Element id:" in text
        assert "p-42" in text

    def test_no_panel_when_locator_evidence_missing(self, tmp_path: Path):
        finding = _finding(
            edit_proposal=EditProposal(
                action_type="EDIT",
                existing_text="x",
                replacement_text="y",
                edit_confidence=0.9,
            ),
            verification=_verification(),
            locator_evidence=None,
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[finding])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Edit Target Evidence" not in text

    def test_no_panel_for_report_only_finding(self, tmp_path: Path):
        # REPORT_ONLY findings never have locator evidence, so the
        # panel never renders.
        finding = Finding(
            severity="MEDIUM",
            fileName="Section_23.docx",
            section="2.1",
            issue="Coordination-only",
            actionType="REPORT_ONLY",
            existingText=None,
            replacementText=None,
            codeReference=None,
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[finding])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Edit Target Evidence" not in text

    def test_not_found_status_renders_warning(self, tmp_path: Path):
        # A locator that failed to find the edit target should still
        # render the panel so the reviewer can see "the edit was not
        # locatable" rather than silently dropping the data.
        finding = _finding(
            edit_proposal=EditProposal(
                action_type="EDIT",
                existing_text="missing text",
                replacement_text="anything",
                edit_confidence=0.85,
            ),
            verification=_verification(),
            locator_evidence={
                "status": "not_found",
                "match_method": "",
                "match_confidence": 0.0,
                "safety_category": "REPORT_ONLY",
                "element_id": "",
            },
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[finding])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Edit Target Evidence" in text
        assert "not_found" in text


# ===========================================================================
# Integration — every CONFIRMED / CORRECTED finding gets a source quote
# ===========================================================================

class TestEvidencePanelSuccessCriteria:
    """Mirrors the success criteria from the chunk plan."""

    def test_confirmed_finding_shows_source_quote_and_verifier_model(
        self, tmp_path: Path
    ):
        confirmed = _finding(
            verification=_verification(
                verdict="CONFIRMED",
                source_quote="CBC §1234 explicitly requires fire sprinklers.",
                model_used="claude-sonnet-4-6",
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[confirmed])), out
        )
        text = _all_text_from(Document(str(out)))
        # Both source quote and verifier model must appear in the
        # output for any CONFIRMED finding.
        assert "CBC §1234 explicitly requires fire sprinklers." in text
        assert "claude-sonnet-4-6" in text

    def test_corrected_finding_shows_source_quote_and_verifier_model(
        self, tmp_path: Path
    ):
        corrected = _finding(
            verification=_verification(
                verdict="CORRECTED",
                correction="Replace with 2025 CBC",
                source_quote="The 2025 CBC supersedes the 2019 cycle.",
                model_used="claude-opus-4-7",
            ),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[corrected])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "The 2025 CBC supersedes the 2019 cycle." in text
        assert "claude-opus-4-7" in text

    def test_auto_edit_candidate_shows_locator_evidence(self, tmp_path: Path):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="2019 CBC",
            replacement_text="2025 CBC",
            edit_confidence=0.92,
        )
        candidate = _finding(
            edit_proposal=proposal,
            verification=_verification(verdict="CONFIRMED"),
            locator_evidence={
                "status": "matched",
                "match_method": "exact",
                "match_confidence": 0.95,
                "safety_category": "AUTO_SAFE",
                "element_id": "p-1",
            },
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[candidate])), out
        )
        text = _all_text_from(Document(str(out)))
        # Locator method + confidence must be visible for the auto-edit
        # candidate.
        assert "Match method:" in text
        assert "Match confidence:" in text
        assert "95%" in text
