"""Tests for extraction of body-walk-invisible content (TRUST_AUDIT P0-6).

python-docx's ``<w:body>`` walk surfaces paragraphs and tables, but a
requirement can also be authored in places the walk never reaches:

* **Text boxes** — text nested in ``<w:txbxContent>`` inside ``<w:drawing>``
  (modern DrawingML) or ``<w:pict>`` (legacy VML) runs. ``Paragraph.text``
  does not descend into them.
* **Footnotes / endnotes** — stored in their own package parts
  (``word/footnotes.xml`` / ``word/endnotes.xml``), not under ``<w:body>``.

Before this change those sources were silently dropped, so a real defect
authored in a callout text box or a footnote produced no finding — a
"miss a real problem" trust gap. The extractor now renders each kind as
its own labeled block after the body, preserving the reconstruction
invariant (the paragraph map's text joins back to ``content``) and
stamping stable element ids (``tb<box>p<para>`` / ``fn<id>p<para>`` /
``en<id>p<para>``, with ``meta:tb`` / ``meta:fn`` / ``meta:en`` delimiters).
A spec with none of these produces byte-identical output to before.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml

from src.input.extractor import (
    _ENDNOTES_CONTENT_TYPE,
    _FOOTNOTES_CONTENT_TYPE,
    _accept_all_paragraph_text,
    extract_text_from_docx,
)

# ---------------------------------------------------------------------------
# OOXML fragment builders
# ---------------------------------------------------------------------------

_FOOTNOTES_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
_ENDNOTES_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
)


def _drawingml_textbox_paragraph(text: str) -> str:
    """A body ``<w:p>`` whose run carries a DrawingML text box (``wps:txbx``)."""
    return (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        "<w:r><w:drawing><wp:inline><a:graphic>"
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp><wps:txbx><w:txbxContent>"
        f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        "</w:txbxContent></wps:txbx></wps:wsp>"
        "</a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>"
    )


def _vml_textbox_paragraph(text: str) -> str:
    """A body ``<w:p>`` whose run carries a legacy VML text box."""
    return (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:v="urn:schemas-microsoft-com:vml">'
        "<w:r><w:pict><v:shape><v:textbox><w:txbxContent>"
        f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )


# ---------------------------------------------------------------------------
# Tracked-changes (revision) fragment builders
# ---------------------------------------------------------------------------

_W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _run(text: str, *, tag: str = "w:t") -> str:
    """A ``<w:r>`` whose text node is ``tag`` (``w:t`` normally, ``w:delText``
    for deleted runs)."""
    return f'<w:r><{tag} xml:space="preserve">{text}</{tag}></w:r>'


def _ins(inner: str) -> str:
    """Wrap runs as a tracked insertion (``<w:ins>``)."""
    return f'<w:ins w:id="1" w:author="A" w:date="2026-01-01T00:00:00Z">{inner}</w:ins>'


def _del(inner: str) -> str:
    """Wrap runs as a tracked deletion (``<w:del>``); inner runs use ``w:delText``."""
    return f'<w:del w:id="2" w:author="A" w:date="2026-01-01T00:00:00Z">{inner}</w:del>'


def _move_from(inner: str) -> str:
    """Wrap runs as the source of a tracked move (removed on accept)."""
    return (
        '<w:moveFrom w:id="3" w:author="A" w:date="2026-01-01T00:00:00Z" '
        f'w:name="m1">{inner}</w:moveFrom>'
    )


def _move_to(inner: str) -> str:
    """Wrap runs as the destination of a tracked move (kept on accept)."""
    return (
        '<w:moveTo w:id="4" w:author="A" w:date="2026-01-01T00:00:00Z" '
        f'w:name="m1">{inner}</w:moveTo>'
    )


def _revision_paragraph(inner: str) -> str:
    """A body ``<w:p>`` (namespace declared once) holding ``inner`` run/revision XML."""
    return f"<w:p {_W_NS}>{inner}</w:p>"


def _drawingml_textbox_with_runs(inner: str) -> str:
    """A DrawingML text box whose single inner paragraph holds ``inner`` runs.

    Like :func:`_drawingml_textbox_paragraph` but the caller supplies the run /
    revision XML, so a text box can carry tracked changes.
    """
    return (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        "<w:r><w:drawing><wp:inline><a:graphic>"
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp><wps:txbx><w:txbxContent>"
        f"<w:p>{inner}</w:p>"
        "</w:txbxContent></wps:txbx></wps:wsp>"
        "</a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>"
    )


def _notes_xml(root_tag: str, note_tag: str, real_notes: list[tuple[str, str]]) -> bytes:
    """Build a footnotes/endnotes part body.

    Always seeds the two structural notes Word emits (``separator`` id -1,
    ``continuationSeparator`` id 0) so the test exercises the type-skipping
    path. ``real_notes`` is a list of ``(id, text)`` authored notes.
    """
    parts = [
        f'<?xml version="1.0"?><w:{root_tag} '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
        f'<w:{note_tag} w:type="separator" w:id="-1">'
        f"<w:p><w:r><w:separator/></w:r></w:p></w:{note_tag}>",
        f'<w:{note_tag} w:type="continuationSeparator" w:id="0">'
        f"<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:{note_tag}>",
    ]
    for note_id, text in real_notes:
        parts.append(
            f'<w:{note_tag} w:id="{note_id}">'
            f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:{note_tag}>"
        )
    parts.append(f"</w:{root_tag}>")
    return "".join(parts).encode("utf-8")


def _attach_part(doc, *, partname: str, content_type: str, rel_type: str, blob: bytes) -> None:
    pkg = doc.part.package
    part = Part(PackURI(partname), content_type, blob, pkg)
    doc.part.relate_to(part, rel_type)


def _build_docx(
    tmp_path: Path,
    *,
    body_paras: list[str] | None = None,
    body_xml: list[str] | None = None,
    textbox_xml: list[str] | None = None,
    footnotes: list[tuple[str, str]] | None = None,
    endnotes: list[tuple[str, str]] | None = None,
    footnotes_blob: bytes | None = None,
    filename: str = "spec.docx",
) -> Path:
    """Assemble an in-memory .docx with the requested supplemental content and
    save it to disk so extraction round-trips through the package layer.

    ``body_xml`` appends raw body-level elements (e.g. ``<w:p>`` carrying
    revision markup) after the plain ``body_paras``.
    """
    doc = Document()
    for text in body_paras or []:
        doc.add_paragraph(text)
    for xml in body_xml or []:
        doc.element.body.append(parse_xml(xml))
    for xml in textbox_xml or []:
        doc.element.body.append(parse_xml(xml))
    if footnotes is not None or footnotes_blob is not None:
        blob = (
            footnotes_blob
            if footnotes_blob is not None
            else _notes_xml("footnotes", "footnote", footnotes or [])
        )
        _attach_part(
            doc,
            partname="/word/footnotes.xml",
            content_type=_FOOTNOTES_CONTENT_TYPE,
            rel_type=_FOOTNOTES_REL_TYPE,
            blob=blob,
        )
    if endnotes is not None:
        _attach_part(
            doc,
            partname="/word/endnotes.xml",
            content_type=_ENDNOTES_CONTENT_TYPE,
            rel_type=_ENDNOTES_REL_TYPE,
            blob=_notes_xml("endnotes", "endnote", endnotes),
        )
    out = tmp_path / filename
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Text boxes
# ---------------------------------------------------------------------------


class TestTextBoxExtraction:
    def test_drawingml_textbox_captured(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["BODY: seismic bracing per CBC."],
            textbox_xml=[_drawingml_textbox_paragraph("R-12 insulation required.")],
        )
        spec = extract_text_from_docx(path)
        assert "R-12 insulation required." in spec.content
        assert "[Text Box] R-12 insulation required." in spec.content
        assert "===== TEXT BOX CONTENT =====" in spec.content
        # Body text still captured.
        assert "BODY: seismic bracing per CBC." in spec.content

    def test_vml_textbox_captured(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            textbox_xml=[_vml_textbox_paragraph("Bronze body valves, 150 psi.")],
        )
        spec = extract_text_from_docx(path)
        assert "[Text Box] Bronze body valves, 150 psi." in spec.content

    def test_multiple_textboxes_distinct_ids_in_order(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            textbox_xml=[
                _drawingml_textbox_paragraph("First box."),
                _vml_textbox_paragraph("Second box."),
            ],
        )
        spec = extract_text_from_docx(path)
        tb_ids = [m.element_id for m in spec.paragraph_map if m.element_type == "textbox"]
        assert tb_ids == ["tb0p0", "tb1p0"]
        # Document order preserved: first box before second in the content.
        assert spec.content.index("First box.") < spec.content.index("Second box.")

    def test_textbox_inside_table_cell_captured(self, tmp_path: Path):
        # A text box anchored inside a table cell is still reached by the
        # body-level descendant search.
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = "Cell text."
        cell._tc.append(parse_xml(_drawingml_textbox_paragraph("Boxed note in cell.")))
        out = tmp_path / "tbl.docx"
        doc.save(out)
        spec = extract_text_from_docx(out)
        assert "Cell text." in spec.content
        assert "[Text Box] Boxed note in cell." in spec.content

    def test_empty_textbox_produces_no_entry(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Body."],
            textbox_xml=[_drawingml_textbox_paragraph("   ")],
        )
        spec = extract_text_from_docx(path)
        assert "===== TEXT BOX CONTENT =====" not in spec.content
        assert not any(m.element_type == "textbox" for m in spec.paragraph_map)

    def test_paragraph_with_text_and_textbox_captures_both_once(self, tmp_path: Path):
        # A run carrying a text box can sit in a paragraph that also has
        # visible text. Both are captured; neither is duplicated.
        mixed = (
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
            ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
            ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
            "<w:r><w:t>Visible inline text.</w:t></w:r>"
            "<w:r><w:drawing><wp:inline><a:graphic>"
            '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
            "<wps:wsp><wps:txbx><w:txbxContent>"
            "<w:p><w:r><w:t>Boxed text.</w:t></w:r></w:p>"
            "</w:txbxContent></wps:txbx></wps:wsp>"
            "</a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>"
        )
        path = _build_docx(tmp_path, textbox_xml=[mixed])
        spec = extract_text_from_docx(path)
        assert spec.content.count("Visible inline text.") == 1
        assert spec.content.count("Boxed text.") == 1
        assert "[Text Box] Boxed text." in spec.content

    def test_clean_doc_has_no_textbox_block(self, tmp_path: Path):
        path = _build_docx(tmp_path, body_paras=["Just body text."])
        spec = extract_text_from_docx(path)
        assert "TEXT BOX CONTENT" not in spec.content
        assert spec.content == "Just body text."


# ---------------------------------------------------------------------------
# Footnotes / endnotes
# ---------------------------------------------------------------------------


class TestFootnoteEndnoteExtraction:
    def test_footnote_captured_separators_excluded(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Body with a footnote ref."],
            footnotes=[("1", "Coordinate with structural S-101.")],
        )
        spec = extract_text_from_docx(path)
        assert "===== FOOTNOTE CONTENT =====" in spec.content
        assert "[Footnote 1] Coordinate with structural S-101." in spec.content
        # The structural separator/continuationSeparator notes (ids -1, 0)
        # carry no authored text and must never reach the review surface.
        assert "[Footnote -1]" not in spec.content
        assert "[Footnote 0]" not in spec.content

    def test_endnote_captured(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Body."],
            endnotes=[("2", "See appendix B fixture schedule.")],
        )
        spec = extract_text_from_docx(path)
        assert "===== ENDNOTE CONTENT =====" in spec.content
        assert "[Endnote 2] See appendix B fixture schedule." in spec.content

    def test_footnote_element_ids(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Body."],
            footnotes=[("1", "First note."), ("3", "Second note.")],
        )
        spec = extract_text_from_docx(path)
        fn_ids = [m.element_id for m in spec.paragraph_map if m.element_type == "footnote"]
        assert fn_ids == ["fn1p0", "fn3p0"]

    def test_absent_notes_parts_produce_no_blocks(self, tmp_path: Path):
        path = _build_docx(tmp_path, body_paras=["Body only."])
        spec = extract_text_from_docx(path)
        assert "FOOTNOTE CONTENT" not in spec.content
        assert "ENDNOTE CONTENT" not in spec.content

    def test_footnotes_part_with_only_separators_emits_nothing(self, tmp_path: Path):
        # A part that has the structural seeds but no authored notes should
        # not produce a (header-less) block.
        path = _build_docx(tmp_path, body_paras=["Body."], footnotes=[])
        spec = extract_text_from_docx(path)
        assert "FOOTNOTE CONTENT" not in spec.content

    def test_malformed_notes_part_does_not_crash(self, tmp_path: Path):
        # A corrupt notes part must never sink the whole extraction; body
        # text is the primary deliverable.
        path = _build_docx(
            tmp_path,
            body_paras=["Body survives."],
            footnotes_blob=b"<not-valid-xml<<<",
        )
        spec = extract_text_from_docx(path)
        assert "Body survives." in spec.content
        assert "FOOTNOTE CONTENT" not in spec.content


# ---------------------------------------------------------------------------
# Invariants across all supplemental kinds
# ---------------------------------------------------------------------------


class TestSupplementalInvariants:
    def test_reconstruction_invariant_holds_with_all_kinds(self, tmp_path: Path):
        # extract_text_from_docx raises ValueError if the paragraph map does
        # not reconstruct the content; a clean return proves the invariant.
        path = _build_docx(
            tmp_path,
            body_paras=["PART 1 GENERAL", "Body requirement."],
            textbox_xml=[_drawingml_textbox_paragraph("Boxed requirement.")],
            footnotes=[("1", "A footnote.")],
            endnotes=[("2", "An endnote.")],
        )
        spec = extract_text_from_docx(path)
        reconstructed = "\n\n".join(m.text for m in spec.paragraph_map)
        assert reconstructed == spec.content

    def test_all_element_ids_unique(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Body one.", "Body two."],
            textbox_xml=[
                _drawingml_textbox_paragraph("Box A."),
                _vml_textbox_paragraph("Box B."),
            ],
            footnotes=[("1", "Note A.")],
            endnotes=[("2", "Note B.")],
        )
        spec = extract_text_from_docx(path)
        ids = [m.element_id for m in spec.paragraph_map]
        assert len(ids) == len(set(ids)), f"duplicate element ids: {ids}"
        # The supplemental delimiters are all present and distinct.
        assert {"meta:tb", "meta:fn", "meta:en"}.issubset(set(ids))

    def test_supplemental_block_order(self, tmp_path: Path):
        # Blocks render in a stable order: text boxes, then footnotes, then
        # endnotes, after the body.
        path = _build_docx(
            tmp_path,
            body_paras=["Body."],
            textbox_xml=[_drawingml_textbox_paragraph("Box.")],
            footnotes=[("1", "Foot.")],
            endnotes=[("2", "End.")],
        )
        spec = extract_text_from_docx(path)
        c = spec.content
        assert (
            c.index("Body.")
            < c.index("TEXT BOX CONTENT")
            < c.index("FOOTNOTE CONTENT")
            < c.index("ENDNOTE CONTENT")
        )

    def test_word_count_includes_supplemental_text(self, tmp_path: Path):
        # Supplemental content contributes to the reviewable word count.
        without = extract_text_from_docx(
            _build_docx(tmp_path, body_paras=["Body."], filename="a.docx")
        )
        with_box = extract_text_from_docx(
            _build_docx(
                tmp_path,
                body_paras=["Body."],
                textbox_xml=[_drawingml_textbox_paragraph("Extra boxed words here.")],
                filename="b.docx",
            )
        )
        assert with_box.word_count > without.word_count


# ---------------------------------------------------------------------------
# Tracked changes (Word revision markup) → Accept-All view
# ---------------------------------------------------------------------------


class TestAcceptAllParagraphHelper:
    """Unit tests for the revision-aware paragraph walk."""

    def _text(self, inner: str) -> str:
        return _accept_all_paragraph_text(parse_xml(_revision_paragraph(inner)))

    def test_insertion_kept(self):
        assert self._text(_run("a ") + _ins(_run("b ")) + _run("c")) == "a b c"

    def test_deletion_dropped(self):
        assert self._text(_run("a ") + _del(_run("b ", tag="w:delText")) + _run("c")) == "a c"

    def test_combined_delete_then_insert(self):
        got = self._text(_run("v") + _del(_run("1", tag="w:delText")) + _ins(_run("2")))
        assert got == "v2"

    def test_move_from_dropped_move_to_kept(self):
        assert self._text(_move_from(_run("x")) + _move_to(_run("y"))) == "y"

    def test_insertion_of_deletion_nets_to_empty(self):
        # An inserted run that was then deleted disappears once both are
        # accepted (the <w:del> inside the <w:ins> is pruned).
        assert self._text(_ins(_del(_run("gone", tag="w:delText")))) == ""

    def test_clean_paragraph_matches_plain_text(self):
        assert self._text(_run("plain only")) == "plain only"

    def test_run_level_tab_and_break_translation(self):
        # Tab / break translation is delegated to CT_R.text, matching
        # python-docx Paragraph.text fidelity.
        got = self._text(_run("a") + "<w:r><w:tab/></w:r>" + _run("b"))
        assert got == "a\tb"


class TestTrackedChangesExtraction:
    def test_body_insertion_included(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_xml=[
                _revision_paragraph(
                    _run("The duct shall be ")
                    + _ins(_run("galvanized steel "))
                    + _run("per spec.")
                )
            ],
        )
        spec = extract_text_from_docx(path)
        assert spec.content == "The duct shall be galvanized steel per spec."
        assert spec.tracked_changes_detected is True

    def test_body_deletion_excluded(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_xml=[
                _revision_paragraph(
                    _run("The pipe shall be ")
                    + _del(_run("copper ", tag="w:delText"))
                    + _run("type L.")
                )
            ],
        )
        spec = extract_text_from_docx(path)
        assert spec.content == "The pipe shall be type L."
        assert "copper" not in spec.content
        assert spec.tracked_changes_detected is True

    def test_combined_edition_change_resolves_to_accept_all(self, tmp_path: Path):
        # The canonical "delete old year, insert new year" redline that today
        # collapses to "Comply with  CBC." under plain python-docx.
        path = _build_docx(
            tmp_path,
            body_xml=[
                _revision_paragraph(
                    _run("Comply with ")
                    + _del(_run("2019", tag="w:delText"))
                    + _ins(_run("2025"))
                    + _run(" CBC.")
                )
            ],
        )
        spec = extract_text_from_docx(path)
        assert spec.content == "Comply with 2025 CBC."

    def test_move_keeps_destination_only(self, tmp_path: Path):
        # A Word move = moveFrom (source, removed) + moveTo (destination, kept).
        path = _build_docx(
            tmp_path,
            body_xml=[
                _revision_paragraph(_run("A ") + _move_from(_run("moved clause ")) + _run("B.")),
                _revision_paragraph(_run("C ") + _move_to(_run("moved clause ")) + _run("D.")),
            ],
        )
        spec = extract_text_from_docx(path)
        assert spec.content == "A B.\n\nC moved clause D."
        assert spec.content.count("moved clause") == 1

    def test_fully_deleted_paragraph_is_skipped(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Kept paragraph."],
            body_xml=[_revision_paragraph(_del(_run("all gone", tag="w:delText")))],
        )
        spec = extract_text_from_docx(path)
        assert spec.content == "Kept paragraph."
        # The empty (fully-deleted) paragraph contributes nothing to the map.
        assert all(m.text for m in spec.paragraph_map)

    def test_clean_doc_unchanged_and_flag_false(self, tmp_path: Path):
        path = _build_docx(tmp_path, body_paras=["Plain one.", "Plain two."])
        spec = extract_text_from_docx(path)
        assert spec.content == "Plain one.\n\nPlain two."
        assert spec.tracked_changes_detected is False

    def test_tracked_change_in_table_cell(self, tmp_path: Path):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        inner = _run("Edition: ") + _del(_run("2019", tag="w:delText")) + _ins(_run("2025"))
        cell.paragraphs[0]._p.getparent().replace(
            cell.paragraphs[0]._p, parse_xml(_revision_paragraph(inner))
        )
        out = tmp_path / "table.docx"
        doc.save(out)
        spec = extract_text_from_docx(out)
        assert "Edition: 2025" in spec.content
        assert "2019" not in spec.content
        assert spec.tracked_changes_detected is True

    def test_tracked_change_in_textbox(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Body."],
            textbox_xml=[
                _drawingml_textbox_with_runs(
                    _run("Note: ")
                    + _del(_run("old ", tag="w:delText"))
                    + _ins(_run("new"))
                )
            ],
        )
        spec = extract_text_from_docx(path)
        assert "[Text Box] Note: new" in spec.content
        assert "old" not in spec.content
        assert spec.tracked_changes_detected is True

    def test_tracked_change_in_footnote_extracted_and_flagged(self, tmp_path: Path):
        # Revisions inside the footnotes part are resolved to accept-all AND
        # trip the advisory flag — the body here is clean, so the flag must
        # come from the note-part scan alone.
        blob = (
            '<?xml version="1.0"?><w:footnotes '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1">'
            "<w:p><w:r><w:separator/></w:r></w:p></w:footnote>"
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>"
            '<w:footnote w:id="1"><w:p>'
            + _run("See ")
            + _del(_run("2019", tag="w:delText"))
            + _ins(_run("2025"))
            + _run(" code.")
            + "</w:p></w:footnote></w:footnotes>"
        ).encode("utf-8")
        path = _build_docx(tmp_path, body_paras=["Body."], footnotes_blob=blob)
        spec = extract_text_from_docx(path)
        assert "[Footnote 1] See 2025 code." in spec.content
        assert "2019" not in spec.content
        assert spec.tracked_changes_detected is True

    def test_tracked_change_in_header(self, tmp_path: Path):
        # The P2 scenario: a redline confined to a header (clean body) must
        # still be resolved to accept-all AND trip the advisory flag, since
        # headers live in a separate package part outside <w:body>.
        doc = Document()
        doc.add_paragraph("Body.")
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        inner = _run("Rev ") + _del(_run("A ", tag="w:delText")) + _ins(_run("B"))
        header.paragraphs[0]._p.getparent().replace(
            header.paragraphs[0]._p, parse_xml(_revision_paragraph(inner))
        )
        out = tmp_path / "hdr.docx"
        doc.save(out)
        spec = extract_text_from_docx(out)
        assert "[Header] Rev B" in spec.content
        assert "Rev A" not in spec.content
        assert spec.tracked_changes_detected is True

    def test_tracked_change_in_footer_flagged(self, tmp_path: Path):
        # Same as the header case, for footers.
        doc = Document()
        doc.add_paragraph("Body.")
        footer = doc.sections[0].footer
        footer.is_linked_to_previous = False
        inner = _run("Page ") + _del(_run("old ", tag="w:delText")) + _ins(_run("new"))
        footer.paragraphs[0]._p.getparent().replace(
            footer.paragraphs[0]._p, parse_xml(_revision_paragraph(inner))
        )
        out = tmp_path / "ftr.docx"
        doc.save(out)
        spec = extract_text_from_docx(out)
        assert "[Footer] Page new" in spec.content
        assert spec.tracked_changes_detected is True

    def test_reconstruction_invariant_holds_with_revisions(self, tmp_path: Path):
        path = _build_docx(
            tmp_path,
            body_paras=["Intro."],
            body_xml=[
                _revision_paragraph(
                    _run("Use ") + _del(_run("2019", tag="w:delText")) + _ins(_run("2025"))
                )
            ],
            textbox_xml=[_drawingml_textbox_with_runs(_run("Box ") + _ins(_run("added")))],
        )
        spec = extract_text_from_docx(path)  # raises if the invariant breaks
        reconstructed = "\n\n".join(m.text for m in spec.paragraph_map)
        assert reconstructed == spec.content
