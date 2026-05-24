"""Chunk 11 tests — web_fetch capability for verification.

Chunk 11 of the Trust Upgrade adds the ``web_fetch`` server tool to the
STANDARD_REASONING and DEEP_REASONING verification modes so the verifier
can pull the full text of a URL when a search snippet is insufficient.
The contract has five surfaces:

* ``api_config.build_web_fetch_tool()`` returns the documented tool dict
  with ``citations`` enabled, ``max_content_tokens`` bounded, and the
  web_search blocklist applied so the two server tools share one
  source-quality policy.
* ``verification_routing.build_verification_tools_from_decision`` appends
  the fetch tool for STANDARD_REASONING / DEEP_REASONING decisions and
  omits it for STRICT_STRUCTURED / LOCAL_SKIP.
* ``VerificationResult`` carries ``web_fetch_requests: int`` and
  ``fetched_sources: list[str]``; both round-trip through the cache
  persist path and the resume-state serialize/deserialize helpers.
* ``_collect_fetch_evidence_detailed`` and ``_web_fetch_count`` extract
  fetch evidence from a message's content blocks and the
  ``usage.server_tool_use`` counter, in parallel to the equivalent
  web_search helpers.
* The verifier system prompt instructs the model when to use web_fetch
  vs. web_search; the per-finding evidence panel in the exported report
  renders the fetch count and a dedicated "Full-text sources consulted"
  sub-section when fetches occurred.
"""
from __future__ import annotations

import json
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from src.core.api_config import (
    DEFAULT_VERIFICATION_MAX_FETCHES,
    WEB_FETCH_MAX_CONTENT_TOKENS,
    build_web_fetch_tool,
)
from src.core.code_cycles import DEFAULT_CYCLE
from src.orchestration.resume_state import (
    deserialize_verification_result,
    serialize_verification_result,
)
from src.output.report_exporter import export_report
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.source_grounding import SearchedSource
from src.verification.verification_cache import (
    VerificationCache,
    _result_to_dict,
)
from src.verification.verification_modes import VerificationMode
from src.verification.verification_routing import (
    _tools_include_web_fetch,
    build_verification_request,
    build_verification_tools_from_decision,
    select_routing,
)
from src.verification.verifier import (
    VerificationResult,
    _apply_source_grounding,
    _collect_fetch_evidence_detailed,
    _get_verification_system_prompt,
    _web_fetch_count,
)
from tests.fixtures.fake_anthropic import (
    FakeMessage,
    FakeServerToolUseBlock,
    FakeTextBlock,
    FakeToolUseBlock,
    FakeUsage,
    FakeWebSearchResultBlock,
    sample_verification_verdict_payload,
)


# ---------------------------------------------------------------------------
# Helpers — fake message blocks for web_fetch (parallel to fake_anthropic.py
# but local to this test file because web_fetch is new and the shared
# fixture module shouldn't grow unbounded for one chunk).
# ---------------------------------------------------------------------------


@dataclass
class FakeWebFetchToolResultBlock:
    """Mimic the ``web_fetch_tool_result`` block shape used by the SDK.

    The Anthropic web_fetch tool returns a single document object inside
    the result block (unlike web_search which returns a list). The
    document carries the fetched URL, title, and content.
    """

    tool_use_id: str = "srvtoolu_fetch_1"
    content: dict[str, Any] = field(
        default_factory=lambda: {
            "type": "web_fetch_result",
            "url": "https://www.nfpa.org/13",
            "document": {
                "url": "https://www.nfpa.org/13",
                "title": "NFPA 13: Standard for the Installation of Sprinkler Systems",
                "content": "Section 10.2.5.2.1 ...",
            },
        }
    )
    type: str = "web_fetch_tool_result"


@dataclass
class FakeWebFetchToolResultErrorBlock:
    """An error-only fetch result block."""

    tool_use_id: str = "srvtoolu_fetch_err"
    content: dict[str, Any] = field(
        default_factory=lambda: {
            "type": "web_fetch_tool_result_error",
            "error_code": "url_not_accessible",
            "error_message": "The URL could not be retrieved.",
        }
    )
    type: str = "web_fetch_tool_result"


def _fake_usage_with_fetches(*, searches: int = 1, fetches: int = 0) -> FakeUsage:
    """FakeUsage with ``server_tool_use.web_fetch_requests`` populated."""
    usage = FakeUsage()
    # The verifier reads ``usage.server_tool_use.web_search_requests`` and
    # ``.web_fetch_requests`` as plain ints. A SimpleNamespace-style
    # sub-object satisfies attribute access.
    class _ServerToolUse:
        pass

    stu = _ServerToolUse()
    stu.web_search_requests = searches
    stu.web_fetch_requests = fetches
    usage.server_tool_use = stu
    return usage


def _verification_response_with_fetch(
    *,
    fetched_url: str = "https://www.nfpa.org/13",
    search_url: str = "https://www.nfpa.org/13",
    payload: dict[str, Any] | None = None,
    fetch_error: bool = False,
    n_fetches: int = 1,
) -> FakeMessage:
    """Build a fake message that includes both a search and a fetch block.

    Pairs a ``server_tool_use`` block for web_fetch (carries the URL the
    model asked to retrieve) with a ``web_fetch_tool_result`` block
    (which the SDK returns after the fetch completes). The usage
    counter is set to ``n_fetches`` so :func:`_web_fetch_count` reads
    the right value.
    """
    payload = payload if payload is not None else sample_verification_verdict_payload(
        source_quote=(
            "Section 10.2.5.2.1 The maximum distance between sprinklers "
            "shall not exceed 15 ft for ordinary hazard occupancies."
        ),
        grounded_sources=[fetched_url],
    )
    content: list[Any] = [
        # Prior web_search results — fetch can only retrieve URLs that
        # appeared in a prior search.
        FakeServerToolUseBlock(
            name="web_search", input={"query": "NFPA 13 sprinkler spacing"}
        ),
        FakeWebSearchResultBlock(
            content=[
                {
                    "type": "web_search_result",
                    "url": search_url,
                    "title": "NFPA 13",
                    "encrypted_content": "fake",
                }
            ]
        ),
        # The fetch invocation + result.
        FakeServerToolUseBlock(
            id="srvtoolu_fetch_1",
            name="web_fetch",
            input={"url": fetched_url},
        ),
    ]
    if fetch_error:
        content.append(FakeWebFetchToolResultErrorBlock())
    else:
        content.append(
            FakeWebFetchToolResultBlock(
                content={
                    "type": "web_fetch_result",
                    "url": fetched_url,
                    "document": {
                        "url": fetched_url,
                        "title": "NFPA 13",
                        "content": "Full text...",
                    },
                }
            )
        )
    content.append(
        FakeToolUseBlock(name="submit_verification_verdict", input=dict(payload))
    )
    return FakeMessage(
        content=content,
        stop_reason="tool_use",
        usage=_fake_usage_with_fetches(searches=1, fetches=n_fetches),
    )


def _finding(
    *,
    severity: str = "HIGH",
    issue: str = "NFPA 13 spacing requirement",
    code_ref: str | None = "NFPA 13 §10.2.5",
    section: str = "2.1",
    file: str = "Section_21_1000.docx",
    action: str = "EDIT",
    existing: str | None = "max spacing 12 ft",
    replacement: str | None = "max spacing 15 ft",
    verification: VerificationResult | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference=code_ref,
        confidence=0.7,
    )
    f.verification = verification
    return f


# ===========================================================================
# 1. Tool definition (api_config.build_web_fetch_tool)
# ===========================================================================


class TestBuildWebFetchTool:
    def test_tool_type_is_web_fetch_20260209(self):
        """The pinned tool name must match the documented Anthropic value."""
        tool = build_web_fetch_tool()
        assert tool["type"] == "web_fetch_20260209"

    def test_tool_name_is_web_fetch(self):
        tool = build_web_fetch_tool()
        assert tool["name"] == "web_fetch"

    def test_citations_enabled(self):
        """Citations must be on so fetch evidence flows through grounding."""
        tool = build_web_fetch_tool()
        assert tool["citations"] == {"enabled": True}

    def test_default_max_uses(self):
        tool = build_web_fetch_tool()
        assert tool["max_uses"] == DEFAULT_VERIFICATION_MAX_FETCHES
        # Default should be conservative — fetch is the deeper tool.
        assert tool["max_uses"] <= 5

    def test_max_uses_override(self):
        tool = build_web_fetch_tool(max_uses=1)
        assert tool["max_uses"] == 1

    def test_max_content_tokens_capped(self):
        """The content-token cap must be set so one fetch can't blow input."""
        tool = build_web_fetch_tool()
        assert tool["max_content_tokens"] == WEB_FETCH_MAX_CONTENT_TOKENS
        # Cap should be well below the 1M context window so the
        # verification response window is preserved.
        assert tool["max_content_tokens"] <= 100_000

    def test_blocked_domains_share_web_search_blocklist(self):
        """Fetch should not be able to read sources we'd reject from search."""
        from src.core.api_config import build_web_search_tool

        fetch_tool = build_web_fetch_tool()
        search_tool = build_web_search_tool()
        # Each blocklist may have its own list but the content must match —
        # a domain we won't search is a domain we won't fetch either.
        assert set(fetch_tool["blocked_domains"]) == set(
            search_tool["blocked_domains"]
        )


# ===========================================================================
# 2. Tool list inclusion by mode
# ===========================================================================


class TestToolListByMode:
    """``build_verification_tools_from_decision`` attaches the fetch tool
    for STANDARD_REASONING and DEEP_REASONING, omits it for the cheaper
    modes."""

    def _decision_for_mode(self, mode: VerificationMode) -> Any:
        # Synthesize a finding whose routing lands on the target mode.
        # Severity + profile combinations: HIGH default → STANDARD,
        # CRITICAL with California signals → DEEP, GRIPES → STRICT.
        if mode is VerificationMode.STANDARD_REASONING:
            finding = _finding(severity="HIGH", code_ref="NFPA 13 §10")
        elif mode is VerificationMode.DEEP_REASONING:
            finding = _finding(
                severity="CRITICAL",
                code_ref="CBC §1004",
                issue="DSA submittal requirement for California K-12",
            )
        elif mode is VerificationMode.STRICT_STRUCTURED:
            finding = _finding(severity="GRIPES", code_ref=None)
        else:
            finding = _finding(severity="GRIPES", code_ref=None)
        decision = select_routing(
            finding,
            escalated=(mode is VerificationMode.DEEP_REASONING),
            local_skip=False,
        )
        return decision

    def test_standard_reasoning_includes_web_fetch(self):
        decision = self._decision_for_mode(VerificationMode.STANDARD_REASONING)
        assert decision.mode is VerificationMode.STANDARD_REASONING
        tools = build_verification_tools_from_decision(decision)
        assert _tools_include_web_fetch(tools), (
            f"STANDARD_REASONING tools should include web_fetch, got {tools}"
        )

    def test_deep_reasoning_includes_web_fetch(self):
        decision = self._decision_for_mode(VerificationMode.DEEP_REASONING)
        assert decision.mode is VerificationMode.DEEP_REASONING
        tools = build_verification_tools_from_decision(decision)
        assert _tools_include_web_fetch(tools)

    def test_strict_structured_omits_web_fetch(self):
        decision = self._decision_for_mode(VerificationMode.STRICT_STRUCTURED)
        # Some routing branches may bump GRIPES to STANDARD if the
        # finding has signals — guard the assertion.
        if decision.mode is not VerificationMode.STRICT_STRUCTURED:
            pytest.skip(
                f"Routing produced {decision.mode}; STRICT_STRUCTURED not "
                "reachable from this finding shape."
            )
        tools = build_verification_tools_from_decision(decision)
        assert not _tools_include_web_fetch(tools), (
            f"STRICT_STRUCTURED tools should omit web_fetch, got {tools}"
        )

    def test_web_search_still_present_when_fetch_added(self):
        """Adding web_fetch must not displace web_search."""
        decision = self._decision_for_mode(VerificationMode.STANDARD_REASONING)
        tools = build_verification_tools_from_decision(decision)
        has_search = any(
            (t.get("type", "").startswith("web_search") or t.get("name") == "web_search")
            for t in tools
        )
        assert has_search

    def test_verdict_tool_stays_last(self):
        """The verdict tool must stay at the end so ``tools_with_cache``
        attaches its breakpoint to the right tool (the trailing one)."""
        decision = self._decision_for_mode(VerificationMode.STANDARD_REASONING)
        tools = build_verification_tools_from_decision(decision)
        # When the verdict tool is included (default env), it should be
        # the last entry.
        if any(t.get("name") == "submit_verification_verdict" for t in tools):
            assert tools[-1]["name"] == "submit_verification_verdict"


# ===========================================================================
# 3. No beta header on the verification request (web_fetch is GA)
# ===========================================================================


class TestNoBetaHeader:
    """Web fetch is generally available and takes NO ``anthropic-beta``
    header. Regression guard for the 400 ``invalid_request_error`` the
    retired ``web-fetch-2026-02-09`` beta triggered: it crashed every
    verification run because STANDARD_REASONING / DEEP_REASONING are the
    common modes and an unrecognized ``anthropic-beta`` value is rejected
    rather than ignored. The fix attaches the web_fetch tool (which is
    valid and current) without any beta header.
    """

    def test_standard_reasoning_attaches_tool_without_beta_header(self):
        """A fetch-eligible request carries the web_fetch tool but NO
        beta header — the tool dict alone enables the GA feature."""
        finding = _finding(severity="HIGH", code_ref="NFPA 13 §10")
        decision = select_routing(finding, escalated=False, local_skip=False)
        if decision.mode not in (
            VerificationMode.STANDARD_REASONING,
            VerificationMode.DEEP_REASONING,
        ):
            pytest.skip(f"Test requires fetch-eligible mode; got {decision.mode}")
        request = build_verification_request(
            decision,
            prompt="user message",
            system_prompt="system message",
        )
        # The tool IS present...
        assert _tools_include_web_fetch(request.params["tools"])
        # ...but no anthropic-beta header is attached.
        assert "anthropic-beta" not in request.extra_headers
        # Regression guard: ``extra_headers`` is an SDK transport kwarg,
        # not a Messages API field. Embedding it inside the per-request
        # batch ``params`` body triggers ``invalid_request_error: Extra
        # inputs are not permitted``.
        assert "extra_headers" not in request.params

    def test_strict_structured_request_omits_beta_header(self):
        """STRICT_STRUCTURED does not attach web_fetch, and carries no beta."""
        finding = _finding(severity="GRIPES", code_ref=None)
        decision = select_routing(finding, escalated=False, local_skip=False)
        if decision.mode is not VerificationMode.STRICT_STRUCTURED:
            pytest.skip(
                f"Routing produced {decision.mode}; STRICT_STRUCTURED needed."
            )
        request = build_verification_request(
            decision,
            prompt="user message",
            system_prompt="system message",
        )
        assert "anthropic-beta" not in request.extra_headers
        assert "extra_headers" not in request.params


# ===========================================================================
# 4. VerificationResult carries the fetch telemetry fields
# ===========================================================================


class TestVerificationResultFetchFields:
    def test_default_web_fetch_requests_is_zero(self):
        r = VerificationResult(verdict="UNVERIFIED")
        assert r.web_fetch_requests == 0

    def test_default_fetched_sources_is_empty_list(self):
        r = VerificationResult(verdict="UNVERIFIED")
        assert r.fetched_sources == []

    def test_fetched_sources_independence(self):
        """Default lists must not be shared across instances."""
        r1 = VerificationResult(verdict="UNVERIFIED")
        r2 = VerificationResult(verdict="UNVERIFIED")
        r1.fetched_sources.append("https://example.com/")
        assert r2.fetched_sources == []

    def test_round_trip_through_constructor(self):
        r = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://www.nfpa.org/13"],
            accepted_sources=["https://www.nfpa.org/13"],
            source_quote="Section 10.2.5.2.1...",
            web_fetch_requests=2,
            fetched_sources=["https://www.nfpa.org/13", "https://up.codes/cmc"],
        )
        assert r.web_fetch_requests == 2
        assert r.fetched_sources == ["https://www.nfpa.org/13", "https://up.codes/cmc"]


# ===========================================================================
# 5. Evidence collection helpers
# ===========================================================================


class TestCollectFetchEvidenceDetailed:
    def test_extracts_fetched_urls_from_server_tool_use_blocks(self):
        msg = _verification_response_with_fetch(
            fetched_url="https://www.nfpa.org/13"
        )
        detailed, successes, errors = _collect_fetch_evidence_detailed(msg)
        urls = [s.url for s in detailed]
        assert "https://www.nfpa.org/13" in urls
        assert successes == 1
        assert errors == 0

    def test_counts_error_only_blocks(self):
        msg = _verification_response_with_fetch(fetch_error=True)
        _, successes, errors = _collect_fetch_evidence_detailed(msg)
        assert errors == 1
        # An error result still emits the server_tool_use block with the
        # URL, so the URL list may contain it. Success counter is 0.
        assert successes == 0

    def test_no_fetch_blocks_returns_empty(self):
        """A response with only web_search and verdict tool calls should
        report zero fetches and an empty URL list."""
        msg = FakeMessage(
            content=[
                FakeServerToolUseBlock(
                    name="web_search", input={"query": "test"}
                ),
                FakeWebSearchResultBlock(
                    content=[
                        {
                            "type": "web_search_result",
                            "url": "https://example.com/",
                            "title": "Example",
                            "encrypted_content": "fake",
                        }
                    ]
                ),
                FakeToolUseBlock(
                    name="submit_verification_verdict",
                    input=sample_verification_verdict_payload(),
                ),
            ],
            stop_reason="tool_use",
        )
        detailed, successes, errors = _collect_fetch_evidence_detailed(msg)
        assert detailed == []
        assert successes == 0
        assert errors == 0

    def test_url_dedup_via_dedupe_searched_sources(self):
        """Two fetches of the same URL collapse via the shared dedupe helper."""
        from src.verification.source_grounding import dedupe_searched_sources

        msg = FakeMessage(
            content=[
                FakeServerToolUseBlock(
                    id="a",
                    name="web_fetch",
                    input={"url": "https://www.nfpa.org/13"},
                ),
                FakeWebFetchToolResultBlock(
                    tool_use_id="a",
                    content={
                        "type": "web_fetch_result",
                        "url": "https://www.nfpa.org/13",
                        "document": {"url": "https://www.nfpa.org/13"},
                    },
                ),
                FakeServerToolUseBlock(
                    id="b",
                    name="web_fetch",
                    input={"url": "https://www.nfpa.org/13"},
                ),
                FakeWebFetchToolResultBlock(
                    tool_use_id="b",
                    content={
                        "type": "web_fetch_result",
                        "url": "https://www.nfpa.org/13",
                        "document": {"url": "https://www.nfpa.org/13"},
                    },
                ),
            ],
            stop_reason="end_turn",
        )
        detailed, _, _ = _collect_fetch_evidence_detailed(msg)
        deduped = dedupe_searched_sources(detailed)
        assert len(deduped) == 1
        assert deduped[0].url == "https://www.nfpa.org/13"


class TestWebFetchCount:
    def test_reads_usage_counter(self):
        msg = _verification_response_with_fetch(n_fetches=3)
        assert _web_fetch_count(msg) == 3

    def test_missing_usage_returns_zero(self):
        msg = FakeMessage(content=[], stop_reason="end_turn", usage=FakeUsage())
        # FakeUsage has no server_tool_use attr by default.
        assert _web_fetch_count(msg) == 0

    def test_message_with_no_usage_returns_zero(self):
        class _BareMessage:
            content: list = []
            stop_reason = "end_turn"

        assert _web_fetch_count(_BareMessage()) == 0


# ===========================================================================
# 6. Source grounding accepts fetched URLs
# ===========================================================================


class TestGroundingAcceptsFetchedUrls:
    def test_fetched_url_grounds_citation_when_search_does_not(self):
        """A CONFIRMED verdict citing a URL only present in the fetched
        list (not in searched) should still pass grounding — the model
        did pull the page, so the URL is real evidence."""
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://www.nfpa.org/13"],
            source_quote="Section 10.2.5.2.1...",
        )
        searched = []  # no search snippets
        fetched = [SearchedSource(url="https://www.nfpa.org/13", title="NFPA 13")]
        out = _apply_source_grounding(result, searched=searched, fetched=fetched)
        # Citation should be accepted because the URL is in the fetched pool.
        assert "https://www.nfpa.org/13" in out.accepted_sources
        # Verdict should not have been downgraded.
        assert out.verdict == "CONFIRMED"

    def test_no_fetched_param_preserves_legacy_behavior(self):
        """Callers that don't pass ``fetched`` see the same partition logic
        they had before — the parameter is optional with default None."""
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://www.nfpa.org/13"],
            source_quote="quote",
        )
        searched = [SearchedSource(url="https://www.nfpa.org/13")]
        out = _apply_source_grounding(result, searched=searched)
        assert out.verdict == "CONFIRMED"
        assert "https://www.nfpa.org/13" in out.accepted_sources

    def test_searched_sources_field_unchanged_by_fetched(self):
        """Fetched URLs must not leak into ``searched_sources`` — that
        field is the report's snippet-source list. Fetched URLs render
        in their own panel."""
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://search.example/"],
            source_quote="q",
        )
        searched = [SearchedSource(url="https://search.example/")]
        fetched = [SearchedSource(url="https://fetched.example/")]
        out = _apply_source_grounding(result, searched=searched, fetched=fetched)
        assert "https://fetched.example/" not in out.searched_sources
        assert "https://search.example/" in out.searched_sources

    def test_pool_dedup_when_url_in_both(self):
        """A URL present in both searched and fetched lists should not
        cause issues in the validation pool."""
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://www.nfpa.org/13"],
            source_quote="q",
        )
        url = "https://www.nfpa.org/13"
        searched = [SearchedSource(url=url)]
        fetched = [SearchedSource(url=url)]
        out = _apply_source_grounding(result, searched=searched, fetched=fetched)
        assert out.verdict == "CONFIRMED"
        assert url in out.accepted_sources


# ===========================================================================
# 7. System prompt mentions web_fetch
# ===========================================================================


class TestSystemPromptIncludesWebFetch:
    def test_prompt_documents_web_fetch_tool(self):
        prompt = _get_verification_system_prompt(DEFAULT_CYCLE)
        assert "web_fetch" in prompt
        # The prompt should explain the relationship between search and
        # fetch (the model can only fetch URLs from prior search results).
        assert "search" in prompt.lower()

    def test_prompt_warns_about_fetch_cost(self):
        """The model should be told to reserve fetch for high-stakes claims."""
        prompt = _get_verification_system_prompt(DEFAULT_CYCLE)
        # Either "high-stakes" or "expensive" should appear so the model
        # doesn't fetch indiscriminately.
        lower = prompt.lower()
        assert (
            "high-stakes" in lower
            or "more expensive" in lower
            or "reserve" in lower
        )

    def test_prompt_instructs_source_quote_from_fetched_content(self):
        """When the model fetches a page it should pull source_quote from
        the fetched body, not the original snippet."""
        prompt = _get_verification_system_prompt(DEFAULT_CYCLE)
        # The exact wording is intentionally checked loosely so wording
        # tweaks don't break the test; substring matching the key idea.
        lower = prompt.lower()
        assert "source_quote" in lower
        assert "fetched" in lower

    def test_prompt_consistent_with_and_without_verdict_tool(self):
        """The web_fetch instructions should appear regardless of whether
        the verdict tool is included — fetch availability is gated by the
        request's tool list, not the verdict-tool flag."""
        with_tool = _get_verification_system_prompt(
            DEFAULT_CYCLE, include_verdict_tool=True
        )
        without_tool = _get_verification_system_prompt(
            DEFAULT_CYCLE, include_verdict_tool=False
        )
        assert "web_fetch" in with_tool
        assert "web_fetch" in without_tool


# ===========================================================================
# 8. Cache persistence round-trips fetch telemetry
# ===========================================================================


class TestCachePersistsFetchFields:
    def _grounded_result_with_fetches(self) -> VerificationResult:
        return VerificationResult(
            verdict="CONFIRMED",
            explanation="NFPA 13 §10.2.5 confirms the spacing limit.",
            sources=["https://www.nfpa.org/13"],
            accepted_sources=["https://www.nfpa.org/13"],
            grounded=True,
            model_used="claude-sonnet-4-6",
            source_quote="The maximum distance between sprinklers shall not exceed 15 ft.",
            web_search_requests=2,
            successful_source_count=1,
            web_fetch_requests=1,
            fetched_sources=["https://www.nfpa.org/13"],
        )

    def test_result_to_dict_includes_fetch_fields(self):
        result = self._grounded_result_with_fetches()
        d = _result_to_dict(result)
        assert d["web_fetch_requests"] == 1
        assert d["fetched_sources"] == ["https://www.nfpa.org/13"]

    def test_disk_round_trip(self, tmp_path: Path):
        """Save + reload preserves fetch telemetry."""
        cache = VerificationCache()
        result = self._grounded_result_with_fetches()
        finding = _finding(
            severity="HIGH",
            code_ref="NFPA 13 §10.2.5",
            issue="NFPA 13 spacing",
        )
        cache.put(finding, cycle=DEFAULT_CYCLE, result=result)
        cache_path = tmp_path / "cache.json"
        cache.save_to_disk(cache_path)

        # Reload into a fresh cache.
        reloaded = VerificationCache()
        reloaded.load_from_disk(cache_path)
        hit = reloaded.get(finding, cycle=DEFAULT_CYCLE)
        assert hit is not None
        assert hit.web_fetch_requests == 1
        assert hit.fetched_sources == ["https://www.nfpa.org/13"]
        # Cache status should flip to hit on the replay path.
        assert hit.cache_status == "hit"

    def test_legacy_cache_row_without_fetch_keys_loads_with_defaults(
        self, tmp_path: Path
    ):
        """Pre-Chunk-11 cache entries (v3 schema, no fetch keys) must
        load with empty fetch fields instead of crashing on the missing
        key. Telemetry fields are non-breaking by design."""
        from src.verification.verification_cache import _CACHE_SCHEMA_VERSION

        cache_path = tmp_path / "cache.json"
        legacy_payload = {
            "version": _CACHE_SCHEMA_VERSION,
            "saved_at": time.time(),
            "entries": {
                "legacy-key-1": {
                    "created_ts": time.time(),
                    "result": {
                        "verdict": "CONFIRMED",
                        "explanation": "legacy",
                        "sources": ["https://www.nfpa.org/13"],
                        "accepted_sources": ["https://www.nfpa.org/13"],
                        "searched_sources": ["https://www.nfpa.org/13"],
                        "cited_sources": ["https://www.nfpa.org/13"],
                        "rejected_sources": [],
                        "grounded": True,
                        "model_used": "claude-sonnet-4-6",
                        "verification_profile": "code_standard",
                        "verification_mode": "standard_reasoning",
                        "source_quote": "legacy quote",
                        "web_search_requests": 1,
                        "successful_source_count": 1,
                        "search_error_count": 0,
                        # web_fetch_requests / fetched_sources absent.
                    },
                }
            },
        }
        cache_path.write_text(json.dumps(legacy_payload), encoding="utf-8")
        reloaded = VerificationCache()
        loaded = reloaded.load_from_disk(cache_path)
        assert loaded == 1
        # The reloaded entry should default to 0 / [] for the missing
        # fields without erroring out.
        # Fetch the entry through the public get() API.
        # We don't have the original finding so we look at the internal
        # entries dict directly.
        entries = list(reloaded._entries.values())
        assert len(entries) == 1
        stored = entries[0].result
        assert stored.web_fetch_requests == 0
        assert stored.fetched_sources == []


# ===========================================================================
# 9. Resume state round-trips fetch telemetry
# ===========================================================================


class TestResumeStateRoundTrip:
    def test_serialize_includes_fetch_fields(self):
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            source_quote="quote",
            web_fetch_requests=2,
            fetched_sources=["https://a.example/", "https://b.example/"],
        )
        payload = serialize_verification_result(result)
        assert payload is not None
        assert payload["web_fetch_requests"] == 2
        assert payload["fetched_sources"] == [
            "https://a.example/",
            "https://b.example/",
        ]

    def test_deserialize_restores_fetch_fields(self):
        result = VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            source_quote="quote",
            web_fetch_requests=2,
            fetched_sources=["https://a.example/", "https://b.example/"],
        )
        round_trip = deserialize_verification_result(
            serialize_verification_result(result)
        )
        assert round_trip is not None
        assert round_trip.web_fetch_requests == 2
        assert round_trip.fetched_sources == [
            "https://a.example/",
            "https://b.example/",
        ]

    def test_legacy_payload_without_fetch_keys_loads_with_defaults(self):
        """A resume payload predating Chunk 11 (no fetch keys) must
        deserialize cleanly with 0 / []."""
        legacy = {
            "verdict": "CONFIRMED",
            "explanation": "legacy",
            "sources": ["https://www.nfpa.org/13"],
            "grounded": True,
            "source_quote": "legacy quote",
            "web_search_requests": 1,
            # No web_fetch_requests / fetched_sources keys.
        }
        result = deserialize_verification_result(legacy)
        assert result is not None
        assert result.web_fetch_requests == 0
        assert result.fetched_sources == []


# ===========================================================================
# 10. Report rendering — evidence panel surfaces fetch telemetry
# ===========================================================================


def _docx_text(path: Path) -> str:
    """Read the body + table text of a .docx file as a single string."""
    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    parts.append(paragraph.text)
    return "\n".join(parts)


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report.

    Mirrors the pattern used by ``test_chunk_4_evidence_panel`` so the
    fetch-rendering tests share one shim shape with the rest of the
    suite.
    """

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


class TestEvidencePanelRendersFetchTelemetry:
    def _review_with_finding(self, finding: Finding) -> ReviewResult:
        return ReviewResult(findings=[finding])

    def _verification(
        self,
        *,
        web_search_requests: int = 2,
        web_fetch_requests: int = 0,
        fetched_sources: list[str] | None = None,
    ) -> VerificationResult:
        return VerificationResult(
            verdict="CONFIRMED",
            grounded=True,
            sources=["https://www.nfpa.org/13"],
            accepted_sources=["https://www.nfpa.org/13"],
            source_quote=(
                "Section 10.2.5.2.1 The maximum distance between sprinklers "
                "shall not exceed 15 ft."
            ),
            explanation="NFPA 13 §10.2.5 confirms the spacing requirement.",
            model_used="claude-sonnet-4-6",
            verification_mode="standard_reasoning",
            web_search_requests=web_search_requests,
            successful_source_count=1,
            web_fetch_requests=web_fetch_requests,
            fetched_sources=list(fetched_sources or []),
        )

    def test_no_fetches_renders_plain_budget_line(self, tmp_path: Path):
        verification = self._verification(
            web_search_requests=2, web_fetch_requests=0
        )
        finding = _finding(verification=verification)
        review = self._review_with_finding(finding)
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review), out)
        text = _docx_text(out)
        # Plain budget line on the common path. HIGH severity has a
        # 7-call ceiling per ``_SEVERITY_MAX_USES``.
        assert "2 of 7 searches used" in text
        # The fetch sub-section should NOT appear.
        assert "Full-text sources consulted" not in text
        assert "Full-page fetches" not in text

    def test_with_fetches_renders_combined_budget_line_and_fetch_section(
        self, tmp_path: Path
    ):
        verification = self._verification(
            web_search_requests=2,
            web_fetch_requests=2,
            fetched_sources=[
                "https://www.nfpa.org/13",
                "https://up.codes/cmc/section-1004",
            ],
        )
        finding = _finding(verification=verification)
        review = self._review_with_finding(finding)
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review), out)
        text = _docx_text(out)
        # Combined budget line includes the fetch count.
        assert "Full-page fetches" in text
        assert "Searches: 2" in text
        # The dedicated fetch sub-section appears with both URLs.
        assert "Full-text sources consulted" in text
        assert "https://www.nfpa.org/13" in text
        assert "https://up.codes/cmc/section-1004" in text

    def test_fetch_section_omitted_when_fetched_sources_empty(
        self, tmp_path: Path
    ):
        """Even with web_fetch_requests > 0, if the URL list is empty
        (shouldn't happen in practice but the defensive default matters),
        the report should not render an empty Full-text section."""
        verification = self._verification(
            web_search_requests=2,
            web_fetch_requests=1,
            fetched_sources=[],
        )
        finding = _finding(verification=verification)
        review = self._review_with_finding(finding)
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review), out)
        text = _docx_text(out)
        # The header should NOT appear when the list is empty — empty
        # sub-sections are confusing for reviewers.
        assert "Full-text sources consulted" not in text
