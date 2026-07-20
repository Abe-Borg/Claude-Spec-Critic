"""Behavior pins for the ``datacenter_architectural`` module's routing + report surfaces.

Complements the byte-exact prompt goldens
(``test_golden_datacenter_arch_surfaces.py``) with decision pins: verification
routing (a CRITICAL building-official / accessibility finding rides the
deep-reasoning path under the arch cycle, and the *same* finding routes
differently under the default California cycle), cross-check chunk assignment
over the Divisions 03–14 map, and the domain-worded report surfaces (arch
title / phrase, generic jurisdiction-free cycle wording, and — this module's
novel surface — the pinned-editions paragraph being ABSENT, because the arch
cycle pins zero standards).

Hermetic: no API key, no network.
"""
from __future__ import annotations

import pytest

from src.modules import DATACENTER_ARCHITECTURAL
from src.review.reviewer import Finding


def _finding(issue: str, *, severity: str = "CRITICAL", filename: str = "08 11 13 - Doors.docx") -> Finding:
    return Finding(
        severity=severity,
        fileName=filename,
        section="1.01",
        issue=issue,
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference=None,
    )


# ---------------------------------------------------------------------------
# Verification routing under the arch cycle
# ---------------------------------------------------------------------------


class TestDatacenterArchRouting:
    def test_building_official_critical_routes_jurisdictional_deep(self):
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        finding = _finding(
            "The building official requires a witnessed envelope mock-up review before approval."
        )
        routed = select_routing(
            finding, local_skip=False, cycle=DATACENTER_ARCHITECTURAL.cycle
        )
        assert routed.profile is VerificationProfile.JURISDICTIONAL
        assert routed.mode is VerificationMode.DEEP_REASONING

    def test_accessibility_critical_routes_jurisdictional_deep(self):
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        finding = _finding(
            "Barrier-free accessibility clearances at the main entrance conflict with the provincial requirement."
        )
        routed = select_routing(
            finding, local_skip=False, cycle=DATACENTER_ARCHITECTURAL.cycle
        )
        assert routed.profile is VerificationProfile.JURISDICTIONAL
        assert routed.mode is VerificationMode.DEEP_REASONING

    def test_same_finding_routes_differently_under_california(self):
        # The arch jurisdictional vocabulary knows "building official"; the
        # default California vocabulary does not, so the same finding is a
        # plain constructability claim there (standard reasoning, not deep).
        from src.verification.verification_modes import VerificationMode
        from src.verification.verification_profiles import VerificationProfile
        from src.verification.verification_routing import select_routing

        finding = _finding(
            "The building official requires a witnessed envelope mock-up review before approval."
        )
        default_routed = select_routing(finding, local_skip=False)
        assert default_routed.profile is VerificationProfile.CONSTRUCTABILITY
        assert default_routed.mode is VerificationMode.STANDARD_REASONING

    def test_manufacturer_and_code_standard_keywords_classify(self):
        from src.verification.verification_profiles import (
            VerificationProfile,
            classify_finding_profile,
        )

        kw = DATACENTER_ARCHITECTURAL.profile_keywords
        assert (
            classify_finding_profile(
                _finding(
                    "Provide a Kawneer curtain-wall system per the datasheet.",
                    severity="MEDIUM",
                ),
                keywords=kw,
            )
            is VerificationProfile.MANUFACTURER
        )
        assert (
            classify_finding_profile(
                _finding(
                    "Cites an ASTM test method that does not exist.",
                    severity="HIGH",
                ),
                keywords=kw,
            )
            is VerificationProfile.CODE_STANDARD
        )


# ---------------------------------------------------------------------------
# Cross-check chunk assignment (Divisions 03–14 map)
# ---------------------------------------------------------------------------


class TestDatacenterArchChunkAssignment:
    def test_module_groups_drive_assignment(self):
        from src.cross_check.cross_checker import _assign_chunk, _chunk_label

        groups = DATACENTER_ARCHITECTURAL.cross_check_chunk_groups
        assert _assign_chunk("03 30 00 - Cast-in-Place Concrete.docx", groups) == "shell_structure"
        assert _assign_chunk("05 50 00 - Metal Fabrications.docx", groups) == "shell_structure"
        assert _assign_chunk("07 84 13 - Penetration Firestopping.docx", groups) == "envelope"
        assert _assign_chunk("08 11 13 - Hollow Metal Doors and Frames.docx", groups) == "openings"
        assert _assign_chunk("09 91 23 - Interior Painting.docx", groups) == "interiors_specialties"
        assert _assign_chunk("14 21 00 - Electric Traction Elevators.docx", groups) == "special_conveying"
        # Divisions with no arch chunk group pool into the reserved general —
        # including Division 01 and any MEP sections mixed into the package.
        assert _assign_chunk("01 33 00 - Submittal Procedures.docx", groups) == "general"
        assert _assign_chunk("23 05 00 - Common HVAC.docx", groups) == "general"
        assert _chunk_label("openings", groups) == "Division 08 — Openings"


# ---------------------------------------------------------------------------
# Report surfaces
# ---------------------------------------------------------------------------


class TestDatacenterArchReportSurfaces:
    def test_methodology_note_renders_arch_phrase_and_no_pinned_editions(self):
        from docx import Document

        from src.output.report_exporter import _write_methodology_note

        doc = Document()
        _write_methodology_note(
            doc,
            cycle_label=DATACENTER_ARCHITECTURAL.cycle.label,
            module=DATACENTER_ARCHITECTURAL,
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "relevant to hyperscale data-center architectural projects." in text
        # No jurisdiction label -> generic cycle sentence (no California).
        assert "This review used dc-arch-ibc-2024 code cycle references." in text
        # standards=() -> the pinned-editions paragraph is ABSENT (the engine
        # omits it entirely; adopted editions come from the per-run research
        # profile and render in the Jurisdiction & Client Requirements section).
        assert "per the dc-arch-ibc-2024 cycle:" not in text
        assert "California" not in text

    def test_title_block_renders_architectural_title_and_bare_cycle(self):
        from docx import Document

        from src.output.report_exporter import _write_title_block

        class _StubReview:
            model = "claude-opus-4-8"

        doc = Document()
        _write_title_block(
            doc,
            _StubReview(),
            ["08 11 13 - Doors.docx"],
            cycle_label=DATACENTER_ARCHITECTURAL.cycle.label,
            module=DATACENTER_ARCHITECTURAL,
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Spec Critic — Architectural Specification Review Report" in text
        # jurisdiction_label is empty -> bare "Code Cycle: dc-arch-ibc-2024".
        assert "Code Cycle: dc-arch-ibc-2024" in text
        assert "California" not in text

    def test_alerts_render_generic_jurisdiction_free_headings(self):
        from docx import Document

        from src.output.report_exporter import _write_alerts

        doc = Document()
        _write_alerts(
            doc,
            [],
            [],
            code_cycle_alerts=[{"filename": "08 11 13 - Doors.docx", "context": "2018 IBC"}],
            invalid_code_cycle_alerts=[{"filename": "08 11 13 - Doors.docx", "context": "2019 IBC"}],
            module=DATACENTER_ARCHITECTURAL,
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        # No jurisdiction word in the headings.
        assert "Stale Code Cycle References" in text
        assert "Invalid Code Cycle Years" in text
        assert "California" not in text
        # Published I-code years come from the shared plausible_cycle_years.
        assert "2009, 2012, 2015, 2018, 2021, 2024" in text
