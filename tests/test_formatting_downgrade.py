"""Phase 3 / Step 3.1 — Span-aware formatting-loss detection.

The legacy ``_formatting_downgrade`` helper downgraded every partial
replacement on a paragraph with 2+ distinct format runs, regardless of
whether the replacement span actually crossed those runs. The fix is
to look at the per-run format signatures recorded on
``ParagraphMapping.run_format_map`` and downgrade only when the
replacement span itself crosses distinct formatting — so an edit that
lands entirely inside one uniformly-formatted region of a richly-
formatted paragraph stays AUTO_SAFE, but an edit that crosses runs
with different formatting downgrades to AUTO_WITH_CAUTION.

These tests pin both axes:

1. The extractor populates ``ParagraphMapping.run_format_map`` for
   body paragraphs with per-run (start, end, signature) tuples in
   stripped-text coordinates.
2. ``_classify_locator_safety`` (via ``_formatting_downgrade``) uses
   the run_format_map when present and falls back to the coarse
   ``distinct_formatting_runs`` count for legacy mappings without a
   map (resume-state payloads from before Step 3.1).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from src.editing.edit_candidates import (
    SAFETY_AUTO_SAFE,
    SAFETY_AUTO_WITH_CAUTION,
    SAFETY_MANUAL_REVIEW,
)
from src.editing.edit_locator import (
    EditLocation,
    _classify_locator_safety,
)
from src.input.extractor import ParagraphMapping, extract_text_from_docx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _signature(*, bold: bool = False, italic: bool = False) -> tuple:
    """Build a format signature tuple matching extractor.py's shape."""
    return (bold, italic, False, None, None, "")


def _mapping_with_runs(
    text: str,
    runs: list[tuple[int, int, tuple]] | None,
    *,
    distinct: int | None = None,
) -> ParagraphMapping:
    """Build a ParagraphMapping with the given run_format_map.

    ``distinct`` defaults to the number of unique signatures in
    ``runs`` (or 1 when runs is None, mimicking the extractor's
    legacy single-run behavior).
    """
    if distinct is None:
        if runs:
            distinct = len({sig for _, _, sig in runs})
        else:
            distinct = 1
    return ParagraphMapping(
        body_index=0,
        element_type="paragraph",
        text=text,
        table_index=None,
        row_index=None,
        cell_index=None,
        run_count=len(runs) if runs else 1,
        distinct_formatting_runs=distinct,
        run_format_map=runs,
    )


def _location(
    mapping: ParagraphMapping,
    *,
    match_start: int,
    match_end: int,
    method: str = "exact",
    confidence: float = 1.0,
) -> EditLocation:
    return EditLocation(
        mapping=mapping,
        match_start=match_start,
        match_end=match_end,
        matched_text=mapping.text[match_start:match_end],
        match_confidence=confidence,
        match_method=method,
    )


# ---------------------------------------------------------------------------
# Span-aware downgrade behavior
# ---------------------------------------------------------------------------


class TestSpanAwareFormattingDowngrade:
    def test_span_crossing_multi_format_runs_downgrades(self):
        """``[normal, bold, normal]`` with an EDIT spanning all three → caution."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        bold = _signature(bold=True)
        runs = [
            (0, 12, normal),         # "Comply with "
            (12, 19, bold),          # "NFPA 13"
            (19, len(text), normal),  # " for sprinkler design."
        ]
        mapping = _mapping_with_runs(text, runs)
        # EDIT replaces the whole sentence; the span crosses all three
        # distinctly-formatted runs.
        location = _location(mapping, match_start=0, match_end=len(text))
        # Whole-paragraph match path still routes to MANUAL_REVIEW.
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="Comply with NFPA 13 for sprinkler density requirements.",
                cross_paragraph=False,
            )
            == SAFETY_MANUAL_REVIEW
        )

    def test_partial_span_crossing_multi_format_runs_downgrades(self):
        """Partial replacement that crosses runs with distinct signatures → caution."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        bold = _signature(bold=True)
        runs = [
            (0, 12, normal),
            (12, 19, bold),
            (19, len(text), normal),
        ]
        mapping = _mapping_with_runs(text, runs)
        # EDIT span runs from inside the normal prefix into the bold
        # token. Not a whole-paragraph match, but it crosses 2+
        # distinct format signatures, so we lose formatting if we
        # silently replace.
        location = _location(mapping, match_start=8, match_end=19)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="with NFPA 13",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_WITH_CAUTION
        )

    def test_span_inside_one_format_run_stays_auto_safe(self):
        """The plan's third acceptance criterion: variation outside the span is preserved."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        bold = _signature(bold=True)
        runs = [
            (0, 12, normal),
            (12, 19, bold),
            (19, len(text), normal),
        ]
        mapping = _mapping_with_runs(text, runs)
        # EDIT replaces only "for sprinkler design." — entirely inside
        # the trailing normal run. The bold "NFPA 13" run is outside
        # the span and survives untouched.
        match_start = text.index("for sprinkler")
        match_end = len(text)
        location = _location(mapping, match_start=match_start, match_end=match_end)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="for sprinkler density requirements.",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_SAFE
        )

    def test_span_inside_bold_run_stays_auto_safe(self):
        """Span fully inside the bold run also stays AUTO_SAFE."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        bold = _signature(bold=True)
        runs = [
            (0, 12, normal),
            (12, 19, bold),
            (19, len(text), normal),
        ]
        mapping = _mapping_with_runs(text, runs)
        # EDIT replaces "13" inside the bold token.
        match_start = text.index("13")
        match_end = match_start + 2
        location = _location(mapping, match_start=match_start, match_end=match_end)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="13R",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_SAFE
        )

    def test_uniform_paragraph_partial_replacement_stays_auto_safe(self):
        """Single-format paragraph keeps AUTO_SAFE for any partial EDIT."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        runs = [(0, len(text), normal)]
        mapping = _mapping_with_runs(text, runs)
        location = _location(mapping, match_start=0, match_end=12)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="Refer to ",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_SAFE
        )

    def test_whole_paragraph_multi_format_still_manual_review(self):
        """Whole-paragraph EDIT on a multi-format paragraph → MANUAL_REVIEW."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        bold = _signature(bold=True)
        runs = [
            (0, 12, normal),
            (12, 19, bold),
            (19, len(text), normal),
        ]
        mapping = _mapping_with_runs(text, runs)
        location = _location(mapping, match_start=0, match_end=len(text))
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="Rewrite of the sentence.",
                cross_paragraph=False,
            )
            == SAFETY_MANUAL_REVIEW
        )

    def test_add_action_not_downgraded_even_with_multi_format(self):
        """ADD actions are exempt from the formatting-loss downgrade."""
        text = "Comply with NFPA 13 for sprinkler design."
        normal = _signature()
        bold = _signature(bold=True)
        runs = [
            (0, 12, normal),
            (12, 19, bold),
            (19, len(text), normal),
        ]
        mapping = _mapping_with_runs(text, runs)
        location = _location(mapping, match_start=0, match_end=len(text))
        # ADD path stays AUTO_SAFE because the anchor's formatting is
        # not mutated; the inserted paragraph builds its own runs.
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="ADD",
                locations=[location],
                replacement_text="Additional sibling paragraph.",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_SAFE
        )


# ---------------------------------------------------------------------------
# Legacy / fallback behavior
# ---------------------------------------------------------------------------


class TestLegacyFallback:
    def test_legacy_mapping_without_run_format_map_uses_coarse_check(self):
        """``run_format_map=None`` (resume payload) falls back to coarse downgrade."""
        text = "Comply with NFPA 13 for sprinkler design."
        # Mapping reports 2+ distinct format runs but does not carry
        # the per-run map (a legacy resume-state payload from before
        # Step 3.1). The downgrade should still fire on partial
        # replacements because we cannot prove the span is uniform.
        mapping = _mapping_with_runs(text, None, distinct=2)
        location = _location(mapping, match_start=12, match_end=19)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="NFPA 13R",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_WITH_CAUTION
        )

    def test_legacy_mapping_with_single_format_skips_downgrade(self):
        text = "Comply with NFPA 13 for sprinkler design."
        mapping = _mapping_with_runs(text, None, distinct=1)
        location = _location(mapping, match_start=12, match_end=19)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="NFPA 13R",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_SAFE
        )

    def test_empty_run_format_map_falls_back_to_coarse_check(self):
        """An empty list (no runs recorded) also falls back to the coarse check."""
        text = "Comply with NFPA 13 for sprinkler design."
        # ``distinct_formatting_runs=2`` says the paragraph has multi-
        # format runs but the per-run map is empty. The legacy path
        # should still downgrade conservatively.
        mapping = _mapping_with_runs(text, [], distinct=2)
        location = _location(mapping, match_start=12, match_end=19)
        assert (
            _classify_locator_safety(
                status="matched",
                action_type="EDIT",
                locations=[location],
                replacement_text="NFPA 13R",
                cross_paragraph=False,
            )
            == SAFETY_AUTO_WITH_CAUTION
        )


# ---------------------------------------------------------------------------
# Extractor wiring
# ---------------------------------------------------------------------------


class TestExtractorPopulatesRunFormatMap:
    def test_uniform_paragraph_emits_single_signature(self, tmp_path: Path):
        source = tmp_path / "spec.docx"
        doc = Document()
        doc.add_paragraph("Comply with project standards.")
        doc.save(source)

        spec = extract_text_from_docx(source)
        body_paragraphs = [
            m for m in spec.paragraph_map or [] if m.element_type == "paragraph"
        ]
        assert body_paragraphs, "Expected at least one body paragraph"
        mapping = body_paragraphs[0]
        assert mapping.run_format_map is not None
        # One run → one entry covering the whole stripped text.
        assert len(mapping.run_format_map) == 1
        start, end, _signature = mapping.run_format_map[0]
        assert start == 0
        assert end == len(mapping.text)

    def test_multi_format_paragraph_emits_multiple_signatures(self, tmp_path: Path):
        source = tmp_path / "spec.docx"
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Comply with ")
        bold = para.add_run("NFPA 13")
        bold.bold = True
        para.add_run(" for sprinkler design.")
        doc.save(source)

        spec = extract_text_from_docx(source)
        body_paragraphs = [
            m for m in spec.paragraph_map or [] if m.element_type == "paragraph"
        ]
        mapping = body_paragraphs[0]
        assert mapping.run_format_map is not None
        signatures = {sig for _, _, sig in mapping.run_format_map}
        assert len(signatures) >= 2, "Expected distinct format signatures"
        assert mapping.distinct_formatting_runs >= 2
        # The bold run's stripped-text offsets must enclose "NFPA 13".
        bold_signature_entries = [
            entry for entry in mapping.run_format_map if entry[2][0] is True
        ]
        assert bold_signature_entries, "Expected a bold-signature entry"
        bold_start, bold_end, _ = bold_signature_entries[0]
        assert mapping.text[bold_start:bold_end] == "NFPA 13"

    def test_run_format_map_offsets_align_with_stripped_text(self, tmp_path: Path):
        source = tmp_path / "spec.docx"
        doc = Document()
        para = doc.add_paragraph()
        # Leading whitespace inside the first run is uncommon but
        # extractor.py strips para.text — make sure the offsets we
        # record still anchor to the stripped representation.
        para.add_run("   leading-whitespace ")
        emph = para.add_run("italic-token")
        emph.italic = True
        para.add_run(" trailing")
        doc.save(source)

        spec = extract_text_from_docx(source)
        body_paragraphs = [
            m for m in spec.paragraph_map or [] if m.element_type == "paragraph"
        ]
        mapping = body_paragraphs[0]
        assert mapping.text == "leading-whitespace italic-token trailing"
        # Every recorded run's stripped slice must be a substring of
        # mapping.text — no off-by-one against the leading whitespace.
        for start, end, _sig in mapping.run_format_map or []:
            assert 0 <= start <= end <= len(mapping.text)
        # And the italic run's slice should be "italic-token".
        italic_entries = [
            entry for entry in (mapping.run_format_map or []) if entry[2][1] is True
        ]
        assert italic_entries, "Expected an italic-signature entry"
        i_start, i_end, _ = italic_entries[0]
        assert mapping.text[i_start:i_end] == "italic-token"

    def test_table_cell_mapping_has_empty_or_none_run_format_map(self, tmp_path: Path):
        """Table cell mappings flatten multiple paragraphs/runs; no per-run map."""
        source = tmp_path / "spec.docx"
        doc = Document()
        doc.add_paragraph("PART 2 PRODUCTS")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Cell content"
        doc.save(source)

        spec = extract_text_from_docx(source)
        cells = [
            m for m in spec.paragraph_map or [] if m.element_type == "table_cell"
        ]
        assert cells, "Expected at least one table-cell mapping"
        # Table cell mappings don't drive the formatting-downgrade
        # path — the field is allowed to be None or empty.
        assert cells[0].run_format_map in (None, [])
