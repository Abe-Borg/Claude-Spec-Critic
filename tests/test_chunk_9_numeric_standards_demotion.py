"""Chunk 9 tests — numeric/standards CORRECTED demotion to MANUAL_EDIT.

Chunk 9 of the Trust Upgrade closes the highest-risk class of
auto-edits: CORRECTED verdicts whose proposed replacement rewrites a
numeric quantity, a standards-body reference, or a §-section reference.
A wrong specific value (5 ft → 8 ft instead of 6 ft) would propagate
silently into the spec, so the surgical mitigation routes any such
edit to MANUAL_EDIT_CANDIDATE regardless of composite confidence.

The contract has four surfaces, each covered below:

* ``numeric_or_standards_demotion_reason(finding)`` — a pure helper
  that returns the canonical rationale string when the demotion
  applies and ``None`` otherwise. The three regex patterns
  (numeric-with-unit, standards prefix, §-section reference) are
  exercised against representative replacement strings, and the
  negative paths (wrong verdict, wrong action type, missing proposal,
  stylistic-only replacement) are pinned so future schema or routing
  tweaks can't silently widen the demotion gate.
* ``classify_edit_action`` integration — a high-confidence supportive
  finding whose replacement matches the demotion patterns is routed
  to MANUAL_EDIT_CANDIDATE; a finding with the same composite but
  whose replacement is purely stylistic still clears the bar to
  AUTO_EDIT_CANDIDATE.
* Report rendering — the inline "Edit demoted:" annotation shows the
  rationale next to the status line when the override fires, and is
  absent for findings that didn't get demoted.
* Helper / classifier consistency — the same helper that
  ``classify_edit_action`` consults is reused by the report so the
  rendered reason matches the routing decision exactly.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.output.report_exporter import export_report
from src.output.report_status import (
    EditActionLabel,
    NUMERIC_STANDARDS_DEMOTION_REASON,
    classify_edit_action,
    numeric_or_standards_demotion_reason,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Numeric correction",
    confidence: float = 0.6,
    action: str = "EDIT",
    existing: str | None = "old text",
    replacement: str | None = "new text",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
    suppression_reason: str | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
        suppression_reason=suppression_reason,
    )
    f.verification = verification
    return f


def _corrected_verification(
    *,
    grounded: bool = True,
    correction: str = "The correct value is 6 ft.",
) -> VerificationResult:
    # Grounded CORRECTED requires at least one accepted source so the
    # status classifies as VERIFIED_CONTRADICTED (a supportive status).
    return VerificationResult(
        verdict="CORRECTED",
        explanation="Web search shows the spec's value differs from the standard.",
        sources=["https://www.nfpa.org/codes"] if grounded else [],
        correction=correction,
        grounded=grounded,
    )


def _confirmed_verification() -> VerificationResult:
    return VerificationResult(
        verdict="CONFIRMED",
        explanation="Verified against the standard.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
    )


def _high_confidence_proposal(
    *,
    action: str = "EDIT",
    existing: str | None = "old",
    replacement: str | None = "new",
    edit_confidence: float = 0.95,
) -> EditProposal:
    return EditProposal(
        action_type=action,
        existing_text=existing,
        replacement_text=replacement,
        edit_confidence=edit_confidence,
    )


# ---------------------------------------------------------------------------
# 1. numeric_or_standards_demotion_reason — positive matches
# ---------------------------------------------------------------------------


class TestDemotionHelperNumericUnits:
    """Replacement text containing digits + an engineering unit token.

    The regex matches a numeric quantity (with optional decimal) and a
    unit drawn from the project's K-12 mechanical/plumbing vocabulary
    (gpm, cfm, psi, ft, in, mm, cm, m, hp, kw, °F, °C, °). Each unit
    listed in the plan gets at least one positive case so the
    alternation can't silently regress.
    """

    @pytest.mark.parametrize(
        "replacement",
        [
            "Provide 5 gpm fixture flow rate.",
            "Discharge 250 cfm at the diffuser.",
            "Maintain 8 psi minimum residual.",
            "Set clearance to 6 ft on all sides.",
            "Allow 12 in of working space.",
            "Use 50 mm pipe.",
            "Allow 30 cm minimum.",
            "Maintain 1.5 m clearance.",
            "Pump shall be 5 hp minimum.",
            "Motor sized for 7.5 kw service.",
            "Heat to 180°F.",
            "Operate below -10°C.",
            "Bend at 90° angle.",
        ],
    )
    def test_numeric_unit_replacements_match(self, replacement: str):
        f = _finding(
            edit_proposal=_high_confidence_proposal(replacement=replacement),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_decimal_value_with_unit_matches(self):
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Set head to 12.5 psi.",
            ),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_no_space_between_value_and_unit_matches(self):
        # ``\s*`` allows zero whitespace; "5ft" is a common shorthand
        # in specs even if "5 ft" is preferred.
        f = _finding(
            edit_proposal=_high_confidence_proposal(replacement="Allow 5ft clearance."),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_case_insensitive_unit_matches(self):
        # Model output may upper-case unit tokens for emphasis;
        # IGNORECASE keeps the gate from silently slipping.
        f = _finding(
            edit_proposal=_high_confidence_proposal(replacement="Flow rate of 5 GPM required."),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_unit_inside_word_does_not_match(self):
        # ``ft`` appears inside "after" — the word-boundary anchor in
        # the regex prevents a spurious match. Same for "5 meters"
        # vs "5 m".
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Check 1 after the inspection cycle.",
            ),
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_meters_word_does_not_trigger_bare_m(self):
        # "5 meters" must NOT match the bare ``m`` alternative because
        # the word boundary stops the match at the letter after ``m``.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Provide 5 meters of pipe.",
            ),
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None


class TestDemotionHelperStandardsPrefix:
    """Replacement text containing a standards-body prefix + a number.

    Each acronym listed in the plan gets a positive case so a typo in
    the alternation would be caught immediately.
    """

    @pytest.mark.parametrize(
        "replacement",
        [
            "Comply with NFPA 13 latest edition.",
            "Reference ASCE 7-22 for wind loads.",
            "Per ASHRAE 90.1 envelope requirements.",
            "Per CBC 2025 occupancy classifications.",
            "Per CMC 2025 ventilation tables.",
            "Per CPC 2025 venting tables.",
            "Per CEC 2025 conductor sizing.",
            "Per CALGreen 2025 water-use limits.",
            "Per IAPMO 1 plumbing testing.",
            "Per ASTM A53 pipe specification.",
            "Per ANSI 117 accessibility.",
            "Per UL 300 fire-suppression listing.",
            "Per API 650 tank standard.",
            "Per AWWA C151 ductile iron pipe.",
            "Per AISC 360 steel design.",
            "Per ICC 500 storm shelters.",
        ],
    )
    def test_each_standards_prefix_matches(self, replacement: str):
        f = _finding(
            edit_proposal=_high_confidence_proposal(replacement=replacement),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_lowercase_standards_prefix_matches(self):
        # IGNORECASE — the model may emit "nfpa" lowercase in prose.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="comply with nfpa 72 fire alarm requirements.",
            ),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_calgreen_mixed_case_matches(self):
        # "CALGreen" is mixed case in the plan's literal regex;
        # IGNORECASE lets every case variation match.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Per calgreen 5.303 outdoor water use.",
            ),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_prefix_without_number_does_not_match(self):
        # "NFPA references" alone is not a specific standard number;
        # the regex requires ``\s+\d+`` after the prefix so no match.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Update the NFPA references throughout.",
            ),
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_prefix_embedded_in_word_does_not_match(self):
        # "PRENFPA 13" or similar — the leading ``\b`` boundary
        # prevents false matches when the prefix is part of another
        # token (e.g. a hyphenated identifier).
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="See PRENFPA 13 draft document.",
            ),
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_standards_revision_change_matches(self):
        # NFPA 13 → NFPA 13R is the canonical "wrong specific value"
        # scenario the plan calls out. Both forms contain a standards
        # prefix + number, so the proposal triggers the demotion.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                existing="Per NFPA 13.",
                replacement="Per NFPA 13R.",
            ),
            verification=_corrected_verification(
                correction="Should reference NFPA 13R."
            ),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )


class TestDemotionHelperSectionReference:
    """Replacement text containing a §-prefixed multi-part section reference.

    The regex requires at least one dot-separated continuation
    (``\\d+(\\.\\d+)+``) so a bare ``§ 1234`` would not match — the
    intent is to catch the deeply-nested code references where a wrong
    sub-section would be hard to spot inline.
    """

    @pytest.mark.parametrize(
        "replacement",
        [
            "Per §1234.5 of the code.",
            "Comply with § 906.1.1.",
            "See §202.3.4.5 for definitions.",
        ],
    )
    def test_section_reference_replacements_match(self, replacement: str):
        f = _finding(
            edit_proposal=_high_confidence_proposal(replacement=replacement),
            verification=_corrected_verification(),
        )
        assert (
            numeric_or_standards_demotion_reason(f)
            == NUMERIC_STANDARDS_DEMOTION_REASON
        )

    def test_bare_section_number_without_subsection_does_not_match(self):
        # ``§ 1234`` alone — no dot — does not match the plan's regex.
        # Reviewers typically catch single-level section references at
        # a glance; the gate targets deeply nested forms.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Per § 1234 of the code.",
            ),
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None


# ---------------------------------------------------------------------------
# 2. numeric_or_standards_demotion_reason — negative paths
# ---------------------------------------------------------------------------


class TestDemotionHelperNegativePaths:
    """Guards that keep the demotion gate narrow.

    The helper must return ``None`` in every case where the plan's
    "Non-goals" section explicitly excludes demotion. Otherwise the
    gate would widen and stylistic / coordination edits would
    unnecessarily land on manual review.
    """

    def test_confirmed_verdict_is_not_demoted(self):
        # Plan "Non-goals": CONFIRMED edits keep the existing
        # confidence-based routing. The model said the text is correct
        # (or proposed a non-numeric stylistic change), so there's no
        # asymmetric numeric-value risk.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Maintain 6 ft clearance.",
            ),
            verification=_confirmed_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_unverified_verdict_is_not_demoted(self):
        # UNVERIFIED already routes to manual via the supportive-status
        # filter; the demotion helper just doesn't fire here.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Provide 8 psi residual pressure.",
            ),
            verification=VerificationResult(verdict="UNVERIFIED", grounded=False),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_disputed_verdict_is_not_demoted(self):
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Per NFPA 13 latest edition.",
            ),
            verification=VerificationResult(verdict="DISPUTED", grounded=False),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_no_verification_is_not_demoted(self):
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                replacement="Per ASHRAE 62.1 ventilation.",
            ),
            verification=None,
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_no_edit_proposal_is_not_demoted(self):
        # REPORT_ONLY findings have no proposal to demote; the helper
        # returns None so ``classify_edit_action`` keeps the existing
        # REPORT_ONLY routing.
        f = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_add_action_is_not_demoted(self):
        # Plan "Non-goals": ADD and DELETE actions are lower-frequency
        # and explicitly out of scope for this chunk.
        proposal = EditProposal(
            action_type="ADD",
            anchor_text="Following the existing requirement,",
            insert_position="after",
            replacement_text="Maintain 5 ft clearance.",
            edit_confidence=0.95,
        )
        f = _finding(
            action="ADD",
            existing=None,
            replacement="Maintain 5 ft clearance.",
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_delete_action_is_not_demoted(self):
        # DELETE actions have no replacement text to match against;
        # the empty/whitespace check returns None before regex even runs.
        proposal = EditProposal(
            action_type="DELETE",
            existing_text="Per NFPA 13.",
            replacement_text=None,
            edit_confidence=0.95,
        )
        f = _finding(
            action="DELETE",
            existing="Per NFPA 13.",
            replacement=None,
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_stylistic_replacement_without_numbers_is_not_demoted(self):
        # The canonical "should auto-edit" case: a CORRECTED edit that
        # rewrites wording but doesn't touch a number or standards
        # reference. The plan's success criterion #3.
        f = _finding(
            edit_proposal=_high_confidence_proposal(
                existing="shall be installed",
                replacement="must be installed",
            ),
            verification=_corrected_verification(
                correction="Per the standard's modal-verb style guide."
            ),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_empty_replacement_text_is_not_demoted(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="",
            edit_confidence=0.95,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None

    def test_whitespace_only_replacement_is_not_demoted(self):
        proposal = EditProposal(
            action_type="EDIT",
            existing_text="old",
            replacement_text="   \n  ",
            edit_confidence=0.95,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert numeric_or_standards_demotion_reason(f) is None


# ---------------------------------------------------------------------------
# 3. classify_edit_action integration
# ---------------------------------------------------------------------------


class TestClassifyEditActionDemotion:
    """The override fires inside ``classify_edit_action``.

    The plan's success criteria are the canonical scenarios:

    * ``"5 ft" → "8 ft"`` (numeric quantity) → MANUAL_EDIT_CANDIDATE.
    * ``"NFPA 13" → "NFPA 13R"`` (standards reference) →
      MANUAL_EDIT_CANDIDATE.
    * ``"shall be installed" → "must be installed"`` (stylistic) →
      AUTO_EDIT_CANDIDATE.
    """

    def test_numeric_correction_routes_to_manual_despite_high_composite(self):
        proposal = _high_confidence_proposal(
            existing="Maintain 5 ft clearance.",
            replacement="Maintain 8 ft clearance.",
            edit_confidence=1.0,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(
                correction="Should be 8 ft per the standard."
            ),
        )
        # Sanity: status is supportive and composite would otherwise
        # clear the default 0.7 floor — proves the demotion fired
        # rather than the existing low-composite branch.
        from src.output.report_status import (
            ReportStatus,
            classify_status,
            composite_edit_confidence,
            auto_edit_confidence_floor,
        )
        assert classify_status(f) is ReportStatus.VERIFIED_CONTRADICTED
        assert composite_edit_confidence(f) >= auto_edit_confidence_floor()

        assert classify_edit_action(f) is EditActionLabel.MANUAL_EDIT_CANDIDATE

    def test_standards_revision_correction_routes_to_manual(self):
        # Plan success criterion #2: "NFPA 13" → "NFPA 13R" must land
        # on MANUAL_EDIT_CANDIDATE. The leading prefix + number match
        # triggers the demotion before composite is even evaluated.
        proposal = _high_confidence_proposal(
            existing="Per NFPA 13.",
            replacement="Per NFPA 13R.",
            edit_confidence=1.0,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(
                correction="Spec is residential; NFPA 13R applies."
            ),
        )
        assert classify_edit_action(f) is EditActionLabel.MANUAL_EDIT_CANDIDATE

    def test_section_reference_correction_routes_to_manual(self):
        proposal = _high_confidence_proposal(
            existing="Comply with § 906.1.1.",
            replacement="Comply with § 906.1.2.",
            edit_confidence=1.0,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert classify_edit_action(f) is EditActionLabel.MANUAL_EDIT_CANDIDATE

    def test_stylistic_correction_still_auto_edits(self):
        # Plan success criterion #3: a non-numeric stylistic CORRECTED
        # edit must continue to auto-apply when composite clears the
        # floor. The override gate must not widen accidentally.
        proposal = _high_confidence_proposal(
            existing="shall be installed",
            replacement="must be installed",
            edit_confidence=0.95,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(
                correction="Per the spec's modal-verb style.",
            ),
        )
        assert classify_edit_action(f) is EditActionLabel.AUTO_EDIT_CANDIDATE

    def test_confirmed_numeric_edit_still_uses_composite_gate(self):
        # CONFIRMED edits with numeric replacement text must NOT be
        # demoted — the plan's "Non-goals" section is explicit. A
        # confirmed numeric value at high composite still auto-edits.
        # (CONFIRMED + matching existing text is unusual but the rule
        # is symmetric: the demotion is verdict-conditional.)
        proposal = _high_confidence_proposal(
            existing="Maintain 6 ft clearance.",
            replacement="Maintain 6 ft clearance.",
            edit_confidence=0.95,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_confirmed_verification(),
        )
        assert classify_edit_action(f) is EditActionLabel.AUTO_EDIT_CANDIDATE

    def test_disputed_numeric_proposal_already_manual_via_status(self):
        # DISPUTED would already route to manual via the supportive
        # status filter — the demotion check is redundant but safe.
        proposal = _high_confidence_proposal(
            existing="Per NFPA 13.",
            replacement="Per NFPA 13R.",
            edit_confidence=0.95,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=VerificationResult(verdict="DISPUTED", grounded=False),
        )
        assert classify_edit_action(f) is EditActionLabel.MANUAL_EDIT_CANDIDATE

    def test_suppression_still_beats_numeric_demotion(self):
        # Suppression is the highest-priority branch; a suppressed
        # finding stays SUPPRESSED even when the numeric demotion would
        # otherwise fire. The label hierarchy must be preserved.
        proposal = _high_confidence_proposal(
            existing="Maintain 5 ft clearance.",
            replacement="Maintain 8 ft clearance.",
            edit_confidence=0.95,
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(),
            suppression_reason="dropped by upstream-disputed filter",
        )
        assert classify_edit_action(f) is EditActionLabel.SUPPRESSED


# ---------------------------------------------------------------------------
# 4. Report rendering — inline annotation
# ---------------------------------------------------------------------------


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report (mirrors Chunk N)."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cross_check_result=None,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or [review_result.findings[0].fileName]
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestReportInlineAnnotation:
    """The exporter surfaces the demotion reason inline.

    Reviewers should not have to expand the Edit Target Evidence panel
    to see why an auto-edit was demoted — the composite line would
    otherwise show a passing (≥ threshold) number and create the wrong
    expectation. A small italic "Edit demoted:" note immediately under
    the status line communicates the override at a glance.
    """

    def test_export_renders_demotion_annotation(self, tmp_path: Path):
        demoted = _finding(
            severity="HIGH",
            file="Section_22_1000.docx",
            section="2.1",
            issue="Clearance value disagrees with NFPA",
            edit_proposal=_high_confidence_proposal(
                existing="Maintain 5 ft clearance.",
                replacement="Maintain 8 ft clearance.",
                edit_confidence=0.95,
            ),
            verification=_corrected_verification(
                correction="Should be 8 ft per NFPA 13.",
            ),
        )
        review = ReviewResult(findings=[demoted])
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review), out)
        doc = Document(str(out))
        text = _all_text_from(doc)

        # Inline annotation prefix + canonical rationale must appear.
        assert "Edit demoted:" in text
        assert NUMERIC_STANDARDS_DEMOTION_REASON in text
        # And the finding must show the manual-edit label (not auto).
        assert "Manual edit candidate" in text

    def test_export_omits_annotation_when_not_demoted(self, tmp_path: Path):
        # A stylistic CORRECTED edit (no numeric/standards content)
        # still auto-edits; the annotation must NOT appear so the
        # reviewer is not misled into thinking an override fired.
        stylistic = _finding(
            severity="HIGH",
            file="Section_22_1000.docx",
            section="2.1",
            issue="Modal verb style",
            edit_proposal=_high_confidence_proposal(
                existing="shall be installed",
                replacement="must be installed",
                edit_confidence=0.95,
            ),
            verification=_corrected_verification(
                correction="Per the spec's modal-verb style.",
            ),
        )
        review = ReviewResult(findings=[stylistic])
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review), out)
        doc = Document(str(out))
        text = _all_text_from(doc)

        assert "Edit demoted:" not in text
        assert "Auto-edit candidate" in text

    def test_export_omits_annotation_for_confirmed_findings(self, tmp_path: Path):
        # Plan "Non-goals" — CONFIRMED findings never demote even when
        # the replacement text contains numbers. The annotation must
        # be absent for these.
        confirmed = _finding(
            severity="HIGH",
            file="Section_22_1000.docx",
            section="2.1",
            issue="Confirmed numeric requirement",
            edit_proposal=_high_confidence_proposal(
                existing="Maintain 6 ft clearance.",
                replacement="Maintain 6 ft clearance.",
                edit_confidence=0.95,
            ),
            verification=_confirmed_verification(),
        )
        review = ReviewResult(findings=[confirmed])
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=review), out)
        doc = Document(str(out))
        text = _all_text_from(doc)

        assert "Edit demoted:" not in text


# ---------------------------------------------------------------------------
# 5. Helper / classifier consistency
# ---------------------------------------------------------------------------


class TestHelperClassifierConsistency:
    """The renderer must use the same helper the classifier consults.

    If the two paths diverge (e.g. the renderer recomputes the regex
    independently), a future regex tweak in one place but not the
    other would silently desync the displayed reason from the actual
    routing decision. Pinning the contract here catches that.
    """

    def test_demoted_findings_carry_a_reason_string(self):
        # Every finding routed to MANUAL_EDIT_CANDIDATE by the numeric
        # override must produce a non-empty reason from the same
        # helper — i.e. the routing decision and the renderable
        # rationale are always consistent.
        proposal = _high_confidence_proposal(
            existing="Maintain 5 ft clearance.",
            replacement="Maintain 8 ft clearance.",
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert classify_edit_action(f) is EditActionLabel.MANUAL_EDIT_CANDIDATE
        reason = numeric_or_standards_demotion_reason(f)
        assert reason
        assert reason == NUMERIC_STANDARDS_DEMOTION_REASON

    def test_auto_edit_findings_have_no_demotion_reason(self):
        proposal = _high_confidence_proposal(
            existing="shall be installed",
            replacement="must be installed",
        )
        f = _finding(
            edit_proposal=proposal,
            verification=_corrected_verification(),
        )
        assert classify_edit_action(f) is EditActionLabel.AUTO_EDIT_CANDIDATE
        assert numeric_or_standards_demotion_reason(f) is None
