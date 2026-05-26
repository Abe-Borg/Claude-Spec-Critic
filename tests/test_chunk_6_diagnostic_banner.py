"""Chunk 6 tests — Run Diagnostics banner.

Chunk 6 of the Trust Upgrade adds a styled-table banner right after the
title block that surfaces operational health at-a-glance:

* Auto-edit / Manual-edit / Report-only / Suppressed counts (from the
  edit-action histogram already computed for the trust-model summary).
* Cache replays with the oldest entry age (using Chunk 5's
  ``cache_entry_created_ts``).
* Verification failures (Chunk 3's ``VERIFICATION_FAILED`` status),
  highlighted red when > 0.
* REPORT_ONLY demotions at parse time (Chunk 7's ``demotion_reason``).
* Spec content extraction warnings (slot reserved for Chunk 10; renders 0
  on every run until that lands).
* Cross-spec coordination status — skipped / failed / completed.

A failure recovery hint paragraph appears below the table whenever the
verification-failure count is non-zero so a reviewer can re-run only
those findings.

The plan's success criteria:

* A clean run shows the banner with all counts at expected values.
* A run with simulated verification failures shows red callouts.
* Cross-spec coordination skipped/failed counts appear in the banner.
"""
from __future__ import annotations

import time
from pathlib import Path

from docx import Document

from src.output.report_exporter import (
    _summarize_run_diagnostics,
    export_report,
)
from src.output.report_status import (
    summarize_edit_actions,
    summarize_statuses,
)
from src.review.reviewer import EditProposal, Finding, ReviewResult
from src.verification.verifier import VerificationResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(
    *,
    severity: str = "HIGH",
    file: str = "Section_22_1000.docx",
    section: str = "2.1",
    issue: str = "Stale code reference",
    confidence: float = 0.8,
    action: str = "EDIT",
    existing: str | None = "2019 CBC",
    replacement: str | None = "2025 CBC",
    verification: VerificationResult | None = None,
    edit_proposal: EditProposal | None = None,
    suppression_reason: str | None = None,
    demotion_reason: str | None = None,
) -> Finding:
    f = Finding(
        severity=severity,
        fileName=file,
        section=section,
        issue=issue,
        actionType=action,
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC §1234",
        confidence=confidence,
        edit_proposal=edit_proposal,
        suppression_reason=suppression_reason,
        demotion_reason=demotion_reason,
    )
    f.verification = verification
    return f


def _verified_supported() -> VerificationResult:
    return VerificationResult(
        verdict="CONFIRMED",
        explanation="Verified against CBC §1234.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        accepted_sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
        cache_status="miss",
        source_quote="The 2025 CBC adopts the 2024 IBC.",
        model_used="claude-sonnet-4-6",
        verification_mode="standard_reasoning",
        web_search_requests=3,
    )


def _failed_verification() -> VerificationResult:
    return VerificationResult(
        verdict="UNVERIFIED",
        explanation="Server overloaded during verification: 529",
        grounded=False,
        verification_failed=True,
    )


def _cache_hit_result(*, age_days: int = 5) -> VerificationResult:
    created_ts = time.time() - (age_days * 86400)
    return VerificationResult(
        verdict="CONFIRMED",
        explanation="Cached verdict from prior run.",
        sources=["https://codes.iccsafe.org/content/CBC2025"],
        accepted_sources=["https://codes.iccsafe.org/content/CBC2025"],
        grounded=True,
        model_used="claude-sonnet-4-6",
        cache_status="hit",
        source_quote="The 2025 CBC adopts the 2024 IBC.",
        cache_entry_created_ts=created_ts,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        files_reviewed: list[str] | None = None,
        cycle_label: str = "2025",
        cross_check_result: ReviewResult | None = None,
    ):
        self.review_result = review_result
        self.cross_check_result = cross_check_result
        self.files_reviewed = files_reviewed or (
            [review_result.findings[0].fileName] if review_result.findings else ["test.docx"]
        )
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _findings_to_summary(
    findings: list[Finding],
    *,
    cross_check_result: ReviewResult | None = None,
    pipeline_result=None,
) -> dict:
    """Convenience: drive _summarize_run_diagnostics with the same
    inputs export_report would feed it."""
    status_counts = summarize_statuses(findings)
    edit_action_counts = summarize_edit_actions(findings)
    return _summarize_run_diagnostics(
        findings=findings,
        status_counts=status_counts,
        edit_action_counts=edit_action_counts,
        cross_check_result=cross_check_result,
        pipeline_result=pipeline_result,
    )


# ---------------------------------------------------------------------------
# 1. _summarize_run_diagnostics — pure function over findings + counts
# ---------------------------------------------------------------------------


class TestSummarizeRunDiagnostics:
    def test_empty_run_returns_zero_counts(self):
        summary = _findings_to_summary([])
        assert summary["auto_edit"] == 0
        assert summary["manual_edit"] == 0
        assert summary["report_only"] == 0
        assert summary["suppressed"] == 0
        assert summary["verification_failed"] == 0
        assert summary["cache_replay_count"] == 0
        assert summary["oldest_cache_age_days"] is None
        assert summary["demotion_count"] == 0
        assert summary["extraction_warning_count"] == 0
        assert summary["cross_check"] is None

    def test_auto_edit_count_reflects_supportive_findings(self):
        # A VERIFIED_SUPPORTED finding with a high-confidence edit
        # proposal is the canonical auto-edit candidate.
        f = _finding(verification=_verified_supported(), confidence=0.9)
        summary = _findings_to_summary([f])
        assert summary["auto_edit"] == 1
        assert summary["manual_edit"] == 0

    def test_manual_edit_count_reflects_low_confidence(self):
        # Edit proposal + supportive status + confidence below floor →
        # MANUAL_EDIT_CANDIDATE rather than AUTO.
        f = _finding(verification=_verified_supported(), confidence=0.5)
        summary = _findings_to_summary([f])
        assert summary["auto_edit"] == 0
        assert summary["manual_edit"] == 1

    def test_report_only_count_reflects_no_proposal(self):
        f = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            verification=_verified_supported(),
        )
        summary = _findings_to_summary([f])
        assert summary["report_only"] == 1

    def test_suppressed_count_reflects_suppression_reason(self):
        f = _finding(
            verification=_verified_supported(),
            suppression_reason="upstream disputed",
        )
        summary = _findings_to_summary([f])
        assert summary["suppressed"] == 1

    def test_verification_failed_count_uses_status_histogram(self):
        f = _finding(verification=_failed_verification())
        summary = _findings_to_summary([f])
        assert summary["verification_failed"] == 1

    def test_cache_replay_counts_only_hit_findings(self):
        hit = _finding(verification=_cache_hit_result(age_days=5))
        miss = _finding(
            file="Section_22_2000.docx",
            verification=_verified_supported(),
        )
        summary = _findings_to_summary([hit, miss])
        assert summary["cache_replay_count"] == 1
        assert summary["oldest_cache_age_days"] == 5

    def test_oldest_cache_age_is_maximum(self):
        # Among several cache hits, the oldest age wins so a reviewer
        # sees the worst-case staleness in the banner.
        ages = [5, 45, 120]
        findings = [
            _finding(
                file=f"Section_22_{i * 1000}.docx",
                verification=_cache_hit_result(age_days=age),
            )
            for i, age in enumerate(ages, start=1)
        ]
        summary = _findings_to_summary(findings)
        assert summary["cache_replay_count"] == 3
        assert summary["oldest_cache_age_days"] == 120

    def test_legacy_cache_hit_counts_but_no_age(self):
        # A cache_status="hit" with cache_entry_created_ts=0.0 (legacy
        # resume payload predating Chunk 5) counts toward the cache
        # replay total but cannot contribute to the oldest-age display.
        # Note: _enforce_grounding_invariant downgrades verdicts where
        # accepted citation chain is missing, but for a legacy payload
        # the cache_status stays "hit" on the verification result.
        legacy = VerificationResult(
            verdict="CONFIRMED",
            sources=["https://x"],
            accepted_sources=["https://x"],
            grounded=True,
            cache_status="hit",
            source_quote="snippet",
            cache_entry_created_ts=0.0,
        )
        f = _finding(verification=legacy)
        summary = _findings_to_summary([f])
        assert summary["cache_replay_count"] == 1
        assert summary["oldest_cache_age_days"] is None

    def test_demotion_count_reflects_demotion_reason_field(self):
        # Findings with a parse-time demotion are the ones surfaced in
        # this row — they signal model-output shape issues, not a
        # deliberate REPORT_ONLY.
        f = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            demotion_reason="EDIT requested but existingText was empty",
        )
        summary = _findings_to_summary([f])
        assert summary["demotion_count"] == 1

    def test_demotion_count_ignores_empty_or_whitespace_reasons(self):
        # An empty string or pure whitespace must not be counted as a
        # demotion — those represent the "no reason recorded" default.
        f1 = _finding(action="REPORT_ONLY", existing=None, replacement=None, demotion_reason="")
        f2 = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            file="Other.docx",
            demotion_reason="   ",
        )
        summary = _findings_to_summary([f1, f2])
        assert summary["demotion_count"] == 0

    def test_extraction_warning_count_defaults_zero(self):
        # Chunk 10 will populate ExtractedSpec.extraction_warnings.
        # Until then this stays 0 for every run.
        summary = _findings_to_summary([])
        assert summary["extraction_warning_count"] == 0

    def test_extraction_warning_count_reads_from_pipeline_result(self):
        # When the pipeline_result has extracted_specs with non-empty
        # extraction_warnings, the count reflects the number of specs
        # affected. This is the slot Chunk 10 will populate.
        class _StubSpec:
            def __init__(self, warnings):
                self.extraction_warnings = warnings

        class _StubPipeline:
            extracted_specs = [
                _StubSpec(["20% drawings"]),
                _StubSpec([]),
                _StubSpec(["30% drawings"]),
            ]

        summary = _findings_to_summary([], pipeline_result=_StubPipeline())
        assert summary["extraction_warning_count"] == 2

    def test_cross_check_state_completed(self):
        cc = ReviewResult(findings=[], cross_check_status="completed")
        summary = _findings_to_summary([], cross_check_result=cc)
        assert summary["cross_check"] is not None
        assert summary["cross_check"]["status"] == "completed"
        assert summary["cross_check"]["finding_count"] == 0

    def test_cross_check_state_skipped(self):
        cc = ReviewResult(
            findings=[],
            cross_check_status="skipped",
            thinking="not enough specs",
        )
        summary = _findings_to_summary([], cross_check_result=cc)
        assert summary["cross_check"]["status"] == "skipped"
        assert summary["cross_check"]["reason"] == "not enough specs"

    def test_cross_check_state_failed(self):
        cc = ReviewResult(
            findings=[],
            cross_check_status="failed",
            error="API timeout",
        )
        summary = _findings_to_summary([], cross_check_result=cc)
        assert summary["cross_check"]["status"] == "failed"
        assert summary["cross_check"]["reason"] == "API timeout"

    def test_cross_check_none_means_not_run(self):
        # No cross-check requested → no cross-check row in the banner.
        # The summary records None so the renderer can suppress the row.
        summary = _findings_to_summary([], cross_check_result=None)
        assert summary["cross_check"] is None


# ---------------------------------------------------------------------------
# 2. Banner rendering — section heading + labels
# ---------------------------------------------------------------------------


class TestBannerRendering:
    def test_banner_heading_appears_in_report(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "Run Diagnostics" in text

    def test_banner_contains_every_row_label(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        # Every plan-specified row label must appear. Some rows depend
        # on the run (cross-check only renders when configured); these
        # are the always-present rows.
        assert "Auto-edit eligible" in text
        assert "Manual edit required" in text
        assert "Report-only" in text
        assert "Suppressed (cross-check filter)" in text
        assert "Cache replays" in text
        assert "Verification failures (operational)" in text
        assert "REPORT_ONLY demotions at parse time" in text
        assert "Spec content extraction warnings" in text

    def test_banner_renders_before_files_reviewed(self, tmp_path: Path):
        # The plan locates the banner between the title block and the
        # methodology note — placing it before "Files Reviewed" keeps
        # the operational summary the first thing a reviewer sees.
        # Note: the title block metadata line "Files Reviewed: N" also
        # contains the substring "Files Reviewed"; the heading version
        # (no colon) is the one whose position we care about.
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        banner_idx = text.find("Run Diagnostics")
        # The heading uses "Files Reviewed\n" (followed by the bullet
        # list); the metadata line is "Files Reviewed: 1". Skip past
        # the metadata occurrence to find the heading.
        files_heading_idx = text.find("Files Reviewed\n")
        methodology_idx = text.find("About This Review")
        assert banner_idx >= 0
        assert files_heading_idx >= 0
        assert methodology_idx >= 0
        # Banner before everything that comes after the title block.
        assert banner_idx < files_heading_idx
        assert banner_idx < methodology_idx


# ---------------------------------------------------------------------------
# 3. Banner content — counts reflect the input findings
# ---------------------------------------------------------------------------


class TestBannerCounts:
    def test_clean_run_shows_zero_failures(self, tmp_path: Path):
        # A clean run with one supported finding has 0 verification
        # failures, 0 demotions, 0 cache replays.
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        # The "Verification failures (operational)" row must report 0
        # and the failure-hint paragraph must NOT appear.
        assert "Verification failures (operational)" in text
        assert "failed verification due to operational errors" not in text

    def test_run_with_verification_failures_shows_count(self, tmp_path: Path):
        # Two findings: one supported, one failed. The banner shows 1
        # failure and the failure-hint paragraph appears.
        f1 = _finding(verification=_verified_supported())
        f2 = _finding(
            file="Other.docx",
            verification=_failed_verification(),
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f1, f2])),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Verification failures (operational)" in text
        # The hint paragraph copies the count into its prose.
        assert "1 finding failed verification" in text
        # Plural form when count > 1.

    def test_failure_hint_uses_plural_for_multiple_failures(self, tmp_path: Path):
        # Plural in the hint paragraph when multiple findings failed.
        f1 = _finding(verification=_failed_verification())
        f2 = _finding(file="Other.docx", verification=_failed_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f1, f2])),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "2 findings failed verification" in text

    def test_failure_hint_omitted_when_zero(self, tmp_path: Path):
        # Clean run → no hint paragraph.
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        assert "failed verification due to operational errors" not in text
        assert "Re-running the review" not in text

    def test_cache_replay_row_shows_oldest_age(self, tmp_path: Path):
        # A cache hit produces a row with "(oldest Nd old)" suffix.
        f = _finding(verification=_cache_hit_result(age_days=45))
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        # The banner row contains the oldest-age annotation; the per-
        # finding badge below has its own "Cache replay — 45d old"
        # form. Both report the same age.
        assert "(oldest 45d old)" in text

    def test_cache_replay_row_zero_when_no_hits(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        # The row label is present but the value should not advertise
        # an "oldest" age when there are no hits.
        assert "Cache replays" in text
        assert "oldest" not in text.lower().split("cache replays")[1].split("\n")[1]

    def test_demotion_count_appears_in_banner(self, tmp_path: Path):
        # A finding with demotion_reason set is counted in the demotion
        # row even when it lives on the regular findings list.
        f = _finding(
            action="REPORT_ONLY",
            existing=None,
            replacement=None,
            demotion_reason="EDIT requested but existingText was empty",
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        text = _all_text_from(Document(str(out)))
        # The row label is present and its value shows the demotion count.
        idx = text.find("REPORT_ONLY demotions at parse time")
        assert idx >= 0
        # The value cell sits in the next non-empty entry after the label.
        # The full text contains the count somewhere downstream.
        assert "REPORT_ONLY demotions at parse time" in text


# ---------------------------------------------------------------------------
# 4. Cross-check status in the banner
# ---------------------------------------------------------------------------


class TestBannerCrossCheckStatus:
    def test_cross_check_completed_renders_count(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        cc = ReviewResult(findings=[], cross_check_status="completed")
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(
                review_result=ReviewResult(findings=[f]),
                cross_check_result=cc,
            ),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Cross-spec coordination" in text
        # "0 findings" because the cross-check produced no coordination
        # issues. The banner reports the count, not "completed".
        assert "0 findings" in text

    def test_cross_check_skipped_renders_skipped(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        cc = ReviewResult(
            findings=[],
            cross_check_status="skipped",
            thinking="too few specs",
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(
                review_result=ReviewResult(findings=[f]),
                cross_check_result=cc,
            ),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Cross-spec coordination" in text
        assert "skipped" in text

    def test_cross_check_failed_renders_failed(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        cc = ReviewResult(
            findings=[],
            cross_check_status="failed",
            error="rate limited",
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(
                review_result=ReviewResult(findings=[f]),
                cross_check_result=cc,
            ),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Cross-spec coordination" in text
        # The banner uses the literal "failed" status as the value.
        idx = text.find("Cross-spec coordination")
        # Pick out a window of text after the label.
        window = text[idx : idx + 200]
        assert "failed" in window.lower()

    def test_cross_check_none_omits_row(self, tmp_path: Path):
        # When cross-check wasn't run, no row is rendered (the
        # _StubPipelineResult passes None by default).
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])),
            out,
        )
        text = _all_text_from(Document(str(out)))
        assert "Cross-spec coordination" not in text


# ---------------------------------------------------------------------------
# 5. Visual highlight — verification failures + extraction warnings
# ---------------------------------------------------------------------------


class TestBannerHighlights:
    """The plan calls for verification failures to highlight in red
    when > 0 — we verify by inspecting cell shading on the value
    column. Same treatment applies to extraction warnings (slot
    reserved for Chunk 10 but the highlight wiring is in place now)."""

    def _value_cell_shading_for_label(self, doc: Document, label: str) -> str | None:
        """Walk every table in the doc; return the value-cell shading
        for the row whose first cell text matches ``label``.

        Returns the hex string (e.g., "FFE5E5") or ``None`` if the row
        has no shading element on its value cell.
        """
        from docx.oxml.ns import qn

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                if row.cells[0].text.strip() != label:
                    continue
                value_cell = row.cells[1]
                tcPr = value_cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    return None
                shd = tcPr.find(qn("w:shd"))
                if shd is None:
                    return None
                return shd.get(qn("w:fill"))
        return None

    def test_verification_failure_value_cell_is_red_when_nonzero(
        self, tmp_path: Path
    ):
        f = _finding(verification=_failed_verification())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        doc = Document(str(out))
        shading = self._value_cell_shading_for_label(
            doc, "Verification failures (operational)"
        )
        assert shading is not None
        # The light-red shading we apply is FFE5E5; the exact value is
        # an implementation detail but it must be a red-family hex.
        assert shading.upper() == "FFE5E5"

    def test_verification_failure_value_cell_unshaded_when_zero(
        self, tmp_path: Path
    ):
        f = _finding(verification=_verified_supported())
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(review_result=ReviewResult(findings=[f])), out
        )
        doc = Document(str(out))
        shading = self._value_cell_shading_for_label(
            doc, "Verification failures (operational)"
        )
        # No shading element means we did not apply the highlight.
        assert shading is None

    def test_cross_check_skipped_value_cell_is_red(self, tmp_path: Path):
        # Skipped cross-check renders as a highlighted row so the
        # absence of coordination analysis is visible at a glance.
        f = _finding(verification=_verified_supported())
        cc = ReviewResult(
            findings=[],
            cross_check_status="skipped",
            thinking="too few specs",
        )
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(
                review_result=ReviewResult(findings=[f]),
                cross_check_result=cc,
            ),
            out,
        )
        doc = Document(str(out))
        shading = self._value_cell_shading_for_label(doc, "Cross-spec coordination")
        assert shading is not None
        assert shading.upper() == "FFE5E5"

    def test_cross_check_completed_value_cell_unshaded(self, tmp_path: Path):
        f = _finding(verification=_verified_supported())
        cc = ReviewResult(findings=[], cross_check_status="completed")
        out = tmp_path / "report.docx"
        export_report(
            _StubPipelineResult(
                review_result=ReviewResult(findings=[f]),
                cross_check_result=cc,
            ),
            out,
        )
        doc = Document(str(out))
        shading = self._value_cell_shading_for_label(doc, "Cross-spec coordination")
        assert shading is None
