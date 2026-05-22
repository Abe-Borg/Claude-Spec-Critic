"""Phase 3 / Step 3.2 — Known-pattern formatting restoration.

When a partial-replacement EDIT crosses runs with distinct character
formatting, ``_replace_in_paragraph`` collapses the affected runs into
the first run's formatting and silently destroys inline emphasis on
tokens like ``NFPA 13`` that used to be bold. The restoration pass
scans the post-mutation replacement span for tokens matching a small
registry of recognized standard / code references and re-applies bold
formatting to each match by splitting the containing run.

This module pins three layers:

1. ``known_pattern_spans`` — pure function that returns
   ``(start, end)`` offsets for every recognized pattern in a string.
2. ``restore_known_formatting_enabled`` — opt-in env-var kill switch
   (default off so the feature ships dormant; operators flip it on
   once their docs prove the pattern registry is conservative
   enough).
3. End-to-end wiring through ``apply_edits_to_spec`` so an EDIT that
   replaces a sentence containing a bolded ``NFPA 13`` rebolds the
   token in the output document.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.editing.edit_locator import EditLocation, LocatorResult
from src.editing.replacement_style import (
    KNOWN_BOLD_PATTERNS,
    known_pattern_spans,
    restore_known_formatting_enabled,
)
from src.editing.spec_editor import apply_edits_to_spec, build_edit_actions
from src.input.extractor import ParagraphMapping, extract_text_from_docx
from src.review.reviewer import Finding


# ---------------------------------------------------------------------------
# known_pattern_spans
# ---------------------------------------------------------------------------


class TestKnownPatternSpans:
    def test_matches_standards_body_with_number(self):
        text = "Comply with NFPA 13 for sprinkler design."
        spans = known_pattern_spans(text)
        assert (12, 19) in spans

    def test_matches_multiple_standards(self):
        text = "Provide ASCE 7-22 and NFPA 13 references."
        spans = known_pattern_spans(text)
        # Both tokens recognized.
        assert any(text[s:e] == "ASCE 7-22" for s, e in spans)
        assert any(text[s:e] == "NFPA 13" for s, e in spans)

    def test_matches_california_code_with_year(self):
        text = "Comply with CBC 2025 seismic provisions."
        spans = known_pattern_spans(text)
        assert any(text[s:e].startswith("CBC ") for s, e in spans)

    def test_matches_csi_section_number(self):
        text = "See Section 23 21 13 for piping requirements."
        spans = known_pattern_spans(text)
        assert any(text[s:e] == "Section 23 21 13" for s, e in spans)

    def test_no_matches_in_prose(self):
        text = "Provide adequate sprinkler coverage throughout."
        assert known_pattern_spans(text) == []

    def test_case_insensitive_matching(self):
        text = "Comply with nfpa 13 for sprinkler design."
        spans = known_pattern_spans(text)
        assert any(text[s:e].lower() == "nfpa 13" for s, e in spans)

    def test_empty_string_no_matches(self):
        assert known_pattern_spans("") == []

    def test_overlapping_matches_merged(self):
        """Two patterns matching the same token produce one merged span."""
        # "NFPA 13" matches twice if any pattern is duplicated; the
        # de-duplication should keep the union, never overlapping.
        text = "Comply with NFPA 13 for sprinkler design."
        spans = known_pattern_spans(text)
        # Spans are sorted and non-overlapping.
        for i in range(1, len(spans)):
            assert spans[i - 1][1] <= spans[i][0]

    def test_patterns_registry_nonempty(self):
        """The registry must have at least one compiled pattern."""
        assert len(KNOWN_BOLD_PATTERNS) >= 1


# ---------------------------------------------------------------------------
# Env-var kill switch
# ---------------------------------------------------------------------------


class TestRestoreKnownFormattingEnabled:
    def test_default_disabled(self, monkeypatch: pytest.MonkeyPatch):
        monkeypatch.delenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", raising=False)
        assert restore_known_formatting_enabled() is False

    @pytest.mark.parametrize("value", ["1", "true", "yes", "on", "TRUE", " on "])
    def test_enabled_via_env(self, monkeypatch: pytest.MonkeyPatch, value: str):
        monkeypatch.setenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", value)
        assert restore_known_formatting_enabled() is True

    @pytest.mark.parametrize("value", ["0", "false", "FALSE", "No", "off"])
    def test_disabled_via_env(self, monkeypatch: pytest.MonkeyPatch, value: str):
        monkeypatch.setenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", value)
        assert restore_known_formatting_enabled() is False


# ---------------------------------------------------------------------------
# Integration: apply_edits_to_spec restores bold on known patterns
# ---------------------------------------------------------------------------


def _locator_result_for_mapping(
    *,
    mapping: ParagraphMapping,
    match_start: int,
    match_end: int,
    matched_text: str,
    replacement_text: str,
    action: str = "EDIT",
) -> LocatorResult:
    location = EditLocation(
        mapping=mapping,
        match_start=match_start,
        match_end=match_end,
        matched_text=matched_text,
        match_confidence=1.0,
        match_method="exact",
    )
    return LocatorResult(
        finding=Finding(
            severity="HIGH",
            fileName="spec.docx",
            section="1.0",
            issue="Issue",
            actionType=action,
            existingText=matched_text,
            replacementText=replacement_text,
            codeReference="Code",
            confidence=0.9,
        ),
        status="matched",
        locations=[location],
        replacement_text=replacement_text,
        action_type=action,
        warning=None,
    )


def _make_paragraph_with_bold_token(
    source: Path,
    *,
    prefix: str = "Comply with ",
    bold_token: str = "NFPA 13",
    suffix: str = " for sprinkler design.",
) -> str:
    """Save a docx whose paragraph carries three runs ``[prefix, BOLD token, suffix]``.

    Returns the full paragraph text so callers can compute offsets.
    """
    doc = Document()
    para = doc.add_paragraph()
    para.add_run(prefix)
    bold_run = para.add_run(bold_token)
    bold_run.bold = True
    para.add_run(suffix)
    doc.save(source)
    return prefix + bold_token + suffix


def test_restoration_rebolds_known_token_when_env_var_enabled(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.setenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", "1")
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    full_text = _make_paragraph_with_bold_token(source)
    spec = extract_text_from_docx(source)
    mapping = next(m for m in spec.paragraph_map or [] if m.element_type == "paragraph")
    new_sentence = "Comply with NFPA 13 for sprinkler density requirements."

    result = _locator_result_for_mapping(
        mapping=mapping,
        match_start=0,
        match_end=len(full_text),
        matched_text=full_text,
        replacement_text=new_sentence,
    )

    actions = build_edit_actions([result], allow_caution=True)
    # The locator downgrades whole-paragraph multi-format EDITs to
    # MANUAL_REVIEW, so build_edit_actions strips this result. For the
    # purpose of this integration test we exercise the restoration
    # pass on a directly-constructed action, which is the same path
    # the apply layer uses; build a partial-span action instead so
    # the locator's downgrade still routes it through.
    if not actions:
        # Fall back to a partial replacement that crosses the bold run
        # — that's the call site Step 3.2 is designed for.
        partial_start = full_text.index("Comply")
        partial_end = full_text.index("for") + len("for")
        new_partial = "Comply with NFPA 13 for"
        result = _locator_result_for_mapping(
            mapping=mapping,
            match_start=partial_start,
            match_end=partial_end,
            matched_text=full_text[partial_start:partial_end],
            replacement_text=new_partial,
        )
        actions = build_edit_actions([result], allow_caution=True)
        new_sentence = new_partial + full_text[partial_end:]

    report = apply_edits_to_spec(source, output, actions)
    assert report.edits_applied == 1

    saved = Document(output)
    paragraph = saved.paragraphs[0]
    assert paragraph.text == new_sentence
    # Bold runs in the saved document must include "NFPA 13".
    bold_run_texts = [run.text for run in paragraph.runs if run.bold]
    assert any("NFPA 13" in text for text in bold_run_texts), (
        f"Expected 'NFPA 13' to be bold; got runs={[(r.text, r.bold) for r in paragraph.runs]}"
    )
    assert report.known_pattern_formatting_restored_count >= 1


def test_restoration_noop_when_env_var_disabled(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    monkeypatch.delenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", raising=False)
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    full_text = _make_paragraph_with_bold_token(source)
    spec = extract_text_from_docx(source)
    mapping = next(m for m in spec.paragraph_map or [] if m.element_type == "paragraph")

    # Partial replacement that crosses the bold run.
    partial_start = full_text.index("Comply")
    partial_end = full_text.index("for") + len("for")
    new_partial = "Comply with NFPA 13 for"
    result = _locator_result_for_mapping(
        mapping=mapping,
        match_start=partial_start,
        match_end=partial_end,
        matched_text=full_text[partial_start:partial_end],
        replacement_text=new_partial,
    )

    actions = build_edit_actions([result], allow_caution=True)
    report = apply_edits_to_spec(source, output, actions)
    assert report.edits_applied == 1

    saved = Document(output)
    paragraph = saved.paragraphs[0]
    # No restoration ran; the bold attribute on "NFPA 13" was lost
    # because the replacement landed in the first (normal) run.
    bold_run_texts = [run.text for run in paragraph.runs if run.bold]
    assert not any("NFPA 13" in text for text in bold_run_texts)
    assert report.known_pattern_formatting_restored_count == 0


def test_restoration_skips_paragraphs_with_uniform_formatting(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """Single-run paragraph; restoration shouldn't fire even with env var on.

    No formatting was ever lost in the first place — the replacement
    span doesn't cross distinct runs.
    """
    monkeypatch.setenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", "1")
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("Comply with general project standards for piping.")
    doc.save(source)
    spec = extract_text_from_docx(source)
    mapping = next(m for m in spec.paragraph_map or [] if m.element_type == "paragraph")

    full_text = mapping.text
    new_text = "Comply with NFPA 13 for piping."
    result = _locator_result_for_mapping(
        mapping=mapping,
        match_start=0,
        match_end=len(full_text),
        matched_text=full_text,
        replacement_text=new_text,
    )
    # Whole-paragraph replacement on a uniformly-formatted paragraph is
    # AUTO_SAFE; the action survives.
    actions = build_edit_actions([result], allow_caution=True)
    report = apply_edits_to_spec(source, output, actions)
    assert report.edits_applied == 1
    # The new sentence contains "NFPA 13" but the original paragraph
    # had no bold runs to "restore" — leave the result alone.
    assert report.known_pattern_formatting_restored_count == 0
    saved = Document(output)
    assert all(run.bold in (None, False) for run in saved.paragraphs[0].runs)


def test_restoration_counter_aggregates_into_diagnostics(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    """``execute_edit_plan`` rolls the per-spec counter into the diagnostics report."""
    monkeypatch.setenv("SPEC_CRITIC_RESTORE_KNOWN_FORMATTING", "1")

    from src.editing.apply_edits import execute_edit_plan
    from src.orchestration.diagnostics import DiagnosticsReport

    source = tmp_path / "spec.docx"
    full_text = _make_paragraph_with_bold_token(source)
    spec = extract_text_from_docx(source)
    spec.filename = "spec.docx"

    # Partial replacement that crosses the bold run so the locator
    # downgrades to AUTO_WITH_CAUTION (still auto-eligible).
    partial_start = full_text.index("Comply")
    partial_end = full_text.index("for") + len("for")
    new_partial = "Comply with NFPA 13 for"
    finding = Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1.0",
        issue="Sentence rewrite",
        actionType="EDIT",
        existingText=full_text[partial_start:partial_end],
        replacementText=new_partial,
        codeReference="NFPA 13",
        confidence=0.9,
    )

    diagnostics = DiagnosticsReport()
    reports = execute_edit_plan(
        selected_finding_indices=[0],
        all_findings=[finding],
        cross_check_findings=[],
        extracted_specs=[spec],
        source_paths=[source],
        output_dir=tmp_path / "out",
        diagnostics=diagnostics,
    )

    assert len(reports) == 1
    assert reports[0].edits_applied == 1
    assert reports[0].known_pattern_formatting_restored_count >= 1
    assert diagnostics.known_pattern_formatting_restored_count >= 1
    text = diagnostics.to_text()
    assert "Known-pattern formatting restored" in text
