from src.editing.edit_locator import locate_edit, locate_edits
from src.input.extractor import ParagraphMapping
from src.review.reviewer import Finding
from src.verification.verifier import VerificationResult


def _mapping(text: str, *, idx: int, element_type: str = "paragraph") -> ParagraphMapping:
    return ParagraphMapping(
        body_index=idx,
        element_type=element_type,
        text=text,
        table_index=0 if element_type == "table_cell" else None,
        row_index=0 if element_type == "table_cell" else None,
        cell_index=None,
    )


def _finding(existing: str | None, replacement: str | None = "new") -> Finding:
    return Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="2.1.A",
        issue="Issue",
        actionType="EDIT",
        existingText=existing,
        replacementText=replacement,
        codeReference="CBC",
        confidence=0.9,
    )


def test_locate_edit_exact_match_single_paragraph():
    paragraph_map = [_mapping("Install copper piping with Type L wall thickness.", idx=0)]
    result = locate_edit(_finding("Type L wall thickness"), paragraph_map)

    assert result.status == "matched"
    assert len(result.locations) == 1
    assert result.locations[0].match_confidence == 1.0
    assert result.locations[0].match_method == "exact"


def test_locate_edit_normalized_match_handles_whitespace_and_case():
    paragraph_map = [_mapping("Provide  FIRE\u00a0RATED    assemblies", idx=0)]
    result = locate_edit(_finding("provide fire rated assemblies"), paragraph_map)

    assert result.status == "matched"
    assert result.locations[0].match_method == "normalized"
    assert result.locations[0].match_confidence == 0.90


def test_locate_edit_fuzzy_match_paraphrase():
    paragraph_map = [_mapping("Provide 2-hour fire-resistance rated shaft enclosures.", idx=0)]
    result = locate_edit(_finding("Provide two hour fire resistance shaft enclosure."), paragraph_map)

    assert result.status == "matched"
    assert result.locations[0].match_method == "fuzzy"
    assert 0.80 <= result.locations[0].match_confidence <= 1.0


def test_locate_edit_ambiguous_when_text_appears_multiple_times():
    paragraph_map = [
        _mapping("Use non-shrink grout at equipment bases.", idx=0),
        _mapping("Use non-shrink grout at equipment bases and supports.", idx=1),
    ]
    result = locate_edit(_finding("Use non-shrink grout at equipment bases"), paragraph_map)

    assert result.status == "ambiguous"
    assert len(result.locations) == 2


def test_locate_edit_not_found():
    paragraph_map = [_mapping("No related text here.", idx=0)]
    result = locate_edit(_finding("Completely absent sentence"), paragraph_map)

    assert result.status == "not_found"


def test_replacement_resolution_from_verification_verdicts():
    paragraph_map = [_mapping("Target text here", idx=0)]

    corrected = _finding("Target text", replacement="orig")
    corrected.verification = VerificationResult(verdict="CORRECTED", correction="better", explanation="", sources=[])

    confirmed = _finding("Target text", replacement="orig")
    confirmed.verification = VerificationResult(verdict="CONFIRMED", correction="ignored", explanation="", sources=[])

    disputed = _finding("Target text", replacement="orig")
    disputed.verification = VerificationResult(verdict="DISPUTED", correction="ignored", explanation="", sources=[])

    corrected_result = locate_edit(corrected, paragraph_map)
    confirmed_result = locate_edit(confirmed, paragraph_map)
    disputed_result = locate_edit(disputed, paragraph_map)

    assert corrected_result.replacement_text == "better"
    assert confirmed_result.replacement_text == "orig"
    assert disputed_result.replacement_text is None


def test_short_text_exact_confidence_is_lowered_and_section_anchored_preferred():
    paragraph_map = [
        _mapping("1.0 GENERAL", idx=0),
        _mapping("Voltage shall be 208V", idx=1),
        _mapping("2.0 PRODUCTS", idx=2),
        _mapping("Voltage shall be 480V", idx=3),
    ]
    finding = _finding("208V")
    finding.section = "1.0"

    result = locate_edit(finding, paragraph_map)

    assert result.status == "matched"
    assert result.locations[0].match_method == "section_anchored"
    assert result.locations[0].match_confidence <= 0.70


def test_section_anchored_match_prefers_csi_heading_scope_for_long_text():
    duplicate = "Material and Thickness: Multilayer, multicolor, plastic labels."
    paragraph_map = [
        _mapping("PART 2 - PRODUCTS", idx=0),
        _mapping("2.01 EQUIPMENT LABELS", idx=1),
        _mapping(duplicate, idx=2),
        _mapping("2.04 DUCT LABELS", idx=3),
        _mapping(duplicate, idx=4),
    ]
    finding = _finding(duplicate)
    finding.section = "PRODUCTS > DUCT LABELS"

    result = locate_edit(finding, paragraph_map)

    assert result.status == "matched"
    assert result.locations[0].mapping.body_index == 4
    assert result.locations[0].match_method == "section_anchored"


def test_none_existing_text_returns_not_found_warning():
    paragraph_map = [_mapping("Anything", idx=0)]
    result = locate_edit(_finding(None), paragraph_map)

    assert result.status == "not_found"
    assert result.warning is not None


def test_cross_paragraph_match_returns_multiple_locations_with_warning():
    paragraph_map = [
        _mapping("PART 1 - GENERAL", idx=0),
        _mapping("Provide submittals within ten days.", idx=1),
        _mapping("Submit operation and maintenance manuals.", idx=2),
    ]
    finding = _finding("Provide submittals within ten days.\n\nSubmit operation and maintenance manuals.")

    result = locate_edit(finding, paragraph_map)

    assert result.status == "matched"
    assert len(result.locations) == 2
    assert result.warning is not None


def test_cross_paragraph_multiple_windows_route_to_manual_review():
    """Step 4.3: when more than one window matches a cross-paragraph
    existingText exactly, the locator refuses to guess and routes the
    finding to manual review.

    The model emitted a 2-paragraph existingText that appears in the
    document in two places (e.g., the same boilerplate ASCE 7
    cross-reference reused in two sections). All cross-paragraph
    matches carry the same flat 0.88 confidence, so the previous
    behavior of picking the first by insertion order would have
    silently chosen one of the two — a coin flip on which paragraph
    actually gets edited. The fix: status="ambiguous",
    safety_category=SAFETY_MANUAL_REVIEW, warning specifically calls
    out the cross-paragraph multi-window case.
    """
    from src.editing.edit_candidates import SAFETY_MANUAL_REVIEW

    paragraph_map = [
        _mapping("PART 1 - GENERAL", idx=0),
        _mapping("Provide submittals within ten days.", idx=1),
        _mapping("Submit operation and maintenance manuals.", idx=2),
        _mapping("PART 2 - PRODUCTS", idx=3),
        # Same 2-paragraph window repeated verbatim — a second valid
        # match site.
        _mapping("Provide submittals within ten days.", idx=4),
        _mapping("Submit operation and maintenance manuals.", idx=5),
    ]
    finding = _finding(
        "Provide submittals within ten days.\n\nSubmit operation and maintenance manuals."
    )

    result = locate_edit(finding, paragraph_map)

    assert result.status == "ambiguous"
    assert result.safety_category == SAFETY_MANUAL_REVIEW
    # Warning text must distinguish "cross-paragraph multiple matches"
    # from the single-match case so users / diagnostics readers
    # immediately understand why this needs manual review.
    assert result.warning is not None
    lower = result.warning.lower()
    assert "cross-paragraph" in lower or "multiple" in lower
    assert "manual review" in lower


def test_cross_paragraph_single_window_still_matches():
    """Step 4.3: when exactly one cross-paragraph window matches, the
    locator behavior is unchanged — status="matched" and the edit can
    proceed under AUTO_WITH_CAUTION (existing safety category).
    """
    paragraph_map = [
        _mapping("PART 1 - GENERAL", idx=0),
        _mapping("Provide submittals within ten days.", idx=1),
        _mapping("Submit operation and maintenance manuals.", idx=2),
    ]
    finding = _finding(
        "Provide submittals within ten days.\n\nSubmit operation and maintenance manuals."
    )

    result = locate_edit(finding, paragraph_map)

    # Single window match — unchanged behavior.
    assert result.status == "matched"
    assert len(result.locations) == 2
    assert result.warning is not None
    # The cross_paragraph_ambiguous flag stays False on the single-window
    # path so the diagnostics counter only counts the truly ambiguous
    # multi-window subset.
    assert result.cross_paragraph_ambiguous is False


def test_cross_paragraph_multiple_windows_flag_set():
    """Step 4.3: the new ``cross_paragraph_ambiguous`` flag is True on
    multi-window matches so apply_edits.execute_edit_plan can count
    them into the diagnostics rollup.
    """
    paragraph_map = [
        _mapping("PART 1 - GENERAL", idx=0),
        _mapping("Provide submittals within ten days.", idx=1),
        _mapping("Submit operation and maintenance manuals.", idx=2),
        _mapping("PART 2 - PRODUCTS", idx=3),
        _mapping("Provide submittals within ten days.", idx=4),
        _mapping("Submit operation and maintenance manuals.", idx=5),
    ]
    finding = _finding(
        "Provide submittals within ten days.\n\nSubmit operation and maintenance manuals."
    )

    result = locate_edit(finding, paragraph_map)

    assert result.status == "ambiguous"
    assert result.cross_paragraph_ambiguous is True
    # No edit action should be produced from build_edit_actions.
    from src.editing.spec_editor import build_edit_actions

    actions = build_edit_actions([result])
    assert actions == []


def test_cross_paragraph_ambiguity_counted_into_diagnostics(tmp_path):
    """Step 4.3 wiring: ``apply_edits.execute_edit_plan`` increments
    ``DiagnosticsReport.cross_paragraph_ambiguity_routed_to_manual_count``
    once per cross-paragraph multi-window ambiguous finding, and the
    "AUTO-APPLY QUALITY" section of ``to_text()`` surfaces the count.
    """
    from docx import Document

    from src.editing.apply_edits import execute_edit_plan
    from src.input.extractor import extract_text_from_docx
    from src.orchestration.diagnostics import DiagnosticsReport
    from src.review.reviewer import Finding

    # Build a docx with two identical 2-paragraph windows so the
    # cross-paragraph match returns 2 windows and the locator routes
    # to manual review.
    source = tmp_path / "spec.docx"
    doc = Document()
    doc.add_paragraph("PART 1 - GENERAL")
    doc.add_paragraph("Provide submittals within ten days.")
    doc.add_paragraph("Submit operation and maintenance manuals.")
    doc.add_paragraph("PART 2 - PRODUCTS")
    doc.add_paragraph("Provide submittals within ten days.")
    doc.add_paragraph("Submit operation and maintenance manuals.")
    doc.save(source)

    spec = extract_text_from_docx(source)
    spec.filename = "spec.docx"

    finding = Finding(
        severity="MEDIUM",
        fileName="spec.docx",
        section="1.0",
        issue="Update boilerplate",
        actionType="EDIT",
        existingText=(
            "Provide submittals within ten days.\n\n"
            "Submit operation and maintenance manuals."
        ),
        replacementText=(
            "Provide submittals within fourteen days.\n\n"
            "Submit operation and maintenance manuals."
        ),
        codeReference="CBC 2025",
        confidence=0.9,
    )

    diagnostics = DiagnosticsReport()
    reports = execute_edit_plan(
        selected_finding_indices=[0],
        all_findings=[finding],
        cross_check_findings=[],
        extracted_specs=[spec],
        source_paths=[source],
        output_dir=tmp_path / "out",
        diagnostics=diagnostics,
    )

    # The locator routed to manual review, so no edit was applied — the
    # report's outcomes are empty (no actions were built) but the
    # diagnostics counter ticked up.
    assert diagnostics.cross_paragraph_ambiguity_routed_to_manual_count == 1
    text = diagnostics.to_text()
    assert "AUTO-APPLY QUALITY" in text
    assert "Cross-paragraph ambiguity routed to manual" in text
    # Sanity: no edit landed (every applied count is zero or the
    # finding turned into a write-the-doc-unchanged copy — either way
    # the multi-window window cannot have been auto-applied).
    if reports:
        for r in reports:
            assert r.edits_applied == 0


def test_table_row_matching_supports_individual_cell_lookup():
    paragraph_map = [_mapping("R1C1 | Allowance Amount | $10,000", idx=0, element_type="table_cell")]
    result = locate_edit(_finding("Allowance Amount"), paragraph_map)

    assert result.status == "matched"
    assert result.locations[0].matched_text == "Allowance Amount"


def test_locate_edits_batch_helper():
    paragraph_map = [_mapping("Paragraph one", idx=0), _mapping("Paragraph two", idx=1)]
    findings = [_finding("Paragraph one"), _finding("missing")]

    results = locate_edits(findings, paragraph_map)

    assert len(results) == 2
    assert results[0].status == "matched"
    assert results[1].status == "not_found"


# ---------------------------------------------------------------------------
# Phase 5 / Step 5.1 — Verifier correction sanity check at the locator.
#
# When verdict is ``CORRECTED`` and ``verification.correction`` is non-empty,
# the locator previously used the correction verbatim as the replacement
# text. The verifier's prompt asks for "1-2 sentences explaining the verdict
# and the corrected reference text" — that's explanation, not clean
# replacement text. The locator now runs a sanity check and falls back to
# ``proposal.replacement_text`` when the correction does not look
# replaceable. The verifier's correction is preserved on the
# ``VerificationResult`` for the report; only the applied edit text changes.
# ---------------------------------------------------------------------------


def test_clean_corrected_correction_still_used_as_replacement():
    """A short, prose-only CORRECTED.correction still wins over replacement_text.

    Pins the existing behavior for clean corrections (the legacy
    ``test_replacement_resolution_from_verification_verdicts`` case) so the
    new sanity check does not over-fire on the common path.
    """
    paragraph_map = [_mapping("Target text here", idx=0)]
    finding = _finding("Target text", replacement="orig")
    finding.verification = VerificationResult(
        verdict="CORRECTED",
        correction="ASCE 7-22",
        explanation="",
        sources=[],
    )

    result = locate_edit(finding, paragraph_map)

    assert result.replacement_text == "ASCE 7-22"
    assert getattr(result, "correction_rejected_as_replacement", False) is False


def test_paragraph_length_correction_falls_back_to_replacement_text():
    """A ~10× longer correction triggers the sanity check fallback."""
    paragraph_map = [_mapping("Use ASCE 7-16 for seismic design.", idx=0)]
    finding = _finding("Use ASCE 7-16", replacement="Use ASCE 7-22")
    finding.verification = VerificationResult(
        verdict="CORRECTED",
        correction=(
            "The applicable standard is ASCE 7-22 because the 2025 California "
            "Building Code adopted that revision through Title 24 Part 2, and "
            "the 7-16 reference is obsolete for projects permitted after "
            "January 1 2026."
        ),
        explanation="",
        sources=[],
    )

    result = locate_edit(finding, paragraph_map)

    # The applied edit uses the model's original replacement, not the
    # verifier's explanatory paragraph.
    assert result.replacement_text == "Use ASCE 7-22"
    assert result.correction_rejected_as_replacement is True
    # The verifier's correction is preserved on the result for the report.
    assert finding.verification.correction.startswith("The applicable standard")


def test_parenthetical_citation_correction_falls_back_to_replacement_text():
    """``(per CBC § 1613.1)``-style citations are explanatory."""
    paragraph_map = [_mapping("Use ASCE 7-16 for seismic design.", idx=0)]
    finding = _finding("Use ASCE 7-16", replacement="Use ASCE 7-22")
    finding.verification = VerificationResult(
        verdict="CORRECTED",
        correction="Use ASCE 7-22 (per CBC § 1613.1).",
        explanation="",
        sources=[],
    )

    result = locate_edit(finding, paragraph_map)

    assert result.replacement_text == "Use ASCE 7-22"
    assert result.correction_rejected_as_replacement is True


def test_url_in_correction_falls_back_to_replacement_text():
    """URLs in the correction never belong in spec body text."""
    paragraph_map = [_mapping("Use ASCE 7-16 for seismic design.", idx=0)]
    finding = _finding("Use ASCE 7-16", replacement="Use ASCE 7-22")
    finding.verification = VerificationResult(
        verdict="CORRECTED",
        correction="See https://www.iccsafe.org/asce-7-22.",
        explanation="",
        sources=[],
    )

    result = locate_edit(finding, paragraph_map)

    assert result.replacement_text == "Use ASCE 7-22"
    assert result.correction_rejected_as_replacement is True


def test_env_var_preserves_legacy_verbatim_behavior(
    monkeypatch,
):
    """``SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT=1`` skips the sanity check."""
    monkeypatch.setenv(
        "SPEC_CRITIC_USE_VERIFIER_CORRECTION_AS_REPLACEMENT", "1"
    )
    paragraph_map = [_mapping("Use ASCE 7-16 for seismic design.", idx=0)]
    finding = _finding("Use ASCE 7-16", replacement="Use ASCE 7-22")
    finding.verification = VerificationResult(
        verdict="CORRECTED",
        correction="Use ASCE 7-22 (per CBC § 1613.1).",
        explanation="",
        sources=[],
    )

    result = locate_edit(finding, paragraph_map)

    # Env var on → legacy verbatim path: use the parenthetical correction.
    assert result.replacement_text == "Use ASCE 7-22 (per CBC § 1613.1)."
    assert result.correction_rejected_as_replacement is False


def test_locator_result_correction_rejected_default_false_on_legacy_construction():
    """LocatorResult constructed without the new field still loads safely.

    Resume-state payloads from before Step 5.1 don't carry the new
    ``correction_rejected_as_replacement`` field; the dataclass default
    keeps them loading cleanly.
    """
    from src.editing.edit_locator import EditLocation, LocatorResult

    finding = _finding("foo", replacement="bar")
    mapping = _mapping("foo", idx=0)
    location = EditLocation(
        mapping=mapping,
        match_start=0,
        match_end=3,
        matched_text="foo",
        match_confidence=1.0,
        match_method="exact",
    )
    result = LocatorResult(
        finding=finding,
        status="matched",
        locations=[location],
        replacement_text="bar",
        action_type="EDIT",
        warning=None,
    )
    assert result.correction_rejected_as_replacement is False


def test_correction_rejection_counted_into_diagnostics(tmp_path):
    """End-to-end: a paragraph-length correction is rejected and counted.

    Pins the wiring from
    ``LocatorResult.correction_rejected_as_replacement`` through
    ``apply_edits.execute_edit_plan`` and into
    ``DiagnosticsReport.verifier_correction_rejected_as_replacement_count``.
    """
    from docx import Document

    from src.editing.apply_edits import execute_edit_plan
    from src.input.extractor import extract_text_from_docx
    from src.orchestration.diagnostics import DiagnosticsReport

    source = tmp_path / "spec.docx"
    doc = Document()
    doc.add_paragraph("Use ASCE 7-16 for seismic design.")
    doc.save(source)

    spec = extract_text_from_docx(source)
    spec.filename = "spec.docx"

    finding = Finding(
        severity="HIGH",
        fileName="spec.docx",
        section="1.0",
        issue="Cite the current edition of ASCE 7.",
        actionType="EDIT",
        existingText="ASCE 7-16",
        replacementText="ASCE 7-22",
        codeReference="ASCE 7",
        confidence=0.9,
    )
    finding.verification = VerificationResult(
        verdict="CORRECTED",
        correction=(
            "The applicable standard is ASCE 7-22 because the 2025 "
            "California Building Code adopted that revision through Title "
            "24 Part 2, and the 7-16 reference is obsolete for projects "
            "permitted after January 1 2026."
        ),
        explanation="",
        sources=["https://example.com/asce-7-22"],
    )

    diagnostics = DiagnosticsReport()
    reports = execute_edit_plan(
        selected_finding_indices=[0],
        all_findings=[finding],
        cross_check_findings=[],
        extracted_specs=[spec],
        source_paths=[source],
        output_dir=tmp_path / "out",
        diagnostics=diagnostics,
    )

    assert len(reports) == 1
    # The edit still applied — we kept the CORRECTED verdict — but with
    # the model's original replacement text, not the verifier's paragraph.
    assert reports[0].edits_applied == 1
    saved_text = (
        Document(reports[0].output_path).paragraphs[0].text
    )
    assert "ASCE 7-22" in saved_text
    # The applied replacement did NOT include the verifier's explanatory
    # paragraph.
    assert "California Building Code adopted" not in saved_text

    # Counter rolled up at both layers.
    assert reports[0].verifier_correction_rejected_as_replacement_count == 1
    assert diagnostics.verifier_correction_rejected_as_replacement_count == 1
    # And the AUTO-APPLY QUALITY rollup surfaces it.
    text = diagnostics.to_text()
    assert "AUTO-APPLY QUALITY" in text
    assert "Verifier correction rejected as replacement" in text
