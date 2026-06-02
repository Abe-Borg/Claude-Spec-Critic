"""Tests for surfacing review-stage failures in the final artifact (P0-1).

A spec whose individual review *failed* (truncated / parse-error /
errored / no result) produces zero findings — exactly like a spec that
was reviewed and found clean. Without an explicit signal the two are
indistinguishable in the exported report, so a partially-failed run
reads as a fully-clean one. That is the single place a compliance-review
tool can actively mislead.

These tests pin the fix end to end:

1. **Data plane** — ``collect_review_batch_results`` records the failed
   spec in ``truncated_specs`` and ``finalize_batch_result`` carries it
   onto ``PipelineResult.failed_review_specs``.
2. **Summary** — ``_summarize_run_diagnostics`` rolls the failed specs
   into ``failed_review_count`` / ``failed_review_specs``.
3. **Report** — ``export_report`` distinguishes the failed spec: a red
   banner row, a recovery hint naming it, a corrected "Files Reviewed:
   {reviewed} of {submitted}" line, and a red annotation on the bullet.
   A clean run keeps the original (unhighlighted, "Files Reviewed: N")
   shape.
"""
from __future__ import annotations

import time
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.batch.batch import BatchJob
from src.orchestration import pipeline as pl
from src.orchestration.pipeline import (
    BatchSubmission,
    PipelineResult,
    collect_review_batch_results,
    finalize_batch_result,
)
from src.output.report_exporter import _summarize_run_diagnostics, export_report
from src.output.report_status import summarize_edit_actions, summarize_statuses
from src.review.reviewer import Finding, ReviewResult


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(*, file: str = "A.docx", issue: str = "Issue") -> Finding:
    return Finding(
        severity="HIGH",
        fileName=file,
        section="2.1",
        issue=issue,
        actionType="REPORT_ONLY",
        existingText=None,
        replacementText=None,
        codeReference=None,
    )


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _value_cell_shading_for_label(doc: Document, label: str) -> str | None:
    """Return the value-cell fill hex for the banner row whose first cell
    matches ``label`` (or ``None`` when the cell carries no shading)."""
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            if row.cells[0].text.strip() != label:
                continue
            tcPr = row.cells[1]._tc.find(qn("w:tcPr"))
            if tcPr is None:
                return None
            shd = tcPr.find(qn("w:shd"))
            if shd is None:
                return None
            return shd.get(qn("w:fill"))
    return None


_FAILED_ROW = "Specs that failed review (not reviewed)"


# ---------------------------------------------------------------------------
# 1. Data plane: collect -> finalize carries the failed spec through
# ---------------------------------------------------------------------------


class TestDataPlane:
    def _submission(self) -> BatchSubmission:
        job = BatchJob(
            batch_id="batch-1",
            job_type="review",
            request_map={
                "review__a__0": {"filename": "A.docx", "index": 0, "type": "review"},
                "review__b__1": {"filename": "B.docx", "index": 1, "type": "review"},
                "review__c__2": {"filename": "C.docx", "index": 2, "type": "review"},
            },
            created_at=time.time(),
        )
        return BatchSubmission(
            job=job,
            files_reviewed=["A.docx", "B.docx", "C.docx"],
            review_request_ids=["review__a__0", "review__b__1", "review__c__2"],
            model="claude-opus-4-8",
            # No prepared specs => the repair batch is skipped (it needs the
            # original extracted specs), so the truncated result stays
            # truncated instead of triggering a real API call.
            prepared_specs=None,
        )

    def test_truncated_spec_flows_to_pipeline_result(self, monkeypatch):
        # Spec B's review came back incomplete (token-limit truncation);
        # A and C completed cleanly.
        results = {
            "review__a__0": ReviewResult(findings=[_finding(file="A.docx")], parse_status="complete"),
            "review__b__1": ReviewResult(findings=[], parse_status="incomplete"),
            "review__c__2": ReviewResult(findings=[_finding(file="C.docx")], parse_status="complete"),
        }
        monkeypatch.setattr(
            pl, "retrieve_review_results", lambda job, *, model: dict(results)
        )

        state = collect_review_batch_results(self._submission())
        assert state.truncated_specs == ["B.docx"]
        # The combined result carries an honest spec-error summary.
        assert state.review_result.error
        assert "B.docx" in state.review_result.error

        result = finalize_batch_result(state)
        assert result.failed_review_specs == ["B.docx"]
        # files_reviewed still lists every submitted spec.
        assert result.files_reviewed == ["A.docx", "B.docx", "C.docx"]

    def test_clean_run_has_no_failed_specs(self, monkeypatch):
        results = {
            "review__a__0": ReviewResult(findings=[_finding(file="A.docx")], parse_status="complete"),
            "review__b__1": ReviewResult(findings=[_finding(file="B.docx")], parse_status="complete"),
            "review__c__2": ReviewResult(findings=[_finding(file="C.docx")], parse_status="complete"),
        }
        monkeypatch.setattr(
            pl, "retrieve_review_results", lambda job, *, model: dict(results)
        )
        state = collect_review_batch_results(self._submission())
        assert state.truncated_specs == []
        result = finalize_batch_result(state)
        assert result.failed_review_specs == []


# ---------------------------------------------------------------------------
# 2. Summary rollup
# ---------------------------------------------------------------------------


class TestSummarize:
    def _summary(self, pipeline_result):
        findings = list(pipeline_result.review_result.findings)
        return _summarize_run_diagnostics(
            findings=findings,
            status_counts=summarize_statuses(findings),
            edit_action_counts=summarize_edit_actions(findings),
            cross_check_result=None,
            pipeline_result=pipeline_result,
        )

    def test_failed_specs_counted_and_named(self):
        pr = PipelineResult(
            review_result=ReviewResult(findings=[_finding(file="A.docx")]),
            files_reviewed=["A.docx", "B.docx", "C.docx"],
            failed_review_specs=["B.docx", "C.docx"],
        )
        summary = self._summary(pr)
        assert summary["failed_review_count"] == 2
        assert summary["failed_review_specs"] == ["B.docx", "C.docx"]

    def test_clean_run_reports_zero(self):
        pr = PipelineResult(
            review_result=ReviewResult(findings=[_finding(file="A.docx")]),
            files_reviewed=["A.docx"],
        )
        summary = self._summary(pr)
        assert summary["failed_review_count"] == 0
        assert summary["failed_review_specs"] == []

    def test_legacy_pipeline_result_without_field_defaults_to_zero(self):
        # A duck-typed pipeline result that predates the field must not
        # crash the summary; the defensive getattr resolves it to 0.
        class _Legacy:
            review_result = ReviewResult(findings=[])

        summary = _summarize_run_diagnostics(
            findings=[],
            status_counts={},
            edit_action_counts={},
            cross_check_result=None,
            pipeline_result=_Legacy(),
        )
        assert summary["failed_review_count"] == 0
        assert summary["failed_review_specs"] == []


# ---------------------------------------------------------------------------
# 3. Report rendering — the exported artifact distinguishes the failure
# ---------------------------------------------------------------------------


class TestReportRendering:
    def _partial_failure_result(self) -> PipelineResult:
        return PipelineResult(
            review_result=ReviewResult(
                findings=[_finding(file="A.docx")],
                error="1 spec(s) had errors: B.docx: Review response truncated",
                model="claude-opus-4-8",
            ),
            files_reviewed=["A.docx", "B.docx", "C.docx"],
            failed_review_specs=["B.docx"],
            cycle_label="2025",
        )

    def _clean_result(self) -> PipelineResult:
        return PipelineResult(
            review_result=ReviewResult(
                findings=[_finding(file="A.docx")], model="claude-opus-4-8"
            ),
            files_reviewed=["A.docx", "B.docx", "C.docx"],
            cycle_label="2025",
        )

    def test_banner_row_present_and_highlighted_on_failure(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(self._partial_failure_result(), out)
        doc = Document(str(out))
        text = _all_text_from(doc)
        assert _FAILED_ROW in text
        shading = _value_cell_shading_for_label(doc, _FAILED_ROW)
        assert shading is not None and shading.upper() == "FFE5E5"

    def test_failure_hint_names_the_spec(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(self._partial_failure_result(), out)
        text = _all_text_from(Document(str(out)))
        assert "1 spec failed review" in text
        assert "B.docx" in text
        # The hint warns that no-findings does not imply compliant.
        assert "does NOT mean" in text

    def test_files_reviewed_count_corrected_on_failure(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(self._partial_failure_result(), out)
        text = _all_text_from(Document(str(out)))
        # 3 submitted, 1 failed => 2 reviewed.
        assert "Files Reviewed: 2 of 3 (1 failed review)" in text

    def test_failed_bullet_annotated(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(self._partial_failure_result(), out)
        text = _all_text_from(Document(str(out)))
        assert "B.docx — review failed (not reviewed)" in text

    def test_clean_run_banner_row_unhighlighted(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(self._clean_result(), out)
        doc = Document(str(out))
        # Row still renders (banner shape is stable) but carries no
        # highlight and shows 0.
        assert _FAILED_ROW in _all_text_from(doc)
        assert _value_cell_shading_for_label(doc, _FAILED_ROW) is None

    def test_clean_run_keeps_plain_files_reviewed_line(self, tmp_path: Path):
        out = tmp_path / "report.docx"
        export_report(self._clean_result(), out)
        text = _all_text_from(Document(str(out)))
        assert "Files Reviewed: 3" in text
        # The corrected "{reviewed} of {submitted}" form and the per-bullet
        # failure annotation appear only on a partial-failure run. (The
        # banner *row label* legitimately contains "failed review", so we
        # assert the specific failure markers rather than that substring.)
        assert " of 3 (" not in text
        assert "review failed (not reviewed)" not in text
