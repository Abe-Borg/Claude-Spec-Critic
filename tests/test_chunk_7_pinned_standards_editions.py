"""Chunk 7 / Trust Upgrade tests — pinned standards editions.

The Trust Upgrade Chunk 7 pins the NFPA, ASHRAE, IAPMO, and UL editions
adopted by California for the current cycle. The contract spans four
surfaces:

* ``CodeCycle`` carries the new edition fields, with ``CALIFORNIA_2025``
  populated with the adoption matrix.
* The reviewer system prompt and user message reference the pinned
  editions so the model knows which editions to compare against.
* The verifier system prompt renders a "Pinned standards editions"
  block before the search budget, with explicit instructions to flag
  edition drift.
* The methodology note in the exported report enumerates the pinned
  editions so reviewers see which editions drove the verdicts.

This module covers the plan's success criteria:

* The verifier system prompt contains the pinned editions.
* A test finding citing NFPA 13 receives a verdict that references the
  correct adopted edition (covered by the prompt-content assertions;
  end-to-end model behavior is exercised in the eval harness).
* The report methodology note enumerates the pinned editions.
"""
from __future__ import annotations

from dataclasses import replace
from pathlib import Path

from docx import Document

from src.core.code_cycles import (
    AVAILABLE_CYCLES,
    CALIFORNIA_2025,
    CodeCycle,
    DEFAULT_CYCLE,
)
from src.output.report_exporter import (
    _render_pinned_editions_note,
    _write_methodology_note,
    export_report,
)
from src.review.prompts import get_single_spec_user_message, get_system_prompt
from src.review.reviewer import Finding, ReviewResult
from src.verification.verifier import (
    _get_verification_system_prompt,
    _pinned_standards_lines,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _finding(file: str = "test.docx") -> Finding:
    return Finding(
        severity="HIGH",
        fileName=file,
        section="2.1",
        issue="Stale code reference",
        actionType="EDIT",
        existingText="2019 CBC",
        replacementText="2025 CBC",
        codeReference="CBC §1234",
        confidence=0.8,
    )


class _StubPipelineResult:
    """Minimal duck-typed PipelineResult for export_report."""

    def __init__(
        self,
        *,
        review_result: ReviewResult,
        cycle_label: str = "2025",
    ):
        self.review_result = review_result
        self.cross_check_result = None
        self.files_reviewed = (
            [review_result.findings[0].fileName] if review_result.findings else ["test.docx"]
        )
        self.leed_alerts = []
        self.placeholder_alerts = []
        self.cycle_label = cycle_label
        self.total_elapsed_seconds = 1.0


def _all_text_from(doc: Document) -> str:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ===========================================================================
# 1. CodeCycle dataclass carries the new fields
# ===========================================================================


class TestCodeCycleFields:
    def test_codecycle_has_nfpa_fields(self):
        # Every NFPA standard from the plan is on the dataclass.
        for attr in ("nfpa13", "nfpa14", "nfpa20", "nfpa24", "nfpa25", "nfpa72"):
            assert hasattr(CALIFORNIA_2025, attr), f"missing field: {attr}"

    def test_codecycle_has_ashrae_fields(self):
        for attr in ("ashrae_62_1", "ashrae_90_1", "ashrae_15"):
            assert hasattr(CALIFORNIA_2025, attr), f"missing field: {attr}"

    def test_codecycle_has_iapmo_field(self):
        assert hasattr(CALIFORNIA_2025, "iapmo_tsc")

    def test_codecycle_has_ul_listing_editions(self):
        # UL editions are stored as a tuple of (standard, edition) pairs
        # (not a dict) so the dataclass stays hashable under frozen=True.
        assert isinstance(CALIFORNIA_2025.ul_listing_editions, tuple)
        # Every UL standard from the plan is present.
        standards = {s for s, _ in CALIFORNIA_2025.ul_listing_editions}
        assert "UL 300" in standards
        assert "UL 555" in standards
        assert "UL 555S" in standards
        assert "UL 268" in standards
        assert "UL 1479" in standards

    def test_codecycle_is_hashable(self):
        # Frozen dataclasses with hashable fields should remain hashable.
        # Switching ul_listing_editions to a tuple-of-tuples preserves this
        # because a dict field would break __hash__.
        assert hash(CALIFORNIA_2025) == hash(CALIFORNIA_2025)
        # Round-trip through a set / dict key works.
        assert CALIFORNIA_2025 in {CALIFORNIA_2025}

    def test_california_2025_has_nonempty_pinned_editions(self):
        # The adoption matrix is populated — pinning isn't merely a
        # schema change with no data.
        assert CALIFORNIA_2025.nfpa13
        assert CALIFORNIA_2025.nfpa72
        assert CALIFORNIA_2025.ashrae_62_1
        assert CALIFORNIA_2025.ashrae_90_1
        assert CALIFORNIA_2025.ul_listing_editions

    def test_codecycle_field_defaults_are_empty(self):
        # A future cycle that doesn't populate these defaults should
        # still construct fine.
        c = CodeCycle(
            label="future",
            cbc="2028",
            cmc="2028",
            cpc="2028",
            energy_code="2028",
            calgreen="2028",
            asce7="7-22",
            asce7_previous="7-16",
        )
        assert c.nfpa13 == ""
        assert c.ashrae_62_1 == ""
        assert c.ul_listing_editions == ()


# ===========================================================================
# 2. Reviewer system prompt + user message reference pinned editions
# ===========================================================================


class TestReviewerPromptPinnedEditions:
    def test_system_prompt_mentions_nfpa_editions(self):
        sp = get_system_prompt(CALIFORNIA_2025)
        # Code edition misalignment category now cites NFPA + ASHRAE.
        assert f"NFPA 13 {CALIFORNIA_2025.nfpa13}" in sp
        assert f"NFPA 72 {CALIFORNIA_2025.nfpa72}" in sp

    def test_system_prompt_mentions_ashrae_editions(self):
        sp = get_system_prompt(CALIFORNIA_2025)
        assert f"ASHRAE 62.1 {CALIFORNIA_2025.ashrae_62_1}" in sp
        assert f"ASHRAE 90.1 {CALIFORNIA_2025.ashrae_90_1}" in sp

    def test_user_message_mentions_pinned_editions(self):
        um = get_single_spec_user_message(
            "spec content", "23 21 13 - Hydronic.docx", cycle=CALIFORNIA_2025
        )
        assert "NFPA 13" in um
        assert "NFPA 72" in um
        assert "ASHRAE 62.1" in um
        assert "ASHRAE 90.1" in um

    def test_system_prompt_stable_across_calls(self):
        # Cache breakpoint invariant — adding pinned editions must keep
        # the prompt byte-stable across calls.
        assert get_system_prompt(CALIFORNIA_2025) == get_system_prompt(CALIFORNIA_2025)

    def test_system_prompt_changes_when_pinned_edition_changes(self):
        # Two cycles that differ only in pinned NFPA 13 edition should
        # produce different system prompts.
        alt = replace(CALIFORNIA_2025, nfpa13="2025 edition")
        sp_a = get_system_prompt(CALIFORNIA_2025)
        sp_b = get_system_prompt(alt)
        assert sp_a != sp_b


# ===========================================================================
# 3. Verifier system prompt block
# ===========================================================================


class TestVerifierPinnedEditionsBlock:
    def test_pinned_lines_render_each_populated_field(self):
        lines = _pinned_standards_lines(CALIFORNIA_2025)
        joined = "\n".join(lines)
        # Header is present.
        assert "Pinned standards editions" in joined
        # Each populated field appears with its label.
        assert f"NFPA 13: {CALIFORNIA_2025.nfpa13}" in joined
        assert f"NFPA 72: {CALIFORNIA_2025.nfpa72}" in joined
        assert f"ASHRAE 62.1: {CALIFORNIA_2025.ashrae_62_1}" in joined
        assert f"ASHRAE 90.1: {CALIFORNIA_2025.ashrae_90_1}" in joined
        # IAPMO TSC + UL editions surface too.
        assert "IAPMO Uniform Plumbing TSC" in joined
        assert "UL 300:" in joined
        assert "UL 555:" in joined

    def test_pinned_lines_include_drift_instruction(self):
        lines = _pinned_standards_lines(CALIFORNIA_2025)
        joined = "\n".join(lines)
        # The plan requires the prompt to explicitly tell the model to
        # flag a different edition than the one pinned.
        assert "flag" in joined.lower()
        assert "edition" in joined.lower()

    def test_pinned_lines_empty_when_no_editions(self):
        # A cycle with no pinned standards renders an empty block so
        # the prompt doesn't make a claim that isn't true.
        bare = CodeCycle(
            label="bare",
            cbc="2025",
            cmc="2025",
            cpc="2025",
            energy_code="2025",
            calgreen="2025",
            asce7="7-22",
            asce7_previous="7-16",
        )
        assert _pinned_standards_lines(bare) == []

    def test_pinned_lines_skip_empty_fields(self):
        # A partial cycle (only NFPA 13 populated) renders only that
        # field — empty fields are silently dropped.
        partial = replace(
            CodeCycle(
                label="partial",
                cbc="2025",
                cmc="2025",
                cpc="2025",
                energy_code="2025",
                calgreen="2025",
                asce7="7-22",
                asce7_previous="7-16",
            ),
            nfpa13="2022",
        )
        lines = _pinned_standards_lines(partial)
        joined = "\n".join(lines)
        assert "NFPA 13: 2022" in joined
        assert "NFPA 72:" not in joined
        assert "ASHRAE" not in joined

    def test_verifier_system_prompt_contains_pinned_block(self):
        prompt = _get_verification_system_prompt(
            CALIFORNIA_2025, include_verdict_tool=True
        )
        assert "Pinned standards editions" in prompt
        assert "NFPA 13" in prompt
        assert "ASHRAE 62.1" in prompt
        assert "UL 1479" in prompt

    def test_verifier_pinned_block_renders_for_text_fallback_prompt(self):
        # The text-fallback branch (verdict tool disabled) must also
        # carry the pinned editions — the model verifies the same way
        # regardless of which branch it ends up in.
        prompt = _get_verification_system_prompt(
            CALIFORNIA_2025, include_verdict_tool=False
        )
        assert "Pinned standards editions" in prompt
        assert "NFPA 13" in prompt

    def test_verifier_pinned_block_before_search_budget(self):
        # Pinned editions render before the search budget so the model
        # reads them while still framing its first query.
        prompt = _get_verification_system_prompt(
            CALIFORNIA_2025, include_verdict_tool=True
        )
        pinned_idx = prompt.find("Pinned standards editions")
        budget_idx = prompt.find("Search budget")
        assert pinned_idx >= 0
        assert budget_idx >= 0
        assert pinned_idx < budget_idx

    def test_verifier_prompt_stable_across_calls(self):
        # Cache breakpoint invariant — the verifier system prompt must
        # remain byte-stable across calls so cache breakpoints land in
        # the same place.
        a = _get_verification_system_prompt(CALIFORNIA_2025, include_verdict_tool=True)
        b = _get_verification_system_prompt(CALIFORNIA_2025, include_verdict_tool=True)
        assert a == b


# ===========================================================================
# 4. Methodology note enumerates pinned editions
# ===========================================================================


class TestMethodologyNotePinnedEditions:
    def test_render_note_includes_all_populated_editions(self):
        note = _render_pinned_editions_note("2025")
        # Each populated edition appears in the rendered note.
        assert f"NFPA 13 {CALIFORNIA_2025.nfpa13}" in note
        assert f"ASHRAE 62.1 {CALIFORNIA_2025.ashrae_62_1}" in note
        assert "UL 300" in note

    def test_render_note_mentions_cycle_label(self):
        note = _render_pinned_editions_note("2025")
        assert "2025" in note
        assert "California" in note

    def test_render_note_advises_reviewers(self):
        # The note's purpose is to tell reviewers what to do when a
        # finding cites a different edition.
        note = _render_pinned_editions_note("2025")
        assert "reviewed for relevance" in note or "review" in note.lower()

    def test_render_note_empty_when_cycle_has_no_pinning(self):
        # Unknown cycle labels fall back to DEFAULT_CYCLE; if the
        # default happens to have no pinning the note returns "".
        # We can't test that directly with CALIFORNIA_2025 (which is
        # the default and IS populated), but we can verify the empty
        # branch by stubbing AVAILABLE_CYCLES via the dataclass.
        bare = CodeCycle(
            label="bare",
            cbc="2025",
            cmc="2025",
            cpc="2025",
            energy_code="2025",
            calgreen="2025",
            asce7="7-22",
            asce7_previous="7-16",
        )
        # Inject a fake "bare" cycle for the duration of this test.
        AVAILABLE_CYCLES["__test_bare__"] = bare
        try:
            assert _render_pinned_editions_note("__test_bare__") == ""
        finally:
            AVAILABLE_CYCLES.pop("__test_bare__", None)

    def test_render_note_unknown_label_falls_back_to_default(self):
        # An unknown label should not crash — it should use DEFAULT_CYCLE
        # which is CALIFORNIA_2025 (populated). This protects future
        # callers that pass an arbitrary cycle_label string.
        note = _render_pinned_editions_note("not-a-real-cycle")
        assert "NFPA 13" in note

    def test_methodology_note_in_doc_contains_pinned_editions(self, tmp_path: Path):
        # End-to-end: an exported report's methodology section should
        # surface the pinned editions list so reviewers see it inline.
        doc = Document()
        _write_methodology_note(doc, cycle_label="2025")
        text = _all_text_from(doc)
        assert "About This Review" in text
        assert "pinned the following standards editions" in text
        assert "NFPA 13" in text
        assert "ASHRAE 62.1" in text


# ===========================================================================
# 5. End-to-end report export carries pinned editions
# ===========================================================================


class TestExportedReportPinnedEditions:
    def test_exported_report_contains_pinned_editions_paragraph(self, tmp_path: Path):
        f = _finding()
        result = ReviewResult(findings=[f])
        out = tmp_path / "report.docx"
        export_report(_StubPipelineResult(review_result=result), out)
        text = _all_text_from(Document(str(out)))
        # The methodology paragraph appended by Chunk 7 must surface in
        # the final exported document.
        assert "pinned the following standards editions" in text
        # Cross-check that at least one pinned standard actually appears.
        assert "NFPA 13" in text


# ===========================================================================
# 6. DEFAULT_CYCLE remains California 2025 (regression guard)
# ===========================================================================


class TestDefaultCycleInvariant:
    def test_default_cycle_is_california_2025(self):
        # Plan invariant from CLAUDE.md §1: California 2025 is the only
        # supported cycle. Chunk 7 must not introduce a new cycle.
        assert DEFAULT_CYCLE is CALIFORNIA_2025
        assert list(AVAILABLE_CYCLES.keys()) == ["2025"]
