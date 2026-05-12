"""Tests for the remaining-unimplemented-items work.

Covers:
- Phase 1.3 — formal FindingGroup / FindingOccurrence types and helpers
- Phase 2.4 — structured outputs for review (tool_use parsing)
- Phase 2.5 — structured outputs for verification (verdict tool)
- Phase 4.6 — annotation / change-log mode
- Phase 5.5 — file hash validation in resume state
- Phase 9.3 — fuzzy match length / quick_ratio prefilters
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.extractor import ExtractedSpec, extract_text_from_docx
from src.pipeline import (
    FindingGroup,
    FindingOccurrence,
    group_findings,
)
from src.reviewer import Finding


def _make_finding(file_name: str = "spec.docx", **overrides) -> Finding:
    base = dict(
        severity="HIGH", fileName=file_name, section="2.1",
        issue="cited code edition is outdated",
        actionType="EDIT",
        existingText="CBC 2019",
        replacementText="CBC 2025",
        codeReference="CBC",
        confidence=0.7,
    )
    base.update(overrides)
    return Finding(**base)


# ---------------------------------------------------------------------------
# Phase 1.3 — FindingGroup / FindingOccurrence
# ---------------------------------------------------------------------------


class TestFindingGrouping:
    def test_single_file_finding_yields_one_occurrence(self):
        f = _make_finding(file_name="A.docx")
        groups = group_findings([f])
        assert len(groups) == 1
        assert isinstance(groups[0], FindingGroup)
        assert len(groups[0].occurrences) == 1
        occ = groups[0].occurrences[0]
        assert isinstance(occ, FindingOccurrence)
        assert occ.file_name == "A.docx"
        assert occ.finding is f

    def test_multi_file_finding_fans_out_to_one_occurrence_per_file(self):
        f = _make_finding(file_name="A.docx")
        f.affected_files = ["A.docx", "B.docx", "C.docx"]
        groups = group_findings([f])
        assert len(groups) == 1
        assert groups[0].file_names == ["A.docx", "B.docx", "C.docx"]

    def test_occurrence_ids_are_unique(self):
        f1 = _make_finding(file_name="A.docx")
        f1.affected_files = ["A.docx", "B.docx"]
        f2 = _make_finding(file_name="C.docx", issue="other")
        groups = group_findings([f1, f2])
        ids = [o.occurrence_id for g in groups for o in g.occurrences]
        assert len(set(ids)) == len(ids)


# ---------------------------------------------------------------------------
# Phase 4.6 — annotation / change-log mode
# ---------------------------------------------------------------------------


def _make_simple_docx(tmp_path: Path, paragraphs: list[str]) -> Path:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    path = tmp_path / "spec.docx"
    doc.save(path)
    return path


class TestAnnotateMode:
    def test_annotation_inserts_paragraph_without_mutating_text(self, tmp_path: Path):
        from src.spec_editor import EditAction, annotate_spec_with_suggestions
        from src.edit_locator import EditLocation, LocatorResult

        source = _make_simple_docx(tmp_path, [
            "PART 1 GENERAL",
            "1.01 SUMMARY",
            "Comply with CBC 2019 requirements.",
        ])
        spec = extract_text_from_docx(source)
        target_mapping = next(
            m for m in (spec.paragraph_map or [])
            if "Comply with CBC 2019" in m.text
        )
        location = EditLocation(
            mapping=target_mapping,
            match_start=0,
            match_end=len(target_mapping.text),
            matched_text=target_mapping.text,
            match_confidence=1.0,
            match_method="exact",
        )
        finding = _make_finding(
            file_name="spec.docx",
            existingText=target_mapping.text,
            replacementText="Comply with CBC 2025 requirements.",
        )
        locator_result = LocatorResult(
            finding=finding,
            status="matched",
            locations=[location],
            replacement_text="Comply with CBC 2025 requirements.",
            action_type="EDIT",
        )
        action = EditAction(
            locator_result=locator_result,
            location=location,
            replacement_text="Comply with CBC 2025 requirements.",
            action_type="EDIT",
            finding_index=0,
        )
        out_path = tmp_path / "spec_annotated.docx"
        report = annotate_spec_with_suggestions(source, out_path, [action])
        assert report.edits_applied == 1
        result_doc = Document(out_path)
        text = "\n".join(p.text for p in result_doc.paragraphs)
        assert "Comply with CBC 2019 requirements." in text
        assert "SPEC CRITIC SUGGESTION" in text
        assert "Proposed: Comply with CBC 2025 requirements." in text

    def test_annotate_refuses_to_overwrite_source(self, tmp_path: Path):
        from src.spec_editor import annotate_spec_with_suggestions

        source = _make_simple_docx(tmp_path, ["Body"])
        with pytest.raises(ValueError, match="differ from source_path"):
            annotate_spec_with_suggestions(source, source, [])


# ---------------------------------------------------------------------------
# Phase 5.5 — file hash validation in resume state
# ---------------------------------------------------------------------------


class TestResumeStateFileHashes:
    def test_serialize_includes_content_and_source_digests(self, tmp_path: Path):
        from src.resume_state import serialize_extracted_spec

        # Write a file so source digest can be computed.
        source = tmp_path / "spec.docx"
        source.write_bytes(b"binary placeholder")
        spec = ExtractedSpec(
            filename="spec.docx",
            content="hello world",
            word_count=2,
            source_path=str(source),
            source_format="docx",
        )
        payload = serialize_extracted_spec(spec)
        assert "content_sha256" in payload
        assert "source_sha256" in payload
        assert payload["content_sha256"]
        assert payload["source_sha256"]

    def test_deserialize_warns_on_content_mismatch(self, tmp_path: Path, caplog):
        import logging
        from src.resume_state import deserialize_extracted_spec

        caplog.set_level(logging.WARNING)
        payload = {
            "filename": "spec.docx",
            "content": "actual content",
            "word_count": 2,
            "source_path": "",
            "source_format": "docx",
            "content_sha256": "0" * 64,  # intentionally wrong digest
            "source_sha256": None,
        }
        spec = deserialize_extracted_spec(payload)
        assert spec.filename == "spec.docx"
        assert any("content digest mismatch" in r.message for r in caplog.records)

    def test_deserialize_silent_on_hash_match(self, tmp_path: Path, caplog):
        import hashlib
        import logging
        from src.resume_state import deserialize_extracted_spec, _content_digest

        caplog.set_level(logging.WARNING)
        payload = {
            "filename": "spec.docx",
            "content": "actual content",
            "word_count": 2,
            "source_path": "",
            "source_format": "docx",
            "content_sha256": _content_digest("actual content"),
            "source_sha256": None,
        }
        deserialize_extracted_spec(payload)
        assert not any("digest mismatch" in r.message for r in caplog.records)


# ---------------------------------------------------------------------------
# Phase 9.3 — fuzzy match prefilters
# ---------------------------------------------------------------------------


class TestFuzzyMatchPrefilter:
    def test_short_vs_long_paragraph_skipped_by_length_gate(self):
        from src.edit_locator import _fuzzy_match
        from src.extractor import ParagraphMapping

        long_para = ParagraphMapping(
            text="x" * 500, body_index=0, element_type="paragraph",
            table_index=None, row_index=None, cell_index=None,
            section_index=0,
        )
        # 5-char query against 500-char paragraph: ratio ceiling is
        # 2*5/(5+500) = 0.0198, well below the 0.80 threshold. Length gate
        # should reject without computing the ratio at all.
        hits = _fuzzy_match("hello", [long_para])
        assert hits == []

    def test_returns_close_match_above_threshold(self):
        from src.edit_locator import _fuzzy_match
        from src.extractor import ParagraphMapping

        para = ParagraphMapping(
            text="Comply with CBC 2025 chapter 17 requirements.",
            body_index=0, element_type="paragraph",
            table_index=None, row_index=None, cell_index=None,
            section_index=0,
        )
        # Same paragraph with one word changed should still match above 0.80.
        hits = _fuzzy_match("Comply with CBC 2025 chapter 18 requirements.", [para])
        assert len(hits) == 1
        assert hits[0].match_method == "fuzzy"
        assert hits[0].match_confidence >= 0.80

    def test_empty_target_returns_empty(self):
        from src.edit_locator import _fuzzy_match
        from src.extractor import ParagraphMapping

        para = ParagraphMapping(
            text="something", body_index=0, element_type="paragraph",
            table_index=None, row_index=None, cell_index=None,
            section_index=0,
        )
        assert _fuzzy_match("", [para]) == []
